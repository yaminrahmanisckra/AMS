import os
# Force OpenBLAS to use single thread to prevent hang during imports (common on macOS + Flask/Eventlet)
os.environ['OPENBLAS_NUM_THREADS'] = '1'

import json
import csv
from datetime import datetime
from io import BytesIO, StringIO
from dotenv import load_dotenv
from flask import Flask, render_template, redirect, url_for, flash, request, send_file, Response, jsonify, current_app, session, make_response
from flask_login import LoginManager, current_user, login_required
from extensions import db, migrate, mail
from user_models import User
from error_handler import register_error_handlers, setup_error_logging, check_dependencies, check_file_permissions, get_system_info
from role_utils import (
    ADMIN_ROLE,
    ROLE_LABELS,
    ROLE_CHOICES,
    NON_ADMIN_ROLE_CHOICES,
    DEFAULT_TEACHING_ROLE,
    has_teacher_privileges,
    parse_roles,
    get_primary_role,
    is_admin,
    validate_role_selection,
    serialize_roles,
)
from blueprints.class_management.models import (
    ExamPaperEvaluation,
    ExamScrutinizerInvite,
    StudentFeedbackLink,
    StudentFeedbackResponse,
    Teacher,
    CourseOutline,
    StudentNotification,
)
from blueprints.student_management.models import Student
from blueprints.course_management.models import Course, DutyAssignment, Curriculum, CurriculumYearTerm, StudentCourseRegistration, SessionArchive, ActiveSemesterConfig, CourseSessionAssignment
from blueprints.remuneration_management.models import RemunerationForm

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:  # pragma: no cover - optional dependency
    Workbook = None
    Alignment = Font = Border = Side = get_column_letter = None

from flask_socketio import SocketIO, emit, join_room, leave_room

load_dotenv()

# Set library path for WeasyPrint on macOS before creating app
import os
import platform
import sys

if platform.system() == 'Darwin':  # macOS
    homebrew_lib_path = '/opt/homebrew/lib'
    if os.path.exists(homebrew_lib_path):
        # Set both DYLD_LIBRARY_PATH and DYLD_FALLBACK_LIBRARY_PATH for better compatibility
        current_dyld = os.environ.get('DYLD_LIBRARY_PATH', '')
        current_fallback = os.environ.get('DYLD_FALLBACK_LIBRARY_PATH', '')
        
        if homebrew_lib_path not in current_dyld:
            os.environ['DYLD_LIBRARY_PATH'] = f"{homebrew_lib_path}:{current_dyld}" if current_dyld else homebrew_lib_path
        
        # DYLD_FALLBACK_LIBRARY_PATH is more reliable on newer macOS versions
        if homebrew_lib_path not in current_fallback:
            fallback_paths = [homebrew_lib_path]
            if current_fallback:
                fallback_paths.extend(current_fallback.split(':'))
            os.environ['DYLD_FALLBACK_LIBRARY_PATH'] = ':'.join(fallback_paths)
        
        # Set PKG_CONFIG_PATH to help pkg-config find libraries
        pkg_config_path = '/opt/homebrew/lib/pkgconfig'
        if os.path.exists(pkg_config_path):
            current_pkg = os.environ.get('PKG_CONFIG_PATH', '')
            if pkg_config_path not in current_pkg:
                os.environ['PKG_CONFIG_PATH'] = f"{pkg_config_path}:{current_pkg}" if current_pkg else pkg_config_path
        
        # Pre-load required libraries using ctypes before WeasyPrint imports
        # This works around SIP restrictions on DYLD_LIBRARY_PATH
        try:
            import ctypes
            from ctypes import util as ctypes_util
            
            # Monkey-patch ctypes.util.find_library to check Homebrew path first
            # This helps cffi (used by WeasyPrint) find the libraries
            original_find_library = ctypes_util.find_library
            
            def patched_find_library(name):
                # Try Homebrew path first
                # WeasyPrint looks for libraries like 'gobject-2.0-0' (without lib prefix)
                # Map common library names to their actual filenames
                lib_mappings = {
                    'gobject-2.0-0': 'libgobject-2.0.0.dylib',
                    'gobject-2.0': 'libgobject-2.0.dylib',
                    'glib-2.0-0': 'libglib-2.0.0.dylib',
                    'glib-2.0': 'libglib-2.0.dylib',
                    'cairo': 'libcairo.2.dylib',
                    'pango-1.0-0': 'libpango-1.0.0.dylib',
                    'pango-1.0': 'libpango-1.0.dylib',
                    'gdk_pixbuf-2.0-0': 'libgdk_pixbuf-2.0.0.dylib',
                    'gdk_pixbuf-2.0': 'libgdk_pixbuf-2.0.dylib',
                }
                
                # Check if we have a mapping for this library
                if name in lib_mappings:
                    lib_file = lib_mappings[name]
                    lib_path = os.path.join(homebrew_lib_path, lib_file)
                    if os.path.exists(lib_path):
                        return lib_path
                
                # Also try common patterns
                for pattern in [f'lib{name}.dylib', f'lib{name}.0.dylib', f'{name}.dylib']:
                    lib_path = os.path.join(homebrew_lib_path, pattern)
                    if os.path.exists(lib_path):
                        return lib_path
                
                # Fall back to original find_library
                result = original_find_library(name)
                if result:
                    return result
                
                # Last resort: check Homebrew path with the name as-is
                lib_path = os.path.join(homebrew_lib_path, name)
                if os.path.exists(lib_path):
                    return lib_path
                
                return None
            
            # Apply the monkey patch
            ctypes_util.find_library = patched_find_library
            
            # Also try to patch cffi's library loading if available
            # This helps WeasyPrint's cffi-based library loading
            try:
                import cffi
                # Pre-load libraries so they're available to cffi's dlopen
                # cffi uses the system's dlopen which should find pre-loaded libraries
                pass  # The pre-loading below should make libraries available
            except ImportError:
                pass  # cffi not available yet, that's okay
            
            # List of libraries that WeasyPrint needs, in dependency order
            # Load base libraries first, then dependent ones
            required_libs = [
                'libglib-2.0.0.dylib',      # Base GLib library
                'libgobject-2.0.0.dylib',   # GObject (depends on glib)
                'libgmodule-2.0.0.dylib',   # GModule (depends on glib)
                'libgthread-2.0.0.dylib',   # GThread (depends on glib)
                'libcairo.2.dylib',         # Cairo graphics library
                'libgdk_pixbuf-2.0.0.dylib', # GDK Pixbuf
                'libpango-1.0.0.dylib',     # Pango text layout
                'libpangocairo-1.0.0.dylib', # Pango-Cairo integration
            ]
            
            # Try to load libraries from Homebrew path
            for lib_name in required_libs:
                lib_path = os.path.join(homebrew_lib_path, lib_name)
                if os.path.exists(lib_path):
                    try:
                        ctypes.CDLL(lib_path, mode=ctypes.RTLD_GLOBAL)
                    except (OSError, AttributeError):
                        # Library might already be loaded or have dependencies
                        # Try loading with RTLD_LOCAL if RTLD_GLOBAL fails
                        try:
                            ctypes.CDLL(lib_path)
                        except:
                            pass
                else:
                    # Try alternative name (without version)
                    alt_name = lib_name.replace('.0.dylib', '.dylib').replace('.2.dylib', '.dylib')
                    alt_path = os.path.join(homebrew_lib_path, alt_name)
                    if os.path.exists(alt_path):
                        try:
                            ctypes.CDLL(alt_path, mode=ctypes.RTLD_GLOBAL)
                        except (OSError, AttributeError):
                            try:
                                ctypes.CDLL(alt_path)
                            except:
                                pass
        except Exception:
            # ctypes might not be available, but that's okay
            # The error will be caught when WeasyPrint tries to import
            pass

# xhtml2pdf helper function removed - conflicts with reportlab 4.0.7
# PDF export routes now redirect to DOCX export or show error message

def create_app():
    app = Flask(__name__)

    @app.template_filter('date')
    def date_format_filter(value, format='%Y'):
        if value == 'now':
            return datetime.utcnow().strftime(format)
        return value

    app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'a_very_secret_default_key')
    app.config['TEMPLATES_AUTO_RELOAD'] = False
    app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 31536000
    
    # Session cookie configuration for ngrok compatibility
    # Don't set domain to allow cookies to work with any domain (including ngrok)
    app.config['SESSION_COOKIE_HTTPONLY'] = True
    app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
    app.config['SESSION_COOKIE_SECURE'] = False  # Set to True only for HTTPS in production
    app.config['PERMANENT_SESSION_LIFETIME'] = 86400  # 24 hours

    # Database configuration - Check environment variable first, fallback to SQLite
    basedir = os.path.abspath(os.path.dirname(__file__))
    database_url = os.getenv('DATABASE_URL')
    if database_url:
        app.config['SQLALCHEMY_DATABASE_URI'] = database_url
        # MySQL connection pool settings to prevent "MySQL server has gone away" errors
        # Optimized for cPanel with reduced connection pool size and timeouts
        app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
            'pool_pre_ping': True,  # Check connection health before using
            'pool_recycle': 1800,   # Recycle connections after 30 minutes (1800 seconds) - reduced from 1 hour
            'pool_size': 5,         # Number of connections to keep in pool - reduced from 10 for cPanel
            'max_overflow': 5,      # Maximum number of connections beyond pool_size - reduced from 20
            'pool_timeout': 20,     # Timeout when getting connection from pool - reduced from 30
            'connect_args': {
                'connect_timeout': 5,   # Connection timeout in seconds - reduced from 10
                'read_timeout': 20,     # Read timeout in seconds - reduced from 30
                'write_timeout': 20,    # Write timeout in seconds - reduced from 30
            }
        }
    else:
        # Fallback to SQLite if DATABASE_URL not set
        db_path = os.path.join(basedir, 'instance', 'academic_management.db')
        os.makedirs(os.path.dirname(db_path), exist_ok=True)
        app.config['SQLALCHEMY_DATABASE_URI'] = f"sqlite:///{db_path}"
        
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

    app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
    app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
    app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'True') == 'True'
    # cPanel/local relay often uses port 25 with both TLS and SSL off; honor MAIL_USE_SSL from env
    app.config['MAIL_USE_SSL'] = os.getenv('MAIL_USE_SSL', 'False') == 'True'
    app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
    app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
    app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER', app.config['MAIL_USERNAME'])
    # Notification channel (e.g. noreply@) — separate from MAIL_* used for password recovery (recovery@)
    def _env_bool(name, fallback):
        v = os.getenv(name)
        if v is None:
            return fallback
        return str(v).strip().lower() in ('1', 'true', 'yes', 'on')

    app.config['NOTIFICATION_MAIL_SERVER'] = os.getenv('NOTIFICATION_MAIL_SERVER') or app.config['MAIL_SERVER']
    _notif_port = os.getenv('NOTIFICATION_MAIL_PORT')
    app.config['NOTIFICATION_MAIL_PORT'] = int(_notif_port) if _notif_port else app.config['MAIL_PORT']
    app.config['NOTIFICATION_MAIL_USE_TLS'] = _env_bool(
        'NOTIFICATION_MAIL_USE_TLS', app.config['MAIL_USE_TLS']
    )
    app.config['NOTIFICATION_MAIL_USE_SSL'] = _env_bool(
        'NOTIFICATION_MAIL_USE_SSL', app.config['MAIL_USE_SSL']
    )
    app.config['NOTIFICATION_MAIL_USERNAME'] = os.getenv('NOTIFICATION_MAIL_USERNAME')
    app.config['NOTIFICATION_MAIL_PASSWORD'] = os.getenv('NOTIFICATION_MAIL_PASSWORD')
    app.config['NOTIFICATION_MAIL_SENDER'] = os.getenv(
        'NOTIFICATION_MAIL_SENDER', os.getenv('NOTIFICATION_MAIL_USERNAME')
    )
    # Re-merge from os.environ so cPanel/Passenger values always win over any stale defaults
    for k in list(os.environ.keys()):
        if not k.startswith('NOTIFICATION_MAIL_'):
            continue
        val = os.environ.get(k)
        if val is None or str(val).strip() == '':
            continue
        if k.endswith('_PASSWORD'):
            app.config[k] = str(val).rstrip('\r\n')
        else:
            app.config[k] = str(val).strip()
    # Optional: public site URL (no trailing slash) for emails when reverse-proxy hides the real host
    app.config['PUBLIC_APP_URL'] = (os.getenv('PUBLIC_APP_URL') or '').strip().rstrip('/')
    app.config['DEFAULT_STUDENT_PASSWORD'] = os.getenv('DEFAULT_STUDENT_PASSWORD', 'Student@123')
    
    # OpenAI API configuration
    app.config['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY', '')
    
    # File upload configuration
    upload_folder = os.path.join(basedir, 'static', 'uploads', 'user_photos')
    os.makedirs(upload_folder, exist_ok=True)
    app.config['UPLOAD_FOLDER'] = upload_folder
    app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
    app.config['ALLOWED_EXTENSIONS'] = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
    
    # Production logging configuration - reduce verbose logging overhead
    import logging
    if os.getenv('FLASK_ENV') == 'production' or os.getenv('CPANEL'):
        app.logger.setLevel(logging.WARNING)  # Only log warnings and errors
        logging.getLogger('werkzeug').setLevel(logging.WARNING)
    
    mail.init_app(app)

    db.init_app(app)
    migrate.init_app(app, db)
    
    login_manager = LoginManager()
    login_manager.login_view = 'auth.login' 
    login_manager.init_app(app)

    @login_manager.user_loader
    def load_user(user_id):
        try:
            return User.query.get(int(user_id))
        except (ValueError, TypeError):
            return None

    @app.before_request
    def attach_active_role():
        if not current_user.is_authenticated:
            session.pop('active_role', None)
            return
        active_role = session.get('active_role')
        stored_roles = set(parse_roles(current_user.role))
        if active_role:
            if active_role == ADMIN_ROLE and ADMIN_ROLE not in stored_roles:
                active_role = None
            elif active_role != ADMIN_ROLE and active_role not in stored_roles:
                active_role = None
        if not active_role:
            session.pop('active_role', None)
        current_user.active_role = active_role

    @app.teardown_appcontext
    def shutdown_session(exception=None):
        """Close database session after each request to prevent memory leaks"""
        db.session.remove()

    from blueprints.class_management.routes import class_management_bp
    from blueprints.result_management.routes import result_management_bp
    from blueprints.routine_management.routes import routine_management_bp
    from blueprints.auth.routes import auth_bp
    from blueprints.student_management import student_management_bp
    from blueprints.course_management import course_management_bp
    from blueprints.leave_application import leave_application_bp

    app.register_blueprint(class_management_bp, url_prefix='/class-management')
    app.register_blueprint(result_management_bp, url_prefix='/result-management')
    app.register_blueprint(routine_management_bp, url_prefix='/routine-management')
    app.register_blueprint(auth_bp)
    app.register_blueprint(student_management_bp, url_prefix='/student-management')
    app.register_blueprint(course_management_bp, url_prefix='/course-management')
    app.register_blueprint(leave_application_bp, url_prefix='/leave-application')
    from blueprints.academic_calendar import academic_calendar_bp
    app.register_blueprint(academic_calendar_bp, url_prefix='/academic-calendar')
    from blueprints.curriculator import curriculator_bp
    app.register_blueprint(curriculator_bp, url_prefix='/curriculator')
    try:
        try:
            from blueprints.self_assessment import self_assessment_bp
        except ImportError:
            from blueprints.self_assessment import self_assessment as self_assessment_bp
        app.register_blueprint(self_assessment_bp, url_prefix='/self-assessment')
    except Exception as e:
        import logging
        import traceback
        logging.warning('Self Assessment blueprint not loaded: %s', e)
        logging.warning(traceback.format_exc())
    
    @app.route('/debug-self-assessment')
    def debug_self_assessment():
        """Temporary: দেখুন Self Assessment ব্লুপ্রিন্ট কেন লোড হচ্ছে না। ব্যবহারের পর রুটটি মুছে ফেলুন।"""
        try:
            try:
                from blueprints.self_assessment import self_assessment_bp
            except ImportError:
                from blueprints.self_assessment import self_assessment as self_assessment_bp
            return '<h2>OK</h2><p>Blueprint imports successfully. If the card still shows setup message, restart the app and clear browser cache.</p>'
        except Exception as e:
            import traceback
            tb = traceback.format_exc()
            return '<h2>Self Assessment load error</h2><pre style="white-space:pre-wrap;background:#f5f5f5;padding:1em;">' + str(e) + '\n\n' + tb + '</pre>'
    
    # Service Worker route for PWA
    @app.route('/sw.js')
    def service_worker():
        response = send_file('static/sw.js', mimetype='application/javascript')
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response
    
    # Manifest route for PWA
    @app.route('/manifest.json')
    def manifest():
        response = send_file('static/manifest.json', mimetype='application/manifest+json')
        response.headers['Cache-Control'] = 'public, max-age=3600'
        return response
    
    # Initialize SocketIO for WebSocket support
    from utils.websocket_events import init_socketio
    socketio = SocketIO(app, cors_allowed_origins="*", async_mode='threading', logger=False, engineio_logger=False)
    init_socketio(socketio)
    
    # WebSocket connection handlers
    @socketio.on('connect', namespace='/')
    def handle_connect(auth):
        """Handle client connection"""
        # Flask-Login session is not available in WebSocket context
        # We'll authenticate via session cookie in the request context
        # For now, allow connection - authentication will be handled via session
        return True
    
    @socketio.on('disconnect', namespace='/')
    def handle_disconnect():
        """Handle client disconnection"""
        from flask_login import current_user
        if current_user.is_authenticated:
            leave_room(f'user_{current_user.id}')
    
    @socketio.on('join_session', namespace='/')
    def handle_join_session(data):
        """Join a session room for live updates"""
        from flask_login import current_user
        if current_user.is_authenticated:
            session_id = data.get('session_id')
            if session_id:
                join_room(f'session_{session_id}')
                emit('joined_session', {'session_id': session_id})
    
    @socketio.on('leave_session', namespace='/')
    def handle_leave_session(data):
        """Leave a session room"""
        from flask_login import current_user
        if current_user.is_authenticated:
            session_id = data.get('session_id')
            if session_id:
                leave_room(f'session_{session_id}')
                emit('left_session', {'session_id': session_id})
    
    @socketio.on('join_result_session', namespace='/')
    def handle_join_result_session(data):
        """Join a result session room for live updates"""
        from flask_login import current_user
        if current_user.is_authenticated:
            session_id = data.get('session_id')
            if session_id:
                join_room(f'result_session_{session_id}')
                emit('joined_result_session', {'session_id': session_id})
    
    @socketio.on('leave_result_session', namespace='/')
    def handle_leave_result_session(data):
        """Leave a result session room"""
        from flask_login import current_user
        if current_user.is_authenticated:
            session_id = data.get('session_id')
            if session_id:
                leave_room(f'result_session_{session_id}')
                emit('left_result_session', {'session_id': session_id})
    
    # Store socketio in app context for access in routes
    app.socketio = socketio
    
    # Register error handlers for cPanel deployment
    register_error_handlers(app)
    
    # Setup logging
    setup_error_logging()
    
    # WeasyPrint availability check removed from startup to prevent hanging
    # WeasyPrint will be imported lazily when actually needed (in routes)
    # This prevents startup hang issues on macOS
    app.logger.info("WeasyPrint will be imported lazily when needed for PDF generation")

    def _migrate_legacy_roles():
        """Convert legacy 'user' roles into the new teacher role."""
        try:
            legacy_users = User.query.filter_by(role='user').all()
            if not legacy_users:
                return
            for user in legacy_users:
                user.role = DEFAULT_TEACHING_ROLE
            db.session.commit()
            app.logger.info("Migrated %s legacy user role(s) to '%s'.", len(legacy_users), DEFAULT_TEACHING_ROLE)
        except Exception as exc:
            app.logger.warning('Role migration skipped: %s', exc)

    with app.app_context():
        _migrate_legacy_roles()

    def _is_exam_committee_chief():
        """Check if current user is assigned as Exam Committee Chief"""
        if not current_user.is_authenticated:
            return False
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return False
        return DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            assigned_teacher_id=teacher.id,
            status='active'
        ).count() > 0

    def _is_exam_committee_member():
        """Check if current user is an Exam Committee Internal Member"""
        if not current_user.is_authenticated:
            return False
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return False
        return DutyAssignment.query.filter_by(
            duty_type='exam_committee_member',
            assigned_teacher_id=teacher.id,
            status='active'
        ).count() > 0

    def _get_exam_committee_assignments(teacher, academic_session=None, year=None, term=None):
        """Return active chief/member assignments for a teacher, optionally scoped by session/year/term."""
        if not teacher:
            return [], []
        chief_query = DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            assigned_teacher_id=teacher.id,
            status='active'
        )
        member_query = DutyAssignment.query.filter_by(
            duty_type='exam_committee_member',
            assigned_teacher_id=teacher.id,
            status='active'
        )
        if academic_session is not None:
            chief_query = chief_query.filter(DutyAssignment.academic_session == academic_session)
            member_query = member_query.filter(DutyAssignment.academic_session == academic_session)
        if year is not None:
            chief_query = chief_query.filter(DutyAssignment.year == year)
            member_query = member_query.filter(DutyAssignment.year == year)
        if term is not None:
            chief_query = chief_query.filter(DutyAssignment.term == term)
            member_query = member_query.filter(DutyAssignment.term == term)
        return chief_query.all(), member_query.all()

    def _is_tabulator():
        """Check if current user has an active tabulator assignment"""
        if not current_user.is_authenticated:
            return False
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return False
        return DutyAssignment.query.filter_by(
            duty_type='tabulator',
            assigned_teacher_id=teacher.id,
            status='active'
        ).count() > 0

    def get_visible_dashboard_items():
        """Returns a dict of which dashboard items should be visible for the current user"""
        if not current_user.is_authenticated:
            return {}
        
        is_officer = 'officer' in parse_roles(current_user.role)
        is_teaching_assistant = 'teaching_assistant' in parse_roles(current_user.role)
        is_teacher = 'teacher' in parse_roles(current_user.role) and not is_teaching_assistant and not is_officer
        is_head = 'head' in parse_roles(current_user.role) or 'dean' in parse_roles(current_user.role)
        head_active = (current_user.active_role == 'head') if hasattr(current_user, 'active_role') else False
        
        # Check for course registration review
        show_course_registration_review = False
        teacher = _current_teacher()
        if teacher:
            show_course_registration_review = DutyAssignment.query.filter_by(
                duty_type='course_coordinator',
                assigned_teacher_id=teacher.id,
                status='active'
            ).count() > 0
        
        return {
            'show_class_management': not is_teaching_assistant and not is_officer and not head_active,
            'show_result_management': (is_head or _is_tabulator()) and not is_teaching_assistant and not is_officer,
            'show_routine_management': (is_teaching_assistant or (not is_officer)) and not head_active,
            'show_exam_evaluation': not is_teaching_assistant and not is_officer and not head_active,
            'show_students_management': (is_teaching_assistant or not is_teacher) and not is_officer,
            'show_curriculum_management': (is_teaching_assistant or not is_teacher) and not is_officer,
            'show_remuneration': not is_teaching_assistant and not head_active,
            'show_academic_calendar': True,  # Always visible
            'show_course_registration': show_course_registration_review and not is_teaching_assistant,
            'show_exam_committee_management': (_is_exam_committee_chief() or _is_exam_committee_member()) and not is_teaching_assistant,
        }

    @app.context_processor
    def inject_role_helpers():
        from datetime import date, timedelta
        visible_items = get_visible_dashboard_items() if current_user.is_authenticated else {}
        # Student notifications: inject on every page so bell shows on dashboard and elsewhere
        student_notification_count = 0
        student_notifications = []
        if current_user.is_authenticated and parse_roles(current_user.role) and 'student' in parse_roles(current_user.role):
            try:
                unread = StudentNotification.query.filter_by(
                    user_id=current_user.id
                ).filter(
                    StudentNotification.read_at.is_(None)
                ).order_by(StudentNotification.created_at.desc()).limit(20).all()
                student_notification_count = len(unread)
                student_notifications = [
                    {'id': n.id, 'type': n.type, 'title': n.title, 'link_url': n.link_url, 'created_at': n.created_at}
                    for n in unread
                ]
            except Exception:
                pass
        return {
            'ROLE_LABELS': ROLE_LABELS,
            'ROLE_CHOICES': NON_ADMIN_ROLE_CHOICES,
            'ALL_ROLE_CHOICES': ROLE_CHOICES,
            'has_teacher_privileges': has_teacher_privileges,
            'parse_roles': parse_roles,
            'get_primary_role': get_primary_role,
            'is_admin': is_admin,
            'is_exam_committee_chief': _is_exam_committee_chief,
            'is_exam_committee_member': _is_exam_committee_member,
            'is_tabulator': _is_tabulator,
            'date': date,
            'timedelta': timedelta,
            'datetime': datetime,
            'student_notification_count': student_notification_count,
            'student_notifications': student_notifications,
            **visible_items,  # Add visible dashboard items to context
        }

    @app.route('/')
    def index():
        if not current_user.is_authenticated:
            return redirect(url_for('auth.login'))
        
        if is_admin(current_user):
            return redirect(url_for('admin_dashboard'))

        roles = parse_roles(current_user.role)
        if 'head' in roles or 'dean' in roles:
            return redirect(url_for('head_dashboard'))

        if getattr(current_user, 'active_role', None) == 'student':
            return redirect(url_for('student_dashboard'))

        show_course_registration_review = False
        teacher = _current_teacher()
        if teacher:
            show_course_registration_review = DutyAssignment.query.filter_by(
                duty_type='course_coordinator',
                assigned_teacher_id=teacher.id,
                status='active'
            ).count() > 0

        show_self_assessment = False
        if 'self_assessment.index' in current_app.view_functions:
            try:
                from blueprints.self_assessment.routes import is_psac_member_or_head
                show_self_assessment = is_psac_member_or_head()
            except Exception:
                pass

        # Leave Application card: show for teacher-privilege users on main dashboard
        show_leave_application = has_teacher_privileges(current_user)

        response = make_response(render_template(
            'dashboard.html',
            show_course_registration_review=show_course_registration_review,
            show_self_assessment=show_self_assessment,
            show_leave_application=show_leave_application,
        ))
        # Add cache-control headers to prevent browser caching of user-specific content
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, private'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response

    @app.route('/test-form', methods=['GET', 'POST'])
    def test_form():
        test_result = None
        if request.method == 'POST':
            test_result = {
                'name': request.form.get('test_name'),
                'email': request.form.get('test_email'),
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            flash('Test form submitted successfully!', 'success')
        return render_template('test_form.html', test_result=test_result)

    @app.route('/simple-test', methods=['GET', 'POST'])
    def simple_test():
        result = None
        if request.method == 'POST':
            result = f"Name: {request.form.get('test_name')}, Email: {request.form.get('test_email')}"
        return render_template('simple_test.html', result=result)

    def _generate_teacher_short_name(base: str) -> str:
        base = (base or 'teacher').split(' ')[0]
        base = ''.join(ch for ch in base.lower() if ch.isalnum()) or 'teacher'
        base = base[:10]
        candidate = base
        counter = 1
        while Teacher.query.filter_by(short_name=candidate).first():
            suffix = str(counter)
            candidate = f"{base[:10-len(suffix)]}{suffix}"
            counter += 1
        return candidate

    def _ensure_teacher_for_user(user):
        if not user or not has_teacher_privileges(user):
            return None
        # Canonical mapping: prefer explicit users.teacher_id when available.
        teacher = None
        user_teacher_id = getattr(user, 'teacher_id', None)
        if user_teacher_id:
            teacher = Teacher.query.get(user_teacher_id)

        if not teacher:
            teacher = Teacher.query.filter_by(name=user.full_name).first()

        if not teacher and getattr(user, 'full_name', None):
            normalized_name = str(user.full_name).strip().casefold()
            teacher = next(
                (
                    t for t in Teacher.query.all()
                    if str(t.name or '').strip().casefold() == normalized_name
                ),
                None
            )

        if teacher:
            if hasattr(user, 'teacher_id') and user.teacher_id != teacher.id:
                user.teacher_id = teacher.id
                db.session.commit()
            return teacher
        short_name = _generate_teacher_short_name(user.username or user.full_name or 'teacher')
        teacher = Teacher(name=user.full_name, short_name=short_name)
        db.session.add(teacher)
        if hasattr(user, 'teacher_id'):
            db.session.flush()
            user.teacher_id = teacher.id
        db.session.commit()
        return teacher

    def _current_teacher():
        if not current_user.is_authenticated:
            return None
        return _ensure_teacher_for_user(current_user)

    def _teacher_identity_ids(teacher):
        """Return all teacher IDs that likely represent the same person."""
        if not teacher:
            return set()
        identity_ids = {teacher.id}
        normalized_name = str(teacher.name or '').strip().casefold()
        if not normalized_name:
            return identity_ids
        for candidate in Teacher.query.all():
            if str(candidate.name or '').strip().casefold() == normalized_name:
                identity_ids.add(candidate.id)
        return identity_ids

    def _require_teacher_privileges():
        if not has_teacher_privileges(current_user):
            flash('This feature is restricted to teaching staff.', 'danger')
            return redirect(url_for('index'))
        return None

    @app.route('/feedback/<string:code>', methods=['GET', 'POST'])
    def student_feedback_form(code):
        link = StudentFeedbackLink.query.filter_by(access_code=code.upper()).first()
        status = 'active'
        if not link:
            status = 'not_found'
        elif link.expires_at and datetime.utcnow() > link.expires_at:
            status = 'expired'

        session_info = link.session if link else None
        prefill = {
            'academic_session': session_info.academic_session if session_info else '',
            'course_title': session_info.course_name if session_info else '',
            'course_code': session_info.course_code if session_info else '',
        }

        if status != 'active':
            return render_template(
                'student_feedback_form.html',
                status=status,
                link=None,
                submitted=False,
                prefill=prefill,
            )

        cookie_key = f'fb_{code}'
        already_submitted = request.cookies.get(cookie_key) == '1' and not link.allow_multiple

        if request.method == 'POST' and not already_submitted:
            academic_info = {
                'academic_session': request.form.get('academic_session', '').strip(),
                'course_title': request.form.get('course_title', '').strip(),
                'course_code': request.form.get('course_code', '').strip(),
            }

            section_a = {
                'course_structure': request.form.get('course_structure'),
                'course_goals': request.form.get('course_goals'),
                'course_content_guidance': request.form.get('course_content_guidance'),
                'course_interest': request.form.get('course_interest'),
            }

            section_b = {
                'course_plan_discussed': request.form.get('course_plan_discussed'),
                'guidelines_received': request.form.get('guidelines_received'),
                'assessment_helpful': request.form.get('assessment_helpful'),
                'feedback_timely': request.form.get('feedback_timely'),
            }
            teaching_methods = request.form.getlist('teaching_methods')
            other_method = request.form.get('teaching_methods_other', '').strip()
            if other_method:
                teaching_methods.append(other_method)
            section_b['teaching_methods'] = teaching_methods

            section_c = {
                'study_time': request.form.get('study_time'),
                'attendance_percent': request.form.get('attendance_percent'),
            }
            effort_focus = request.form.getlist('effort_focus')
            other_effort = request.form.get('effort_other', '').strip()
            if other_effort:
                effort_focus.append(other_effort)
            section_c['effort_focus'] = effort_focus

            section_d = {
                'likes': request.form.get('likes', '').strip(),
                'challenges': request.form.get('challenges', '').strip(),
                'suggestions': request.form.get('suggestions', '').strip(),
            }

            if not all([section_d['likes'], section_d['challenges'], section_d['suggestions']]):
                flash('Section D (Praise and Suggestions)-এর সবগুলো ঘর পূরণ করা বাধ্যতামূলক।', 'warning')
            # Ensure at least one core question is answered
            elif not any(section_a.values()) and not any(section_b.values()) and not effort_focus and not section_d['likes'] and not section_d['challenges'] and not section_d['suggestions']:
                flash('কমপক্ষে একটি প্রশ্নের উত্তর দিন।', 'warning')
            else:
                payload = {
                    'academic_info': academic_info,
                    'section_a': section_a,
                    'section_b': section_b,
                    'section_c': section_c,
                    'section_d': section_d,
                    'user_agent': request.headers.get('User-Agent'),
                }
                response = StudentFeedbackResponse(
                    feedback_link_id=link.id,
                    payload=json.dumps(payload),
                )
                db.session.add(response)
                db.session.commit()
                resp = redirect(url_for('student_feedback_form', code=code, submitted='1'))
                if not link.allow_multiple:
                    resp.set_cookie(cookie_key, '1', max_age=60 * 60 * 24 * 90)
                return resp

        submitted = request.args.get('submitted') == '1' or already_submitted
        return render_template(
            'student_feedback_form.html',
            status='active',
            link=link,
            submitted=submitted,
            already_submitted=already_submitted,
            prefill=prefill,
        )

    def _normalize_section(value):
        """Normalize section labels for comparison (blank -> FULL)."""
        value = (value or '').strip()
        if not value:
            return 'FULL'
        return value.upper()

    def _assignment_matches_entry(assignment, entry):
        """Return True if a DutyAssignment matches a given exam entry."""
        if not assignment or assignment.duty_type != 'scrutinizer' or not entry:
            return False

        def _norm(text):
            return (text or '').strip().lower()

        if assignment.course_code and _norm(entry.course_code) != _norm(assignment.course_code):
            return False
        if assignment.year and _norm(entry.year) != _norm(assignment.year):
            return False
        if assignment.term and _norm(entry.term) != _norm(assignment.term):
            return False
        if assignment.batch and _norm(entry.batch) != _norm(assignment.batch):
            return False
        if assignment.academic_session and _norm(entry.academic_session) != _norm(assignment.academic_session):
            return False
        return True

    def _apply_scrutinizer_assignment(assignment):
        """Assign a teacher as scrutinizer to matching exam entries."""
        if not assignment or assignment.duty_type != 'scrutinizer' or not assignment.assigned_teacher_id:
            return
        entries = ExamPaperEvaluation.query.filter(ExamPaperEvaluation.archived.is_(False)).all()
        changed = False
        for entry in entries:
            if _assignment_matches_entry(assignment, entry):
                entry.assigned_scrutinizer_id = assignment.assigned_teacher_id
                changed = True
        if changed:
            db.session.commit()

    def _clear_scrutinizer_assignment(assignment):
        """Remove a teacher from matching exam entries when assignment is removed."""
        if not assignment or assignment.duty_type != 'scrutinizer' or not assignment.assigned_teacher_id:
            return
        entries = ExamPaperEvaluation.query.filter(
            ExamPaperEvaluation.assigned_scrutinizer_id == assignment.assigned_teacher_id
        ).all()
        changed = False
        for entry in entries:
            if _assignment_matches_entry(assignment, entry):
                entry.assigned_scrutinizer_id = None
                changed = True
        if changed:
            db.session.commit()

    def _auto_assign_scrutinizer_for_entry(entry):
        """Auto assign scrutinizer when an exam entry is created."""
        if not entry or not entry.course_code:
            return
        assignments = DutyAssignment.query.filter_by(duty_type='scrutinizer', status='active').all()
        for assignment in assignments:
            if assignment.assigned_teacher_id and _assignment_matches_entry(assignment, entry):
                entry.assigned_scrutinizer_id = assignment.assigned_teacher_id
                db.session.commit()
                break

    def _has_section_conflict(course_code, new_section, exclude_entry_id=None):
        """Return an error message if the course already has conflicting sections assigned."""
        if not course_code:
            return None
        normalized_new = _normalize_section(new_section)
        query = ExamPaperEvaluation.query.filter(
            ExamPaperEvaluation.course_code == course_code,
            ExamPaperEvaluation.archived.is_(False)
        )
        if exclude_entry_id:
            query = query.filter(ExamPaperEvaluation.id != exclude_entry_id)
        existing_entries = query.all()
        existing_sections = {_normalize_section(entry.section) for entry in existing_entries}

        if 'FULL' in existing_sections:
            return 'This course already has an assigned marks entry. Delete the existing entry before adding another.'
        if normalized_new == 'FULL' and existing_sections:
            return 'This course already has section-level assignments. Delete them before assigning a full-course entry.'
        if normalized_new in existing_sections:
            return f'Section "{new_section or "Full"}" is already assigned. Delete the existing entry before reassigning.'
        if len(existing_sections) >= 2:
            return 'Both sections are already assigned to teachers. Delete an existing section before adding another.'
        return None

    @app.route('/exam-evaluation', methods=['GET', 'POST'])
    @login_required
    def exam_evaluation():
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        if request.method == 'POST':
            batch = (request.form.get('batch') or '').strip() or None
            academic_session = request.form.get('academic_session', '').strip() or None
            course_id = request.form.get('course_id', type=int)
            course_name = request.form.get('course_name')
            course_code = request.form.get('course_code')
            discipline = request.form.get('discipline', 'Law')
            school = request.form.get('school', 'Law')
            year = (request.form.get('year') or '').strip() or None
            term = (request.form.get('term') or '').strip() or None
            section = request.form.get('section')
            program_level = request.form.get('program_level', 'ug')

            # If course_id is provided, fetch course details from Course model
            try:
                from blueprints.course_management.models import Course, DutyAssignment
                if course_id and Course:
                    course = Course.query.get(course_id)
                    if course:
                        course_code = course.course_code
                        course_name = course.course_name
            except:
                pass

            if not course_name or not course_code:
                flash('Course name and course code are required.', 'danger')
            elif not academic_session or not year or not term:
                current_app.logger.warning(
                    'Rejected exam entry due to missing running context: '
                    f'academic_session={academic_session}, year={year}, term={term}, course_code={course_code}'
                )
                flash('Academic Session, Year, and Term are required to keep retake processing in the running term.', 'danger')
            else:
                # First check if semester is active
                try:
                    from utils.semester_utils import is_semester_active
                    if not is_semester_active(academic_session, year, term, batch=batch):
                        flash('This semester is not active. Please activate it in Active Semester Management first.', 'danger')
                        return redirect(url_for('exam_evaluation'))
                except Exception as e:
                    current_app.logger.error(f'Error checking active semester: {e}', exc_info=True)
                    # Continue with chief check if error
                
                chief_exists = DutyAssignment.query.filter_by(
                    duty_type='exam_committee_chief',
                    academic_session=academic_session,
                    year=year,
                    term=term,
                    status='active'
                ).first()
                if not chief_exists:
                    flash('No Exam Committee Chief is assigned for this session/year/term. Please assign a chief before submitting marksheets.', 'danger')
                    return redirect(url_for('exam_evaluation'))

                conflict_message = _has_section_conflict(course_code, section)
                if conflict_message:
                    flash(conflict_message, 'danger')
                    return redirect(url_for('exam_evaluation'))
                owner_teacher = _current_teacher()
                record = ExamPaperEvaluation(
                    course_name=course_name,
                    course_code=course_code,
                    academic_session=academic_session,
                    batch=batch,
                    discipline=discipline,
                    school=school,
                    year=year,
                    term=term,
                    section=section,
                    program_level=program_level,
                    owner_teacher_id=owner_teacher.id if owner_teacher else None
                )
                db.session.add(record)
                db.session.commit()
                _auto_assign_scrutinizer_for_entry(record)
                flash('Exam paper evaluation entry saved successfully!', 'success')
                return redirect(url_for('exam_evaluation'))

        current_teacher = _current_teacher()
        current_teacher_ids = _teacher_identity_ids(current_teacher)
        committee_assignment_entry_ids = set()

        from blueprints.class_management.models import ExamPaperEvaluatorAssignment
        from blueprints.course_management.models import StudentCourseRegistration, Course, DutyAssignment, CurriculumYearTerm

        teacher_assignments = []
        active_semester_keys = None
        if not is_admin(current_user):
            try:
                from utils.semester_utils import get_active_semesters, _normalize_year_term
                active_semesters = get_active_semesters(batch=None)
                active_semester_keys = {
                    (
                        str(sem.academic_session or '').strip().casefold(),
                        _normalize_year_term(sem.year),
                        _normalize_year_term(sem.term),
                    )
                    for sem in active_semesters
                }
            except Exception:
                active_semester_keys = None

        if current_teacher:
            # Simple rule: evaluator page follows committee assignments directly.
            try:
                teacher_assignments = ExamPaperEvaluatorAssignment.query.filter(
                    ExamPaperEvaluatorAssignment.assigned_teacher_id.in_(list(current_teacher_ids))
                ).all()

                if active_semester_keys is not None:
                    def _assignment_key(assignment):
                        try:
                            from utils.semester_utils import _normalize_year_term
                            return (
                                str(assignment.academic_session or '').strip().casefold(),
                                _normalize_year_term(assignment.year),
                                _normalize_year_term(assignment.term),
                            )
                        except Exception:
                            return (
                                str(assignment.academic_session or '').strip().casefold(),
                                str(assignment.year or '').strip().casefold(),
                                str(assignment.term or '').strip().casefold(),
                            )

                    teacher_assignments = [
                        assignment for assignment in teacher_assignments
                        if _assignment_key(assignment) in active_semester_keys
                    ]
                repair_made = False

                for assignment in teacher_assignments:
                    part_label = (assignment.part or '').strip().upper()
                    if part_label not in {'A', 'B'}:
                        continue

                    course = Course.query.get(assignment.course_id) if assignment.course_id else None
                    if not course:
                        continue

                    linked_entry = assignment.exam_paper_evaluation
                    linked_ok = bool(
                        linked_entry
                        and linked_entry.owner_teacher_id in current_teacher_ids
                        and not bool(linked_entry.archived)
                        and (linked_entry.course_code or '').strip() == (course.course_code or '').strip()
                        and (linked_entry.academic_session or '').strip() == (assignment.academic_session or '').strip()
                        and (linked_entry.year or '').strip() == (assignment.year or '').strip()
                        and (linked_entry.term or '').strip() == (assignment.term or '').strip()
                        and (linked_entry.section or '').strip() == f'Part {part_label}'
                    )
                    if linked_ok:
                        committee_assignment_entry_ids.add(linked_entry.id)
                        continue

                    existing_entry = ExamPaperEvaluation.query.filter(
                        ExamPaperEvaluation.owner_teacher_id.in_(list(current_teacher_ids)),
                        ExamPaperEvaluation.course_code == (course.course_code or '').strip(),
                        ExamPaperEvaluation.academic_session == (assignment.academic_session or '').strip(),
                        ExamPaperEvaluation.year == (assignment.year or '').strip(),
                        ExamPaperEvaluation.term == (assignment.term or '').strip(),
                        ExamPaperEvaluation.section == f'Part {part_label}',
                        ExamPaperEvaluation.archived.is_(False)
                    ).order_by(ExamPaperEvaluation.created_at.desc()).first()

                    if existing_entry:
                        assignment.exam_paper_evaluation_id = existing_entry.id
                        committee_assignment_entry_ids.add(existing_entry.id)
                        repair_made = True
                        continue

                    resolved_batch = ''
                    duty_ctx = DutyAssignment.query.filter(
                        DutyAssignment.status == 'active',
                        DutyAssignment.academic_session == (assignment.academic_session or '').strip(),
                        DutyAssignment.year == (assignment.year or '').strip(),
                        DutyAssignment.term == (assignment.term or '').strip(),
                        DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member']),
                        DutyAssignment.batch.isnot(None),
                        DutyAssignment.batch != ''
                    ).order_by(DutyAssignment.created_at.desc()).first()
                    if duty_ctx and duty_ctx.batch:
                        resolved_batch = duty_ctx.batch.strip()
                    if not resolved_batch:
                        cfg = CurriculumYearTerm.query.filter_by(
                            academic_session=(assignment.academic_session or '').strip(),
                            year=(assignment.year or '').strip(),
                            term=(assignment.term or '').strip()
                        ).filter(
                            CurriculumYearTerm.batch.isnot(None),
                            CurriculumYearTerm.batch != '',
                            CurriculumYearTerm.batch != 'None'
                        ).order_by(CurriculumYearTerm.updated_at.desc()).first()
                        if cfg and cfg.batch:
                            resolved_batch = cfg.batch.split(',')[0].strip() if ',' in cfg.batch else cfg.batch.strip()

                    created_entry = ExamPaperEvaluation(
                        course_name=course.course_name,
                        course_code=course.course_code,
                        academic_session=assignment.academic_session,
                        batch=resolved_batch or None,
                        year=assignment.year,
                        term=assignment.term,
                        section=f'Part {part_label}',
                        program_level=course.category or 'ug',
                        owner_teacher_id=current_teacher.id,
                        submitted_to_committee=False
                    )
                    db.session.add(created_entry)
                    db.session.flush()
                    assignment.exam_paper_evaluation_id = created_entry.id
                    committee_assignment_entry_ids.add(created_entry.id)
                    repair_made = True

                if repair_made:
                    db.session.commit()
            except Exception as sync_exc:
                db.session.rollback()
                current_app.logger.warning(
                    f'Failed to sync committee assignments for teacher_id={current_teacher.id}: {sync_exc}',
                    exc_info=True
                )

        if current_teacher and teacher_assignments:
            assigned_query = ExamPaperEvaluation.query.filter(
                ExamPaperEvaluation.id.in_(list(committee_assignment_entry_ids)),
            )
            # Even in committee-driven mode, keep only active semester rows.
            try:
                from utils.semester_utils import filter_by_active_semester
                if filter_by_active_semester and not is_admin(current_user):
                    assigned_query = filter_by_active_semester(
                        assigned_query,
                        ExamPaperEvaluation,
                        batch=None,
                        admin_override=False
                    )
            except ImportError:
                pass

            entries = assigned_query.filter(
                ExamPaperEvaluation.archived.is_(False)
            ).order_by(ExamPaperEvaluation.created_at.desc()).all()
            archived_entries = assigned_query.filter(
                ExamPaperEvaluation.archived.is_(True)
            ).order_by(ExamPaperEvaluation.created_at.desc()).all()
        else:
            base_query = ExamPaperEvaluation.query
            if current_teacher:
                base_query = base_query.filter(ExamPaperEvaluation.owner_teacher_id.in_(list(current_teacher_ids)))
            else:
                base_query = base_query.filter_by(owner_teacher_id=None)

            # Fallback (no committee assignment): keep legacy active semester filtering.
            try:
                from utils.semester_utils import filter_by_active_semester
                if filter_by_active_semester and not is_admin(current_user):
                    base_query = filter_by_active_semester(base_query, ExamPaperEvaluation, batch=None, admin_override=False)
            except ImportError:
                pass

            entries = base_query.filter_by(archived=False).order_by(ExamPaperEvaluation.created_at.desc()).all()
            archived_entries = base_query.filter_by(archived=True).order_by(ExamPaperEvaluation.created_at.desc()).all()

        evaluator_assigned_entry_ids = set(committee_assignment_entry_ids)
        scrutiny_entries = []
        scrutiny_invites_map = {}
        if current_teacher:
            owned_ids = {entry.id for entry in entries}
            base_scrutiny_query = ExamPaperEvaluation.query.filter(
                ExamPaperEvaluation.archived.is_(False),
                ExamPaperEvaluation.assigned_scrutinizer_id == current_teacher.id,
                ExamPaperEvaluation.submitted_to_committee.is_(True)
            )
            
            # Apply active semester filtering
            try:
                from utils.semester_utils import filter_by_active_semester
                if filter_by_active_semester and not is_admin(current_user):
                    base_scrutiny_query = filter_by_active_semester(
                        base_scrutiny_query, 
                        ExamPaperEvaluation, 
                        batch=None, 
                        admin_override=False
                    )
            except ImportError:
                pass
            
            scrutiny_entries = base_scrutiny_query.order_by(ExamPaperEvaluation.created_at.desc()).all()
            scrutiny_entries = [entry for entry in scrutiny_entries if entry.id not in owned_ids]
            
            # Get invite information for each scrutiny entry to check completion status
            if scrutiny_entries:
                entry_ids = [entry.id for entry in scrutiny_entries]
                invites = ExamScrutinizerInvite.query.filter(
                    ExamScrutinizerInvite.exam_entry_id.in_(entry_ids),
                    ExamScrutinizerInvite.scrutinizer_teacher_id == current_teacher.id,
                    ExamScrutinizerInvite.status == 'accepted'
                ).all()
                scrutiny_invites_map = {inv.exam_entry_id: inv for inv in invites}

        # Show "(Retake)" only when this entry maps to a merge-off retake registration.
        retake_remarks = {'retake', 're-retake', 're retake', 'reretake'}
        all_visible_entries = entries + archived_entries + scrutiny_entries
        for entry in all_visible_entries:
            entry.show_retake_label = False

        all_visible_entry_ids = [entry.id for entry in all_visible_entries if getattr(entry, 'id', None)]
        assignment_by_entry_id = {}
        if all_visible_entry_ids:
            all_assignments = ExamPaperEvaluatorAssignment.query.filter(
                ExamPaperEvaluatorAssignment.exam_paper_evaluation_id.isnot(None),
                ExamPaperEvaluatorAssignment.exam_paper_evaluation_id.in_(all_visible_entry_ids)
            ).all()
            assignment_by_entry_id = {
                assignment.exam_paper_evaluation_id: assignment
                for assignment in all_assignments
                if assignment.exam_paper_evaluation_id
            }

        assignment_course_ids = {
            assignment.course_id
            for assignment in assignment_by_entry_id.values()
            if getattr(assignment, 'course_id', None)
        }
        registrations = []
        if assignment_course_ids:
            registrations.extend(
                StudentCourseRegistration.query.filter(
                    StudentCourseRegistration.course_id.in_(assignment_course_ids),
                    StudentCourseRegistration.use_relevant_for_committee.is_(False)
                ).all()
            )

        registration_by_id = {}
        for registration in registrations:
            if registration and getattr(registration, 'id', None):
                registration_by_id[registration.id] = registration

        retake_course_keys = set()
        for registration in registration_by_id.values():
            remark_value = str(registration.remark or '').strip().lower()
            if remark_value not in retake_remarks:
                continue

            relevant_context = (
                str(registration.relevant_academic_session or '').strip(),
                str(registration.relevant_year or '').strip(),
                str(registration.relevant_term or '').strip()
            )
            running_context = (
                str(registration.academic_session or '').strip(),
                str(registration.year or '').strip(),
                str(registration.term or '').strip()
            )
            registration_context = relevant_context if all(relevant_context) else running_context
            if not all(registration_context):
                continue

            if registration.course_id:
                retake_course_keys.add((registration.course_id, *registration_context))

        for entry in all_visible_entries:
            assignment = assignment_by_entry_id.get(entry.id)
            if not assignment:
                continue

            assignment_context = (
                str(assignment.academic_session or '').strip(),
                str(assignment.year or '').strip(),
                str(assignment.term or '').strip()
            )
            if not all(assignment_context):
                continue

            if (assignment.course_id, *assignment_context) in retake_course_keys:
                entry.show_retake_label = True
        
        # Determine if current user is evaluator (owner) - if yes, hide scrutinizer info
        # Evaluators (owners) should not see scrutinizer info, but scrutinizers and admins should
        hide_scrutinizer_info = False
        if current_teacher:
            # is_admin is already imported at the top of the file
            is_admin_user = is_admin(current_user)
            is_head = hasattr(current_user, 'active_role') and current_user.active_role == 'head'
            
            # Admins and heads should always see scrutinizer info
            if is_admin_user or is_head:
                hide_scrutinizer_info = False
            # If user has entries as owner (evaluator), hide scrutinizer info
            # Note: entries are already filtered by owner_teacher_id, so if entries exist, user is evaluator
            elif entries:
                # User is an evaluator (owner) - hide scrutinizer info
                # Exception: If they are also a scrutinizer (have scrutiny_entries), they should see it
                # But since they're viewing their own entries page, they're primarily an evaluator, so hide it
                hide_scrutinizer_info = True
        
        teacher_map = {}
        teacher_ids = {e.assigned_scrutinizer_id for e in entries if e.assigned_scrutinizer_id}
        if teacher_ids:
            teacher_map = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(teacher_ids)).all()}
        
        # Get batches for dropdown
        batches = []
        try:
            from blueprints.student_management.models import Student
            if Student:
                all_batches = db.session.query(Student.batch).distinct().filter(Student.batch.isnot(None)).order_by(Student.batch.desc()).all()
                batches = [batch[0] for batch in all_batches]
        except:
            batches = []
        
        # Get curricula and year/term configurations
        curricula = []
        curriculum_configs = {}
        available_sessions = set()
        curriculum_configs_json = '{}'
        try:
            from blueprints.course_management.models import Curriculum, CurriculumYearTerm
            from utils.semester_utils import get_active_semesters
            
            # Get active semesters first
            active_semesters = get_active_semesters(batch=None)
            
            if not active_semesters:
                # No active semester - return empty curricula
                curricula = []
                curriculum_configs_json = '{}'
                available_sessions = []
            else:
                # Load all curricula
                all_curricula = Curriculum.query.order_by(Curriculum.created_at.desc()).all()
                
                # Filter curricula to only include those with active semester configs
                for curriculum in all_curricula:
                    configs_query = curriculum.year_term_configs.order_by(
                        CurriculumYearTerm.year.asc(),
                        CurriculumYearTerm.term.asc()
                    )
                    configs = []
                    for config in configs_query.all():
                        # Check if this config matches any active semester
                        matches_active = False
                        for sem in active_semesters:
                            if (config.academic_session == sem.academic_session and
                                config.year == sem.year and
                                config.term == sem.term):
                                # Check batch match: active sem batch None means all batches, or exact match
                                if sem.batch is None or config.batch == sem.batch:
                                    matches_active = True
                                    break
                        
                        if matches_active:
                            configs.append({
                                'curriculum_id': curriculum.id,
                                'year': config.year,
                                'term': config.term,
                                'batch': config.batch,
                                'academic_session': config.academic_session
                            })
                            if config.academic_session:
                                available_sessions.add(config.academic_session)
                    
                    # Only add curriculum if it has at least one active semester config
                    if configs:
                        curricula.append(curriculum)
                        curriculum_configs[curriculum.id] = configs
                
                curriculum_configs_json = json.dumps(curriculum_configs)
        except Exception as e:
            current_app.logger.error(f'Error loading curricula: {e}', exc_info=True)
            curricula = []
            curriculum_configs_json = '{}'
        available_sessions = sorted(available_sessions)
        
        return render_template('exam_evaluation.html',
                               entries=entries,
                               archived_entries=archived_entries,
                               teacher_map=teacher_map,
                               batches=batches,
                               academic_sessions=available_sessions,
                               curricula=curricula,
                               curriculum_configs_json=curriculum_configs_json,
                               scrutiny_entries=scrutiny_entries,
                               scrutiny_invites_map=scrutiny_invites_map,
                               hide_scrutinizer_info=hide_scrutinizer_info,
                               evaluator_assigned_entry_ids=evaluator_assigned_entry_ids)

    @app.route('/exam-evaluation/<int:entry_id>/submit-to-committee', methods=['POST'])
    @login_required
    def submit_exam_entry_to_committee(entry_id):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        entry = ExamPaperEvaluation.query.get_or_404(entry_id)
        current_teacher = _current_teacher()
        current_teacher_ids = _teacher_identity_ids(current_teacher)
        if not current_teacher or entry.owner_teacher_id not in current_teacher_ids:
            flash('Only the entry owner can submit to the Exam Committee.', 'danger')
            return redirect(url_for('exam_evaluation'))
        if entry.submitted_to_committee:
            flash('This marksheet has already been submitted to the Exam Committee.', 'info')
            return redirect(url_for('exam_evaluation'))
        entry.submitted_to_committee = True
        entry.submitted_at = datetime.utcnow()
        db.session.commit()
        flash('Marksheet submitted to Exam Committee for scrutiny.', 'success')
        return redirect(url_for('exam_evaluation'))

    @app.route('/exam-evaluation/<int:entry_id>/unsubmit', methods=['POST'])
    @login_required
    def unsubmit_exam_entry(entry_id):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        entry = ExamPaperEvaluation.query.get_or_404(entry_id)
        current_teacher = _current_teacher()
        current_teacher_ids = _teacher_identity_ids(current_teacher)
        if not current_teacher or entry.owner_teacher_id not in current_teacher_ids:
            flash('Only the entry owner can unsubmit a marksheet.', 'danger')
            return redirect(url_for('exam_evaluation'))
        if not entry.submitted_to_committee:
            flash('This marksheet has not been submitted yet.', 'info')
            return redirect(url_for('exam_evaluation'))
        entry.submitted_to_committee = False
        entry.submitted_at = None
        db.session.commit()
        flash('Submission withdrawn. You can edit the marksheet again.', 'success')
        return redirect(url_for('exam_evaluation'))

    @app.route('/exam-evaluation/<int:entry_id>/edit', methods=['GET', 'POST'])
    @login_required
    def edit_exam_evaluation(entry_id):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        entry = ExamPaperEvaluation.query.get_or_404(entry_id)
        if request.method == 'POST':
            course_code = request.form.get('course_code')
            section = request.form.get('section')
            conflict_message = _has_section_conflict(course_code, section, exclude_entry_id=entry.id) if course_code else None
            if conflict_message:
                flash(conflict_message, 'danger')
                return redirect(url_for('exam_evaluation'))
            entry.course_name = request.form.get('course_name')
            entry.course_code = request.form.get('course_code')
            entry.batch = request.form.get('batch')
            entry.academic_session = request.form.get('academic_session')
            entry.discipline = request.form.get('discipline')
            entry.school = request.form.get('school')
            entry.year = request.form.get('year')
            entry.term = request.form.get('term')
            entry.section = request.form.get('section')
            entry.program_level = request.form.get('program_level', entry.program_level)
            db.session.commit()
            flash('Entry updated successfully.', 'success')
            return redirect(url_for('exam_evaluation'))
        return render_template('exam_evaluation_edit.html', entry=entry)

    @app.route('/exam-evaluation/<int:entry_id>/delete', methods=['POST'])
    @login_required
    def delete_exam_evaluation(entry_id):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        entry = ExamPaperEvaluation.query.get_or_404(entry_id)
        invites = entry.scrutinizer_invites.all() if hasattr(entry, 'scrutinizer_invites') else []
        for invite in invites:
            db.session.delete(invite)
        db.session.delete(entry)
        db.session.commit()
        flash('Entry deleted successfully.', 'success')
        return redirect(url_for('exam_evaluation'))

    @app.route('/exam-evaluation/<int:entry_id>/archive', methods=['POST'])
    @login_required
    def archive_exam_evaluation(entry_id):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        entry = ExamPaperEvaluation.query.get_or_404(entry_id)
        undo = request.form.get('undo')
        try:
            entry.archived = False if undo else True
            db.session.commit()
            if undo:
                flash('Entry restored successfully.', 'success')
            else:
                flash('Entry archived successfully.', 'success')
        except Exception as exc:
            db.session.rollback()
            flash(f'Failed to update archive status: {exc}', 'danger')
        return redirect(url_for('exam_evaluation'))
    
    def _flatten_marks_payload(data):
        questions = data.get('questions') or []
        headers = ['Serial', 'Code']
        question_keys = []

        for question in questions:
            q_label = question.get('label', '')
            parts = question.get('parts') or []
            if not parts:
                headers.append(q_label)
                question_keys.append((q_label, None))
                continue
            if len(parts) == 1:
                part_label = parts[0].get('label', 'Marks')
                headers.append(q_label)
                question_keys.append((q_label, part_label))
            else:
                for part in parts:
                    part_label = part.get('label', '')
                    headers.append(f"{q_label} ({part_label})")
                    question_keys.append((q_label, part_label))

        headers.append('Total')
        headers.append('Student ID')

        rows = []
        for index, row in enumerate(data.get('rows', []), start=1):
            row_values = [index, row.get('code', '')]
            marks_map = row.get('marks') or {}
            for question_label, part_label in question_keys:
                question_marks = marks_map.get(question_label, {})
                if isinstance(question_marks, dict):
                    value = question_marks.get(part_label, '')
                else:
                    value = question_marks
                row_values.append(value)
            row_values.append(row.get('total', ''))
            row_values.append(row.get('student_id', ''))
            rows.append(row_values)

        return headers, rows

    def _generate_marks_pdf(entry, data):
        headers, rows = _flatten_marks_payload(data)

        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            leftMargin=18 * mm,
            rightMargin=18 * mm,
        )
        styles = getSampleStyleSheet()
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Heading1'],
            alignment=1,
            fontSize=16,
            leading=18,
            textColor=colors.HexColor('#0a3d62'),
            spaceAfter=4,
        )
        info_label_style = ParagraphStyle(
            'InfoLabel',
            parent=styles['Normal'],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#0a3d62'),
        )
        info_value_style = ParagraphStyle(
            'InfoValue',
            parent=styles['Normal'],
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#111111'),
        )

        elements = []
        elements.append(Paragraph('Khulna University', header_style))
        elements.append(Spacer(1, 6))

        info_table = Table(
            [
                [
                    Paragraph('Course', info_label_style),
                    Paragraph(entry.course_name or '', info_value_style),
                    Paragraph('Course Code', info_label_style),
                    Paragraph(entry.course_code or '', info_value_style),
                ],
                [
                    Paragraph('Discipline', info_label_style),
                    Paragraph(entry.discipline or '', info_value_style),
                    Paragraph('School', info_label_style),
                    Paragraph(entry.school or '', info_value_style),
                ],
                [
                    Paragraph('Year', info_label_style),
                    Paragraph(entry.year or '', info_value_style),
                    Paragraph('Term', info_label_style),
                    Paragraph(entry.term or '', info_value_style),
                ],
                [
                    Paragraph('Section', info_label_style),
                    Paragraph(entry.section or '', info_value_style),
                    Paragraph('', info_label_style),
                    Paragraph('', info_value_style),
                ],
            ],
            colWidths=[32 * mm, 46 * mm, 32 * mm, 46 * mm],
        )
        info_table.setStyle(
            TableStyle(
                [
                    ('GRID', (0, 0), (-1, -1), 0.25, colors.black),
                    ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('SPAN', (2, 3), (3, 3)),
                ]
            )
        )
        elements.append(info_table)
        elements.append(Spacer(1, 12))

        table_data = [headers] + rows
        total_index = headers.index('Total') if 'Total' in headers else max(len(headers) - 2, 0)
        table_header_style = ParagraphStyle(
            'TableHeader',
            parent=styles['Normal'],
            fontSize=8,
            leading=10,
            alignment=1,
            wordWrap='CJK',
        )
        table_cell_center = ParagraphStyle(
            'TableCellCenter',
            parent=styles['Normal'],
            fontSize=8,
            leading=10,
            alignment=1,
            wordWrap='CJK',
        )
        table_cell_left = ParagraphStyle(
            'TableCellLeft',
            parent=styles['Normal'],
            fontSize=8,
            leading=10,
            alignment=0,
            wordWrap='CJK',
        )
        table_cell_bold = ParagraphStyle(
            'TableCellBold',
            parent=table_cell_center,
            fontName='Helvetica-Bold',
        )

        def shorten_header(text):
            if not text:
                return ''
            value = str(text)
            lower = value.lower()
            if lower.startswith('question'):
                remainder = value[8:].lstrip()
                return f"Q {remainder}".strip()
            return value

        formatted_table_data = []
        for row_index, row_values in enumerate(table_data):
            formatted_row = []
            for col_index, value in enumerate(row_values):
                text = '' if value is None else str(value)
                if row_index == 0:
                    text = shorten_header(text)
                if row_index == 0:
                    formatted_row.append(Paragraph(text, table_header_style))
                else:
                    if col_index == 0 or col_index == len(row_values) - 1:
                        formatted_row.append(Paragraph(text, table_cell_left))
                    elif col_index == total_index:
                        formatted_row.append(Paragraph(text, table_cell_bold))
                    else:
                        formatted_row.append(Paragraph(text, table_cell_center))
            formatted_table_data.append(formatted_row)
        table_data = formatted_table_data
        available_width = doc.width
        weights = []
        for header in headers:
            lowered = str(header).strip().lower()
            if lowered in {'serial', '#'}:
                weights.append(0.7)
            elif lowered == 'code':
                weights.append(1.4)
            elif lowered == 'student id':
                weights.append(1.6)
            elif lowered == 'total':
                weights.append(1.2)
            else:
                weights.append(1.0)
        total_weight = sum(weights) or 1
        col_widths = [available_width * (w / total_weight) for w in weights]
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table_style = [
            ('GRID', (0, 0), (-1, -1), 0.25, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.whitesmoke),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('ALIGN', (0, 1), (0, -1), 'LEFT'),
            ('ALIGN', (-1, 1), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 1), (-2, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
            ('FONTNAME', (total_index, 0), (total_index, -1), 'Helvetica-Bold'),
        ]
        table.setStyle(TableStyle(table_style))
        elements.append(table)
        elements.append(Spacer(1, 18))

        signature_table = Table(
            [
                ['Name of Examiner:', '', 'Signature Date:', ''],
                ['Name of Scrutinizer:', '', 'Signature Date:', ''],
            ],
            colWidths=[30 * mm, 58 * mm, 26 * mm, 60 * mm],
            rowHeights=[18 * mm, 18 * mm],
        )
        signature_table.setStyle(
            TableStyle(
                [
                    ('BOX', (0, 0), (-1, -1), 0, colors.white),
                    ('INNERGRID', (0, 0), (-1, -1), 0, colors.white),
                    ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ]
            )
        )
        elements.append(signature_table)

        class NumberedCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []

            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                total_pages = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self.draw_page_number(total_pages)
                    canvas.Canvas.showPage(self)
                canvas.Canvas.save(self)

            def draw_page_number(self, page_count):
                page = self.getPageNumber()
                label = f"Page {page} of {page_count}"
                self.setFont("Helvetica", 9)
                width, _ = self._pagesize
                self.drawRightString(width - 18 * mm, 12 * mm, label)

        doc.build(elements, canvasmaker=NumberedCanvas)
        buffer.seek(0)
        return buffer

    def _generate_marks_excel(entry, data):
        headers, rows = _flatten_marks_payload(data)

        if Workbook is None:
            csv_buffer = StringIO()
            writer = csv.writer(csv_buffer)
            writer.writerow(headers)
            for row in rows:
                writer.writerow(row)
            byte_buffer = BytesIO(csv_buffer.getvalue().encode('utf-8-sig'))
            byte_buffer.seek(0)
            return byte_buffer, 'text/csv', 'csv'

        wb = Workbook()
        ws = wb.active
        ws.title = 'Marks'

        title_font = Font(bold=True, size=14) if Font else None
        header_font = Font(bold=True) if Font else None
        center = Alignment(horizontal='center', vertical='center') if Alignment else None
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin'),
        ) if Border and Side else None

        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        ws.cell(row=1, column=1, value='Khulna University')
        if title_font:
            ws.cell(row=1, column=1).font = title_font
        if center:
            ws.cell(row=1, column=1).alignment = center

        metadata_rows = [
            ('Course', entry.course_name or '', 'Course Code', entry.course_code or ''),
            ('Discipline', entry.discipline or '', 'School', entry.school or ''),
            ('Year', entry.year or '', 'Term', entry.term or ''),
            ('Section', entry.section or '', '', ''),
        ]

        for meta in metadata_rows:
            ws.append(meta)
            row_idx = ws.max_row
            if header_font:
                ws.cell(row=row_idx, column=1).font = header_font
                ws.cell(row=row_idx, column=3).font = header_font

        ws.append([])

        ws.append(headers)
        header_row = ws.max_row
        if header_font and center:
            for col_idx in range(1, len(headers) + 1):
                cell = ws.cell(row=header_row, column=col_idx)
                cell.font = header_font
                cell.alignment = center
                if border:
                    cell.border = border

        for row_values in rows:
            ws.append(row_values)
            row_idx = ws.max_row
            for col_idx in range(1, len(headers) + 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                if center:
                    cell.alignment = center
                if border:
                    cell.border = border

        # Adjust column widths
        for idx, header in enumerate(headers, start=1):
            length = max(len(str(header)), 10)
            for row_values in rows:
                length = max(length, len(str(row_values[idx - 1])))
            column_letter = get_column_letter(idx) if get_column_letter else chr(64 + idx)
            ws.column_dimensions[column_letter].width = length + 2

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'xlsx'

    @app.route('/exam-evaluation/<int:entry_id>/marks', methods=['GET', 'POST'])
    @login_required
    def exam_marks_entry(entry_id):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        entry = ExamPaperEvaluation.query.get_or_404(entry_id)
        # Refresh to ensure we have the latest data from database
        db.session.refresh(entry)
        role = request.args.get('role')

        current_teacher = _current_teacher()
        if current_teacher:
            if entry.assigned_scrutinizer_id == current_teacher.id:
                role = 'scrutinizer'
            elif entry.owner_teacher_id == current_teacher.id:
                role = 'evaluator'

        if role not in {'evaluator', 'scrutinizer'}:
            role = 'evaluator'
        
        # Prevent evaluator from editing marks if entry is submitted
        if role == 'evaluator' and entry.submitted_to_committee:
            flash('Cannot edit marks after submission. Please unsubmit the entry first.', 'warning')
            return redirect(url_for('exam_evaluation'))

        if request.method == 'POST':
            # Prevent evaluator from saving marks if entry is submitted
            if role == 'evaluator' and entry.submitted_to_committee:
                flash('Cannot save marks after submission. Please unsubmit the entry first.', 'warning')
                return redirect(url_for('exam_evaluation'))
            
            action = request.form.get('action_type', 'save')
            payload = request.form.get('marks_payload')
            if not payload:
                flash('No marks data received.', 'warning')
                return redirect(url_for('exam_marks_entry', entry_id=entry_id, role=role))

            try:
                data = json.loads(payload)
            except json.JSONDecodeError:
                flash('Failed to process marks data (invalid format).', 'danger')
                return redirect(url_for('exam_marks_entry', entry_id=entry_id, role=role))

            if action == 'pdf':
                pdf_buffer = _generate_marks_pdf(entry, data)
                # Format: Course Code_Course Name_Part_Marksheet.pdf
                import re
                def sanitize_filename(text):
                    """Remove or replace special characters for safe filename"""
                    if not text:
                        return ''
                    # Replace spaces with underscores
                    text = text.replace(' ', '_')
                    # Remove or replace special characters that might cause issues
                    text = re.sub(r'[<>:"/\\|?*]', '', text)
                    # Remove multiple consecutive underscores
                    text = re.sub(r'_+', '_', text)
                    return text.strip('_')
                
                course_code = sanitize_filename(entry.course_code or 'exam')
                course_name = sanitize_filename(entry.course_name or '')
                part = sanitize_filename(entry.section if entry.section else 'Full')
                filename = f"{course_code}_{course_name}_Part_{part}_Marksheet.pdf"
                pdf_data = pdf_buffer.getvalue()
                return Response(
                    pdf_data,
                    mimetype='application/pdf',
                    headers={
                        'Content-Disposition': f'attachment; filename=\"{filename}\"',
                        'Content-Length': str(len(pdf_data)),
                        'Cache-Control': 'no-cache, no-store, must-revalidate',
                        'Pragma': 'no-cache',
                        'Expires': '0',
                    },
                )
            if action == 'excel':
                excel_buffer, mimetype, extension = _generate_marks_excel(entry, data)
                filename = f"{entry.course_code or 'exam'}_marksheet.{extension}"
                excel_data = excel_buffer.getvalue()
                return Response(
                    excel_data,
                    mimetype=mimetype,
                    headers={
                        'Content-Disposition': f'attachment; filename=\"{filename}\"',
                        'Content-Length': str(len(excel_data)),
                        'Cache-Control': 'no-cache, no-store, must-revalidate',
                        'Pragma': 'no-cache',
                        'Expires': '0',
                    },
                )

            # Validate for duplicate Student IDs before saving
            rows = data.get('rows', [])
            student_ids_seen = {}
            duplicates = []
            
            for idx, row in enumerate(rows):
                student_id = str(row.get('student_id', '')).strip()
                if student_id:  # Only check non-empty Student IDs
                    if student_id in student_ids_seen:
                        # Found duplicate
                        first_row = student_ids_seen[student_id]
                        duplicates.append({
                            'student_id': student_id,
                            'row1': first_row + 1,  # 1-based row number
                            'row2': idx + 1
                        })
                    else:
                        student_ids_seen[student_id] = idx
            
            if duplicates:
                error_message = 'Duplicate Student ID found:\n\n'
                for dup in duplicates:
                    error_message += f"Student ID \"{dup['student_id']}\" appears in row {dup['row1']} and row {dup['row2']}\n"
                error_message += '\nPlease fix duplicate Student IDs before saving.'
                flash(error_message, 'error')
                return redirect(url_for('exam_marks_entry', entry_id=entry_id, role=role))

            entry.marks_data = json.dumps(data)
            db.session.commit()
            flash(f"Marks entry saved for {len(rows)} students.", 'success')
            return redirect(url_for('exam_marks_entry', entry_id=entry_id, role=role))

        initial_data = {}
        if entry.marks_data:
            try:
                initial_data = json.loads(entry.marks_data)
            except json.JSONDecodeError:
                initial_data = {}

        # Auto-populate discipline and school from teacher's institute if missing
        # Also normalize existing values (e.g., "Law Discipline, KU" -> "Law")
        needs_update = False
        
        # Normalize discipline: extract first word if it contains multiple words
        if entry.discipline and ' ' in entry.discipline:
            discipline_parts = entry.discipline.split()
            if discipline_parts:
                entry.discipline = discipline_parts[0]
                needs_update = True
        elif not entry.discipline:
            if entry.owner_teacher and entry.owner_teacher.institute:
                institute = entry.owner_teacher.institute
                institute_parts = institute.split()
                if institute_parts:
                    entry.discipline = institute_parts[0]
                else:
                    entry.discipline = 'Law'
            else:
                entry.discipline = 'Law'
            needs_update = True
        
        # Normalize school: extract first word if it contains multiple words or is "Law Discipline, KU"
        if entry.school:
            if ' ' in entry.school or entry.school == 'Law Discipline, KU':
                school_parts = entry.school.split()
                if school_parts:
                    entry.school = school_parts[0]
                else:
                    entry.school = 'Law'
                needs_update = True
        else:
            if entry.owner_teacher and entry.owner_teacher.institute:
                institute = entry.owner_teacher.institute
                institute_parts = institute.split()
                if institute_parts:
                    entry.school = institute_parts[0]
                else:
                    entry.school = 'Law'
            else:
                entry.school = 'Law'
            needs_update = True
        
        if needs_update:
            db.session.commit()

        return render_template('exam_evaluation_marks.html', entry=entry, role=role, initial_marks_data=initial_data)

    @app.route('/exam-evaluation/<int:entry_id>/marks/auto-save', methods=['POST'])
    @login_required
    def exam_marks_auto_save(entry_id):
        """Auto-save exam marks via AJAX"""
        if not has_teacher_privileges(current_user):
            return jsonify({'success': False, 'message': 'Not authorized'}), 403
        try:
            entry = ExamPaperEvaluation.query.get_or_404(entry_id)
            
            # Handle both JSON and form data
            if request.is_json:
                json_data = request.get_json()
                # Extract marks_payload from JSON if it exists
                payload = json_data.get('marks_payload')
                if not payload:
                    return jsonify({'success': False, 'message': 'No marks data received'}), 400
                try:
                    data = json.loads(payload)
                except (json.JSONDecodeError, TypeError):
                    # If payload is already a dict, use it directly
                    if isinstance(payload, dict):
                        data = payload
                    else:
                        return jsonify({'success': False, 'message': 'Invalid JSON format'}), 400
            else:
                payload = request.form.get('marks_payload')
                if not payload:
                    return jsonify({'success': False, 'message': 'No marks data received'}), 400
                try:
                    data = json.loads(payload)
                except json.JSONDecodeError:
                    return jsonify({'success': False, 'message': 'Invalid JSON format'}), 400
            
            # Save to database
            entry.marks_data = json.dumps(data)
            db.session.add(entry)  # Ensure entry is in session
            db.session.commit()
            db.session.refresh(entry)  # Refresh to ensure data is persisted
            return jsonify({'success': True, 'message': f'Marks saved automatically for {len(data.get("rows", []))} students'})
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error auto-saving exam marks for entry {entry_id}: {e}")
            return jsonify({'success': False, 'message': f'Error saving: {str(e)}'}), 500

    @app.route('/exam-evaluation/<int:entry_id>/scrutinizer', methods=['GET', 'POST'])
    @login_required
    def exam_assign_scrutinizer(entry_id):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        entry = ExamPaperEvaluation.query.get_or_404(entry_id)
        current_teacher = _current_teacher()

        if not current_teacher:
            flash('No teacher profile found for your account.', 'warning')
            return redirect(url_for('exam_evaluation'))

        if entry.owner_teacher_id and entry.owner_teacher_id != current_teacher.id:
            flash('Only the owner of this entry can manage scrutinizer assignments.', 'danger')
            return redirect(url_for('exam_evaluation'))

        if not entry.owner_teacher_id:
            entry.owner_teacher_id = current_teacher.id
            db.session.commit()

        if request.method == 'POST':
            action = request.form.get('action')
            try:
                if action == 'invite':
                    scrutinizer_id = int(request.form.get('scrutinizer_teacher_id', 0))
                    remarks = request.form.get('remarks')
                    if scrutinizer_id == current_teacher.id:
                        flash('You cannot invite yourself as scrutinizer.', 'warning')
                    else:
                        existing = ExamScrutinizerInvite.query.filter_by(
                            exam_entry_id=entry.id,
                            scrutinizer_teacher_id=scrutinizer_id
                        ).order_by(ExamScrutinizerInvite.created_at.desc()).first()
                        if existing and existing.status in {'invited', 'accepted'}:
                            flash('This teacher already has an active invitation.', 'info')
                        else:
                            invite = ExamScrutinizerInvite(
                                exam_entry_id=entry.id,
                                inviter_teacher_id=current_teacher.id,
                                scrutinizer_teacher_id=scrutinizer_id,
                                remarks=remarks or ''
                            )
                            db.session.add(invite)
                            db.session.commit()
                            flash('Scrutinizer invitation sent successfully.', 'success')
                elif action == 'cancel':
                    invite_id = int(request.form.get('invite_id', 0))
                    invite = ExamScrutinizerInvite.query.get_or_404(invite_id)
                    if invite.exam_entry_id != entry.id:
                        flash('Invalid invitation.', 'danger')
                    else:
                        invite.status = 'cancelled'
                        invite.responded_at = datetime.utcnow()
                        if entry.assigned_scrutinizer_id == invite.scrutinizer_teacher_id:
                            entry.assigned_scrutinizer_id = None
                        db.session.commit()
                        flash('Invitation cancelled.', 'info')
                elif action == 'clear_assigned':
                    entry.assigned_scrutinizer_id = None
                    db.session.commit()
                    flash('Assigned scrutinizer removed.', 'info')
            except Exception as exc:
                db.session.rollback()
                flash(f'Failed to process request: {exc}', 'danger')
            return redirect(url_for('exam_assign_scrutinizer', entry_id=entry_id))

        invites = ExamScrutinizerInvite.query.filter_by(exam_entry_id=entry.id).order_by(ExamScrutinizerInvite.created_at.desc()).all()
        teacher_ids = {inv.scrutinizer_teacher_id for inv in invites if inv.scrutinizer_teacher_id}
        if entry.owner_teacher_id:
            teacher_ids.add(entry.owner_teacher_id)

        from role_utils import get_teachers_excluding_head
        teacher_map = {t.id: t for t in get_teachers_excluding_head()}

        eligible_users = User.query.filter(User.role != ADMIN_ROLE).order_by(User.full_name.asc()).all()
        available_teachers = []
        seen_teacher_ids = set()
        for user in eligible_users:
            if current_user and user.id == current_user.id:
                continue
            teacher = _ensure_teacher_for_user(user)
            if not teacher or teacher.id == current_teacher.id:
                continue
            if teacher.id in seen_teacher_ids:
                continue
            seen_teacher_ids.add(teacher.id)
            teacher_map.setdefault(teacher.id, teacher)
            available_teachers.append(teacher)

        for teacher_id in list(teacher_ids):
            if teacher_id and teacher_id not in teacher_map:
                teacher = Teacher.query.get(teacher_id)
                if teacher:
                    teacher_map[teacher_id] = teacher

        if not available_teachers:
            flash('No other teachers are available to invite yet. Please add another teacher from Class Management first.', 'info')
        assigned_scrutinizer = teacher_map.get(entry.assigned_scrutinizer_id) if entry.assigned_scrutinizer_id else None

        return render_template(
            'exam_evaluation_scrutinizer.html',
            entry=entry,
            available_teachers=available_teachers,
            invites=invites,
            teacher_map=teacher_map,
            assigned_scrutinizer=assigned_scrutinizer
        )

    @app.route('/exam-evaluation/scrutinizer/invitations/<int:invite_id>/<string:action>', methods=['POST'])
    @login_required
    def exam_scrutinizer_invitation_action(invite_id, action):
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        invite = ExamScrutinizerInvite.query.get_or_404(invite_id)
        teacher = _current_teacher()
        if not teacher or teacher.id != invite.scrutinizer_teacher_id:
            flash('You are not authorized to respond to this invitation.', 'danger')
            return redirect(url_for('class_management.my_invitations'))

        # Handle complete/incomplete toggle
        if action == 'toggle_complete':
            if invite.status != 'accepted':
                flash('Only accepted invitations can be marked as complete/incomplete.', 'warning')
                # Redirect to next parameter or default to invitations page
                next_url = request.form.get('next') or request.args.get('next')
                if next_url:
                    return redirect(next_url)
                return redirect(url_for('class_management.my_invitations'))
            try:
                was_complete = invite.is_complete
                invite.is_complete = not invite.is_complete
                
                # If marking as incomplete (was True, now False), clear marks from Result Management
                if not invite.is_complete and was_complete:
                    # #region agent log
                    import json as json_module
                    try:
                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                            f.write(json_module.dumps({'location': 'app.py:toggle_incomplete', 'message': 'Toggle incomplete triggered clear', 'data': {'invite_id': invite_id, 'exam_entry_id': invite.exam_entry_id, 'was_complete': was_complete, 'now_complete': False}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'M'}) + '\n')
                    except: pass
                    # #endregion
                    try:
                        from blueprints.result_management.routes import clear_exam_marks_from_result_management
                        clear_result = clear_exam_marks_from_result_management(invite.exam_entry_id)
                        
                        # #region agent log
                        try:
                            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                f.write(json_module.dumps({'location': 'app.py:clear_result', 'message': 'Clear result received', 'data': {'success': clear_result.get('success'), 'marks_cleared': clear_result.get('marks_cleared'), 'message': clear_result.get('message'), 'errors': clear_result.get('errors', [])}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'M'}) + '\n')
                        except: pass
                        # #endregion
                        
                        if clear_result['success']:
                            current_app.logger.info(f'Auto-clear successful for exam_entry {invite.exam_entry_id}: {clear_result["message"]}')
                            flash(f'Status updated to Incomplete. Marks cleared from Result Management. {clear_result["marks_cleared"]} marks cleared.', 'info')
                        else:
                            # Log error but don't fail the toggle action
                            current_app.logger.error(f'Auto-clear failed for exam_entry {invite.exam_entry_id}: {clear_result["message"]}. Errors: {clear_result.get("errors", [])}')
                            flash(f'Status updated to Incomplete, but marks clear had issues. Please check Result Management manually.', 'warning')
                    except Exception as clear_exc:
                        # Log error but don't fail the toggle action
                        current_app.logger.error(f'Error during auto-clear for exam_entry {invite.exam_entry_id}: {str(clear_exc)}', exc_info=True)
                        flash(f'Status updated to Incomplete, but marks clear failed. Please check Result Management manually.', 'warning')
                
                # If marking as complete (was False, now True), sync marks to Result Management
                elif invite.is_complete and not was_complete:
                    # #region agent log
                    import json as json_module
                    try:
                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                            f.write(json_module.dumps({'location': 'app.py:toggle_complete', 'message': 'Toggle complete triggered sync', 'data': {'invite_id': invite_id, 'exam_entry_id': invite.exam_entry_id, 'was_complete': was_complete, 'now_complete': True}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'A'}) + '\n')
                    except: pass
                    # #endregion
                    try:
                        from blueprints.result_management.routes import sync_exam_marks_to_result_management
                        sync_result = sync_exam_marks_to_result_management(invite.exam_entry_id)
                        
                        # #region agent log
                        try:
                            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                f.write(json_module.dumps({'location': 'app.py:sync_result', 'message': 'Sync result received', 'data': {'success': sync_result.get('success'), 'marks_updated': sync_result.get('marks_updated'), 'message': sync_result.get('message'), 'errors': sync_result.get('errors', [])}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'A'}) + '\n')
                        except: pass
                        # #endregion
                        
                        if sync_result['success']:
                            current_app.logger.info(f'Auto-sync successful for exam_entry {invite.exam_entry_id}: {sync_result["message"]}')
                            # Optionally add a flash message, but keep it brief
                            if sync_result.get('session_created') or sync_result.get('subject_created'):
                                flash(f'Status updated to Complete. Marks synced to Result Management. {sync_result["marks_updated"]} marks updated.', 'success')
                            else:
                                flash(f'Status updated to Complete. Marks synced: {sync_result["marks_updated"]} marks updated.', 'success')
                        else:
                            # Log error but don't fail the toggle action
                            current_app.logger.error(f'Auto-sync failed for exam_entry {invite.exam_entry_id}: {sync_result["message"]}. Errors: {sync_result.get("errors", [])}')
                            flash(f'Status updated to Complete, but marks sync had issues. Please check Result Management manually.', 'warning')
                    except Exception as sync_exc:
                        # Log error but don't fail the toggle action
                        current_app.logger.error(f'Error during auto-sync for exam_entry {invite.exam_entry_id}: {str(sync_exc)}', exc_info=True)
                        flash(f'Status updated to Complete, but marks sync failed. Please check Result Management manually.', 'warning')
                
                db.session.commit()
                # Flash messages are already shown by sync/clear functions above
            except Exception as exc:
                db.session.rollback()
                flash(f'Failed to update status: {exc}', 'danger')
            # Redirect to next parameter or default to invitations page
            next_url = request.form.get('next') or request.args.get('next')
            if next_url:
                return redirect(next_url)
            return redirect(url_for('class_management.my_invitations'))

        if action == 'cancel' and invite.status == 'cancelled':
            flash('This invitation has already been cancelled.', 'info')
            return redirect(url_for('class_management.my_invitations'))

        if action != 'cancel' and invite.status not in {'invited', 'accepted'}:
            flash('This invitation has already been processed.', 'info')
            return redirect(url_for('class_management.my_invitations'))

        try:
            if action == 'accept':
                invite.status = 'accepted'
                invite.responded_at = datetime.utcnow()
                invite.exam_entry.assigned_scrutinizer_id = teacher.id
                other_pending = ExamScrutinizerInvite.query.filter(
                    ExamScrutinizerInvite.exam_entry_id == invite.exam_entry_id,
                    ExamScrutinizerInvite.id != invite.id,
                    ExamScrutinizerInvite.status == 'accepted'
                ).all()
                for other in other_pending:
                    other.status = 'cancelled'
                    other.responded_at = datetime.utcnow()
            elif action == 'decline':
                invite.status = 'declined'
                invite.responded_at = datetime.utcnow()
                if invite.exam_entry.assigned_scrutinizer_id == teacher.id:
                    invite.exam_entry.assigned_scrutinizer_id = None
            elif action == 'cancel':
                if invite.exam_entry.assigned_scrutinizer_id == teacher.id:
                    invite.exam_entry.assigned_scrutinizer_id = None
                invite.status = 'cancelled'
                invite.responded_at = datetime.utcnow()
            else:
                flash('Unknown action.', 'danger')
                return redirect(url_for('class_management.my_invitations'))
            db.session.commit()
            flash('Invitation updated.', 'success')
        except Exception as exc:
            db.session.rollback()
            flash(f'Failed to update invitation: {exc}', 'danger')

        return redirect(url_for('class_management.my_invitations'))

    @app.route('/test-simple-pdf')
    def test_simple_pdf():
        """Simple PDF test endpoint"""
        try:
            from flask import Response
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            
            elements = []
            elements.append(Paragraph("Test PDF", styles['Title']))
            elements.append(Paragraph(f"Generated at: {datetime.now()}", styles['Normal']))
            
            doc.build(elements)
            buffer.seek(0)
            
            return Response(
                buffer.getvalue(),
                mimetype='application/pdf',
                headers={
                    'Content-Disposition': 'attachment; filename="test.pdf"',
                    'Content-Length': str(len(buffer.getvalue()))
                }
            )
        except Exception as e:
            return f"Error: {str(e)}", 500

    @app.route('/admin')
    @login_required
    def admin_dashboard():
        if not is_admin(current_user):
            flash('You do not have permission to access this page.', 'danger')
            return redirect(url_for('index'))
        search = request.args.get('search', '').strip()
        teacher_filter = request.args.get('teacher_filter', 'all').strip().lower()  # all | internal | external
        if teacher_filter not in ('all', 'internal', 'external'):
            teacher_filter = 'all'
        query = User.query
        if search:
            like = f"%{search}%"
            query = query.filter(
                (User.username.ilike(like)) |
                (User.full_name.ilike(like)) |
                (User.email.ilike(like))
            )
        users = query.order_by(User.id.asc()).all()
        student_users = [u for u in users if 'student' in parse_roles(u.role)]
        other_users = [u for u in users if 'student' not in parse_roles(u.role)]
        
        # Fetch teacher information for users with teacher role (prefer user.teacher_id link, else match by name)
        user_teachers = {}
        teacher_users = [u for u in other_users if 'teacher' in parse_roles(u.role)]
        if teacher_users:
            for user in teacher_users:
                teacher = None
                if getattr(user, 'teacher_id', None):
                    teacher = Teacher.query.get(user.teacher_id)
                if not teacher:
                    teacher = Teacher.query.filter_by(name=user.full_name).first()
                if teacher:
                    user_teachers[user.id] = teacher
        
        # Filter by External/Internal teacher category
        if teacher_filter == 'external':
            other_users = [u for u in other_users if user_teachers.get(u.id) and getattr(user_teachers[u.id], 'is_external', False)]
        elif teacher_filter == 'internal':
            other_users = [u for u in other_users if u.id not in user_teachers or not getattr(user_teachers.get(u.id), 'is_external', False)]
        
        response = make_response(render_template(
            'admin_dashboard.html',
            users=other_users,
            student_users=student_users,
            role_labels=ROLE_LABELS,
            role_choices=NON_ADMIN_ROLE_CHOICES,
            search_query=search,
            user_teachers=user_teachers,
            teacher_filter=teacher_filter
        ))
        # Add cache-control headers to prevent browser caching of user-specific content
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, private'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response

    @app.route('/admin/role-privileges')
    @login_required
    def admin_role_privileges():
        if not is_admin(current_user):
            flash('You do not have permission to access this page.', 'danger')
            return redirect(url_for('index'))
        return render_template('admin_role_privileges.html', role_labels=ROLE_LABELS)

    @app.route('/admin/active-semester')
    @login_required
    def admin_active_semester():
        """Active Semester Management Page"""
        if not is_admin(current_user):
            flash('You do not have permission to access this page.', 'danger')
            return redirect(url_for('index'))
        
        from utils.semester_utils import get_active_semester_info
        
        # Get current active semesters
        active_semesters = get_active_semester_info()
        
        # Get available sessions from CurriculumYearTerm
        available_sessions_data = db.session.query(
            CurriculumYearTerm.academic_session,
            CurriculumYearTerm.year,
            CurriculumYearTerm.term,
            CurriculumYearTerm.batch
        ).filter(
            CurriculumYearTerm.academic_session.isnot(None)
        ).distinct().order_by(
            CurriculumYearTerm.academic_session.desc(),
            CurriculumYearTerm.year.asc(),
            CurriculumYearTerm.term.asc()
        ).all()
        
        # Group by academic_session, year, term (base from CurriculumYearTerm)
        sessions_dict = {}
        for row in available_sessions_data:
            key = f"{row[0]}|{row[1]}|{row[2]}"
            if key not in sessions_dict:
                sessions_dict[key] = {
                    'academic_session': row[0],
                    'year': row[1],
                    'term': row[2],
                    'batches': []
                }
            if row[3] and row[3] not in sessions_dict[key]['batches']:
                sessions_dict[key]['batches'].append(row[3])

        # Merge batches from CourseSessionAssignment too (newly assigned batches may exist only here)
        assignment_rows = db.session.query(
            CourseSessionAssignment.academic_session,
            CourseSessionAssignment.year,
            CourseSessionAssignment.term,
            CourseSessionAssignment.batch
        ).filter(
            CourseSessionAssignment.academic_session.isnot(None),
            CourseSessionAssignment.year.isnot(None),
            CourseSessionAssignment.term.isnot(None)
        ).distinct().all()

        for row in assignment_rows:
            key = f"{row[0]}|{row[1]}|{row[2]}"
            if key not in sessions_dict:
                sessions_dict[key] = {
                    'academic_session': row[0],
                    'year': row[1],
                    'term': row[2],
                    'batches': []
                }
            if row[3] and row[3] not in sessions_dict[key]['batches']:
                sessions_dict[key]['batches'].append(row[3])

        # Merge batches from actual student registrations (retake/re-retake batches may exist only here)
        registration_batch_rows = db.session.query(
            StudentCourseRegistration.academic_session,
            StudentCourseRegistration.year,
            StudentCourseRegistration.term,
            Student.batch
        ).join(
            Student, Student.id == StudentCourseRegistration.student_id
        ).filter(
            StudentCourseRegistration.academic_session.isnot(None),
            StudentCourseRegistration.year.isnot(None),
            StudentCourseRegistration.term.isnot(None),
            Student.batch.isnot(None),
            Student.batch != ''
        ).distinct().all()

        for row in registration_batch_rows:
            key = f"{row[0]}|{row[1]}|{row[2]}"
            if key not in sessions_dict:
                sessions_dict[key] = {
                    'academic_session': row[0],
                    'year': row[1],
                    'term': row[2],
                    'batches': []
                }
            if row[3] and row[3] not in sessions_dict[key]['batches']:
                sessions_dict[key]['batches'].append(row[3])
        
        available_sessions = list(sessions_dict.values())
        for session_data in available_sessions:
            session_data['batches'] = sorted(session_data['batches'])
        
        # Get unique academic sessions for dropdown (to avoid duplicates)
        unique_academic_sessions = sorted(set(s['academic_session'] for s in available_sessions if s['academic_session']))
        
        # Get all active semester configs (including history)
        all_configs = ActiveSemesterConfig.query.order_by(
            ActiveSemesterConfig.activated_at.desc()
        ).limit(50).all()
        
        return render_template('admin/active_semester.html',
                             active_semesters=active_semesters,
                             available_sessions=available_sessions,
                             unique_academic_sessions=unique_academic_sessions,
                             history=all_configs)

    @app.route('/admin/active-semester/set', methods=['POST'])
    @login_required
    def admin_set_active_semester():
        """API endpoint to set active semester"""
        if not is_admin(current_user):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        batch = data.get('batch', '').strip() or None
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            from utils.semester_utils import set_active_semester
            
            activated_by = current_user.full_name or current_user.username
            new_config = set_active_semester(
                academic_session=academic_session,
                year=year,
                term=term,
                batch=batch,
                activated_by=activated_by,
                deactivate_others=False
            )
            
            return jsonify({
                'success': True,
                'message': f'Active semester set to {academic_session} - {year} - {term}' + (f' (Batch: {batch})' if batch else ''),
                'semester': new_config.to_dict()
            })
        except Exception as e:
            current_app.logger.error(f'Error setting active semester: {e}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

    @app.route('/admin/active-semester/list', methods=['GET'])
    @login_required
    def admin_list_active_semesters():
        """API endpoint to list active semesters"""
        if not is_admin(current_user):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        try:
            from utils.semester_utils import get_active_semester_info
            
            batch = request.args.get('batch', '').strip() or None
            active_semesters = get_active_semester_info(batch=batch)
            
            return jsonify({
                'success': True,
                'active_semesters': active_semesters
            })
        except Exception as e:
            current_app.logger.error(f'Error listing active semesters: {e}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

    @app.route('/admin/active-semester/deactivate', methods=['POST'])
    @login_required
    def admin_deactivate_semester():
        """API endpoint to deactivate a semester"""
        if not is_admin(current_user):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        batch = data.get('batch', '').strip() or None
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            from utils.semester_utils import deactivate_semester
            
            success = deactivate_semester(
                academic_session=academic_session,
                year=year,
                term=term,
                batch=batch
            )
            
            if success:
                batch_str = f' (Batch: {batch})' if batch else ''
                return jsonify({
                    'success': True,
                    'message': f'Semester {academic_session} - {year} - {term}{batch_str} deactivated successfully.'
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Semester not found or already inactive.'
                }), 404
        except Exception as e:
            current_app.logger.error(f'Error deactivating semester: {e}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

    @app.route('/admin/active-semester/preview-deletion', methods=['POST'])
    @login_required
    def admin_preview_deletion():
        """API endpoint to preview what will be deleted (counts only, no actual deletion)"""
        if not is_admin(current_user):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        batch = data.get('batch', '').strip() or None
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            # Check if semester is active
            active_check = ActiveSemesterConfig.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term,
                batch=batch,
                is_active=True
            ).first()
            
            if active_check:
                return jsonify({
                    'success': False,
                    'message': 'Cannot delete data for an active semester. Please deactivate it first.'
                }), 400
            
            # Import all necessary models
            from blueprints.class_management.models import (
                Session, ClassAttendance, ClassStudent, CourseReview,
                EvaluationInvite, EvaluationSubmission, StudentFeedbackLink,
                StudentFeedbackResponse, CourseOutline, ClassSplitInvite,
                ExamPaperEvaluation
            )
            from blueprints.result_management.models import (
                RSession, RMark, RCourseRegistration, RSubject, RStudent
            )
            from blueprints.course_management.models import (
                StudentCourseRegistration, CourseRegistrationInvite,
                CourseSessionAssignment, DutyAssignment
            )
            
            # Try to import BatchCustomEvent (may not exist)
            try:
                from blueprints.academic_calendar.models import BatchCustomEvent
            except ImportError:
                BatchCustomEvent = None
            
            # Count records that will be deleted
            counts = {}
            
            # Class Management: Get sessions matching criteria
            session_query = Session.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            )
            if batch:
                # If batch is specified, we need to check CourseSessionAssignment
                from blueprints.course_management.models import CourseSessionAssignment
                assignment_ids = [a.session_id for a in CourseSessionAssignment.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term,
                    batch=batch
                ).all() if a.session_id]
                if assignment_ids:
                    session_query = session_query.filter(Session.id.in_(assignment_ids))
                else:
                    session_query = session_query.filter(False)  # No matching sessions
            
            matching_sessions = session_query.all()
            session_ids = [s.id for s in matching_sessions]
            
            counts['sessions'] = len(session_ids)
            
            if session_ids:
                # Count child records
                counts['class_attendance'] = ClassAttendance.query.filter(
                    ClassAttendance.session_id.in_(session_ids)
                ).count()
                
                counts['class_students'] = ClassStudent.query.filter(
                    ClassStudent.session_id.in_(session_ids)
                ).count()
                
                counts['course_reviews'] = CourseReview.query.filter(
                    CourseReview.session_id.in_(session_ids)
                ).count()
                
                counts['evaluation_invites'] = EvaluationInvite.query.filter(
                    EvaluationInvite.session_id.in_(session_ids)
                ).count()
                
                counts['evaluation_submissions'] = EvaluationSubmission.query.filter(
                    EvaluationSubmission.session_id.in_(session_ids)
                ).count()
                
                feedback_link_ids = [link.id for link in StudentFeedbackLink.query.filter(
                    StudentFeedbackLink.session_id.in_(session_ids)
                ).all()]
                counts['student_feedback_links'] = len(feedback_link_ids)
                
                if feedback_link_ids:
                    counts['student_feedback_responses'] = StudentFeedbackResponse.query.filter(
                        StudentFeedbackResponse.feedback_link_id.in_(feedback_link_ids)
                    ).count()
                else:
                    counts['student_feedback_responses'] = 0
                
                counts['course_outlines'] = CourseOutline.query.filter(
                    CourseOutline.session_id.in_(session_ids)
                ).count()
                
                counts['class_split_invites'] = ClassSplitInvite.query.filter(
                    ClassSplitInvite.inviter_session_id.in_(session_ids)
                ).count()
                
                if BatchCustomEvent:
                    counts['batch_custom_events'] = BatchCustomEvent.query.filter(
                        BatchCustomEvent.session_id.in_(session_ids)
                    ).count()
                else:
                    counts['batch_custom_events'] = 0
            else:
                counts['class_attendance'] = 0
                counts['class_students'] = 0
                counts['course_reviews'] = 0
                counts['evaluation_invites'] = 0
                counts['evaluation_submissions'] = 0
                counts['student_feedback_links'] = 0
                counts['student_feedback_responses'] = 0
                counts['course_outlines'] = 0
                counts['class_split_invites'] = 0
                counts['batch_custom_events'] = 0
            
            # Exam Paper Evaluation
            exam_query = ExamPaperEvaluation.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            )
            if batch:
                exam_query = exam_query.filter_by(batch=batch)
            counts['exam_paper_evaluations'] = exam_query.count()
            
            # Result Management: RSession uses 'name' field for academic_session
            result_session_query = RSession.query.filter_by(
                name=academic_session,
                year=year,
                term=term
            )
            if batch:
                result_session_query = result_session_query.filter_by(batch=batch)
            
            matching_result_sessions = result_session_query.all()
            result_session_ids = [rs.id for rs in matching_result_sessions]
            counts['result_sessions'] = len(result_session_ids)
            
            if result_session_ids:
                counts['r_marks'] = RMark.query.filter(
                    RMark.student_id.in_(
                        db.session.query(RStudent.id).filter(
                            RStudent.session_id.in_(result_session_ids)
                        )
                    )
                ).count()
                
                counts['r_course_registrations'] = RCourseRegistration.query.filter(
                    RCourseRegistration.student_id.in_(
                        db.session.query(RStudent.id).filter(
                            RStudent.session_id.in_(result_session_ids)
                        )
                    )
                ).count()
                
                counts['r_subjects'] = RSubject.query.filter(
                    RSubject.session_id.in_(result_session_ids)
                ).count()
            else:
                counts['r_marks'] = 0
                counts['r_course_registrations'] = 0
                counts['r_subjects'] = 0
            
            # Course Management
            registration_query = StudentCourseRegistration.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            )
            if batch:
                # Batch filtering for registrations would need to check student batch
                # For now, count all for the session/year/term
                pass
            counts['student_course_registrations'] = registration_query.count()
            
            counts['course_registration_invites'] = CourseRegistrationInvite.query.filter(
                CourseRegistrationInvite.registration_id.in_(
                    db.session.query(StudentCourseRegistration.id).filter_by(
                        academic_session=academic_session,
                        year=year,
                        term=term
                    )
                )
            ).count()
            
            assignment_query = CourseSessionAssignment.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            )
            if batch:
                assignment_query = assignment_query.filter_by(batch=batch)
            counts['course_session_assignments'] = assignment_query.count()
            
            duty_query = DutyAssignment.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            )
            if batch:
                duty_query = duty_query.filter_by(batch=batch)
            counts['duty_assignments'] = duty_query.count()
            
            # Calculate total
            total = sum(counts.values())
            
            return jsonify({
                'success': True,
                'counts': counts,
                'total': total
            })
        except Exception as e:
            current_app.logger.error(f'Error previewing deletion: {e}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

    @app.route('/admin/active-semester/delete-old-data', methods=['POST'])
    @login_required
    def admin_delete_old_data():
        """API endpoint to delete all data from old/inactive semesters"""
        if not is_admin(current_user):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        batch = data.get('batch', '').strip() or None
        confirmation = data.get('confirmation', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        if confirmation != 'DELETE':
            return jsonify({'success': False, 'message': 'Confirmation text "DELETE" is required'}), 400
        
        try:
            # Check if semester is active
            active_check = ActiveSemesterConfig.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term,
                batch=batch,
                is_active=True
            ).first()
            
            if active_check:
                return jsonify({
                    'success': False,
                    'message': 'Cannot delete data for an active semester. Please deactivate it first.'
                }), 400
            
            # Import all necessary models
            from blueprints.class_management.models import (
                Session, ClassAttendance, ClassStudent, CourseReview,
                EvaluationInvite, EvaluationSubmission, StudentFeedbackLink,
                StudentFeedbackResponse, CourseOutline, ClassSplitInvite,
                ExamPaperEvaluation
            )
            from blueprints.result_management.models import (
                RSession, RMark, RCourseRegistration, RSubject, RStudent
            )
            from blueprints.course_management.models import (
                StudentCourseRegistration, CourseRegistrationInvite,
                CourseSessionAssignment, DutyAssignment
            )
            
            # Try to import BatchCustomEvent (may not exist)
            try:
                from blueprints.academic_calendar.models import BatchCustomEvent
            except ImportError:
                BatchCustomEvent = None
            
            deletion_counts = {}
            
            try:
                # ===== CLASS MANAGEMENT DATA =====
                
                # Get sessions matching criteria
                session_query = Session.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                )
                if batch:
                    # If batch is specified, check CourseSessionAssignment
                    assignment_ids = [a.session_id for a in CourseSessionAssignment.query.filter_by(
                        academic_session=academic_session,
                        year=year,
                        term=term,
                        batch=batch
                    ).all() if a.session_id]
                    if assignment_ids:
                        session_query = session_query.filter(Session.id.in_(assignment_ids))
                    else:
                        session_query = session_query.filter(False)  # No matching sessions
                
                matching_sessions = session_query.all()
                session_ids = [s.id for s in matching_sessions]
                
                if session_ids:
                    # Delete child records first (respecting foreign key constraints)
                    
                    # 1. Delete student feedback responses (via feedback links)
                    feedback_link_ids = [link.id for link in StudentFeedbackLink.query.filter(
                        StudentFeedbackLink.session_id.in_(session_ids)
                    ).all()]
                    if feedback_link_ids:
                        deleted = StudentFeedbackResponse.query.filter(
                            StudentFeedbackResponse.feedback_link_id.in_(feedback_link_ids)
                        ).delete(synchronize_session=False)
                        deletion_counts['student_feedback_responses'] = deleted
                    
                    # 2. Delete student feedback links
                    deleted = StudentFeedbackLink.query.filter(
                        StudentFeedbackLink.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['student_feedback_links'] = deleted
                    
                    # 3. Delete batch custom events (if exists)
                    if BatchCustomEvent:
                        deleted = BatchCustomEvent.query.filter(
                            BatchCustomEvent.session_id.in_(session_ids)
                        ).delete(synchronize_session=False)
                        deletion_counts['batch_custom_events'] = deleted
                    
                    # 4. Delete course outlines
                    deleted = CourseOutline.query.filter(
                        CourseOutline.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['course_outlines'] = deleted
                    
                    # 5. Delete evaluation submissions
                    deleted = EvaluationSubmission.query.filter(
                        EvaluationSubmission.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['evaluation_submissions'] = deleted
                    
                    # 6. Delete evaluation invites
                    deleted = EvaluationInvite.query.filter(
                        EvaluationInvite.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['evaluation_invites'] = deleted
                    
                    # 7. Delete course reviews
                    deleted = CourseReview.query.filter(
                        CourseReview.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['course_reviews'] = deleted
                    
                    # 8. Delete split course invites
                    deleted = ClassSplitInvite.query.filter(
                        ClassSplitInvite.inviter_session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['class_split_invites'] = deleted
                    
                    # 9. Delete class attendance
                    deleted = ClassAttendance.query.filter(
                        ClassAttendance.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['class_attendance'] = deleted
                    
                    # 10. Delete class students
                    deleted = ClassStudent.query.filter(
                        ClassStudent.session_id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['class_students'] = deleted
                    
                    # 11. Delete sessions
                    deleted = Session.query.filter(
                        Session.id.in_(session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['sessions'] = deleted
                
                # ===== EXAM PAPER EVALUATION =====
                exam_query = ExamPaperEvaluation.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                )
                if batch:
                    exam_query = exam_query.filter_by(batch=batch)
                deleted = exam_query.delete(synchronize_session=False)
                deletion_counts['exam_paper_evaluations'] = deleted
                
                # ===== RESULT MANAGEMENT DATA =====
                # RSession uses 'name' field for academic_session
                result_session_query = RSession.query.filter_by(
                    name=academic_session,
                    year=year,
                    term=term
                )
                if batch:
                    result_session_query = result_session_query.filter_by(batch=batch)
                
                matching_result_sessions = result_session_query.all()
                result_session_ids = [rs.id for rs in matching_result_sessions]
                
                if result_session_ids:
                    # Get RStudent IDs for these sessions
                    r_student_ids = [s.id for s in RStudent.query.filter(
                        RStudent.session_id.in_(result_session_ids)
                    ).all()]
                    
                    if r_student_ids:
                        # Delete RMark (via student_id)
                        deleted = RMark.query.filter(
                            RMark.student_id.in_(r_student_ids)
                        ).delete(synchronize_session=False)
                        deletion_counts['r_marks'] = deleted
                        
                        # Delete RCourseRegistration (via student_id)
                        deleted = RCourseRegistration.query.filter(
                            RCourseRegistration.student_id.in_(r_student_ids)
                        ).delete(synchronize_session=False)
                        deletion_counts['r_course_registrations'] = deleted
                    
                    # Delete RSubject
                    deleted = RSubject.query.filter(
                        RSubject.session_id.in_(result_session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['r_subjects'] = deleted
                    
                    # Delete RStudent (cascade will handle marks/registrations, but we already deleted them)
                    deleted = RStudent.query.filter(
                        RStudent.session_id.in_(result_session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['r_students'] = deleted
                    
                    # Delete RSession
                    deleted = RSession.query.filter(
                        RSession.id.in_(result_session_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['result_sessions'] = deleted
                
                # ===== COURSE MANAGEMENT DATA =====
                
                # Delete student course registrations
                registration_query = StudentCourseRegistration.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                )
                # Note: Batch filtering for registrations is complex (would need student batch check)
                # For now, delete all for the session/year/term
                registration_ids = [r.id for r in registration_query.all()]
                
                if registration_ids:
                    # Delete course registration invites
                    deleted = CourseRegistrationInvite.query.filter(
                        CourseRegistrationInvite.registration_id.in_(registration_ids)
                    ).delete(synchronize_session=False)
                    deletion_counts['course_registration_invites'] = deleted
                
                deleted = registration_query.delete(synchronize_session=False)
                deletion_counts['student_course_registrations'] = deleted
                
                # Delete course session assignments
                assignment_query = CourseSessionAssignment.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                )
                if batch:
                    assignment_query = assignment_query.filter_by(batch=batch)
                deleted = assignment_query.delete(synchronize_session=False)
                deletion_counts['course_session_assignments'] = deleted
                
                # Delete duty assignments
                duty_query = DutyAssignment.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                )
                if batch:
                    duty_query = duty_query.filter_by(batch=batch)
                deleted = duty_query.delete(synchronize_session=False)
                deletion_counts['duty_assignments'] = deleted
                
                # Commit transaction
                db.session.commit()
                
                # Log deletion
                total_deleted = sum(deletion_counts.values())
                current_app.logger.info(
                    f'Admin {current_user.username} deleted old semester data: '
                    f'{academic_session} - {year} - {term} (Batch: {batch or "All"}). '
                    f'Total records deleted: {total_deleted}. Details: {deletion_counts}'
                )
                
                return jsonify({
                    'success': True,
                    'message': f'Successfully deleted {total_deleted} records for {academic_session} - {year} - {term}' + (f' (Batch: {batch})' if batch else ''),
                    'deletion_counts': deletion_counts,
                    'total': total_deleted
                })
                
            except Exception as e:
                db.session.rollback()
                raise e
                
        except Exception as e:
            current_app.logger.error(f'Error deleting old semester data: {e}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

    @app.route('/student/dashboard')
    @login_required
    def student_dashboard():
        roles = parse_roles(current_user.role)
        if 'student' not in roles and 'teaching_assistant' not in roles:
            flash('Student dashboard is available only for student accounts.', 'danger')
            return redirect(url_for('index'))
        response = make_response(render_template('student/dashboard.html'))
        # Add cache-control headers to prevent browser caching of user-specific content
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, private'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response

    @app.route('/head/dashboard')
    @login_required
    def head_dashboard():
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            flash('Head dashboard is available only for Head or Dean accounts.', 'danger')
            return redirect(url_for('index'))
        head_cards = [
            {'title': 'Assign Duties', 'desc': 'Assign Course Co-Ordinator, Exam Committee Chief, Teaching Assistant, and other duties',
             'icon': 'fas fa-user-tie', 'color': '#ffc107', 'bg': 'rgba(255,193,7,0.18)', 'url': url_for('assign_duties')},
            {'title': 'Result Management', 'desc': 'Handle student results, subjects, and marks',
             'icon': 'fas fa-poll', 'color': '#198754', 'bg': 'rgba(25,135,84,0.15)', 'url': url_for('result_management.index')},
            {'title': 'Students Management', 'desc': 'Manage student information and records',
             'icon': 'fas fa-users', 'color': '#dc3545', 'bg': 'rgba(220,53,69,0.15)', 'url': url_for('student_management.index')},
            {'title': 'Curriculum Management', 'desc': 'Manage curricula and courses',
             'icon': 'fas fa-book', 'color': '#6f42c1', 'bg': 'rgba(111,66,193,0.15)', 'url': url_for('course_management.index')},
            {'title': 'Curriculator', 'desc': 'Syllabus (Part A–D): collaboratively develop and download DOCX/PDF',
             'icon': 'fas fa-book-open', 'color': '#0d6efd', 'bg': 'rgba(13,110,253,0.15)', 'url': url_for('curriculator.index')},
            {'title': 'Course Registration', 'desc': 'Review and finalize student course registrations',
             'icon': 'fas fa-clipboard-check', 'color': '#0dcaf0', 'bg': 'rgba(13,202,240,0.15)', 'url': url_for('course_management.coordinator_registrations')},
            {'title': 'Exam Committee Archive', 'desc': 'View archived examination committees and their members',
             'icon': 'fas fa-archive', 'color': '#fd7e14', 'bg': 'rgba(253,126,20,0.15)', 'url': url_for('head_exam_committee_archive')},
            {'title': 'Academic Calendar', 'desc': 'View academic calendar with holidays, events, and important dates',
             'icon': 'fas fa-calendar', 'color': '#6366f1', 'bg': 'rgba(99,102,241,0.15)', 'url': url_for('academic_calendar.index')},
        ]
        # Self Assessment card: লিংক সবসময় /self-assessment (blueprints/self_assessment)
        if 'self_assessment.index' in current_app.view_functions:
            self_assessment_url = url_for('self_assessment.index')
        else:
            self_assessment_url = '/self-assessment'
        head_cards.append({
            'title': 'Self Assessment',
            'desc': 'PSAC Committee – Alumni, Employer, Faculty, Non-Academic Staff, and Student surveys',
            'icon': 'fas fa-clipboard-check', 'color': '#059669', 'bg': 'rgba(5,150,105,0.15)',
            'url': self_assessment_url
        })
        response = make_response(render_template('head/dashboard.html', head_cards=head_cards))
        # Add cache-control headers to prevent browser caching of user-specific content
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, private'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response

    @app.route('/head/exam-committee-archive')
    @login_required
    def head_exam_committee_archive():
        """View archived examination committees"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            flash('This page is available only for Head or Dean accounts.', 'danger')
            return redirect(url_for('index'))
        
        # Get all archived committee assignments grouped by session/year/term
        archived_assignments = DutyAssignment.query.filter(
            DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member']),
            DutyAssignment.status == 'archived'
        ).order_by(
            DutyAssignment.academic_session.desc(),
            DutyAssignment.year.asc(),
            DutyAssignment.term.asc()
        ).all()
        
        # Group by session/year/term
        archived_committees = {}
        for assignment in archived_assignments:
            key = f"{assignment.academic_session}|{assignment.year}|{assignment.term}"
            if key not in archived_committees:
                archived_committees[key] = {
                    'academic_session': assignment.academic_session,
                    'year': assignment.year,
                    'term': assignment.term,
                    'chief': None,
                    'internal_members': [],
                    'external_members': []
                }
            
            if assignment.duty_type == 'exam_committee_chief' and assignment.assigned_teacher:
                chief_designation = ''
                chief_institute = ''
                try:
                    chief_data = json.loads(assignment.remarks) if assignment.remarks else {}
                    if chief_data.get('type') == 'chief':
                        chief_designation = chief_data.get('designation', '')
                        chief_institute = chief_data.get('institute', '')
                except:
                    pass
                
                archived_committees[key]['chief'] = {
                    'name': assignment.assigned_teacher.name,
                    'designation': chief_designation or assignment.assigned_teacher.designation or 'Head, Law Discipline, KU',
                    'institute': chief_institute or assignment.assigned_teacher.institute or 'Law Discipline, KU'
                }
            elif assignment.duty_type == 'exam_committee_member':
                if assignment.assigned_teacher_id:
                    # Internal member
                    member_teacher = assignment.assigned_teacher
                    if member_teacher:
                        member_designation = ''
                        member_institute = ''
                        try:
                            member_data = json.loads(assignment.remarks) if assignment.remarks else {}
                            if member_data.get('type') == 'internal':
                                member_designation = member_data.get('designation', '')
                                member_institute = member_data.get('institute', '')
                        except:
                            pass
                        
                        archived_committees[key]['internal_members'].append({
                            'name': member_teacher.name,
                            'designation': member_designation or member_teacher.designation or 'Assistant Professor',
                            'institute': member_institute or member_teacher.institute or 'Law Discipline, KU'
                        })
                else:
                    # External member
                    try:
                        external_info = json.loads(assignment.remarks) if assignment.remarks else {}
                        if external_info.get('type') == 'external':
                            archived_committees[key]['external_members'].append({
                                'name': external_info.get('name', ''),
                                'designation': external_info.get('designation', ''),
                                'institute': external_info.get('institute', '')
                            })
                    except:
                        pass
        
        # Convert to list for template
        committees_list = list(archived_committees.values())
        
        return render_template('head/exam_committee_archive.html', committees=committees_list)

    @app.route('/head/restore-exam-committee', methods=['POST'])
    @login_required
    def restore_exam_committee():
        """Restore an archived examination committee back to active status"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'This action is available only for Head or Dean accounts.'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            # Find all archived committee assignments for this session/year/term
            archived_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member']),
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'archived'
            ).all()
            
            if not archived_assignments:
                # Check if there are any active assignments for this session/year/term
                active_assignments = DutyAssignment.query.filter(
                    DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member']),
                    DutyAssignment.academic_session == academic_session,
                    DutyAssignment.year == year,
                    DutyAssignment.term == term,
                    DutyAssignment.status == 'active'
                ).all()
                
                if active_assignments:
                    return jsonify({
                        'success': False,
                        'message': f'Cannot restore: An active committee already exists for {academic_session} - {year} - {term}. Please archive or deactivate the active committee first.'
                    }), 400
                else:
                    return jsonify({
                        'success': False,
                        'message': f'No archived committee found for {academic_session} - {year} - {term}'
                    }), 404
            
            # Check if there are any active assignments for this session/year/term
            # If yes, archive them first before restoring
            active_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member']),
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'active'
            ).all()
            
            if active_assignments:
                # Archive active assignments first (don't restore them, just archive)
                for assignment in active_assignments:
                    assignment.status = 'archived'
                db.session.commit()
                # Don't add active assignments to archived_assignments - we only want to restore the originally archived ones
            
            # Restore all archived assignments to active status
            restored_count = 0
            for assignment in archived_assignments:
                assignment.status = 'active'
                restored_count += 1
            
            # Restore tabulators for this session/year/term
            archived_tabulators = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'tabulator',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'archived'
            ).all()
            
            for tabulator in archived_tabulators:
                tabulator.status = 'active'
                restored_count += 1
            
            # Restore scrutinizers for this session/year/term
            archived_scrutinizers = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'scrutinizer',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'archived'
            ).all()
            
            for scrutinizer in archived_scrutinizers:
                scrutinizer.status = 'active'
                restored_count += 1
            
            # Restore remuneration forms for this session/year/term
            from blueprints.remuneration_management.models import RemunerationForm
            import json
            
            archived_remuneration_forms = RemunerationForm.query.filter_by(
                status='archived'
            ).all()
            
            restored_forms_count = 0
            for form in archived_remuneration_forms:
                # Check if form data matches the session/year/term
                try:
                    if form.form_data:
                        form_data = json.loads(form.form_data)
                        form_session = form_data.get('academic_session') or form.academic_year
                        form_year = form_data.get('year') or form.year
                        form_term = form_data.get('term') or form.term
                        
                        if (form_session == academic_session and 
                            form_year == year and 
                            form_term == term):
                            form.status = 'draft'
                            form.archived_at = None
                            restored_forms_count += 1
                except:
                    # If form_data parsing fails, check direct fields
                    if (form.academic_year == academic_session and 
                        form.year == year and 
                        form.term == term):
                        form.status = 'draft'
                        form.archived_at = None
                        restored_forms_count += 1
            
            restored_count += restored_forms_count
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Committee, tabulators, scrutinizers, and remuneration forms restored successfully for {academic_session} - {year} - {term}. {restored_count} item(s) restored.'
            })
        
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error restoring committee: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error restoring committee: {str(e)}'}), 500

    @app.route('/head/session-archive')
    @login_required
    def head_session_archive():
        """Session Archive - Archive and restore complete academic sessions"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            flash('This page is available only for Head or Dean accounts.', 'danger')
            return redirect(url_for('index'))
        
        # Get all archived sessions
        archived_sessions = SessionArchive.query.order_by(
            SessionArchive.archived_at.desc()
        ).all()
        
        # Get available sessions to archive (from CurriculumYearTerm)
        from blueprints.course_management.models import CurriculumYearTerm
        available_sessions = db.session.query(
            CurriculumYearTerm.academic_session,
            CurriculumYearTerm.year,
            CurriculumYearTerm.term,
            CurriculumYearTerm.batch
        ).distinct().filter(
            CurriculumYearTerm.academic_session.isnot(None)
        ).order_by(
            CurriculumYearTerm.academic_session.desc(),
            CurriculumYearTerm.year.asc(),
            CurriculumYearTerm.term.asc()
        ).all()
        
        # Group available sessions
        sessions_to_archive = {}
        for session_row in available_sessions:
            key = f"{session_row[0]}|{session_row[1]}|{session_row[2]}"
            if key not in sessions_to_archive:
                sessions_to_archive[key] = {
                    'academic_session': session_row[0],
                    'year': session_row[1],
                    'term': session_row[2],
                    'batches': []
                }
            if session_row[3] and session_row[3] not in sessions_to_archive[key]['batches']:
                sessions_to_archive[key]['batches'].append(session_row[3])
        
        sessions_list = list(sessions_to_archive.values())
        
        return render_template('head/session_archive.html',
                             archived_sessions=archived_sessions,
                             available_sessions=sessions_list)

    @app.route('/head/session-archive/api/archive', methods=['POST'])
    @login_required
    def archive_session_api():
        """Archive a complete academic session with all related data"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        description = data.get('description', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            # Check if already archived
            existing = SessionArchive.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term,
                is_active=True
            ).first()
            
            if existing:
                return jsonify({
                    'success': False,
                    'message': f'Session {academic_session} - {year} - {term} is already archived.'
                }), 400
            
            # Collect all related data
            archive_data = {
                'academic_session': academic_session,
                'year': year,
                'term': term,
                'archived_at': datetime.utcnow().isoformat(),
                'class_management': {},
                'result_management': {},
                'course_registrations': [],
                'exam_evaluations': {},
                'academic_calendar': {},
                'duty_assignments': [],
                'course_session_assignments': []
            }
            
            # Archive Class Management data
            try:
                from blueprints.class_management.models import Session as ClassSession, ClassStudent, ClassAttendance, CourseOutline
                from blueprints.course_management.models import CourseSessionAssignment
                
                # First, find sessions via CourseSessionAssignment (most reliable method)
                assignments = CourseSessionAssignment.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).all()
                
                session_ids_from_assignments = [a.session_id for a in assignments if a.session_id]
                
                # Use session_ids from assignments if available (most reliable)
                if session_ids_from_assignments:
                    class_sessions = ClassSession.query.filter(
                        ClassSession.id.in_(session_ids_from_assignments),
                        ClassSession.archived == False
                    ).all()
                    current_app.logger.info(f'Found {len(class_sessions)} sessions via CourseSessionAssignment (session_ids: {session_ids_from_assignments})')
                    
                    # Also try direct match for any sessions that might not be in assignments
                    direct_sessions = ClassSession.query.filter(
                        ClassSession.academic_session == academic_session,
                        ClassSession.year == year,
                        ClassSession.term == term,
                        ClassSession.archived == False
                    ).all()
                    
                    # Add any direct matches that weren't in assignments
                    assignment_ids = set(session_ids_from_assignments)
                    additional_sessions = [s for s in direct_sessions if s.id not in assignment_ids]
                    if additional_sessions:
                        class_sessions.extend(additional_sessions)
                        current_app.logger.info(f'Added {len(additional_sessions)} additional sessions from direct match')
                else:
                    # Fallback: use direct match if no assignments found
                    class_sessions = ClassSession.query.filter(
                        ClassSession.academic_session == academic_session,
                        ClassSession.year == year,
                        ClassSession.term == term,
                        ClassSession.archived == False
                    ).all()
                    current_app.logger.info(f'No CourseSessionAssignments found, using direct match: found {len(class_sessions)} sessions')
                
                archive_data['class_management'] = {
                    'sessions': []
                }
                
                current_app.logger.info(f'Found {len(class_sessions)} class sessions to archive for {academic_session} - {year} - {term}')
                
                for s in class_sessions:
                    try:
                        outline = CourseOutline.query.filter_by(session_id=s.id).first()
                        # Collect detailed student data with assessment marks
                        students_data = []
                        for cs in s.students:
                            student_dict = {
                                'student_id': cs.student_id,
                                'name': cs.name,
                                'added_at': cs.added_at.isoformat() if cs.added_at else None,
                                # Assessment marks
                                'assessment1': cs.assessment1,
                                'assessment2': cs.assessment2,
                                'assessment3': cs.assessment3,
                                'assessment4': cs.assessment4,
                                'assessment_total': cs.assessment_total,
                                'assessment_avg': cs.assessment_avg,
                                'assessment_total_40': cs.assessment_total_40,
                                'sessional_report': cs.sessional_report,
                                'sessional_viva': cs.sessional_viva,
                                'assessment_absent': cs.assessment_absent  # JSON string
                            }
                            students_data.append(student_dict)
                        
                        # Collect detailed attendance data
                        attendances_data = []
                        for ca in s.attendances:
                            attendance_dict = {
                                'student_id': ca.student_id,
                                'date': ca.date.isoformat() if ca.date else None,
                                'status': ca.status,
                                'remark': ca.remark if hasattr(ca, 'remark') else None
                            }
                            attendances_data.append(attendance_dict)
                        
                        session_data = {
                            'id': s.id,
                            'course_code': s.course_code,
                            'course_name': s.course_name,
                            'teacher_id': s.teacher_id,
                            'course_type': s.course_type,
                            'category': s.category,
                            'course_scope': s.course_scope,
                            'created_at': s.created_at.isoformat() if s.created_at else None,
                            'students': students_data,
                            'attendances': attendances_data,
                            'course_outline': None,
                            'assessment_revealed': s.assessment_revealed  # JSON string for reveal status
                        }
                        
                        if outline:
                            session_data['course_outline'] = {
                                'teacher_id': outline.teacher_id,
                                'course_objectives': outline.course_objectives,
                                'course_summary': outline.course_summary,
                                'lesson_plan': outline.lesson_plan,
                                'created_at': outline.created_at.isoformat() if outline.created_at else None
                            }
                        
                        archive_data['class_management']['sessions'].append(session_data)
                        
                        # Mark session as archived (remove from live view)
                        s.archived = True
                        current_app.logger.debug(f'Marked session {s.id} ({s.course_name}) as archived')
                    except Exception as session_error:
                        current_app.logger.error(f'Error archiving individual session {s.id}: {session_error}', exc_info=True)
                        # Continue with other sessions even if one fails
                        continue
                
                # Ensure all archived flags are persisted
                db.session.flush()
                current_app.logger.info(f'Marked {len(class_sessions)} class sessions as archived (flushed to DB)')
                
            except Exception as e:
                current_app.logger.error(f'Error archiving class management data: {e}', exc_info=True)
                # Don't re-raise - allow other archive operations to continue
            
            # Archive Result Management data
            try:
                from blueprints.result_management.models import RSession, RStudent, RSubject, RMark
                
                result_sessions = RSession.query.filter_by(
                    name=academic_session,
                    year=year,
                    term=term
                ).filter_by(is_archived=False).all()
                
                archive_data['result_management'] = {
                    'sessions': []
                }
                
                for rs in result_sessions:
                    # Get marks for this session
                    marks_list = []
                    for student in rs.students:
                        for subject in rs.subjects:
                            mark = RMark.query.filter_by(
                                student_id=student.id,
                                subject_id=subject.id
                            ).first()
                            if mark:
                                marks_list.append({
                                    'student_id': student.id,
                                    'subject_id': subject.id,
                                    'marks': mark.marks,
                                    'grade': mark.grade
                                })
                    
                    session_data = {
                        'id': rs.id,
                        'batch': rs.batch,
                        'curriculum_id': rs.curriculum_id,
                        'students': [{'id': s.id, 'student_id': s.student_id, 'name': s.name} for s in rs.students],
                        'subjects': [{'id': sub.id, 'code': sub.code, 'name': sub.name, 'credit': sub.credit} for sub in rs.subjects],
                        'marks': marks_list
                    }
                    archive_data['result_management']['sessions'].append(session_data)
                    
                    # Mark result session as archived (remove from live view)
                    rs.is_archived = True
                
                # Ensure all archived flags are persisted
                if result_sessions:
                    db.session.flush()
                    current_app.logger.info(f'Marked {len(result_sessions)} result sessions as archived (flushed to DB)')
                    
            except Exception as e:
                current_app.logger.error(f'Error archiving result management data: {e}', exc_info=True)
            
            # Archive Course Registrations
            try:
                registrations = StudentCourseRegistration.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).filter(StudentCourseRegistration.status != 'archived').all()
                
                archive_data['course_registrations'] = []
                for reg in registrations:
                    archive_data['course_registrations'].append({
                        'id': reg.id,
                        'student_id': reg.student_id,
                        'course_id': reg.course_id,
                        'course_code': reg.course_code,
                        'course_name': reg.course_name,
                        'credit': reg.credit,
                        'course_type': reg.course_type,
                        'nature': reg.nature,
                        'remark': reg.remark,
                        'carry_on': reg.carry_on,
                        'status': reg.status,
                        'registered_by': reg.registered_by,
                        'created_at': reg.created_at.isoformat() if reg.created_at else None
                    })
                    # Mark registration as archived (remove from live view)
                    reg.status = 'archived'
                
                # Ensure all archived flags are persisted
                if registrations:
                    db.session.flush()
                    current_app.logger.info(f'Marked {len(registrations)} course registrations as archived (flushed to DB)')
                    
            except Exception as e:
                current_app.logger.error(f'Error archiving course registrations: {e}', exc_info=True)
            
            # Archive Duty Assignments
            try:
                duty_assignments = DutyAssignment.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).filter(DutyAssignment.status != 'archived').all()
                
                archive_data['duty_assignments'] = []
                for da in duty_assignments:
                    archive_data['duty_assignments'].append({
                        'id': da.id,
                        'course_id': da.course_id,
                        'course_code': da.course_code,
                        'course_name': da.course_name,
                        'batch': da.batch,
                        'duty_type': da.duty_type,
                        'assigned_teacher_id': da.assigned_teacher_id,
                        'student_id': da.student_id,
                        'remarks': da.remarks,
                        'status': da.status,
                        'created_at': da.created_at.isoformat() if da.created_at else None
                    })
                    # Mark duty assignment as archived (remove from live view)
                    da.status = 'archived'
                
                # Ensure all archived flags are persisted
                if duty_assignments:
                    db.session.flush()
                    current_app.logger.info(f'Marked {len(duty_assignments)} duty assignments as archived (flushed to DB)')
                    
            except Exception as e:
                current_app.logger.error(f'Error archiving duty assignments: {e}', exc_info=True)
            
            # Archive Course Session Assignments (store data but keep assignments - they're needed for reference)
            try:
                from blueprints.course_management.models import CourseSessionAssignment
                course_assignments = CourseSessionAssignment.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).all()
                
                archive_data['course_session_assignments'] = [{
                    'id': ca.id,
                    'course_id': ca.course_id,
                    'curriculum_id': ca.curriculum_id,
                    'teacher_id': ca.teacher_id,
                    'section': ca.section,
                    'batch': ca.batch,
                    'session_created': ca.session_created,
                    'session_id': ca.session_id,
                    'created_at': ca.created_at.isoformat() if ca.created_at else None
                } for ca in course_assignments]
                # Note: CourseSessionAssignments are kept as reference, not marked as archived
            except Exception as e:
                current_app.logger.warning(f'Error archiving course session assignments: {e}')
            
            # Archive Exam Paper Evaluations
            try:
                from blueprints.class_management.models import ExamPaperEvaluation
                exam_entries = ExamPaperEvaluation.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).filter_by(archived=False).all()
                
                archive_data['exam_evaluations'] = [{
                    'id': ee.id,
                    'course_code': ee.course_code,
                    'course_name': ee.course_name,
                    'batch': ee.batch,
                    'section': ee.section,
                    'program_level': ee.program_level,
                    'marks_data': ee.marks_data,
                    'owner_teacher_id': ee.owner_teacher_id,
                    'assigned_scrutinizer_id': ee.assigned_scrutinizer_id,
                    'submitted_to_committee': ee.submitted_to_committee,
                    'created_at': ee.created_at.isoformat() if ee.created_at else None
                } for ee in exam_entries]
                
                # Mark exam evaluations as archived (remove from live view)
                for ee in exam_entries:
                    ee.archived = True
                
                # Ensure all archived flags are persisted
                if exam_entries:
                    db.session.flush()
                    current_app.logger.info(f'Marked {len(exam_entries)} exam paper evaluations as archived (flushed to DB)')
                    
            except Exception as e:
                current_app.logger.error(f'Error archiving exam paper evaluations: {e}', exc_info=True)
            
            # Archive Academic Calendar Events
            try:
                from blueprints.academic_calendar.models import AcademicCalendarEvent
                calendar_events = AcademicCalendarEvent.query.filter_by(
                    academic_session=academic_session
                ).all()
                
                archive_data['academic_calendar'] = {
                    'events': [{
                        'id': ev.id,
                        'title': ev.title,
                        'description': ev.description,
                        'event_date': ev.event_date.isoformat() if ev.event_date else None,
                        'event_type': ev.event_type,
                        'created_at': ev.created_at.isoformat() if ev.created_at else None
                    } for ev in calendar_events]
                }
                # Note: Calendar events are kept for historical reference, typically not hidden
            except Exception as e:
                current_app.logger.error(f'Error archiving academic calendar events: {e}', exc_info=True)
            
            # Archive Routine/Schedules
            try:
                from blueprints.routine_management.models import Routine
                routines = Routine.query.filter_by(
                    year=year,
                    term=term
                ).all()
                
                archive_data['routines'] = [{
                    'id': r.id,
                    'day': r.day,
                    'time_slot': r.time_slot,
                    'room_number': r.room_number,
                    'course_code': r.course_code,
                    'teacher_short_name': r.teacher_short_name,
                    'part': r.part,
                    'is_shared': r.is_shared,
                    'shared_with': r.shared_with,
                    'teacher_id': r.teacher_id,
                    'year': r.year,
                    'term': r.term
                } for r in routines]
                
                current_app.logger.info(f'Archived {len(routines)} routine entries for {year} - {term}')
            except Exception as e:
                current_app.logger.error(f'Error archiving routines: {e}', exc_info=True)
            
            # Archive Remuneration Forms
            try:
                from blueprints.remuneration_management.models import RemunerationForm
                # Find remuneration forms matching the session/year/term
                # Forms have academic_year (which is the session), year, and term fields
                remuneration_forms = RemunerationForm.query.filter(
                    RemunerationForm.academic_year == academic_session,
                    RemunerationForm.year == year,
                    RemunerationForm.term == term,
                    RemunerationForm.status != 'archived'
                ).all()
                
                archive_data['remuneration_forms'] = []
                for form in remuneration_forms:
                    form_dict = form.to_dict() if hasattr(form, 'to_dict') else {
                        'id': form.id,
                        'user_id': form.user_id,
                        'status': form.status,
                        'title': form.title,
                        'applicant_name': form.applicant_name,
                        'designation': form.designation,
                        'year': form.year,
                        'term': form.term,
                        'academic_year': form.academic_year,
                        'total_amount': form.total_amount,
                        'created_at': form.created_at.isoformat() if form.created_at else None,
                        'form_data': form.form_data  # JSON string
                    }
                    archive_data['remuneration_forms'].append(form_dict)
                    
                    # Mark form as archived
                    form.status = 'archived'
                    form.archived_at = datetime.utcnow()
                
                if remuneration_forms:
                    db.session.flush()
                    current_app.logger.info(f'Marked {len(remuneration_forms)} remuneration forms as archived (flushed to DB)')
                else:
                    current_app.logger.info(f'No remuneration forms found for {academic_session} - {year} - {term}')
                    
            except Exception as e:
                current_app.logger.error(f'Error archiving remuneration forms: {e}', exc_info=True)
            
            # Archive Committee Members, Tabulators, and Scrutinizers (already archived via DutyAssignment, but store for reference)
            # Note: These are already stored in duty_assignments above, but we add explicit committee info
            try:
                # Committee members are stored as DutyAssignment with duty_type='committee_member'
                committee_assignments = [da for da in archive_data.get('duty_assignments', []) if da.get('duty_type') == 'committee_member']
                tabulator_assignments = [da for da in archive_data.get('duty_assignments', []) if da.get('duty_type') == 'tabulator']
                scrutinizer_assignments = [da for da in archive_data.get('duty_assignments', []) if da.get('duty_type') == 'scrutinizer']
                
                archive_data['committee_info'] = {
                    'committee_members': committee_assignments,
                    'tabulators': tabulator_assignments,
                    'scrutinizers': scrutinizer_assignments,
                    'total_committee': len(committee_assignments),
                    'total_tabulators': len(tabulator_assignments),
                    'total_scrutinizers': len(scrutinizer_assignments)
                }
                
                current_app.logger.info(f'Archived committee info: {len(committee_assignments)} members, {len(tabulator_assignments)} tabulators, {len(scrutinizer_assignments)} scrutinizers')
            except Exception as e:
                current_app.logger.error(f'Error archiving committee info: {e}', exc_info=True)
            
            # Create archive record
            archive = SessionArchive(
                academic_session=academic_session,
                year=year,
                term=term,
                archive_data=json.dumps(archive_data, default=str),
                archived_by=current_user.full_name or current_user.username,
                description=description,
                is_active=True
            )
            
            db.session.add(archive)
            
            # Log before commit for debugging
            archived_class_sessions_count = len(archive_data.get('class_management', {}).get('sessions', []))
            archived_result_sessions_count = len(archive_data.get('result_management', {}).get('sessions', []))
            archived_registrations_count = len(archive_data.get('course_registrations', []))
            archived_exam_entries_count = len(archive_data.get('exam_evaluations', []))
            archived_duty_assignments_count = len(archive_data.get('duty_assignments', []))
            
            current_app.logger.info(f'Archiving session {academic_session} - {year} - {term}: '
                                  f'{archived_class_sessions_count} class sessions, '
                                  f'{archived_result_sessions_count} result sessions, '
                                  f'{archived_registrations_count} registrations, '
                                  f'{archived_exam_entries_count} exam entries, '
                                  f'{archived_duty_assignments_count} duty assignments')
            
            db.session.commit()
            
            current_app.logger.info(f'Session {academic_session} - {year} - {term} archived successfully. All data marked as archived in database.')
            
            return jsonify({
                'success': True,
                'message': f'Session {academic_session} - {year} - {term} archived successfully. All related data has been archived and removed from live view.'
            })
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error archiving session: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error archiving session: {str(e)}'}), 500

    @app.route('/head/session-archive/api/delete/<int:archive_id>', methods=['DELETE'])
    @login_required
    def delete_archive_api(archive_id):
        """Delete an archived session"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        try:
            archive = SessionArchive.query.get(archive_id)
            if not archive:
                return jsonify({'success': False, 'message': 'Archive not found'}), 404
            
            academic_session = archive.academic_session
            year = archive.year
            term = archive.term
            
            # Delete the archive record
            db.session.delete(archive)
            db.session.commit()
            
            current_app.logger.info(f'Archive {archive_id} ({academic_session} - {year} - {term}) deleted by {current_user.full_name or current_user.username}')
            
            return jsonify({
                'success': True,
                'message': f'Archive for {academic_session} - {year} - {term} has been permanently deleted.'
            })
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error deleting archive: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error deleting archive: {str(e)}'}), 500

    @app.route('/head/session-archive/api/details/<int:archive_id>', methods=['GET'])
    @login_required
    def archive_details_api(archive_id):
        """Get archive details"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        try:
            archive = SessionArchive.query.get(archive_id)
            if not archive:
                return jsonify({'success': False, 'message': 'Archive not found'}), 404
            
            # Parse archive data
            try:
                archive_data = json.loads(archive.archive_data)
            except json.JSONDecodeError as e:
                current_app.logger.error(f'Error parsing archive data: {e}')
                return jsonify({'success': False, 'message': f'Error parsing archive data: {str(e)}'}), 500
            
            # Get all teachers from archived sessions for dropdown
            teachers_map = {}
            teachers_list = []
            teacher_ids = set()
            
            # Extract teacher_ids from class management sessions
            if archive_data.get('class_management') and archive_data['class_management'].get('sessions'):
                for session in archive_data['class_management']['sessions']:
                    if session.get('teacher_id'):
                        teacher_ids.add(session['teacher_id'])
            
            # Also extract from exam evaluations
            if archive_data.get('exam_evaluations'):
                for ee in archive_data['exam_evaluations']:
                    if ee.get('owner_teacher_id'):
                        teacher_ids.add(ee['owner_teacher_id'])
            
            # Fetch teacher details (teachers might not exist if database was cleared)
            from blueprints.class_management.models import Teacher
            if teacher_ids:
                try:
                    teachers = Teacher.query.filter(Teacher.id.in_(teacher_ids)).all()
                    found_teacher_ids = {t.id for t in teachers}
                    
                    for teacher in teachers:
                        teachers_map[teacher.id] = {
                            'id': teacher.id,
                            'name': teacher.name,
                            'short_name': teacher.short_name,
                            'designation': teacher.designation or ''
                        }
                        teachers_list.append({
                            'id': teacher.id,
                            'name': teacher.name,
                            'short_name': teacher.short_name,
                            'designation': teacher.designation or ''
                        })
                    
                    # Add placeholder entries for teachers that don't exist in database
                    for teacher_id in teacher_ids:
                        if teacher_id not in found_teacher_ids:
                            teachers_list.append({
                                'id': teacher_id,
                                'name': f'Teacher ID {teacher_id}',
                                'short_name': f'T{teacher_id}',
                                'designation': ''
                            })
                except Exception as teacher_error:
                    current_app.logger.warning(f'Error fetching teachers: {teacher_error}')
                    # If teachers query fails, create placeholder entries from teacher_ids
                    for teacher_id in teacher_ids:
                        teachers_list.append({
                            'id': teacher_id,
                            'name': f'Teacher ID {teacher_id}',
                            'short_name': f'T{teacher_id}',
                            'designation': ''
                        })
            
            return jsonify({
                'success': True,
                'archive_data': archive_data,
                'teachers': teachers_list,
                'metadata': {
                    'archived_by': archive.archived_by,
                    'archived_at': archive.archived_at.isoformat() if archive.archived_at else None,
                    'restored_at': archive.restored_at.isoformat() if archive.restored_at else None,
                    'restored_by': archive.restored_by,
                    'description': archive.description
                }
            })
            
        except Exception as e:
            current_app.logger.error(f'Error getting archive details: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error loading archive details: {str(e)}'}), 500

    @app.route('/head/session-archive/api/restore', methods=['POST'])
    @login_required
    def restore_session_api():
        """Restore an archived session (safely archives current session first)"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        archive_id = data.get('archive_id')
        
        if not archive_id:
            return jsonify({'success': False, 'message': 'Archive ID is required'}), 400
        
        try:
            archive = SessionArchive.query.get(archive_id)
            if not archive:
                return jsonify({'success': False, 'message': 'Archive not found'}), 404
            
            if not archive.is_active:
                return jsonify({'success': False, 'message': 'This archive has already been restored'}), 400
            
            academic_session = archive.academic_session
            year = archive.year
            term = archive.term
            
            # Step 1: Archive current active session data to prevent data loss
            current_app.logger.info(f'Step 1: Archiving current active session before restoring: {academic_session} - {year} - {term}')
            
            # Check if there's any active data for this session/year/term
            from blueprints.class_management.models import Session as ClassSession
            active_class_sessions = ClassSession.query.filter(
                ClassSession.academic_session == academic_session,
                ClassSession.year == year,
                ClassSession.term == term,
                ClassSession.archived == False
            ).all()
            
            if active_class_sessions:
                # Archive current active session automatically
                current_app.logger.info(f'Found {len(active_class_sessions)} active class sessions. Archiving them first...')
                
                # Create a backup archive for current active data
                current_archive_data = {
                    'academic_session': academic_session,
                    'year': year,
                    'term': term,
                    'class_management': {'sessions': []},
                    'note': 'Auto-archived before restore operation'
                }
                
                # Archive current class sessions
                for s in active_class_sessions:
                    s.archived = True
                    current_app.logger.info(f'Marked active session {s.id} ({s.course_name}) as archived before restore')
                
                # Create backup archive record
                backup_archive = SessionArchive(
                    academic_session=academic_session,
                    year=year,
                    term=term,
                    archive_data=json.dumps(current_archive_data, default=str),
                    archived_by=current_user.full_name or current_user.username,
                    description=f'Auto-archived before restoring archive #{archive_id}',
                    is_active=True
                )
                db.session.add(backup_archive)
                db.session.flush()
                current_app.logger.info(f'Created backup archive {backup_archive.id} for current active data')
            
            # Step 2: Parse archive data
            archive_data = json.loads(archive.archive_data)
            current_app.logger.info(f'Step 2: Parsed archive data for restoration')
            
            # Step 3: Restore Class Management data
            restored_count = 0
            try:
                from blueprints.class_management.models import Session as ClassSession, ClassStudent, ClassAttendance, CourseOutline
                
                if archive_data.get('class_management') and archive_data['class_management'].get('sessions'):
                    for session_data in archive_data['class_management']['sessions']:
                        try:
                            # Check if session already exists (by ID)
                            existing_session = ClassSession.query.get(session_data.get('id'))
                            
                            if existing_session:
                                # Update existing session (unarchive it)
                                existing_session.archived = False
                                current_app.logger.info(f'Unarchived existing session {existing_session.id} ({existing_session.course_name})')
                                
                                # Restore students
                                if session_data.get('students'):
                                    for student_data in session_data['students']:
                                        student = ClassStudent.query.filter_by(
                                            session_id=existing_session.id,
                                            student_id=student_data.get('student_id')
                                        ).first()
                                        
                                        if not student:
                                            # Create new student record
                                            student = ClassStudent(
                                                student_id=student_data.get('student_id'),
                                                name=student_data.get('name', ''),
                                                session_id=existing_session.id,
                                                teacher_id=existing_session.teacher_id
                                            )
                                            db.session.add(student)
                                        
                                        # Restore assessment marks
                                        if student_data.get('assessment1') is not None:
                                            student.assessment1 = student_data.get('assessment1')
                                        if student_data.get('assessment2') is not None:
                                            student.assessment2 = student_data.get('assessment2')
                                        if student_data.get('assessment3') is not None:
                                            student.assessment3 = student_data.get('assessment3')
                                        if student_data.get('assessment4') is not None:
                                            student.assessment4 = student_data.get('assessment4')
                                        if student_data.get('assessment_total') is not None:
                                            student.assessment_total = student_data.get('assessment_total')
                                        if student_data.get('assessment_avg') is not None:
                                            student.assessment_avg = student_data.get('assessment_avg')
                                        if student_data.get('assessment_total_40') is not None:
                                            student.assessment_total_40 = student_data.get('assessment_total_40')
                                        if student_data.get('sessional_report') is not None:
                                            student.sessional_report = student_data.get('sessional_report')
                                        if student_data.get('sessional_viva') is not None:
                                            student.sessional_viva = student_data.get('sessional_viva')
                                        if student_data.get('assessment_absent'):
                                            student.assessment_absent = student_data.get('assessment_absent')
                                
                                # Restore attendances
                                if session_data.get('attendances'):
                                    for att_data in session_data['attendances']:
                                        # Check if attendance already exists
                                        if att_data.get('date'):
                                            try:
                                                try:
                                                    att_date = datetime.fromisoformat(att_data['date'].replace('Z', '+00:00'))
                                                except:
                                                    try:
                                                        att_date = datetime.strptime(att_data['date'], '%Y-%m-%d')
                                                    except:
                                                        att_date = None
                                                
                                                if att_date:
                                                    attendance = ClassAttendance.query.filter_by(
                                                        session_id=existing_session.id,
                                                        student_id=att_data.get('student_id'),
                                                        date=att_date.date() if hasattr(att_date, 'date') else att_date
                                                    ).first()
                                                    
                                                    if not attendance:
                                                        attendance = ClassAttendance(
                                                            session_id=existing_session.id,
                                                            student_id=att_data.get('student_id'),
                                                            date=att_date.date() if hasattr(att_date, 'date') else att_date,
                                                            status=att_data.get('status', 'present'),
                                                            teacher_id=existing_session.teacher_id
                                                        )
                                                        if att_data.get('remark'):
                                                            attendance.remark = att_data.get('remark')
                                                        db.session.add(attendance)
                                            except Exception as att_error:
                                                current_app.logger.warning(f'Error restoring attendance: {att_error}')
                                
                                # Restore course outline
                                if session_data.get('course_outline'):
                                    outline = CourseOutline.query.filter_by(session_id=existing_session.id).first()
                                    if not outline:
                                        outline = CourseOutline(
                                            session_id=existing_session.id,
                                            teacher_id=session_data['course_outline'].get('teacher_id', existing_session.teacher_id)
                                        )
                                        db.session.add(outline)
                                    
                                    outline.course_objectives = session_data['course_outline'].get('course_objectives')
                                    outline.course_summary = session_data['course_outline'].get('course_summary')
                                    outline.lesson_plan = session_data['course_outline'].get('lesson_plan')
                                
                                # Restore assessment_revealed status
                                if session_data.get('assessment_revealed'):
                                    existing_session.assessment_revealed = session_data.get('assessment_revealed')
                                
                                restored_count += 1
                        except Exception as session_error:
                            current_app.logger.error(f'Error restoring session {session_data.get("id")}: {session_error}', exc_info=True)
                            continue
                
                db.session.flush()
                current_app.logger.info(f'Step 3: Restored {restored_count} class sessions')
            except Exception as e:
                current_app.logger.error(f'Error restoring class management data: {e}', exc_info=True)
            
            # Step 4: Restore Result Management data
            try:
                from blueprints.result_management.models import RSession, RStudent, RSubject, RMark
                
                if archive_data.get('result_management') and archive_data['result_management'].get('sessions'):
                    for rs_data in archive_data['result_management']['sessions']:
                        try:
                            existing_rsession = RSession.query.get(rs_data.get('id'))
                            if existing_rsession:
                                existing_rsession.is_archived = False
                                # Restore marks if needed
                                if rs_data.get('marks'):
                                    for mark_data in rs_data['marks']:
                                        mark = RMark.query.filter_by(
                                            student_id=mark_data.get('student_id'),
                                            subject_id=mark_data.get('subject_id')
                                        ).first()
                                        if mark:
                                            mark.marks = mark_data.get('marks')
                                            mark.grade = mark_data.get('grade')
                        except Exception as rs_error:
                            current_app.logger.error(f'Error restoring result session {rs_data.get("id")}: {rs_error}', exc_info=True)
                            continue
                
                db.session.flush()
                current_app.logger.info(f'Step 4: Restored result management data')
            except Exception as e:
                current_app.logger.error(f'Error restoring result management data: {e}', exc_info=True)
            
            # Step 5: Restore Course Registrations
            try:
                if archive_data.get('course_registrations'):
                    for reg_data in archive_data['course_registrations']:
                        reg = StudentCourseRegistration.query.get(reg_data.get('id'))
                        if reg:
                            reg.status = 'active'  # Unarchive
            except Exception as e:
                current_app.logger.error(f'Error restoring course registrations: {e}', exc_info=True)
            
            # Step 6: Restore Exam Paper Evaluations
            try:
                from blueprints.class_management.models import ExamPaperEvaluation
                if archive_data.get('exam_evaluations'):
                    for ee_data in archive_data['exam_evaluations']:
                        ee = ExamPaperEvaluation.query.get(ee_data.get('id'))
                        if ee:
                            ee.archived = False
            except Exception as e:
                current_app.logger.error(f'Error restoring exam evaluations: {e}', exc_info=True)
            
            # Step 7: Restore Duty Assignments
            try:
                if archive_data.get('duty_assignments'):
                    for da_data in archive_data['duty_assignments']:
                        da = DutyAssignment.query.get(da_data.get('id'))
                        if da:
                            da.status = 'active'  # Unarchive
            except Exception as e:
                current_app.logger.error(f'Error restoring duty assignments: {e}', exc_info=True)
            
            # Step 8: Restore Remuneration Forms
            try:
                from blueprints.remuneration_management.models import RemunerationForm
                if archive_data.get('remuneration_forms'):
                    for form_data in archive_data['remuneration_forms']:
                        form = RemunerationForm.query.get(form_data.get('id'))
                        if form:
                            form.status = 'draft'  # Unarchive
                            form.archived_at = None
            except Exception as e:
                current_app.logger.error(f'Error restoring remuneration forms: {e}', exc_info=True)
            
            # Step 9: Mark archive as restored
            archive.is_active = False
            archive.restored_at = datetime.utcnow()
            archive.restored_by = current_user.full_name or current_user.username
            
            db.session.commit()
            
            current_app.logger.info(f'Session {academic_session} - {year} - {term} restored successfully. Restored {restored_count} class sessions.')
            
            return jsonify({
                'success': True,
                'message': f'Session {academic_session} - {year} - {term} restored successfully. {restored_count} class sessions restored. Current active data was automatically archived for safety.'
            })
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error restoring session: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error restoring session: {str(e)}'}), 500

    @app.route('/head/assign-duties')
    @login_required
    def assign_duties():
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            flash('This page is available only for Head or Dean accounts.', 'danger')
            return redirect(url_for('index'))
        
        # Get only internal teachers for duty assignment dropdown (excluding Head and external)
        from role_utils import get_teachers_excluding_head
        teachers = get_teachers_excluding_head(external_only=False)
        
        # Get all courses
        courses = Course.query.order_by(Course.course_code.asc()).all()
        
        # Get all students (for teaching assistant selection)
        students = Student.query.order_by(Student.student_id.asc()).all()
        
        # Show only relevant duty assignments (exclude tabulator/scrutinizer)
        visible_duty_types = ['course_coordinator', 'exam_committee_chief', 'routine_maker', 'teaching_assistant']
        assignments = DutyAssignment.query.filter(
            DutyAssignment.status == 'active',
            DutyAssignment.duty_type.in_(visible_duty_types)
        ).order_by(DutyAssignment.created_at.desc()).all()
        
        # Filter by active semester for exam_committee_chief assignments
        try:
            from utils.semester_utils import is_semester_active
            filtered_assignments = []
            for assignment in assignments:
                if assignment.duty_type == 'exam_committee_chief':
                    # Check if semester is active
                    if assignment.academic_session and assignment.year and assignment.term:
                        if is_semester_active(assignment.academic_session, assignment.year, assignment.term, batch=assignment.batch):
                            filtered_assignments.append(assignment)
                    # If no semester info, exclude it
                else:
                    # Not exam_committee_chief, include it
                    filtered_assignments.append(assignment)
            assignments = filtered_assignments
        except ImportError:
            pass
        
        # Get distinct academic sessions from curriculum year/term configuration
        academic_sessions = []
        try:
            session_rows = db.session.query(CurriculumYearTerm.academic_session).filter(
                CurriculumYearTerm.academic_session.isnot(None)
            ).distinct().order_by(CurriculumYearTerm.academic_session.desc()).all()
            academic_sessions = [row[0] for row in session_rows if row[0]]
        except Exception:
            academic_sessions = []
        
        # Get distinct batches from students
        batches = []
        try:
            all_batches = db.session.query(Student.batch).distinct().filter(
                Student.batch.isnot(None)
            ).order_by(Student.batch.desc()).all()
            batches = [batch[0] for batch in all_batches if batch[0]]
        except Exception:
            batches = []
        
        duty_label_map = {
            'course_coordinator': 'Course Co-Ordinator',
            'exam_committee_chief': 'Exam Committee Chief',
            'teaching_assistant': 'Teaching Assistant'
        }
        
        # Format remarks for display (parse JSON if needed)
        import json
        formatted_assignments = []
        for assignment in assignments:
            if assignment.remarks:
                try:
                    # Try to parse as JSON
                    if assignment.remarks.strip().startswith('{'):
                        parsed = json.loads(assignment.remarks)
                        formatted_remarks = []
                        if parsed.get('type') == 'chief':
                            formatted_remarks.append(f"<strong>Designation:</strong> {parsed.get('designation', 'N/A')}")
                            if parsed.get('institute'):
                                formatted_remarks.append(f"<strong>Institute:</strong> {parsed.get('institute')}")
                        elif parsed.get('type') == 'external':
                            formatted_remarks.append(f"<strong>Name:</strong> {parsed.get('name', 'N/A')}")
                            formatted_remarks.append(f"<strong>Designation:</strong> {parsed.get('designation', 'N/A')}")
                            if parsed.get('institute'):
                                formatted_remarks.append(f"<strong>Institute:</strong> {parsed.get('institute')}")
                        else:
                            for key, value in parsed.items():
                                formatted_remarks.append(f"<strong>{key.replace('_', ' ').title()}:</strong> {value}")
                        # Add formatted_remarks as a property
                        assignment.formatted_remarks = '<br>'.join(formatted_remarks)
                    else:
                        assignment.formatted_remarks = assignment.remarks
                except (json.JSONDecodeError, AttributeError, ValueError):
                    assignment.formatted_remarks = assignment.remarks
            else:
                assignment.formatted_remarks = None
            formatted_assignments.append(assignment)
        
        return render_template('head/assign_duties.html',
                             teachers=teachers,
                             students=students,
                             courses=courses,
                             assignments=formatted_assignments,
                             academic_sessions=academic_sessions,
                             batches=batches,
                             duty_label_map=duty_label_map)

    @app.route('/head/assign-duties/api/year-term', methods=['GET'])
    @login_required
    def head_assign_duty_year_term():
        """Get Year and Term options based on selected Academic Session"""
        from blueprints.class_management.models import Session
        
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        session_name = request.args.get('session', '').strip()
        if not session_name:
            return jsonify({'success': False, 'message': 'Session is required'}), 400
        
        try:
            # Get distinct year and term combinations from CurriculumYearTerm for this session
            year_term_rows = db.session.query(
                CurriculumYearTerm.year,
                CurriculumYearTerm.term
            ).filter(
                CurriculumYearTerm.academic_session == session_name
            ).distinct().order_by(
                CurriculumYearTerm.year.asc(),
                CurriculumYearTerm.term.asc()
            ).all()
            
            # Also check Session model for year/term combinations
            session_rows = db.session.query(
                Session.year,
                Session.term
            ).filter(
                Session.academic_session == session_name
            ).distinct().all()
            
            # Combine and deduplicate
            year_term_set = set()
            for year, term in year_term_rows:
                if year and term:
                    year_term_set.add((year, term))
            for year, term in session_rows:
                if year and term:
                    year_term_set.add((year, term))
            
            # Convert to sorted list
            year_term_list = sorted(list(year_term_set), key=lambda x: (x[0], x[1]))
            
            years = sorted(list(set([yt[0] for yt in year_term_list if yt[0]])), key=lambda x: (
                'First' if x == 'First' else
                'Second' if x == 'Second' else
                'Third' if x == 'Third' else
                'Fourth' if x == 'Fourth' else
                'Fifth' if x == 'Fifth' else
                'LLM' if x == 'LLM' else x
            ))
            
            terms = sorted(list(set([yt[1] for yt in year_term_list if yt[1]])), key=lambda x: (
                'First' if x == 'First' else
                'Second' if x == 'Second' else
                'Thesis Term' if x == 'Thesis Term' else x
            ))
            
            return jsonify({
                'success': True,
                'years': years,
                'terms': terms
            })
        except Exception as e:
            current_app.logger.error(f'Error getting year-term options: {e}', exc_info=True)
            return jsonify({'success': False, 'message': 'Error fetching year-term options'}), 500

    @app.route('/head/assign-duties/api/courses', methods=['GET'])
    @login_required
    def head_assign_duty_courses():
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles and not (_is_exam_committee_chief() or _is_exam_committee_member()):
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        batch = request.args.get('batch', '').strip()
        if not batch:
            return jsonify({'success': False, 'message': 'Batch is required'}), 400
        
        try:
            curriculum_ids = []
            if Curriculum:
                for curriculum in Curriculum.query.all():
                    if batch in curriculum.get_batches_list():
                        curriculum_ids.append(curriculum.id)
            
            course_query = Course.query
            if curriculum_ids:
                course_query = course_query.filter(Course.curriculum_id.in_(curriculum_ids))
            
            courses = course_query.order_by(Course.course_code.asc()).all()
            course_payload = [{
                'id': course.id,
                'code': course.course_code,
                'name': course.course_name,
                'year': course.display_year,
                'term': course.display_term
            } for course in courses]
            
            return jsonify({'success': True, 'courses': course_payload})
        except Exception as exc:
            current_app.logger.error(f'Failed to load courses for batch {batch}: {exc}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to load courses'}), 500

    @app.route('/head/assign-duties/api/assign', methods=['POST'])
    @login_required
    def assign_duty_api():
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        course_id = data.get('course_id')
        if course_id == '' or course_id is None:
            course_id = None
        else:
            try:
                course_id = int(course_id)
            except (TypeError, ValueError):
                course_id = None
        raw_course_ids = data.get('course_ids') or []
        if isinstance(raw_course_ids, str):
            raw_course_ids = [raw_course_ids]
        course_ids = []
        for cid in raw_course_ids:
            try:
                course_ids.append(int(cid))
            except (TypeError, ValueError):
                continue
        course_code = (data.get('course_code', '') or '').strip()
        course_name = (data.get('course_name', '') or '').strip()
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        batch = data.get('batch', '').strip()
        # Normalize empty batch to None
        if batch == '':
            batch = None
        duty_type = data.get('duty_type', '').strip()
        
        teacher_id = data.get('teacher_id')
        if teacher_id == '' or teacher_id is None:
            teacher_id = None
        else:
            try:
                teacher_id = int(teacher_id)
            except (TypeError, ValueError):
                teacher_id = None
        
        student_id = data.get('student_id')
        if student_id == '' or student_id is None:
            student_id = None
        else:
            try:
                student_id = int(student_id)
            except (TypeError, ValueError):
                student_id = None
        
        remarks = data.get('remarks', '').strip()
        
        if course_id and Course and not course_code:
            course_obj = Course.query.get(course_id)
            if course_obj:
                course_code = course_obj.course_code
                course_name = course_obj.course_name

        course_payloads = []
        if course_ids and Course:
            selected_courses = Course.query.filter(Course.id.in_(course_ids)).all()
            for course in selected_courses:
                course_payloads.append({
                    'course_id': course.id,
                    'course_code': (course.course_code or '').upper(),
                    'course_name': course.course_name
                })

        course_code = (course_code or '').strip().upper() or None

        exam_entry_ids = data.get('exam_entry_ids') or []
        if isinstance(exam_entry_ids, str):
            exam_entry_ids = [exam_entry_ids]

        if duty_type not in {'course_coordinator', 'exam_committee_chief', 'tabulator', 'teaching_assistant', 'routine_maker'}:
            return jsonify({'success': False, 'message': 'Unsupported duty type'}), 400
        
        if duty_type in {'course_coordinator', 'exam_committee_chief', 'tabulator', 'routine_maker'} and not teacher_id:
            return jsonify({'success': False, 'message': 'Please select a teacher'}), 400
        
        if duty_type == 'course_coordinator':
            if not batch:
                return jsonify({'success': False, 'message': 'Batch is required for Course Co-Ordinator assignment'}), 400
        elif duty_type == 'tabulator':
            missing = []
            if not academic_session:
                missing.append('Academic Session')
            if not year:
                missing.append('Year')
            if not term:
                missing.append('Term')
            if missing:
                return jsonify({'success': False, 'message': f'Tabulator assignment requires: {", ".join(missing)}'}), 400
        elif duty_type == 'teaching_assistant':
            if not student_id:
                return jsonify({'success': False, 'message': 'Please select a student for Teaching Assistant duty'}), 400
        elif duty_type == 'exam_committee_chief':
            missing = []
            if not academic_session:
                missing.append('Academic Session')
            if not year:
                missing.append('Year')
            if not term:
                missing.append('Term')
            if missing:
                return jsonify({'success': False, 'message': f'Exam Committee Chief assignment requires: {", ".join(missing)}'}), 400
            
            # Check if semester is active
            # For Exam Committee Chief, batch is not required, so pass None
            try:
                from utils.semester_utils import is_semester_active
                # Exam Committee Chief doesn't need batch, so check without batch filter
                is_active = is_semester_active(academic_session, year, term, batch=None)
                current_app.logger.info(f'Checking active semester for Exam Committee Chief: session={academic_session}, year={year}, term={term}, is_active={is_active}')
                if not is_active:
                    return jsonify({
                        'success': False, 
                        'message': 'This semester is not active. Please activate it in Active Semester Management first.'
                    }), 400
            except Exception as e:
                current_app.logger.error(f'Error checking active semester: {e}', exc_info=True)
                # Continue anyway - don't block if there's an error
        
        try:
            # Non-scrutinizer duties
            filter_kwargs = {
                'duty_type': duty_type,
                'status': 'active'
            }
            if duty_type == 'course_coordinator':
                filter_kwargs.update({
                    'assigned_teacher_id': teacher_id,
                    'batch': batch or None
                })
            elif duty_type == 'tabulator':
                filter_kwargs.update({
                    'assigned_teacher_id': teacher_id,
                    'academic_session': academic_session or None,
                    'year': year or None,
                    'term': term or None
                })
            elif duty_type == 'teaching_assistant':
                filter_kwargs.update({
                    'student_id': student_id
                })
            elif duty_type == 'exam_committee_chief':
                filter_kwargs.update({
                    'assigned_teacher_id': teacher_id,
                    'academic_session': academic_session or None,
                    'year': year or None,
                    'term': term or None
                })
            elif duty_type == 'routine_maker':
                filter_kwargs.update({
                    'assigned_teacher_id': teacher_id
                })
            
            existing = DutyAssignment.query.filter_by(**filter_kwargs).first()
            if existing:
                return jsonify({'success': False, 'message': 'This duty assignment already exists'}), 400
            
            assignment = DutyAssignment(
                course_id=course_id,
                course_code=course_code,
                course_name=course_name or None,
                academic_session=academic_session or None,
                year=year or None,
                term=term or None,
                batch=batch if duty_type == 'course_coordinator' else None,
                duty_type=duty_type,
                assigned_teacher_id=teacher_id if duty_type != 'teaching_assistant' else None,
                student_id=student_id if duty_type == 'teaching_assistant' else None,
                assigned_by_id=current_user.id,
                remarks=remarks or None,
                status='active'
            )
            db.session.add(assignment)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Duty assigned successfully',
                'assignment_id': assignment.id
            })
        except Exception as exc:
            db.session.rollback()
            current_app.logger.error(f'Failed to assign duty: {exc}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to assign duty'}), 500

    @app.route('/head/assign-duties/api/archive-committee/<int:assignment_id>', methods=['POST'])
    @login_required
    def archive_exam_committee_from_assign_duties(assignment_id):
        """Archive an Exam Committee Chief assignment and all related data"""
        roles = parse_roles(current_user.role)
        if 'head' not in roles and 'dean' not in roles:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            # Get the chief assignment
            chief_assignment = DutyAssignment.query.filter_by(
                id=assignment_id,
                duty_type='exam_committee_chief',
                status='active'
            ).first()
            
            if not chief_assignment:
                return jsonify({'success': False, 'message': 'Exam Committee Chief assignment not found'}), 404
            
            # Verify session/year/term match
            if (chief_assignment.academic_session != academic_session or 
                chief_assignment.year != year or 
                chief_assignment.term != term):
                return jsonify({'success': False, 'message': 'Session/Year/Term mismatch'}), 400
            
            archived_count = 0
            
            # Archive the chief assignment
            chief_assignment.status = 'archived'
            archived_count += 1
            
            # Archive all committee members for this session/year/term
            committee_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'active'
            ).all()
            
            for assignment in committee_assignments:
                assignment.status = 'archived'
                archived_count += 1
            
            # Archive tabulators for this session/year/term assigned by this chief
            tabulator_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'tabulator',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.assigned_by_id == chief_assignment.assigned_teacher_id,
                DutyAssignment.status == 'active'
            ).all()
            
            for tabulator in tabulator_assignments:
                tabulator.status = 'archived'
                archived_count += 1
            
            # Archive scrutinizers for this session/year/term assigned by this chief
            scrutinizer_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'scrutinizer',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.assigned_by_id == chief_assignment.assigned_teacher_id,
                DutyAssignment.status == 'active'
            ).all()
            
            for scrutinizer in scrutinizer_assignments:
                scrutinizer.status = 'archived'
                archived_count += 1
            
            # Archive remuneration forms for this session/year/term
            from blueprints.remuneration_management.models import RemunerationForm
            from datetime import datetime
            import json
            
            # Get the chief's user_id
            chief_teacher = chief_assignment.assigned_teacher
            if chief_teacher:
                # Find user by teacher name
                chief_user = User.query.filter_by(full_name=chief_teacher.name).first()
                if chief_user:
                    remuneration_forms = RemunerationForm.query.filter_by(
                        user_id=chief_user.id,
                        status='draft'
                    ).all()
                    
                    for form in remuneration_forms:
                        # Check if form data matches the session/year/term
                        try:
                            if form.form_data:
                                form_data = json.loads(form.form_data)
                                form_session = form_data.get('academic_session') or form.academic_year
                                form_year = form_data.get('year') or form.year
                                form_term = form_data.get('term') or form.term
                                
                                if (form_session == academic_session and 
                                    form_year == year and 
                                    form_term == term):
                                    form.status = 'archived'
                                    form.archived_at = datetime.utcnow()
                                    archived_count += 1
                            else:
                                # If form_data parsing fails, check direct fields
                                if (form.academic_year == academic_session and 
                                    form.year == year and 
                                    form.term == term):
                                    form.status = 'archived'
                                    form.archived_at = datetime.utcnow()
                                    archived_count += 1
                        except:
                            pass
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Exam Committee archived successfully for {academic_session} - {year} - {term}. {archived_count} item(s) archived.'
            })
        
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error archiving exam committee: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error archiving exam committee: {str(e)}'}), 500

    @app.route('/head/assign-duties/api/remove/<int:assignment_id>', methods=['POST'])
    @login_required
    def remove_duty_api(assignment_id):
        roles = parse_roles(current_user.role)
        # Check if user is Head/Dean or Exam Committee Chief
        is_head = 'head' in roles or 'dean' in roles
        is_chief = _is_exam_committee_chief()
        is_member = _is_exam_committee_member()

        if not is_head and not is_chief and not is_member:
            return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        try:
            assignment = DutyAssignment.query.get_or_404(assignment_id)
            
            # If Exam Committee Chief/Member, only allow removing assignments they created
            if (is_chief or is_member) and not is_head:
                if assignment.assigned_by_id != current_user.id:
                    return jsonify({'success': False, 'message': 'You can only remove assignments you created'}), 403
                # Only allow removing tabulator and scrutinizer assignments
                if assignment.duty_type not in {'tabulator', 'scrutinizer'}:
                    return jsonify({'success': False, 'message': 'You can only remove Tabulator and Scrutinizer assignments'}), 403
            
            # If removing Exam Committee Chief, delete all related data
            if assignment.duty_type == 'exam_committee_chief' and is_head:
                academic_session = assignment.academic_session
                year = assignment.year
                term = assignment.term
                chief_teacher_id = assignment.assigned_teacher_id
                
                # Archive/Delete all committee members for this session/year/term
                # Delete ALL members regardless of status to ensure complete cleanup
                committee_members = DutyAssignment.query.filter(
                    DutyAssignment.duty_type == 'exam_committee_member',
                    DutyAssignment.academic_session == academic_session,
                    DutyAssignment.year == year,
                    DutyAssignment.term == term
                ).all()
                for member in committee_members:
                    member.status = 'inactive'
                    # Force update to ensure it's marked as inactive
                    db.session.add(member)
                
                # Archive/Delete all tabulators for this session/year/term assigned by this chief
                tabulators = DutyAssignment.query.filter(
                    DutyAssignment.duty_type == 'tabulator',
                    DutyAssignment.academic_session == academic_session,
                    DutyAssignment.year == year,
                    DutyAssignment.term == term,
                    DutyAssignment.assigned_by_id == chief_teacher_id,
                    DutyAssignment.status == 'active'
                ).all()
                for tabulator in tabulators:
                    tabulator.status = 'inactive'
                
                # Archive/Delete all scrutinizers for this session/year/term assigned by this chief
                scrutinizers = DutyAssignment.query.filter(
                    DutyAssignment.duty_type == 'scrutinizer',
                    DutyAssignment.academic_session == academic_session,
                    DutyAssignment.year == year,
                    DutyAssignment.term == term,
                    DutyAssignment.assigned_by_id == chief_teacher_id,
                    DutyAssignment.status == 'active'
                ).all()
                for scrutinizer in scrutinizers:
                    scrutinizer.status = 'inactive'
                    _clear_scrutinizer_assignment(scrutinizer)
                
                # Archive remuneration forms for this session/year/term
                from blueprints.remuneration_management.models import RemunerationForm
                from datetime import datetime
                import json
                
                chief_teacher = assignment.assigned_teacher
                if chief_teacher:
                    chief_user = User.query.filter_by(full_name=chief_teacher.name).first()
                    if chief_user:
                        remuneration_forms = RemunerationForm.query.filter_by(
                            user_id=chief_user.id,
                            status='draft'
                        ).all()
                        
                        for form in remuneration_forms:
                            try:
                                if form.form_data:
                                    form_data = json.loads(form.form_data)
                                    form_session = form_data.get('academic_session') or form.academic_year
                                    form_year = form_data.get('year') or form.year
                                    form_term = form_data.get('term') or form.term
                                    
                                    if (form_session == academic_session and 
                                        form_year == year and 
                                        form_term == term):
                                        form.status = 'archived'
                                        form.archived_at = datetime.utcnow()
                                else:
                                    if (form.academic_year == academic_session and 
                                        form.year == year and 
                                        form.term == term):
                                        form.status = 'archived'
                                        form.archived_at = datetime.utcnow()
                            except:
                                pass
            
            was_scrutinizer = assignment.duty_type == 'scrutinizer'
            assignment.status = 'inactive'
            db.session.commit()
            if was_scrutinizer:
                _clear_scrutinizer_assignment(assignment)
            
            return jsonify({'success': True, 'message': 'Duty assignment removed successfully'})
        except Exception as exc:
            db.session.rollback()
            current_app.logger.error(f'Failed to remove duty: {exc}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to remove duty assignment'}), 500

    @app.route('/exam-committee-management')
    @login_required
    def exam_committee_management():
        """Intermediate page showing Exam Committee Chief and Member cards"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            flash('Teacher profile not found.', 'danger')
            return redirect(url_for('index'))
        
        # Check if user has any active chief assignments
        active_chief_assignments = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_chief',
            DutyAssignment.assigned_teacher_id == teacher.id,
            DutyAssignment.status == 'active'
        ).count()
        is_chief = active_chief_assignments > 0
        
        is_member = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_member',
            DutyAssignment.assigned_teacher_id == teacher.id,
            DutyAssignment.status == 'active'
        ).first() is not None
        
        if not is_chief and not is_member:
            flash('You are not assigned as Exam Committee Chief or Member.', 'danger')
            return redirect(url_for('index'))
        
        # Get chief assignments - only active ones, no duplicates
        chief_assignments = []
        chief_saved_forms = {}
        if is_chief:
            # Get all active assignments, then deduplicate by session/year/term (keep latest)
            all_chief_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_chief',
                DutyAssignment.assigned_teacher_id == teacher.id,
                DutyAssignment.status == 'active'  # Only active assignments
            ).order_by(
                DutyAssignment.academic_session.desc(),
                DutyAssignment.year.asc(),
                DutyAssignment.term.asc(),
                DutyAssignment.updated_at.desc(),
                DutyAssignment.created_at.desc()
            ).all()
            
            # Deduplicate: keep only the latest assignment for each session/year/term combination
            # Also mark older duplicates as inactive
            seen_combinations = {}
            duplicates_to_inactivate = []
            
            for assignment in all_chief_assignments:
                key = f"{assignment.academic_session}_{assignment.year}_{assignment.term}"
                if key not in seen_combinations:
                    seen_combinations[key] = assignment
                    chief_assignments.append(assignment)
                else:
                    # Duplicate found - determine which one to keep
                    existing = seen_combinations[key]
                    keep_new = False
                    
                    if assignment.updated_at and existing.updated_at:
                        if assignment.updated_at > existing.updated_at:
                            keep_new = True
                    elif assignment.updated_at and not existing.updated_at:
                        keep_new = True
                    elif assignment.created_at and existing.created_at:
                        if assignment.created_at > existing.created_at:
                            keep_new = True
                    
                    if keep_new:
                        # Keep new one, mark old one as inactive
                        duplicates_to_inactivate.append(existing)
                        chief_assignments.remove(existing)
                        chief_assignments.append(assignment)
                        seen_combinations[key] = assignment
                    else:
                        # Keep existing one, mark new one as inactive
                        duplicates_to_inactivate.append(assignment)
            
            # Inactivate duplicate assignments
            if duplicates_to_inactivate:
                for dup in duplicates_to_inactivate:
                    dup.status = 'inactive'
                try:
                    db.session.commit()
                except:
                    db.session.rollback()
            
            # Sort again after deduplication
            chief_assignments.sort(key=lambda x: (
                x.academic_session or '',
                x.year or '',
                x.term or ''
            ), reverse=True)
            
            # Get saved remuneration forms for chief
            from blueprints.remuneration_management.models import RemunerationForm
            for assignment in chief_assignments:
                forms = RemunerationForm.query.filter_by(
                    user_id=current_user.id,
                    academic_year=assignment.academic_session,
                    year=assignment.year,
                    term=assignment.term,
                    status='draft'
                ).order_by(RemunerationForm.created_at.desc()).all()
                key = f"{assignment.academic_session}_{assignment.year}_{assignment.term}"
                chief_saved_forms[key] = forms
        
        # Get member assignments and saved forms if user is a member - no duplicates
        member_assignments = []
        member_saved_forms = {}
        if is_member:
            # Get all active member assignments, then deduplicate by session/year/term
            # Only fetch active assignments - inactive/archived ones should not appear
            all_member_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.assigned_teacher_id == teacher.id,
                DutyAssignment.status.in_(['active'])  # Explicitly only active
            ).order_by(
                DutyAssignment.academic_session.desc(),
                DutyAssignment.year.asc(),
                DutyAssignment.term.asc(),
                DutyAssignment.updated_at.desc(),
                DutyAssignment.created_at.desc()
            ).all()
            
            # Deduplicate: keep only the latest assignment for each session/year/term combination
            # Also mark older duplicates as inactive
            seen_combinations = {}
            duplicates_to_inactivate = []
            
            for assignment in all_member_assignments:
                key = f"{assignment.academic_session}_{assignment.year}_{assignment.term}"
                if key not in seen_combinations:
                    seen_combinations[key] = assignment
                    member_assignments.append(assignment)
                else:
                    # Duplicate found - determine which one to keep
                    existing = seen_combinations[key]
                    keep_new = False
                    
                    if assignment.updated_at and existing.updated_at:
                        if assignment.updated_at > existing.updated_at:
                            keep_new = True
                    elif assignment.updated_at and not existing.updated_at:
                        keep_new = True
                    elif assignment.created_at and existing.created_at:
                        if assignment.created_at > existing.created_at:
                            keep_new = True
                    
                    if keep_new:
                        # Keep new one, mark old one as inactive
                        duplicates_to_inactivate.append(existing)
                        member_assignments.remove(existing)
                        member_assignments.append(assignment)
                        seen_combinations[key] = assignment
                    else:
                        # Keep existing one, mark new one as inactive
                        duplicates_to_inactivate.append(assignment)
            
            # Inactivate duplicate assignments
            if duplicates_to_inactivate:
                for dup in duplicates_to_inactivate:
                    dup.status = 'inactive'
                try:
                    db.session.commit()
                except:
                    db.session.rollback()
            
            # Sort again after deduplication
            member_assignments.sort(key=lambda x: (
                x.academic_session or '',
                x.year or '',
                x.term or ''
            ), reverse=True)
            
            # Get saved remuneration forms for this member - show Chief's forms for same assignment
            from blueprints.remuneration_management.models import RemunerationForm
            for assignment in member_assignments:
                # Normalize assignment values for comparison
                assignment_session = str(assignment.academic_session or '').strip()
                assignment_year = str(assignment.year or '').strip()
                assignment_term = str(assignment.term or '').strip()
                
                # Find the Chief for this assignment
                chief_assignment = DutyAssignment.query.filter_by(
                    duty_type='exam_committee_chief',
                    academic_session=assignment.academic_session,
                    year=assignment.year,
                    term=assignment.term,
                    status='active'
                ).first()
                
                # Get Chief's saved forms for this assignment
                chief_forms = []
                if chief_assignment and chief_assignment.assigned_teacher:
                    chief_user = User.query.filter_by(full_name=chief_assignment.assigned_teacher.name).first()
                    if chief_user:
                        # Get all forms for chief user
                        all_chief_forms = RemunerationForm.query.filter_by(
                            user_id=chief_user.id
                        ).order_by(RemunerationForm.created_at.desc()).all()
                        
                        # Filter forms that match this assignment (case-insensitive, trimmed)
                        for form in all_chief_forms:
                            form_session = str(form.academic_year or '').strip()
                            form_year = str(form.year or '').strip()
                            form_term = str(form.term or '').strip()
                            
                            # Match session (case-insensitive, exact or contains)
                            session_match = False
                            if form_session and assignment_session:
                                form_session_lower = form_session.lower()
                                assignment_session_lower = assignment_session.lower()
                                session_match = (form_session_lower == assignment_session_lower or 
                                               form_session_lower in assignment_session_lower or 
                                               assignment_session_lower in form_session_lower)
                            
                            # Match year (case-insensitive, exact or contains)
                            year_match = False
                            if form_year and assignment_year:
                                form_year_lower = form_year.lower()
                                assignment_year_lower = assignment_year.lower()
                                year_match = (form_year_lower == assignment_year_lower or 
                                            form_year_lower in assignment_year_lower or 
                                            assignment_year_lower in form_year_lower)
                            
                            # Match term (case-insensitive, exact or contains)
                            term_match = False
                            if form_term and assignment_term:
                                form_term_lower = form_term.lower()
                                assignment_term_lower = assignment_term.lower()
                                term_match = (form_term_lower == assignment_term_lower or 
                                            form_term_lower in assignment_term_lower or 
                                            assignment_term_lower in form_term_lower)
                            
                            if session_match and year_match and term_match:
                                chief_forms.append(form)
                
                key = f"{assignment.academic_session}_{assignment.year}_{assignment.term}"
                member_saved_forms[key] = chief_forms
                
                # Debug logging
                if chief_forms:
                    current_app.logger.info(f"Found {len(chief_forms)} chief forms for member {current_user.id}, assignment: {key}")
                else:
                    current_app.logger.info(f"No chief forms found for member {current_user.id}, assignment: {key}")
        
        return render_template('exam_committee_management/index.html', 
                             is_chief=is_chief, 
                             is_member=is_member,
                             chief_assignments=chief_assignments,
                             chief_saved_forms=chief_saved_forms,
                             member_assignments=member_assignments,
                             member_saved_forms=member_saved_forms)

    @app.route('/exam-committee-member/dashboard')
    @login_required
    def exam_committee_member_dashboard():
        """Dashboard for Exam Committee Members to view their assignments and remuneration forms"""
        # Check if current user is assigned as Exam Committee Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            flash('Teacher profile not found.', 'warning')
            return redirect(url_for('index'))
        
        # Get all active member assignments
        member_assignments = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_member',
            DutyAssignment.assigned_teacher_id == teacher.id,
            DutyAssignment.status == 'active'
        ).order_by(
            DutyAssignment.academic_session.desc(),
            DutyAssignment.year.asc(),
            DutyAssignment.term.asc()
        ).all()
        
        if not member_assignments:
            flash('You are not assigned as Exam Committee Member.', 'danger')
            return redirect(url_for('index'))
        
        # Get saved remuneration forms for this member
        from blueprints.remuneration_management.models import RemunerationForm
        saved_forms = {}
        for assignment in member_assignments:
            forms = RemunerationForm.query.filter_by(
                user_id=current_user.id,
                academic_year=assignment.academic_session,
                year=assignment.year,
                term=assignment.term,
                status='draft'
            ).order_by(RemunerationForm.created_at.desc()).all()
            key = f"{assignment.academic_session}_{assignment.year}_{assignment.term}"
            saved_forms[key] = forms
        
        return render_template('exam_committee_member/dashboard.html',
                             member_assignments=member_assignments,
                             saved_forms=saved_forms)

    @app.route('/exam-committee-chief/dashboard')
    @login_required
    def exam_committee_chief_dashboard():
        """Dashboard for Exam Committee Chief to manage tabulators and scrutinizers"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            flash('Teacher profile not found.', 'warning')
            return redirect(url_for('index'))
        
        # Get session/year/term from URL parameters
        url_session = request.args.get('session', '').strip()
        url_year = request.args.get('year', '').strip()
        url_term = request.args.get('term', '').strip()
        
        # Find matching chief assignment
        chief_assignment = None
        if url_session and url_year and url_term:
            # Find assignment matching URL parameters
            chief_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                assigned_teacher_id=teacher.id,
                academic_session=url_session,
                year=url_year,
                term=url_term,
                status='active'
            ).first()
        
        # If no match found, use first assignment (for backward compatibility)
        if not chief_assignment:
            chief_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                assigned_teacher_id=teacher.id,
                status='active'
            ).first()
        
        # Check if user is an internal member (not external)
        member_assignments = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_member',
            DutyAssignment.assigned_teacher_id == teacher.id,
            DutyAssignment.status == 'active'
        ).all()
        
        # Filter out external members (those without assigned_teacher_id)
        internal_member_assignments = [ma for ma in member_assignments if ma.assigned_teacher_id]
        
        if not chief_assignment and not internal_member_assignments:
            flash('You are not assigned as Exam Committee Chief or Member.', 'danger')
            return redirect(url_for('index'))
        
        # Get all teachers for assignment (excluding Head of the Discipline)
        from role_utils import get_teachers_excluding_head
        teachers = get_teachers_excluding_head()
        # External teachers/options for External Members dropdown.
        # Source: Admin Users list (teacher-role users linked to Teacher and marked is_external=True).
        external_teacher_options = []
        seen_external_names = set()
        user_teacher_rows = db.session.query(User, Teacher).join(
            Teacher, User.teacher_id == Teacher.id
        ).filter(
            User.teacher_id.isnot(None),
            Teacher.is_external.is_(True)
        ).order_by(User.full_name.asc()).all()

        for user_row, teacher_row in user_teacher_rows:
            # Keep only teacher-role users from admin user list.
            if 'teacher' not in parse_roles(getattr(user_row, 'role', None)):
                continue
            ext_name = (user_row.full_name or teacher_row.name or '').strip()
            if not ext_name:
                continue
            key = ext_name.lower()
            if key in seen_external_names:
                continue
            seen_external_names.add(key)
            external_teacher_options.append({
                'name': ext_name,
                'designation': teacher_row.designation or '',
                'institute': teacher_row.institute or ''
            })
        # Convert teachers to dictionaries for JSON serialization
        teachers_dict = [{'id': t.id, 'name': t.name} for t in teachers]
        
        # Get distinct academic sessions
        from blueprints.class_management.models import Session
        sessions = db.session.query(Session.academic_session).distinct().filter(
            Session.academic_session.isnot(None)
        ).order_by(Session.academic_session.desc()).all()
        academic_sessions = [s[0] for s in sessions if s[0]]

        # Determine active committee context.
        # If the user is not chief for this context, allow internal-member context.
        member_context_assignment = None
        if url_session and url_year and url_term:
            member_context_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_member',
                assigned_teacher_id=teacher.id,
                academic_session=url_session,
                year=url_year,
                term=url_term,
                status='active'
            ).first()
        if not member_context_assignment and internal_member_assignments:
            member_context_assignment = internal_member_assignments[0]

        # Use chief context first, otherwise member context.
        current_session = None
        current_year = None
        current_term = None
        current_batch = None
        active_context_assignment = chief_assignment or member_context_assignment
        chief_assignment_details = chief_assignment
        if active_context_assignment:
            current_session = active_context_assignment.academic_session
            current_year = active_context_assignment.year
            current_term = active_context_assignment.term
            current_batch = active_context_assignment.batch

        # For member-only users, still load chief details for the same committee context.
        if not chief_assignment_details and current_session and current_year and current_term:
            chief_assignment_details = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                academic_session=current_session,
                year=current_year,
                term=current_term,
                status='active'
            ).first()
            if chief_assignment_details and chief_assignment_details.batch and not current_batch:
                current_batch = chief_assignment_details.batch
        
        # Get assigned tabulators - filter by current committee's session/year/term
        tabulators_query = DutyAssignment.query.filter_by(
            duty_type='tabulator',
            status='active',
            assigned_by_id=current_user.id
        )
        # If we have current session/year/term from chief assignment, filter by them
        if current_session and current_year and current_term:
            tabulators_query = tabulators_query.filter_by(
                academic_session=current_session,
                year=current_year,
                term=current_term
            )
        tabulators = tabulators_query.order_by(DutyAssignment.created_at.desc()).all()
        
        # Get assigned scrutinizers - filter by current committee's session/year/term
        scrutinizers_query = DutyAssignment.query.filter_by(
            duty_type='scrutinizer',
            status='active',
            assigned_by_id=current_user.id
        )
        # If we have current session/year/term from chief assignment, filter by them
        if current_session and current_year and current_term:
            scrutinizers_query = scrutinizers_query.filter_by(
                academic_session=current_session,
                year=current_year,
                term=current_term
            )
        scrutinizers = scrutinizers_query.order_by(DutyAssignment.created_at.desc()).all()
        
        available_entries = []
        if current_session and current_year and current_term:
            # Check if semester is active before fetching entries
            try:
                from utils.semester_utils import is_semester_active
                if not is_semester_active(current_session, current_year, current_term, batch=None):
                    # Semester is inactive, skip fetching entries
                    available_entries = []
                else:
                    submitted_entries = ExamPaperEvaluation.query.filter(
                        ExamPaperEvaluation.archived.is_(False),
                        ExamPaperEvaluation.submitted_to_committee.is_(True),
                        ExamPaperEvaluation.academic_session == current_session,
                        ExamPaperEvaluation.year == current_year,
                        ExamPaperEvaluation.term == current_term,
                        ExamPaperEvaluation.assigned_scrutinizer_id.is_(None)  # Only show entries without assigned scrutinizer
                    ).order_by(ExamPaperEvaluation.created_at.desc()).all()
                    for entry in submitted_entries:
                        owner_name = entry.owner_teacher.name if entry.owner_teacher else 'N/A'
                        section_label = entry.section or 'Full'
                        available_entries.append({
                            'id': entry.id,
                            'course_code': entry.course_code,
                            'course_name': entry.course_name,
                            'section': section_label,
                            'owner': owner_name,
                            'owner_teacher_id': entry.owner_teacher_id  # Include owner ID for filtering
                        })
            except Exception as e:
                current_app.logger.error(f'Error checking active semester: {e}', exc_info=True)
                # Continue with fetching entries if error
                submitted_entries = ExamPaperEvaluation.query.filter(
                    ExamPaperEvaluation.archived.is_(False),
                    ExamPaperEvaluation.submitted_to_committee.is_(True),
                    ExamPaperEvaluation.academic_session == current_session,
                    ExamPaperEvaluation.year == current_year,
                    ExamPaperEvaluation.term == current_term,
                    ExamPaperEvaluation.assigned_scrutinizer_id.is_(None)  # Only show entries without assigned scrutinizer
                ).order_by(ExamPaperEvaluation.created_at.desc()).all()
                for entry in submitted_entries:
                    owner_name = entry.owner_teacher.name if entry.owner_teacher else 'N/A'
                    section_label = entry.section or 'Full'
                    available_entries.append({
                        'id': entry.id,
                        'course_code': entry.course_code,
                        'course_name': entry.course_name,
                        'section': section_label,
                        'owner': owner_name,
                        'owner_teacher_id': entry.owner_teacher_id  # Include owner ID for filtering
                    })
        
        # Get saved committee members for current session/year/term
        saved_committee_members = {
            'internal': [],
            'external': [],
            'remarks': None
        }
        if current_session and current_year and current_term:
            # Get all active committee members for this session/year/term
            # Don't filter by assigned_by_id because restored committees may have different assigned_by_id
            committee_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == current_session,
                DutyAssignment.year == current_year,
                DutyAssignment.term == current_term,
                DutyAssignment.status == 'active'
            ).all()
            seen_internal_ids = set()
            seen_external_names = set()
            
            for assignment in committee_assignments:
                if assignment.assigned_teacher_id:
                    if assignment.assigned_teacher_id in seen_internal_ids:
                        continue
                    seen_internal_ids.add(assignment.assigned_teacher_id)
                    # Internal member
                    designation = ''
                    institute = ''
                    # Try to get designation from remarks JSON
                    try:
                        member_info = json.loads(assignment.remarks) if assignment.remarks else {}
                        if member_info.get('type') == 'internal':
                            designation = member_info.get('designation', '')
                            institute = member_info.get('institute', '')
                    except:
                        pass
                    
                    # Get designation and institute from Teacher table if not in remarks
                    teacher = assignment.assigned_teacher
                    if teacher:
                        if not designation and teacher.designation:
                            designation = teacher.designation
                        if not institute and teacher.institute:
                            institute = teacher.institute or 'Law Discipline, KU'
                    
                    saved_committee_members['internal'].append({
                        'id': assignment.assigned_teacher_id,
                        'name': assignment.assigned_teacher.name if assignment.assigned_teacher else 'N/A',
                        'designation': designation,
                        'institute': institute or 'Law Discipline, KU'
                    })
                else:
                    # External member (info stored in remarks as JSON)
                    try:
                        external_info = json.loads(assignment.remarks) if assignment.remarks else {}
                        if external_info.get('type') == 'external':
                            external_name = (external_info.get('name') or '').strip()
                            if not external_name:
                                continue
                            external_name_key = external_name.lower()
                            if external_name_key in seen_external_names:
                                continue
                            seen_external_names.add(external_name_key)
                            saved_committee_members['external'].append({
                                'name': external_name,
                                'designation': external_info.get('designation', ''),
                                'institute': external_info.get('institute', '')
                            })
                    except:
                        pass
                
                # Get remarks from first internal member assignment (not external/internal JSON)
                if assignment.assigned_teacher_id and saved_committee_members['remarks'] is None:
                    try:
                        member_info = json.loads(assignment.remarks) if assignment.remarks else {}
                        if member_info.get('type') != 'internal':
                            if assignment.remarks and not assignment.remarks.strip().startswith('{'):
                                saved_committee_members['remarks'] = assignment.remarks
                    except:
                        if assignment.remarks and not assignment.remarks.strip().startswith('{'):
                            saved_committee_members['remarks'] = assignment.remarks
        
        # Get chief info with designation
        chief_info = None
        if chief_assignment_details:
            designation = ''
            # Try to get chief designation from remarks JSON
            try:
                chief_data = json.loads(chief_assignment_details.remarks) if chief_assignment_details.remarks else {}
                if chief_data.get('type') == 'chief':
                    designation = chief_data.get('designation', '')
            except:
                pass
            
            # Get from Teacher table if not in remarks
            chief_teacher = chief_assignment_details.assigned_teacher
            institute = 'Law Discipline, KU'
            if chief_teacher:
                if not designation and chief_teacher.designation:
                    designation = chief_teacher.designation
                if chief_teacher.institute:
                    institute = chief_teacher.institute
            
            chief_info = {
                'id': chief_assignment_details.assigned_teacher_id,
                'name': chief_teacher.name if chief_teacher else 'N/A',
                'designation': designation,
                'institute': institute
            }
        
        # Get saved custom remuneration forms (only draft, not archived) - filter by current committee's session/year/term
        from blueprints.remuneration_management.models import RemunerationForm
        remuneration_forms_query = RemunerationForm.query.filter_by(
            user_id=current_user.id,
            status='draft'
        )
        # If we have current session/year/term from chief assignment, filter by them
        if current_session and current_year and current_term:
            remuneration_forms_query = remuneration_forms_query.filter_by(
                academic_year=current_session,
                year=current_year,
                term=current_term
            )
        saved_remuneration_forms = remuneration_forms_query.order_by(RemunerationForm.created_at.desc()).limit(20).all()
        
        # Prepare member assignments data for cards (only internal members)
        member_assignments_data = []
        for assignment in internal_member_assignments:
            member_assignments_data.append({
                'id': assignment.id,
                'academic_session': assignment.academic_session or 'N/A',
                'year': assignment.year or 'N/A',
                'term': assignment.term or 'N/A'
            })
        
        return render_template('exam_committee_chief/dashboard.html',
                             teachers=teachers,
                             external_teachers=external_teacher_options,
                             teachers_dict=teachers_dict,  # For JSON serialization in template
                             tabulators=tabulators,
                             scrutinizers=scrutinizers,
                             academic_sessions=academic_sessions,
                             chief_assignment=chief_assignment_details,
                             current_session=current_session,
                             current_year=current_year,
                             current_term=current_term,
                             current_batch=current_batch,
                             available_entries=available_entries,
                             saved_committee_members=saved_committee_members,
                             chief_info=chief_info,
                             saved_remuneration_forms=saved_remuneration_forms,
                             is_chief=bool(chief_assignment),
                             member_assignments=member_assignments_data)

    @app.route('/exam-committee-chief/assign', methods=['POST'])
    @login_required
    def exam_committee_chief_assign():
        """Assign tabulator or scrutinizer by Exam Committee Chief/Internal Member"""
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        chief_assignment = chief_assignments[0] if chief_assignments else None
        member_assignment = member_assignments[0] if member_assignments else None
        base_assignment = chief_assignment or member_assignment

        if not base_assignment:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        data = request.get_json() or {}
        duty_type = data.get('duty_type', '').strip()
        teacher_id = data.get('teacher_id')
        # Auto-fill session, year, term from current committee assignment if not provided
        academic_session = data.get('academic_session', '').strip() or (base_assignment.academic_session if base_assignment else '')
        year = data.get('year', '').strip() or (base_assignment.year if base_assignment else '')
        term = data.get('term', '').strip() or (base_assignment.term if base_assignment else '')
        batch = (data.get('batch', '') or '').strip()
        exam_entry_ids = data.get('exam_entry_ids') or []
        if isinstance(exam_entry_ids, str):
            exam_entry_ids = [exam_entry_ids]
        course_id = data.get('course_id')
        if course_id:
            try:
                course_id = int(course_id)
            except (TypeError, ValueError):
                course_id = None
        course_code = (data.get('course_code', '') or '').strip()
        course_name = (data.get('course_name', '') or '').strip()
        remarks = data.get('remarks', '').strip()
        if course_id and Course and not course_code:
            course_obj = Course.query.get(course_id)
            if course_obj:
                course_code = course_obj.course_code
                course_name = course_obj.course_name
        course_code = course_code.upper() if course_code else None
        
        if duty_type not in {'tabulator', 'scrutinizer'}:
            return jsonify({'success': False, 'message': 'Only Tabulator and Scrutinizer can be assigned by Exam Committee Chief/Member'}), 400
        
        if not teacher_id:
            return jsonify({'success': False, 'message': 'Please select a teacher'}), 400
        
        if duty_type == 'tabulator':
            if not academic_session or not year or not term:
                return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required for Tabulator'}), 400
        elif duty_type == 'scrutinizer':
            missing = []
            if not academic_session:
                missing.append('Academic Session')
            if not year:
                missing.append('Year')
            if not term:
                missing.append('Term')
            if missing:
                return jsonify({'success': False, 'message': f'Scrutinizer assignment requires: {", ".join(missing)}'}), 400
            if not exam_entry_ids:
                return jsonify({'success': False, 'message': 'Please select at least one submitted course for Scrutinizer duty.'}), 400
        
        try:
            if duty_type == 'scrutinizer':
                created_assignments = []
                for entry_id in exam_entry_ids:
                    try:
                        entry_id = int(entry_id)
                    except (TypeError, ValueError):
                        continue
                    exam_entry = ExamPaperEvaluation.query.get(entry_id)
                    if not exam_entry:
                        return jsonify({'success': False, 'message': 'Invalid submitted course selected.'}), 400
                    if not exam_entry.submitted_to_committee:
                        return jsonify({'success': False, 'message': f'{exam_entry.course_code} has not been submitted yet.'}), 400
                    if exam_entry.assigned_scrutinizer_id:
                        return jsonify({'success': False, 'message': f'{exam_entry.course_code} already has a Scrutinizer.'}), 400
                    if exam_entry.owner_teacher_id and exam_entry.owner_teacher_id == teacher_id:
                        return jsonify({'success': False, 'message': f'{exam_entry.course_code} - আপনি নিজে যে সাবজেক্টের এক্সাম পেপার ইভালুয়েট করেছেন, সেই সাবজেক্টের স্ক্রুটিনাইজার হতে পারবেন না।'}), 400

                    assignment = DutyAssignment(
                        course_code=exam_entry.course_code,
                        course_name=exam_entry.course_name,
                        academic_session=exam_entry.academic_session or academic_session or None,
                        year=exam_entry.year or year or None,
                        term=exam_entry.term or term or None,
                        batch=exam_entry.batch or None,
                        duty_type='scrutinizer',
                        assigned_teacher_id=teacher_id,
                        assigned_by_id=current_user.id,
                        remarks=remarks or None,
                        status='active'
                    )
                    db.session.add(assignment)
                    created_assignments.append((assignment, exam_entry))

                db.session.commit()

                invite_count = 0
                for assignment, exam_entry in created_assignments:
                    exam_entry.assigned_scrutinizer_id = assignment.assigned_teacher_id
                    invite = ExamScrutinizerInvite(
                        exam_entry_id=exam_entry.id,
                        inviter_teacher_id=teacher.id,
                        scrutinizer_teacher_id=assignment.assigned_teacher_id,
                        status='invited'
                    )
                    db.session.add(invite)
                    invite_count += 1
                db.session.commit()

                return jsonify({
                    'success': True,
                    'message': f'Scrutinizer assigned to {invite_count} submitted course(s).'
                })

            # Check if assignment already exists
            filter_kwargs = {
                'duty_type': duty_type,
                'assigned_teacher_id': teacher_id,
                'status': 'active'
            }
            if duty_type == 'tabulator':
                filter_kwargs.update({
                    'academic_session': academic_session or None,
                    'year': year or None,
                    'term': term or None
                })
            elif duty_type == 'scrutinizer':
                filter_kwargs.update({
                    'academic_session': academic_session or None,
                    'year': year or None,
                    'term': term or None,
                    'batch': batch or None,
                    'course_code': course_code,
                })
            
            existing = DutyAssignment.query.filter_by(**filter_kwargs).first()
            if existing:
                return jsonify({'success': False, 'message': 'This duty assignment already exists'}), 400
            
            assignment = DutyAssignment(
                academic_session=academic_session or None,
                year=year or None,
                term=term or None,
                batch=batch or None,
                course_id=course_id,
                course_code=course_code,
                course_name=course_name or None,
                duty_type=duty_type,
                assigned_teacher_id=teacher_id,
                assigned_by_id=current_user.id,
                remarks=remarks or None,
                status='active'
            )
            db.session.add(assignment)
            db.session.commit()

            if assignment.duty_type == 'scrutinizer':
                _apply_scrutinizer_assignment(assignment)
            
            return jsonify({
                'success': True,
                'message': f'{duty_type.title()} assigned successfully',
                'assignment_id': assignment.id
            })
        except Exception as exc:
            db.session.rollback()
            current_app.logger.error(f'Failed to assign duty: {exc}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to assign duty'}), 500

    @app.route('/exam-committee-chief/save-committee-members', methods=['POST'])
    @login_required
    def save_committee_members():
        """Save examination committee members"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        chief_assignment = chief_assignments[0] if chief_assignments else None
        if not chief_assignment and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        chief = data.get('chief')  # {id, name, designation}
        internal_members = data.get('internal_members') or []  # [{id, name, designation, institute}]
        external_members = data.get('external_members') or []  # [{name, designation, institute}]
        remarks = data.get('remarks', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        if not chief and (not internal_members or not isinstance(internal_members, list)) and (not external_members or not isinstance(external_members, list)):
            return jsonify({'success': False, 'message': 'Please add at least the Chief or one committee member (internal or external)'}), 400
        
        try:
            # Normalize payload to avoid accidental duplicate entries from repeated UI actions.
            normalized_internal_members = []
            seen_internal_member_ids = set()
            for member in (internal_members if isinstance(internal_members, list) else []):
                member_id = member.get('id')
                if not member_id or member_id in seen_internal_member_ids:
                    continue
                seen_internal_member_ids.add(member_id)
                normalized_internal_members.append(member)

            normalized_external_members = []
            seen_external_names = set()
            for member in (external_members if isinstance(external_members, list) else []):
                name = (member.get('name') or '').strip()
                if not name:
                    continue
                name_key = name.lower()
                if name_key in seen_external_names:
                    continue
                seen_external_names.add(name_key)
                normalized_external_members.append({
                    'name': name,
                    'designation': (member.get('designation') or '').strip(),
                    'institute': (member.get('institute') or '').strip()
                })

            # First, deactivate existing committee members for this session/year/term
            existing_members = DutyAssignment.query.filter_by(
                duty_type='exam_committee_member',
                academic_session=academic_session,
                year=year,
                term=term,
                status='active'
            ).all()
            
            for existing in existing_members:
                existing.status = 'inactive'
            
            # Save chief info separately (we'll store it as a special assignment or in remarks)
            # For now, we'll store chief designation in the chief assignment's remarks field
            if chief and chief.get('id'):
                chief_teacher = Teacher.query.get(chief.get('id'))
                if chief_teacher:
                    # Update chief assignment with designation
                    chief_assignment_update = DutyAssignment.query.filter_by(
                        duty_type='exam_committee_chief',
                        assigned_teacher_id=chief.get('id'),
                        academic_session=academic_session,
                        year=year,
                        term=term
                    ).first()
                    if chief_assignment_update:
                        chief_info_json = json.dumps({
                            'type': 'chief',
                            'designation': chief.get('designation', '')
                        })
                        chief_assignment_update.remarks = chief_info_json
            
            # Save chief designation in chief assignment remarks
            if chief and chief.get('id'):
                chief_assignment_update = DutyAssignment.query.filter_by(
                    duty_type='exam_committee_chief',
                    assigned_teacher_id=chief.get('id'),
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).first()
                if chief_assignment_update:
                    chief_info_json = json.dumps({
                        'type': 'chief',
                        'designation': chief.get('designation', '')
                    })
                    chief_assignment_update.remarks = chief_info_json
            
            # Create assignments for internal members (with teacher_id and designation)
            for member in normalized_internal_members:
                member_id = member.get('id')
                if not member_id:
                    continue
                
                # Check if teacher exists
                member_teacher = Teacher.query.get(member_id)
                if not member_teacher:
                    continue
                
                # Store member info with designation in remarks as JSON
                member_info = {
                    'type': 'internal',
                    'designation': member.get('designation', ''),
                    'institute': member.get('institute', '')
                }
                member_remarks = json.dumps(member_info) if (member.get('designation') or member.get('institute')) else (remarks or '')
                
                # Create or reactivate assignment
                assignment = DutyAssignment.query.filter_by(
                    duty_type='exam_committee_member',
                    assigned_teacher_id=member_id,
                    academic_session=academic_session,
                    year=year,
                    term=term,
                    assigned_by_id=current_user.id
                ).first()
                
                if assignment:
                    assignment.status = 'active'
                    assignment.remarks = member_remarks
                else:
                    assignment = DutyAssignment(
                        duty_type='exam_committee_member',
                        assigned_teacher_id=member_id,
                        academic_session=academic_session,
                        year=year,
                        term=term,
                        assigned_by_id=current_user.id,
                        remarks=member_remarks,
                        status='active'
                    )
                    db.session.add(assignment)
            
            # Create assignments for external members (without teacher_id, store info in remarks as JSON)
            for member in normalized_external_members:
                name = member.get('name', '').strip()
                designation = member.get('designation', '').strip()
                institute = member.get('institute', '').strip()
                
                if not name or not designation or not institute:
                    continue
                
                # Store external member info as JSON in remarks
                external_info = {
                    'type': 'external',
                    'name': name,
                    'designation': designation,
                    'institute': institute
                }
                external_remarks = json.dumps(external_info)
                
                # For external members, we use name as identifier in remarks
                # Create new assignment (external members don't have teacher_id)
                assignment = DutyAssignment(
                    duty_type='exam_committee_member',
                    assigned_teacher_id=None,  # External members have no teacher_id
                    academic_session=academic_session,
                    year=year,
                    term=term,
                    assigned_by_id=current_user.id,
                    remarks=external_remarks,  # Store external info as JSON
                    status='active'
                )
                db.session.add(assignment)
            
            db.session.commit()
            
            internal_count = len(normalized_internal_members)
            external_count = len(normalized_external_members)
            total_count = internal_count + external_count
            
            return jsonify({
                'success': True,
                'message': f'Committee members saved successfully ({internal_count} internal, {external_count} external) for {academic_session} - {year} - {term}'
            })
        
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error saving committee members: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to save committee members'}), 500

    @app.route('/exam-committee-chief/archive-committee-members', methods=['POST'])
    @login_required
    def archive_committee_members():
        """Archive all committee members for a session/year/term"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        if not chief_assignments and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            # Find all committee members for this session/year/term (both active and inactive)
            committee_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.assigned_by_id == current_user.id
            ).all()
            
            # Archive all assignments (set status to 'archived')
            archived_count = 0
            for assignment in committee_assignments:
                if assignment.status != 'archived':
                    assignment.status = 'archived'
                    archived_count += 1
            
            # Also archive the chief assignment for this session/year/term if it exists
            chief_assignment_to_archive = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_chief',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'active'
            ).first()
            
            if chief_assignment_to_archive:
                chief_assignment_to_archive.status = 'archived'
                archived_count += 1
            
            # Archive tabulators for this session/year/term assigned by current chief
            tabulator_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'tabulator',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.assigned_by_id == current_user.id,
                DutyAssignment.status == 'active'
            ).all()
            
            for tabulator in tabulator_assignments:
                tabulator.status = 'archived'
                archived_count += 1
            
            # Archive scrutinizers for this session/year/term assigned by current chief
            scrutinizer_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'scrutinizer',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.assigned_by_id == current_user.id,
                DutyAssignment.status == 'active'
            ).all()
            
            for scrutinizer in scrutinizer_assignments:
                scrutinizer.status = 'archived'
                archived_count += 1
            
            # Archive remuneration forms for this session/year/term
            from blueprints.remuneration_management.models import RemunerationForm
            from datetime import datetime
            import json
            
            remuneration_forms = RemunerationForm.query.filter_by(
                user_id=current_user.id,
                status='draft'
            ).all()
            
            archived_forms_count = 0
            for form in remuneration_forms:
                # Check if form data matches the session/year/term
                try:
                    if form.form_data:
                        form_data = json.loads(form.form_data)
                        form_session = form_data.get('academic_session') or form.academic_year
                        form_year = form_data.get('year') or form.year
                        form_term = form_data.get('term') or form.term
                        
                        if (form_session == academic_session and 
                            form_year == year and 
                            form_term == term):
                            form.status = 'archived'
                            form.archived_at = datetime.utcnow()
                            archived_forms_count += 1
                except:
                    # If form_data parsing fails, check direct fields
                    if (form.academic_year == academic_session and 
                        form.year == year and 
                        form.term == term):
                        form.status = 'archived'
                        form.archived_at = datetime.utcnow()
                        archived_forms_count += 1
            
            archived_count += archived_forms_count
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Committee, tabulators, scrutinizers, and remuneration forms archived successfully for {academic_session} - {year} - {term}. {archived_count} item(s) archived.'
            })
        
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error archiving committee members: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error archiving committee members: {str(e)}'}), 500

    @app.route('/exam-committee-chief/restore-committee-members', methods=['POST'])
    @login_required
    def restore_committee_members():
        """Restore an archived examination committee back to active status"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        if not chief_assignments and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            # First, check if there are any active assignments for this session/year/term
            # If yes, we cannot restore - must archive or deactivate active committee first
            active_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member']),
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'active'
            ).all()
            
            if active_assignments:
                return jsonify({
                    'success': False,
                    'message': f'Cannot restore: An active committee already exists for {academic_session} - {year} - {term}. Please archive or deactivate the active committee first.'
                }), 400
            
            # Ensure requester is related to this committee in archived state too.
            archived_role_assignment = DutyAssignment.query.filter(
                DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member']),
                DutyAssignment.assigned_teacher_id == teacher.id,
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'archived'
            ).first()

            # Find archived chief assignment for this session/year/term
            archived_chief_assignment = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_chief',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'archived'
            ).first()

            if not archived_chief_assignment or not archived_role_assignment:
                return jsonify({
                    'success': False,
                    'message': f'No archived committee found for {academic_session} - {year} - {term} that you can restore.'
                }), 404
            
            # Find all archived member assignments for this session/year/term
            archived_member_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'archived'
            ).all()
            
            # Restore all archived assignments to active status
            restored_count = 0
            archived_chief_assignment.status = 'active'
            restored_count += 1
            
            for assignment in archived_member_assignments:
                assignment.status = 'active'
                restored_count += 1
            
            # Restore tabulators for this session/year/term assigned by current chief
            archived_tabulators = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'tabulator',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.assigned_by_id == current_user.id,
                DutyAssignment.status == 'archived'
            ).all()
            
            for tabulator in archived_tabulators:
                tabulator.status = 'active'
                restored_count += 1
            
            # Restore scrutinizers for this session/year/term assigned by current chief
            archived_scrutinizers = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'scrutinizer',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.assigned_by_id == current_user.id,
                DutyAssignment.status == 'archived'
            ).all()
            
            for scrutinizer in archived_scrutinizers:
                scrutinizer.status = 'active'
                restored_count += 1
            
            # Restore remuneration forms for this session/year/term
            from blueprints.remuneration_management.models import RemunerationForm
            import json
            
            archived_remuneration_forms = RemunerationForm.query.filter_by(
                user_id=current_user.id,
                status='archived'
            ).all()
            
            restored_forms_count = 0
            for form in archived_remuneration_forms:
                # Check if form data matches the session/year/term
                try:
                    if form.form_data:
                        form_data = json.loads(form.form_data)
                        form_session = form_data.get('academic_session') or form.academic_year
                        form_year = form_data.get('year') or form.year
                        form_term = form_data.get('term') or form.term
                        
                        if (form_session == academic_session and 
                            form_year == year and 
                            form_term == term):
                            form.status = 'draft'
                            form.archived_at = None
                            restored_forms_count += 1
                except:
                    # If form_data parsing fails, check direct fields
                    if (form.academic_year == academic_session and 
                        form.year == year and 
                        form.term == term):
                        form.status = 'draft'
                        form.archived_at = None
                        restored_forms_count += 1
            
            restored_count += restored_forms_count
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Committee, tabulators, scrutinizers, and remuneration forms restored successfully for {academic_session} - {year} - {term}. {restored_count} item(s) restored.'
            })
        
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error restoring committee members: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error restoring committee members: {str(e)}'}), 500

    @app.route('/exam-committee-chief/get-subjects-for-evaluator', methods=['GET'])
    @login_required
    def get_subjects_for_evaluator():
        """Get subjects from curriculum for Exam Paper Evaluator assignment"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
        
        chief_assignment = DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            assigned_teacher_id=teacher.id,
            status='active'
        ).first()
        
        member_assignment = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_member',
            DutyAssignment.assigned_teacher_id == teacher.id,
            DutyAssignment.status == 'active'
        ).first()
        
        if not chief_assignment and not member_assignment:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        academic_session = request.args.get('academic_session', '').strip()
        year = request.args.get('year', '').strip()
        term = request.args.get('term', '').strip()
        
        if not all([academic_session, year, term]):
            return jsonify({'success': False, 'message': 'Academic session, year, and term are required'}), 400
        
        # Check if this semester is active
        try:
            from utils.semester_utils import is_semester_active
            batch = request.args.get('batch', '').strip() or None
            
            if not is_semester_active(academic_session, year, term, batch=batch):
                return jsonify({
                    'success': False,
                    'message': f'This semester ({academic_session} - {year} - {term}) is not active. Please activate it in Active Semester Management first.'
                }), 400
        except Exception as e:
            current_app.logger.error(f'Error checking active semester: {e}', exc_info=True)
            # Continue anyway - don't block if there's an error
        
        try:
            from blueprints.course_management.models import CurriculumYearTerm, Course, StudentCourseRegistration
            from blueprints.class_management.models import ExamPaperEvaluatorAssignment
            
            # Find curriculum year/term configs matching the criteria
            configs = CurriculumYearTerm.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            ).all()
            
            current_app.logger.info(f'Found {len(configs)} configs for session={academic_session}, year={year}, term={term}')
            
            if not configs:
                current_app.logger.info(
                    f'No curriculum configuration found for session={academic_session}, year={year}, term={term}; '
                    'continuing with separate-retake fallback rows if available.'
                )
                curriculum_ids = []
                all_courses = []
            else:
                # Get all curriculum IDs
                curriculum_ids = [config.curriculum_id for config in configs]

                # Get all courses for these curricula - only Theory courses
                all_courses = Course.query.filter(
                    Course.curriculum_id.in_(curriculum_ids),
                    Course.offered == True,
                    Course.course_type == 'Theory'
                ).order_by(Course.course_code).all()
            
            # Normalize year/term for comparison
            def normalize_label(label, is_term=False):
                if not label:
                    return ''
                label = str(label).strip()
                lower_label = label.lower()
                if is_term:
                    for suffix in [' term', 'semester', ' sem']:
                        if lower_label.endswith(suffix):
                            lower_label = lower_label[:-len(suffix)].strip()
                    term_alias = {
                        '1': 'first',
                        '1st': 'first',
                        'first': 'first',
                        '2': 'second',
                        '2nd': 'second',
                        'second': 'second',
                        '3': 'third',
                        '3rd': 'third',
                        'third': 'third',
                        'thesis': 'thesis',
                        'thesis term': 'thesis',
                    }
                    return term_alias.get(lower_label, lower_label)

                for suffix in [' year', 'yr', ' years']:
                    if lower_label.endswith(suffix):
                        lower_label = lower_label[:-len(suffix)].strip()
                year_alias = {
                    '1': 'first',
                    '1st': 'first',
                    'first': 'first',
                    '2': 'second',
                    '2nd': 'second',
                    'second': 'second',
                    '3': 'third',
                    '3rd': 'third',
                    'third': 'third',
                    '4': 'fourth',
                    '4th': 'fourth',
                    'fourth': 'fourth',
                    '5': 'fifth',
                    '5th': 'fifth',
                    'fifth': 'fifth',
                    'llm': 'fifth',
                }
                return year_alias.get(lower_label, lower_label)
            
            normalized_year = normalize_label(year, is_term=False)
            normalized_term = normalize_label(term, is_term=True)
            
            # Filter courses by year/term
            matching_courses = []
            for course in all_courses:
                course_year = normalize_label(course.display_year or course.year or '', is_term=False)
                course_term = normalize_label(course.display_term or course.term or '', is_term=True)
                if course_year == normalized_year and course_term == normalized_term:
                    matching_courses.append(course)
            
            current_app.logger.info(f'Found {len(matching_courses)} matching courses out of {len(all_courses)} total courses')

            # Merge duplicate regular subjects (same course_code across multiple curricula/syllabi)
            # so committee UI treats them as one effective subject in this context.
            merged_courses_by_code = {}
            all_matching_course_ids = []
            matching_courses_by_id = {}
            regular_course_ids_by_code = {}
            for course in matching_courses:
                code_key = (course.course_code or '').strip()
                if not code_key:
                    continue
                all_matching_course_ids.append(course.id)
                matching_courses_by_id[course.id] = course
                regular_course_ids_by_code.setdefault(code_key, set()).add(course.id)
                existing = merged_courses_by_code.get(code_key)
                # Prefer later-created (higher id) entry as canonical representative.
                if existing is None or (course.id and existing.id and course.id > existing.id):
                    merged_courses_by_code[code_key] = course

            # Add separate retake-original subjects ONLY when merge is disabled.
            # Rule selected by user:
            # - merge ON  -> shared relevant-subject assignment (no extra row)
            # - merge OFF -> show original retake subject as a separate row here
            separate_retake_course_codes = set()
            separate_retake_entries = {}
            try:
                from types import SimpleNamespace
                from sqlalchemy import func as sa_func, or_
                retake_remarks = {'retake', 're-retake', 're retake', 'reretake'}
                separate_retake_regs = StudentCourseRegistration.query.filter(
                    StudentCourseRegistration.status.in_(['finalized', 'approved']),
                    or_(
                        StudentCourseRegistration.use_relevant_for_committee.is_(False),
                        StudentCourseRegistration.use_relevant_for_committee.is_(None)
                    ),
                    StudentCourseRegistration.course_code.isnot(None),
                    StudentCourseRegistration.course_code != ''
                ).all()

                for reg in separate_retake_regs:
                    remark_value = str(reg.remark or '').strip().lower()
                    if remark_value not in retake_remarks:
                        continue
                    reg_relevant_session = str(reg.relevant_academic_session or '').strip().casefold()
                    if reg_relevant_session != str(academic_session or '').strip().casefold():
                        continue
                    reg_relevant_year = normalize_label(reg.relevant_year or '', is_term=False)
                    reg_relevant_term = normalize_label(reg.relevant_term or '', is_term=True)
                    if normalized_year and reg_relevant_year != normalized_year:
                        continue
                    if normalized_term and reg_relevant_term != normalized_term:
                        continue

                    code_key = (reg.course_code or '').strip()
                    if not code_key:
                        continue

                    candidate_course = None
                    if reg.course_id:
                        candidate_course = Course.query.get(reg.course_id)
                    # Fallback: prefer latest course row by code as ID carrier.
                    if not candidate_course:
                        candidate_course = Course.query.filter(
                            Course.course_code == code_key
                        ).order_by(Course.id.desc()).first()
                    if not candidate_course:
                        normalized_code = ''.join(ch for ch in code_key.lower() if not ch.isspace())
                        candidate_course = Course.query.filter(
                            sa_func.replace(sa_func.lower(Course.course_code), ' ', '') == normalized_code
                        ).order_by(Course.id.desc()).first()
                    if not candidate_course:
                        candidate_course = SimpleNamespace(
                            id=None,
                            course_code=code_key,
                            course_name=(reg.course_name or code_key),
                            content_section_a='',
                            content_section_b=''
                        )

                    separate_retake_course_codes.add(code_key)
                    entry_key = f"course:{getattr(candidate_course, 'id', None)}" if getattr(candidate_course, 'id', None) else f"code:{code_key}"
                    if entry_key not in separate_retake_entries:
                        separate_retake_entries[entry_key] = {
                            'course': candidate_course,
                            'source_session': str(reg.academic_session or '').strip(),
                            'source_year': str(reg.year or '').strip(),
                            'source_term': str(reg.term or '').strip(),
                        }
                    if getattr(candidate_course, 'id', None):
                        all_matching_course_ids.append(candidate_course.id)
                        matching_courses_by_id[candidate_course.id] = candidate_course
            except Exception as separate_retake_error:
                current_app.logger.warning(
                    f'Could not include separate retake subjects in evaluator list: {separate_retake_error}',
                    exc_info=True
                )

            merged_courses = list(merged_courses_by_code.values())
            
            # Get existing evaluator assignments
            existing_assignments = ExamPaperEvaluatorAssignment.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            ).all()
            
            assigned_map = {}  # {(course_id, part): teacher_id}
            question_setter_map = {}  # {(course_id, part): question_setter_id}
            is_same_map = {}  # {(course_id, part): is_same_person}
            assignment_id_map = {}  # {(course_id, part): assignment_id}

            assignment_course_ids = [a.course_id for a in existing_assignments if a.course_id]
            assignment_courses = {}
            if assignment_course_ids:
                assignment_courses = {
                    c.id: c for c in Course.query.filter(Course.id.in_(assignment_course_ids)).all()
                }

            for assignment in existing_assignments:
                if not assignment.course_id:
                    continue
                key = (assignment.course_id, assignment.part)
                # Keep latest assignment when duplicates exist.
                if key not in assignment_id_map or (assignment.id and assignment.id > assignment_id_map[key]):
                    assigned_map[key] = assignment.assigned_teacher_id
                    question_setter_map[key] = assignment.question_setter_id
                    is_same_map[key] = assignment.is_same_person
                    assignment_id_map[key] = assignment.id

            # Build default teacher mapping from curriculum/class course assignments.
            # Rules:
            # - Full course assignment (no section): Part A defaults to that teacher; Part B remains unassigned.
            # - Split course assignments: Section A -> Part A default, Section B -> Part B default.
            default_part_a_teacher = {}  # {course_code: teacher_id}
            default_part_b_teacher = {}  # {course_code: teacher_id}
            if all_matching_course_ids:
                course_ids = all_matching_course_ids
                session_assignments = CourseSessionAssignment.query.filter(
                    CourseSessionAssignment.course_id.in_(course_ids),
                    CourseSessionAssignment.academic_session == academic_session,
                    CourseSessionAssignment.year == year,
                    CourseSessionAssignment.term == term
                ).all()

                for sess_assign in session_assignments:
                    if not sess_assign.teacher_id:
                        continue
                    sess_course = matching_courses_by_id.get(sess_assign.course_id)
                    if not sess_course:
                        continue
                    code_key = (sess_course.course_code or '').strip()
                    if not code_key:
                        continue
                    raw_section = (sess_assign.section or '').strip().upper()
                    # Full/blank section assignment -> only Part A default
                    if not raw_section:
                        if code_key not in default_part_a_teacher:
                            default_part_a_teacher[code_key] = sess_assign.teacher_id
                        continue
                    if raw_section == 'A' and code_key not in default_part_a_teacher:
                        default_part_a_teacher[code_key] = sess_assign.teacher_id
                    elif raw_section == 'B' and code_key not in default_part_b_teacher:
                        default_part_b_teacher[code_key] = sess_assign.teacher_id
            
            subject_entries = []
            for course in merged_courses:
                code_key = (course.course_code or '').strip()
                lookup_ids = list(regular_course_ids_by_code.get(code_key, []))
                if getattr(course, 'id', None) and course.id not in lookup_ids:
                    lookup_ids.append(course.id)
                subject_entries.append({
                    'course': course,
                    'lookup_ids': lookup_ids,
                    'is_separate_retake_subject': False,
                    'source_context': None,
                })

            for separate_entry in separate_retake_entries.values():
                separate_course = separate_entry['course']
                subject_entries.append({
                    'course': separate_course,
                    'lookup_ids': [separate_course.id] if getattr(separate_course, 'id', None) else [],
                    'is_separate_retake_subject': True,
                    'source_context': {
                        'academic_session': separate_entry['source_session'],
                        'year': separate_entry['source_year'],
                        'term': separate_entry['source_term'],
                    }
                })

            def resolve_assignment(lookup_ids, part):
                chosen = {
                    'assigned_teacher_id': None,
                    'question_setter_id': None,
                    'is_same_person': False,
                    'assignment_id': None
                }
                best_assignment_id = None
                for cid in lookup_ids or []:
                    key = (cid, part)
                    assignment_id = assignment_id_map.get(key)
                    if not assignment_id:
                        continue
                    if best_assignment_id is None or assignment_id > best_assignment_id:
                        best_assignment_id = assignment_id
                        chosen = {
                            'assigned_teacher_id': assigned_map.get(key),
                            'question_setter_id': question_setter_map.get(key),
                            'is_same_person': is_same_map.get(key, False),
                            'assignment_id': assignment_id
                        }
                return chosen

            subjects = []
            for entry in subject_entries:
                course = entry['course']
                code_key = (course.course_code or '').strip()
                part_a_assignment = resolve_assignment(entry['lookup_ids'], 'A')
                part_b_assignment = resolve_assignment(entry['lookup_ids'], 'B')
                course_name = course.course_name
                if entry['is_separate_retake_subject']:
                    course_name = f"{course_name} [Retake Separate]"
                    if entry['source_context'] and all(entry['source_context'].values()):
                        course_name = (
                            f"{course_name} "
                            f"(Source: {entry['source_context']['academic_session']} / "
                            f"{entry['source_context']['year']} / {entry['source_context']['term']})"
                        )
                
                subjects.append({
                    'course_id': getattr(course, 'id', None),
                    'course_code': course.course_code,
                    'course_name': course_name,
                    'is_separate_retake_subject': entry['is_separate_retake_subject'] or (code_key in separate_retake_course_codes),
                    'is_assignable': bool(getattr(course, 'id', None)),
                    'part_a': {
                        'assigned': bool(part_a_assignment['assigned_teacher_id']),
                        'assigned_teacher_id': part_a_assignment['assigned_teacher_id'],
                        'assignment_id': part_a_assignment['assignment_id'],
                        'question_setter_id': part_a_assignment['question_setter_id'],
                        'is_same_person': part_a_assignment['is_same_person'],
                        'default_assigned_teacher_id': default_part_a_teacher.get(code_key),
                        'default_question_setter_id': default_part_a_teacher.get(code_key),
                        'has_content': bool(course.content_section_a)
                    },
                    'part_b': {
                        'assigned': bool(part_b_assignment['assigned_teacher_id']),
                        'assigned_teacher_id': part_b_assignment['assigned_teacher_id'],
                        'assignment_id': part_b_assignment['assignment_id'],
                        'question_setter_id': part_b_assignment['question_setter_id'],
                        'is_same_person': part_b_assignment['is_same_person'],
                        'default_assigned_teacher_id': default_part_b_teacher.get(code_key),
                        'default_question_setter_id': default_part_b_teacher.get(code_key),
                        'has_content': bool(course.content_section_b)
                    }
                })
            
            return jsonify({
                'success': True,
                'subjects': subjects
            })
        except Exception as exc:
            current_app.logger.error(f'Failed to get subjects: {exc}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to get subjects'}), 500
    
    @app.route('/exam-committee-chief/assign-evaluator', methods=['POST'])
    @login_required
    def assign_evaluator():
        """Assign Exam Paper Evaluator to a subject part"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
        
        chief_assignment = DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            assigned_teacher_id=teacher.id,
            status='active'
        ).first()
        
        member_assignment = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_member',
            DutyAssignment.assigned_teacher_id == teacher.id,
            DutyAssignment.status == 'active'
        ).first()
        
        if not chief_assignment and not member_assignment:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        data = request.get_json() or {}
        course_id = data.get('course_id')
        part = data.get('part', '').strip().upper()  # 'A' or 'B'
        assigned_teacher_id = data.get('assigned_teacher_id')
        question_setter_id = data.get('question_setter_id')
        is_same_person = data.get('is_same_person', False)
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        batch = data.get('batch', '').strip() or None
        
        # If is_same_person is True, set question_setter_id to assigned_teacher_id
        if is_same_person:
            question_setter_id = assigned_teacher_id
        
        # Convert course_id and assigned_teacher_id to int
        try:
            if course_id:
                course_id = int(course_id)
            if assigned_teacher_id:
                assigned_teacher_id = int(assigned_teacher_id)
            if question_setter_id:
                question_setter_id = int(question_setter_id)
        except (ValueError, TypeError):
            return jsonify({'success': False, 'message': 'Invalid course_id or assigned_teacher_id'}), 400
        
        if not all([course_id, part, assigned_teacher_id, academic_session, year, term]):
            return jsonify({'success': False, 'message': 'All fields are required'}), 400
        
        if part not in ['A', 'B']:
            return jsonify({'success': False, 'message': 'Part must be A or B'}), 400
        
        try:
            from blueprints.course_management.models import Course
            from blueprints.class_management.models import ExamPaperEvaluatorAssignment

            # Backward-compatible fallback for older clients that don't send batch.
            if not batch and academic_session and year and term:
                context_assignment = DutyAssignment.query.filter(
                    DutyAssignment.assigned_teacher_id == teacher.id,
                    DutyAssignment.status == 'active',
                    DutyAssignment.academic_session == academic_session,
                    DutyAssignment.year == year,
                    DutyAssignment.term == term,
                    DutyAssignment.duty_type.in_(['exam_committee_chief', 'exam_committee_member'])
                ).order_by(DutyAssignment.created_at.desc()).first()
                if context_assignment and context_assignment.batch:
                    batch = context_assignment.batch.strip()

            # Final fallback: use first configured curriculum batch for this session/year/term.
            if not batch and academic_session and year and term:
                cfg = CurriculumYearTerm.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).filter(
                    CurriculumYearTerm.batch.isnot(None),
                    CurriculumYearTerm.batch != '',
                    CurriculumYearTerm.batch != 'None'
                ).order_by(CurriculumYearTerm.updated_at.desc()).first()
                if cfg and cfg.batch:
                    batch = cfg.batch.split(',')[0].strip() if ',' in cfg.batch else cfg.batch.strip()

            current_app.logger.info(
                f'assign_evaluator payload: course_id={course_id}, part={part}, '
                f'assigned_teacher_id={assigned_teacher_id}, question_setter_id={question_setter_id}, '
                f'academic_session={academic_session}, year={year}, term={term}, batch={batch}'
            )
            
            course = Course.query.get(course_id)
            if not course:
                return jsonify({'success': False, 'message': 'Course not found'}), 404
            
            # Check if already assigned for this exact course row in the context.
            # This allows separate-retake rows (different course_id, same code) to be assigned independently.
            existing = ExamPaperEvaluatorAssignment.query.filter(
                ExamPaperEvaluatorAssignment.course_id == course_id,
                ExamPaperEvaluatorAssignment.part == part,
                ExamPaperEvaluatorAssignment.academic_session == academic_session,
                ExamPaperEvaluatorAssignment.year == year,
                ExamPaperEvaluatorAssignment.term == term
            ).first()
            
            if existing:
                return jsonify({'success': False, 'message': 'This subject part is already assigned'}), 400
            
            # Create ExamPaperEvaluation entry for the assigned teacher
            exam_evaluation = ExamPaperEvaluation(
                course_name=course.course_name,
                course_code=course.course_code,
                academic_session=academic_session,
                batch=batch,
                year=year,
                term=term,
                section=f'Part {part}',
                program_level=course.category or 'ug',
                owner_teacher_id=assigned_teacher_id,
                submitted_to_committee=False
            )
            db.session.add(exam_evaluation)
            db.session.flush()  # Get the ID
            
            # Create assignment record
            assignment = ExamPaperEvaluatorAssignment(
                course_id=course_id,
                part=part,
                assigned_teacher_id=assigned_teacher_id,
                question_setter_id=question_setter_id,
                is_same_person=is_same_person,
                academic_session=academic_session,
                year=year,
                term=term,
                exam_paper_evaluation_id=exam_evaluation.id,
                assigned_by_id=current_user.id
            )
            db.session.add(assignment)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Evaluator assigned successfully to {course.course_code} Part {part}',
                'assignment_id': assignment.id
            })
        except Exception as exc:
            db.session.rollback()
            error_message = str(exc)
            current_app.logger.error(f'Failed to assign evaluator: {exc}', exc_info=True)
            return jsonify({
                'success': False, 
                'message': f'Failed to assign evaluator: {error_message}'
            }), 500
    
    @app.route('/exam-committee-chief/get-evaluator-assignments', methods=['GET'])
    @login_required
    def get_evaluator_assignments():
        """Get evaluator assignments for Remuneration Form auto-population"""
        academic_session = request.args.get('academic_session', '').strip()
        year = request.args.get('year', '').strip()
        term = request.args.get('term', '').strip()
        
        if not all([academic_session, year, term]):
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            from blueprints.class_management.models import ExamPaperEvaluatorAssignment, ExamPaperEvaluation
            from blueprints.course_management.models import Course
            
            # Get all assignments for this session/year/term
            assignments = ExamPaperEvaluatorAssignment.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            ).all()
            
            # Build a map: {(course_code, part): {question_setter_name, evaluator_name, script_count}}
            assignment_map = {}
            for assignment in assignments:
                course = Course.query.get(assignment.course_id)
                if not course:
                    continue
                
                # Get question setter name
                question_setter_name = None
                if assignment.question_setter_id:
                    question_setter = Teacher.query.get(assignment.question_setter_id)
                    if question_setter:
                        question_setter_name = question_setter.name
                
                # Get evaluator name
                evaluator_name = None
                if assignment.assigned_teacher_id:
                    evaluator = Teacher.query.get(assignment.assigned_teacher_id)
                    if evaluator:
                        evaluator_name = evaluator.name
                
                # Get script count for this evaluator/course/part
                # Count ALL submitted scripts for this teacher/course/part combination
                script_count = 0
                if assignment.assigned_teacher_id:
                    course_code = course.course_code.strip()
                    part = assignment.part.strip().upper()
                    section_text = f"Part {part}"
                    
                    # Find ALL submitted exam evaluations for this teacher/course/part
                    # Only count entries that are submitted to committee
                    submitted_entries = ExamPaperEvaluation.query.filter(
                        ExamPaperEvaluation.owner_teacher_id == assignment.assigned_teacher_id,
                        ExamPaperEvaluation.course_code == course_code,
                        ExamPaperEvaluation.section == section_text,
                        ExamPaperEvaluation.academic_session == academic_session,
                        ExamPaperEvaluation.year == year,
                        ExamPaperEvaluation.term == term,
                        ExamPaperEvaluation.submitted_to_committee == True
                    ).all()
                    
                    # Sum up script counts from ALL submitted entries
                    for entry in submitted_entries:
                        if entry.marks_data:
                            try:
                                marks_data = json.loads(entry.marks_data) if isinstance(entry.marks_data, str) else entry.marks_data
                                if isinstance(marks_data, dict):
                                    # Check if marks_data has a "rows" array (new structure)
                                    if 'rows' in marks_data and isinstance(marks_data['rows'], list):
                                        # Count number of student rows
                                        script_count += len(marks_data['rows'])
                                    else:
                                        # Fallback: if marks_data is a direct dict of students (old structure)
                                        # Count number of student keys (excluding metadata keys)
                                        student_keys = [k for k in marks_data.keys() if k not in ['questions', 'rows']]
                                        script_count += len(student_keys)
                            except (json.JSONDecodeError, TypeError):
                                pass
                
                # Use course_code as key (normalize it)
                course_code = course.course_code.strip()
                part = assignment.part.strip().upper()
                course_full_text = f"{course_code} - {course.course_name}"
                
                # Store by course_code and full course text (JSON stringified for JavaScript compatibility)
                key_by_code = json.dumps([course_code, part])
                key_by_full = json.dumps([course_full_text, part])
                
                assignment_data = {
                    'question_setter_name': question_setter_name,
                    'evaluator_name': evaluator_name,
                    'script_count': script_count,
                    'course_code': course_code,
                    'course_name': course.course_name
                }
                
                assignment_map[key_by_code] = assignment_data
                assignment_map[key_by_full] = assignment_data
            
            return jsonify({
                'success': True,
                'assignments': assignment_map
            })
        except Exception as exc:
            current_app.logger.error(f'Failed to get evaluator assignments: {exc}', exc_info=True)
            return jsonify({
                'success': False,
                'message': f'Failed to get evaluator assignments: {str(exc)}'
            }), 500

    @app.route('/exam-committee-chief/evaluator-assignments-pdf', methods=['GET'])
    @login_required
    def exam_committee_chief_evaluator_assignments_pdf():
        """Download Question Setter + Evaluator assignments as PDF."""
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            flash('Teacher profile not found.', 'warning')
            return redirect(url_for('exam_committee_chief_dashboard'))

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        if not chief_assignments and not member_assignments:
            flash('You are not assigned as Exam Committee Chief or Member.', 'danger')
            return redirect(url_for('index'))

        academic_session = request.args.get('academic_session', '').strip()
        year = request.args.get('year', '').strip()
        term = request.args.get('term', '').strip()

        if not all([academic_session, year, term]):
            flash('Academic session, year, and term are required for PDF export.', 'warning')
            return redirect(url_for('exam_committee_chief_dashboard'))

        allowed_combinations = {
            (str(a.academic_session or '').strip(), str(a.year or '').strip(), str(a.term or '').strip())
            for a in (chief_assignments + member_assignments)
            if a.academic_session and a.year and a.term
        }
        requested_combo = (academic_session, year, term)
        if requested_combo not in allowed_combinations:
            flash('You are not authorized to export this context.', 'danger')
            return redirect(url_for('exam_committee_chief_dashboard'))

        try:
            from io import BytesIO
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from blueprints.class_management.models import ExamPaperEvaluatorAssignment
            from blueprints.course_management.models import Course, StudentCourseRegistration
            from sqlalchemy import or_

            assignments = ExamPaperEvaluatorAssignment.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            ).all()

            # Normalize by exact course row for assignments, then project into PDF rows.
            course_ids = [a.course_id for a in assignments if a.course_id]
            courses_by_id = {}
            if course_ids:
                courses_by_id = {
                    c.id: c for c in Course.query.filter(Course.id.in_(course_ids)).all()
                }

            teacher_ids = set()
            for assignment in assignments:
                if assignment.assigned_teacher_id:
                    teacher_ids.add(assignment.assigned_teacher_id)
                if assignment.question_setter_id:
                    teacher_ids.add(assignment.question_setter_id)
            teachers_by_id = {}
            if teacher_ids:
                teachers_by_id = {
                    t.id: t for t in Teacher.query.filter(Teacher.id.in_(list(teacher_ids))).all()
                }

            deduped_by_course_part = {}
            for assignment in assignments:
                if not assignment.course_id:
                    continue
                part_key = (assignment.part or '').strip().upper()
                if part_key not in {'A', 'B'}:
                    continue
                key = (assignment.course_id, part_key)
                if key not in deduped_by_course_part or (assignment.id and assignment.id > deduped_by_course_part[key].id):
                    deduped_by_course_part[key] = assignment

            # Build one row per subject:
            # Subject Code | Subject Name | Part A (QS + Evaluator) | Part B (QS + Evaluator)
            course_rows = {}

            # Regular rows: keep one merged row by subject code.
            for (course_id, part), assignment in deduped_by_course_part.items():
                course = courses_by_id.get(course_id)
                if not course:
                    continue
                course_code = (course.course_code or '').strip()
                if not course_code:
                    continue
                if course_code not in course_rows:
                    course_rows[course_code] = {
                        'course_code': course_code,
                        'course_name': (course.course_name or '-') or '-',
                        'A': None,
                        'B': None,
                        'is_retake': False
                    }
                existing = course_rows[course_code].get(part)
                if existing is None or ((assignment.id or 0) > (existing.id or 0)):
                    course_rows[course_code][part] = assignment

            # Separate-retake rows: include in PDF even when unassigned.
            def normalize_label(label, is_term=False):
                if not label:
                    return ''
                value = str(label).strip().lower()
                if is_term:
                    for suffix in [' term', 'semester', ' sem']:
                        if value.endswith(suffix):
                            value = value[:-len(suffix)].strip()
                    term_alias = {
                        '1': 'first', '1st': 'first', 'first': 'first',
                        '2': 'second', '2nd': 'second', 'second': 'second',
                        '3': 'third', '3rd': 'third', 'third': 'third',
                        'thesis': 'thesis', 'thesis term': 'thesis',
                    }
                    return term_alias.get(value, value)
                for suffix in [' year', 'yr', ' years']:
                    if value.endswith(suffix):
                        value = value[:-len(suffix)].strip()
                year_alias = {
                    '1': 'first', '1st': 'first', 'first': 'first',
                    '2': 'second', '2nd': 'second', 'second': 'second',
                    '3': 'third', '3rd': 'third', 'third': 'third',
                    '4': 'fourth', '4th': 'fourth', 'fourth': 'fourth',
                    '5': 'fifth', '5th': 'fifth', 'fifth': 'fifth',
                    'llm': 'fifth',
                }
                return year_alias.get(value, value)

            normalized_year = normalize_label(year, is_term=False)
            normalized_term = normalize_label(term, is_term=True)
            separate_retake_rows = {}
            retake_remarks = {'retake', 're-retake', 're retake', 'reretake'}
            separate_retake_regs = StudentCourseRegistration.query.filter(
                StudentCourseRegistration.status.in_(['finalized', 'approved']),
                or_(
                    StudentCourseRegistration.use_relevant_for_committee.is_(False),
                    StudentCourseRegistration.use_relevant_for_committee.is_(None)
                ),
                StudentCourseRegistration.course_code.isnot(None),
                StudentCourseRegistration.course_code != ''
            ).all()

            for reg in separate_retake_regs:
                remark_value = str(reg.remark or '').strip().lower()
                if remark_value not in retake_remarks:
                    continue
                reg_session = str(reg.relevant_academic_session or '').strip().casefold()
                if reg_session != str(academic_session or '').strip().casefold():
                    continue
                reg_year = normalize_label(reg.relevant_year or '', is_term=False)
                reg_term = normalize_label(reg.relevant_term or '', is_term=True)
                if normalized_year and reg_year != normalized_year:
                    continue
                if normalized_term and reg_term != normalized_term:
                    continue

                course_code = (reg.course_code or '').strip()
                if not course_code:
                    continue
                retake_key = f"retake:{reg.course_id}" if reg.course_id else f"retake:{course_code.lower()}"
                if retake_key in separate_retake_rows:
                    continue

                matched_course = Course.query.get(reg.course_id) if reg.course_id else None
                if not matched_course:
                    matched_course = Course.query.filter(
                        Course.course_code == course_code
                    ).order_by(Course.id.desc()).first()
                if not matched_course:
                    # Show row in PDF even if source mapping/course row is missing.
                    separate_retake_rows[retake_key] = {
                        'course_code': course_code,
                        'course_name': ((reg.course_name or course_code).strip() + ' (Retake)').strip(),
                        'A': None,
                        'B': None,
                        'is_retake': True
                    }
                    continue

                part_a_assignment = deduped_by_course_part.get((matched_course.id, 'A'))
                part_b_assignment = deduped_by_course_part.get((matched_course.id, 'B'))
                separate_retake_rows[retake_key] = {
                    'course_code': (matched_course.course_code or course_code).strip(),
                    'course_name': ((matched_course.course_name or reg.course_name or course_code).strip() + ' (Retake)').strip(),
                    'A': part_a_assignment,
                    'B': part_b_assignment,
                    'is_retake': True
                }

            def _format_part_cell(part_assignment):
                if not part_assignment:
                    return 'Question Setter: -\nEvaluator: -'
                evaluator = teachers_by_id.get(part_assignment.assigned_teacher_id)
                question_setter = teachers_by_id.get(part_assignment.question_setter_id) if part_assignment.question_setter_id else None
                evaluator_name = evaluator.name if evaluator else '-'
                if part_assignment.is_same_person:
                    question_setter_name = f'{evaluator_name} (Same)'
                else:
                    question_setter_name = question_setter.name if question_setter else '-'
                return f'Question Setter: {question_setter_name}\nEvaluator: {evaluator_name}'

            rows = []
            all_rows_for_pdf = list(course_rows.values()) + list(separate_retake_rows.values())
            for row_info in sorted(all_rows_for_pdf, key=lambda item: (item.get('course_code') or '', item.get('is_retake', False))):
                rows.append([
                    row_info['course_code'],
                    row_info['course_name'],
                    _format_part_cell(row_info.get('A')),
                    _format_part_cell(row_info.get('B')),
                ])

            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=landscape(A4),
                topMargin=14 * mm,
                bottomMargin=14 * mm,
                leftMargin=14 * mm,
                rightMargin=14 * mm,
            )

            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'EvaluatorAssignmentPdfTitle',
                parent=styles['Heading2'],
                alignment=1,
                textColor=colors.HexColor('#0a3d62'),
                spaceAfter=4,
            )
            meta_style = ParagraphStyle(
                'EvaluatorAssignmentPdfMeta',
                parent=styles['Normal'],
                alignment=1,
                fontSize=9,
                textColor=colors.HexColor('#333333'),
                spaceAfter=10,
            )
            cell_header_style = ParagraphStyle(
                'EvaluatorAssignmentCellHeader',
                parent=styles['Normal'],
                alignment=1,
                fontName='Helvetica-Bold',
                fontSize=9,
                leading=11,
                wordWrap='CJK',
            )
            cell_style = ParagraphStyle(
                'EvaluatorAssignmentCell',
                parent=styles['Normal'],
                fontSize=8.5,
                leading=11,
                wordWrap='CJK',
            )
            cell_center_style = ParagraphStyle(
                'EvaluatorAssignmentCellCenter',
                parent=cell_style,
                alignment=1,
            )

            elements = [
                Paragraph('Assign Exam Paper Question Setter & Evaluator', title_style),
                Paragraph(
                    f'Academic Session: {academic_session} | Year: {year} | Term: {term}',
                    meta_style
                ),
            ]

            from html import escape
            table_data = [[
                Paragraph('Subject Code', cell_header_style),
                Paragraph('Subject Name', cell_header_style),
                Paragraph('Part A (Question Setter + Evaluator)', cell_header_style),
                Paragraph('Part B (Question Setter + Evaluator)', cell_header_style),
            ]]
            if rows:
                for row in rows:
                    table_data.append([
                        Paragraph(escape(str(row[0] or '-')), cell_center_style),
                        Paragraph(escape(str(row[1] or '-')), cell_style),
                        Paragraph(escape(str(row[2] or '-')).replace('\n', '<br/>'), cell_style),
                        Paragraph(escape(str(row[3] or '-')).replace('\n', '<br/>'), cell_style),
                    ])
            else:
                table_data.append([
                    Paragraph('-', cell_center_style),
                    Paragraph('No assignments found for this context.', cell_style),
                    Paragraph('-', cell_center_style),
                    Paragraph('-', cell_center_style),
                ])

            table = Table(
                table_data,
                colWidths=[28 * mm, 72 * mm, 82 * mm, 82 * mm],
                repeatRows=1
            )
            table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.35, colors.HexColor('#333333')),
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eef3fb')),
                ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 5),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#fafcff')]),
            ]))
            elements.append(Spacer(1, 4))
            elements.append(table)

            doc.build(elements)
            buffer.seek(0)
            pdf_data = buffer.getvalue()

            safe_session = academic_session.replace('/', '-').replace(' ', '_')
            safe_year = year.replace(' ', '_')
            safe_term = term.replace(' ', '_')
            filename = f'Exam_Paper_Assignments_{safe_session}_{safe_year}_{safe_term}.pdf'

            return Response(
                pdf_data,
                mimetype='application/pdf',
                headers={
                    'Content-Disposition': f'attachment; filename="{filename}"',
                    'Content-Length': str(len(pdf_data)),
                    'Cache-Control': 'no-cache, no-store, must-revalidate',
                    'Pragma': 'no-cache',
                    'Expires': '0',
                }
            )
        except Exception as exc:
            current_app.logger.error(f'Failed to export evaluator assignments PDF: {exc}', exc_info=True)
            flash('Failed to export evaluator assignment PDF.', 'danger')
            return redirect(url_for(
                'exam_committee_chief_dashboard',
                session=academic_session,
                year=year,
                term=term
            ))
    
    @app.route('/exam-committee-chief/get-scrutinizers-with-scripts', methods=['GET'])
    @login_required
    def get_scrutinizers_with_scripts():
        """Get scrutinizers with their script counts for the current session/year/term"""
        academic_session = request.args.get('academic_session', '').strip()
        year = request.args.get('year', '').strip()
        term = request.args.get('term', '').strip()
        
        if not all([academic_session, year, term]):
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            from blueprints.class_management.models import ExamPaperEvaluation
            
            # Get all submitted exam entries for this session/year/term that have assigned scrutinizers
            submitted_entries = ExamPaperEvaluation.query.filter(
                ExamPaperEvaluation.academic_session == academic_session,
                ExamPaperEvaluation.year == year,
                ExamPaperEvaluation.term == term,
                ExamPaperEvaluation.submitted_to_committee == True,
                ExamPaperEvaluation.assigned_scrutinizer_id.isnot(None)
            ).all()
            
            # Build a map: {teacher_id: {'name': teacher_name, 'script_count': total_count}}
            scrutinizer_map = {}
            
            for entry in submitted_entries:
                if not entry.assigned_scrutinizer_id or not entry.marks_data:
                    continue
                
                # Get scrutinizer
                scrutinizer = Teacher.query.get(entry.assigned_scrutinizer_id)
                if not scrutinizer:
                    continue
                
                teacher_id = scrutinizer.id
                
                # Count scripts from marks_data
                try:
                    marks_data = json.loads(entry.marks_data) if isinstance(entry.marks_data, str) else entry.marks_data
                    script_count = 0
                    
                    if isinstance(marks_data, dict):
                        # Check if marks_data has a "rows" array (new structure)
                        if 'rows' in marks_data and isinstance(marks_data['rows'], list):
                            # Count number of student rows
                            script_count = len(marks_data['rows'])
                        else:
                            # Fallback: if marks_data is a direct dict of students (old structure)
                            # Count number of student keys (excluding metadata keys)
                            student_keys = [k for k in marks_data.keys() if k not in ['questions', 'rows']]
                            script_count = len(student_keys)
                    
                    # Add to total for this scrutinizer
                    if teacher_id in scrutinizer_map:
                        scrutinizer_map[teacher_id]['script_count'] += script_count
                    else:
                        scrutinizer_map[teacher_id] = {
                            'name': scrutinizer.name,
                            'designation': scrutinizer.designation or '',
                            'institute': scrutinizer.institute or 'Law Discipline, KU',
                            'script_count': script_count
                        }
                except (json.JSONDecodeError, TypeError):
                    pass
            
            # Convert to list
            scrutinizers = list(scrutinizer_map.values())
            
            return jsonify({
                'success': True,
                'scrutinizers': scrutinizers
            })
        except Exception as exc:
            current_app.logger.error(f'Failed to get scrutinizers with scripts: {exc}', exc_info=True)
            return jsonify({
                'success': False,
                'message': f'Failed to get scrutinizers with scripts: {str(exc)}'
            }), 500
    
    @app.route('/exam-committee-chief/unassign-evaluator', methods=['POST'])
    @login_required
    def unassign_evaluator():
        """Unassign Exam Paper Evaluator from a subject part"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
        
        chief_assignment = DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            assigned_teacher_id=teacher.id,
            status='active'
        ).first()
        
        member_assignment = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_member',
            DutyAssignment.assigned_teacher_id == teacher.id,
            DutyAssignment.status == 'active'
        ).first()
        
        if not chief_assignment and not member_assignment:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        data = request.get_json() or {}
        assignment_id = data.get('assignment_id')
        course_id = data.get('course_id')
        part = data.get('part', '').strip().upper()
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        
        try:
            from blueprints.class_management.models import ExamPaperEvaluatorAssignment
            
            # Find assignment by ID or by course_id/part/session/year/term
            if assignment_id:
                assignment = ExamPaperEvaluatorAssignment.query.get(assignment_id)
            elif course_id and part and academic_session and year and term:
                assignment = ExamPaperEvaluatorAssignment.query.filter_by(
                    course_id=course_id,
                    part=part,
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).first()
            else:
                return jsonify({'success': False, 'message': 'Either assignment_id or course_id/part/session/year/term is required'}), 400
            
            if not assignment:
                return jsonify({'success': False, 'message': 'Assignment not found'}), 404
            
            # Delete the ExamPaperEvaluation entry and related records
            if assignment.exam_paper_evaluation_id:
                exam_eval = ExamPaperEvaluation.query.get(assignment.exam_paper_evaluation_id)
                if exam_eval:
                    # Delete related ExamScrutinizerInvite records first
                    from blueprints.class_management.models import ExamScrutinizerInvite
                    ExamScrutinizerInvite.query.filter_by(exam_entry_id=exam_eval.id).delete()
                    # Now delete the exam evaluation entry
                    db.session.delete(exam_eval)
            
            # Delete the assignment
            db.session.delete(assignment)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': 'Evaluator unassigned successfully'
            })
        except Exception as exc:
            db.session.rollback()
            error_message = str(exc)
            current_app.logger.error(f'Failed to unassign evaluator: {exc}', exc_info=True)
            return jsonify({
                'success': False, 
                'message': f'Failed to unassign evaluator: {error_message}'
            }), 500

    @app.route('/exam-committee-chief/reset-committee-members', methods=['POST'])
    @login_required
    def reset_committee_members():
        """Reset/Delete all committee members for a session/year/term"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        if not chief_assignments and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        data = request.get_json() or {}
        academic_session = data.get('academic_session', '').strip()
        year = data.get('year', '').strip()
        term = data.get('term', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Academic Session, Year, and Term are required'}), 400
        
        try:
            # Find all committee members for this session/year/term
            committee_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == academic_session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'active'
            ).all()
            
            # Deactivate all assignments
            for assignment in committee_assignments:
                assignment.status = 'inactive'
            
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Committee members reset successfully for {academic_session} - {year} - {term}'
            })
        
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error resetting committee members: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to reset committee members'}), 500

    @app.route('/exam-committee-chief/custom-remuneration')
    @app.route('/exam-committee-chief/custom-remuneration/<int:form_id>')
    @login_required
    def exam_committee_chief_custom_remuneration(form_id=None):
        """Custom Remuneration Statement form for Exam Committee Chief"""
        # Check if current user is assigned as Exam Committee Chief
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            flash('Teacher profile not found.', 'warning')
            return redirect(url_for('index'))
        
        # Get all active exam committee chief assignments for this teacher
        chief_assignments = DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            assigned_teacher_id=teacher.id,
            status='active'
        ).all()
        
        # Get all active exam committee member assignments for this teacher
        member_assignments = DutyAssignment.query.filter_by(
            duty_type='exam_committee_member',
            assigned_teacher_id=teacher.id,
            status='active'
        ).all()
        
        # If form_id is provided, check if user can access it (either as chief or member of same committee)
        if form_id:
            form_entry = RemunerationForm.query.filter_by(id=form_id).first()
            if form_entry:
                # Check if user owns the form or is a member of the same committee
                can_access = False
                if form_entry.user_id == current_user.id:
                    can_access = True
                elif member_assignments:
                    # Check if current user is a member of the same committee
                    for member_assignment in member_assignments:
                        if (member_assignment.academic_session == form_entry.academic_year and
                            member_assignment.year == form_entry.year and
                            member_assignment.term == form_entry.term):
                            can_access = True
                            break
                
                if not can_access:
                    flash('You do not have permission to access this form.', 'danger')
                    return redirect(url_for('index'))
        
        if not chief_assignments and not member_assignments:
            flash('You are not assigned as Exam Committee Chief or Member.', 'danger')
            return redirect(url_for('index'))
        
        # Get unique session/year/term combinations from assignments (both chief and member)
        allowed_combinations = set()
        for assignment in chief_assignments:
            if assignment.academic_session and assignment.year and assignment.term:
                allowed_combinations.add((
                    str(assignment.academic_session),
                    str(assignment.year),
                    str(assignment.term)
                ))
        for assignment in member_assignments:
            if assignment.academic_session and assignment.year and assignment.term:
                allowed_combinations.add((
                    str(assignment.academic_session),
                    str(assignment.year),
                    str(assignment.term)
                ))
        
        # Determine which committee to load based on:
        # 1. URL parameters (session, year, term)
        # 2. Form data if form_id is provided
        # 3. Otherwise use first assignment (for backward compatibility)
        url_session = request.args.get('session', '').strip()
        url_year = request.args.get('year', '').strip()
        url_term = request.args.get('term', '').strip()
        
        # Load saved form data if form_id is provided
        saved_data = None
        saved_examination_committee = None
        form_entry = None
        if form_id:
            # Allow loading Chief's form if user is a member of the same committee
            form_entry = RemunerationForm.query.filter_by(id=form_id).first()
            if form_entry:
                # Use form's session/year/term to find matching assignment
                url_session = form_entry.academic_year or url_session
                url_year = form_entry.year or url_year
                url_term = form_entry.term or url_term
                # Check if user owns the form or is a member of the same committee
                can_access = False
                if form_entry.user_id == current_user.id:
                    can_access = True
                elif member_assignments:
                    # Check if current user is a member of the same committee
                    for member_assignment in member_assignments:
                        if (member_assignment.academic_session == form_entry.academic_year and
                            member_assignment.year == form_entry.year and
                            member_assignment.term == form_entry.term):
                            can_access = True
                            break
                
                if can_access:
                    try:
                        saved_data = json.loads(form_entry.form_data) if form_entry.form_data else {}
                        saved_data['form_id'] = form_entry.id
                        saved_examination_committee = saved_data.get('examination_committee', [])
                    except:
                        saved_data = {}
        
        # Get all teachers and convert to JSON-serializable format (excluding Head of the Discipline)
        from role_utils import get_teachers_excluding_head
        teachers_query = get_teachers_excluding_head()
        teachers = []
        teachers_data = []  # For JSON serialization in JavaScript
        for teacher in teachers_query:
            teachers.append(teacher)  # Keep for template rendering (dropdowns)
            teachers_data.append({
                'name': teacher.name or '',
                'designation': teacher.designation or '',
                'institute': teacher.institute or 'Law Discipline, KU'
            })
        
        # Get curriculum data (Year, Term, Batch, Session)
        from blueprints.course_management.models import CurriculumYearTerm
        
        # Filter curriculum data to only show session/year/term combinations where user is chief
        if allowed_combinations:
            # Build filter conditions for allowed combinations
            from sqlalchemy import or_
            filter_conditions = []
            for session, year, term in allowed_combinations:
                filter_conditions.append(
                    db.and_(
                        CurriculumYearTerm.academic_session == session,
                        CurriculumYearTerm.year == year,
                        CurriculumYearTerm.term == term
                    )
                )
            
            if filter_conditions:
                curriculum_query = CurriculumYearTerm.query.filter(or_(*filter_conditions))
            else:
                curriculum_query = CurriculumYearTerm.query.filter(False)  # No results
        else:
            # If no specific assignments, use all (fallback)
            curriculum_query = CurriculumYearTerm.query
        
        # Get unique academic sessions from filtered curriculum
        curriculum_sessions = curriculum_query.with_entities(
            CurriculumYearTerm.academic_session
        ).distinct().filter(
            CurriculumYearTerm.academic_session.isnot(None)
        ).order_by(CurriculumYearTerm.academic_session.desc()).all()
        academic_sessions = [s[0] for s in curriculum_sessions if s[0]]
        
        # Get all unique years, terms, and batches from filtered curriculum
        all_years = curriculum_query.with_entities(
            CurriculumYearTerm.year
        ).distinct().filter(
            CurriculumYearTerm.year.isnot(None)
        ).order_by(CurriculumYearTerm.year.asc()).all()
        unique_years = sorted(set([y[0] for y in all_years if y[0]]))
        
        all_terms = curriculum_query.with_entities(
            CurriculumYearTerm.term
        ).distinct().filter(
            CurriculumYearTerm.term.isnot(None)
        ).order_by(CurriculumYearTerm.term.asc()).all()
        unique_terms = sorted(set([t[0] for t in all_terms if t[0]]))
        
        all_batches = curriculum_query.with_entities(
            CurriculumYearTerm.batch
        ).distinct().filter(
            CurriculumYearTerm.batch.isnot(None)
        ).order_by(CurriculumYearTerm.batch.desc()).all()
        unique_batches = [b[0] for b in all_batches if b[0]]
        
        # Get all curriculum configs for dynamic filtering (filtered)
        curriculum_configs = curriculum_query.order_by(
            CurriculumYearTerm.academic_session.desc(),
            CurriculumYearTerm.year.asc(),
            CurriculumYearTerm.term.asc()
        ).all()
        
        # Convert to JSON-serializable format (handle None values)
        configs_data = []
        for config in curriculum_configs:
            configs_data.append({
                'academic_session': str(config.academic_session) if config.academic_session else '',
                'year': str(config.year) if config.year else '',
                'term': str(config.term) if config.term else '',
                'batch': str(config.batch) if config.batch else ''
            })
        
        # Find the matching chief assignment based on URL params or form data
        chief_assignment = None
        if url_session and url_year and url_term:
            # Find assignment matching URL parameters
            current_app.logger.info(f'Looking for chief assignment: session={url_session}, year={url_year}, term={url_term}')
            for assignment in chief_assignments:
                assignment_session = str(assignment.academic_session) if assignment.academic_session else ''
                assignment_year = str(assignment.year) if assignment.year else ''
                assignment_term = str(assignment.term) if assignment.term else ''
                current_app.logger.info(f'Checking assignment: session={assignment_session}, year={assignment_year}, term={assignment_term}')
                if (assignment_session == url_session and
                    assignment_year == url_year and
                    assignment_term == url_term):
                    chief_assignment = assignment
                    current_app.logger.info(f'✅ Found matching chief assignment: ID={assignment.id}')
                    break
        
        # If no match found, use first assignment (for backward compatibility)
        if not chief_assignment and chief_assignments:
            chief_assignment = chief_assignments[0]
            current_app.logger.warning(f'⚠️ No matching assignment found, using first assignment: session={chief_assignment.academic_session}, year={chief_assignment.year}, term={chief_assignment.term}')
        
        # Set current values from matched assignment or URL params
        current_session = url_session or (chief_assignment.academic_session if chief_assignment else None)
        current_year = url_year or (chief_assignment.year if chief_assignment else None)
        current_term = url_term or (chief_assignment.term if chief_assignment else None)
        
        current_app.logger.info(f'Final current values: session={current_session}, year={current_year}, term={current_term}')
        
        # Get current batch from curriculum if available
        current_batch = None
        if current_session and current_year and current_term:
            config = CurriculumYearTerm.query.filter_by(
                academic_session=current_session,
                year=current_year,
                term=current_term
            ).first()
            if config:
                current_batch = config.batch
        
        # Get examination committee members (Chief + Internal + External)
        examination_committee = []
        
        # Add Chief (Chairman)
        if chief_assignment and chief_assignment.assigned_teacher:
            chief_teacher = chief_assignment.assigned_teacher
            chief_designation = ''
            chief_institute = ''
            
            # Try to get from chief assignment remarks (from committee setup)
            try:
                chief_data = json.loads(chief_assignment.remarks) if chief_assignment.remarks else {}
                if chief_data.get('type') == 'chief':
                    chief_designation = chief_data.get('designation', '')
                    chief_institute = chief_data.get('institute', '')
            except:
                pass
            
            # Fall back to Teacher table if not in remarks
            if not chief_designation and chief_teacher.designation:
                chief_designation = chief_teacher.designation
            if not chief_institute and chief_teacher.institute:
                chief_institute = chief_teacher.institute
            
            # Final defaults
            if not chief_designation:
                chief_designation = 'Head, Law Discipline, KU'
            if not chief_institute:
                chief_institute = 'Law Discipline, KU'
            
            examination_committee.append({
                'name': chief_teacher.name,
                'designation': f'{chief_designation}, {chief_institute}',
                'position': 'Chairman'
            })
        
        # Get committee members for current session/year/term
        if current_session and current_year and current_term:
            committee_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == current_session,
                DutyAssignment.year == current_year,
                DutyAssignment.term == current_term,
                DutyAssignment.status == 'active',
                DutyAssignment.assigned_by_id == current_user.id
            ).all()
            
            for assignment in committee_assignments:
                if assignment.assigned_teacher_id:
                    # Internal member
                    teacher = assignment.assigned_teacher
                    if teacher:
                        member_designation = ''
                        member_institute = ''
                        
                        # Try to get from assignment remarks (from committee setup)
                        try:
                            member_data = json.loads(assignment.remarks) if assignment.remarks else {}
                            if member_data.get('type') == 'internal':
                                member_designation = member_data.get('designation', '')
                                member_institute = member_data.get('institute', '')
                        except:
                            pass
                        
                        # Fall back to Teacher table if not in remarks
                        if not member_designation and teacher.designation:
                            member_designation = teacher.designation
                        if not member_institute and teacher.institute:
                            member_institute = teacher.institute
                        
                        # Final defaults
                        if not member_designation:
                            member_designation = 'Assistant Professor'
                        if not member_institute:
                            member_institute = 'Law Discipline, KU'
                        
                        examination_committee.append({
                            'name': teacher.name,
                            'designation': f'{member_designation}, {member_institute}',
                            'position': 'Member'
                        })
                else:
                    # External member (info stored in remarks as JSON)
                    try:
                        external_info = json.loads(assignment.remarks) if assignment.remarks else {}
                        if external_info.get('type') == 'external':
                            designation = f"{external_info.get('designation', '')}, {external_info.get('institute', '')}"
                            examination_committee.append({
                                'name': external_info.get('name', ''),
                                'designation': designation.strip(', '),
                                'position': 'Ext. Member'
                            })
                    except:
                        pass
        
        # Log saved_data for debugging
        if saved_data:
            current_app.logger.info(f'Saved data found for form_id {form_id}: {list(saved_data.keys())}')
            if 'question_preparation' in saved_data:
                current_app.logger.info(f'Question preparation items: {len(saved_data.get("question_preparation", []))}')
            if 'script_examination' in saved_data:
                current_app.logger.info(f'Script examination items: {len(saved_data.get("script_examination", []))}')
        else:
            current_app.logger.info(f'No saved_data found for form_id {form_id}')
        
        # Convert allowed combinations to list for template
        allowed_combinations_list = [{'session': s, 'year': y, 'term': t} for s, y, t in allowed_combinations]
        
        return render_template('exam_committee_chief/custom_remuneration.html',
                             teachers=teachers,
                             teachers_data=teachers_data,  # JSON-serializable version
                             academic_sessions=academic_sessions,
                             unique_years=unique_years,
                             unique_terms=unique_terms,
                             unique_batches=unique_batches,
                             curriculum_configs=configs_data,
                             current_session=current_session,
                             current_year=current_year,
                             current_term=current_term,
                             current_batch=current_batch,
                             examination_committee=saved_examination_committee if saved_examination_committee else examination_committee,
                             saved_data=saved_data,
                             allowed_combinations=allowed_combinations_list)

    @app.route('/exam-committee-chief/custom-remuneration/save', methods=['POST'])
    @login_required
    def exam_committee_chief_custom_remuneration_save():
        """Save custom remuneration form data to database"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        if not chief_assignments and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        try:
            data = request.get_json()
            if not data:
                return jsonify({'success': False, 'message': 'No data provided'}), 400
            
            from blueprints.remuneration_management.models import RemunerationForm
            
            form_id = data.get('form_id')  # For editing existing form
            session = data.get('session', '')
            year = data.get('year', '')
            term = data.get('term', '')
            allowed_combinations = {
                (str(a.academic_session), str(a.year), str(a.term))
                for a in (chief_assignments + member_assignments)
                if a.academic_session and a.year and a.term
            }
            submitted_combination = (str(session), str(year), str(term))
            if session and year and term and submitted_combination not in allowed_combinations:
                return jsonify({
                    'success': False,
                    'message': f'You are not authorized for Session: {session}, Year: {year}, Term: {term}.'
                }), 403
            
            # Generate title
            title_parts = []
            
            if session:
                title_parts.append(session)
            if year:
                title_parts.append(year)
            if term:
                title_parts.append(term)
            
            if title_parts:
                title = f"Custom Remuneration - {' '.join(title_parts)}"
            else:
                title = f"Custom Remuneration - {current_user.full_name}"
            
            # Create or update form entry
            form_entry = None
            if form_id:
                # Edit existing form (allow same-committee member editing too)
                form_entry = RemunerationForm.query.filter_by(id=form_id).first()
                if not form_entry:
                    return jsonify({'success': False, 'message': 'Form not found'}), 404
                can_edit = form_entry.user_id == current_user.id
                if not can_edit:
                    for assignment in member_assignments:
                        if (
                            assignment.academic_session == form_entry.academic_year and
                            assignment.year == form_entry.year and
                            assignment.term == form_entry.term
                        ):
                            can_edit = True
                            break
                if not can_edit:
                    return jsonify({'success': False, 'message': 'You do not have permission to edit this form'}), 403
            else:
                # Create new form
                form_entry = RemunerationForm(
                    user_id=current_user.id,
                    status='draft'
                )
                db.session.add(form_entry)
            
            # Update form fields
            form_entry.title = title
            form_entry.academic_year = session if session else ''
            form_entry.year = year if year else ''
            form_entry.term = term if term else ''
            form_entry.exam_start_date = data.get('exam_start_date', '')
            form_entry.exam_end_date = data.get('exam_end_date', '')
            
            # Calculate total amount if provided
            total_amount_str = data.get('total_amount', '0') or '0'
            try:
                total_amount = float(total_amount_str.replace(',', '').replace('৳', '').strip())
            except:
                total_amount = 0.0
            form_entry.total_amount = total_amount
            form_entry.total_in_words = data.get('total_in_words', '')
            
            # Save all form data as JSON (includes all sections)
            try:
                form_data_json = json.dumps(data)
                form_entry.form_data = form_data_json
            except Exception as json_error:
                current_app.logger.error(f'Error serializing form data to JSON: {str(json_error)}')
                return jsonify({'success': False, 'message': f'Error serializing data: {str(json_error)}'}), 500
            
            # Commit to database
            try:
                db.session.commit()
                current_app.logger.info(f'Custom remuneration form saved successfully. Form ID: {form_entry.id}, User ID: {current_user.id}')
            except Exception as commit_error:
                db.session.rollback()
                current_app.logger.error(f'Error committing to database: {str(commit_error)}', exc_info=True)
                return jsonify({'success': False, 'message': f'Database error: {str(commit_error)}'}), 500
            
            return jsonify({
                'success': True,
                'message': 'Custom remuneration form saved successfully',
                'form_id': form_entry.id
            })
        except Exception as e:
            db.session.rollback()
            error_msg = str(e)
            current_app.logger.error(f'Error saving custom remuneration form: {error_msg}', exc_info=True)
            # Log the full traceback
            import traceback
            current_app.logger.error(f'Traceback: {traceback.format_exc()}')
            return jsonify({'success': False, 'message': f'Failed to save form: {error_msg}'}), 500

    @app.route('/exam-committee-chief/custom-remuneration/save-simple', methods=['POST'])
    @login_required
    def exam_committee_chief_custom_remuneration_save_simple():
        """Simplified save route that accepts JSON"""
        try:
            # Accept JSON
            if request.is_json:
                data = request.get_json()
            elif request.content_type and 'application/json' in request.content_type:
                data = request.get_json()
            else:
                # Try to parse as JSON anyway
                try:
                    data = request.get_json(force=True)
                except:
                    # Fallback to form-data
                    form_data_str = request.form.get('form_data', '')
                    if form_data_str:
                        data = json.loads(form_data_str)
                    else:
                        return jsonify({'success': False, 'message': 'No data provided. Expected JSON.'}), 400
            
            if not data:
                return jsonify({'success': False, 'message': 'No data provided'}), 400
            
            current_app.logger.info(f'Received save request with data keys: {list(data.keys()) if data else "None"}')
            
            # Check if current user is assigned as Exam Committee Chief or Member
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if not teacher:
                return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
            
            # Get all active exam committee chief assignments for this teacher
            chief_assignments = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                assigned_teacher_id=teacher.id,
                status='active'
            ).all()
            
            # Get all active exam committee member assignments for this teacher
            member_assignments = DutyAssignment.query.filter_by(
                duty_type='exam_committee_member',
                assigned_teacher_id=teacher.id,
                status='active'
            ).all()
            
            if not chief_assignments and not member_assignments:
                return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
            
            # Get allowed session/year/term combinations (both chief and member)
            allowed_combinations = set()
            for assignment in chief_assignments:
                if assignment.academic_session and assignment.year and assignment.term:
                    allowed_combinations.add((
                        str(assignment.academic_session),
                        str(assignment.year),
                        str(assignment.term)
                    ))
            for assignment in member_assignments:
                if assignment.academic_session and assignment.year and assignment.term:
                    allowed_combinations.add((
                        str(assignment.academic_session),
                        str(assignment.year),
                        str(assignment.term)
                    ))
            
            # Validate that the submitted session/year/term is allowed
            session = data.get('session', '')
            year = data.get('year', '')
            term = data.get('term', '')
            
            if session and year and term:
                submitted_combination = (str(session), str(year), str(term))
                if submitted_combination not in allowed_combinations:
                    return jsonify({
                        'success': False,
                        'message': f'You are not authorized to create remuneration statement for Session: {session}, Year: {year}, Term: {term}. You can only create statements for the session/year/term combinations you are assigned as Exam Committee Chief.'
                    }), 403
            
            from blueprints.remuneration_management.models import RemunerationForm
            
            # Extract form_id - handle both string and int
            form_id_str = data.get('form_id') or request.form.get('form_id') or ''
            form_id = None
            if form_id_str:
                try:
                    form_id = int(str(form_id_str).strip())
                except (ValueError, TypeError):
                    form_id = None
            
            # Generate title
            session = data.get('session', '')
            year = data.get('year', '')
            term = data.get('term', '')
            title_parts = []
            if session:
                title_parts.append(session)
            if year:
                title_parts.append(year)
            if term:
                title_parts.append(term)
            
            title = f"Custom Remuneration - {' '.join(title_parts)}" if title_parts else f"Custom Remuneration - {current_user.full_name}"
            
            # Validate user_id is available
            if not current_user or not current_user.id:
                current_app.logger.error('Current user or user ID is missing')
                return jsonify({'success': False, 'message': 'User authentication error'}), 401
            
            # Create or update form entry
            if form_id:
                # Allow members to update Chief's forms
                form_entry = RemunerationForm.query.filter_by(id=form_id).first()
                if not form_entry:
                    current_app.logger.warning(f'Form not found for update. form_id: {form_id}')
                    return jsonify({'success': False, 'message': 'Form not found'}), 404
                
                # Check if user owns the form or is a member of the same committee
                can_edit = False
                if form_entry.user_id == current_user.id:
                    can_edit = True
                elif member_assignments:
                    # Check if current user is a member of the same committee
                    for member_assignment in member_assignments:
                        if (member_assignment.academic_session == form_entry.academic_year and
                            member_assignment.year == form_entry.year and
                            member_assignment.term == form_entry.term):
                            can_edit = True
                            break
                
                if not can_edit:
                    return jsonify({'success': False, 'message': 'You do not have permission to edit this form'}), 403
                
                # Ensure user_id is set (should already be set, but double-check)
                if not form_entry.user_id:
                    form_entry.user_id = current_user.id
                
                # If member is editing Chief's form, keep the original user_id (Chief's ID)
                # This way the form remains associated with the Chief
            else:
                # Create new form - chiefs and internal members are both allowed.
                form_entry = RemunerationForm(user_id=current_user.id, status='draft')
                db.session.add(form_entry)
            
            # Update fields - ensure all fields have safe defaults
            form_entry.title = title or f"Custom Remuneration - {current_user.full_name}"
            form_entry.academic_year = session or ''
            form_entry.year = year or ''
            form_entry.term = term or ''
            form_entry.exam_start_date = data.get('exam_start_date') or ''
            form_entry.exam_end_date = data.get('exam_end_date') or ''
            
            # Ensure user_id is set (critical for database constraint)
            if not form_entry.user_id:
                form_entry.user_id = current_user.id
            
            # Ensure status is set
            if not form_entry.status:
                form_entry.status = 'draft'
            
            # Total amount
            total_amount_str = data.get('total_amount', '0') or '0'
            try:
                total_amount = float(str(total_amount_str).replace(',', '').replace('৳', '').strip() or '0')
            except (ValueError, TypeError) as e:
                current_app.logger.warning(f'Error parsing total_amount: {e}, defaulting to 0.0')
                total_amount = 0.0
            form_entry.total_amount = total_amount
            form_entry.total_in_words = data.get('total_in_words') or ''
            
            # Save as JSON - Remove form_id from data before saving
            try:
                # Create a copy of data without form_id
                data_to_save = {k: v for k, v in data.items() if k != 'form_id'}
                
                # Ensure all data is JSON-serializable
                serializable_data = {}
                for key, value in data_to_save.items():
                    try:
                        # Handle None values
                        if value is None:
                            serializable_data[key] = None
                            continue
                        
                        # Test if serializable
                        json.dumps(value, default=str, ensure_ascii=False)
                        serializable_data[key] = value
                    except (TypeError, ValueError) as e:
                        # Convert non-serializable to string
                        current_app.logger.warning(f'Converting non-serializable value for key {key}: {str(e)}')
                        try:
                            serializable_data[key] = str(value) if value is not None else None
                        except Exception as str_error:
                            current_app.logger.error(f'Error converting {key} to string: {str_error}')
                            serializable_data[key] = None
                
                form_entry.form_data = json.dumps(serializable_data, ensure_ascii=False, default=str)
                current_app.logger.info(f'✅ Form data serialized. Size: {len(form_entry.form_data)} chars, Keys: {len(serializable_data)}')
            except Exception as e:
                current_app.logger.error(f'❌ JSON serialization error: {str(e)}')
                import traceback
                current_app.logger.error(f'Full traceback: {traceback.format_exc()}')
                db.session.rollback()
                return jsonify({'success': False, 'message': f'Data format error: {str(e)}'}), 500
            
            # Validate before commit
            if not form_entry.user_id:
                current_app.logger.error('Validation failed: user_id is missing before commit')
                db.session.rollback()
                return jsonify({'success': False, 'message': 'Validation error: user_id is required'}), 500
            
            # Commit with retry mechanism
            max_commit_retries = 3
            commit_success = False
            last_error = None
            
            for attempt in range(max_commit_retries):
                try:
                    db.session.commit()
                    current_app.logger.info(f'✅ Form saved successfully. ID: {form_entry.id}, User: {current_user.id}, Title: {title}')
                    commit_success = True
                    break
                except Exception as e:
                    last_error = e
                    db.session.rollback()
                    import traceback
                    error_traceback = traceback.format_exc()
                    current_app.logger.error(f'❌ Database commit error (attempt {attempt + 1}/{max_commit_retries}): {str(e)}')
                    
                    # If not the last attempt, refresh the object and retry
                    if attempt < max_commit_retries - 1:
                        try:
                            db.session.refresh(form_entry)
                            current_app.logger.info(f'Retrying commit (attempt {attempt + 2})...')
                            continue
                        except:
                            pass
                    
                    current_app.logger.error(f'Full traceback: {error_traceback}')
            
            if not commit_success:
                # Provide more helpful error message
                error_msg = str(last_error)
                if 'NOT NULL constraint' in error_msg:
                    return jsonify({'success': False, 'message': f'Database constraint error: Missing required field. Details: {error_msg}'}), 500
                elif 'UNIQUE constraint' in error_msg:
                    return jsonify({'success': False, 'message': f'Database constraint error: Duplicate entry. Details: {error_msg}'}), 500
                elif 'MySQL server has gone away' in error_msg or 'Lost connection' in error_msg:
                    return jsonify({'success': False, 'message': 'Database connection lost. Please try again.'}), 500
                else:
                    return jsonify({'success': False, 'message': f'Database error: {error_msg}'}), 500
            
            return jsonify({
                'success': True,
                'message': 'Form saved successfully',
                'form_id': form_entry.id
            })
            
        except json.JSONDecodeError as e:
            db.session.rollback()
            current_app.logger.error(f'JSON decode error: {str(e)}')
            return jsonify({'success': False, 'message': f'Invalid JSON format: {str(e)}'}), 400
        except Exception as e:
            db.session.rollback()
            error_msg = str(e)
            import traceback
            error_traceback = traceback.format_exc()
            current_app.logger.error(f'❌ Unexpected error in simple save: {error_msg}')
            current_app.logger.error(f'Full traceback: {error_traceback}')
            return jsonify({'success': False, 'message': f'Unexpected error: {error_msg}. Please check server logs for details.'}), 500

    @app.route('/exam-committee-chief/custom-remuneration/load/<int:form_id>', methods=['GET'])
    @login_required
    def exam_committee_chief_custom_remuneration_load(form_id):
        """Load saved custom remuneration form data"""
        # Check if current user is assigned as Exam Committee Chief
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
        
        # Check if user is Chief or Member
        chief_assignment = DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            assigned_teacher_id=teacher.id,
            status='active'
        ).first()
        
        member_assignments = DutyAssignment.query.filter_by(
            duty_type='exam_committee_member',
            assigned_teacher_id=teacher.id,
            status='active'
        ).all()
        
        if not chief_assignment and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        try:
            form_entry = RemunerationForm.query.filter_by(id=form_id).first()
            if not form_entry:
                return jsonify({'success': False, 'message': 'Form not found'}), 404
            
            # Check if user owns the form or is a member of the same committee
            can_access = False
            if form_entry.user_id == current_user.id:
                can_access = True
            elif member_assignments:
                # Check if current user is a member of the same committee
                for member_assignment in member_assignments:
                    if (member_assignment.academic_session == form_entry.academic_year and
                        member_assignment.year == form_entry.year and
                        member_assignment.term == form_entry.term):
                        can_access = True
                        break
            
            if not can_access:
                return jsonify({'success': False, 'message': 'You do not have permission to access this form'}), 403
            
            # Load from JSON data
            if form_entry.form_data:
                try:
                    saved_data = json.loads(form_entry.form_data)
                    saved_data['form_id'] = form_entry.id
                    # Merge exam dates from database fields (in case they're not in JSON)
                    if form_entry.exam_start_date:
                        saved_data['exam_start_date'] = form_entry.exam_start_date
                    if form_entry.exam_end_date:
                        saved_data['exam_end_date'] = form_entry.exam_end_date
                    return jsonify({'success': True, 'data': saved_data})
                except Exception as e:
                    current_app.logger.error(f'Error parsing form_data JSON: {str(e)}')
                    return jsonify({'success': False, 'message': 'Error loading form data'}), 500
            else:
                return jsonify({'success': False, 'message': 'No form data found'}), 404
                
        except Exception as e:
            current_app.logger.error(f'Error loading custom remuneration form: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Failed to load form: {str(e)}'}), 500

    @app.route('/exam-committee-chief/custom-remuneration/get-by-session', methods=['GET'])
    @login_required
    def exam_committee_chief_custom_remuneration_get_by_session():
        """Get Remuneration Statement Form data by session, year, and term for Bill Form"""
        try:
            # Check if current user is assigned as Exam Committee Chief or Member
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if not teacher:
                return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
            
            # Get parameters
            session = request.args.get('session', '').strip()
            year = request.args.get('year', '').strip()
            term = request.args.get('term', '').strip()
            
            if not session or not year or not term:
                return jsonify({'success': False, 'message': 'Session, Year, and Term are required'}), 400
            
            # Check if user is Chief
            chief_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                assigned_teacher_id=teacher.id,
                status='active'
            ).first()
            
            # Check if user is Member of the same committee
            member_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_member',
                assigned_teacher_id=teacher.id,
                academic_session=session,
                year=year,
                term=term,
                status='active'
            ).first()
            
            # Allow all teachers to access this endpoint for auto-population
            # Removed restriction - all teachers can now access remuneration statement data
            
            # Get parameters (duplicate check removed - already validated above)
            from blueprints.remuneration_management.models import RemunerationForm
            
            # Get the Exam Committee Chief's form for this session/year/term
            # All teachers can access this for auto-population
            form_entry = None
            
            # First, try to get the current user's form if they are the Chief
            if chief_assignment:
                form_entry = RemunerationForm.query.filter_by(
                    user_id=current_user.id,
                    academic_year=session,
                    year=year,
                    term=term
                ).order_by(RemunerationForm.id.desc()).first()
            
            # If not found, try to find the Chief's form for this session/year/term (works for all teachers)
            if not form_entry:
                chief_for_committee = DutyAssignment.query.filter_by(
                    duty_type='exam_committee_chief',
                    academic_session=session,
                    year=year,
                    term=term,
                    status='active'
                ).first()
                
                if chief_for_committee and chief_for_committee.assigned_teacher:
                    chief_user = User.query.filter_by(full_name=chief_for_committee.assigned_teacher.name).first()
                    if chief_user:
                        form_entry = RemunerationForm.query.filter_by(
                            user_id=chief_user.id,
                            academic_year=session,
                            year=year,
                            term=term
                        ).order_by(RemunerationForm.id.desc()).first()
            
            if not form_entry or not form_entry.form_data:
                return jsonify({
                    'success': True,
                    'data': None,
                    'message': 'No Remuneration Statement Form found for this session/year/term'
                })
            
            # Parse and return the form data
            try:
                saved_data = json.loads(form_entry.form_data)
                saved_data['form_id'] = form_entry.id
                # Include exam dates from database fields (in case they're not in JSON)
                if form_entry.exam_start_date:
                    saved_data['exam_start_date'] = form_entry.exam_start_date
                if form_entry.exam_end_date:
                    saved_data['exam_end_date'] = form_entry.exam_end_date
                return jsonify({
                    'success': True,
                    'data': saved_data,
                    'message': 'Remuneration Statement Form data loaded successfully'
                })
            except Exception as e:
                current_app.logger.error(f'Error parsing form_data JSON: {str(e)}')
                return jsonify({'success': False, 'message': 'Error loading form data'}), 500
                
        except Exception as e:
            current_app.logger.error(f'Error getting custom remuneration form by session: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Failed to get form: {str(e)}'}), 500

    @app.route('/exam-committee-chief/custom-remuneration/get-committee', methods=['GET'])
    @login_required
    def exam_committee_chief_get_committee():
        """Get examination committee data for specific session/year/term"""
        try:
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if not teacher:
                return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
            
            session = request.args.get('session', '').strip()
            year = request.args.get('year', '').strip()
            term = request.args.get('term', '').strip()
            
            if not session or not year or not term:
                return jsonify({'success': False, 'message': 'Session, Year, and Term are required'}), 400
            
            chief_assignments, member_assignments = _get_exam_committee_assignments(
                teacher,
                academic_session=session,
                year=year,
                term=term
            )
            if not chief_assignments and not member_assignments:
                return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief/Member for this session/year/term'}), 403

            chief_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                academic_session=session,
                year=year,
                term=term,
                status='active'
            ).first()
            
            # Build examination committee
            examination_committee = []
            
            # Add Chief (Chairman)
            if chief_assignment.assigned_teacher:
                chief_teacher = chief_assignment.assigned_teacher
                chief_designation = ''
                chief_institute = ''
                
                try:
                    chief_data = json.loads(chief_assignment.remarks) if chief_assignment.remarks else {}
                    if chief_data.get('type') == 'chief':
                        chief_designation = chief_data.get('designation', '')
                        chief_institute = chief_data.get('institute', '')
                except:
                    pass
                
                if not chief_designation and chief_teacher.designation:
                    chief_designation = chief_teacher.designation
                if not chief_institute and chief_teacher.institute:
                    chief_institute = chief_teacher.institute
                
                if not chief_designation:
                    chief_designation = 'Head, Law Discipline, KU'
                if not chief_institute:
                    chief_institute = 'Law Discipline, KU'
                
                examination_committee.append({
                    'name': chief_teacher.name,
                    'designation': f'{chief_designation}, {chief_institute}',
                    'position': 'Chairman'
                })
            
            # Get committee members
            committee_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == session,
                DutyAssignment.year == year,
                DutyAssignment.term == term,
                DutyAssignment.status == 'active'
            ).all()
            
            for assignment in committee_assignments:
                if assignment.assigned_teacher_id:
                    teacher = assignment.assigned_teacher
                    if teacher:
                        member_designation = ''
                        member_institute = ''
                        
                        try:
                            member_data = json.loads(assignment.remarks) if assignment.remarks else {}
                            if member_data.get('type') == 'internal':
                                member_designation = member_data.get('designation', '')
                                member_institute = member_data.get('institute', '')
                        except:
                            pass
                        
                        if not member_designation and teacher.designation:
                            member_designation = teacher.designation
                        if not member_institute and teacher.institute:
                            member_institute = teacher.institute
                        
                        if not member_designation:
                            member_designation = 'Assistant Professor'
                        if not member_institute:
                            member_institute = 'Law Discipline, KU'
                        
                        examination_committee.append({
                            'name': teacher.name,
                            'designation': f'{member_designation}, {member_institute}',
                            'position': 'Member'
                        })
                else:
                    try:
                        external_info = json.loads(assignment.remarks) if assignment.remarks else {}
                        if external_info.get('type') == 'external':
                            designation = f"{external_info.get('designation', '')}, {external_info.get('institute', '')}"
                            examination_committee.append({
                                'name': external_info.get('name', ''),
                                'designation': designation.strip(', '),
                                'position': 'Ext. Member'
                            })
                    except:
                        pass
            
            return jsonify({
                'success': True,
                'committee': examination_committee
            })
            
        except Exception as e:
            current_app.logger.error(f'Error getting committee data: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Failed to get committee data: {str(e)}'}), 500

    @app.route('/exam-committee-chief/custom-remuneration/delete/<int:form_id>', methods=['POST'])
    @login_required
    def exam_committee_chief_custom_remuneration_delete(form_id):
        """Delete a custom remuneration form"""
        try:
            # Check if current user is assigned as Exam Committee Chief or Member
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if not teacher:
                return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

            chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
            if not chief_assignments and not member_assignments:
                return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
            
            from blueprints.remuneration_management.models import RemunerationForm
            
            # Get the form
            form_entry = RemunerationForm.query.filter_by(id=form_id).first()
            if not form_entry:
                return jsonify({'success': False, 'message': 'Form not found'}), 404

            can_delete = form_entry.user_id == current_user.id
            if not can_delete:
                for assignment in member_assignments:
                    if (
                        assignment.academic_session == form_entry.academic_year and
                        assignment.year == form_entry.year and
                        assignment.term == form_entry.term
                    ):
                        can_delete = True
                        break
            if not can_delete:
                return jsonify({'success': False, 'message': 'You do not have permission to delete this form'}), 403
            
            # Delete the form
            db.session.delete(form_entry)
            db.session.commit()
            
            current_app.logger.info(f'Custom remuneration form deleted. Form ID: {form_id}, User ID: {current_user.id}')
            
            return jsonify({
                'success': True,
                'message': 'Form deleted successfully'
            })
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error deleting custom remuneration form: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Failed to delete form: {str(e)}'}), 500

    @app.route('/exam-committee-chief/custom-remuneration/export-pdf', methods=['POST'])
    @login_required
    def exam_committee_chief_custom_remuneration_export_pdf():
        """Export custom remuneration form to PDF matching scanned copy format"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404

        chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
        if not chief_assignments and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
        
        # Get allowed session/year/term combinations
        allowed_combinations = set()
        for assignment in (chief_assignments + member_assignments):
            if assignment.academic_session and assignment.year and assignment.term:
                allowed_combinations.add((
                    str(assignment.academic_session),
                    str(assignment.year),
                    str(assignment.term)
                ))
        
        try:
            from io import BytesIO
            from flask import send_file
            import os
            import json
            
            data = request.get_json()
            if not data:
                return jsonify({'error': 'No data provided'}), 400
            
            # Get form data
            program = data.get('program', '')
            session = data.get('session', '')
            year = data.get('year', '')
            term = data.get('term', '')
            batch = data.get('batch', '')
            total_amount = data.get('total_amount', '0') or '0'
            total_in_words = data.get('total_in_words', '')
            
            # Validate that the submitted session/year/term is allowed
            if session and year and term:
                submitted_combination = (str(session), str(year), str(term))
                if submitted_combination not in allowed_combinations:
                    return jsonify({
                        'success': False,
                        'message': f'You are not authorized to export remuneration statement for Session: {session}, Year: {year}, Term: {term}. You can only export statements for the session/year/term combinations where you are assigned as Exam Committee Chief/Member.'
                    }), 403
            
            # Get chief name for signature from examination committee
            chief_name = ''  # Will be empty if not found in form data
            chief_title = 'Head'
            examination_committee = data.get('examination_committee', [])
            for member in examination_committee:
                if member.get('position') == 'Chairman':
                    chief_name = member.get('name', '')
                    # Extract title from designation (e.g., "Head, Law Discipline, KU" -> "Head")
                    designation = member.get('designation', '')
                    if designation:
                        parts = designation.split(',')
                        if parts:
                            chief_title = parts[0].strip()
                    break
            
            # Calculate unique courses for Section 2 (Question Preparation)
            question_preparation = data.get('question_preparation', [])
            unique_courses_qp = set()
            qp_total_questions = 0
            # Process question preparation data - extract course codes and calculate totals
            processed_qp = []
            for item in question_preparation:
                course_full = item.get('course', '') or ''
                course_code = ''
                course_name = ''
                
                # Try to parse as JSON first (if it's stored as JSON string)
                try:
                    if course_full.strip().startswith('{'):
                        course_data = json.loads(course_full)
                        course_code = course_data.get('course_code', '') or ''
                        course_name = course_data.get('course_name', '') or ''
                        # Also get section from JSON if not in item
                        if not item.get('section') and course_data.get('section'):
                            item['section'] = course_data.get('section')
                except:
                    pass
                
                # If JSON parsing didn't work or didn't provide course_code, try string split
                if not course_code:
                    if ' - ' in course_full:
                        course_code = course_full.split(' - ', 1)[0].strip()
                        course_name = course_full.split(' - ', 1)[1].strip()
                    else:
                        course_code = course_full.strip()
                        course_name = ''
                
                if course_code:
                    unique_courses_qp.add(course_code)
                
                # Calculate total questions
                try:
                    qty = int(item.get('questions', '1') or '1')
                    qp_total_questions += qty
                except:
                    qp_total_questions += 1
                
                # Prepare processed item for template
                processed_qp.append({
                    'course_code': course_code,
                    'course_name': course_name,
                    'section': item.get('section', '') or 'Full',
                    'teacher': item.get('teacher', '') or '',
                    'designation': item.get('designation', '') or '',
                    'questions': item.get('questions', '1') or '1'
                })
            
            # Calculate unique courses and total scripts for Section 3 (Script Examination)
            script_examination_raw = data.get('script_examination', [])
            unique_courses_se = set()
            total_scripts_se = 0
            processed_se = []
            for item in script_examination_raw:
                course_full = item.get('course', '') or ''
                course_code = ''
                course_name = ''
                
                # Try to parse as JSON first
                try:
                    if course_full.strip().startswith('{'):
                        course_data = json.loads(course_full)
                        course_code = course_data.get('course_code', '') or ''
                        course_name = course_data.get('course_name', '') or ''
                        if not item.get('section') and course_data.get('section'):
                            item['section'] = course_data.get('section')
                except:
                    pass
                
                # If JSON parsing didn't work, try string split
                if not course_code:
                    if ' - ' in course_full:
                        course_code = course_full.split(' - ', 1)[0].strip()
                        course_name = course_full.split(' - ', 1)[1].strip()
                    else:
                        course_code = course_full.strip()
                        course_name = ''
                
                if course_code:
                    unique_courses_se.add(course_code)
                
                try:
                    scripts = int(item.get('scripts', 0) or 0)
                    total_scripts_se += scripts
                except:
                    pass
                
                processed_se.append({
                    'course_code': course_code,
                    'course_name': course_name,
                    'section': item.get('section', '') or 'Full',
                    'teacher': item.get('teacher', '') or '',
                    'designation': item.get('designation', '') or '',
                    'scripts': item.get('scripts', '0') or '0'
                })
            
            # Calculate unique courses for Section 4 (Class Test)
            class_test_raw = data.get('class_test', [])
            unique_courses_ct = set()
            processed_ct = []
            for item in class_test_raw:
                course_full = item.get('course', '') or ''
                course_code = ''
                course_name = ''
                
                # Try to parse as JSON first
                try:
                    if course_full.strip().startswith('{'):
                        course_data = json.loads(course_full)
                        course_code = course_data.get('course_code', '') or ''
                        course_name = course_data.get('course_name', '') or ''
                        if not item.get('section') and course_data.get('section'):
                            item['section'] = course_data.get('section')
                except:
                    pass
                
                # If JSON parsing didn't work, try string split
                if not course_code:
                    if ' - ' in course_full:
                        course_code = course_full.split(' - ', 1)[0].strip()
                        course_name = course_full.split(' - ', 1)[1].strip()
                    else:
                        course_code = course_full.strip()
                        course_name = ''
                
                if course_code:
                    unique_courses_ct.add(course_code)
                
                processed_ct.append({
                    'course_code': course_code,
                    'course_name': course_name,
                    'section': item.get('section', '') or 'Full',
                    'teacher': item.get('teacher', '') or '',
                    'designation': item.get('designation', '') or '',
                    'class_test': item.get('class_test', '0') or '0',
                    'students': item.get('students', '0') or '0'
                })
            
            # Prepare data for template
            template_data = {
                'program': program,
                'session': session,
                'year': year,
                'term': term,
                'batch': batch,
                'total_amount': total_amount,
                'total_in_words': total_in_words,
                'chief_name': chief_name,
                'chief_title': chief_title,
                'examination_committee': examination_committee,
                'question_preparation': processed_qp,
                'unique_qp_courses_count': len(unique_courses_qp),
                'qp_total_questions': qp_total_questions,
                'script_examination': processed_se,
                'unique_se_courses_count': len(unique_courses_se),
                'total_scripts_se': total_scripts_se,
                'class_test': processed_ct,
                'unique_ct_courses_count': len(unique_courses_ct),
                'moderation_committee': data.get('moderation_committee', []),
                'sessional_assessment': data.get('sessional_assessment', []),
                'sessional_viva': data.get('sessional_viva', []),
                'script_scrutiny': data.get('script_scrutiny', []),
                'tabulation': data.get('tabulation', []),
                'coding_decoding': data.get('coding_decoding', []),
                'invigilation': data.get('invigilation', []),
                'thesis_supervision': data.get('thesis_supervision', []),
                'viva': data.get('viva', []),
                'question_typing': data.get('question_typing', []),
                'question_photocopy': data.get('question_photocopy', []),
                'official_support': data.get('official_support', [])
            }
            
            # Process sessional sections to extract course_code from course field
            def process_course_field(items, course_key='course'):
                processed = []
                for item in items or []:
                    course_full = str(item.get(course_key, '') or '')
                    course_code = ''
                    
                    # Try to parse as JSON first
                    try:
                        if course_full.strip().startswith('{'):
                            course_data = json.loads(course_full)
                            course_code = course_data.get('course_code', '') or ''
                    except:
                        pass
                    
                    # If JSON parsing didn't work, try string split
                    if not course_code:
                        if ' - ' in course_full:
                            course_code = course_full.split(' - ', 1)[0].strip()
                        else:
                            course_code = course_full.strip()
                    
                    processed_item = dict(item)
                    processed_item['course_code'] = course_code
                    processed.append(processed_item)
                return processed
            
            template_data['sessional_assessment'] = process_course_field(template_data.get('sessional_assessment', []))
            template_data['sessional_viva'] = process_course_field(template_data.get('sessional_viva', []))
            
            # Ensure all list fields are lists (not None)
            for key in ['examination_committee', 'moderation_committee', 'sessional_assessment', 
                       'sessional_viva', 'script_scrutiny', 'tabulation', 'coding_decoding',
                       'invigilation', 'thesis_supervision', 'viva', 'question_typing', 'question_photocopy', 'official_support']:
                if template_data[key] is None:
                    template_data[key] = []
            
            # Build list of sections with data only, for sequential numbering
            sections_list = []
            
            # Section 1: Examination Committee
            if template_data.get('examination_committee') and len(template_data['examination_committee']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'examination_committee',
                    'title': 'Examination Committee',
                    'data': template_data['examination_committee']
                })
            
            # Section 2: Question Preparation
            if template_data.get('question_preparation') and len(template_data['question_preparation']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'question_preparation',
                    'title': f'Question Preparation: ({template_data["unique_qp_courses_count"]} Courses {template_data["qp_total_questions"]} Questions)',
                    'data': template_data['question_preparation'],
                    'unique_courses_count': template_data['unique_qp_courses_count'],
                    'total_questions': template_data['qp_total_questions']
                })
            
            # Section 3: Script Examination
            if template_data.get('script_examination') and len(template_data['script_examination']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'script_examination',
                    'title': f'Answer Script Examination: ({template_data["unique_se_courses_count"]} Courses {template_data["total_scripts_se"]} Scripts)',
                    'data': template_data['script_examination'],
                    'unique_courses_count': template_data['unique_se_courses_count'],
                    'total_scripts': template_data['total_scripts_se']
                })
            
            # Section 4: Class Test
            if template_data.get('class_test') and len(template_data['class_test']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'class_test',
                    'title': f'Class Test: ({template_data["unique_ct_courses_count"]} Courses)',
                    'data': template_data['class_test'],
                    'unique_courses_count': template_data['unique_ct_courses_count']
                })
            
            # Section 5: Moderation Committee
            if template_data.get('moderation_committee') and len(template_data['moderation_committee']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'moderation_committee',
                    'title': 'Moderation Committee',
                    'data': template_data['moderation_committee']
                })
            
            # Section 6: Sessional Assessment
            if template_data.get('sessional_assessment') and len(template_data['sessional_assessment']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'sessional_assessment',
                    'title': f'Sessional Assessment: ({len(template_data["sessional_assessment"])} Course{"s" if len(template_data["sessional_assessment"]) != 1 else ""})',
                    'data': template_data['sessional_assessment']
                })
            
            # Section 7: Sessional Viva
            if template_data.get('sessional_viva') and len(template_data['sessional_viva']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'sessional_viva',
                    'title': f'Sessional viva: ({len(template_data["sessional_viva"])} Course{"s" if len(template_data["sessional_viva"]) != 1 else ""})',
                    'data': template_data['sessional_viva']
                })
            
            # Section 8: Script Scrutiny
            if template_data.get('script_scrutiny') and len(template_data['script_scrutiny']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'script_scrutiny',
                    'title': 'Answer Script Scrutiny',
                    'data': template_data['script_scrutiny']
                })
            
            # Section 9: Tabulation
            if template_data.get('tabulation') and len(template_data['tabulation']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'tabulation',
                    'title': 'Tabulation',
                    'data': template_data['tabulation']
                })
            
            # Section 10: Coding & Decoding
            if template_data.get('coding_decoding') and len(template_data['coding_decoding']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'coding_decoding',
                    'title': 'Coding & Decoding',
                    'data': template_data['coding_decoding']
                })
            
            # Section 11: Invigilation
            if template_data.get('invigilation') and len(template_data['invigilation']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'invigilation',
                    'title': 'Chief Invigilation/Invigilation',
                    'data': template_data['invigilation']
                })
            
            # Section 12: Thesis Supervision
            if template_data.get('thesis_supervision') and len(template_data['thesis_supervision']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'thesis_supervision',
                    'title': 'Thesis Supervision',
                    'data': template_data['thesis_supervision']
                })
            
            # Section 13: Viva
            if template_data.get('viva') and len(template_data['viva']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'viva',
                    'title': 'Viva',
                    'data': template_data['viva']
                })
            
            # Section 14: Question Typing
            if template_data.get('question_typing') and len(template_data['question_typing']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'question_typing',
                    'title': f'Question Typing: ({len(template_data["question_typing"])} Courses)',
                    'data': template_data['question_typing']
                })
            
            # Section 15: Question Photocopy
            if template_data.get('question_photocopy') and len(template_data['question_photocopy']) > 0:
                sections_list.append({
                    'number': len(sections_list) + 1,
                    'type': 'question_photocopy',
                    'title': f'Question Photocopy: ({len(template_data["question_photocopy"])} Courses)',
                    'data': template_data['question_photocopy']
                })
            
            # Add sections_list to template_data
            template_data['sections_list'] = sections_list
            
            # Generate PDF using WeasyPrint with HTML template
            from flask import render_template
            from weasyprint import HTML, CSS
            
            # Convert year/term to ordinal words for display
            year_words = {
                '1': 'First', '2': 'Second', '3': 'Third', '4': 'Fourth',
                '5': 'Fifth', '6': 'Sixth', '7': 'Seventh', '8': 'Eighth'
            }
            term_words = {
                '1': 'First', '2': 'Second', '3': 'Third', '4': 'Fourth'
            }
            year_display = year_words.get(str(year), str(year))
            term_display = term_words.get(str(term), str(term))
            
            # Update template_data with display values (template expects "First Year", "First Term" format)
            template_data['year'] = f"{year_display} Year"
            template_data['term'] = f"{term_display} Term"
            
            # Render HTML template with data
            html_content = render_template(
                'exam_committee_chief/custom_remuneration_pdf.html',
                **template_data
            )
            
            # Generate PDF using WeasyPrint
            pdf_buffer = BytesIO()
            
            # Create HTML object
            html_obj = HTML(string=html_content, base_url=request.url_root)
            
            # Write PDF
            html_obj.write_pdf(pdf_buffer, presentational_hints=True)
            pdf_buffer.seek(0)
            pdf_data = pdf_buffer.getvalue()
            
            # Enhanced headers for cPanel compatibility
            response = Response(
                pdf_data,
                mimetype='application/pdf',
                headers={
                    'Content-Disposition': f'attachment; filename="Custom_Remuneration_Statement.pdf"; filename*=UTF-8\'\'Custom_Remuneration_Statement.pdf',
                    'Content-Length': str(len(pdf_data)),
                    'Cache-Control': 'no-cache, no-store, must-revalidate',
                    'Pragma': 'no-cache',
                    'Expires': '0',
                    'X-Content-Type-Options': 'nosniff',
                    'X-Frame-Options': 'DENY'
                }
            )
            
            return response
            
        except Exception as e:
            current_app.logger.error(f'Error generating custom remuneration PDF: {str(e)}', exc_info=True)
            import traceback
            error_trace = traceback.format_exc()
            current_app.logger.error(f'Full traceback: {error_trace}')
            return jsonify({'success': False, 'error': f'Failed to generate PDF: {str(e)}', 'message': str(e)}), 500

    @app.route('/exam-committee-chief/get-courses-with-assignments', methods=['GET'])
    @login_required
    def get_courses_with_assignments():
        """Get courses with teacher assignments for a given session/year/term"""
        academic_session = request.args.get('session', '').strip()
        year = request.args.get('year', '').strip()
        term = request.args.get('term', '').strip()
        course_type_filter = request.args.get('course_type', '').strip().lower()  # 'theory', 'sessional', or empty for all
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Session, Year, and Term are required'}), 400
        
        try:
            from blueprints.course_management.models import CourseSessionAssignment, Course, CurriculumYearTerm
            
            # Normalize year/term for matching (remove "Year" and "Term" suffixes if present)
            def normalize_label(label):
                if not label:
                    return ''
                label = str(label).strip()
                for suffix in [' Year', ' Term', 'Year', 'Term']:
                    if label.lower().endswith(suffix.lower()):
                        label = label[:-len(suffix)].strip()
                return label
            
            normalized_year = normalize_label(year)
            normalized_term = normalize_label(term)
            
            # Get all course assignments for this session/year/term
            # Try exact match first
            assignments = CourseSessionAssignment.query.filter(
                CourseSessionAssignment.academic_session == academic_session
            ).all()
            
            # Filter by normalized year/term
            matching_assignments = []
            for assignment in assignments:
                assign_year = normalize_label(assignment.year or '')
                assign_term = normalize_label(assignment.term or '')
                
                if assign_year == normalized_year and assign_term == normalized_term:
                    matching_assignments.append(assignment)
            
            current_app.logger.info(f'Found {len(matching_assignments)} assignments for session={academic_session}, year={year} (normalized: {normalized_year}), term={term} (normalized: {normalized_term})')
            
            courses_data = []
            for assignment in matching_assignments:
                if assignment.course:
                    course_type_raw = assignment.course.course_type
                    course_type = (course_type_raw or '').strip().lower()
                    
                    # Apply course_type filter if specified
                    if course_type_filter:
                        if course_type != course_type_filter:
                            current_app.logger.debug(f'🚫 SKIPPING course (filter: {course_type_filter}): {assignment.course.course_code} - {assignment.course.course_name} (type: "{course_type_raw}" -> "{course_type}")')
                            continue
                    else:
                        # Default: Only Theory courses (for backward compatibility)
                        if course_type != 'theory':
                            current_app.logger.info(f'🚫 SKIPPING non-theory course: {assignment.course.course_code} - {assignment.course.course_name} (type: "{course_type_raw}" -> "{course_type}")')
                            continue
                    
                    current_app.logger.debug(f'✅ INCLUDING {course_type} course: {assignment.course.course_code} - {assignment.course.course_name}')
                    
                    # Get teacher designation and institute (allow courses without teachers for Class Test)
                    if assignment.teacher:
                        teacher = assignment.teacher
                        teacher_name = teacher.name
                        designation = teacher.designation or 'Assistant Professor'
                        institute = teacher.institute or 'Law Discipline, KU'
                    else:
                        teacher_name = 'Not Assigned'
                        designation = 'Assistant Professor'
                        institute = 'Law Discipline, KU'
                    
                    courses_data.append({
                        'course_id': assignment.course_id,
                        'course_code': assignment.course.course_code,
                        'course_name': assignment.course.course_name,
                        'section': assignment.section or 'Full',
                        'teacher_id': assignment.teacher_id,
                        'teacher_name': teacher_name,
                        'designation': designation,
                        'institute': institute,
                        'designation_institute': f'{designation}, {institute}',
                        'course_type': assignment.course.course_type  # Include for frontend verification
                    })
            
            # If no assignments found, try to get courses from CurriculumYearTerm
            if not courses_data:
                current_app.logger.info(f'No CourseSessionAssignment found, trying CurriculumYearTerm for session={academic_session}, year={year}, term={term}')
                
                # Find curriculum configs matching the criteria
                configs = CurriculumYearTerm.query.filter_by(
                    academic_session=academic_session
                ).all()
                
                # Filter by normalized year/term
                matching_configs = []
                for config in configs:
                    config_year = normalize_label(config.year or '')
                    config_term = normalize_label(config.term or '')
                    
                    if config_year == normalized_year and config_term == normalized_term:
                        matching_configs.append(config)
                
                if matching_configs:
                    # Get all courses for these curricula
                    curriculum_ids = [config.curriculum_id for config in matching_configs]
                    all_courses = Course.query.filter(
                        Course.curriculum_id.in_(curriculum_ids)
                    ).order_by(Course.course_code).all()
                    
                    # Filter courses by course_type if specified
                    filtered_courses_by_type = []
                    for course in all_courses:
                        course_type_raw = course.course_type
                        course_type = (course_type_raw or '').strip().lower()
                        
                        # Apply course_type filter
                        if course_type_filter:
                            if course_type == course_type_filter:
                                filtered_courses_by_type.append(course)
                                current_app.logger.debug(f'✅ INCLUDING {course_type} course: {course.course_code} - {course.course_name}')
                            else:
                                current_app.logger.debug(f'🚫 SKIPPING course (filter: {course_type_filter}): {course.course_code} - {course.course_name} (type: "{course_type_raw}" -> "{course_type}")')
                        else:
                            # Default: Only Theory courses (for backward compatibility)
                            if course_type == 'theory':
                                filtered_courses_by_type.append(course)
                                current_app.logger.debug(f'✅ INCLUDING theory course: {course.course_code} - {course.course_name}')
                            else:
                                current_app.logger.info(f'🚫 SKIPPING non-theory course: {course.course_code} - {course.course_name} (type: "{course_type_raw}" -> "{course_type}")')
                    
                    current_app.logger.info(f'Found {len(filtered_courses_by_type)} {course_type_filter or "theory"} courses from CurriculumYearTerm after filtering')
                    
                    # Filter courses by their year/term to match the selected year/term
                    filtered_courses = []
                    for course in filtered_courses_by_type:
                        course_year = normalize_label(course.display_year or '')
                        course_term = normalize_label(course.display_term or '')
                        
                        # Match course year/term with selected year/term
                        if course_year == normalized_year and course_term == normalized_term:
                            filtered_courses.append(course)
                    
                    current_app.logger.info(f'Found {len(filtered_courses)} {course_type_filter or "theory"} courses matching year={normalized_year}, term={normalized_term}')
                    
                    for course in filtered_courses:
                        # Try to find teacher assignment for this course matching the session/year/term
                        course_assignment = CourseSessionAssignment.query.filter_by(
                            course_id=course.id,
                            academic_session=academic_session
                        ).first()
                        
                        # If no assignment found for this session, try to find any assignment for this course
                        if not course_assignment:
                            course_assignment = CourseSessionAssignment.query.filter_by(
                                course_id=course.id
                            ).first()
                        
                        teacher_name = 'Not Assigned'
                        designation = 'Assistant Professor'
                        institute = 'Law Discipline, KU'
                        section = 'Full'
                        
                        if course_assignment and course_assignment.teacher:
                            teacher = course_assignment.teacher
                            teacher_name = teacher.name
                            designation = teacher.designation or 'Assistant Professor'
                            institute = teacher.institute or 'Law Discipline, KU'
                            section = course_assignment.section or 'Full'
                        
                        # FINAL STRICT CHECK: Only add courses matching the filter
                        course_type_raw = course.course_type
                        course_type_check = (course_type_raw or '').strip().lower()
                        
                        # Apply course_type filter if specified
                        should_add = False
                        if course_type_filter:
                            should_add = (course_type_check == course_type_filter)
                        else:
                            # Default: Only theory courses (backward compatibility)
                            should_add = (course_type_check == 'theory')
                        
                        if should_add:
                            courses_data.append({
                                'course_id': course.id,
                                'course_code': course.course_code,
                                'course_name': course.course_name,
                                'section': section,
                                'teacher_id': course_assignment.teacher_id if course_assignment else None,
                                'teacher_name': teacher_name,
                                'designation': designation,
                                'institute': institute,
                                'designation_institute': f'{designation}, {institute}',
                                'course_type': course.course_type  # Include for frontend verification
                            })
                            current_app.logger.debug(f'✅ ADDED {course_type_check} course: {course.course_code} - {course.course_name}')
                        else:
                            current_app.logger.warning(f'🚫 FINAL FILTER: Rejected course: {course.course_code} - {course.course_name} (type: "{course_type_raw}" -> "{course_type_check}", filter: {course_type_filter or "theory"})')
            
            # Log summary with course types to verify filtering
            course_types_summary = {}
            for course in courses_data:
                # Get actual course_type from database for logging
                try:
                    from blueprints.course_management.models import Course
                    db_course = Course.query.get(course['course_id'])
                    if db_course:
                        course_type = db_course.course_type or 'Unknown'
                        course_types_summary[course_type] = course_types_summary.get(course_type, 0) + 1
                except:
                    pass
            
            current_app.logger.info(f'📊 FINAL SUMMARY: Returning {len(courses_data)} courses')
            if course_types_summary:
                for ctype, count in course_types_summary.items():
                    current_app.logger.info(f'   - {ctype}: {count} courses')
            else:
                current_app.logger.info(f'   - No courses found')
            
            # Final verification: ensure all courses match the requested course_type filter
            if course_type_filter:
                # When a specific course_type is requested, verify courses match that type
                mismatched_codes = []
                for course in courses_data:
                    try:
                        from blueprints.course_management.models import Course
                        db_course = Course.query.get(course['course_id'])
                        if db_course:
                            course_type = (db_course.course_type or '').strip().lower()
                            if course_type != course_type_filter:
                                mismatched_codes.append(f"{course['course_code']} (type: {db_course.course_type}, expected: {course_type_filter})")
                    except:
                        pass
                
                if mismatched_codes:
                    current_app.logger.warning(f'⚠️ Found {len(mismatched_codes)} courses that don\'t match filter {course_type_filter}: {", ".join(mismatched_codes)}')
                    # Filter them out as a safety measure
                    courses_data = [c for c in courses_data if c['course_code'] not in [nc.split(' (')[0] for nc in mismatched_codes]]
                    current_app.logger.info(f'✅ Filtered out mismatched courses. Now returning {len(courses_data)} {course_type_filter} courses only.')
            else:
                # Default: Only allow theory courses (backward compatibility)
                non_theory_codes = []
                for course in courses_data:
                    try:
                        from blueprints.course_management.models import Course
                        db_course = Course.query.get(course['course_id'])
                        if db_course:
                            course_type = (db_course.course_type or '').strip().lower()
                            if course_type != 'theory':
                                non_theory_codes.append(f"{course['course_code']} (type: {db_course.course_type})")
                    except:
                        pass
                
                if non_theory_codes:
                    current_app.logger.error(f'❌ ERROR: Found {len(non_theory_codes)} non-theory courses in response: {", ".join(non_theory_codes)}')
                    # Filter them out as a safety measure
                    courses_data = [c for c in courses_data if c['course_code'] not in [nc.split(' (')[0] for nc in non_theory_codes]]
                    current_app.logger.info(f'✅ Filtered out non-theory courses. Now returning {len(courses_data)} theory courses only.')
            
            return jsonify({
                'success': True,
                'courses': courses_data
            })
        except Exception as e:
            current_app.logger.error(f'Error fetching courses with assignments: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error fetching courses: {str(e)}'}), 500

    @app.route('/exam-committee-chief/get-student-count', methods=['GET'])
    @login_required
    def exam_committee_chief_get_student_count():
        """Get total registered students (by batch) for Statement of Remuneration.

        Statement টেবিল ৪/৬/৭–এ আমরা Section A/B আলাদা না করে,
        নির্দিষ্ট batch এবং course-এর জন্য মোট ছাত্র-সংখ্যা চাই।
        """
        course_code = request.args.get('course_code')
        academic_session = request.args.get('academic_session')
        year = request.args.get('year')
        term = request.args.get('term')
        batch = request.args.get('batch')

        if not course_code or not academic_session:
            return jsonify({'success': False, 'message': 'Course code and academic session are required'}), 400

        try:
            from blueprints.course_management.models import StudentCourseRegistration
            from blueprints.student_management.models import Student
            from sqlalchemy import func, or_, and_

            # Extract bare course code (e.g. "0421 28 Law 3201 - ..." → "0421 28 Law 3201")
            if ' - ' in course_code:
                parts = course_code.split(' - ', 1)
                course_code = parts[0].strip()
            course_code = (course_code or '').strip()

            def normalize_label(label):
                if not label:
                    return ''
                label = str(label).strip()
                for suffix in [' Year', ' Term', 'Year', 'Term']:
                    if label.lower().endswith(suffix.lower()):
                        label = label[:-len(suffix)].strip()
                return label

            normalized_year = normalize_label(year) if year else None
            normalized_term = normalize_label(term) if term else None

            # Build query:
            # - direct registration match on running context, OR
            # - relevant-course mapping match on relevant context.
            base_query = db.session.query(
                func.count(func.distinct(StudentCourseRegistration.student_id))
            ).filter(
                StudentCourseRegistration.status == 'finalized'
            )

            direct_clause = and_(
                StudentCourseRegistration.course_code == course_code,
                StudentCourseRegistration.academic_session == academic_session
            )
            relevant_clause = and_(
                StudentCourseRegistration.relevant_course_code == course_code,
                StudentCourseRegistration.relevant_academic_session == academic_session,
                StudentCourseRegistration.use_relevant_for_committee.is_(True)
            )

            if normalized_year:
                direct_clause = and_(direct_clause, StudentCourseRegistration.year == normalized_year)
                relevant_clause = and_(relevant_clause, StudentCourseRegistration.relevant_year == normalized_year)
            if normalized_term:
                direct_clause = and_(direct_clause, StudentCourseRegistration.term == normalized_term)
                relevant_clause = and_(relevant_clause, StudentCourseRegistration.relevant_term == normalized_term)

            base_query = base_query.filter(or_(direct_clause, relevant_clause))
            if batch:
                base_query = base_query.join(
                    Student,
                    Student.id == StudentCourseRegistration.student_id
                ).filter(Student.batch == batch)

            total_registered = base_query.scalar() or 0

            current_app.logger.info(
                f'[Chief student-count] course={course_code}, session={academic_session}, '
                f'year={normalized_year}, term={normalized_term}, batch={batch}, '
                f'total_registered={total_registered}'
            )

            return jsonify({'success': True, 'student_count': total_registered})

        except Exception as e:
            current_app.logger.error(f'Error fetching student count (chief): {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': f'Error fetching student count: {str(e)}'}), 500

    @app.route('/api/get-teacher-info/<int:teacher_id>', methods=['GET'])
    @login_required
    def get_teacher_info(teacher_id):
        """Get teacher designation and institute by teacher ID"""
        try:
            teacher = Teacher.query.get(teacher_id)
            if not teacher:
                return jsonify({'success': False, 'message': 'Teacher not found'}), 404
            
            return jsonify({
                'success': True,
                'teacher': {
                    'id': teacher.id,
                    'name': teacher.name,
                    'designation': teacher.designation or '',
                    'institute': teacher.institute or 'Law Discipline, KU'
                }
            })
        except Exception as e:
            current_app.logger.error(f'Error fetching teacher info: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Error fetching teacher information'}), 500

    @app.route('/exam-committee-member/remuneration')
    @login_required
    def exam_committee_member_remuneration():
        """Custom Remuneration Statement form for Exam Committee Internal Members"""
        # Check if current user is assigned as Exam Committee Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            flash('Teacher profile not found.', 'warning')
            return redirect(url_for('index'))
        
        # Get URL parameters for specific assignment or form_id
        session_param = request.args.get('session', '').strip()
        year_param = request.args.get('year', '').strip()
        term_param = request.args.get('term', '').strip()
        form_id_param = request.args.get('form_id', '').strip()
        
        # If form_id is provided, load saved form
        saved_data = None
        if form_id_param and form_id_param.isdigit():
            from blueprints.remuneration_management.models import RemunerationForm
            saved_form = RemunerationForm.query.filter_by(
                id=int(form_id_param),
                user_id=current_user.id
            ).first()
            
            if saved_form:
                # Verify this form belongs to a valid member assignment
                member_check = DutyAssignment.query.filter(
                    DutyAssignment.duty_type == 'exam_committee_member',
                    DutyAssignment.assigned_teacher_id == teacher.id,
                    DutyAssignment.academic_session == saved_form.academic_year,
                    DutyAssignment.year == saved_form.year,
                    DutyAssignment.term == saved_form.term,
                    DutyAssignment.status == 'active'
                ).first()
                
                if member_check:
                    try:
                        saved_data = json.loads(saved_form.form_data) if saved_form.form_data else None
                        session_param = saved_form.academic_year or session_param
                        year_param = saved_form.year or year_param
                        term_param = saved_form.term or term_param
                    except:
                        pass
        
        # Find member assignment based on parameters or get first one
        if session_param and year_param and term_param:
            member_assignment = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.assigned_teacher_id == teacher.id,
                DutyAssignment.academic_session == session_param,
                DutyAssignment.year == year_param,
                DutyAssignment.term == term_param,
                DutyAssignment.status == 'active'
            ).first()
        else:
            member_assignment = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.assigned_teacher_id == teacher.id,
                DutyAssignment.status == 'active'
            ).first()
        
        if not member_assignment:
            flash('You are not assigned as Exam Committee Member for the specified session/year/term.', 'danger')
            return redirect(url_for('exam_committee_member_dashboard'))
        
        # Get all teachers (excluding Head of the Discipline)
        from role_utils import get_teachers_excluding_head
        teachers = get_teachers_excluding_head()
        
        # Get academic sessions
        from blueprints.class_management.models import Session
        sessions = db.session.query(Session.academic_session).distinct().filter(
            Session.academic_session.isnot(None)
        ).order_by(Session.academic_session.desc()).all()
        academic_sessions = [s[0] for s in sessions if s[0]]
        
        current_session = member_assignment.academic_session if member_assignment else None
        current_year = member_assignment.year if member_assignment else None
        current_term = member_assignment.term if member_assignment else None
        
        # Get examination committee members (Chief + Internal + External)
        examination_committee = []
        
        # Add Chief (Chairman)
        if member_assignment.academic_session and member_assignment.year and member_assignment.term:
            chief_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                academic_session=member_assignment.academic_session,
                year=member_assignment.year,
                term=member_assignment.term,
                status='active'
            ).first()
            
            if chief_assignment and chief_assignment.assigned_teacher:
                chief_teacher = chief_assignment.assigned_teacher
                chief_designation = ''
                chief_institute = ''
                try:
                    chief_data = json.loads(chief_assignment.remarks) if chief_assignment.remarks else {}
                    if chief_data.get('type') == 'chief':
                        chief_designation = chief_data.get('designation', '')
                        chief_institute = chief_data.get('institute', '')
                except:
                    pass
                
                examination_committee.append({
                    'name': chief_teacher.name or '',
                    'designation': chief_designation or (chief_teacher.designation if chief_teacher.designation else '') or 'Head, Law Discipline, KU',
                    'institute': chief_institute or (chief_teacher.institute if chief_teacher.institute else '') or 'Law Discipline, KU',
                    'position': 'Chairman'
                })
            
            # Get committee members
            committee_assignments = DutyAssignment.query.filter(
                DutyAssignment.duty_type == 'exam_committee_member',
                DutyAssignment.academic_session == member_assignment.academic_session,
                DutyAssignment.year == member_assignment.year,
                DutyAssignment.term == member_assignment.term,
                DutyAssignment.status == 'active'
            ).all()
            
            for assignment in committee_assignments:
                if assignment.assigned_teacher_id:
                    # Internal member
                    member_teacher = assignment.assigned_teacher
                    if member_teacher:
                        member_designation = ''
                        member_institute = ''
                        try:
                            member_data = json.loads(assignment.remarks) if assignment.remarks else {}
                            if member_data.get('type') == 'internal':
                                member_designation = member_data.get('designation', '')
                                member_institute = member_data.get('institute', '')
                        except:
                            pass
                        
                        examination_committee.append({
                            'name': member_teacher.name or '',
                            'designation': member_designation or (member_teacher.designation if member_teacher.designation else '') or 'Assistant Professor',
                            'institute': member_institute or (member_teacher.institute if member_teacher.institute else '') or 'Law Discipline, KU',
                            'position': 'Member'
                        })
                else:
                    # External member (info stored in remarks as JSON)
                    try:
                        external_info = json.loads(assignment.remarks) if assignment.remarks else {}
                        if external_info.get('type') == 'external':
                            designation = external_info.get('designation', '') or ''
                            institute = external_info.get('institute', '') or ''
                            name = external_info.get('name', '') or ''
                            if name:  # Only add if name exists
                                examination_committee.append({
                                    'name': name,
                                    'designation': designation,
                                    'institute': institute,
                                    'position': 'Ext. Member'
                                })
                    except:
                        pass
        
        # Get all teachers and convert to JSON-serializable format
        teachers_data = []
        for teacher in teachers:
            teachers_data.append({
                'name': teacher.name or '',
                'designation': teacher.designation or '',
                'institute': teacher.institute or 'Law Discipline, KU'
            })
        
        # Get curriculum data for members (they can see all, but we'll filter based on their assignment)
        from blueprints.course_management.models import CurriculumYearTerm
        
        # Get curriculum configs for the member's session/year/term
        curriculum_configs = []
        if current_session and current_year and current_term:
            configs = CurriculumYearTerm.query.filter_by(
                academic_session=current_session,
                year=current_year,
                term=current_term
            ).all()
            
            for config in configs:
                curriculum_configs.append({
                    'academic_session': str(config.academic_session) if config.academic_session else '',
                    'year': str(config.year) if config.year else '',
                    'term': str(config.term) if config.term else '',
                    'batch': str(config.batch) if config.batch else ''
                })
        
        # For exam committee members, they don't have restrictions like chiefs
        # So allowed_combinations is empty (they can work with their assigned session/year/term)
        allowed_combinations = []
        if current_session and current_year and current_term:
            allowed_combinations = [{
                'session': str(current_session),
                'year': str(current_year),
                'term': str(current_term)
            }]
        
        # Get unique years and terms for dropdowns (based on member's session if available)
        unique_years = []
        unique_terms = []
        unique_batches = []
        if current_session:
            all_years = CurriculumYearTerm.query.filter_by(
                academic_session=current_session
            ).with_entities(CurriculumYearTerm.year).distinct().filter(
                CurriculumYearTerm.year.isnot(None)
            ).order_by(CurriculumYearTerm.year.asc()).all()
            unique_years = sorted(set([y[0] for y in all_years if y[0]]))
            
            all_terms = CurriculumYearTerm.query.filter_by(
                academic_session=current_session
            ).with_entities(CurriculumYearTerm.term).distinct().filter(
                CurriculumYearTerm.term.isnot(None)
            ).order_by(CurriculumYearTerm.term.asc()).all()
            unique_terms = sorted(set([t[0] for t in all_terms if t[0]]))
            
            all_batches = CurriculumYearTerm.query.filter_by(
                academic_session=current_session,
                year=current_year,
                term=current_term
            ).with_entities(CurriculumYearTerm.batch).distinct().filter(
                CurriculumYearTerm.batch.isnot(None)
            ).order_by(CurriculumYearTerm.batch.desc()).all()
            unique_batches = [b[0] for b in all_batches if b[0]]
        
        current_batch = None
        if current_session and current_year and current_term:
            config = CurriculumYearTerm.query.filter_by(
                academic_session=current_session,
                year=current_year,
                term=current_term
            ).first()
            if config:
                current_batch = config.batch
        
        return render_template('exam_committee_chief/custom_remuneration.html',
                             teachers=teachers,
                             teachers_data=teachers_data,  # JSON-serializable version
                             academic_sessions=academic_sessions,
                             unique_years=unique_years,
                             unique_terms=unique_terms,
                             unique_batches=unique_batches,
                             curriculum_configs=curriculum_configs,
                             current_session=current_session,
                             current_year=current_year,
                             current_term=current_term,
                             current_batch=current_batch,
                             examination_committee=examination_committee,
                             saved_data=saved_data,  # Load saved form data if form_id provided
                             allowed_combinations=allowed_combinations)

    @app.route('/api/get-teacher-by-name', methods=['GET'])
    @login_required
    def get_teacher_by_name():
        """Get teacher designation and institute by teacher name"""
        try:
            name = request.args.get('name', '').strip()
            if not name:
                return jsonify({'success': False, 'message': 'Name is required'}), 400
            
            teacher = Teacher.query.filter_by(name=name).first()
            if not teacher:
                return jsonify({'success': False, 'message': 'Teacher not found'}), 404
            
            return jsonify({
                'success': True,
                'teacher': {
                    'id': teacher.id,
                    'name': teacher.name,
                    'designation': teacher.designation or '',
                    'institute': teacher.institute or 'Law Discipline, KU'
                }
            })
        except Exception as e:
            current_app.logger.error(f'Error fetching teacher by name: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Error fetching teacher information'}), 500

    @app.route('/exam-committee-chief/get-committee-members', methods=['GET'])
    @login_required
    def get_committee_members():
        """Get examination committee members for a given session/year/term"""
        # Check if current user is assigned as Exam Committee Chief or Member
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not teacher:
            return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
        
        academic_session = request.args.get('session', '').strip()
        year = request.args.get('year', '').strip()
        term = request.args.get('term', '').strip()
        
        if not academic_session or not year or not term:
            return jsonify({'success': False, 'message': 'Session, Year, and Term are required'}), 400

        chief_assignments, member_assignments = _get_exam_committee_assignments(
            teacher,
            academic_session=academic_session,
            year=year,
            term=term
        )
        if not chief_assignments and not member_assignments:
            return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief/Member for this session/year/term'}), 403

        chief_assignment = DutyAssignment.query.filter_by(
            duty_type='exam_committee_chief',
            academic_session=academic_session,
            year=year,
            term=term,
            status='active'
        ).first()
        
        # Check if semester is active
        try:
            from utils.semester_utils import is_semester_active
            if not is_semester_active(academic_session, year, term, batch=None):
                # Semester is inactive, return empty committee
                return jsonify({
                    'success': True,
                    'committee': [],
                    'message': f'This semester ({academic_session} - {year} - {term}) is not active.'
                })
        except Exception as e:
            current_app.logger.error(f'Error checking active semester: {e}', exc_info=True)
            # Continue anyway
        
        examination_committee = []
        
        # Add Chief (Chairman) - check if this session/year/term matches the chief's assignment
        if chief_assignment and (chief_assignment.academic_session == academic_session and 
            chief_assignment.year == year and 
            chief_assignment.term == term and
            chief_assignment.assigned_teacher):
            chief_teacher = chief_assignment.assigned_teacher
            chief_designation = ''
            chief_institute = ''
            
            # Try to get from chief assignment remarks
            try:
                chief_data = json.loads(chief_assignment.remarks) if chief_assignment.remarks else {}
                if chief_data.get('type') == 'chief':
                    chief_designation = chief_data.get('designation', '')
                    chief_institute = chief_data.get('institute', '')
            except:
                pass
            
            # Fall back to Teacher table
            if not chief_designation and chief_teacher.designation:
                chief_designation = chief_teacher.designation
            if not chief_institute and chief_teacher.institute:
                chief_institute = chief_teacher.institute
            
            # Final defaults
            if not chief_designation:
                chief_designation = 'Head, Law Discipline, KU'
            if not chief_institute:
                chief_institute = 'Law Discipline, KU'
            
            examination_committee.append({
                'name': chief_teacher.name,
                'designation': f'{chief_designation}, {chief_institute}',
                'position': 'Chairman'
            })
        
        # Get committee members
        committee_assignments = DutyAssignment.query.filter(
            DutyAssignment.duty_type == 'exam_committee_member',
            DutyAssignment.academic_session == academic_session,
            DutyAssignment.year == year,
            DutyAssignment.term == term,
            DutyAssignment.status == 'active'
        ).all()
        
        for assignment in committee_assignments:
            if assignment.assigned_teacher_id:
                # Internal member
                teacher_member = assignment.assigned_teacher
                if teacher_member:
                    member_designation = ''
                    member_institute = ''
                    
                    # Try to get from assignment remarks
                    try:
                        member_data = json.loads(assignment.remarks) if assignment.remarks else {}
                        if member_data.get('type') == 'internal':
                            member_designation = member_data.get('designation', '')
                            member_institute = member_data.get('institute', '')
                    except:
                        pass
                    
                    # Fall back to Teacher table
                    if not member_designation and teacher_member.designation:
                        member_designation = teacher_member.designation
                    if not member_institute and teacher_member.institute:
                        member_institute = teacher_member.institute
                    
                    # Final defaults
                    if not member_designation:
                        member_designation = 'Assistant Professor'
                    if not member_institute:
                        member_institute = 'Law Discipline, KU'
                    
                    examination_committee.append({
                        'name': teacher_member.name,
                        'designation': f'{member_designation}, {member_institute}',
                        'position': 'Member'
                    })
            else:
                # External member (info stored in remarks as JSON)
                try:
                    external_info = json.loads(assignment.remarks) if assignment.remarks else {}
                    if external_info.get('type') == 'external':
                        designation = f"{external_info.get('designation', '')}, {external_info.get('institute', '')}"
                        examination_committee.append({
                            'name': external_info.get('name', ''),
                            'designation': designation.strip(', '),
                            'position': 'Ext. Member'
                        })
                except:
                    pass
        
        return jsonify({
            'success': True,
            'committee': examination_committee
        })

    @app.route('/officer/exam-info')
    @login_required
    def officer_exam_info():
        """Exam Info page for officers"""
        roles = parse_roles(current_user.role)
        if 'officer' not in roles and not is_admin(current_user):
            flash('This page is only accessible to officers.', 'danger')
            return redirect(url_for('index'))
        
        return render_template('officer/exam_info.html')

    @app.route('/remuneration')
    @login_required
    def remuneration_portal():
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        # Initialize default values for graceful degradation
        teachers = []
        all_configs = []
        academic_sessions = []
        all_years = []
        all_terms = []
        years_terms_map = {}
        saved_data = None
        has_warning = False
        
        try:
            # Fetch teachers for name dropdown (excluding Head of the Discipline)
            try:
                from blueprints.class_management.models import Teacher
                from role_utils import get_teachers_excluding_head
                teachers = get_teachers_excluding_head()
            except Exception as e:
                current_app.logger.error(f'Error fetching teachers: {e}', exc_info=True)
                teachers = []
                has_warning = True
            
            # Fetch all curriculum year/term configurations and aggregate data
            try:
                all_configs = CurriculumYearTerm.query.all()
            except Exception as e:
                current_app.logger.error(f'Error fetching curriculum configs: {e}', exc_info=True)
                all_configs = []
                has_warning = True
            
            # Collect all unique academic sessions
            try:
                academic_sessions = sorted(list(set(
                    config.academic_session for config in all_configs 
                    if config.academic_session
                )))
            except Exception as e:
                current_app.logger.error(f'Error processing academic sessions: {e}', exc_info=True)
                academic_sessions = []
                has_warning = True
            
            # Collect all unique years
            try:
                all_years = sorted(list(set(
                    config.year for config in all_configs
                )))
            except Exception as e:
                current_app.logger.error(f'Error processing years: {e}', exc_info=True)
                all_years = []
                has_warning = True
            
            # Collect all unique terms
            try:
                all_terms = sorted(list(set(
                    config.term for config in all_configs
                    if config.term
                )))
            except Exception as e:
                current_app.logger.error(f'Error processing terms: {e}', exc_info=True)
                all_terms = []
                has_warning = True
            
            # Build year-term mapping (all curricula combined)
            try:
                for config in all_configs:
                    year = config.year
                    term = config.term
                    if year not in years_terms_map:
                        years_terms_map[year] = []
                    if term not in years_terms_map[year]:
                        years_terms_map[year].append(term)
                
                # Sort terms for each year
                for year in years_terms_map:
                    years_terms_map[year] = sorted(years_terms_map[year])
            except Exception as e:
                current_app.logger.error(f'Error building year-term mapping: {e}', exc_info=True)
                years_terms_map = {}
                has_warning = True
            
            # Load saved form data from session if exists
            try:
                saved_data = session.get('remuneration_form_data', None)
            except Exception as e:
                current_app.logger.error(f'Error loading saved form data: {e}', exc_info=True)
                saved_data = None
                has_warning = True
            
        except Exception as e:
            # Catch any unexpected errors
            current_app.logger.error(f'Unexpected error in remuneration_portal: {e}', exc_info=True)
            has_warning = True
        
        # Show warning if any errors occurred
        if has_warning:
            flash('Some data could not be loaded. Please refresh the page if needed.', 'warning')
        
        # Define remuneration rates based on the scanned document
        remuneration_rates = {
            '1': [  # প্রশ্নপত্র প্রণয়ন
                {'label': 'স্নাতক (প্রতি প্রশ্নপত্র)', 'value': '2300'},
                {'label': 'স্নাতকোত্তর/এমফিল/পিএইচডি (প্রতি প্রশ্নপত্র)', 'value': '2400'}
            ],
            '2': [  # প্রশ্নপত্র মডারেশন
                {'label': 'সর্বোচ্চ', 'value': '2400'},
                {'label': 'সর্বনিম্ন', 'value': '1500'}
            ],
            '3': [  # উত্তরপত্র পরীক্ষণ
                {'label': 'স্নাতক - অর্ধ পত্র (প্রতি উত্তরপত্র)', 'value': '80'},
                {'label': 'স্নাতক - ক্লাস টেস্ট/টার্ম পেপার (প্রতি পরীক্ষার্থী)', 'value': '30'},
                {'label': 'স্নাতকোত্তর - অর্ধ পত্র (প্রতি উত্তরপত্র)', 'value': '100'},
                {'label': 'স্নাতকোত্তর - ক্লাস টেস্ট/টার্ম পেপার (প্রতি পরীক্ষার্থী)', 'value': '80'},
                {'label': 'স্নাতকোত্তর - মিড টার্ম পূর্ণপত্র (প্রতি উত্তরপত্র)', 'value': '60'},
                {'label': 'স্নাতকোত্তর - টার্ম ফাইনাল পূর্ণপত্র (প্রতি উত্তরপত্র)', 'value': '160'},
                {'label': 'ন্যূনতম (প্রতি কোর্স)', 'value': '600'}
            ],
            '4': [  # ক্লাস টেস্ট/টার্ম পেপার/ হোম ওয়ার্ক/ এ্যাসাইনমেন্ট
                {'label': 'স্নাতক - ক্লাস টেস্ট/টার্ম পেপার (প্রতি পরীক্ষার্থী)', 'value': '30'},
                {'label': 'স্নাতকোত্তর - ক্লাস টেস্ট/টার্ম পেপার (প্রতি পরীক্ষার্থী)', 'value': '40'}
            ],
            '5': [  # সেশনাল
                {'label': 'প্রজেক্ট পেপার/এ্যাসাইনমেন্ট (প্রতি পরীক্ষার্থী)', 'value': '230'},
                {'label': 'ফিল্ড ওয়ার্ক/সার্ভে ওয়ার্ক (প্রতি পরীক্ষার্থী)', 'value': '300'},
                {'label': 'মৌখিক পরীক্ষা (প্রতি পরীক্ষার্থী)', 'value': '50'},
                {'label': 'ল্যাব কর্মকর্তা', 'value': '200'},
                {'label': '৩য় শ্রেণির কর্মচারী', 'value': '150'},
                {'label': '৪র্থ শ্রেণির কর্মচারী', 'value': '110'}
            ],
            '6': [  # সেশনাল মৌখিক পরীক্ষা
                {'label': 'প্রতি পরীক্ষার্থী', 'value': '50'}
            ],
            '7': [  # প্রফেশনাল এ্যাটাসমেন্ট/ইন্ডাস্ট্রিয়াল
                {'label': 'সুপারভিশন ও রিপোর্ট পরীক্ষণ (প্রতি পরীক্ষার্থী)', 'value': '100'}
            ],
            '8': [  # উত্তরপত্র নিরীক্ষণ
                {'label': 'প্রতি উত্তরপত্র', 'value': '8'}
            ],
            '9': [  # টেবুলেশন
                {'label': 'কোর্স ভিত্তিক (প্রতি কোর্স)', 'value': '200'},
                {'label': 'পরীক্ষার্থী ভিত্তিক (প্রতি পরীক্ষার্থী)', 'value': '40'}
            ],
            '9a': [  # টেবুলেশন - কোর্স ভিত্তিক
                {'label': 'কোর্স ভিত্তিক (প্রতি কোর্স)', 'value': '200'}
            ],
            '9b': [  # টেবুলেশন - পরীক্ষার্থী ভিত্তিক
                {'label': 'পরীক্ষার্থী ভিত্তিক (প্রতি পরীক্ষার্থী)', 'value': '40'}
            ],
            '10': [  # প্রশ্নপত্র প্রস্তুতকরণ
                {'label': 'অংকনসহ অন্যান্য কাজ (প্রতি প্রশ্নপত্র)', 'value': '250'},
                {'label': 'ফটোকপি (প্রতি প্রশ্নপত্র)', 'value': '7'}
            ],
            '10a': [  # প্রশ্নপত্র প্রস্তুতকরণ - অংকন
                {'label': 'অংকনসহ অন্যান্য কাজ (প্রতি প্রশ্নপত্র)', 'value': '250'}
            ],
            '10b': [  # প্রশ্নপত্র প্রস্তুতকরণ - ফটোকপি
                {'label': 'ফটোকপি (প্রতি প্রশ্নপত্র)', 'value': '7'}
            ],
            '11': [  # পরীক্ষা কমিটির সভাপতি/সদস্য
                {'label': 'স্নাতক - সভাপতি (প্রতি টার্ম)', 'value': '2500'},
                {'label': 'স্নাতক - সদস্য (প্রতি টার্ম)', 'value': '1000'},
                {'label': 'স্নাতকোত্তর - সভাপতি (প্রতি টার্ম)', 'value': '3000'},
                {'label': 'স্নাতকোত্তর - সদস্য (প্রতি টার্ম)', 'value': '1000'}
            ],
            '12': [  # চীফ ইনভিজিলেশন / ইনভিজিলেশন
                {'label': 'প্রধান তদারকী (প্রতি ঘন্টা)', 'value': '600'},
                {'label': 'অন্যান্য তদারকী (প্রতি ঘন্টা)', 'value': '500'}
            ],
            '12a': [  # চীফ ইনভিজিলেশন
                {'label': 'চীফ ইনভিজিলেশন', 'value': '1800'}
            ],
            '12b': [  # ইনভিজিলেশন
                {'label': 'ইনভিজিলেশন', 'value': '1500'}
            ],
            '15': [  # কোডিং/ডিকোডিং
                {'label': 'পরীক্ষার্থী প্রতি', 'value': '30'}
            ],
            '13': [  # থিসিস
                # পরীক্ষণ
                {'label': 'স্নাতক - থিসিস/প্রজেক্ট মূল্যায়ন (প্রতি পরীক্ষার্থী)', 'value': '1200'},
                {'label': 'স্নাতকোত্তর - ডিজারটেশন মূল্যায়ন (প্রতি পরীক্ষার্থী)', 'value': '2500'},
                {'label': 'পিএইচডি - ডিজারটেশন মূল্যায়ন (প্রতি পরীক্ষার্থী)', 'value': '10000'}
            ],
            '13a': [  # থিসিস - পরীক্ষণ
                {'label': 'স্নাতক - থিসিস/প্রজেক্ট মূল্যায়ন (প্রতি পরীক্ষার্থী)', 'value': '1200'},
                {'label': 'স্নাতকোত্তর - ডিজারটেশন মূল্যায়ন (প্রতি পরীক্ষার্থী)', 'value': '2500'},
                {'label': 'পিএইচডি - ডিজারটেশন মূল্যায়ন (প্রতি পরীক্ষার্থী)', 'value': '10000'}
            ],
            '13b': [  # থিসিস - সুপারভিশন
                {'label': 'স্নাতক - থিসিস/প্রজেক্ট সুপারভিশন (প্রতি পরীক্ষার্থী)', 'value': '2000'},
                {'label': 'স্নাতকোত্তর - ডিজারটেশন সুপারভিশন (প্রতি পরীক্ষার্থী)', 'value': '5000'},
                {'label': 'স্নাতকোত্তর - প্রজেক্ট সুপারভিশন (প্রতি পরীক্ষার্থী)', 'value': '2500'},
                {'label': 'পিএইচডি - ডিজারটেশন সুপারভিশন (প্রতি পরীক্ষার্থী)', 'value': '35000'}
            ],
            '13c': [  # থিসিস - কো-সুপারভিশন
                {'label': 'স্নাতকোত্তর - কো-সুপারভিশন (প্রতি পরীক্ষার্থী)', 'value': '1500'},
                {'label': 'পিএইচডি - কো-সুপারভিশন (প্রতি পরীক্ষার্থী)', 'value': '15000'}
            ],
            '13d': [  # থিসিস - মৌখিক পরীক্ষা
                {'label': 'স্নাতক - ফাইনাল ডিফেন্স/মৌখিক (প্রতি পরীক্ষার্থী)', 'value': '120'},
                {'label': 'স্নাতকোত্তর - ফাইনাল ডিফেন্স/মৌখিক (প্রতি পরীক্ষার্থী)', 'value': '500'},
                {'label': 'পিএইচডি - ফাইনাল ডিফেন্স/মৌখিক (প্রতি পরীক্ষার্থী)', 'value': '2000'}
            ],
            '14': [  # ভাইভা
                {'label': 'পরীক্ষার্থী প্রতি ৫০টাকা', 'value': '50'}
            ],
            '16': [  # কোডিং/ডিকোডিং
                {'label': 'প্রতি খাতা', 'value': '30'}
            ]
        }
        
        # Set logo path directly - always use Images folder
        logo_filename = 'KU_logo_2.png'
        logo_folder = 'Images'  # Direct path - file should be at static/Images/KU_logo_2.png
        
        return render_template('remuneration_placeholder.html', 
                             teachers=teachers, 
                             academic_sessions=academic_sessions,
                             all_years=all_years,
                             all_terms=all_terms,
                             years_terms_map=years_terms_map,
                             saved_data=saved_data,
                             remuneration_rates=remuneration_rates,
                             logo_filename=logo_filename,
                             logo_folder=logo_folder)

    @app.route('/remuneration/api/save', methods=['POST'])
    @login_required
    def remuneration_save():
        """Save form data to database"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            data = request.get_json()
            if not data:
                return jsonify({'success': False, 'message': 'No data provided'}), 400
            
            form_id = data.get('form_id')  # For editing existing form
            title = data.get('title', '')  # Optional title
            
            # Extract main fields
            form_entry = None
            if form_id:
                # Edit existing form
                form_entry = RemunerationForm.query.filter_by(id=form_id, user_id=current_user.id).first()
                if not form_entry:
                    return jsonify({'success': False, 'message': 'Form not found'}), 404
            else:
                # Create new form
                form_entry = RemunerationForm(
                    user_id=current_user.id,
                    status='draft'
                )
                db.session.add(form_entry)
            
            # Auto-generate title if not provided
            if not title or title.strip() == '':
                # Build title using Year, Term, and Session
                title_parts = []
                
                # Get Year, Term, and Session
                year = data.get('year', '').strip()
                term = data.get('term', '').strip()
                session = data.get('academic_year', '').strip()
                
                # Add Year if available
                if year:
                    title_parts.append(f"Year: {year}")
                
                # Add Term if available
                if term:
                    title_parts.append(f"Term: {term}")
                
                # Add Session if available
                if session:
                    title_parts.append(f"Session: {session}")
                
                # Generate title
                if title_parts:
                    title = ' - '.join(title_parts)
                else:
                    applicant_name = data.get('applicant_name', '')
                    title = f"Remuneration Form - {applicant_name or 'Untitled'}"
            
            # Update form fields
            form_entry.title = title
            form_entry.applicant_name = data.get('applicant_name', '')
            form_entry.designation = data.get('designation', '')
            form_entry.address = data.get('address', '')
            form_entry.discipline = data.get('discipline', '')
            form_entry.exam_discipline = data.get('exam_discipline', '')
            form_entry.year = data.get('year', '')
            form_entry.term = data.get('term', '')
            form_entry.academic_year = data.get('academic_year', '')
            form_entry.exam_start_date = data.get('exam_start_date', '')
            form_entry.exam_end_date = data.get('exam_end_date', '')
            form_entry.voucher_no = data.get('voucher_no', '')
            form_entry.voucher_date = data.get('voucher_date', '')
            form_entry.total_amount = float(data.get('total_amount', 0) or 0)
            form_entry.total_in_words = data.get('total_in_words', '')
            form_entry.bank_account = data.get('bank_account', '')
            form_entry.bank_advice_no = data.get('bank_advice_no', '')
            form_entry.payment_date = data.get('payment_date', '')
            
            # Save all form data as JSON
            form_entry.form_data = json.dumps(data)
            
            db.session.commit()
            
            return jsonify({
                'success': True, 
                'message': 'Form data saved successfully',
                'form_id': form_entry.id
            })
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error saving form data: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to save form data'}), 500

    @app.route('/remuneration/api/load', methods=['GET'])
    @login_required
    def remuneration_load():
        """Load saved form data from database"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            form_id = request.args.get('form_id')
            academic_year = request.args.get('academic_year') or request.args.get('session', '').strip()
            year = request.args.get('year', '').strip()
            term = request.args.get('term', '').strip()
            
            if form_id:
                # Load specific form
                form_entry = RemunerationForm.query.filter_by(id=form_id, user_id=current_user.id).first()
                if not form_entry:
                    return jsonify({'success': False, 'message': 'Form not found'}), 404
                
                # Load from JSON data or reconstruct from fields
                if form_entry.form_data:
                    try:
                        saved_data = json.loads(form_entry.form_data)
                    except:
                        saved_data = form_entry.to_dict()
                else:
                    saved_data = form_entry.to_dict()
                
                return jsonify({'success': True, 'data': saved_data})
            elif academic_year and year and term:
                # Load form by session/year/term for current user
                form_entry = RemunerationForm.query.filter_by(
                    user_id=current_user.id,
                    academic_year=academic_year,
                    year=year,
                    term=term
                ).order_by(RemunerationForm.id.desc()).first()
                
                if form_entry and form_entry.form_data:
                    try:
                        saved_data = json.loads(form_entry.form_data)
                        saved_data['form_id'] = form_entry.id
                        # Include exam dates from database fields if not in JSON
                        if form_entry.exam_start_date:
                            saved_data['exam_start_date'] = form_entry.exam_start_date
                        if form_entry.exam_end_date:
                            saved_data['exam_end_date'] = form_entry.exam_end_date
                        return jsonify({
                            'success': True,
                            'data': saved_data,
                            'message': 'Form data loaded successfully'
                        })
                    except Exception as e:
                        current_app.logger.error(f'Error parsing form_data JSON: {str(e)}')
                        return jsonify({'success': False, 'message': 'Error loading form data'}), 500
                else:
                    return jsonify({
                        'success': True,
                        'data': None,
                        'message': 'No saved form found for this session/year/term'
                    })
            else:
                # Load from session as fallback
                saved_data = session.get('remuneration_form_data', None)
                if saved_data:
                    return jsonify({'success': True, 'data': saved_data})
                else:
                    return jsonify({'success': False, 'message': 'No saved data found'})
        except Exception as e:
            current_app.logger.error(f'Error loading form data: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to load form data'}), 500

    @app.route('/remuneration/api/clear', methods=['POST'])
    @login_required
    def remuneration_clear():
        """Clear saved form data from session"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            session.pop('remuneration_form_data', None)
            return jsonify({'success': True, 'message': 'Saved data cleared successfully'})
        except Exception as e:
            current_app.logger.error(f'Error clearing form data: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to clear form data'}), 500

    @app.route('/remuneration/list')
    @login_required
    def remuneration_list():
        """List all saved remuneration forms"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            status_filter = request.args.get('status', 'all')  # 'all', 'draft', 'archived'
            
            query = RemunerationForm.query.filter_by(user_id=current_user.id)
            
            if status_filter == 'draft':
                query = query.filter_by(status='draft')
            elif status_filter == 'archived':
                query = query.filter_by(status='archived')
            
            forms = query.order_by(RemunerationForm.created_at.desc()).all()
            
            forms_list = []
            for f in forms:
                # Extract session and course names from form_data if available
                academic_session = f.academic_year or ''
                course_names = []
                
                # Try to get course names from form_data JSON
                if f.form_data:
                    try:
                        form_data_json = json.loads(f.form_data)
                        course_sections = form_data_json.get('course_sections', {})
                        if course_sections:
                            all_course_names = []
                            for row_id, pairs in course_sections.items():
                                if isinstance(pairs, list):
                                    for pair in pairs:
                                        course = pair.get('course', '') if isinstance(pair, dict) else ''
                                        if course and ' - ' in course:
                                            course_name = course.split(' - ', 1)[1].strip()
                                            if course_name:
                                                all_course_names.append(course_name)
                            # Remove duplicates
                            seen = set()
                            for name in all_course_names:
                                if name and name not in seen:
                                    course_names.append(name)
                                    seen.add(name)
                    except:
                        pass
                
                forms_list.append({
                    'id': f.id,
                    'title': f.title or f'Remuneration Form #{f.id}',
                    'course_names': ', '.join(course_names) if course_names else '',
                    'applicant_name': f.applicant_name,
                    'status': f.status,
                    'academic_session': academic_session,
                    'total_amount': f.total_amount,
                    'created_at': f.created_at.strftime('%Y-%m-%d %H:%M') if f.created_at else '',
                    'updated_at': f.updated_at.strftime('%Y-%m-%d %H:%M') if f.updated_at else ''
                })
            
            return render_template('remuneration_list.html', forms=forms_list, status_filter=status_filter)
        except Exception as e:
            current_app.logger.error(f'Error listing forms: {str(e)}', exc_info=True)
            flash('Error loading forms list', 'error')
            return redirect(url_for('remuneration_portal'))

    @app.route('/remuneration/edit/<int:form_id>')
    @login_required
    def remuneration_edit(form_id):
        """Edit an existing remuneration form"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            form_entry = RemunerationForm.query.filter_by(id=form_id, user_id=current_user.id).first()
            if not form_entry:
                flash('Form not found', 'error')
                return redirect(url_for('remuneration_list'))
            
            # Load form data
            if form_entry.form_data:
                try:
                    saved_data = json.loads(form_entry.form_data)
                except:
                    saved_data = form_entry.to_dict()
            else:
                saved_data = form_entry.to_dict()
            
            # Ensure Year, Term, and Session are included in saved_data
            if not saved_data.get('year') and form_entry.year:
                saved_data['year'] = form_entry.year
            if not saved_data.get('term') and form_entry.term:
                saved_data['term'] = form_entry.term
            if not saved_data.get('academic_year') and form_entry.academic_year:
                saved_data['academic_year'] = form_entry.academic_year
            
            saved_data['form_id'] = form_entry.id
            
            # Fetch teachers for name dropdown (excluding Head and deleted accounts)
            from role_utils import get_teachers_excluding_head
            teachers = get_teachers_excluding_head()
            
            # Fetch all curriculum year/term configurations and aggregate data
            all_configs = CurriculumYearTerm.query.all()
            
            # Collect all unique academic sessions
            academic_sessions = sorted(list(set(
                config.academic_session for config in all_configs 
                if config.academic_session
            )))
            
            # Collect all unique years
            all_years = sorted(list(set(
                config.year for config in all_configs
            )))
            
            # Collect all unique terms
            all_terms = sorted(list(set(
                config.term for config in all_configs
                if config.term
            )))
            
            # Build year-term mapping (all curricula combined)
            years_terms_map = {}
            for config in all_configs:
                year = config.year
                term = config.term
                if year not in years_terms_map:
                    years_terms_map[year] = []
                if term not in years_terms_map[year]:
                    years_terms_map[year].append(term)
            
            # Sort terms for each year
            for year in years_terms_map:
                years_terms_map[year] = sorted(years_terms_map[year])
            
            # Define remuneration rates (same as in main portal)
            remuneration_rates = {
                '1': [{'label': 'স্নাতক (প্রতি প্রশ্নপত্র)', 'value': '2300'}, {'label': 'স্নাতকোত্তর/এমফিল/পিএইচডি (প্রতি প্রশ্নপত্র)', 'value': '2400'}],
                '2': [{'label': 'সর্বোচ্চ', 'value': '2400'}],
                '3': [{'label': 'প্রতি খাতা', 'value': '400'}],
                '4': [{'label': 'প্রতি ছাত্র', 'value': '50'}],
                '5': [{'label': 'প্রতি কোর্স', 'value': '3000'}],
                '6': [{'label': 'প্রতি কোর্স', 'value': '3000'}],
                '7': [{'label': 'কাস্টম হার', 'value': ''}],
                '8': [{'label': 'প্রতি খাতা', 'value': '100'}],
                '9': [{'label': 'প্রতি কোর্স', 'value': '2500'}],
                '10': [{'label': 'প্রতি প্রশ্নপত্র', 'value': '500'}],
                '10a': [{'label': 'অংকনসহ অন্যান্য কাজ (প্রতি প্রশ্নপত্র)', 'value': '250'}],
                '10b': [{'label': 'ফটোকপি (প্রতি প্রশ্নপত্র)', 'value': '7'}],
                '11': [{'label': 'সভাপতি', 'value': '5000'}, {'label': 'সদস্য', 'value': '3000'}],
                '12': [{'label': 'চীফ ইনভিজিলেশন', 'value': '1800'}, {'label': 'ইনভিজিলেশন', 'value': '1500'}],
                '12a': [{'label': 'চীফ ইনভিজিলেশন', 'value': '1800'}],
                '12b': [{'label': 'ইনভিজিলেশন', 'value': '1500'}],
                '15': [{'label': 'পরীক্ষার্থী প্রতি', 'value': '30'}],
                '13': [{'label': 'কাস্টম হার', 'value': ''}],
                '13a': [{'label': 'কাস্টম হার', 'value': ''}],
                '13b': [{'label': 'কাস্টম হার', 'value': ''}],
                '13c': [{'label': 'কাস্টম হার', 'value': ''}],
                '13d': [{'label': 'কাস্টম হার', 'value': ''}],
                '14': [{'label': 'পরীক্ষার্থী প্রতি ৫০টাকা', 'value': '50'}],
                '16': [{'label': 'প্রতি খাতা', 'value': '30'}]
            }
            
            logo_filename = 'KU_logo_2.png'
            logo_folder = 'Images'
            
            return render_template('remuneration_placeholder.html', 
                                 teachers=teachers, 
                                 academic_sessions=academic_sessions,
                                 all_years=all_years,
                                 all_terms=all_terms,
                                 years_terms_map=years_terms_map,
                                 saved_data=saved_data,
                                 remuneration_rates=remuneration_rates,
                                 logo_filename=logo_filename,
                                 logo_folder=logo_folder)
        except Exception as e:
            current_app.logger.error(f'Error loading form for edit: {str(e)}', exc_info=True)
            flash('Error loading form', 'error')
            return redirect(url_for('remuneration_list'))

    @app.route('/remuneration/api/archive/<int:form_id>', methods=['POST'])
    @login_required
    def remuneration_archive(form_id):
        """Archive a remuneration form"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            form_entry = RemunerationForm.query.filter_by(id=form_id, user_id=current_user.id).first()
            if not form_entry:
                return jsonify({'success': False, 'message': 'Form not found'}), 404
            
            form_entry.status = 'archived'
            form_entry.archived_at = datetime.utcnow()
            db.session.commit()
            
            return jsonify({'success': True, 'message': 'Form archived successfully'})
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error archiving form: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to archive form'}), 500

    @app.route('/remuneration/api/unarchive/<int:form_id>', methods=['POST'])
    @login_required
    def remuneration_unarchive(form_id):
        """Unarchive a remuneration form"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            form_entry = RemunerationForm.query.filter_by(id=form_id, user_id=current_user.id).first()
            if not form_entry:
                return jsonify({'success': False, 'message': 'Form not found'}), 404
            
            form_entry.status = 'draft'
            form_entry.archived_at = None
            db.session.commit()
            
            return jsonify({'success': True, 'message': 'Form unarchived successfully'})
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error unarchiving form: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to unarchive form'}), 500

    @app.route('/remuneration/api/delete/<int:form_id>', methods=['POST'])
    @login_required
    def remuneration_delete(form_id):
        """Delete a remuneration form"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            form_entry = RemunerationForm.query.filter_by(id=form_id, user_id=current_user.id).first()
            if not form_entry:
                return jsonify({'success': False, 'message': 'Form not found'}), 404
            
            db.session.delete(form_entry)
            db.session.commit()
            
            return jsonify({'success': True, 'message': 'Form deleted successfully'})
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error deleting form: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to delete form'}), 500

    @app.route('/remuneration/api/courses', methods=['GET'])
    @login_required
    def remuneration_get_courses():
        """Get courses based on academic session, year, and term"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        academic_session = request.args.get('academic_session')
        year = request.args.get('year')
        term = request.args.get('term')
        
        if not all([academic_session, year, term]):
            return jsonify({'success': False, 'message': 'Academic session, year, and term are required'}), 400
        
        try:
            from sqlalchemy import or_

            # Find curriculum year/term configs matching the criteria
            try:
                configs = CurriculumYearTerm.query.filter_by(
                    academic_session=academic_session,
                    year=year,
                    term=term
                ).all()
            except Exception as e:
                current_app.logger.error(f'Error querying CurriculumYearTerm: {str(e)}', exc_info=True)
                return jsonify({'success': False, 'message': 'Database error while fetching curriculum configs'}), 500

            all_courses = []
            if configs:
                # Get all curriculum IDs
                curriculum_ids = [config.curriculum_id for config in configs]

                # Get all courses for these curricula
                try:
                    all_courses = Course.query.filter(
                        Course.curriculum_id.in_(curriculum_ids)
                    ).order_by(Course.course_code).all()
                except Exception as e:
                    current_app.logger.error(f'Error querying Course: {str(e)}', exc_info=True)
                    return jsonify({'success': False, 'message': 'Database error while fetching courses'}), 500
            
            # Filter courses by year/term (check both direct match and display_year/display_term)
            # Normalize year/term for comparison (remove 'Year'/'Term' suffix if present)
            def normalize_label(label):
                if not label:
                    return ''
                label = str(label).strip()
                for suffix in [' Year', ' Term', 'Year', 'Term']:
                    if label.lower().endswith(suffix.lower()):
                        label = label[:-len(suffix)].strip()
                return label
            
            normalized_year = normalize_label(year)
            normalized_term = normalize_label(term)
            
            matching_courses = []
            for course in all_courses:
                # Check direct match first
                direct_year_match = (course.year and normalize_label(course.year) == normalized_year) or (not course.year and normalize_label(course.display_year) == normalized_year)
                direct_term_match = (course.term and normalize_label(course.term) == normalized_term) or (not course.term and normalize_label(course.display_term) == normalized_term)
                
                # Also check display_year and display_term
                display_year_match = normalize_label(course.display_year) == normalized_year
                display_term_match = normalize_label(course.display_term) == normalized_term
                
                # Match if either direct or display matches
                if (direct_year_match or display_year_match) and (direct_term_match or display_term_match):
                    matching_courses.append(course)

            # Add separate (merge OFF) retake courses that belong to this relevant context.
            retake_remarks = {'retake', 're-retake', 're retake', 'reretake'}
            separate_retake_regs = StudentCourseRegistration.query.filter(
                StudentCourseRegistration.status.in_(['finalized', 'approved']),
                or_(
                    StudentCourseRegistration.use_relevant_for_committee.is_(False),
                    StudentCourseRegistration.use_relevant_for_committee.is_(None)
                ),
                StudentCourseRegistration.course_code.isnot(None),
                StudentCourseRegistration.course_code != ''
            ).all()

            separate_retake_entries = {}
            for reg in separate_retake_regs:
                remark_value = str(reg.remark or '').strip().lower()
                if remark_value not in retake_remarks:
                    continue
                reg_session = str(reg.relevant_academic_session or '').strip().casefold()
                if reg_session != str(academic_session or '').strip().casefold():
                    continue
                reg_year = normalize_label(reg.relevant_year)
                reg_term = normalize_label(reg.relevant_term)
                if reg_year != normalized_year or reg_term != normalized_term:
                    continue

                course_code = (reg.course_code or '').strip()
                if not course_code:
                    continue

                candidate_course = Course.query.get(reg.course_id) if reg.course_id else None
                if not candidate_course:
                    candidate_course = Course.query.filter_by(course_code=course_code).order_by(Course.id.desc()).first()

                if candidate_course:
                    course_type = candidate_course.course_type or ''
                    base_name = candidate_course.course_name or reg.course_name or course_code
                else:
                    course_type = reg.course_type or ''
                    base_name = reg.course_name or course_code

                retake_key = f"{course_code.lower()}::{base_name.strip().lower()}"
                if retake_key in separate_retake_entries:
                    continue
                separate_retake_entries[retake_key] = {
                    'id': reg.course_id or None,
                    'course_code': course_code,
                    'course_name': f"{base_name.strip()} (Retake)",
                    'course_type': course_type
                }
            
            courses_data = [{
                'id': course.id,
                'course_code': course.course_code,
                'course_name': course.course_name,
                'course_type': course.course_type or ''
            } for course in matching_courses]

            # Keep separate retake subjects as distinct options in remuneration applicable lists.
            courses_data.extend(separate_retake_entries.values())
            
            return jsonify({'success': True, 'courses': courses_data})
            
        except Exception as e:
            current_app.logger.error(f'Error fetching courses: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to fetch courses'}), 500

    @app.route('/remuneration/api/course-assignments', methods=['GET'])
    @login_required
    def remuneration_get_course_assignments():
        """Get course assignments with teacher and section info for Table 4 (Class Test)"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        academic_session = request.args.get('academic_session')
        year = request.args.get('year')
        term = request.args.get('term')
        
        if not all([academic_session, year, term]):
            return jsonify({'success': False, 'message': 'Academic session, year, and term are required'}), 400
        
        try:
            from blueprints.course_management.models import CourseSessionAssignment, Course, CurriculumYearTerm
            
            # Find curriculum year/term configs matching the criteria
            configs = CurriculumYearTerm.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            ).all()
            
            if not configs:
                return jsonify({'success': True, 'assignments': []})
            
            # Get all curriculum IDs
            curriculum_ids = [config.curriculum_id for config in configs]
            
            # Get course assignments for these curricula
            assignments = CourseSessionAssignment.query.filter(
                CourseSessionAssignment.academic_session == academic_session,
                CourseSessionAssignment.year == year,
                CourseSessionAssignment.term == term,
                CourseSessionAssignment.curriculum_id.in_(curriculum_ids)
            ).join(Course).filter(
                db.or_(
                    Course.course_type == 'Theory',
                    Course.course_type == 'theory',
                    Course.course_type.is_(None),
                    Course.course_type == ''
                )  # Only theory courses (or assume theory if not specified)
            ).all()
            
            # Group assignments by course to handle split courses (A and B count as 1)
            assignments_by_course = {}
            for assignment in assignments:
                course_id = assignment.course_id
                if course_id not in assignments_by_course:
                    assignments_by_course[course_id] = []
                assignments_by_course[course_id].append(assignment)
            
            assignments_data = []
            for course_id, course_assignments in assignments_by_course.items():
                course = course_assignments[0].course
                
                # Check if this is a split course (has both A and B sections)
                sections = [a.section for a in course_assignments if a.section]
                has_a = 'A' in sections
                has_b = 'B' in sections
                has_full = any(a.section is None or a.section == '' for a in course_assignments)
                
                # If split course (A and B), create one entry per section
                if has_a and has_b:
                    # Section A
                    assignment_a = next((a for a in course_assignments if a.section == 'A'), None)
                    if assignment_a and assignment_a.teacher:
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'A',
                            'teacher_name': assignment_a.teacher.name,
                            'teacher_designation': assignment_a.teacher.designation or '',
                            'teacher_institute': assignment_a.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
                    
                    # Section B
                    assignment_b = next((a for a in course_assignments if a.section == 'B'), None)
                    if assignment_b and assignment_b.teacher:
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'B',
                            'teacher_name': assignment_b.teacher.name,
                            'teacher_designation': assignment_b.teacher.designation or '',
                            'teacher_institute': assignment_b.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
                elif has_full:
                    # Full course
                    assignment_full = next((a for a in course_assignments if a.section is None or a.section == ''), None)
                    if assignment_full and assignment_full.teacher:
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'Full',
                            'teacher_name': assignment_full.teacher.name,
                            'teacher_designation': assignment_full.teacher.designation or '',
                            'teacher_institute': assignment_full.teacher.institute or 'Law Discipline, KU',
                            'is_split': False
                        })
                else:
                    # Single section (A or B only)
                    for assignment in course_assignments:
                        if assignment.teacher:
                            assignments_data.append({
                                'course_id': course.id,
                                'course_code': course.course_code,
                                'course_name': course.course_name,
                                'section': assignment.section or 'Full',
                                'teacher_name': assignment.teacher.name,
                                'teacher_designation': assignment.teacher.designation or '',
                                'teacher_institute': assignment.teacher.institute or 'Law Discipline, KU',
                                'is_split': False
                            })
            
            return jsonify({'success': True, 'assignments': assignments_data})
            
        except Exception as e:
            current_app.logger.error(f'Error fetching course assignments: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to fetch course assignments'}), 500

    @app.route('/remuneration/api/sessional-course-assignments', methods=['GET'])
    @login_required
    def remuneration_get_sessional_course_assignments():
        """Get sessional course assignments with teacher and section info for Table 6 and 7"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        academic_session = request.args.get('academic_session')
        year = request.args.get('year')
        term = request.args.get('term')
        
        if not all([academic_session, year, term]):
            return jsonify({'success': False, 'message': 'Academic session, year, and term are required'}), 400
        
        try:
            from blueprints.course_management.models import CourseSessionAssignment, Course, CurriculumYearTerm
            
            # Find curriculum year/term configs matching the criteria
            configs = CurriculumYearTerm.query.filter_by(
                academic_session=academic_session,
                year=year,
                term=term
            ).all()
            
            if not configs:
                return jsonify({'success': True, 'assignments': []})
            
            # Get all curriculum IDs
            curriculum_ids = [config.curriculum_id for config in configs]
            
            # Get course assignments for sessional courses
            assignments = CourseSessionAssignment.query.filter(
                CourseSessionAssignment.academic_session == academic_session,
                CourseSessionAssignment.year == year,
                CourseSessionAssignment.term == term,
                CourseSessionAssignment.curriculum_id.in_(curriculum_ids)
            ).join(Course).filter(
                db.or_(
                    Course.course_type == 'Sessional',
                    Course.course_type == 'sessional',
                    db.func.lower(Course.course_type).contains('sessional')
                )  # Only sessional courses
            ).all()
            
            # Group assignments by course to handle split courses (A and B count as 1)
            assignments_by_course = {}
            for assignment in assignments:
                course_id = assignment.course_id
                if course_id not in assignments_by_course:
                    assignments_by_course[course_id] = []
                assignments_by_course[course_id].append(assignment)
            
            assignments_data = []
            for course_id, course_assignments in assignments_by_course.items():
                course = course_assignments[0].course
                
                # Check if this is a split course (has both A and B sections)
                sections = [a.section for a in course_assignments if a.section]
                has_a = 'A' in sections
                has_b = 'B' in sections
                has_full = any(a.section is None or a.section == '' for a in course_assignments)
                
                # If split course (A and B), create one entry per section
                if has_a and has_b:
                    # Section A
                    assignment_a = next((a for a in course_assignments if a.section == 'A'), None)
                    if assignment_a and assignment_a.teacher:
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'A',
                            'teacher_name': assignment_a.teacher.name,
                            'teacher_designation': assignment_a.teacher.designation or '',
                            'teacher_institute': assignment_a.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
                    
                    # Section B
                    assignment_b = next((a for a in course_assignments if a.section == 'B'), None)
                    if assignment_b and assignment_b.teacher:
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'B',
                            'teacher_name': assignment_b.teacher.name,
                            'teacher_designation': assignment_b.teacher.designation or '',
                            'teacher_institute': assignment_b.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
                elif has_full:
                    # Full course - create A and B entries
                    assignment_full = next((a for a in course_assignments if a.section is None or a.section == ''), None)
                    if assignment_full and assignment_full.teacher:
                        # Create A section
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'A',
                            'teacher_name': assignment_full.teacher.name,
                            'teacher_designation': assignment_full.teacher.designation or '',
                            'teacher_institute': assignment_full.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
                        # Create B section
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'B',
                            'teacher_name': assignment_full.teacher.name,
                            'teacher_designation': assignment_full.teacher.designation or '',
                            'teacher_institute': assignment_full.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
                else:
                    # Single section (A or B only) - create both A and B with same teacher
                    assignment = course_assignments[0]
                    if assignment.teacher:
                        # Create A section
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'A',
                            'teacher_name': assignment.teacher.name,
                            'teacher_designation': assignment.teacher.designation or '',
                            'teacher_institute': assignment.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
                        # Create B section
                        assignments_data.append({
                            'course_id': course.id,
                            'course_code': course.course_code,
                            'course_name': course.course_name,
                            'section': 'B',
                            'teacher_name': assignment.teacher.name,
                            'teacher_designation': assignment.teacher.designation or '',
                            'teacher_institute': assignment.teacher.institute or 'Law Discipline, KU',
                            'is_split': True
                        })
            
            return jsonify({'success': True, 'assignments': assignments_data})
            
        except Exception as e:
            current_app.logger.error(f'Error fetching sessional course assignments: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to fetch sessional course assignments'}), 500

    @app.route('/remuneration/api/tabulators', methods=['GET'])
    @login_required
    def remuneration_get_tabulators():
        """Get assigned tabulators for Table 9 (Tabulation)"""
        import json
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        academic_session = request.args.get('academic_session')
        year = request.args.get('year')
        term = request.args.get('term')
        
        if not all([academic_session, year, term]):
            return jsonify({'success': False, 'message': 'Academic session, year, and term are required'}), 400
        
        try:
            from blueprints.course_management.models import DutyAssignment
            from blueprints.class_management.models import Teacher
            
            # Get current user's teacher record
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if not teacher:
                return jsonify({'success': False, 'message': 'Teacher profile not found'}), 404
            
            # Check if current user is Exam Committee Chief or Internal Member
            chief_assignments, member_assignments = _get_exam_committee_assignments(teacher)
            if not chief_assignments and not member_assignments:
                return jsonify({'success': False, 'message': 'You are not assigned as Exam Committee Chief or Member'}), 403
            
            # Get tabulators assigned by this Exam Committee Chief for the given session/year/term
            # Simplified query - use filter_by for exact matches
            tabulator_assignments = DutyAssignment.query.filter_by(
                duty_type='tabulator',
                assigned_by_id=current_user.id,
                status='active',
                academic_session=academic_session,
                year=year,
                term=term
            ).all()
            
            tabulators_data = []
            for assignment in tabulator_assignments:
                if assignment.assigned_teacher:
                    tabulators_data.append({
                        'name': assignment.assigned_teacher.name,
                        'designation': assignment.assigned_teacher.designation or '',
                        'institute': assignment.assigned_teacher.institute or 'Law Discipline, KU'
                    })
            
            return jsonify({'success': True, 'tabulators': tabulators_data})
            
        except Exception as e:
            current_app.logger.error(f'Error fetching tabulators: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to fetch tabulators'}), 500

    @app.route('/remuneration/api/student-count', methods=['GET'])
    @login_required
    def remuneration_get_student_count():
        """Get total number of students for a course assigned to current teacher's account"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        course_code = request.args.get('course_code')
        academic_session = request.args.get('academic_session')
        year = request.args.get('year')
        term = request.args.get('term')
        batch = request.args.get('batch')  # Optional: restrict to specific batch
        section = request.args.get('section', '').strip().upper()  # A, B, Full, or empty
        
        if not course_code or not academic_session:
            return jsonify({'success': False, 'message': 'Course code and academic session are required'}), 400
        
        # Map section to course_scope for Session filtering
        course_scope = None
        if section == 'A':
            course_scope = 'part_a'
        elif section == 'B':
            course_scope = 'part_b'
        elif section == 'FULL' or section == '':
            course_scope = 'full'
        # If section is something else, treat as 'full'
        if course_scope is None:
            course_scope = 'full'
        
        try:
            # Import models
            from blueprints.course_management.models import StudentCourseRegistration
            from blueprints.student_management.models import Student
            from blueprints.class_management.models import Teacher, Session, ClassStudent
            from sqlalchemy import func, or_, and_
            import re
            
            # Extract course code only (remove course name if present in format "CODE - Name")
            original_course_code = course_code
            course_name = None
            if ' - ' in course_code:
                parts = course_code.split(' - ', 1)
                course_code = parts[0].strip()
                if len(parts) > 1:
                    course_name = parts[1].strip()
            
            # Clean course code
            course_code = course_code.strip()
            
            # Normalize year/term labels
            def normalize_label(label):
                if not label:
                    return ''
                label = str(label).strip()
                for suffix in [' Year', ' Term', 'Year', 'Term']:
                    if label.lower().endswith(suffix.lower()):
                        label = label[:-len(suffix)].strip()
                return label
            
            normalized_year = normalize_label(year) if year else None
            normalized_term = normalize_label(term) if term else None
            
            # PRIMARY: Count registered students from StudentCourseRegistration
            # Filter by direct original-course context OR relevant-course mapping context.
            # and (optionally) student.batch when batch is provided.
            base_query = db.session.query(
                func.count(func.distinct(StudentCourseRegistration.student_id))
            ).filter(
                StudentCourseRegistration.status == 'finalized'
            )

            direct_clause = and_(
                StudentCourseRegistration.course_code == course_code,
                StudentCourseRegistration.academic_session == academic_session
            )
            relevant_clause = and_(
                StudentCourseRegistration.relevant_course_code == course_code,
                StudentCourseRegistration.relevant_academic_session == academic_session,
                StudentCourseRegistration.use_relevant_for_committee.is_(True)
            )

            if batch:
                base_query = base_query.join(
                    Student,
                    Student.id == StudentCourseRegistration.student_id
                ).filter(Student.batch == batch)
            
            # Add year filter if provided
            if normalized_year:
                direct_clause = and_(direct_clause, StudentCourseRegistration.year == normalized_year)
                relevant_clause = and_(relevant_clause, StudentCourseRegistration.relevant_year == normalized_year)
            
            # Add term filter if provided
            if normalized_term:
                direct_clause = and_(direct_clause, StudentCourseRegistration.term == normalized_term)
                relevant_clause = and_(relevant_clause, StudentCourseRegistration.relevant_term == normalized_term)

            base_query = base_query.filter(or_(direct_clause, relevant_clause))
            
            # Get total registered count
            total_registered_count = base_query.scalar() or 0
            
            current_app.logger.info(
                f'Registered students for course_code={course_code}, session={academic_session}, '
                f'year={normalized_year}, term={normalized_term}, total={total_registered_count}'
            )
            
            # If section is specified (A or B), try to filter by Session course_scope
            if section in ['A', 'B'] and total_registered_count > 0:
                # Find Sessions matching this course and section
                session_query = Session.query.filter(
                    Session.course_code == course_code,
                    Session.academic_session == academic_session,
                    Session.course_scope == course_scope
                )
                
                if normalized_year:
                    session_query = session_query.filter(Session.year == normalized_year)
                if normalized_term:
                    session_query = session_query.filter(Session.term == normalized_term)
                
                matched_sessions = session_query.all()
                
                if matched_sessions:
                    # Count students in these specific sessions
                    session_ids = [s.id for s in matched_sessions]
                    section_count = db.session.query(
                        func.count(func.distinct(ClassStudent.student_id))
                    ).filter(
                        ClassStudent.session_id.in_(session_ids)
                    ).scalar() or 0
                    
                    current_app.logger.info(
                        f'Section {section} (course_scope={course_scope}): Found {len(matched_sessions)} sessions, '
                        f'{section_count} students in sessions'
                    )
                    
                    # Use section count if available, otherwise use total registered count
                    if section_count > 0:
                        effective_count = max(section_count, total_registered_count)
                        return jsonify({
                            'success': True,
                            'student_count': effective_count
                        })
                    else:
                        # If no students in sessions but we have registered students,
                        # return total registered (sessions might not be set up yet)
                        current_app.logger.info(
                            f'No students in sessions for section {section}, using total registered count'
                        )
                        return jsonify({
                            'success': True,
                            'student_count': total_registered_count
                        })
                else:
                    # No sessions found for this section, return total registered
                    current_app.logger.info(
                        f'No sessions found for section {section}, using total registered count'
                    )
                    return jsonify({
                        'success': True,
                        'student_count': total_registered_count
                    })
            
            # For Full section or no section specified, return total registered count
            if total_registered_count > 0:
                return jsonify({
                    'success': True,
                    'student_count': total_registered_count
                })
            
            # Fallback: Try flexible matching if exact match didn't work
            # Try partial course code match
            if total_registered_count == 0:
                current_app.logger.info(
                    f'No exact match found. Trying flexible matching...'
                )
                
                # Try matching by course name if available
                if course_name:
                    name_query = db.session.query(
                        func.count(func.distinct(StudentCourseRegistration.student_id))
                    ).filter(
                        StudentCourseRegistration.course_name.like(f'%{course_name}%'),
                        StudentCourseRegistration.academic_session == academic_session,
                        StudentCourseRegistration.status == 'finalized'
                    )

                    if batch:
                        name_query = name_query.join(
                            Student,
                            Student.id == StudentCourseRegistration.student_id
                        ).filter(Student.batch == batch)
                    
                    if normalized_year:
                        name_query = name_query.filter(StudentCourseRegistration.year == normalized_year)
                    if normalized_term:
                        name_query = name_query.filter(StudentCourseRegistration.term == normalized_term)
                    
                    count = name_query.scalar() or 0
                    
                    if count > 0:
                        current_app.logger.info(
                            f'Found {count} students by course name match'
                        )
                        return jsonify({
                            'success': True,
                            'student_count': count
                        })
                
                # Try without year/term filters (just session and course code)
                fallback_query = db.session.query(
                    func.count(func.distinct(StudentCourseRegistration.student_id))
                ).filter(
                    StudentCourseRegistration.course_code == course_code,
                    StudentCourseRegistration.academic_session == academic_session,
                    StudentCourseRegistration.status == 'finalized'
                )

                if batch:
                    fallback_query = fallback_query.join(
                        Student,
                        Student.id == StudentCourseRegistration.student_id
                    ).filter(Student.batch == batch)
                
                count = fallback_query.scalar() or 0
                
                if count > 0:
                    current_app.logger.info(
                        f'Found {count} students without year/term filter'
                    )
                    return jsonify({
                        'success': True,
                        'student_count': count
                    })
            
            # Final log
            current_app.logger.info(
                f'FINAL: code="{course_code}", name="{course_name}", '
                f'session="{academic_session}", year="{normalized_year}", term="{normalized_term}", '
                f'section="{section}", FINAL COUNT={total_registered_count}'
            )
            
            # Return count (even if 0, to indicate no registered students found)
            return jsonify({
                'success': True,
                'student_count': total_registered_count
            })
            
        except Exception as e:
            current_app.logger.error(f'Error fetching student count: {str(e)}', exc_info=True)
            return jsonify({'success': False, 'message': 'Failed to fetch student count'}), 500

    @app.route('/remuneration/export-docx', methods=['POST'])
    @login_required
    def remuneration_export_docx():
        """Export remuneration form to DOCX in legal size (8.5" x 14")"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_SECTION
            from docx.oxml.ns import qn
            
            # Create document
            doc = Document()
            
            # Set page size to Legal (8.5" x 14")
            section = doc.sections[0]
            section.page_height = Cm(35.56)  # 14 inches
            section.page_width = Cm(21.59)  # 8.5 inches
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            
            # Get form data
            data = request.form.to_dict()
            
            # Helper function to get value or empty string
            def get_val(key, default=''):
                return data.get(key, default) or default
            
            # Header
            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.add_run('পরীক্ষা নিয়ন্ত্রকের কার্যালয়')
            run.font.size = Pt(10)
            run.font.name = 'Noto Sans Bengali'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
            
            heading2 = doc.add_paragraph()
            heading2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = heading2.add_run('খুলনা বিশ্ববিদ্যালয়')
            run2.font.size = Pt(12)
            run2.font.bold = True
            run2.font.name = 'Noto Sans Bengali'
            run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
            
            heading3 = doc.add_paragraph()
            heading3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run3 = heading3.add_run('পরীক্ষা পারিতোষিক বিল ফরম')
            run3.font.size = Pt(11)
            run3.font.bold = True
            run3.font.name = 'Noto Sans Bengali'
            run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
            
            note = doc.add_paragraph()
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_note = note.add_run('(প্রতি বর্ষের প্রতি টার্মের জন্য পৃথক বিল ফরম ব্যবহার করতে হবে)')
            run_note.font.size = Pt(8)
            run_note.font.name = 'Noto Sans Bengali'
            run_note._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
            
            # Voucher box (right aligned)
            voucher_para = doc.add_paragraph()
            voucher_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            voucher_text = f'ভাউচার নং: {get_val("voucher_no")}\nতারিখ: {get_val("voucher_date")}'
            run_v = voucher_para.add_run(voucher_text)
            run_v.font.size = Pt(9)
            run_v.font.name = 'Noto Sans Bengali'
            run_v._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
            
            # Personal details table
            details_table = doc.add_table(rows=6, cols=4)
            details_table.style = 'Table Grid'
            
            details_data = [
                ('নাম:', get_val('applicant_name'), '', ''),
                ('পদবী:', get_val('designation'), 'ডিসিপ্লিন / বিভাগ:', get_val('discipline')),
                ('ঠিকানা:', get_val('address'), '', ''),
                ('যে ডিসিপ্লিনের পরীক্ষা:', get_val('exam_discipline'), 'বর্ষ:', get_val('year')),
                ('শিক্ষাবর্ষ:', get_val('academic_year'), 'টার্ম:', get_val('term')),
                ('পরীক্ষা অনুষ্ঠানের তারিখ:', f'{get_val("exam_start_date")} থেকে', f'{get_val("exam_end_date")} পর্যন্ত', '')
            ]
            
            for i, row_data in enumerate(details_data):
                row = details_table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Noto Sans Bengali'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                            run.font.size = Pt(9)
            
            # Work details section
            work_heading = doc.add_paragraph()
            work_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_work = work_heading.add_run('পরীক্ষা সংক্রান্ত কাজের বিবরণ')
            run_work.font.size = Pt(10)
            run_work.font.bold = True
            run_work.font.name = 'Noto Sans Bengali'
            run_work._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
            
            # Work details table
            work_table = doc.add_table(rows=1, cols=7)
            work_table.style = 'Table Grid'
            
            # Header row
            headers = ['ক্রমিক নম্বর', 'বিবরণ', 'কোর্স নম্বর', 'প্রশ্ন/খাতা/ছাত্র/কোর্স পরীক্ষক/দিনের সংখ্যা', 'অর্ধ/ পূর্ণপত্র', 'পারিতোষিকের হার', 'মোট টাকা']
            header_row = work_table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
                        run.font.name = 'Noto Sans Bengali'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
            
            # Job rows
            jobs = [
                'প্রশ্নপত্র প্রণয়ন',
                'প্রশ্নপত্র মডারেশন',
                'উত্তরপত্র পরীক্ষণ',
                'ক্লাস টেস্ট/টার্ম পেপার/ হোম ওয়ার্ক/ এ্যাসাইনমেন্ট',
                'সেশনাল',
                'সেশনাল মৌখিক পরীক্ষা',
                'প্রফেশনাল এ্যাটাসমেন্ট/ইন্ডাস্ট্রিয়াল (ট্রেনিং/এ্যাটাসমেন্ট)',
                'উত্তরপত্র নিরীক্ষণ',
                'টেবুলেশন',
                'প্রশ্নপত্র প্রস্তুতকরণ (অংকন, স্টেনসিল কাটা ও ঘুরানো)',
                'পরীক্ষা কমিটির সভাপতি/সদস্য',
                'চীফ ইনভিজিলেশন / ইনভিজিলেশন',
                'থিসিস',
                'অন্যান্য'
            ]
            
            for idx, job in enumerate(jobs, 1):
                row = work_table.add_row()
                row.cells[0].text = str(idx)
                row.cells[1].text = job
                row.cells[2].text = get_val(f'course_no_{idx}')
                row.cells[3].text = get_val(f'quantity_{idx}')
                row.cells[4].text = get_val(f'paper_type_{idx}')
                row.cells[5].text = get_val(f'rate_{idx}')
                row.cells[6].text = get_val(f'amount_{idx}')
                
                # Format cells
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Noto Sans Bengali'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                            run.font.size = Pt(8)
                
                # Center align serial number
                row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add thesis sub-items
                if idx == 13:
                    sub_items = [
                        ('পরীক্ষণ', '13a'),
                        ('সুপারভিশন (থিসিস/প্রজেক্ট থিসিস/ইন্টার্নশীপ রিপোর্ট)', '13b'),
                        ('কো-সুপারভিশন', '13c'),
                        ('মৌখিক পরীক্ষা', '13d')
                    ]
                    for sub_job, sub_key in sub_items:
                        sub_row = work_table.add_row()
                        sub_row.cells[0].text = ''
                        sub_row.cells[1].text = sub_job
                        sub_row.cells[2].text = get_val(f'course_no_{sub_key}')
                        sub_row.cells[3].text = get_val(f'quantity_{sub_key}')
                        sub_row.cells[4].text = get_val(f'paper_type_{sub_key}')
                        sub_row.cells[5].text = get_val(f'rate_{sub_key}')
                        sub_row.cells[6].text = get_val(f'amount_{sub_key}')
                        
                        for cell in sub_row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Noto Sans Bengali'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                                    run.font.size = Pt(8)
            
            # Total row
            total_row = work_table.add_row()
            total_row.cells[0].text = ''
            total_row.cells[1].text = ''
            total_row.cells[2].text = ''
            total_row.cells[3].text = ''
            total_row.cells[4].text = ''
            total_row.cells[5].text = 'সর্বমোট টাকার পরিমাণ'
            total_row.cells[6].text = get_val('total_amount')
            
            for cell in total_row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = 'Noto Sans Bengali'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                        run.font.size = Pt(8)
            
            # Total in words
            total_words_para = doc.add_paragraph()
            total_words_para.add_run(f'সর্বমোট টাকার পরিমাণ (কথায়): {get_val("total_in_words")}')
            for run in total_words_para.runs:
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(9)
            
            # Bank declaration
            bank_para = doc.add_paragraph()
            bank_text = f'এই বিলের প্রাপ্য অর্থ অগ্রণী ব্যাংক, খুলনা বিশ্ববিদ্যালয় শাখায় আমার নামে রক্ষিত নং {get_val("bank_account")} হিসাবে/চেকের মাধ্যমে পরিশোধের অনুরোধ করছি এবং এই মর্মে অঙ্গীকার করছি যে, এই বিলে আমি কোন অতিরিক্ত অর্থ দাবী করিনি। যদি ভবিষ্যতে এই বিলে কোন আপত্তি উত্থাপিত হয় তাহলে গৃহীত অতিরিক্ত অর্থ ফেরৎ দিতে বাধ্য থাকব।'
            bank_para.add_run(bank_text)
            for run in bank_para.runs:
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(8)
            
            revenue_note = doc.add_paragraph()
            revenue_note.add_run('বিঃ দ্রঃ- প্রত্যেক বিলে রাজস্ব টিকিট লাগাতে হবে।')
            for run in revenue_note.runs:
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(8)
            
            # Signatures
            sig_table = doc.add_table(rows=1, cols=3)
            sig_table.style = 'Table Grid'
            sig_headers = ['ডিসিপ্লিন প্রধান\n(স্বাক্ষর ও সিল)', 'সভাপতি, পরীক্ষা কমিটি\n(স্বাক্ষর ও সিল)', 'প্রাপকের স্বাক্ষর ও তারিখ']
            for i, sig_text in enumerate(sig_headers):
                cell = sig_table.rows[0].cells[i]
                cell.text = sig_text
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Noto Sans Bengali'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                        run.font.size = Pt(9)
            
            # Exam Controller section
            controller_heading = doc.add_paragraph()
            controller_heading.add_run('পরীক্ষা নিয়ন্ত্রকের কার্যালয়ের ব্যবহারের জন্য')
            for run in controller_heading.runs:
                run.font.bold = True
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(9)
            
            controller_note = doc.add_paragraph()
            controller_note.add_run('পরীক্ষার পারিতোষিকের হার এবং ডিসিপ্লিন থেকে প্রাপ্ত স্টেটমেন্ট অনুযায়ী বিলসমূহ নিরীক্ষান্তে বিলের অর্থ পরিশোধের জন্য সুপারিশ করা হলো।')
            for run in controller_note.runs:
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(8)
            
            controller_table = doc.add_table(rows=3, cols=3)
            controller_table.style = 'Table Grid'
            controller_data = [
                ('বিল নিরীক্ষক/সেকশন অফিসার/সহকারী পরীক্ষা নিয়ন্ত্রক', get_val('auditor_sign'), get_val('auditor_date')),
                ('উপ-পরীক্ষা নিয়ন্ত্রক', get_val('deputy_sign'), get_val('deputy_date')),
                ('পরীক্ষা নিয়ন্ত্রক', get_val('controller_sign'), get_val('controller_date'))
            ]
            
            for i, row_data in enumerate(controller_data):
                row = controller_table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Noto Sans Bengali'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                            run.font.size = Pt(8)
            
            # Finance section
            finance_heading = doc.add_paragraph()
            finance_heading.add_run('অর্থ ও হিসাব বিভাগের ব্যবহারের জন্য')
            for run in finance_heading.runs:
                run.font.bold = True
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(9)
            
            finance_note = doc.add_paragraph()
            finance_note.add_run(f'পরীক্ষান্তে বর্ণিত পারিতোষিক বিল বাবদ {get_val("finance_amount_words")} কথায়: ({get_val("finance_amount")}) মাত্র পরিশোধের জন্য ছাড়া হলো।')
            for run in finance_note.runs:
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(8)
            
            finance_table = doc.add_table(rows=3, cols=3)
            finance_table.style = 'Table Grid'
            finance_data = [
                ('সেকশন অফিসার/সহকারী পরিচালক', get_val('section_officer_sign'), get_val('section_officer_date')),
                ('উপ-পরিচালক', get_val('deputy_director_sign'), get_val('deputy_director_date')),
                ('পরিচালক', get_val('director_sign'), get_val('director_date'))
            ]
            
            for i, row_data in enumerate(finance_data):
                row = finance_table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Noto Sans Bengali'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                            run.font.size = Pt(8)
            
            audit_note1 = doc.add_paragraph()
            audit_note1.add_run('এই বিল পরিশোধে কোন আপত্তি নেই নিরীক্ষান্তে')
            for run in audit_note1.runs:
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(8)
            
            audit_note2 = doc.add_paragraph()
            audit_note2.add_run('টাকার বিলটি পরিশোধের সুপারিশ করা হলো।')
            for run in audit_note2.runs:
                run.font.name = 'Noto Sans Bengali'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                run.font.size = Pt(8)
            
            audit_table = doc.add_table(rows=2, cols=3)
            audit_table.style = 'Table Grid'
            audit_data = [
                ('সহকারী পরিচালক (অডিট)', get_val('audit_assistant_sign'), get_val('audit_assistant_date')),
                ('উপ-পরিচালক/ প্রধান (অডিট সেল)', get_val('audit_head_sign'), get_val('audit_head_date'))
            ]
            
            for i, row_data in enumerate(audit_data):
                row = audit_table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Noto Sans Bengali'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                            run.font.size = Pt(8)
            
            # Bank advice
            bank_advice_table = doc.add_table(rows=2, cols=2)
            bank_advice_table.style = 'Table Grid'
            bank_advice_data = [
                ('ব্যাংক এ্যাডভাইস/ চেক নং:-', get_val('bank_advice_no')),
                ('তারিখ:', get_val('payment_date'))
            ]
            
            for i, row_data in enumerate(bank_advice_data):
                row = bank_advice_table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Noto Sans Bengali'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Bengali')
                            run.font.size = Pt(8)
            
            # Save to BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return send_file(
                buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='Exam_Remuneration_Form.docx'
            )
            
        except Exception as e:
            current_app.logger.error(f'Error generating DOCX: {str(e)}', exc_info=True)
            return jsonify({'error': 'Failed to generate document'}), 500

    @app.route('/static/Fonts/kalpurush.ttf')
    @app.route('/static/fonts/kalpurush.ttf')
    def serve_kalpurush_font():
        """Serve Kalpurush font file for WeasyPrint PDF generation"""
        from flask import send_file
        import os
        
        font_paths = [
            os.path.join(current_app.root_path, 'static', 'Fonts', 'kalpurush.ttf'),
            os.path.join(current_app.root_path, 'static', 'fonts', 'kalpurush.ttf'),
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                return send_file(font_path, mimetype='font/ttf')
        
        return 'Font not found', 404

    @app.route('/remuneration/export-pdf', methods=['POST'])
    @login_required
    def remuneration_export_pdf():
        """Export remuneration form to PDF in legal size (8.5" x 14") using WeasyPrint"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        try:
            # Check if WeasyPrint is available
            try:
                from weasyprint import HTML, CSS
            except ImportError as import_err:
                current_app.logger.error(f'WeasyPrint import error: {str(import_err)}', exc_info=True)
                return jsonify({
                    'error': 'PDF generation not available',
                    'message': f'WeasyPrint is not installed or dependencies are missing: {str(import_err)}'
                }), 500
            except Exception as weasy_err:
                current_app.logger.error(f'WeasyPrint initialization error: {str(weasy_err)}', exc_info=True)
                return jsonify({
                    'error': 'PDF generation failed',
                    'message': f'WeasyPrint initialization error: {str(weasy_err)}'
                }), 500
            
            import os
            
            # Get form data
            data = request.form.to_dict()
            
            # Helper function to get value or empty string
            def get_val(key, default=''):
                return data.get(key, default) or default
            
            # Helper function to get course-section pairs for a row
            def get_course_sections(row_id):
                """Get course-section pairs for a given row (without duplicates)"""
                courses = request.form.getlist(f'course_no_{row_id}[]')
                sections = request.form.getlist(f'section_{row_id}[]')
                # Use a set to track unique course-section pairs
                seen_pairs = set()
                unique_pairs = []
                
                for i, course in enumerate(courses):
                    section = sections[i] if i < len(sections) else ''
                    if course:
                        # Extract course code only (in case format is "CODE - Name")
                        course_code = course.split(' - ')[0].strip() if ' - ' in course else course.strip()
                        section_label = section if section else 'Full'
                        pair_key = (course_code, section_label)
                        # Only add if we haven't seen this exact course-section combination
                        if pair_key not in seen_pairs:
                            seen_pairs.add(pair_key)
                            unique_pairs.append(f'{course_code} ({section_label})')
                
                return ', '.join(unique_pairs) if unique_pairs else ''
            
            # Prepare jobs data for template
            jobs = [
                'প্রশ্নপত্র প্রণয়ন',
                'প্রশ্নপত্র মডারেশন',
                'উত্তরপত্র পরীক্ষণ',
                'ক্লাস টেস্ট/টার্ম পেপার/ হোম ওয়ার্ক/ এ্যাসাইনমেন্ট',
                'সেশনাল',
                'সেশনাল মৌখিক পরীক্ষা',
                'প্রফেশনাল এ্যাটাসমেন্ট/ইন্ডাস্ট্রিয়াল (ট্রেনিং/এ্যাটাসমেন্ট)',
                'উত্তরপত্র নিরীক্ষণ',
                'টেবুলেশন',
                'প্রশ্নপত্র প্রস্তুতকরণ (অংকন, স্টেনসিল কাটা ও ঘুরানো)',
                'পরীক্ষা কমিটির সভাপতি/সদস্য',
                'চীফ ইনভিজিলেশন / ইনভিজিলেশন',
                'থিসিস',
                'ভাইভা',
                'কোডিং/ডিকোডিং',
                'অন্যান্য'
            ]
            
            jobs_data = []
            serial = 1
            
            for idx, job in enumerate(jobs, 1):
                # Get rate from dropdown or custom input - use exactly what's in the form
                rate_val = get_val(f'rate_{idx}')
                if rate_val == 'custom':
                    rate_val = get_val(f'rate_custom_{idx}')
                
                # For row 3, show breakdown format: "4 × 80 = 320"
                # If calculated amount < 600, rate should be fixed at 600
                quantity_display = get_val(f'quantity_{idx}')
                if idx == 3:
                    # Try to get breakdown from row3_breakdown field
                    row3_breakdown = get_val('row3_breakdown')
                    if row3_breakdown:
                        # Parse breakdown text and format it
                        breakdown_lines = row3_breakdown.strip().split('\n')
                        formatted_breakdowns = []
                        for line in breakdown_lines:
                            # Extract scripts, rate, and calculated amount from breakdown
                            # Format might be: "4 × 80 = 320" or "4 × 80 = 320 < 600 → 600"
                            if '×' in line and '=' in line:
                                # Clean up the line to extract calculation
                                parts = line.split('=')
                                if len(parts) >= 2:
                                    left_part = parts[0].strip()
                                    right_part = parts[1].strip()
                                    # Extract scripts and rate
                                    if '×' in left_part:
                                        calc_parts = left_part.split('×')
                                        if len(calc_parts) == 2:
                                            scripts = calc_parts[0].strip()
                                            rate = calc_parts[1].strip()
                                            # Extract calculated amount
                                            calc_amount = right_part.split('<')[0].strip()
                                            try:
                                                scripts_num = float(scripts)
                                                rate_num = float(rate)
                                                calc_amount_num = float(calc_amount)
                                                # Format: "4 × 80 = 320"
                                                formatted_breakdowns.append(f'{int(scripts_num)} × {int(rate_num)} = {int(calc_amount_num)}')
                                            except (ValueError, TypeError):
                                                formatted_breakdowns.append(line.strip())
                                    else:
                                        formatted_breakdowns.append(line.strip())
                                else:
                                    formatted_breakdowns.append(line.strip())
                            else:
                                formatted_breakdowns.append(line.strip())
                        
                        if formatted_breakdowns:
                            quantity_display = ' '.join(formatted_breakdowns)
                    
                    # If no breakdown found, try to extract from quantity field or calculate
                    if not quantity_display or quantity_display == '0':
                        # Get rate and quantity to create breakdown
                        scripts_count = get_val('quantity_3')
                        current_rate = rate_val
                        if scripts_count and current_rate:
                            try:
                                scripts = float(scripts_count)
                                rate = float(current_rate.replace(',', '').replace('৳', '').strip() or '0')
                                if scripts > 0 and rate > 0:
                                    calculated_amount = scripts * rate
                                    # Check if minimum 600 should apply
                                    if calculated_amount < 600:
                                        # Show breakdown and indicate rate is fixed at 600
                                        quantity_display = f'{int(scripts)} × {int(rate)} = {int(calculated_amount)}'
                                        # Rate should be 600 if calculated < 600
                                        rate_val = '600'
                                    else:
                                        quantity_display = f'{int(scripts)} × {int(rate)} = {int(calculated_amount)}'
                            except (ValueError, TypeError):
                                pass
                
                # Use the rate exactly as it appears in the form - but adjust for Row 3 minimum
                if idx == 3:
                    # For Row 3, if calculated amount < 600, rate should be 600
                    scripts_count = get_val('quantity_3')
                    current_rate = rate_val
                    if scripts_count and current_rate:
                        try:
                            scripts = float(scripts_count)
                            rate = float(current_rate.replace(',', '').replace('৳', '').strip() or '0')
                            if scripts > 0 and rate > 0:
                                calculated_amount = scripts * rate
                                if calculated_amount < 600:
                                    rate_val = '600'
                        except (ValueError, TypeError):
                            pass
                
                # For row 4, show calculation process (student_count × multiplier = product)
                if idx == 4:
                    # Row 4 uses arrays for multiple courses: student_count_4[], section_multiplier_4[], quantity_4[]
                    student_counts = request.form.getlist('student_count_4[]')
                    section_multipliers = request.form.getlist('section_multiplier_4[]')
                    quantity_products = request.form.getlist('quantity_4[]')
                    
                    # Build breakdown string for all courses
                    breakdowns = []
                    for i in range(len(student_counts)):
                        if student_counts[i] and section_multipliers[i] and quantity_products[i]:
                            try:
                                student_count = float(student_counts[i])
                                multiplier = float(section_multipliers[i])
                                product = float(quantity_products[i])
                                # Verify calculation is correct
                                if abs(student_count * multiplier - product) < 0.01:  # Allow small floating point differences
                                    breakdowns.append(f'{int(student_count)} × {int(multiplier)} = {int(product)}')
                                else:
                                    # If calculation doesn't match, show corrected version
                                    corrected_product = int(student_count * multiplier)
                                    breakdowns.append(f'{int(student_count)} × {int(multiplier)} = {corrected_product}')
                            except (ValueError, TypeError):
                                # If parsing fails, just show the product
                                if quantity_products[i]:
                                    breakdowns.append(quantity_products[i])
                    
                    if breakdowns:
                        quantity_display = ' '.join(breakdowns)
                    else:
                        # Fallback to total quantity if breakdowns not available
                        quantity_display = get_val(f'quantity_{idx}') or '0'
                
                # For row 2, if course is "ALL", show "All" in PDF (backend calculation uses 1)
                if idx == 2:
                    quantity_val = get_val(f'quantity_{idx}')
                    # Check if course is "ALL" - if so, display "All" in quantity column
                    course_sections_str = get_course_sections(str(idx))
                    if course_sections_str and 'ALL' in course_sections_str.upper():
                        quantity_display = 'All'
                    elif quantity_val == 'All':
                        quantity_display = 'All'
                    else:
                        quantity_display = quantity_val
                
                # For rows 9, 10, 12: main row has no rate/amount
                if idx in [9, 10, 12]:
                    jobs_data.append({
                        'serial': str(serial),
                        'description': job,
                        'courses': '',
                        'quantity': '',
                        'paper_type': '',
                        'rate': '',
                        'amount': ''
                    })
                else:
                    # For row 11, 13: no course picker (empty courses)
                    courses_val = ''
                    if idx == 11:
                        courses_val = ''  # Row 11 has no course picker
                    elif idx == 13:
                        courses_val = ''  # Row 13 has no course picker
                    elif idx == 16:
                        courses_val = get_val('course_custom_16') or ''  # Row 16 uses custom text input
                    else:
                        courses_val = get_course_sections(str(idx))
                    
                    jobs_data.append({
                        'serial': str(serial),
                        'description': job,
                        'courses': courses_val,
                        'quantity': quantity_display,
                        'paper_type': get_val(f'paper_type_{idx}'),
                        'rate': rate_val,
                        'amount': get_val(f'amount_{idx}')
                    })
                
                serial += 1
                
                # Add tabulation sub-items (row 9)
                if idx == 9:
                    sub_items = [
                        ('কোর্স ভিত্তিক', '9a'),
                        ('পরীক্ষার্থী ভিত্তিক', '9b')
                    ]
                    for sub_job, sub_key in sub_items:
                        sub_rate_val = get_val(f'rate_{sub_key}')
                        if sub_rate_val == 'custom':
                            sub_rate_val = get_val(f'rate_custom_{sub_key}')
                        
                        jobs_data.append({
                            'serial': '',
                            'description': sub_job,
                            'courses': '',
                            'quantity': get_val(f'quantity_{sub_key}'),
                            'paper_type': get_val(f'paper_type_{sub_key}'),
                            'rate': sub_rate_val,
                            'amount': get_val(f'amount_{sub_key}')
                        })
                
                # Add question preparation sub-items (row 10)
                if idx == 10:
                    sub_items = [
                        ('অংকন', '10a'),
                        ('ফটোকপি', '10b')
                    ]
                    for sub_job, sub_key in sub_items:
                        sub_rate_val = get_val(f'rate_{sub_key}')
                        if sub_rate_val == 'custom':
                            sub_rate_val = get_val(f'rate_custom_{sub_key}')
                        
                        jobs_data.append({
                            'serial': '',
                            'description': sub_job,
                            'courses': '',
                            'quantity': get_val(f'quantity_{sub_key}'),
                            'paper_type': get_val(f'paper_type_{sub_key}'),
                            'rate': sub_rate_val,
                            'amount': get_val(f'amount_{sub_key}')
                        })
                
                # Add invigilation sub-items (row 12)
                if idx == 12:
                    sub_items = [
                        ('চীফ ইনভিজিলেশন', '12a'),
                        ('ইনভিজিলেশন', '12b')
                    ]
                    for sub_job, sub_key in sub_items:
                        sub_rate_val = get_val(f'rate_{sub_key}')
                        if sub_rate_val == 'custom':
                            sub_rate_val = get_val(f'rate_custom_{sub_key}')
                        
                        jobs_data.append({
                            'serial': '',
                            'description': sub_job,
                            'courses': '',
                            'quantity': get_val(f'quantity_{sub_key}'),
                            'paper_type': get_val(f'paper_type_{sub_key}'),
                            'rate': sub_rate_val,
                            'amount': get_val(f'amount_{sub_key}')
                        })
                
                # Add thesis sub-items
                if idx == 13:
                    sub_items = [
                        ('পরীক্ষণ', '13a'),
                        ('সুপারভিশন (থিসিস/প্রজেক্ট থিসিস/ইন্টার্নশীপ রিপোর্ট)', '13b'),
                        ('কো-সুপারভিশন', '13c'),
                        ('মৌখিক পরীক্ষা', '13d')
                    ]
                    for sub_job, sub_key in sub_items:
                        sub_rate_val = get_val(f'rate_{sub_key}')
                        if sub_rate_val == 'custom':
                            sub_rate_val = get_val(f'rate_custom_{sub_key}')
                        
                        jobs_data.append({
                            'serial': '',
                            'description': sub_job,
                            'courses': '',  # Row 13 sub-rows don't have course picker
                            'quantity': get_val(f'quantity_{sub_key}'),
                            'paper_type': get_val(f'paper_type_{sub_key}'),
                            'rate': sub_rate_val,
                            'amount': get_val(f'amount_{sub_key}')
                        })
                
                # Handle row 16 custom course field (text input)
                if idx == 16:
                    custom_course = get_val('course_custom_16') or ''
                    # Update the last added job_data (row 16) with custom course
                    if jobs_data and jobs_data[-1]['serial'] == str(serial - 1):
                        jobs_data[-1]['courses'] = custom_course
            
            # Get font file path for Bengali font - try multiple paths
            font_path_absolute = None
            font_paths_to_try = [
                os.path.join(current_app.root_path, 'static', 'Fonts', 'kalpurush.ttf'),  # Capital F
                os.path.join(current_app.root_path, 'static', 'fonts', 'kalpurush.ttf'),   # Lowercase f
                os.path.join(current_app.root_path, 'static', 'Fonts', 'Kalpurush.ttf'),   # Capital K
                os.path.join(current_app.root_path, 'static', 'fonts', 'Kalpurush.ttf'),   # Capital K, lowercase f
            ]

            for font_path in font_paths_to_try:
                if os.path.exists(font_path):
                    font_path_absolute = os.path.abspath(font_path)
                    current_app.logger.info(f'Kalpurush font found at: {font_path_absolute}')
                    break

            if not font_path_absolute:
                current_app.logger.warning('Kalpurush font not found in any expected location. Bengali text may not render correctly.')
                current_app.logger.warning(f'Searched paths: {font_paths_to_try}')
            
            # Get logo file path and convert to data URI for PDF
            logo_path_absolute = None
            logo_paths_to_try = [
                os.path.join(current_app.root_path, 'static', 'Images', 'KU_logo_2.png'),  # Capital I
                os.path.join(current_app.root_path, 'static', 'images', 'KU_logo_2.png'),   # Lowercase i
            ]
            
            for logo_path in logo_paths_to_try:
                if os.path.exists(logo_path):
                    logo_path_absolute = os.path.abspath(logo_path)
                    current_app.logger.info(f'Logo found at: {logo_path_absolute}')
                    break
            
            # Convert logo to base64 data URI if found
            logo_data_uri = None
            if logo_path_absolute:
                try:
                    import base64
                    with open(logo_path_absolute, 'rb') as logo_file:
                        logo_data = logo_file.read()
                        logo_base64 = base64.b64encode(logo_data).decode('utf-8')
                        logo_data_uri = f'data:image/png;base64,{logo_base64}'
                        current_app.logger.info('Logo converted to data URI for PDF')
                except Exception as e:
                    current_app.logger.error(f'Error converting logo to data URI: {e}')
                    logo_data_uri = None
            
            # Render HTML template with data
            html_content = render_template(
                'remuneration_pdf_template.html',
                font_path=font_path_absolute,
                logo_data_uri=logo_data_uri,
                voucher_no=get_val('voucher_no'),
                voucher_date=get_val('voucher_date'),
                applicant_name=get_val('applicant_name'),
                designation=get_val('designation'),
                discipline=get_val('discipline') or 'Law',
                address=get_val('address'),
                exam_discipline=get_val('exam_discipline') or 'Law',
                year=get_val('year'),
                academic_year=get_val('academic_year'),
                term=get_val('term'),
                exam_start_date=get_val('exam_start_date'),
                exam_end_date=get_val('exam_end_date'),
                jobs_data=jobs_data,
                total_amount=get_val('total_amount'),
                total_in_words=get_val('total_in_words'),
                bank_account=get_val('bank_account'),
                auditor_sign=get_val('auditor_sign'),
                deputy_sign=get_val('deputy_sign'),
                controller_sign=get_val('controller_sign'),
                finance_amount=get_val('finance_amount'),
                finance_amount_words=get_val('finance_amount_words'),
                section_officer_sign=get_val('section_officer_sign'),
                deputy_director_sign=get_val('deputy_director_sign'),
                director_sign=get_val('director_sign'),
                audit_amount=get_val('audit_amount'),
                audit_assistant_sign=get_val('audit_assistant_sign'),
                audit_head_sign=get_val('audit_head_sign'),
                bank_advice_no=get_val('bank_advice_no'),
                payment_date=get_val('payment_date')
            )
            
            # Generate PDF using WeasyPrint
            pdf_buffer = BytesIO()
            
            # Create CSS for PDF - slightly increased spacing, single page
            css_string = """
            @page {
                size: 8.5in 14in; /* Legal size */
                margin: 0.18in 0.28in; /* Slightly increased margins */
            }
            * {
                page-break-inside: avoid !important;
                page-break-after: avoid !important;
                page-break-before: avoid !important;
            }
            body {
                margin: 0 !important;
                padding: 0 !important;
                font-size: 0.64rem !important; /* Increased from 0.56rem (7pt) to 0.64rem (8pt) */
                line-height: 1.18 !important; /* Slightly more breathing room */
                font-family: 'Kalpurush', sans-serif !important;
            }
            .rem-wrapper {
                margin: 0 !important;
                padding: 0 !important;
            }
            .rem-sheet {
                padding: 10px 17px !important; /* Slightly increased padding */
                margin: 0 !important;
                border: none !important;
                border-radius: 0 !important;
            }
            .rem-heading {
                margin-bottom: 0.15rem !important;
                gap: 0.7rem !important;
            }
            .rem-heading-logo img {
                height: 38px !important; /* Larger logo */
            }
            .rem-heading-content .text-muted {
                font-size: 0.62rem !important;
                margin-bottom: 0.04rem !important;
            }
            .rem-heading h4 {
                font-size: 0.78rem !important;
                margin-bottom: 0.015rem !important;
            }
            .rem-heading-content > div {
                font-size: 0.72rem !important;
                margin-bottom: 0.015rem !important;
            }
            .rem-heading-content small {
                font-size: 0.56rem !important;
                margin-top: 0.04rem !important;
            }
            .voucher-box {
                padding: 0.12rem 0.35rem 0.015rem 0.35rem !important;
                font-size: 0.56rem !important;
                width: 210px !important;
            }
            .voucher-box div {
                padding: 0.04rem 0 !important;
            }
            .voucher-box span {
                min-width: 52px !important;
                font-size: 0.56rem !important;
            }
            .voucher-box input {
                font-size: 0.56rem !important;
                padding: 0.06rem 0.16rem !important;
            }
            .meta-grid {
                margin-top: 0.25rem !important;
                margin-bottom: 0.25rem !important;
            }
            .meta-grid td {
                padding: 0.18rem 0.32rem !important;
                font-size: 0.56rem !important; /* Reverted to original 7pt */
            }
            .meta-label {
                font-size: 0.56rem !important; /* Reverted to original 7pt */
                width: 142px !important;
            }
            .meta-grid input,
            .meta-grid select {
                font-size: 0.56rem !important; /* Reverted to original 7pt */
                padding: 0.11rem 0.22rem !important;
            }
            .rem-table {
                margin: 0.25rem 0 !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
            }
            .rem-table th,
            .rem-table td {
                padding: 0.18rem 0.27rem !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                line-height: 1.12 !important;
            }
            .rem-table th {
                padding: 0.22rem 0.27rem !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
            }
            .section-title {
                margin-top: 0.65rem !important;
                margin-bottom: 0.35rem !important;
                font-size: 0.71rem !important;
            }
            .signature-box {
                min-height: 38px !important;
                padding: 0.22rem !important;
                font-size: 0.56rem !important;
            }
            .signature-box span {
                margin-top: 0.55rem !important;
                padding-top: 0.18rem !important;
            }
            .controller-signature-section,
            .finance-signature-section,
            .audit-approval-section {
                margin-top: 0.25rem !important;
                margin-bottom: 0.28rem !important;
                gap: 0.25rem !important;
            }
            .controller-signature-box,
            .finance-signature-box,
            .audit-signature-box-single {
                min-height: 32px !important;
                padding: 0.17rem 0.16rem !important;
            }
            .controller-designation,
            .finance-designation,
            .audit-designation {
                font-size: 0.51rem !important;
                margin-bottom: 0.27rem !important;
                padding-bottom: 0.16rem !important;
            }
            .controller-signature-line,
            .finance-signature-line,
            .audit-signature-line {
                margin-top: 0.22rem !important;
                font-size: 0.51rem !important;
            }
            .foot-table {
                margin: 0.25rem 0 0.18rem 0 !important;
            }
            .foot-table td {
                padding: 0.18rem 0.32rem !important;
                font-size: 0.56rem !important; /* Reverted to original 7pt */
            }
            .foot-table input {
                font-size: 0.56rem !important; /* Reverted to original 7pt */
                padding: 0.11rem 0.22rem !important;
            }
            .info-note,
            .statement-note,
            .finance-release-note,
            .audit-approval-text {
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                margin: 0.18rem 0 !important;
                line-height: 1.22 !important;
            }
            .info-note input,
            .statement-note input {
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                padding: 0.04rem 0.11rem !important;
            }
            .bank-declaration {
                padding: 0.27rem !important;
                margin-top: 0.25rem !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                line-height: 1.22 !important;
            }
            .revenue-ticket {
                width: 18mm !important;
                height: 18mm !important;
                margin-left: 0.35rem !important;
            }
            """
            
            # Initialize html_content for font injection
            html_content_final = html_content
            
            if font_path_absolute:
                # Inject @font-face rule directly into HTML <head> section
                # This works better with WeasyPrint 52.5 than CSS object
                try:
                    import base64
                    with open(font_path_absolute, 'rb') as font_file:
                        font_data = font_file.read()
                        font_base64 = base64.b64encode(font_data).decode('utf-8')
                    
                    # Try application/font-sfnt MIME type for TTF fonts
                    font_face_rule = f"""
        <style>
        @font-face {{
            font-family: 'Kalpurush';
            src: url(data:application/font-sfnt;base64,{font_base64}) format('truetype');
            font-weight: normal;
            font-style: normal;
        }}
        @font-face {{
            font-family: 'Kalpurush';
            src: url(data:application/font-sfnt;base64,{font_base64}) format('truetype');
            font-weight: bold;
            font-style: normal;
        }}
        </style>
        """
                    
                    # Inject into HTML <head> section
                    if '</head>' in html_content_final:
                        html_content_final = html_content_final.replace('</head>', font_face_rule + '</head>')
                    elif '<head>' in html_content_final:
                        html_content_final = html_content_final.replace('<head>', '<head>' + font_face_rule)
                    else:
                        # No head tag, add it at the beginning
                        html_content_final = font_face_rule + html_content_final
                    
                    # Also add to CSS string for redundancy
                    font_css = """
        * {
            font-family: 'Kalpurush', sans-serif !important;
        }
        body {
            font-family: 'Kalpurush', sans-serif !important;
        }
        .english-text {
            font-family: 'Tahoma', 'Arial', sans-serif !important;
        }
        """
                    css_string = font_css + css_string
                    current_app.logger.info('Kalpurush font embedded as base64 in HTML <head>')
                except Exception as e:
                    current_app.logger.error(f'Failed to embed font as base64: {e}', exc_info=True)
                    # Fallback: try file:// URL (works on some systems)
                    if os.name == 'nt':  # Windows
                        font_url = f"file:///{font_path_absolute.replace(os.sep, '/').replace(':', '')}"
                    else:  # macOS/Linux
                        font_url = f"file://{font_path_absolute}"
                    
                    font_face_rule = f"""
        <style>
        @font-face {{
            font-family: 'Kalpurush';
            src: url('{font_url}') format('truetype');
            font-weight: normal;
            font-style: normal;
        }}
        @font-face {{
            font-family: 'Kalpurush';
            src: url('{font_url}') format('truetype');
            font-weight: bold;
            font-style: normal;
        }}
        </style>
        """
                    
                    if '</head>' in html_content_final:
                        html_content_final = html_content_final.replace('</head>', font_face_rule + '</head>')
                    elif '<head>' in html_content_final:
                        html_content_final = html_content_final.replace('<head>', '<head>' + font_face_rule)
                    
                    font_css = """
        * {
            font-family: 'Kalpurush', sans-serif !important;
        }
        body {
            font-family: 'Kalpurush', sans-serif !important;
        }
        .english-text {
            font-family: 'Tahoma', 'Arial', sans-serif !important;
        }
        """
                    css_string = font_css + css_string
            
            # Use the final HTML content (with font injected)
            html_content = html_content_final
            
            # Create HTML object
            html_obj = HTML(string=html_content, base_url=request.url_root)
            
            # Create CSS object if font CSS exists
            css_obj = CSS(string=css_string) if css_string else None
            
            # Write PDF with CSS
            if css_obj:
                html_obj.write_pdf(pdf_buffer, stylesheets=[css_obj], presentational_hints=True)
            else:
                html_obj.write_pdf(pdf_buffer, presentational_hints=True)
            pdf_buffer.seek(0)
            pdf_data = pdf_buffer.getvalue()
            
            # Enhanced headers for cPanel compatibility
            response = Response(
                pdf_data,
                mimetype='application/pdf',
                headers={
                    'Content-Disposition': f'attachment; filename="Exam_Remuneration_Form.pdf"; filename*=UTF-8\'\'Exam_Remuneration_Form.pdf',
                    'Content-Length': str(len(pdf_data)),
                    'Cache-Control': 'no-cache, no-store, must-revalidate',
                    'Pragma': 'no-cache',
                    'Expires': '0',
                    'X-Content-Type-Options': 'nosniff',
                    'X-Frame-Options': 'DENY'
                }
            )
            
            return response
            
        except Exception as e:
            error_msg = str(e)
            error_type = type(e).__name__
            current_app.logger.error(f'Error generating PDF: {error_type}: {error_msg}', exc_info=True)
            
            # Provide more helpful error message for debugging
            if 'WeasyPrint' in error_msg or 'weasyprint' in error_msg.lower():
                detailed_msg = f'WeasyPrint error: {error_msg}. Please ensure WeasyPrint and its dependencies (cairo, pango, etc.) are installed.'
            elif 'Template' in error_msg or 'template' in error_msg.lower():
                detailed_msg = f'Template error: {error_msg}. Please check if remuneration_pdf_template.html exists.'
            elif 'Permission' in error_msg or 'permission' in error_msg.lower():
                detailed_msg = f'Permission error: {error_msg}. Please check file permissions and paths.'
            else:
                detailed_msg = f'PDF generation failed: {error_msg}'
            
            return jsonify({
                'error': 'Failed to generate PDF document',
                'message': detailed_msg,
                'type': error_type
            }), 500
        finally:
            # Clean up memory after PDF generation to prevent memory accumulation
            import gc
            gc.collect()

    @app.route('/remuneration/api/teachers', methods=['GET'])
    @login_required
    def remuneration_get_teachers():
        """Get list of teachers for bulk download selection - filtered by statement participants"""
        restriction = _require_teacher_privileges()
        if restriction:
            return restriction
        
        session_val = request.args.get('session', '').strip()
        year_val = request.args.get('year', '').strip()
        term_val = request.args.get('term', '').strip()
        
        if not session_val or not year_val or not term_val:
            return jsonify({
                'success': False,
                'message': 'সেশন, বর্ষ ও টার্ম সিলেক্ট করুন'
            }), 400
        
        try:
            from blueprints.class_management.models import Teacher
            from role_utils import get_teachers_excluding_head
            
            # Fetch statement data
            statement_data = _get_remuneration_statement_data(session_val, year_val, term_val)
            if not statement_data:
                # If no statement data, return empty list with message
                return jsonify({
                    'success': True,
                    'teachers': [],
                    'message': 'এই সেশন, বর্ষ ও টার্মের জন্য Statement of Remuneration পাওয়া যায়নি'
                })
            
            # Collect all teacher names from statement data
            teacher_names_in_statement = set()
            
            # Check all relevant statement sections
            sections_to_check = [
                'question_preparation',
                'moderation_committee',  # For Row 2
                'script_examination',
                'class_test',
                'sessional_assessment',
                'sessional_viva',
                'professional_attachment',
                'script_scrutiny',
                'tabulation',
                'invigilation',
                'question_typing',
                'question_photocopy',
                'examination_committee',
                'thesis_supervision',
                'viva',
                'coding_decoding'  # For Row 15
            ]
            
            for section in sections_to_check:
                items = statement_data.get(section, [])
                if not items:
                    continue
                
                for item in items:
                    # Extract teacher name from various field names
                    teacher_name = (
                        item.get('teacher') or
                        item.get('name') or
                        item.get('teacher_name') or
                        ''
                    )
                    if teacher_name:
                        teacher_names_in_statement.add(str(teacher_name).strip())
            
            if not teacher_names_in_statement:
                return jsonify({
                    'success': True,
                    'teachers': [],
                    'message': 'Statement-এ কোনো শিক্ষকের তথ্য পাওয়া যায়নি'
                })
            
            current_app.logger.info(
                f'Found {len(teacher_names_in_statement)} teachers in statement: {session_val}/{year_val}/{term_val}'
            )
            
            # Get all teachers (excluding head)
            all_teachers = get_teachers_excluding_head()
            
            # Filter teachers that appear in statement
            teachers_list = []
            for teacher in all_teachers:
                teacher_name = (teacher.name or '').strip()
                if not teacher_name:
                    continue
                
                # Check if this teacher matches any name in statement (flexible matching)
                matches = False
                for stmt_name in teacher_names_in_statement:
                    if _matches_teacher_name(stmt_name, teacher_name):
                        matches = True
                        break
                
                if matches:
                    teachers_list.append({
                        'id': teacher.id,
                        'name': teacher.name,
                        'designation': teacher.designation or '',
                        'institute': teacher.institute or ''
                    })
            
            current_app.logger.info(
                f'Filtered to {len(teachers_list)} teachers matching statement data'
            )
            
            return jsonify({
                'success': True,
                'teachers': teachers_list
            })
        except Exception as e:
            current_app.logger.error(f'Error fetching teachers: {str(e)}', exc_info=True)
            return jsonify({
                'success': False,
                'message': 'শিক্ষক লোড করতে সমস্যা হয়েছে'
            }), 500

    def _get_remuneration_statement_data(session, year, term):
        """Fetch Exam Committee Chief's remuneration statement data for session/year/term."""
        try:
            from blueprints.remuneration_management.models import RemunerationForm
            from blueprints.course_management.models import DutyAssignment
            import json

            chief_assignment = DutyAssignment.query.filter_by(
                duty_type='exam_committee_chief',
                academic_session=session,
                year=year,
                term=term,
                status='active'
            ).first()

            if chief_assignment and chief_assignment.assigned_teacher:
                chief_user = User.query.filter_by(full_name=chief_assignment.assigned_teacher.name).first()
                if chief_user:
                    form_entry = RemunerationForm.query.filter_by(
                        user_id=chief_user.id,
                        academic_year=session,
                        year=year,
                        term=term
                    ).order_by(RemunerationForm.id.desc()).first()

                    if form_entry and form_entry.form_data:
                        return json.loads(form_entry.form_data)
        except Exception as exc:
            current_app.logger.error(f'Failed to load statement data: {exc}', exc_info=True)
        return None

    def _has_remuneration_line_items(form_data):
        """Return True if any row quantities/amounts exist in form_data (non-zero)."""
        if not isinstance(form_data, dict):
            return False

        def _has_value(val):
            if val is None:
                return False
            text = str(val).strip()
            if text == '' or text == '0' or text == '0.0' or text == '0.00':
                return False
            # If numeric and > 0, treat as value
            try:
                return float(text) > 0
            except Exception:
                return True

        for idx in range(1, 17):
            qty = form_data.get(f'quantity_{idx}', '')
            amt = form_data.get(f'amount_{idx}', '')
            rate = form_data.get(f'rate_{idx}', '')
            if _has_value(qty) or _has_value(amt) or _has_value(rate):
                return True

        # Handle split rows (12a/12b, 10a/10b, etc.)
        for key in ['quantity_12a', 'quantity_12b', 'amount_12a', 'amount_12b',
                    'quantity_10a', 'quantity_10b', 'amount_10a', 'amount_10b']:
            if _has_value(form_data.get(key, '')):
                return True
        return False

    def _auto_populate_remuneration_data(teacher, session, year, term):
        """Auto-populate remuneration form data from statement or assignments"""
        form_data = {}
        total_amount = 0
        jobs_rows = []
        
        try:
            from blueprints.class_management.models import ExamPaperEvaluatorAssignment, ExamPaperEvaluation
            from blueprints.course_management.models import Course, DutyAssignment
            import json

            current_app.logger.info(
                f'Auto-populating remuneration data for {teacher.name}, session={session}, year={year}, term={term}'
            )

            statement_data = _get_remuneration_statement_data(session, year, term)

            def matches_teacher(item_teacher):
                if not item_teacher:
                    return False
                item_teacher = str(item_teacher).strip().lower()
                target = (teacher.name or '').strip().lower()
                return item_teacher == target or item_teacher in target or target in item_teacher

            def safe_int(value, default=0):
                try:
                    return int(float(value))
                except Exception:
                    return default

            def is_postgraduate():
                yr = (year or '').upper()
                return any(token in yr for token in ['LLM', 'MASTER', 'M.PHIL', 'MPHIL', 'PHD'])

            is_pg = is_postgraduate()

            if statement_data:
                # Pass through exam dates if available in statement data
                if statement_data.get('exam_start_date'):
                    form_data['exam_start_date'] = statement_data.get('exam_start_date')
                if statement_data.get('exam_end_date'):
                    form_data['exam_end_date'] = statement_data.get('exam_end_date')

                # Row 1: Question Preparation - use course-level PG/UG, not year-level
                items = [i for i in statement_data.get('question_preparation', []) if matches_teacher(i.get('teacher'))]
                if items:
                    courses = []
                    qty = 0
                    for i in items:
                        course = i.get('course') or i.get('course_code') or ''
                        section = i.get('section') or ''
                        if course:
                            course_text = f"{course} ({section})" if section else course
                            courses.append(course_text)
                        qty += max(safe_int(i.get('questions', 1)), 1)
                    # Determine rate by course code (e.g. 042 = UG), not by program year (LLM)
                    first_course = (items[0].get('course') or items[0].get('course_code') or '').strip()
                    if ' - ' in first_course:
                        first_course = first_course.split(' - ')[0].strip()
                    is_pg_row1 = _is_postgraduate_course(first_course, year=None) if first_course else is_pg
                    rate = 2400 if is_pg_row1 else 2300
                    amount = qty * rate
                    form_data['course_section_1'] = ', '.join(courses)
                    form_data['quantity_1'] = str(qty)
                    form_data['rate_1'] = str(rate)
                    form_data['amount_1'] = str(amount)
                    total_amount += amount
                    jobs_rows.append({
                        'description': 'প্রশ্নপত্র প্রণয়ন',
                        'courses': form_data.get('course_section_1', ''),
                        'quantity': form_data.get('quantity_1', ''),
                        'paper_type': '',
                        'rate': form_data.get('rate_1', ''),
                        'amount': form_data.get('amount_1', '')
                    })

                # Row 3: Script Examination
                items = [i for i in statement_data.get('script_examination', []) if matches_teacher(i.get('teacher'))]
                if items:
                    courses = []
                    total_scripts = 0
                    per_course_amounts = []
                    rate = 100 if is_pg else 80
                    for i in items:
                        course = i.get('course') or i.get('course_code') or ''
                        section = i.get('section') or ''
                        scripts = safe_int(i.get('scripts', 0))
                        if course:
                            course_text = f"{course} ({section})" if section else course
                            courses.append(course_text)
                        total_scripts += scripts
                        if scripts > 0:
                            amount = scripts * rate
                            per_course_amounts.append(max(amount, 600))
                    if total_scripts > 0:
                        form_data['course_section_3'] = ', '.join(courses)
                        form_data['quantity_3'] = str(total_scripts)
                        form_data['rate_3'] = str(rate)
                        amount = sum(per_course_amounts) if per_course_amounts else total_scripts * rate
                        form_data['amount_3'] = str(amount)
                        total_amount += amount
                        jobs_rows.append({
                            'description': 'উত্তরপত্র পরীক্ষণ',
                            'courses': form_data.get('course_section_3', ''),
                            'quantity': form_data.get('quantity_3', ''),
                            'paper_type': '',
                            'rate': form_data.get('rate_3', ''),
                            'amount': form_data.get('amount_3', '')
                        })

                # Row 4: Class Test
                items = [i for i in statement_data.get('class_test', []) if matches_teacher(i.get('teacher') or i.get('teacher_name'))]
                if items:
                    courses = []
                    total_students = 0
                    for i in items:
                        course = i.get('course') or i.get('course_code') or ''
                        section = i.get('section') or ''
                        students = safe_int(i.get('students', 0))
                        if course:
                            course_text = f"{course} ({section})" if section else course
                            courses.append(course_text)
                        total_students += students
                    rate = 80 if is_pg else 30
                    amount = total_students * rate
                    form_data['course_section_4'] = ', '.join(courses)
                    form_data['quantity_4'] = str(total_students)
                    form_data['rate_4'] = str(rate)
                    form_data['amount_4'] = str(amount)
                    total_amount += amount
                    jobs_rows.append({
                        'description': 'ক্লাস টেস্ট/টার্ম পেপার/ হোম ওয়ার্ক/ এ্যাসাইনমেন্ট',
                        'courses': form_data.get('course_section_4', ''),
                        'quantity': form_data.get('quantity_4', ''),
                        'paper_type': '',
                        'rate': form_data.get('rate_4', ''),
                        'amount': form_data.get('amount_4', '')
                    })

                # Row 5: Sessional Assessment
                items = [i for i in statement_data.get('sessional_assessment', []) if matches_teacher(i.get('teacher'))]
                if items:
                    courses = []
                    total_students = 0
                    for i in items:
                        course = i.get('course') or ''
                        students = safe_int(i.get('students', 0))
                        if course:
                            courses.append(course)
                        total_students += students
                    rate = 230  # default rate
                    amount = total_students * rate
                    form_data['course_section_5'] = ', '.join(courses)
                    form_data['quantity_5'] = str(total_students)
                    form_data['rate_5'] = str(rate)
                    form_data['amount_5'] = str(amount)
                    total_amount += amount
                    jobs_rows.append({
                        'description': 'সেশনাল',
                        'courses': form_data.get('course_section_5', ''),
                        'quantity': form_data.get('quantity_5', ''),
                        'paper_type': '',
                        'rate': form_data.get('rate_5', ''),
                        'amount': form_data.get('amount_5', '')
                    })

                # Row 6: Sessional Viva
                items = [i for i in statement_data.get('sessional_viva', []) if matches_teacher(i.get('teacher'))]
                if items:
                    courses = []
                    total_students = 0
                    for i in items:
                        course = i.get('course') or ''
                        students = safe_int(i.get('students', 0))
                        if course:
                            courses.append(course)
                        total_students += students
                    rate = 50
                    amount = total_students * rate
                    form_data['course_section_6'] = ', '.join(courses)
                    form_data['quantity_6'] = str(total_students)
                    form_data['rate_6'] = str(rate)
                    form_data['amount_6'] = str(amount)
                    total_amount += amount
                    jobs_rows.append({
                        'description': 'সেশনাল মৌখিক পরীক্ষা',
                        'courses': form_data.get('course_section_6', ''),
                        'quantity': form_data.get('quantity_6', ''),
                        'paper_type': '',
                        'rate': form_data.get('rate_6', ''),
                        'amount': form_data.get('amount_6', '')
                    })

                # Row 8: Script Scrutiny
                items = [i for i in statement_data.get('script_scrutiny', []) if matches_teacher(i.get('name'))]
                if items:
                    total_scripts = sum(safe_int(i.get('scripts', 0)) for i in items)
                    if total_scripts > 0:
                        rate = 8
                        amount = total_scripts * rate
                        form_data['quantity_8'] = str(total_scripts)
                        form_data['rate_8'] = str(rate)
                        form_data['amount_8'] = str(amount)
                        total_amount += amount
                        jobs_rows.append({
                            'description': 'উত্তরপত্র নিরীক্ষণ',
                            'courses': '',
                            'quantity': form_data.get('quantity_8', ''),
                            'paper_type': '',
                            'rate': form_data.get('rate_8', ''),
                            'amount': form_data.get('amount_8', '')
                        })

                # Row 9: Tabulation
                items = [i for i in statement_data.get('tabulation', []) if matches_teacher(i.get('name'))]
                if items:
                    course_count = 0
                    courses = []
                    for i in items:
                        course_wise = safe_int(i.get('course_wise', 0))
                        course_count += course_wise if course_wise > 0 else 1
                        if i.get('course_wise'):
                            courses.append(i.get('course_wise'))
                    rate = 200
                    amount = course_count * rate
                    form_data['quantity_9'] = str(course_count)
                    form_data['rate_9'] = str(rate)
                    form_data['amount_9'] = str(amount)
                    total_amount += amount
                    jobs_rows.append({
                        'description': 'টেবুলেশন',
                        'courses': '',
                        'quantity': form_data.get('quantity_9', ''),
                        'paper_type': '',
                        'rate': form_data.get('rate_9', ''),
                        'amount': form_data.get('amount_9', '')
                    })

                # Row 12: Invigilation
                items = [i for i in statement_data.get('invigilation', []) if matches_teacher(i.get('name'))]
                if items:
                    chief_count = sum(safe_int(i.get('chief', 0)) for i in items)
                    inv_count = sum(safe_int(i.get('invigilation', 0)) for i in items)
                    if chief_count > 0:
                        form_data['quantity_12a'] = str(chief_count)
                        form_data['rate_12a'] = '1800'
                        chief_amount = chief_count * 1800
                        form_data['amount_12a'] = str(chief_amount)
                        total_amount += chief_amount
                    if inv_count > 0:
                        form_data['quantity_12b'] = str(inv_count)
                        form_data['rate_12b'] = '1500'
                        inv_amount = inv_count * 1500
                        form_data['amount_12b'] = str(inv_amount)
                        total_amount += inv_amount
                    if chief_count > 0 or inv_count > 0:
                        jobs_rows.append({
                            'description': 'চীফ ইনভিজিলেশন / ইনভিজিলেশন',
                            'courses': '',
                            'quantity': f"চীফ: {chief_count}, ইনভি: {inv_count}" if chief_count and inv_count else (f"চীফ: {chief_count}" if chief_count else f"ইনভি: {inv_count}"),
                            'paper_type': '',
                            'rate': '1800/1500',
                            'amount': str((chief_count * 1800) + (inv_count * 1500))
                        })

            else:
                # Fallback: use evaluator assignments if statement data is not available
                assignments = ExamPaperEvaluatorAssignment.query.filter_by(
                    academic_session=session,
                    year=year,
                    term=term
                ).all()

                # Row 1 and 3 from evaluator assignments
                question_courses = []
                script_courses = []
                total_scripts = 0
                rate_1 = 2400 if is_pg else 2300
                rate_3 = 100 if is_pg else 80

                for assignment in assignments:
                    course = Course.query.get(assignment.course_id)
                    if not course:
                        continue
                    course_code = course.course_code or ''
                    part = assignment.part or ''
                    course_text = f"{course_code} (Part {part})"

                    if assignment.question_setter_id == teacher.id:
                        question_courses.append(course_text)
                    if assignment.assigned_teacher_id == teacher.id:
                        script_courses.append(course_text)

                        # Script count
                        submitted_entries = ExamPaperEvaluation.query.filter(
                            ExamPaperEvaluation.owner_teacher_id == assignment.assigned_teacher_id,
                            ExamPaperEvaluation.course_code == course_code,
                            ExamPaperEvaluation.section == f"Part {part}",
                            ExamPaperEvaluation.academic_session == session,
                            ExamPaperEvaluation.year == year,
                            ExamPaperEvaluation.term == term,
                            ExamPaperEvaluation.submitted_to_committee == True
                        ).all()

                        for entry in submitted_entries:
                            if entry.marks_data:
                                try:
                                    marks_data = json.loads(entry.marks_data) if isinstance(entry.marks_data, str) else entry.marks_data
                                    if isinstance(marks_data, dict):
                                        if 'rows' in marks_data and isinstance(marks_data['rows'], list):
                                            total_scripts += len(marks_data['rows'])
                                        else:
                                            student_keys = [k for k in marks_data.keys() if k not in ['questions', 'rows']]
                                            total_scripts += len(student_keys)
                                except Exception:
                                    pass

                if question_courses:
                    qty = len(question_courses)
                    amount = qty * rate_1
                    form_data['course_section_1'] = ', '.join(question_courses)
                    form_data['quantity_1'] = str(qty)
                    form_data['rate_1'] = str(rate_1)
                    form_data['amount_1'] = str(amount)
                    total_amount += amount

                if script_courses and total_scripts > 0:
                    amount = total_scripts * rate_3
                    form_data['course_section_3'] = ', '.join(script_courses)
                    form_data['quantity_3'] = str(total_scripts)
                    form_data['rate_3'] = str(rate_3)
                    form_data['amount_3'] = str(amount)
                    total_amount += amount
            
            # Store total amount
            form_data['total_amount'] = str(total_amount)
            
            # Convert total to words (simple conversion)
            if total_amount > 0:
                form_data['total_in_words'] = _number_to_words_bengali(total_amount)
            
            # Store jobs rows for direct PDF rendering
            if jobs_rows:
                form_data['_jobs_data'] = jobs_rows

            current_app.logger.info(f'Auto-populated form data for {teacher.name}: {len(jobs_rows)} rows with data, total={total_amount}')
            
        except Exception as e:
            current_app.logger.error(f'Error auto-populating remuneration data for {teacher.name}: {str(e)}', exc_info=True)
            # Return empty form_data on error
            form_data = {}
        
        return form_data
    
    def _number_to_words_bengali(num):
        """Convert number to Bengali words (matches JavaScript numberToBengaliWords)"""
        try:
            if not num or num == 0:
                return 'শূন্য টাকা'
            
            # Convert to float to handle decimals
            num_float = float(num)
            
            ones = ['', 'এক', 'দুই', 'তিন', 'চার', 'পাঁচ', 'ছয়', 'সাত', 'আট', 'নয়', 'দশ', 
                   'এগারো', 'বারো', 'তেরো', 'চৌদ্দ', 'পনেরো', 'ষোল', 'সতেরো', 'আঠারো', 'উনিশ', 'বিশ']
            tens = ['', '', 'বিশ', 'ত্রিশ', 'চল্লিশ', 'পঞ্চাশ', 'ষাট', 'সত্তর', 'আশি', 'নব্বই']
            
            # Special words for numbers 21-99
            special_numbers = {
                21: 'একুশ', 22: 'বাইশ', 23: 'তেইশ', 24: 'চব্বিশ', 25: 'পঁচিশ', 26: 'ছাব্বিশ', 
                27: 'সাতাশ', 28: 'আটাশ', 29: 'ঊনত্রিশ',
                31: 'একত্রিশ', 32: 'বত্রিশ', 33: 'তেত্রিশ', 34: 'চৌত্রিশ', 35: 'পঁয়ত্রিশ', 
                36: 'ছত্রিশ', 37: 'সাঁইত্রিশ', 38: 'আটত্রিশ', 39: 'ঊনচল্লিশ',
                41: 'একচল্লিশ', 42: 'বিয়াল্লিশ', 43: 'তেতাল্লিশ', 44: 'চুয়াল্লিশ', 45: 'পঁয়তাল্লিশ', 
                46: 'ছেচল্লিশ', 47: 'সাতচল্লিশ', 48: 'আটচল্লিশ', 49: 'ঊনপঞ্চাশ',
                51: 'একান্ন', 52: 'বায়ান্ন', 53: 'তিপ্পান্ন', 54: 'চুয়ান্ন', 55: 'পঞ্চান্ন', 
                56: 'ছাপ্পান্ন', 57: 'সাতান্ন', 58: 'আটান্ন', 59: 'ঊনষাট',
                61: 'একষট্টি', 62: 'বাষট্টি', 63: 'তেষট্টি', 64: 'চৌষট্টি', 65: 'পঁয়ষট্টি', 
                66: 'ছেষট্টি', 67: 'সাতষট্টি', 68: 'আটষট্টি', 69: 'ঊনসত্তর',
                71: 'একাত্তর', 72: 'বাহাত্তর', 73: 'তিয়াত্তর', 74: 'চুয়াত্তর', 75: 'পঁচাত্তর', 
                76: 'ছিয়াত্তর', 77: 'সাতাত্তর', 78: 'আটাত্তর', 79: 'ঊনআশি',
                81: 'একাশি', 82: 'বিরাশি', 83: 'তিরাশি', 84: 'চুরাশি', 85: 'পঁচাশি', 
                86: 'ছিয়াশি', 87: 'সাতাশি', 88: 'আটাশি', 89: 'ঊননব্বই',
                91: 'একানব্বই', 92: 'বিরানব্বই', 93: 'তিরানব্বই', 94: 'চুরানব্বই', 95: 'পঁচানব্বই', 
                96: 'ছিয়ানব্বই', 97: 'সাতানব্বই', 98: 'আটানব্বই', 99: 'নিরানব্বই'
            }
            
            def convert_under_100(n):
                """Convert numbers less than 100"""
                if n == 0:
                    return ''
                if n <= 20:
                    return ones[n]
                if n in special_numbers:
                    return special_numbers[n]
                ten = n // 10
                one = n % 10
                if ten > 0 and one > 0:
                    return tens[ten] + ' ' + ones[one]
                elif ten > 0:
                    return tens[ten]
                else:
                    return ones[one]
            
            def convert_under_1000(n):
                """Convert numbers less than 1000"""
                if n == 0:
                    return ''
                hundred = n // 100
                remainder = n % 100
                result = ''
                if hundred > 0:
                    result += convert_under_100(hundred) + 'শত'
                    if remainder > 0:
                        result += ' '
                if remainder > 0:
                    result += convert_under_100(remainder)
                return result
            
            # Handle decimal part (paise)
            num_str = str(num_float)
            parts = num_str.split('.')
            integer_part = int(float(parts[0]))
            decimal_part = 0
            if len(parts) > 1:
                decimal_str = parts[1].ljust(2, '0')[:2]
                decimal_part = int(decimal_str)
            
            if integer_part == 0 and decimal_part == 0:
                return 'শূন্য টাকা'
            
            result = ''
            n = integer_part
            
            # Crore (1,00,00,000)
            if n >= 10000000:
                crore = n // 10000000
                result += convert_under_100(crore) + ' কোটি'
                n %= 10000000
                if n > 0:
                    result += ' '
            
            # Lakh (1,00,000)
            if n >= 100000:
                lakh = n // 100000
                result += convert_under_100(lakh) + ' লক্ষ'
                n %= 100000
                if n > 0:
                    result += ' '
            
            # Thousand (1,000)
            if n >= 1000:
                thousand = n // 1000
                result += convert_under_100(thousand) + ' হাজার'
                n %= 1000
                if n > 0:
                    result += ' '
            
            # Hundreds and below
            if n > 0:
                result += convert_under_1000(n)
            
            result = result.strip()
            if not result:
                result = 'শূন্য'
            
            result += ' টাকা'
            
            # Add paise if decimal part exists
            if decimal_part > 0:
                result += ' এবং '
                paise_words = convert_under_100(decimal_part)
                result += paise_words + ' পয়সা'
            
            return result
        except Exception as e:
            current_app.logger.error(f'Error converting number to Bengali words: {e}', exc_info=True)
            # Fallback to number if conversion fails
            try:
                return f'{int(float(num))} টাকা'
            except:
                return ''

    def _is_postgraduate_course(course_value, year=None):
        """Check if a course is postgraduate (mirrors JS isPostgraduateCourse logic)"""
        if not course_value:
            return False
        
        course_code = str(course_value).split(' - ')[0].strip()
        
        # Method 1: Check course code pattern - year digit 5 typically means LLM (PG)
        import re
        digits = re.findall(r'\d+', course_code)
        if digits and len(digits) > 0:
            last_digits = digits[-1]
            # IMPORTANT: UG course codes are often 3 digits (e.g. "042").
            # In that case, don't let program year (e.g. LLM) force PG.
            if len(last_digits) <= 3:
                return False
            if len(last_digits) >= 4:
                year_digit = int(last_digits[-4]) if len(last_digits) >= 4 else None
                if year_digit == 5:
                    return True
        
        # Method 2: Check if course code contains PG indicators
        # IMPORTANT: don't use generic '5' substring here (causes false positives).
        pg_indicators = ['LLM', 'MPHIL', 'M.PHIL', 'PHD', 'PG']
        upper_code = course_code.upper()
        if any(indicator in upper_code for indicator in pg_indicators):
            return True
        
        # Method 3: Check year field if provided
        if year:
            year_upper = str(year).upper()
            if any(token in year_upper for token in ['LLM', 'PG', 'MPHIL', 'PHD']):
                return True
        
        return False

    def _matches_teacher_name(item_teacher, teacher_name):
        """Flexible teacher name matching (mirrors JS logic)"""
        if not item_teacher or not teacher_name:
            return False
        item_teacher = str(item_teacher).strip()
        teacher_name = str(teacher_name).strip()
        item_lower = item_teacher.lower()
        teacher_lower = teacher_name.lower()
        return (item_teacher == teacher_name or
                item_lower == teacher_lower or
                item_lower in teacher_lower or
                teacher_lower in item_lower)

    def _build_jobs_from_statement(statement_data, teacher_name, year, term, session):
        """Build jobs_data list from statement data, mirroring JS auto-populate logic"""
        if not statement_data:
            return [], 0.0
        
        # Debug: Log statement data structure
        try:
            current_app.logger.info(f'_build_jobs_from_statement called for teacher: {teacher_name}')
            current_app.logger.info(f'Statement data keys: {list(statement_data.keys()) if isinstance(statement_data, dict) else "Not a dict"}')
            if isinstance(statement_data, dict) and 'invigilation' in statement_data:
                invigilation_data = statement_data.get('invigilation', [])
                current_app.logger.info(f'Invigilation array exists with {len(invigilation_data)} items')
                # Log first few items for debugging
                for idx, inv_item in enumerate(invigilation_data[:3]):  # Log first 3 items
                    inv_name = inv_item.get('name') or inv_item.get('teacher') or 'Unknown'
                    inv_chief = inv_item.get('chief')
                    inv_invig = inv_item.get('invigilation')
                    current_app.logger.info(f'  Sample item {idx}: name="{inv_name}", chief={repr(inv_chief)}, invigilation={repr(inv_invig)}')
            else:
                current_app.logger.warning(f'Invigilation array NOT found in statement data!')
        except Exception as e:
            current_app.logger.error(f'Error logging statement data structure: {e}', exc_info=True)
        
        jobs_data = []
        total_amount = 0.0
        
        def safe_int(value, default=0):
            """Safely convert value to int, handling None, empty strings, dashes, and various formats"""
            if value is None:
                return default
            if isinstance(value, (int, float)):
                return int(value)
            if isinstance(value, str):
                value = value.strip()
                # Handle empty string, dash, or other non-numeric values
                if not value or value == '' or value == '-' or value.lower() == 'none':
                    return default
                try:
                    return int(float(value))
                except (ValueError, TypeError):
                    return default
            try:
                return int(float(value))
            except (ValueError, TypeError):
                return default
        
        def safe_float(value, default=0.0):
            try:
                return float(value)
            except Exception:
                return default
        
        def is_pg_year():
            yr = (year or '').upper()
            return any(token in yr for token in ['LLM', 'MASTER', 'M.PHIL', 'MPHIL', 'PHD'])
        
        is_pg = is_pg_year()
        
        # Row 1: প্রশ্নপত্র প্রণয়ন (Question Preparation)
        items = [i for i in statement_data.get('question_preparation', []) 
                 if _matches_teacher_name(i.get('teacher'), teacher_name)]
        if items:
            courses = []
            qty = 0
            for i in items:
                course = i.get('course') or i.get('course_code') or ''
                # Extract course code only (remove course name if present in "CODE - Name" format)
                if ' - ' in course:
                    course = course.split(' - ')[0].strip()
                section = i.get('section') or ''
                if course:
                    course_text = f"{course} ({section})" if section else course
                    courses.append(course_text)
                qty += max(safe_int(i.get('questions', 1)), 1)
            
            # Determine rate based on first course (PG/UG)
            first_course = items[0].get('course') or items[0].get('course_code') or ''
            is_pg_course = _is_postgraduate_course(first_course, year)
            rate = 2400 if is_pg_course else 2300
            amount = qty * rate
            total_amount += amount
            jobs_data.append({
                'description': 'প্রশ্নপত্র প্রণয়ন',
                'courses': ', '.join(courses),
                'quantity': str(qty),
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # Row 2: প্রশ্নপত্র মডারেশন (Question Moderation) - Only for moderation committee members
        # Check if teacher is in moderation committee (mirrors JS isTeacherInModerationCommittee)
        moderation_committee = statement_data.get('moderation_committee', [])
        is_in_moderation_committee = any(
            _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)
            for i in moderation_committee
        )
        if is_in_moderation_committee:
            # When "All" is selected, quantity is 1 (via data-actual-quantity in JS)
            qty = 1
            rate = 2400  # সর্বোচ্চ
            amount = qty * rate
            total_amount += amount
            jobs_data.append({
                'description': 'প্রশ্নপত্র মডারেশন',
                'courses': 'All',  # "All" courses
                'quantity': 'All',  # Display as "All", but calculated as 1
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # Row 3: উত্তরপত্র পরীক্ষণ (Script Examination) - with minimum 600 per course
        items = [i for i in statement_data.get('script_examination', []) 
                 if _matches_teacher_name(i.get('teacher'), teacher_name)]
        if items:
            courses = []
            total_scripts = 0
            per_course_amounts = []
            
            # Group by course-section for minimum calculation
            course_section_map = {}
            for i in items:
                course = i.get('course') or i.get('course_code') or ''
                # Extract course code only (remove course name if present in "CODE - Name" format)
                if ' - ' in course:
                    course = course.split(' - ')[0].strip()
                section = (i.get('section') or 'Full').upper()
                key = f"{course}_{section}"
                if key not in course_section_map:
                    course_section_map[key] = {
                        'course': course,
                        'section': section,
                        'scripts': 0
                    }
                course_section_map[key]['scripts'] += safe_int(i.get('scripts', 0))
            
            for key, data in course_section_map.items():
                course = data['course']
                section = data['section']
                scripts = data['scripts']
                if scripts == 0:
                    continue
                
                if course:
                    course_text = f"{course} ({section})" if section and section != 'FULL' else course
                    courses.append(course_text)
                
                total_scripts += scripts
                
                # Determine rate based on course (PG/UG) and section
                is_pg_course = _is_postgraduate_course(course, year)
                if section == 'FULL' or section == '':
                    rate = 160 if is_pg_course else 80  # Full paper: PG 160, UG 80
                else:
                    rate = 100 if is_pg_course else 80  # Half paper: PG 100, UG 80
                
                # Calculate with minimum 600 per course
                calculated_amount = scripts * rate
                course_amount = max(calculated_amount, 600)
                per_course_amounts.append(course_amount)
            
            if total_scripts > 0:
                amount = sum(per_course_amounts) if per_course_amounts else 0
                total_amount += amount
                
                # Build breakdown string like single PDF: "4 × 80 = 320" for each course
                breakdowns = []
                rates_used = {}  # Track rates and their script counts
                for key, data in course_section_map.items():
                    if data['scripts'] == 0:
                        continue
                    scripts = data['scripts']
                    is_pg_course = _is_postgraduate_course(data['course'], year)
                    section = data['section']
                    if section == 'FULL' or section == '':
                        rate = 160 if is_pg_course else 80
                    else:
                        rate = 100 if is_pg_course else 80
                    # Track rate usage with script count
                    if rate not in rates_used:
                        rates_used[rate] = 0
                    rates_used[rate] += scripts
                    calculated = scripts * rate
                    final_amount = max(calculated, 600)
                    if calculated < 600:
                        breakdowns.append(f'{scripts} × {rate} = {calculated} < 600 → 600')
                    else:
                        breakdowns.append(f'{scripts} × {rate} = {final_amount}')
                
                quantity_display = ' '.join(breakdowns) if breakdowns else str(total_scripts)
                
                # For rate display: if all courses use same rate, show that rate; otherwise show primary rate
                # Match single PDF format - show actual rate used, not "80/100/160"
                if len(rates_used) == 1:
                    rate_display = str(list(rates_used.keys())[0])
                else:
                    # Multiple rates used - show the rate used for the majority of scripts
                    primary_rate = max(rates_used.items(), key=lambda x: x[1])[0]
                    rate_display = str(primary_rate)
                
                jobs_data.append({
                    'description': 'উত্তরপত্র পরীক্ষণ',
                    'courses': ', '.join(courses),
                    'quantity': quantity_display,
                    'paper_type': '',
                    'rate': rate_display,  # Show actual rate used, matching single PDF format
                    'amount': f'{amount:.2f}'
                })
        
        # Row 4: ক্লাস টেস্ট/টার্ম পেপার (Class Test)
        # Row 4 uses multiplier: Full = 4, A or B = 2
        items = [i for i in statement_data.get('class_test', []) 
                 if _matches_teacher_name(i.get('teacher') or i.get('teacher_name'), teacher_name)]
        if items:
            courses = []
            total_quantity = 0  # This will be sum of (students × multiplier) for each course-section
            
            # Group by course-section to handle duplicates properly
            course_section_map = {}
            for i in items:
                course = i.get('course') or i.get('course_code') or ''
                # Extract course code only (remove course name if present in "CODE - Name" format)
                if ' - ' in course:
                    course = course.split(' - ')[0].strip()
                section = (i.get('section') or 'Full').upper()
                students = safe_int(i.get('students') or i.get('student_count', 0))
                
                if not course or students == 0:
                    continue
                
                key = f"{course}_{section}"
                if key not in course_section_map:
                    course_section_map[key] = {
                        'course': course,
                        'section': section,
                        'students': 0
                    }
                # Sum students if same course-section appears multiple times
                course_section_map[key]['students'] += students
            
            # Process each course-section pair
            for key, data in course_section_map.items():
                course = data['course']
                section = data['section']
                students = data['students']
                
                # Determine multiplier based on section
                if section == 'FULL' or section == '':
                    multiplier = 4  # Full section
                else:
                    multiplier = 2  # A or B section
                
                # Calculate quantity for this course-section: students × multiplier
                course_quantity = students * multiplier
                total_quantity += course_quantity
                
                # Add to courses list
                if section and section != 'FULL':
                    course_text = f"{course} ({section})"
                else:
                    course_text = course
                courses.append(course_text)
            
            if total_quantity > 0:
                # Determine rate based on first course (PG/UG)
                first_item = items[0]
                first_course = first_item.get('course') or first_item.get('course_code') or ''
                is_pg_course = _is_postgraduate_course(first_course, year)
                rate = 40 if is_pg_course else 30
                
                # Calculate amount: total_quantity × rate
                amount = total_quantity * rate
                total_amount += amount
                
                # Build breakdown string like single PDF: "10 × 4 = 40" for each course-section
                breakdowns = []
                for key, data in course_section_map.items():
                    if data['students'] == 0:
                        continue
                    students = data['students']
                    section = data['section']
                    multiplier = 4 if (section == 'FULL' or section == '') else 2
                    product = students * multiplier
                    breakdowns.append(f'{students} × {multiplier} = {product}')
                
                quantity_display = ' '.join(breakdowns) if breakdowns else str(total_quantity)
                
                jobs_data.append({
                    'description': 'ক্লাস টেস্ট/টার্ম পেপার/ হোম ওয়ার্ক/ এ্যাসাইনমেন্ট',
                    'courses': ', '.join(courses),
                    'quantity': quantity_display,  # Show breakdown like "10 × 4 = 40"
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
        
        # Row 5: সেশনাল (Sessional Assessment)
        items = [i for i in statement_data.get('sessional_assessment', []) 
                 if _matches_teacher_name(i.get('teacher'), teacher_name)]
        if items:
            courses = []
            total_students = 0
            for i in items:
                course = i.get('course') or i.get('course_code') or ''
                # Extract course code only (remove course name if present in "CODE - Name" format)
                if ' - ' in course:
                    course = course.split(' - ')[0].strip()
                students = safe_int(i.get('students', 0))
                if course:
                    courses.append(course)
                total_students += students
            
            rate = 230  # Default: প্রজেক্ট পেপার/এ্যাসাইনমেন্ট
            amount = total_students * rate
            total_amount += amount
            jobs_data.append({
                'description': 'সেশনাল',
                'courses': ', '.join(courses),
                'quantity': str(total_students),
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # Row 6: সেশনাল মৌখিক পরীক্ষা (Sessional Viva)
        items = [i for i in statement_data.get('sessional_viva', []) 
                 if _matches_teacher_name(i.get('teacher'), teacher_name)]
        if items:
            courses = []
            total_students = 0
            for i in items:
                course = i.get('course') or i.get('course_code') or ''
                # Extract course code only (remove course name if present in "CODE - Name" format)
                if ' - ' in course:
                    course = course.split(' - ')[0].strip()
                students = safe_int(i.get('students', 0))
                if course:
                    courses.append(course)
                total_students += students
            
            rate = 50
            amount = total_students * rate
            total_amount += amount
            jobs_data.append({
                'description': 'সেশনাল মৌখিক পরীক্ষা',
                'courses': ', '.join(courses),
                'quantity': str(total_students),
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # Row 7: প্রফেশনাল এ্যাটাসমেন্ট (Professional Attachment)
        items = [i for i in statement_data.get('professional_attachment', []) 
                 if _matches_teacher_name(i.get('teacher') or i.get('name'), teacher_name)]
        if items:
            courses = []
            total_count = 0
            for i in items:
                course = i.get('course') or i.get('course_code') or ''
                # Extract course code only (remove course name if present in "CODE - Name" format)
                if ' - ' in course:
                    course = course.split(' - ')[0].strip()
                count = safe_int(i.get('count', 0))
                if course:
                    courses.append(course)
                total_count += count
            
            rate = 100  # সুপারভিশন ও রিপোর্ট পরীক্ষণ
            amount = total_count * rate
            total_amount += amount
            jobs_data.append({
                'description': 'প্রফেশনাল এ্যাটাসমেন্ট/ইন্ডাস্ট্রিয়াল (ট্রেনিং/এ্যাটাসমেন্ট)',
                'courses': ', '.join(courses),
                'quantity': str(total_count),
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # Row 8: উত্তরপত্র নিরীক্ষণ (Script Scrutiny)
        items = [i for i in statement_data.get('script_scrutiny', []) 
                 if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        if items:
            total_scripts = sum(safe_int(i.get('scripts', 0)) for i in items)
            if total_scripts > 0:
                rate = 8
                amount = total_scripts * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'উত্তরপত্র নিরীক্ষণ',
                    'courses': '',
                    'quantity': str(total_scripts),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
        
        # Row 9: টেবুলেশন (Tabulation) - course_wise and student_wise separately
        items = [i for i in statement_data.get('tabulation', []) 
                 if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        if items:
            total_course_wise = sum(safe_int(i.get('course_wise', 0)) for i in items)
            total_student_wise = sum(safe_int(i.get('student_wise', 0)) for i in items)
            
            if total_course_wise > 0:
                rate = 200
                amount = total_course_wise * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'টেবুলেশন (কোর্স ভিত্তিক)',
                    'courses': '',
                    'quantity': str(total_course_wise),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
            
            if total_student_wise > 0:
                rate = 40
                amount = total_student_wise * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'টেবুলেশন (পরীক্ষার্থী ভিত্তিক)',
                    'courses': '',
                    'quantity': str(total_student_wise),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
        
        # Row 10: প্রশ্নপত্র প্রস্তুতকরণ (Question Typing/Photocopy)
        # 10a: Drawing
        items = [i for i in statement_data.get('question_typing', []) 
                 if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        total_drawing = sum(safe_int(i.get('questions', 0)) for i in items)
        if total_drawing > 0:
            rate = 250
            amount = total_drawing * rate
            total_amount += amount
            jobs_data.append({
                'description': 'প্রশ্নপত্র প্রস্তুতকরণ (অংকন)',
                'courses': '',
                'quantity': str(total_drawing),
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # 10b: Photocopy
        photocopy_items = statement_data.get('question_photocopy', [])
        if not photocopy_items:
            # If question_photocopy doesn't exist, use question_typing
            photocopy_items = items
        photocopy_items = [i for i in photocopy_items 
                          if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        total_photocopy = sum(safe_int(i.get('questions', 0)) for i in photocopy_items)
        if total_photocopy > 0:
            rate = 7
            amount = total_photocopy * rate
            total_amount += amount
            jobs_data.append({
                'description': 'প্রশ্নপত্র প্রস্তুতকরণ (ফটোকপি)',
                'courses': '',
                'quantity': str(total_photocopy),
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # Row 11: পরীক্ষা কমিটির সভাপতি/সদস্য (Exam Committee)
        items = [i for i in statement_data.get('examination_committee', []) 
                 if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        if items:
            # Check if teacher is chairman/chief
            position = (items[0].get('position') or '').lower()
            is_chairman = any(keyword in position for keyword in ['chief', 'chairman', 'সভাপতি', 'চীফ'])
            
            qty = 1
            if is_pg and is_chairman:
                rate = 3000  # স্নাতকোত্তর - সভাপতি
            elif is_pg and not is_chairman:
                rate = 1000  # স্নাতকোত্তর - সদস্য
            elif not is_pg and is_chairman:
                rate = 2500  # স্নাতক - সভাপতি
            else:
                rate = 1000  # স্নাতক - সদস্য
            
            amount = qty * rate
            total_amount += amount
            jobs_data.append({
                'description': 'পরীক্ষা কমিটির সভাপতি/সদস্য',
                'courses': '',
                'quantity': str(qty),
                'paper_type': '',
                'rate': str(rate),
                'amount': f'{amount:.2f}'
            })
        
        # Row 12: চীফ ইনভিজিলেশন / ইনভিজিলেশন (Invigilation)
        # Match individual download logic exactly: use 'name' or 'teacher' field, handle chief and invigilation separately
        invigilation_data = statement_data.get('invigilation', [])
        
        # Debug: Log all invigilation data first
        try:
            current_app.logger.info(f'Row 12 for {teacher_name}: Total invigilation items in statement: {len(invigilation_data)}')
            if invigilation_data:
                for idx, inv_item in enumerate(invigilation_data):
                    inv_name = inv_item.get('name') or inv_item.get('teacher') or 'Unknown'
                    inv_chief = inv_item.get('chief')
                    inv_invig = inv_item.get('invigilation')
                    current_app.logger.info(f'  Statement item {idx}: name="{inv_name}", chief={repr(inv_chief)} (type={type(inv_chief).__name__}), invigilation={repr(inv_invig)} (type={type(inv_invig).__name__})')
            else:
                current_app.logger.warning(f'Row 12 for {teacher_name}: No invigilation data found in statement!')
        except Exception as e:
            current_app.logger.error(f'Error logging all invigilation data: {e}', exc_info=True)
        
        items = [i for i in invigilation_data 
                 if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        
        # Debug: Log matching results
        try:
            current_app.logger.info(f'Row 12 for {teacher_name}: Matched {len(items)} items after name matching (searching for: "{teacher_name}")')
            if len(items) == 0 and invigilation_data:
                # Log all names for debugging
                all_names = [i.get('name') or i.get('teacher') or 'Unknown' for i in invigilation_data]
                current_app.logger.info(f'  Available names in statement: {all_names}')
        except Exception:
            pass
        
        if items:
            # Debug: Log raw values before conversion
            try:
                current_app.logger.info(f'Row 12 for {teacher_name}: Processing {len(items)} matched invigilation items')
                for idx, item in enumerate(items):
                    raw_chief = item.get('chief')
                    raw_invig = item.get('invigilation')
                    item_name = item.get('name') or item.get('teacher') or 'Unknown'
                    current_app.logger.info(f'  Matched item {idx} ({item_name}): chief={repr(raw_chief)} (type={type(raw_chief).__name__}), invigilation={repr(raw_invig)} (type={type(raw_invig).__name__})')
            except Exception as e:
                current_app.logger.error(f'Error logging matched invigilation items: {e}')
            
            # Calculate chief count - EXACTLY mirror JS: parseInt(item.chief || 0) || 0
            # JS logic: totalChief += parseInt(item.chief || 0) || 0;
            chief_count = 0
            for i in items:
                chief_val = i.get('chief')
                # Step 1: item.chief || 0 (in JS, this means: if falsy, use 0)
                # Falsy values in JS: null, undefined, '', 0, false, NaN
                if not chief_val or chief_val == '' or chief_val == '-' or chief_val == 0:
                    chief_val = 0
                else:
                    # Step 2: parseInt(value) - convert to int
                    # Step 3: result || 0 - if result is falsy (0, NaN), use 0
                    parsed = safe_int(chief_val, 0)
                    chief_val = parsed if parsed > 0 else 0
                
                # Add to total (JS does: totalChief += ...)
                chief_count += chief_val
                
                try:
                    original = i.get('chief')
                    current_app.logger.info(f'  Chief calculation: original={repr(original)}, parsed={chief_val}, total_chief={chief_count}')
                except:
                    pass
            
            # Calculate invigilation count - EXACTLY mirror JS: parseInt(item.invigilation || 0) || 0
            # JS logic: totalInvigilation += parseInt(item.invigilation || 0) || 0;
            inv_count = 0
            for i in items:
                inv_val = i.get('invigilation')
                # Step 1: item.invigilation || 0
                if not inv_val or inv_val == '' or inv_val == '-' or inv_val == 0:
                    inv_val = 0
                else:
                    # Step 2: parseInt(value), Step 3: result || 0
                    parsed = safe_int(inv_val, 0)
                    inv_val = parsed if parsed > 0 else 0
                
                # Add to total
                inv_count += inv_val
            
            # Debug: Log calculated counts
            try:
                current_app.logger.info(f'Row 12 for {teacher_name}: FINAL - chief_count={chief_count}, inv_count={inv_count}')
            except Exception:
                pass
            
            if chief_count > 0:
                rate = 1800
                amount = chief_count * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'চীফ ইনভিজিলেশন',
                    'courses': '',
                    'quantity': str(chief_count),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
                try:
                    current_app.logger.info(f'Row 12: Added chief invigilation for {teacher_name}: {chief_count} × {rate} = {amount:.2f}')
                except Exception:
                    pass
            
            if inv_count > 0:
                rate = 1500
                amount = inv_count * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'ইনভিজিলেশন',
                    'courses': '',
                    'quantity': str(inv_count),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
        
        # Row 13: থিসিস (Thesis Supervision)
        items = [i for i in statement_data.get('thesis_supervision', []) 
                 if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        if items:
            total_examine = sum(safe_int(i.get('examine', 0)) for i in items)
            total_supervision = sum(safe_int(i.get('supervision', 0)) for i in items)
            total_co_supervision = sum(safe_int(i.get('co_supervision', 0)) for i in items)
            total_viva = sum(safe_int(i.get('viva', 0)) for i in items)
            
            # 13a: পরীক্ষণ (Examine)
            if total_examine > 0:
                rate = 2500 if is_pg else 1200
                amount = total_examine * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'থিসিস পরীক্ষণ',
                    'courses': '',
                    'quantity': str(total_examine),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
            
            # 13b: সুপারভিশন (Supervision)
            if total_supervision > 0:
                rate = 5000 if is_pg else 2000
                amount = total_supervision * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'থিসিস সুপারভিশন',
                    'courses': '',
                    'quantity': str(total_supervision),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
            
            # 13c: কো-সুপারভিশন (Co-Supervision) - Only PG
            if total_co_supervision > 0 and is_pg:
                rate = 1500
                amount = total_co_supervision * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'থিসিস কো-সুপারভিশন',
                    'courses': '',
                    'quantity': str(total_co_supervision),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
            
            # 13d: মৌখিক পরীক্ষা (Viva)
            if total_viva > 0:
                rate = 500 if is_pg else 120
                amount = total_viva * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'থিসিস মৌখিক পরীক্ষা',
                    'courses': '',
                    'quantity': str(total_viva),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
        
        # Row 14: ভাইভা (Viva)
        items = [i for i in statement_data.get('viva', []) 
                 if _matches_teacher_name(i.get('teacher') or i.get('name'), teacher_name)]
        if items:
            total_students = sum(safe_int(i.get('students', 0)) for i in items)
            if total_students > 0:
                rate = 50
                amount = total_students * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'ভাইভা',
                    'courses': '',
                    'quantity': str(total_students),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
        
        # Row 15: কোডিং/ডিকোডিং (Coding/Decoding)
        items = [i for i in statement_data.get('coding_decoding', []) 
                 if _matches_teacher_name(i.get('name') or i.get('teacher'), teacher_name)]
        if items:
            total_scripts = sum(safe_int(i.get('scripts', 0)) for i in items)
            if total_scripts > 0:
                rate = 30  # পরীক্ষার্থী প্রতি
                amount = total_scripts * rate
                total_amount += amount
                jobs_data.append({
                    'description': 'কোডিং/ডিকোডিং',
                    'courses': '',
                    'quantity': str(total_scripts),
                    'paper_type': '',
                    'rate': str(rate),
                    'amount': f'{amount:.2f}'
                })
        
        # Format jobs_data to match single PDF structure (with main rows and sub-items)
        formatted_jobs_data = []
        serial = 1
        
        # Define job order and structure - matching single PDF exactly
        job_structure = [
            {'row': 1, 'desc': 'প্রশ্নপত্র প্রণয়ন'},
            {'row': 2, 'desc': 'প্রশ্নপত্র মডারেশন'},
            {'row': 3, 'desc': 'উত্তরপত্র পরীক্ষণ'},
            {'row': 4, 'desc': 'ক্লাস টেস্ট/টার্ম পেপার/ হোম ওয়ার্ক/ এ্যাসাইনমেন্ট'},
            {'row': 5, 'desc': 'সেশনাল'},
            {'row': 6, 'desc': 'সেশনাল মৌখিক পরীক্ষা'},
            {'row': 7, 'desc': 'প্রফেশনাল এ্যাটাসমেন্ট/ইন্ডাস্ট্রিয়াল (ট্রেনিং/এ্যাটাসমেন্ট)'},
            {'row': 8, 'desc': 'উত্তরপত্র নিরীক্ষণ'},
            {'row': 9, 'desc': 'টেবুলেশন', 'has_subs': True, 'subs': [
                {'key': '9a', 'desc': 'কোর্স ভিত্তিক'},
                {'key': '9b', 'desc': 'পরীক্ষার্থী ভিত্তিক'}
            ]},
            {'row': 10, 'desc': 'প্রশ্নপত্র প্রস্তুতকরণ (অংকন, স্টেনসিল কাটা ও ঘুরানো)', 'has_subs': True, 'subs': [
                {'key': '10a', 'desc': 'অংকন'},
                {'key': '10b', 'desc': 'ফটোকপি'}
            ]},
            {'row': 11, 'desc': 'পরীক্ষা কমিটির সভাপতি/সদস্য'},
            {'row': 12, 'desc': 'চীফ ইনভিজিলেশন / ইনভিজিলেশন', 'has_subs': True, 'subs': [
                {'key': '12a', 'desc': 'চীফ ইনভিজিলেশন'},
                {'key': '12b', 'desc': 'ইনভিজিলেশন'}
            ]},
            {'row': 13, 'desc': 'থিসিস', 'has_subs': True, 'subs': [
                {'key': '13a', 'desc': 'পরীক্ষণ'},
                {'key': '13b', 'desc': 'সুপারভিশন (থিসিস/প্রজেক্ট থিসিস/ইন্টার্নশীপ রিপোর্ট)'},
                {'key': '13c', 'desc': 'কো-সুপারভিশন'},
                {'key': '13d', 'desc': 'মৌখিক পরীক্ষা'}
            ]},
            {'row': 14, 'desc': 'ভাইভা'},
            {'row': 15, 'desc': 'কোডিং/ডিকোডিং'},
            {'row': 16, 'desc': 'অন্যান্য'}
        ]
        
        # Create a map of current jobs_data by description/key
        jobs_map = {}
        for job in jobs_data:
            desc = job.get('description', '').strip()
            # Handle sub-items (they have specific descriptions)
            if 'কোর্স ভিত্তিক' in desc:
                jobs_map['9a'] = job
            elif 'পরীক্ষার্থী ভিত্তিক' in desc:
                jobs_map['9b'] = job
            elif 'অংকন' in desc and 'প্রশ্নপত্র প্রস্তুতকরণ' in desc:
                jobs_map['10a'] = job
            elif 'ফটোকপি' in desc and 'প্রশ্নপত্র প্রস্তুতকরণ' in desc:
                jobs_map['10b'] = job
            elif desc == 'চীফ ইনভিজিলেশন':
                jobs_map['12a'] = job
            elif desc == 'ইনভিজিলেশন':
                jobs_map['12b'] = job
            elif 'থিসিস পরীক্ষণ' in desc:
                jobs_map['13a'] = job
            elif 'থিসিস সুপারভিশন' in desc:
                jobs_map['13b'] = job
            elif 'থিসিস কো-সুপারভিশন' in desc:
                jobs_map['13c'] = job
            elif 'থিসিস মৌখিক পরীক্ষা' in desc:
                jobs_map['13d'] = job
            else:
                # Main items - use description as key (exact match)
                jobs_map[desc] = job
        
        # Log all jobs_data for debugging
        try:
            current_app.logger.debug(f'Jobs data for {teacher_name}: {[j.get("description") for j in jobs_data]}')
            current_app.logger.debug(f'Jobs map keys for {teacher_name}: {list(jobs_map.keys())}')
        except Exception:
            pass
        
        # Build formatted jobs_data matching single PDF structure
        for job_def in job_structure:
            row_num = job_def['row']
            main_desc = job_def['desc']
            has_subs = job_def.get('has_subs', False)
            
            # Check if we have data for this row
            has_data = False
            if has_subs:
                # Check if any sub-item has data
                for sub in job_def.get('subs', []):
                    if sub['key'] in jobs_map:
                        has_data = True
                        break
            else:
                # Check if main item has data - try exact match first, then flexible matching
                if main_desc in jobs_map:
                    has_data = True
                else:
                    # Try flexible matching (strip whitespace)
                    for key in list(jobs_map.keys()):
                        if key.strip() == main_desc.strip():
                            # Update jobs_map with correct key
                            jobs_map[main_desc] = jobs_map.pop(key)
                            has_data = True
                            break
            
            # Log for Row 15 specifically for debugging
            if row_num == 15:
                try:
                    current_app.logger.info(
                        f'Row 15 check for {teacher_name}: main_desc="{main_desc}", '
                        f'has_data={has_data}, in_jobs_map={main_desc in jobs_map}, '
                        f'all_keys={list(jobs_map.keys())}'
                    )
                except Exception:
                    pass
            
            if not has_data:
                continue  # Skip rows with no data
            
            # Add main row (for rows 9, 10, 12, 13, main row has empty fields)
            if has_subs:
                formatted_jobs_data.append({
                    'serial': str(serial),
                    'description': main_desc,
                    'courses': '',
                    'quantity': '',
                    'paper_type': '',
                    'rate': '',
                    'amount': ''
                })
            else:
                # Regular row - use data from jobs_map
                job = jobs_map.get(main_desc, {})
                formatted_jobs_data.append({
                    'serial': str(serial),
                    'description': main_desc,
                    'courses': job.get('courses', ''),
                    'quantity': job.get('quantity', ''),
                    'paper_type': job.get('paper_type', ''),
                    'rate': job.get('rate', ''),
                    'amount': job.get('amount', '')
                })
            
            serial += 1
            
            # Add sub-items if applicable
            if has_subs:
                for sub in job_def.get('subs', []):
                    sub_key = sub['key']
                    sub_desc = sub['desc']
                    if sub_key in jobs_map:
                        sub_job = jobs_map[sub_key]
                        formatted_jobs_data.append({
                            'serial': '',  # Sub-items have empty serial
                            'description': sub_desc,
                            'courses': sub_job.get('courses', ''),
                            'quantity': sub_job.get('quantity', ''),
                            'paper_type': sub_job.get('paper_type', ''),
                            'rate': sub_job.get('rate', ''),
                            'amount': sub_job.get('amount', '')
                        })
        
        try:
            current_app.logger.info(
                f'Built {len(formatted_jobs_data)} formatted jobs from statement for {teacher_name}, total={total_amount:.2f}'
            )
        except Exception:
            # If logger is not available, continue without logging
            pass
        
        return formatted_jobs_data, total_amount

    @app.route('/remuneration/bulk-download', methods=['POST'])
    @login_required
    def remuneration_bulk_download():
        """Bulk download auto-generated remuneration forms for selected teachers as a ZIP file"""
        try:
            restriction = _require_teacher_privileges()
            if restriction:
                return restriction
            
            import zipfile
            import io
            import os
            from datetime import datetime
            import json
            
            # Get data from POST request
            try:
                data = request.get_json()
            except Exception as e:
                current_app.logger.error(f'Error parsing JSON request: {e}', exc_info=True)
                return jsonify({
                    'success': False,
                    'message': 'Invalid request data format'
                }), 400
            
            if not data:
                return jsonify({
                    'success': False,
                    'message': 'Invalid request data'
                }), 400
            
            session_val = data.get('session', '').strip()
            year_val = data.get('year', '').strip()
            term_val = data.get('term', '').strip()
            teacher_ids = data.get('teacher_ids', [])
            
            if not session_val or not year_val or not term_val:
                return jsonify({
                    'success': False,
                    'message': 'সেশন, বর্ষ ও টার্ম সিলেক্ট করুন'
                }), 400
            
            if not teacher_ids or len(teacher_ids) == 0:
                return jsonify({
                    'success': False,
                    'message': 'অন্তত একজন শিক্ষক নির্বাচন করুন'
                }), 400
            
            try:
                from blueprints.remuneration_management.models import RemunerationForm
                from blueprints.class_management.models import Teacher
                from sqlalchemy import or_, func
                
                # Normalize values
                session_val_clean = session_val.strip()
                year_val_clean = year_val.strip()
                term_val_clean = term_val.strip()
                
                current_app.logger.info(f'Bulk download: session="{session_val_clean}", year="{year_val_clean}", term="{term_val_clean}", teachers={teacher_ids}')
                
                # Fetch selected teachers
                try:
                    teacher_ids_int = [int(tid) for tid in teacher_ids]
                    teachers = Teacher.query.filter(Teacher.id.in_(teacher_ids_int)).all()
                except (ValueError, TypeError) as e:
                    current_app.logger.error(f'Invalid teacher IDs: {e}')
                    return jsonify({
                        'success': False,
                        'message': 'অবৈধ শিক্ষক ID'
                    }), 400
                
                if not teachers:
                    return jsonify({
                        'success': False,
                        'message': 'নির্বাচিত শিক্ষক পাওয়া যায়নি'
                    }), 404
                
                current_app.logger.info(f'Found {len(teachers)} teachers to generate forms for')
                
                # Fetch statement data first - required for bulk download
                statement_data = _get_remuneration_statement_data(session_val_clean, year_val_clean, term_val_clean)
                if not statement_data:
                    return jsonify({
                        'success': False,
                        'message': 'এই সেশন, বর্ষ ও টার্মের জন্য পরীক্ষা কমিটির চীফের Statement of Remuneration ফর্ম সেভ করা হয়নি। অনুগ্রহ করে প্রথমে Statement ফর্মটি সেভ করুন।'
                    }), 400
                
                current_app.logger.info(f'Using statement data for bulk download: session={session_val_clean}, year={year_val_clean}, term={term_val_clean}')
                
                # Check if WeasyPrint is available
                try:
                    from weasyprint import HTML, CSS
                except ImportError as import_err:
                    current_app.logger.error(f'WeasyPrint import error: {str(import_err)}', exc_info=True)
                    return jsonify({
                        'success': False,
                        'message': 'PDF জেনারেশন উপলব্ধ নেই'
                    }), 500
                
                # Create ZIP buffer
                zip_buffer = io.BytesIO()
                
                # Use ZIP_STORED for better compatibility, or ZIP_DEFLATED with compression level
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED, compresslevel=6) as zip_file:
                    successful_pdfs = 0
                    
                    for teacher in teachers:
                        try:
                            applicant_name = teacher.name
                            
                            # Build jobs_data from statement (ignoring saved forms for bulk output)
                            try:
                                jobs_data, total_amount = _build_jobs_from_statement(
                                    statement_data, teacher.name, year_val_clean, term_val_clean, session_val_clean
                                )
                            except Exception as build_error:
                                current_app.logger.error(
                                    f'Error building jobs from statement for {teacher.name}: {build_error}',
                                    exc_info=True
                                )
                                # Skip this teacher if we can't build their data
                                continue
                            
                            # If no jobs data, skip this teacher
                            if not jobs_data:
                                current_app.logger.info(f'No statement items found for {teacher.name}, skipping')
                                continue
                            
                            # jobs_data already has serial numbers from _build_jobs_from_statement
                            
                            # Get exam dates from statement
                            exam_start_date = statement_data.get('exam_start_date', '')
                            exam_end_date = statement_data.get('exam_end_date', '')
                            
                            # Get bank account from teacher
                            bank_account = teacher.bank_account_no or ''
                            
                            # Convert total to words
                            total_in_words = _number_to_words_bengali(total_amount) if total_amount > 0 else ''
                            
                            # Get font file path for Bengali font - same as single PDF
                            import os
                            font_path_absolute = None
                            font_paths_to_try = [
                                os.path.join(current_app.root_path, 'static', 'Fonts', 'kalpurush.ttf'),
                                os.path.join(current_app.root_path, 'static', 'fonts', 'kalpurush.ttf'),
                                os.path.join(current_app.root_path, 'static', 'Fonts', 'Kalpurush.ttf'),
                                os.path.join(current_app.root_path, 'static', 'fonts', 'Kalpurush.ttf'),
                            ]
                            
                            for font_path in font_paths_to_try:
                                if os.path.exists(font_path):
                                    font_path_absolute = os.path.abspath(font_path)
                                    break
                            
                            # Get logo file path and convert to data URI - same as single PDF
                            logo_path_absolute = None
                            logo_paths_to_try = [
                                os.path.join(current_app.root_path, 'static', 'Images', 'KU_logo_2.png'),
                                os.path.join(current_app.root_path, 'static', 'images', 'KU_logo_2.png'),
                            ]
                            
                            for logo_path in logo_paths_to_try:
                                if os.path.exists(logo_path):
                                    logo_path_absolute = os.path.abspath(logo_path)
                                    break
                            
                            # Convert logo to base64 data URI if found
                            logo_data_uri = None
                            if logo_path_absolute:
                                try:
                                    import base64
                                    with open(logo_path_absolute, 'rb') as logo_file:
                                        logo_data = logo_file.read()
                                        logo_base64 = base64.b64encode(logo_data).decode('utf-8')
                                        logo_data_uri = f'data:image/png;base64,{logo_base64}'
                                except Exception as e:
                                    current_app.logger.error(f'Error converting logo to data URI: {e}')
                                    logo_data_uri = None
                            
                            # Prepare context for template - same as single PDF
                            context = {
                                'font_path': font_path_absolute,
                                'logo_data_uri': logo_data_uri,
                                'voucher_no': '',
                                'voucher_date': '',
                                'applicant_name': applicant_name,
                                'designation': teacher.designation or '',
                                'discipline': 'Law',
                                'address': teacher.institute or 'Law Discipline, KU',
                                'exam_discipline': 'Law',
                                'year': year_val_clean,
                                'academic_year': session_val_clean,
                                'term': term_val_clean,
                                'exam_start_date': exam_start_date,
                                'exam_end_date': exam_end_date,
                                'jobs_data': jobs_data,
                                'total_amount': str(total_amount),
                                'total_in_words': total_in_words,
                                'bank_account': bank_account,
                                'auditor_sign': '',
                                'deputy_sign': '',
                                'controller_sign': '',
                                'finance_amount': '',
                                'finance_amount_words': '',
                                'section_officer_sign': '',
                                'deputy_director_sign': '',
                                'director_sign': '',
                                'audit_amount': '',
                                'audit_assistant_sign': '',
                                'audit_head_sign': '',
                                'bank_advice_no': '',
                                'payment_date': ''
                            }
                            
                            # Render HTML template
                            html_content = render_template('remuneration_pdf_template.html', **context)
                            
                            if not html_content:
                                current_app.logger.warning(f'Empty HTML content for {applicant_name}, skipping')
                                continue
                            
                            # Initialize html_content for font injection - same as single PDF
                            html_content_final = html_content
                            
                            # Inject font as base64 - same as single PDF
                            if font_path_absolute:
                                try:
                                    import base64
                                    with open(font_path_absolute, 'rb') as font_file:
                                        font_data = font_file.read()
                                        font_base64 = base64.b64encode(font_data).decode('utf-8')
                                    
                                    font_face_rule = f"""
        <style>
        @font-face {{
            font-family: 'Kalpurush';
            src: url(data:application/font-sfnt;base64,{font_base64}) format('truetype');
            font-weight: normal;
            font-style: normal;
        }}
        @font-face {{
            font-family: 'Kalpurush';
            src: url(data:application/font-sfnt;base64,{font_base64}) format('truetype');
            font-weight: bold;
            font-style: normal;
        }}
        </style>
        """
                                    
                                    if '</head>' in html_content_final:
                                        html_content_final = html_content_final.replace('</head>', font_face_rule + '</head>')
                                    elif '<head>' in html_content_final:
                                        html_content_final = html_content_final.replace('<head>', '<head>' + font_face_rule)
                                    else:
                                        html_content_final = font_face_rule + html_content_final
                                    
                                except Exception as e:
                                    current_app.logger.error(f'Failed to embed font as base64: {e}', exc_info=True)
                            
                            # Use the final HTML content (with font injected)
                            html_content = html_content_final
                            
                            # Generate PDF - same CSS as single PDF
                            pdf_buffer = io.BytesIO()
                            try:
                                html_obj = HTML(string=html_content, base_url=request.url_root)
                                
                                # CSS for legal size page - EXACT same as single PDF
                                css_string = """
            @page {
                size: 8.5in 14in; /* Legal size */
                margin: 0.18in 0.28in; /* Slightly increased margins */
            }
            * {
                page-break-inside: avoid !important;
                page-break-after: avoid !important;
                page-break-before: avoid !important;
            }
            body {
                margin: 0 !important;
                padding: 0 !important;
                font-size: 0.64rem !important; /* Increased from 0.56rem (7pt) to 0.64rem (8pt) */
                line-height: 1.18 !important; /* Slightly more breathing room */
                font-family: 'Kalpurush', sans-serif !important;
            }
            .rem-wrapper {
                margin: 0 !important;
                padding: 0 !important;
            }
            .rem-sheet {
                padding: 10px 17px !important; /* Slightly increased padding */
                margin: 0 !important;
                border: none !important;
                border-radius: 0 !important;
            }
            .rem-heading {
                margin-bottom: 0.15rem !important;
                gap: 0.7rem !important;
            }
            .rem-heading-logo img {
                height: 38px !important; /* Larger logo */
            }
            .rem-heading-content .text-muted {
                font-size: 0.62rem !important;
                margin-bottom: 0.04rem !important;
            }
            .rem-heading h4 {
                font-size: 0.78rem !important;
                margin-bottom: 0.015rem !important;
            }
            .rem-heading-content > div {
                font-size: 0.72rem !important;
                margin-bottom: 0.015rem !important;
            }
            .rem-heading-content small {
                font-size: 0.56rem !important;
                margin-top: 0.04rem !important;
            }
            .voucher-box {
                padding: 0.12rem 0.35rem 0.015rem 0.35rem !important;
                font-size: 0.56rem !important;
                width: 210px !important;
            }
            .voucher-box div {
                padding: 0.04rem 0 !important;
            }
            .voucher-box span {
                min-width: 52px !important;
                font-size: 0.56rem !important;
            }
            .voucher-box input {
                font-size: 0.56rem !important;
                padding: 0.06rem 0.16rem !important;
            }
            .meta-grid {
                margin-top: 0.25rem !important;
                margin-bottom: 0.25rem !important;
            }
            .meta-grid td {
                padding: 0.18rem 0.32rem !important;
                font-size: 0.56rem !important; /* Reverted to original 7pt */
            }
            .meta-label {
                font-size: 0.56rem !important; /* Reverted to original 7pt */
                width: 142px !important;
            }
            .meta-grid input,
            .meta-grid select {
                font-size: 0.56rem !important; /* Reverted to original 7pt */
                padding: 0.11rem 0.22rem !important;
            }
            .rem-table {
                margin: 0.25rem 0 !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
            }
            .rem-table th,
            .rem-table td {
                padding: 0.18rem 0.27rem !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                line-height: 1.12 !important;
            }
            .rem-table th {
                padding: 0.22rem 0.27rem !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
            }
            .section-title {
                margin-top: 0.65rem !important;
                margin-bottom: 0.35rem !important;
                font-size: 0.71rem !important;
            }
            .signature-box {
                min-height: 38px !important;
                padding: 0.22rem !important;
                font-size: 0.56rem !important;
            }
            .signature-box span {
                margin-top: 0.55rem !important;
                padding-top: 0.18rem !important;
            }
            .controller-signature-section,
            .finance-signature-section,
            .audit-approval-section {
                margin-top: 0.25rem !important;
                margin-bottom: 0.28rem !important;
                gap: 0.25rem !important;
            }
            .controller-signature-box,
            .finance-signature-box,
            .audit-signature-box-single {
                min-height: 32px !important;
                padding: 0.17rem 0.16rem !important;
            }
            .controller-designation,
            .finance-designation,
            .audit-designation {
                font-size: 0.51rem !important;
                margin-bottom: 0.27rem !important;
                padding-bottom: 0.16rem !important;
            }
            .controller-signature-line,
            .finance-signature-line,
            .audit-signature-line {
                margin-top: 0.22rem !important;
                font-size: 0.51rem !important;
            }
            .foot-table {
                margin: 0.25rem 0 0.18rem 0 !important;
            }
            .foot-table td {
                padding: 0.18rem 0.32rem !important;
                font-size: 0.56rem !important; /* Reverted to original 7pt */
            }
            .foot-table input {
                font-size: 0.56rem !important; /* Reverted to original 7pt */
                padding: 0.11rem 0.22rem !important;
            }
            .info-note,
            .statement-note,
            .finance-release-note,
            .audit-approval-text {
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                margin: 0.18rem 0 !important;
                line-height: 1.22 !important;
            }
            .info-note input,
            .statement-note input {
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                padding: 0.04rem 0.11rem !important;
            }
            .bank-declaration {
                padding: 0.27rem !important;
                margin-top: 0.25rem !important;
                font-size: 0.51rem !important; /* Reverted to original 6pt */
                line-height: 1.22 !important;
            }
            .revenue-ticket {
                width: 18mm !important;
                height: 18mm !important;
                margin-left: 0.35rem !important;
            }
            * {
                font-family: 'Kalpurush', sans-serif !important;
            }
            body {
                font-family: 'Kalpurush', sans-serif !important;
            }
            .english-text {
                font-family: 'Tahoma', 'Arial', sans-serif !important;
            }
            """
                                
                                css_obj = CSS(string=css_string)
                                html_obj.write_pdf(pdf_buffer, stylesheets=[css_obj], presentational_hints=True)
                                
                                # Reset buffer position
                                pdf_buffer.seek(0)
                                
                                # Get PDF content
                                pdf_content = pdf_buffer.getvalue()
                                
                                # Validate PDF content
                                if not pdf_content or len(pdf_content) == 0:
                                    current_app.logger.warning(f'Empty PDF generated for {applicant_name}, skipping')
                                    continue
                                
                                # Check if it's a valid PDF (should start with %PDF)
                                if not pdf_content.startswith(b'%PDF'):
                                    current_app.logger.warning(f'Invalid PDF content for {applicant_name} (starts with: {pdf_content[:20]}), skipping')
                                    continue
                                
                            except Exception as pdf_error:
                                current_app.logger.error(f'Error generating PDF for {applicant_name}: {str(pdf_error)}', exc_info=True)
                                continue
                            
                            # Sanitize filename - ensure it's valid for ZIP
                            safe_name = ''.join(c if c.isalnum() or c in ' _-' else '_' for c in applicant_name)
                            filename = f"{safe_name}_{session_val_clean}_{year_val_clean}_{term_val_clean}.pdf"
                            
                            # Add to ZIP - use writestr directly (simpler and more reliable)
                            zip_file.writestr(filename, pdf_content)
                            successful_pdfs += 1
                            
                            current_app.logger.info(f'Added PDF to ZIP: {filename} ({len(pdf_content)} bytes)')
                            
                            current_app.logger.info(f'Generated PDF for {applicant_name}: {len(jobs_data)} jobs, total={total_amount:.2f}')
                            
                        except Exception as form_error:
                            current_app.logger.error(f'Error generating PDF for {applicant_name}: {str(form_error)}', exc_info=True)
                            continue
                    
                    if successful_pdfs == 0:
                        return jsonify({
                            'success': False,
                            'message': 'কোনো PDF জেনারেট করা যায়নি'
                        }), 500
                
                # ZIP file is automatically closed and finalized when exiting the 'with' block
                # Reset buffer position to beginning before reading
                zip_buffer.seek(0)
                
                # Get the ZIP content
                zip_content = zip_buffer.getvalue()
                
                # Generate filename for ZIP
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                zip_filename = f"Remuneration_{session_val_clean}_{year_val_clean}_{term_val_clean}_{timestamp}.zip"
                
                zip_size = len(zip_content)
                
                current_app.logger.info(f'Bulk download: Generated {successful_pdfs} PDFs in {zip_filename}, size={zip_size} bytes')
                
                if not zip_content or len(zip_content) == 0:
                    current_app.logger.error('ZIP buffer is empty!')
                    return jsonify({
                        'success': False,
                        'message': 'ZIP ফাইল তৈরি করতে সমস্যা হয়েছে'
                    }), 500
                
                # Validate ZIP file (should start with PK signature)
                if not zip_content.startswith(b'PK'):
                    current_app.logger.error(f'Invalid ZIP file! Starts with: {zip_content[:10]}')
                    return jsonify({
                        'success': False,
                        'message': 'ZIP ফাইল করাপ্ট হয়েছে'
                    }), 500
                
                return Response(
                    zip_content,
                    mimetype='application/zip',
                    headers={
                        'Content-Disposition': f'attachment; filename="{zip_filename}"',
                        'Content-Length': str(zip_size),
                        'Content-Type': 'application/zip',
                        'Cache-Control': 'no-cache, no-store, must-revalidate',
                        'Pragma': 'no-cache',
                        'Expires': '0'
                    }
                )
                
            except Exception as e:
                current_app.logger.error(f'Bulk download error: {str(e)}', exc_info=True)
                return jsonify({
                    'success': False,
                    'message': f'বাল্ক ডাউনলোড ব্যর্থ: {str(e)}'
                }), 500
            finally:
                import gc
                gc.collect()
        except Exception as outer_error:
            # Catch any errors that occur before the inner try block
            current_app.logger.error(f'Bulk download outer error: {str(outer_error)}', exc_info=True)
            return jsonify({
                'success': False,
                'message': f'বাল্ক ডাউনলোড ব্যর্থ: {str(outer_error)}'
            }), 500

    @app.route('/admin/delete_user/<int:user_id>', methods=['POST'])
    @login_required
    def delete_user(user_id):
        if not is_admin(current_user):
            flash('You do not have permission to perform this action.', 'danger')
            return redirect(url_for('index'))
        
        user_to_delete = User.query.get_or_404(user_id)
        if is_admin(user_to_delete):
            flash('Admin users cannot be deleted.', 'danger')
            return redirect(url_for('admin_dashboard'))
        
        db.session.delete(user_to_delete)
        db.session.commit()
        flash('User deleted successfully.', 'success')
        return redirect(url_for('admin_dashboard'))

    @app.route('/admin/edit_user/<int:user_id>', methods=['GET', 'POST'])
    @login_required
    def admin_edit_user(user_id):
        if not is_admin(current_user):
            flash('You do not have permission to access this page.', 'danger')
            return redirect(url_for('index'))
        
        user_to_edit = User.query.get_or_404(user_id)
        if is_admin(user_to_edit):
            flash('Cannot edit admin users from here.', 'danger')
            return redirect(url_for('admin_dashboard'))

        if request.method == 'POST':
            old_full_name = user_to_edit.full_name
            user_to_edit.full_name = request.form['full_name']
            user_to_edit.email = request.form['email']
            selected_roles = request.form.getlist('roles')
            is_valid, normalized_roles = validate_role_selection(selected_roles)
            if not is_valid:
                flash(normalized_roles, 'danger')
                return redirect(url_for('admin_edit_user', user_id=user_id))
            user_to_edit.role = serialize_roles(normalized_roles)
            
            is_external_list = request.form.getlist('is_external_teacher')
            is_ext_val = '1' in (is_external_list or [])
            # Prefer stored link so we always update the SAME teacher the dashboard shows
            teacher_id_val = getattr(user_to_edit, 'teacher_id', None) or request.form.get('teacher_id', type=int)
            teacher = None
            if teacher_id_val:
                teacher = Teacher.query.get(teacher_id_val)
            if not teacher and 'teacher' in normalized_roles:
                teacher = Teacher.query.filter_by(name=old_full_name).first()
            if not teacher and 'teacher' in normalized_roles:
                teacher = Teacher.query.filter_by(name=user_to_edit.full_name).first()
            if not teacher and 'teacher' in normalized_roles:
                short_name = (user_to_edit.username or user_to_edit.full_name.split()[0].lower())[:10]
                counter = 1
                base_short = short_name
                while Teacher.query.filter_by(short_name=short_name).first():
                    short_name = f"{base_short[:10-len(str(counter))]}{counter}"
                    counter += 1
                teacher = Teacher(name=user_to_edit.full_name, short_name=short_name, is_external=is_ext_val)
                db.session.add(teacher)
            
            if teacher:
                user_to_edit.teacher_id = teacher.id  # persist link so next time we use it
                if 'teacher' in normalized_roles:
                    if teacher.name != user_to_edit.full_name:
                        teacher.name = user_to_edit.full_name
                    designation = request.form.get('designation', '').strip()
                    institute = request.form.get('institute', '').strip()
                    call_sign = request.form.get('call_sign', '').strip()
                    bank_account_no = request.form.get('bank_account_no', '').strip()
                    teacher.designation = designation if designation else None
                    teacher.institute = institute if institute else 'Law Discipline, KU'
                    teacher.call_sign = call_sign if call_sign else None
                    teacher.bank_account_no = bank_account_no if bank_account_no else None
                from sqlalchemy import text
                db.session.execute(
                    text("UPDATE teacher SET is_external = :v WHERE id = :id"),
                    {"v": 1 if is_ext_val else 0, "id": teacher.id}
                )
            elif teacher_id_val:
                from sqlalchemy import text
                db.session.execute(
                    text("UPDATE teacher SET is_external = :v WHERE id = :id"),
                    {"v": 1 if is_ext_val else 0, "id": teacher_id_val}
                )
            
            db.session.commit()
            flash('User updated successfully.', 'success')
            return redirect(url_for('admin_dashboard'))
        
        # Get teacher record for display: use user.teacher_id link first (same teacher dashboard shows), else match by name and save link
        current_roles = parse_roles(user_to_edit.role)
        teacher_record = None
        if 'teacher' in current_roles:
            if getattr(user_to_edit, 'teacher_id', None):
                teacher_record = Teacher.query.get(user_to_edit.teacher_id)
            if not teacher_record:
                teacher_record = Teacher.query.filter_by(name=user_to_edit.full_name).first()
            if not teacher_record and user_to_edit.full_name:
                from sqlalchemy import func
                teacher_record = Teacher.query.filter(
                    func.trim(Teacher.name) == user_to_edit.full_name.strip()
                ).first()
            # If we found teacher by name but user has no link yet, we'll save the link on next POST (when they save)
        
        return render_template(
            'admin_edit_user.html',
            user=user_to_edit,
            role_choices=NON_ADMIN_ROLE_CHOICES,
            current_roles=current_roles,
            teacher_record=teacher_record,
        )

    @app.route('/admin/create_user', methods=['POST'])
    @login_required
    def admin_create_user():
        if not is_admin(current_user):
            flash('You do not have permission to create users.', 'danger')
            return redirect(url_for('index'))
        
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        full_name = request.form.get('full_name', '').strip()
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        selected_roles = request.form.getlist('roles')
        
        # Validation
        if not all([username, email, full_name, password, confirm_password]):
            flash('All fields are required.', 'danger')
            return redirect(url_for('admin_dashboard'))
        
        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return redirect(url_for('admin_dashboard'))
        
        if len(password) < 6:
            flash('Password must be at least 6 characters long.', 'danger')
            return redirect(url_for('admin_dashboard'))
        
        if not selected_roles:
            flash('Please select at least one role/category.', 'danger')
            return redirect(url_for('admin_dashboard'))
        
        # Check if user already exists
        if User.query.filter_by(username=username).first():
            flash('Username already exists.', 'danger')
            return redirect(url_for('admin_dashboard'))
        
        if User.query.filter_by(email=email).first():
            flash('Email already registered.', 'danger')
            return redirect(url_for('admin_dashboard'))
        
        # Validate role selection
        is_valid, normalized_roles = validate_role_selection(selected_roles)
        if not is_valid:
            flash(normalized_roles, 'danger')
            return redirect(url_for('admin_dashboard'))
        
        try:
            # Create user
            user = User(
                username=username,
                email=email,
                full_name=full_name,
                role=serialize_roles(normalized_roles)
            )
            user.set_password(password)
            db.session.add(user)
            db.session.flush()  # Flush to get user ID
            
            # Create teacher record if user has teacher role
            if 'teacher' in normalized_roles or 'dean' in normalized_roles or 'head' in normalized_roles:
                # Generate short_name from username or name
                short_name = (username or full_name.split()[0].lower())[:10]
                counter = 1
                base_short = short_name
                while Teacher.query.filter_by(short_name=short_name).first():
                    short_name = f"{base_short[:10-len(str(counter))]}{counter}"
                    counter += 1
                
                teacher = Teacher(
                    name=full_name,
                    short_name=short_name,
                    designation=request.form.get('designation', '').strip() or None,
                    institute=request.form.get('institute', '').strip() or 'Law Discipline, KU',
                    call_sign=request.form.get('call_sign', '').strip() or None,
                    bank_account_no=request.form.get('bank_account_no', '').strip() or None
                )
                try:
                    if hasattr(teacher, 'is_external'):
                        teacher.is_external = bool(request.form.get('is_external_teacher'))
                except Exception:
                    pass
                db.session.add(teacher)
                db.session.flush()
                if hasattr(user, 'teacher_id'):
                    user.teacher_id = teacher.id
            db.session.commit()
            flash(f'User "{full_name}" ({username}) created successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error creating user: {e}', exc_info=True)
            flash(f'Error creating user: {str(e)}', 'danger')
        
        return redirect(url_for('admin_dashboard'))

    @app.route('/admin/reset_password/<int:user_id>', methods=['GET', 'POST'])
    @login_required
    def admin_reset_password(user_id):
        if not is_admin(current_user):
            flash('You do not have permission to access this page.', 'danger')
            return redirect(url_for('index'))

        user_to_reset = User.query.get_or_404(user_id)
        if is_admin(user_to_reset):
            flash('Cannot reset password for admin users from here.', 'danger')
            return redirect(url_for('admin_dashboard'))

        if request.method == 'POST':
            new_password = request.form['new_password']
            user_to_reset.set_password(new_password)
            db.session.commit()
            flash(f"Password for {user_to_reset.username} has been reset.", 'success')
            return redirect(url_for('admin_dashboard'))
            
        return render_template('admin_reset_password.html', user=user_to_reset)

    @app.route('/debug/system-info')
    def debug_system_info():
        """Debug endpoint to check system information"""
        try:
            deps = check_dependencies()
            permissions = check_file_permissions()
            system_info = get_system_info()
            
            return {
                'dependencies': deps,
                'permissions': permissions,
                'system_info': system_info,
                'status': 'ok'
            }
        except Exception as e:
            return {
                'error': str(e),
                'status': 'error'
            }
    
    @app.route('/profile', methods=['GET', 'POST'])
    @login_required
    def profile():
        if request.method == 'POST':
            full_name = request.form.get('full_name', '').strip()
            email = request.form.get('email', '').strip()
            current_password = request.form.get('current_password')
            new_password = request.form.get('new_password')
            confirm_password = request.form.get('confirm_password')

            # Store old values before updating (for finding related records)
            old_full_name = current_user.full_name
            old_email = current_user.email
            name_changed = False
            email_changed = False

            # Validate and update name
            if full_name and full_name != old_full_name:
                # Check if another user already has this full_name (if needed)
                existing_user = User.query.filter(User.id != current_user.id, User.full_name == full_name).first()
                if existing_user:
                    flash('This name is already in use by another user.', 'danger')
                    db.session.rollback()
                    teacher_record = None
                    user_roles = parse_roles(current_user.role)
                    if 'teacher' in user_roles or 'dean' in user_roles or 'head' in user_roles:
                        teacher_record = Teacher.query.filter_by(name=current_user.full_name).first()
                    return render_template('profile.html', teacher_record=teacher_record)
                current_user.full_name = full_name
                name_changed = True

            # Validate and update email
            if email and email != old_email:
                # Check if email already exists
                existing_user = User.query.filter(User.id != current_user.id, User.email == email).first()
                if existing_user:
                    flash('Email already in use by another user.', 'danger')
                    db.session.rollback()
                    teacher_record = None
                    user_roles = parse_roles(current_user.role)
                    if 'teacher' in user_roles or 'dean' in user_roles or 'head' in user_roles:
                        teacher_record = Teacher.query.filter_by(name=current_user.full_name).first()
                    return render_template('profile.html', teacher_record=teacher_record)
                current_user.email = email
                email_changed = True

            # Handle photo upload
            if 'photo' in request.files:
                photo_file = request.files['photo']
                if photo_file and photo_file.filename:
                    # Check file extension
                    filename = photo_file.filename
                    file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
                    if file_ext not in app.config['ALLOWED_EXTENSIONS']:
                        flash('Invalid file type. Allowed: PNG, JPG, JPEG, GIF, WEBP', 'danger')
                        db.session.rollback()
                        teacher_record = None
                        user_roles = parse_roles(current_user.role)
                        if 'teacher' in user_roles or 'dean' in user_roles or 'head' in user_roles:
                            teacher_record = Teacher.query.filter_by(name=current_user.full_name).first()
                        return render_template('profile.html', teacher_record=teacher_record)
                    
                    # Delete old photo if exists
                    if current_user.photo:
                        old_photo_rel_path = current_user.photo.lstrip('/')
                        old_photo_path = os.path.join(app.root_path, old_photo_rel_path)
                        if os.path.exists(old_photo_path):
                            try:
                                os.remove(old_photo_path)
                            except:
                                pass
                    
                    # Save new photo
                    filename = f"user_{current_user.id}_{int(datetime.utcnow().timestamp())}.{file_ext}"
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    photo_file.save(filepath)
                    
                    # Store relative path in database
                    current_user.photo = f"/static/uploads/user_photos/{filename}"

            # Update password if provided
            if new_password or confirm_password:
                if not current_password or not current_user.check_password(current_password):
                    flash('Current password is incorrect.', 'danger')
                    db.session.rollback()
                    teacher_record = None
                    user_roles = parse_roles(current_user.role)
                    if 'teacher' in user_roles or 'dean' in user_roles or 'head' in user_roles:
                        teacher_record = Teacher.query.filter_by(name=current_user.full_name).first()
                    return render_template('profile.html', teacher_record=teacher_record)
                if new_password != confirm_password:
                    flash('New passwords do not match.', 'danger')
                    db.session.rollback()
                    teacher_record = None
                    user_roles = parse_roles(current_user.role)
                    if 'teacher' in user_roles or 'dean' in user_roles or 'head' in user_roles:
                        teacher_record = Teacher.query.filter_by(name=current_user.full_name).first()
                    return render_template('profile.html', teacher_record=teacher_record)
                if new_password:
                    current_user.set_password(new_password)

            # Update teacher information if user is a teacher
            # IMPORTANT: Find teacher by OLD name first, then update to new name
            user_roles = parse_roles(current_user.role)
            if 'teacher' in user_roles or 'dean' in user_roles or 'head' in user_roles:
                # Use old_full_name to find existing teacher record
                teacher = Teacher.query.filter_by(name=old_full_name).first()
                
                if not teacher:
                    # If no teacher found by old name, try new name (in case name wasn't changed)
                    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
                
                if not teacher:
                    # Create teacher record if it doesn't exist
                    short_name_base = (current_user.username or current_user.full_name.split()[0].lower() if current_user.full_name else 'teacher')
                    short_name_base = ''.join(ch for ch in short_name_base.lower() if ch.isalnum())[:10] or 'teacher'
                    short_name = short_name_base
                    counter = 1
                    while Teacher.query.filter_by(short_name=short_name).first():
                        suffix = str(counter)
                        short_name = f"{short_name_base[:10-len(suffix)]}{suffix}"
                        counter += 1
                    teacher = Teacher(name=current_user.full_name, short_name=short_name)
                    db.session.add(teacher)
                    db.session.flush()
                
                # Update teacher name if it changed (adapts to name change)
                if name_changed and teacher.name != current_user.full_name:
                    teacher.name = current_user.full_name
                    # Note: All existing relationships (class_sessions, assignments, etc.) 
                    # that reference this teacher by ID will automatically adapt
                    # since they use teacher_id foreign key, not name matching
                
                # Update call_sign and bank_account_no
                call_sign = request.form.get('call_sign', '').strip()
                bank_account_no = request.form.get('bank_account_no', '').strip()
                teacher.call_sign = call_sign if call_sign else None
                teacher.bank_account_no = bank_account_no if bank_account_no else None

            # Update theme preference
            theme = request.form.get('theme', 'default').strip()
            if theme:
                current_user.theme = theme

            # Commit all changes
            try:
                db.session.commit()
                flash('Profile updated successfully!', 'success')
                if name_changed:
                    flash(f'Name updated from "{old_full_name}" to "{current_user.full_name}". All related data has been adapted.', 'info')
                if email_changed:
                    flash(f'Email updated from "{old_email}" to "{current_user.email}".', 'info')
            except Exception as e:
                db.session.rollback()
                current_app.logger.error(f'Error updating profile: {str(e)}')
                flash('An error occurred while updating your profile. Please try again.', 'danger')
            
            return redirect(url_for('profile'))
        
        # Get teacher record for display
        teacher_record = None
        user_roles = parse_roles(current_user.role)
        if 'teacher' in user_roles or 'dean' in user_roles or 'head' in user_roles:
            teacher_record = Teacher.query.filter_by(name=current_user.full_name).first()
        
        return render_template('profile.html', teacher_record=teacher_record)

    return app

app = create_app()

if __name__ == '__main__':
    print("Initializing database...")
    with app.app_context():
        try:
            db.create_all()
            print("✓ Database initialized")
        except Exception as e:
            # If database connection fails, log but continue
            # This allows the app to start even if database is temporarily unavailable
            import logging
            logging.warning(f"Could not create database tables: {e}")
            print(f"⚠️  Warning: Database connection issue: {e}")
            print("App will continue, but some features may not work.")
    
    print("Setting up server configuration...")
    
    # Get port from environment variable (for Render) or use default
    port = int(os.environ.get('PORT', 5001))
    
    # Use 0.0.0.0 for production (Render) or network access, 127.0.0.1 for localhost only
    # Set ALLOW_NETWORK_ACCESS=1 in environment to enable network access
    allow_network = os.environ.get('ALLOW_NETWORK_ACCESS', '0') == '1'
    host = '0.0.0.0' if (os.environ.get('RENDER') or allow_network) else '127.0.0.1'
    
    if host == '0.0.0.0':
        print("Getting network IP address...")
        # Try to get local IP address for display
        local_ip = 'YOUR_LOCAL_IP'
        try:
            import socket
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.settimeout(1)  # 1 second timeout to prevent hanging
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
            s.close()
            print(f"✓ Found local IP: {local_ip}")
        except (socket.timeout, socket.error, OSError, Exception) as e:
            # If we can't get IP from socket, try alternative method
            print("⚠️  Could not get IP from network, trying alternative method...")
            try:
                import socket
                hostname = socket.gethostname()
                local_ip = socket.gethostbyname(hostname)
                if local_ip.startswith('127.'):
                    # If we got localhost, try to get real IP from network interfaces
                    local_ip = 'YOUR_LOCAL_IP'
                    print("⚠️  Using localhost - network access may be limited")
                else:
                    print(f"✓ Found IP via hostname: {local_ip}")
            except Exception as e2:
                print(f"⚠️  Could not determine IP address: {e2}")
                print("   Server will still run, but network access URL won't be shown")
            pass
        
        print(f"\n{'='*60}")
        print(f"Server running on ALL network interfaces")
        print(f"Local access: http://127.0.0.1:{port}")
        print(f"Network access: http://{local_ip}:{port}")
        print(f"{'='*60}\n")
    else:
        print(f"\nServer running on http://127.0.0.1:{port} (localhost only)")
        print("To enable network access, set ALLOW_NETWORK_ACCESS=1 environment variable\n")

    # Suppress harmless socket errors (client disconnections)
    import logging
    logging.getLogger('werkzeug').setLevel(logging.ERROR)  # Only show errors
    
    # Suppress socket connection errors (harmless client disconnections)
    import warnings
    warnings.filterwarnings('ignore', category=RuntimeWarning)
    
    try:
        # Verify host binding
        if host == '0.0.0.0':
            print(f"✓ Binding to 0.0.0.0:{port} (all network interfaces)")
            print(f"✓ Network access enabled")
        else:
            print(f"✓ Binding to 127.0.0.1:{port} (localhost only)")
        
        print(f"\n🚀 Server started successfully!")
        if host == '0.0.0.0':
            print(f"📱 Access from mobile/other devices:")
            print(f"   http://{local_ip}:{port}")
            print(f"\n💡 Make sure:")
            print(f"   1. Both devices are on the same WiFi network")
            print(f"   2. Firewall allows Python/port {port}")
            print(f"   3. Use exact URL: http://{local_ip}:{port}\n")
        
        # Run with minimal logging to reduce noise
        import sys
        from werkzeug.serving import WSGIRequestHandler
        
        class QuietHandler(WSGIRequestHandler):
            def log_request(self, *args, **kwargs):
                pass  # Don't log normal requests
        
        # Ensure we're actually binding to 0.0.0.0 for network access
        if host == '0.0.0.0':
            # Double check binding
            import socket
            test_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            test_socket.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            try:
                test_socket.bind(('0.0.0.0', port))
                test_socket.close()
                print(f"✓ Port {port} is available and will bind to 0.0.0.0")
            except OSError as e:
                print(f"⚠️  Port binding test failed: {e}")
                test_socket.close()
        
        print("\n" + "="*60)
        print("🔄 Starting Flask development server...")
        print("   (The server will appear to 'hang' - this is NORMAL!)")
        print("   It's waiting for HTTP requests. Press CTRL+C to stop.")
        print("="*60 + "\n")
        
        # Use SocketIO to run the app (enables WebSocket support)
        app.socketio.run(app, host=host, port=port, use_reloader=False, log_output=False, allow_unsafe_werkzeug=True)
    except OSError as e:
        if "Address already in use" in str(e):
            print(f"\n❌ Error: Port {port} is already in use!")
            print(f"   Please stop the other application or use a different port:")
            print(f"   PORT=5002 ALLOW_NETWORK_ACCESS=1 python3 app.py\n")
        else:
            print(f"\n❌ Error starting server: {e}\n")
        raise
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}\n")
        raise 
    