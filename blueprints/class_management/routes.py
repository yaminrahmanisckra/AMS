from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, current_app, Response, jsonify
from flask_login import login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from sqlalchemy import or_, text, func
from .models import (
    db,
    Teacher,
    Session,
    ClassStudent,
    ClassAttendance,
    ClassSplitInvite,
    CourseReview,
    EvaluationInvite,
    EvaluationSubmission,
    ExamScrutinizerInvite,
    ExamPaperEvaluation,
    StudentFeedbackLink,
    StudentFeedbackResponse,
    CourseOutline,
    CourseQuestionThread,
    CourseQuestionMessage,
    CourseQuestionAttachment,
)
from models import User  # Import the User model from the new models file
try:
    from blueprints.student_management.models import Student
except ImportError:
    Student = None

try:
    from blueprints.course_management.models import Curriculum, Course, CurriculumYearTerm, CourseSessionAssignment, StudentCourseRegistration
except ImportError:
    Curriculum = None
    CourseSessionAssignment = None
    Course = None
    CurriculumYearTerm = None
    StudentCourseRegistration = None
import pandas as pd
import os
from datetime import datetime, date
import secrets
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from collections import Counter, defaultdict
from reportlab.lib.units import inch
import io
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import json
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
# docx imports moved to lazy imports (only when needed) to prevent startup hang
# from docx import Document
# from docx.shared import Pt, Inches
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from uuid import uuid4
from role_utils import has_teacher_privileges, is_admin, parse_roles
try:
    from utils.semester_utils import filter_by_active_semester
except ImportError:
    filter_by_active_semester = None

# WeasyPrint lazy import - only import when needed to prevent startup hang
# Module-level import removed because it causes startup hang on macOS
# This prevents the app from hanging during startup
_WEASYPRINT_HTML = None
_WEASYPRINT_AVAILABLE = None

def _get_weasyprint_html():
    """Lazy import WeasyPrint HTML - only import when actually needed"""
    global _WEASYPRINT_HTML, _WEASYPRINT_AVAILABLE
    
    if _WEASYPRINT_AVAILABLE is None:
        # First time - try to import
        import logging
        import os
        import platform
        import ctypes
        from ctypes import util as ctypes_util
        
        logger = logging.getLogger(__name__)
        
        # Setup library paths for macOS BEFORE importing WeasyPrint
        if platform.system() == 'Darwin':
            homebrew_lib_path = '/opt/homebrew/lib'
            if os.path.exists(homebrew_lib_path):
                # Set environment variables
                os.environ['DYLD_FALLBACK_LIBRARY_PATH'] = f"{homebrew_lib_path}:{os.environ.get('DYLD_FALLBACK_LIBRARY_PATH', '')}"
                os.environ['PKG_CONFIG_PATH'] = f"/opt/homebrew/lib/pkgconfig:{os.environ.get('PKG_CONFIG_PATH', '')}"
                
                # Monkey-patch ctypes.util.find_library
                original_find_library = ctypes_util.find_library
                def patched_find_library(name):
                    lib_mappings = {
                        'gobject-2.0-0': 'libgobject-2.0.0.dylib',
                        'gobject-2.0': 'libgobject-2.0.dylib',
                    }
                    if name in lib_mappings:
                        lib_path = os.path.join(homebrew_lib_path, lib_mappings[name])
                        if os.path.exists(lib_path):
                            return lib_path
                    for pattern in [f'lib{name}.dylib', f'lib{name}.0.dylib']:
                        lib_path = os.path.join(homebrew_lib_path, pattern)
                        if os.path.exists(lib_path):
                            return lib_path
                    result = original_find_library(name)
                    return result if result else None
                
                ctypes_util.find_library = patched_find_library
                
                # Pre-load libraries
                try:
                    lib_path = os.path.join(homebrew_lib_path, 'libgobject-2.0.0.dylib')
                    if os.path.exists(lib_path):
                        ctypes.CDLL(lib_path, mode=ctypes.RTLD_GLOBAL)
                except:
                    pass
        
        try:
            logger.info("Attempting to import WeasyPrint (lazy import)...")
            from weasyprint import HTML
            _WEASYPRINT_HTML = HTML
            _WEASYPRINT_AVAILABLE = True
            logger.info("✓ WeasyPrint imported successfully (lazy import)")
        except ImportError as e:
            _WEASYPRINT_AVAILABLE = False
            _WEASYPRINT_HTML = None
            logger.error(f"✗ WeasyPrint ImportError: {e}")
            logger.warning("PDF generation features will be disabled.")
        except Exception as e:
            _WEASYPRINT_AVAILABLE = False
            _WEASYPRINT_HTML = None
            logger.error(f"✗ WeasyPrint import error: {e}", exc_info=True)
            logger.warning("PDF generation features will be disabled.")
    
    if _WEASYPRINT_AVAILABLE and _WEASYPRINT_HTML is None:
        # Retry if somehow HTML is None but available is True
        try:
            from weasyprint import HTML
            _WEASYPRINT_HTML = HTML
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to re-import WeasyPrint: {e}")
            _WEASYPRINT_AVAILABLE = False
    
    return _WEASYPRINT_HTML if _WEASYPRINT_AVAILABLE else None

def _is_weasyprint_available():
    """Check if WeasyPrint is available (lazy check)"""
    if _WEASYPRINT_AVAILABLE is None:
        _get_weasyprint_html()  # Trigger lazy import
    return _WEASYPRINT_AVAILABLE is True


COURSE_REVIEW_GRADE_ROWS = [
    {'key': 'grade_a_plus', 'scale': '80% and above', 'letter': 'A+'},
    {'key': 'grade_a', 'scale': '75% to less than 80%', 'letter': 'A'},
    {'key': 'grade_a_minus', 'scale': '70% to less than 75%', 'letter': 'A-'},
    {'key': 'grade_b_plus', 'scale': '65% to less than 70%', 'letter': 'B+'},
    {'key': 'grade_b', 'scale': '60% to less than 65%', 'letter': 'B'},
    {'key': 'grade_b_minus', 'scale': '55% to less than 60%', 'letter': 'B-'},
    {'key': 'grade_c_plus', 'scale': '50% to less than 55%', 'letter': 'C+'},
    {'key': 'grade_c', 'scale': '45% to less than 50%', 'letter': 'C'},
    {'key': 'grade_d', 'scale': '40% to less than 45%', 'letter': 'D'},
    {'key': 'grade_f', 'scale': 'Less than 40%', 'letter': 'F'},
    {'key': 'grade_i', 'scale': 'Incomplete', 'letter': 'I'},
    {'key': 'grade_w', 'scale': 'Withdrawal', 'letter': 'W'},
]

COURSE_REVIEW_COMMENT_FIELDS = [
    {'key': 'comment_student_questionnaires', 'label': '1) Student (Course Evaluation) Questionnaires:'},
    {'key': 'comment_external_examiners', 'label': '2) External Examiners or Moderators (if any):'},
    {'key': 'comment_curriculum', 'label': '3) Curriculum: Comment on the continuing appropriateness of the Course curriculum in relation to the intended learning outcomes and its compliance with the National Qualification Framework'},
    {'key': 'comment_assessment', 'label': '4) Assessment: Comment on the continuing effectiveness of method(s) of assessment in relation to the intended learning outcomes (Course objectives)'},
    {'key': 'comment_enhancement', 'label': '5) Enhancement: Comment on the implementation of changes proposed in earlier Faculty Course Review Reports'},
    {'key': 'comment_future_changes', 'label': "6) Outline any changes in the future delivery or structure of the Course that this semester/term's experience may prompt"},
]

SCOPE_FULL = 'full'
SCOPE_PART_A = 'part_a'
SCOPE_PART_B = 'part_b'
SPLIT_PARTS = {SCOPE_PART_A, SCOPE_PART_B}
COURSE_SCOPE_LABELS = {
    SCOPE_FULL: 'Full Course',
    SCOPE_PART_A: 'Part A',
    SCOPE_PART_B: 'Part B',
}

def _generate_feedback_code():
    """Generate a short, URL-friendly code for feedback access."""
    while True:
        raw = secrets.token_urlsafe(8)
        code = ''.join(ch for ch in raw if ch.isalnum()).upper()
        code = code[:10]
        if len(code) < 6:
            continue
        if not StudentFeedbackLink.query.filter_by(access_code=code).first():
            return code


def find_course_from_curriculum(session_course_code, session_course_name=None):
    """
    Find a Course from the curriculum that matches the session's course code or name.
    Handles various formats like "0421 28 Law 4103" -> "Law 4103" or "Law4103"
    
    Returns: Course object if found, None otherwise
    """
    import re
    import logging
    logger = logging.getLogger(__name__)
    
    if not Course:
        return None
    
    course_data = None
    
    def extract_core_code(code_str):
        """Extract the core code pattern (e.g., 'Law4103' without space) from various formats"""
        if not code_str:
            return None
        # Try to find pattern like "Law 4103", "Law4103", "CSE 1101", etc.
        match = re.search(r'([A-Za-z]+)\s*(\d{4})', code_str)
        if match:
            return f"{match.group(1)}{match.group(2)}"  # No space: "Law4103"
        return None
    
    # Extract core code pattern from session course code
    session_core_code = extract_core_code(session_course_code)
    logger.debug(f"find_course_from_curriculum: session_code='{session_course_code}', session_core='{session_core_code}', session_name='{session_course_name}'")
    
    # Try exact match by course code
    if session_course_code:
        course_data = Course.query.filter_by(course_code=session_course_code).first()
        if course_data:
            logger.debug(f"Found by exact course_code match: {course_data.course_code}")
            return course_data
    
    # Try case-insensitive match by course code
    if session_course_code:
        course_data = Course.query.filter(func.lower(Course.course_code) == func.lower(session_course_code)).first()
        if course_data:
            logger.debug(f"Found by case-insensitive course_code match: {course_data.course_code}")
            return course_data
    
    # Try whitespace-normalized match (handles extra spaces, tabs, etc.)
    if session_course_code:
        session_code_normalized = ' '.join(session_course_code.strip().split())  # Normalize whitespace
        all_courses = Course.query.all()
        for course in all_courses:
            if course.course_code:
                curriculum_code_normalized = ' '.join(course.course_code.strip().split())
                if session_code_normalized.lower() == curriculum_code_normalized.lower():
                    logger.debug(f"Found by whitespace-normalized match: {course.course_code}")
                    return course
    
    # Try with extracted course code pattern (with space)
    if session_core_code:
        # Try "Law 4103" format (with space)
        extracted_with_space = re.sub(r'([A-Za-z]+)(\d{4})', r'\1 \2', session_core_code)
        course_data = Course.query.filter(func.lower(Course.course_code) == func.lower(extracted_with_space)).first()
        if course_data:
            logger.debug(f"Found by extracted code with space: {course_data.course_code}")
            return course_data
        
        # Try "Law4103" format (without space)
        course_data = Course.query.filter(func.lower(Course.course_code) == func.lower(session_core_code)).first()
        if course_data:
            logger.debug(f"Found by extracted code without space: {course_data.course_code}")
            return course_data
    
    # Try normalized matching - compare core codes from both session and curriculum
    if session_core_code:
        all_courses = Course.query.all()
        session_core_lower = session_core_code.lower()
        for course in all_courses:
            if course.course_code:
                # Extract core code from curriculum course code
                curriculum_core = extract_core_code(course.course_code)
                if curriculum_core and curriculum_core.lower() == session_core_lower:
                    logger.debug(f"Found by normalized core code match: {course.course_code} (core: {curriculum_core})")
                    return course
    
    # Try partial match - check if curriculum course code is contained in session code
    if session_course_code:
        all_courses = Course.query.all()
        session_code_lower = session_course_code.lower()
        # Also try without spaces for partial matching
        session_code_no_space = session_code_lower.replace(' ', '')
        for course in all_courses:
            if course.course_code:
                course_code_lower = course.course_code.lower()
                course_code_no_space = course_code_lower.replace(' ', '')
                # Check with and without spaces
                if course_code_lower in session_code_lower or course_code_no_space in session_code_no_space:
                    logger.debug(f"Found by partial code match: {course.course_code}")
                    return course
    
    # Try exact match by course name
    if session_course_name:
        course_data = Course.query.filter_by(course_name=session_course_name).first()
        if course_data:
            logger.debug(f"Found by exact course_name match: {course_data.course_name}")
            return course_data
    
    # Try case-insensitive partial match by course name
    if session_course_name:
        course_data = Course.query.filter(func.lower(Course.course_name).like(f'%{session_course_name.lower()}%')).first()
        if course_data:
            logger.debug(f"Found by partial course_name match: {course_data.course_name}")
            return course_data
    
    # Try reverse match (session name contains course name or vice versa)
    if session_course_name:
        all_courses = Course.query.all()
        session_name_lower = session_course_name.lower()
        for course in all_courses:
            if course.course_name:
                course_name_lower = course.course_name.lower()
                if session_name_lower in course_name_lower or course_name_lower in session_name_lower:
                    logger.debug(f"Found by reverse course_name match: {course.course_name}")
                    return course
    
    logger.debug(f"No course found for session_code='{session_course_code}', session_name='{session_course_name}'")
    return None


class_management_bp = Blueprint(
    'class_management', __name__,
    template_folder='templates',
    static_folder='static'
)


@class_management_bp.before_request
def restrict_to_teaching_roles():
    if not current_user.is_authenticated:
        return
    
    # Allow student routes for all authenticated users
    student_routes = [
        'class_management.student_view_scores',
        'class_management.student_course_files',
        'class_management.student_download_course_outline_pdf',
        'class_management.student_download_uploaded_file',
        'class_management.student_create_course_question',
        'class_management.student_reply_course_question',
        'class_management.download_course_question_attachment',
        'class_management.delete_course_question_thread',
        'class_management.delete_course_question_message'
    ]
    if request.endpoint in student_routes:
        return
    
    if has_teacher_privileges(current_user):
        return
    flash('Class Management is available only to teaching staff.', 'danger')
    return redirect(url_for('index'))


def _ensure_current_teacher():
    """Return a Teacher instance for the logged-in user, creating one if needed."""
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if teacher:
        return teacher

    base = (current_user.full_name or 'teacher').split(' ')[0].lower()
    base = ''.join(ch for ch in base if ch.isalnum()) or 'teacher'
    base = base[:10]
    candidate = base
    counter = 1
    while Teacher.query.filter_by(short_name=candidate).first():
        suffix = str(counter)
        candidate = f"{base[:10-len(suffix)]}{suffix}"
        counter += 1

    teacher = Teacher(name=current_user.full_name, short_name=candidate)
    db.session.add(teacher)
    db.session.commit()
    return teacher


def _get_related_sessions(session, include_archived=False):
    """Return all sessions that belong to the same split group."""
    if not session or not session.split_group_id:
        return [session] if session else []

    query = Session.query.filter_by(split_group_id=session.split_group_id)
    if not include_archived:
        query = query.filter_by(archived=False)
    related = query.order_by(Session.id.asc()).all()
    return related or [session]


def _carry_on_assessment_marks(class_student, session):
    """Carry on previous assessment marks for retake students if carry_on is enabled in registration"""
    if not StudentCourseRegistration or not Student:
        return
    
    try:
        # Get student record from Students Management
        student_record = Student.query.filter_by(student_id=class_student.student_id).first()
        if not student_record:
            return
        
        # Find registration for this course and session
        registration = StudentCourseRegistration.query.filter_by(
            student_id=student_record.id,
            course_code=session.course_code,
            academic_session=session.academic_session,
            year=session.year,
            term=session.term
        ).first()
        
        if not registration or not registration.carry_on:
            return
        
        # Only carry on for retake/re-retake students
        if registration.remark not in ['Retake', 'Re-retake']:
            return
        
        # Find previous session with same course_code and student_id
        # Look for sessions with different academic_session/year/term
        previous_sessions = Session.query.filter(
            Session.course_code == session.course_code,
            Session.id != session.id,
            Session.archived == False
        ).order_by(Session.academic_session.desc(), Session.created_at.desc()).all()
        
        for prev_session in previous_sessions:
            # Find student in previous session
            prev_student = ClassStudent.query.filter_by(
                session_id=prev_session.id,
                student_id=class_student.student_id
            ).first()
            
            if prev_student:
                # Copy assessment marks
                if prev_student.assessment1 is not None:
                    class_student.assessment1 = prev_student.assessment1
                if prev_student.assessment2 is not None:
                    class_student.assessment2 = prev_student.assessment2
                if prev_student.assessment3 is not None:
                    class_student.assessment3 = prev_student.assessment3
                if prev_student.assessment4 is not None:
                    class_student.assessment4 = prev_student.assessment4
                
                current_app.logger.info(
                    f'Carried on assessment marks for student {class_student.student_id} '
                    f'from session {prev_session.id} to session {session.id}'
                )
                break  # Only carry from the most recent previous session
    except Exception as e:
        current_app.logger.error(f'Error carrying on assessment marks: {str(e)}', exc_info=True)


def _replicate_student_to_peers(session, source_student, *, old_identifier=None):
    """Create or update a student record across split peer sessions."""
    if not session or not session.split_group_id:
        return

    peer_sessions = [s for s in _get_related_sessions(session) if s.id != session.id]
    for peer in peer_sessions:
        identifier = old_identifier or source_student.student_id
        peer_student = ClassStudent.query.filter_by(session_id=peer.id, student_id=identifier).first()
        if not peer_student:
            peer_student = ClassStudent(
                student_id=source_student.student_id,
                name=source_student.name,
                session_id=peer.id,
                teacher_id=peer.teacher_id
            )
            db.session.add(peer_student)
        else:
            peer_student.student_id = source_student.student_id
            peer_student.name = source_student.name


def _delete_student_from_peers(session, student_identifier):
    """Remove a student from all peer sessions within the split group."""
    if not session or not session.split_group_id:
        return

    for peer in _get_related_sessions(session):
        if peer.id == session.id:
            continue
        peer_student = ClassStudent.query.filter_by(session_id=peer.id, student_id=student_identifier).first()
        if peer_student:
            db.session.delete(peer_student)


def _gather_split_student_map(session):
    """Return related sessions and a map of student_id -> [ClassStudent,...]."""
    related_sessions = _get_related_sessions(session)
    session_ids = [s.id for s in related_sessions if s]
    if not session_ids:
        return related_sessions, {}

    students = ClassStudent.query.filter(ClassStudent.session_id.in_(session_ids)).all()
    student_map = defaultdict(list)
    for stu in students:
        student_map[stu.student_id].append(stu)
    return related_sessions, student_map


def _recalculate_assessment_totals(session):
    """Recompute assessment aggregates across split sessions."""
    if not session or session.course_type != 'theory':
        return

    _, student_map = _gather_split_student_map(session)
    for entries in student_map.values():
        marks = []
        for entry in entries:
            for idx in range(1, 5):
                value = getattr(entry, f'assessment{idx}')
                if value is not None:
                    marks.append(value)
        marks.sort(reverse=True)

        if session.category == 'pg':
            if marks:
                best = marks[:3] if len(marks) >= 3 else marks
                avg = sum(best) / len(best)
                # Convert best 3 sum to 40 marks scale: always use 30 as max (sum / 30) * 40
                best_sum = sum(best)
                total_40 = int(round((best_sum / 30) * 40))  # Round for PG courses
                avg_value = round(avg, 2)
            else:
                avg_value = None
                total_40 = None
            for entry in entries:
                entry.assessment_avg = avg_value
                entry.assessment_total_40 = total_40
                entry.assessment_total = None
        else:
            if marks:
                best = marks[:3] if len(marks) >= 3 else marks
                total = sum(best)  # Keep fraction for UG courses, no rounding
            else:
                total = None
            for entry in entries:
                entry.assessment_total = total
                entry.assessment_avg = None
                entry.assessment_total_40 = None


def _collect_combined_assessment_marks(session):
    """Return a map of student_id -> list of assessment marks across split sessions."""
    _, student_map = _gather_split_student_map(session)
    marks_map = {}
    for student_id, entries in student_map.items():
        values = []
        for entry in entries:
            for idx in range(1, 5):
                val = getattr(entry, f'assessment{idx}')
                if val is not None:
                    values.append(val)
        marks_map[student_id] = values
    return marks_map


def _build_combined_assessment_values(session):
    """
    Combine assessment values from all related sessions in a split course.
    
    Returns:
        value_map: {student_id: {1: val1, 2: val2, 3: val3, 4: val4}}
        ug_best3: {student_id: best3_total}
        pg_avg_map: {student_id: average}
        pg_total_map: {student_id: total_40_scale}
    """
    # Step 1: Get all related sessions (if split course)
    if session.split_group_id:
        related_sessions = Session.query.filter_by(
            split_group_id=session.split_group_id,
            archived=False
        ).all()
    else:
        related_sessions = [session]
    
    if not related_sessions:
        return {}, {}, {}, {}
    
    # Step 2: Get ALL ClassStudent records from ALL related sessions
    session_ids = [s.id for s in related_sessions]
    all_class_students = ClassStudent.query.filter(
        ClassStudent.session_id.in_(session_ids)
    ).all()
    
    # Step 3: Group by student_id (STRING) - same student from different sessions
    # Structure: {student_id: [ClassStudent_from_session_A, ClassStudent_from_session_B, ...]}
    student_groups = defaultdict(list)
    for cs in all_class_students:
        student_groups[cs.student_id].append(cs)
    
    # Step 4: Build combined values for each student
    value_map = {}
    ug_best3 = {}
    pg_avg_map = {}
    pg_total_map = {}
    
    for student_id, student_records in student_groups.items():
        # Initialize: {1: None, 2: None, 3: None, 4: None}
        combined = {1: None, 2: None, 3: None, 4: None}
        
        # Combine values from ALL records (sessions)
        # Example: Part A session has assessment1=10, assessment2=5
        #          Part B session has assessment3=7, assessment4=8
        # Result: {1: 10, 2: 5, 3: 7, 4: 8}
        for record in student_records:
            for idx in [1, 2, 3, 4]:
                val = getattr(record, f'assessment{idx}', None)
                # Set value if it exists and we don't have one yet
                if val is not None and combined[idx] is None:
                    try:
                        combined[idx] = float(val)
                    except (ValueError, TypeError):
                        pass
        
        value_map[student_id] = combined
        
        # Step 5: Calculate Best 3 / PG Total
        valid_marks = [v for v in combined.values() if v is not None]
        valid_marks.sort(reverse=True)  # Descending
        
        if session.category == 'pg':
            if valid_marks:
                best = valid_marks[:3] if len(valid_marks) >= 3 else valid_marks
                avg = sum(best) / len(best)
                pg_avg_map[student_id] = round(avg, 2)
                best_sum = sum(best)
                pg_total_map[student_id] = int(round((best_sum / 30) * 40))  # Round for PG courses
            else:
                pg_avg_map[student_id] = None
                pg_total_map[student_id] = None
        else:
            # UG: Best 3 total
            if valid_marks:
                best = valid_marks[:3] if len(valid_marks) >= 3 else valid_marks
                ug_best3[student_id] = sum(best)  # Keep fraction for UG courses, no rounding
            else:
                ug_best3[student_id] = None
    
    return value_map, ug_best3, pg_avg_map, pg_total_map


def _get_editable_assessment_indices(session):
    """Return which assessment inputs current teacher can edit."""
    if session.course_scope == SCOPE_PART_A:
        return {1, 2}
    if session.course_scope == SCOPE_PART_B:
        return {3, 4}
    return {1, 2, 3, 4}


def _calculate_attendance_mark_from_percentage(percentage):
    """Convert attendance percentage to marks."""
    if percentage >= 90:
        return 10
    if percentage >= 85:
        return 9
    if percentage >= 80:
        return 8
    if percentage >= 75:
        return 7
    if percentage >= 70:
        return 6
    if percentage >= 65:
        return 5
    if percentage >= 60:
        return 4
    return 0


def _build_attendance_summary(session):
    """Aggregate attendance across split sessions."""
    related_sessions = _get_related_sessions(session)
    session_ids = [s.id for s in related_sessions if s]
    if not session_ids:
        return {'total_classes': 0, 'per_student': {}, 'per_session_totals': defaultdict(int), 'related_sessions': related_sessions}

    attendance_records = ClassAttendance.query.filter(
        ClassAttendance.session_id.in_(session_ids)
    ).order_by(ClassAttendance.date.asc(), ClassAttendance.id.asc()).all()

    students = ClassStudent.query.filter(ClassStudent.session_id.in_(session_ids)).all()
    student_lookup = {stu.id: stu for stu in students}

    per_student_counts = defaultdict(lambda: {'present': 0})
    per_session_date_counts = defaultdict(lambda: defaultdict(int))

    for record in attendance_records:
        per_session_date_counts[(record.session_id, record.date)][record.student_id] += 1
        student = student_lookup.get(record.student_id)
        if not student:
            continue
        if record.is_present:
            per_student_counts[student.student_id]['present'] += 1
        per_student_counts[student.student_id]['records'] = per_student_counts[student.student_id].get('records', 0) + 1

    per_session_totals = defaultdict(int)
    total_classes = 0
    for (session_id, _), counts in per_session_date_counts.items():
        class_count = max(counts.values()) if counts else 0
        total_classes += class_count
        per_session_totals[session_id] += class_count

    per_student_result = {}
    # Group students by student_id to handle split courses (multiple sessions, same student_id)
    students_by_id = {}
    for student in students:
        if student.student_id not in students_by_id:
            students_by_id[student.student_id] = []
        students_by_id[student.student_id].append(student)
    
    for student_id, student_list in students_by_id.items():
        stats = per_student_counts.get(student_id, {'present': 0, 'records': 0})
        percentage = (stats['present'] / total_classes * 100) if total_classes else 0
        
        # Check all student records for this student_id to find manual marks
        # (for split courses, manual marks might be on any of the related student records)
        manual_marks = None
        for student in student_list:
            if student.attendance_marks_manual is not None:
                manual_marks = student.attendance_marks_manual
                break
        
        if manual_marks is not None:
            marks = manual_marks
        else:
            marks = _calculate_attendance_mark_from_percentage(percentage)
        per_student_result[student_id] = {
            'present': stats['present'],
            'percentage': percentage,
            'marks': marks,
            'marks_manual': manual_marks is not None
        }

    return {
        'total_classes': total_classes,
        'per_student': per_student_result,
        'per_session_totals': per_session_totals,
        'related_sessions': related_sessions
    }


def _build_split_context(session, attendance_summary=None):
    """Prepare metadata for templates about split courses."""
    if not session or not session.split_group_id or session.course_scope == SCOPE_FULL:
        return None

    peer_info = []
    for peer in _get_related_sessions(session, include_archived=True):
        if not peer or peer.id == session.id:
            continue
        peer_info.append({
            'id': peer.id,
            'teacher_name': peer.teacher.name if peer.teacher else '—',
            'teacher_short': peer.teacher.short_name if peer.teacher else '',
            'course_scope': COURSE_SCOPE_LABELS.get(peer.course_scope, 'Part')
        })

    context = {
        'scope_label': COURSE_SCOPE_LABELS.get(session.course_scope, 'Part'),
        'peers': peer_info
    }

    if attendance_summary:
        totals = attendance_summary.get('per_session_totals', {})
        context['class_totals'] = []
        for related in attendance_summary.get('related_sessions', []):
            if not related:
                continue
            context['class_totals'].append({
                'session_id': related.id,
                'teacher_name': related.teacher.name if related.teacher else '—',
                'teacher_short': related.teacher.short_name if related.teacher else '',
                'classes': totals.get(related.id, 0),
                'is_current': related.id == session.id
            })
        context['total_classes'] = attendance_summary.get('total_classes', 0)

    return context


# Create uploads folder if it doesn't exist
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Q&A uploads folder (separate)
QA_UPLOAD_FOLDER = os.path.join(UPLOAD_FOLDER, 'qa_questions')
if not os.path.exists(QA_UPLOAD_FOLDER):
    os.makedirs(QA_UPLOAD_FOLDER)


def _get_qa_upload_dir(thread_id):
    """Return (and create) per-thread upload folder."""
    thread_dir = os.path.join(QA_UPLOAD_FOLDER, str(thread_id))
    os.makedirs(thread_dir, exist_ok=True)
    return thread_dir


def _save_qa_attachments(files, thread_id):
    """Save attachments and return metadata list."""
    saved = []
    if not files:
        return saved
    upload_dir = _get_qa_upload_dir(thread_id)
    for file in files:
        if not file or not file.filename:
            continue
        safe_name = secure_filename(file.filename)
        if not safe_name:
            continue
        unique_name = f"{uuid4().hex}_{safe_name}"
        file_path = os.path.join(upload_dir, unique_name)
        file.save(file_path)
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else None
        file_type = file.mimetype or os.path.splitext(safe_name)[1].lstrip('.')
        saved.append({
            'file_name': safe_name,
            'file_path': file_path,
            'file_size': file_size,
            'file_type': file_type
        })
    return saved


def _delete_qa_attachments(attachments):
    """Delete attachment files from disk."""
    for attachment in attachments:
        try:
            if attachment.file_path and os.path.exists(attachment.file_path):
                os.remove(attachment.file_path)
        except Exception:
            pass

@class_management_bp.route('/')
@login_required
def index():
    """Main dashboard for class management"""
    teacher = _ensure_current_teacher()
    
    # Start with base query
    query = Session.query.filter_by(
        teacher_id=teacher.id,
        archived=False
    )
    
    # IMPORTANT: Update sessions with academic_session BEFORE filtering
    # This ensures all sessions have academic_session set for proper filtering
    sessions_before_update = query.all()
    
    # Update sessions with academic_session and batch from CourseSessionAssignment if available
    # Also sync all assignments with curriculum year-term config if missing
    # IMPORTANT: Do this BEFORE auto-creating sessions so assignments have correct academic_session
    if CourseSessionAssignment and Curriculum and CurriculumYearTerm:
        try:
            # First, update all assignments that are missing batch/academic_session from curriculum year-term config
            all_assignments = CourseSessionAssignment.query.all()
            assignment_update_count = 0
            for assignment in all_assignments:
                if assignment.curriculum_id and (not assignment.batch or not assignment.academic_session):
                    try:
                        curriculum = Curriculum.query.get(assignment.curriculum_id)
                        if curriculum:
                            year_term_config = curriculum.get_year_term_config(assignment.year, assignment.term)
                            if year_term_config:
                                updated = False
                                if not assignment.batch and year_term_config.batch and year_term_config.batch != 'None':
                                    assignment.batch = year_term_config.batch
                                    updated = True
                                    current_app.logger.info(f'Updated assignment {assignment.id} batch from year-term config: {year_term_config.batch}')
                                if not assignment.academic_session and year_term_config.academic_session:
                                    assignment.academic_session = year_term_config.academic_session
                                    updated = True
                                    current_app.logger.info(f'Updated assignment {assignment.id} academic_session from year-term config: {year_term_config.academic_session}')
                                
                                if updated:
                                    assignment_update_count += 1
                    except Exception as assign_error:
                        current_app.logger.warning(f'Error updating assignment {assignment.id}: {assign_error}')
            
            if assignment_update_count > 0:
                db.session.commit()
                current_app.logger.info(f'Updated {assignment_update_count} assignments with batch/academic_session from curriculum year-term config')
            
            # Now update sessions with academic_session from assignments
            updated_count = 0
            for session in sessions_before_update:
                # Find CourseSessionAssignment for this session
                assignment = CourseSessionAssignment.query.filter_by(session_id=session.id).first()
                if assignment:
                    # Update academic_session if assignment has it and session doesn't
                    if assignment.academic_session and not session.academic_session:
                        session.academic_session = assignment.academic_session
                        updated_count += 1
                        current_app.logger.info(f'Updated session {session.id} ({session.course_name}) with academic_session: {assignment.academic_session}')
                    elif assignment.academic_session and session.academic_session != assignment.academic_session:
                        # Update if different
                        session.academic_session = assignment.academic_session
                        updated_count += 1
                        current_app.logger.info(f'Updated session {session.id} ({session.course_name}) academic_session from {session.academic_session} to {assignment.academic_session}')
            
            if updated_count > 0:
                db.session.commit()
                current_app.logger.info(f'Updated {updated_count} sessions with academic_session from CourseSessionAssignment')
        except Exception as e:
            current_app.logger.error(f'Error updating sessions from CourseSessionAssignment: {str(e)}', exc_info=True)
            db.session.rollback()
    
    # Auto-create missing Sessions from CourseSessionAssignment
    # This ensures all assigned courses appear in Class Management
    # IMPORTANT: Do this AFTER updating assignments with academic_session from curriculum config
    if CourseSessionAssignment and Course:
        try:
            # Find all assignments for this teacher that don't have sessions yet
            missing_assignments = CourseSessionAssignment.query.filter_by(
                teacher_id=teacher.id
            ).filter(
                or_(
                    CourseSessionAssignment.session_created == False,
                    CourseSessionAssignment.session_id.is_(None)
                )
            ).all()
            
            current_app.logger.info(f'[DEBUG] Teacher {teacher.id} ({teacher.name}): Found {len(missing_assignments)} assignments without sessions')
            
            # Log assignment details for debugging
            for idx, assignment in enumerate(missing_assignments, 1):
                course = Course.query.get(assignment.course_id) if assignment.course_id else None
                course_code = course.course_code if course else f'course_id={assignment.course_id}'
                current_app.logger.info(f'[DEBUG] Missing assignment #{idx}: ID={assignment.id}, Course={course_code}, Year={assignment.year}, Term={assignment.term}, AcademicSession={assignment.academic_session}, Batch={assignment.batch}, Section={assignment.section}')
            
            created_count = 0
            for assignment in missing_assignments:
                try:
                    # Get course details
                    course = Course.query.get(assignment.course_id)
                    if not course:
                        current_app.logger.warning(f'Course {assignment.course_id} not found for assignment {assignment.id}')
                        continue
                    
                    # Determine course_scope based on section
                    if assignment.section == 'A':
                        course_scope = SCOPE_PART_A
                    elif assignment.section == 'B':
                        course_scope = SCOPE_PART_B
                    else:
                        course_scope = SCOPE_FULL
                    
                    # Check if a session with similar parameters already exists
                    existing_session = Session.query.filter_by(
                        course_code=course.course_code,
                        teacher_id=teacher.id,
                        year=assignment.year,
                        term=assignment.term,
                        archived=False
                    ).first()
                    
                    if existing_session:
                        # Link the assignment to the existing session
                        assignment.session_id = existing_session.id
                        assignment.session_created = True
                        current_app.logger.info(f'Linked assignment {assignment.id} to existing session {existing_session.id}')
                    else:
                        # Create new Session with academic_session from assignment
                        session_obj = Session(
                            year=assignment.year,
                            term=assignment.term,
                            academic_session=assignment.academic_session,  # Use assignment's academic_session (should be set by now)
                            course_code=course.course_code,
                            course_name=course.course_name,
                            teacher_id=teacher.id,
                            course_type=course.course_type.lower() if course.course_type else 'theory',
                            category=course.category if course.category else 'ug',
                            course_scope=course_scope
                        )
                        db.session.add(session_obj)
                        db.session.flush()  # Get session ID
                        
                        # Link assignment to session
                        assignment.session_id = session_obj.id
                        assignment.session_created = True
                        
                        created_count += 1
                        current_app.logger.info(f'Auto-created session {session_obj.id} from assignment {assignment.id} for course {course.course_code} (academic_session: {assignment.academic_session})')
                
                except Exception as create_error:
                    current_app.logger.error(f'Error auto-creating session for assignment {assignment.id}: {create_error}', exc_info=True)
                    continue
            
            if created_count > 0:
                db.session.commit()
                current_app.logger.info(f'[DEBUG] Auto-created {created_count} sessions from CourseSessionAssignment for teacher {teacher.id}')
            else:
                current_app.logger.info(f'[DEBUG] No new sessions created for teacher {teacher.id} (all assignments already have sessions or no valid assignments)')
        except Exception as e:
            current_app.logger.error(f'[DEBUG] Error auto-creating sessions from CourseSessionAssignment: {str(e)}', exc_info=True)
            db.session.rollback()
    
    # Now apply active semester filtering AFTER updating academic_session and creating sessions
    # Re-query to get updated sessions including newly created ones
    query = Session.query.filter_by(
        teacher_id=teacher.id,
        archived=False
    )
    
    # Log sessions before filtering
    sessions_before_filter = query.all()
    current_app.logger.info(f'[DEBUG] Teacher {teacher.id}: Found {len(sessions_before_filter)} sessions BEFORE active semester filtering')
    for s in sessions_before_filter:
        current_app.logger.info(f'[DEBUG] Session before filter: ID={s.id}, Course={s.course_code} ({s.course_name}), Year={s.year}, Term={s.term}, AcademicSession={s.academic_session}')
    
    # Apply active semester filtering (if not admin and filter function available)
    if filter_by_active_semester and not is_admin(current_user):
        try:
            # Don't filter by batch - get ALL active semesters
            # This ensures courses from all active semesters are shown, not just from a specific batch
            batch = None
            
            # Log active semester configuration
            try:
                from utils.semester_utils import get_active_semesters
                active_semesters = get_active_semesters(batch=None)  # Get ALL active semesters
                active_sem_info = [f"{s.academic_session}-{s.year}-{s.term}-{s.batch or 'ALL'}" for s in active_semesters]
                current_app.logger.info(f'[DEBUG] Active semesters for filtering (all batches): {active_sem_info}')
            except Exception as sem_error:
                current_app.logger.warning(f'[DEBUG] Error getting active semesters: {sem_error}')
            
            # Apply active semester filter (batch=None to get all active semesters)
            query = filter_by_active_semester(query, Session, batch=None, admin_override=False)
            current_app.logger.info(f'[DEBUG] Applied active semester filtering for teacher {teacher.id} (using all active semesters)')
        except Exception as filter_error:
            current_app.logger.error(f'[DEBUG] Error applying active semester filter: {filter_error}', exc_info=True)
            # Don't fail the request, but log the error
    
    sessions = query.order_by(Session.created_at.desc()).all()

    current_app.logger.info(f'[DEBUG] Teacher {teacher.id} ({teacher.name}): Found {len(sessions)} sessions AFTER filtering (was {len(sessions_before_filter)} before filtering)')
    for s in sessions:
        current_app.logger.info(f'[DEBUG] Session after filter: ID={s.id}, Course={s.course_code} ({s.course_name}), Year={s.year}, Term={s.term}, AcademicSession={s.academic_session}')
    for s in sessions:
        current_app.logger.debug(f'Session: ID={s.id}, Name={s.course_name}, Session={s.academic_session}, Year={s.year}, Term={s.term}, Archived={s.archived}, Teacher={s.teacher_id}')

    split_context_map = {session.id: _build_split_context(session) for session in sessions if session.split_group_id}
    # Get teachers excluding Head of the Discipline
    from role_utils import get_teachers_excluding_head
    teachers = get_teachers_excluding_head()
    pending_split_invites = ClassSplitInvite.query.filter_by(invited_teacher_id=teacher.id, status='pending').order_by(ClassSplitInvite.created_at.desc()).all()
    
    # Get all batches from Students Management for the dropdown
    batches = []
    if Student:
        try:
            all_batches = db.session.query(Student.batch).distinct().filter(Student.batch.isnot(None)).order_by(Student.batch.desc()).all()
            batches = [batch[0] for batch in all_batches]
        except Exception:
            batches = []
    
    # Build assignment map for template to access batch and academic_session from CourseSessionAssignment
    assignment_map = {}
    if CourseSessionAssignment and Course:
        try:
            for session in sessions:
                # Try to find assignment by session_id first
                assignment = CourseSessionAssignment.query.filter_by(session_id=session.id).first()
                
                # If not found by session_id, try to find by course_code, teacher_id, year, term
                if not assignment and session.course_code and session.teacher_id and session.year and session.term:
                    try:
                        # Try to match by course_code, teacher_id, year, term
                        # First try exact match
                        assignment = CourseSessionAssignment.query.filter_by(
                            teacher_id=session.teacher_id,
                            year=session.year,
                            term=session.term
                        ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                            Course.course_code == session.course_code
                        ).first()
                        
                        # If not found, try without section matching (for full course sessions)
                        if not assignment:
                            assignment = CourseSessionAssignment.query.filter_by(
                                teacher_id=session.teacher_id,
                                year=session.year,
                                term=session.term
                            ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                                Course.course_code == session.course_code
                            ).filter(
                                or_(
                                    CourseSessionAssignment.section.is_(None),
                                    CourseSessionAssignment.section == ''
                                )
                            ).first()
                        
                        # If found, update the session_id to link them
                        if assignment and not assignment.session_id:
                            assignment.session_id = session.id
                            assignment.session_created = True
                            try:
                                db.session.commit()
                                current_app.logger.info(f'Linked assignment {assignment.id} to session {session.id} for course {session.course_code}')
                            except Exception as commit_error:
                                db.session.rollback()
                                current_app.logger.warning(f'Could not link assignment {assignment.id} to session {session.id}: {commit_error}')
                    except Exception as query_error:
                        current_app.logger.warning(f'Error querying assignment for session {session.id}: {query_error}')
                
                if assignment:
                    # If assignment doesn't have batch/academic_session, try to get from curriculum year-term config
                    batch = assignment.batch
                    academic_session = assignment.academic_session
                    
                    if Curriculum and CurriculumYearTerm and (not batch or not academic_session):
                        try:
                            if assignment.curriculum_id:
                                curriculum = Curriculum.query.get(assignment.curriculum_id)
                                if curriculum:
                                    year_term_config = curriculum.get_year_term_config(assignment.year, assignment.term)
                                    if year_term_config:
                                        if not batch and year_term_config.batch and year_term_config.batch != 'None':
                                            batch = year_term_config.batch
                                            assignment.batch = batch
                                            current_app.logger.info(f'Updated assignment {assignment.id} batch from year-term config: {batch}')
                                        if not academic_session and year_term_config.academic_session:
                                            academic_session = year_term_config.academic_session
                                            assignment.academic_session = academic_session
                                            current_app.logger.info(f'Updated assignment {assignment.id} academic_session from year-term config: {academic_session}')
                                        
                                        if (not assignment.batch and batch) or (not assignment.academic_session and academic_session):
                                            try:
                                                db.session.commit()
                                            except Exception as commit_error:
                                                db.session.rollback()
                                                current_app.logger.warning(f'Could not update assignment {assignment.id}: {commit_error}')
                        except Exception as config_error:
                            current_app.logger.warning(f'Error getting year-term config for assignment {assignment.id}: {config_error}')
                    
                    assignment_map[session.id] = {
                        'batch': batch or '',
                        'academic_session': academic_session or ''
                    }
                    # Also update session's academic_session if assignment has it and session doesn't
                    if academic_session and not session.academic_session:
                        try:
                            session.academic_session = academic_session
                            db.session.commit()
                            current_app.logger.info(f'Updated session {session.id} academic_session from assignment: {academic_session}')
                        except Exception as update_error:
                            db.session.rollback()
                            current_app.logger.warning(f'Could not update session {session.id} academic_session: {update_error}')
                    
                    # Auto-add students from batch if session has no students but batch is available
                    if batch and batch.strip() and batch != 'None' and Student:
                        try:
                            existing_students_count = ClassStudent.query.filter_by(session_id=session.id).count()
                            if existing_students_count == 0:
                                current_app.logger.info(f'Session {session.id} has no students but batch {batch} is available. Attempting to add students...')
                                students_from_batch = Student.query.filter_by(batch=batch).all()
                                if students_from_batch:
                                    added_count = 0
                                    for student in students_from_batch:
                                        # Check if student is registered for this course (finalized registration only)
                                        if StudentCourseRegistration and session.course_code and session.academic_session and session.year and session.term:
                                            registration = StudentCourseRegistration.query.filter_by(
                                                student_id=student.id,
                                                course_code=session.course_code,
                                                academic_session=session.academic_session,
                                                year=session.year,
                                                term=session.term,
                                                status='finalized'
                                            ).first()
                                            
                                            if not registration:
                                                current_app.logger.info(f'Student {student.student_id} ({student.name}) not registered for course {session.course_code}, skipping...')
                                                continue
                                        
                                        class_student = ClassStudent(
                                            student_id=student.student_id,
                                            name=student.name,
                                            session_id=session.id,
                                            teacher_id=session.teacher_id
                                        )
                                        db.session.add(class_student)
                                        db.session.flush()  # Flush to get class_student.id before carry on
                                        
                                        # Carry on assessment marks if enabled in registration
                                        _carry_on_assessment_marks(class_student, session)
                                        
                                        # Replicate to peer sessions for split courses
                                        _replicate_student_to_peers(session, class_student)
                                        
                                        added_count += 1
                                    
                                    if added_count > 0:
                                        db.session.commit()
                                        current_app.logger.info(f'Successfully added {added_count} students from batch {batch} to session {session.id}')
                                else:
                                    current_app.logger.warning(f'No students found in batch {batch} for session {session.id}')
                        except Exception as auto_add_error:
                            db.session.rollback()
                            current_app.logger.error(f'Error auto-adding students to session {session.id} from batch {batch}: {auto_add_error}', exc_info=True)
        except Exception as e:
            current_app.logger.error(f'Error building assignment map: {str(e)}', exc_info=True)
            db.session.rollback()

    # Q&A notifications: count threads with latest message from student
    qa_notification_map = {}
    try:
        from sqlalchemy.orm import selectinload
        session_ids = [s.id for s in sessions]
        if session_ids:
            threads = CourseQuestionThread.query.options(
                selectinload(CourseQuestionThread.messages)
            ).filter(
                CourseQuestionThread.session_id.in_(session_ids),
                CourseQuestionThread.teacher_id == teacher.id
            ).all()
            for thread in threads:
                last_message = None
                if thread.messages:
                    last_message = max(
                        thread.messages,
                        key=lambda m: m.created_at or datetime.min
                    )
                if last_message and last_message.sender_role == 'student':
                    qa_notification_map[thread.session_id] = qa_notification_map.get(thread.session_id, 0) + 1
    except Exception as e:
        current_app.logger.warning(f'Error loading Q&A notifications: {e}')

    return render_template(
        'class_management/index.html',
        sessions=sessions,
        teacher=teacher,
        teacher_display_name=(getattr(current_user, 'full_name', '') or teacher.name or current_user.username),
        teachers=teachers,
        course_scope_labels=COURSE_SCOPE_LABELS,
        split_context_map=split_context_map,
        pending_split_invites=pending_split_invites,
        batches=batches,
        assignment_map=assignment_map,
        qa_notification_map=qa_notification_map
    )

@class_management_bp.route('/create_session', methods=['POST'])
@login_required
def create_session():
    """Create a new session"""
    try:
        teacher = _ensure_current_teacher()
        batch = request.form.get('batch', '').strip()
        curriculum_id = request.form.get('curriculum_id', type=int)
        year = request.form.get('year', '').strip()
        term = request.form.get('term', '').strip()
        academic_session = request.form.get('academic_session', '').strip()
        course_id = request.form.get('course_id', type=int)
        course_code = request.form.get('course_code', '').strip()
        course_name = request.form.get('course_name', '').strip()
        course_type = request.form.get('course_type', 'theory')
        category = request.form.get('category', 'ug')
        course_scope = request.form.get('course_scope', SCOPE_FULL)
        partner_teacher_id = request.form.get('partner_teacher_id')
        
        current_app.logger.info(f'Creating session - batch: {batch}, curriculum_id: {curriculum_id}, year: {year}, term: {term}, course_name: {course_name}, course_code: {course_code}')
        
        # If course_id is provided, fetch course details from Course model
        if course_id and Course:
            course = Course.query.get(course_id)
            if course:
                course_code = course.course_code
                course_name = course.course_name
                course_type = course.course_type.lower()
                category = course.category
                current_app.logger.info(f'Fetched course details from Course model: {course_code} - {course_name}')
        
        if not year or not term:
            flash('Year and term are required!', 'error')
            current_app.logger.warning(f'Missing year or term - year: {year}, term: {term}')
            return redirect(url_for('class_management.index'))
        
        if not course_code or not course_name:
            flash('Course code and course name are required!', 'error')
            current_app.logger.warning(f'Missing course_code or course_name - code: {course_code}, name: {course_name}')
            return redirect(url_for('class_management.index'))

        if course_scope not in COURSE_SCOPE_LABELS:
            flash('Invalid course scope selection.', 'error')
            return redirect(url_for('class_management.index'))

        # Prevent more than two teachers (Part A & Part B) from taking the same course simultaneously
        active_sessions = Session.query.filter(
            Session.course_code == course_code,
            Session.archived.is_(False)
        ).all()
        full_exists = any(s.course_scope == SCOPE_FULL for s in active_sessions)
        part_a_exists = any(s.course_scope == SCOPE_PART_A for s in active_sessions)
        part_b_exists = any(s.course_scope == SCOPE_PART_B for s in active_sessions)

        if full_exists:
            flash('This course already has a full-course session. Delete the existing session before assigning another teacher.', 'error')
            return redirect(url_for('class_management.index'))

        if part_a_exists and part_b_exists:
            flash('Both Part A and Part B are already assigned to teachers. Delete an existing section before adding another teacher.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_FULL and (part_a_exists or part_b_exists):
            flash('This course is already split between teachers. Delete the split sections before creating a full-course session.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_PART_A and part_a_exists:
            flash('Part A is already assigned to another teacher. Delete the existing Part A session before reassigning.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_PART_B and part_b_exists:
            flash('Part B is already assigned to another teacher. Delete the existing Part B session before reassigning.', 'error')
            return redirect(url_for('class_management.index'))

        if course_scope == SCOPE_FULL:
            session_obj = Session(
                year=year,
                term=term,
                academic_session=academic_session,
                course_code=course_code,
                course_name=course_name,
                teacher_id=teacher.id,
                course_type=course_type,
                category=category,
                course_scope=SCOPE_FULL
            )
            db.session.add(session_obj)
            db.session.flush()  # Get session ID before commit
            
            # Automatically add students from Students Management based on batch
            if batch and Student:
                try:
                    students_from_batch = Student.query.filter_by(batch=batch).all()
                    added_count = 0
                    skipped_count = 0
                    
                    # Get existing student IDs for this session to avoid duplicates
                    existing_student_ids = set()
                    
                    for student in students_from_batch:
                        # Check if already exists
                        existing = ClassStudent.query.filter_by(
                            session_id=session_obj.id,
                            student_id=student.student_id
                        ).first()
                        
                        if existing:
                            skipped_count += 1
                            continue
                        
                        class_student = ClassStudent(
                            student_id=student.student_id,
                            name=student.name,
                            session_id=session_obj.id,
                            teacher_id=teacher.id
                        )
                        db.session.add(class_student)
                        db.session.flush()  # Flush to get class_student.id before carry on
                        
                        # Carry on assessment marks if enabled in registration
                        _carry_on_assessment_marks(class_student, session_obj)
                        
                        _replicate_student_to_peers(session_obj, class_student)
                        added_count += 1
                    
                    db.session.commit()
                    
                    # Emit WebSocket event for live update
                    try:
                        from utils.websocket_events import emit_session_created
                        emit_session_created({
                            'session_id': session_obj.id,
                            'course_code': session_obj.course_code,
                            'course_name': session_obj.course_name,
                            'teacher_id': session_obj.teacher_id
                        })
                    except Exception as e:
                        current_app.logger.warning(f'Failed to emit session created event: {e}')
                    
                    if added_count > 0:
                        flash(f'Session created successfully! Automatically added {added_count} students from batch {batch}.', 'success')
                    else:
                        flash('Session created successfully!', 'success')
                        if skipped_count > 0:
                            flash(f'Note: {skipped_count} students were already in the session.', 'info')
                except Exception as e:
                    db.session.rollback()
                    current_app.logger.error(f'Error auto-adding students: {str(e)}', exc_info=True)
                    db.session.commit()  # Commit session even if student addition fails
                    flash('Session created successfully, but there was an error adding students automatically.', 'warning')
            else:
                db.session.commit()
                # Emit WebSocket event for live update
                try:
                    from utils.websocket_events import emit_session_created
                    emit_session_created({
                        'session_id': session_obj.id,
                        'course_code': session_obj.course_code,
                        'course_name': session_obj.course_name,
                        'teacher_id': session_obj.teacher_id
                    })
                except Exception as e:
                    current_app.logger.warning(f'Failed to emit session created event: {e}')
                flash('Session created successfully!', 'success')
            
            current_app.logger.info(f'Session created successfully - ID: {session_obj.id}, Name: {course_name}')
            return redirect(url_for('class_management.index'))

        # Handle split course (Part A/B)
        counterpart_scope = SCOPE_PART_B if course_scope == SCOPE_PART_A else SCOPE_PART_A

        try:
            partner_teacher_id_int = int(partner_teacher_id) if partner_teacher_id else None
        except (TypeError, ValueError):
            partner_teacher_id_int = None

        partner_teacher = None
        if partner_teacher_id_int:
            partner_teacher = Teacher.query.get(partner_teacher_id_int)

        if not partner_teacher:
            flash('Please select the teacher who will take the other part.', 'error')
            return redirect(url_for('class_management.index'))

        if partner_teacher.id == teacher.id:
            flash('Please assign a different teacher for the other part.', 'error')
            return redirect(url_for('class_management.index'))

        split_group_id = str(uuid4())

        current_session = Session(
            year=year,
            term=term,
            academic_session=academic_session,
            course_code=course_code,
            course_name=course_name,
            teacher_id=teacher.id,
            course_type=course_type,
            category=category,
            course_scope=course_scope,
            split_group_id=split_group_id
        )
        db.session.add(current_session)
        db.session.flush()

        invite = ClassSplitInvite(
            split_group_id=split_group_id,
            inviter_session_id=current_session.id,
            inviter_teacher_id=teacher.id,
            invited_teacher_id=partner_teacher.id,
            invited_scope=counterpart_scope,
            status='pending'
        )
        db.session.add(invite)
        
        # Automatically add students from Students Management based on batch
        if batch and Student:
            try:
                students_from_batch = Student.query.filter_by(batch=batch).all()
                added_count = 0
                skipped_count = 0
                not_registered_count = 0
                
                for student in students_from_batch:
                    # Check if already exists
                    existing = ClassStudent.query.filter_by(
                        session_id=current_session.id,
                        student_id=student.student_id
                    ).first()
                    
                    if existing:
                        skipped_count += 1
                        continue
                    
                    # Check if student is registered for this course (finalized registration only)
                    if StudentCourseRegistration and current_session.course_code and current_session.academic_session and current_session.year and current_session.term:
                        registration = StudentCourseRegistration.query.filter_by(
                            student_id=student.id,
                            course_code=current_session.course_code,
                            academic_session=current_session.academic_session,
                            year=current_session.year,
                            term=current_session.term,
                            status='finalized'
                        ).first()
                        
                        if not registration:
                            not_registered_count += 1
                            current_app.logger.info(f'Student {student.student_id} ({student.name}) not registered for course {current_session.course_code}, skipping...')
                            continue
                    
                    class_student = ClassStudent(
                        student_id=student.student_id,
                        name=student.name,
                        session_id=current_session.id,
                        teacher_id=teacher.id
                    )
                    db.session.add(class_student)
                    db.session.flush()  # Flush to get class_student.id before carry on
                    
                    # Carry on assessment marks if enabled in registration
                    _carry_on_assessment_marks(class_student, current_session)
                    
                    _replicate_student_to_peers(current_session, class_student)
                    added_count += 1
                
                # Commit session, invite, and students together
                db.session.commit()
                
                # Emit WebSocket event for live update
                try:
                    from utils.websocket_events import emit_session_created
                    emit_session_created({
                        'session_id': current_session.id,
                        'course_code': current_session.course_code,
                        'course_name': current_session.course_name,
                        'teacher_id': current_session.teacher_id
                    })
                except Exception as e:
                    current_app.logger.warning(f'Failed to emit session created event: {e}')
                
                if added_count > 0:
                    message = f'Split course created. Invitation sent. Automatically added {added_count} students from batch {batch}.'
                    if not_registered_count > 0:
                        message += f' Skipped {not_registered_count} student(s) not registered for this course.'
                    flash(message, 'success')
                else:
                    flash('Split course created. Invitation sent to the selected teacher.', 'success')
                    if skipped_count > 0:
                        flash(f'Note: {skipped_count} students were already in the session.', 'info')
            except Exception as e:
                db.session.rollback()
                current_app.logger.error(f'Error auto-adding students to split course: {str(e)}', exc_info=True)
                # Still commit the session and invite even if student addition fails
                db.session.add(current_session)
                db.session.add(invite)
                db.session.commit()
                flash('Split course created. Invitation sent, but there was an error adding students automatically.', 'warning')
        else:
            # Commit session and invite
            db.session.commit()
            flash('Split course created. Invitation sent to the selected teacher.', 'success')
        
        current_app.logger.info(f'Split session created successfully - ID: {current_session.id}, Name: {course_name}')
        return redirect(url_for('class_management.index'))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error creating session: {str(e)}', exc_info=True)
        flash(f'Error creating session: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))


@class_management_bp.route('/split_invites/<int:invite_id>/accept', methods=['POST'])
@login_required
def accept_split_invite(invite_id):
    teacher = _ensure_current_teacher()
    invite = ClassSplitInvite.query.get_or_404(invite_id)
    if invite.invited_teacher_id != teacher.id:
        flash('You are not authorized to respond to this invitation.', 'error')
        return redirect(url_for('class_management.index'))
    if invite.status != 'pending':
        flash('This invitation has already been processed.', 'info')
        return redirect(url_for('class_management.index'))

    inviter_session = Session.query.get(invite.inviter_session_id)
    if not inviter_session:
        invite.status = 'declined'
        invite.responded_at = datetime.utcnow()
        db.session.commit()
        flash('The original course is no longer available.', 'error')
        return redirect(url_for('class_management.index'))

    new_session = Session(
        year=inviter_session.year,
        term=inviter_session.term,
        academic_session=inviter_session.academic_session,
        course_code=inviter_session.course_code,
        course_name=inviter_session.course_name,
        teacher_id=teacher.id,
        course_type=inviter_session.course_type,
        category=inviter_session.category,
        course_scope=invite.invited_scope,
        split_group_id=invite.split_group_id
    )
    db.session.add(new_session)
    db.session.flush()

    inviter_students = ClassStudent.query.filter_by(session_id=inviter_session.id).all()
    for stu in inviter_students:
        clone = ClassStudent(
            student_id=stu.student_id,
            name=stu.name,
            session_id=new_session.id,
            teacher_id=teacher.id
        )
        db.session.add(clone)

    invite.status = 'accepted'
    invite.responded_at = datetime.utcnow()
    db.session.commit()
    flash('Invitation accepted. The course has been added to your dashboard.', 'success')
    return redirect(url_for('class_management.index'))


@class_management_bp.route('/split_invites/<int:invite_id>/decline', methods=['POST'])
@login_required
def decline_split_invite(invite_id):
    teacher = _ensure_current_teacher()
    invite = ClassSplitInvite.query.get_or_404(invite_id)
    if invite.invited_teacher_id != teacher.id:
        flash('You are not authorized to respond to this invitation.', 'error')
        return redirect(url_for('class_management.index'))
    if invite.status != 'pending':
        flash('This invitation has already been processed.', 'info')
        return redirect(url_for('class_management.index'))

    invite.status = 'declined'
    invite.responded_at = datetime.utcnow()
    db.session.commit()
    flash('Invitation declined.', 'info')
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/upload_students/<int:session_id>', methods=['POST'])
@login_required
def upload_students(session_id):
    """Upload students from Excel file"""
    teacher = _ensure_current_teacher()
    
    if 'file' not in request.files:
        flash('No file uploaded!', 'error')
        return redirect(url_for('class_management.index'))
    
    file = request.files['file']
    if file.filename == '':
        flash('No file selected!', 'error')
        return redirect(url_for('class_management.index'))
    
    if not file.filename.endswith('.xlsx'):
        flash('Please upload an Excel file!', 'error')
        return redirect(url_for('class_management.index'))
    
    try:
        df = pd.read_excel(file)
        # Clean and normalize column names
        df.columns = [col.strip().lower().replace(' ', '_') for col in df.columns]
        if 'student_id' not in df.columns or 'name' not in df.columns:
            raise Exception('Excel file must have columns: Student ID, Name')
        
        session = Session.query.get_or_404(session_id)
        for _, row in df.iterrows():
            student = ClassStudent(
                student_id=str(row['student_id']),
                name=row['name'],
                session_id=session.id,
                teacher_id=session.teacher_id or teacher.id
            )
            db.session.add(student)
            _replicate_student_to_peers(session, student)
        db.session.commit()
        flash('Students uploaded successfully!', 'success')
    except Exception as e:
        flash(f'Error uploading students: {str(e)}', 'error')
    
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/take_attendance/<int:session_id>', methods=['GET', 'POST'])
@login_required
def take_attendance(session_id):
    """Take or update attendance for a session."""
    session = Session.query.get_or_404(session_id)
    students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()
    
    if request.method == 'POST':
        try:
            date_val = datetime.strptime(request.form.get('date'), '%Y-%m-%d').date()
            double_class = request.form.get('double_class') == '1'
            
            # Overwrite logic: Delete existing records for this date first
            ClassAttendance.query.filter_by(session_id=session_id, date=date_val).delete()
            
            # Add new records
            for student in students:
                is_present = request.form.get(f'student_{student.id}') == 'present'
                num_classes = 2 if double_class else 1
                for _ in range(num_classes):
                    db.session.add(ClassAttendance(
                        date=date_val,
                        is_present=is_present,
                        student_id=student.id,
                        session_id=session_id,
                        teacher_id=session.teacher_id
                    ))
            
            db.session.commit()
            # Emit WebSocket event for live update
            try:
                from utils.websocket_events import emit_attendance_update
                emit_attendance_update(session_id, {
                    'date': date_val.isoformat(),
                    'double_class': double_class
                })
            except Exception as e:
                current_app.logger.warning(f'Failed to emit attendance update event: {e}')
            flash('Attendance saved successfully!', 'success')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error saving attendance for session {session_id}: {e}")
            flash(f'Error saving attendance: {str(e)}', 'error')
            return redirect(url_for('class_management.take_attendance', session_id=session_id, date=request.form.get('date')))

    # GET request logic
    try:
        selected_date_str = request.args.get('date', date.today().strftime('%Y-%m-%d'))
        selected_date = datetime.strptime(selected_date_str, '%Y-%m-%d').date()

        # Fetch existing records for the selected date
        existing_records = ClassAttendance.query.filter_by(
            session_id=session_id,
            date=selected_date
        ).all()
        
        # Prepare data for the template
        attendance_status = {}
        is_double_class = False
        if existing_records:
            # Check if it was a double class
            student_counts = defaultdict(int)
            for record in existing_records:
                student_counts[record.student_id] += 1
            if student_counts and max(student_counts.values()) > 1:
                is_double_class = True
            
            # Get the attendance status for each student
            for student in students:
                # A student is marked 'present' if they have at least one present record on that day
                is_present = any(r.is_present for r in existing_records if r.student_id == student.id)
                attendance_status[student.id] = is_present

        return render_template('class_management/take_attendance.html', 
                                session=session, 
                                students=students, 
                                today=selected_date_str,
                                attendance_status=attendance_status,
                                is_double_class=is_double_class,
                                split_meta=_build_split_context(session))
    except Exception as e:
        current_app.logger.error(f"Error loading attendance page for session {session_id}: {e}")
        flash(f'Error loading attendance page: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/view_attendance/<int:session_id>')
@login_required
def view_attendance(session_id):
    """View attendance for a session and display a detailed report."""
    session = Session.query.get_or_404(session_id)
    
    # Check user permissions: admin/head can view any session, regular teachers only their own
    user_roles = set(parse_roles(getattr(current_user, 'role', '')))
    if getattr(current_user, 'active_role', None):
        user_roles = set(parse_roles(current_user.active_role))
    can_view_all = is_admin(current_user) or 'head' in user_roles or 'dean' in user_roles
    
    # If not admin/head, check if this session belongs to the current teacher
    if not can_view_all:
        current_teacher = _ensure_current_teacher()
        if not current_teacher or session.teacher_id != current_teacher.id:
            flash('You do not have permission to view attendance for this session.', 'danger')
            return redirect(url_for('class_management.index'))
    
    attendance_summary = _build_attendance_summary(session)
    part_total_classes = None
    if session.course_scope in SPLIT_PARTS:
        part_total_classes = attendance_summary.get('per_session_totals', {}).get(session.id, 0)
    
    # Check if this is a split course
    is_split_course = session.split_group_id and session.course_scope in SPLIT_PARTS
    related_sessions = []
    if is_split_course:
        related_sessions = _get_related_sessions(session, include_archived=True)
    
    # For split courses, prepare separate attendance data for each teacher
    teacher_attendance_data = []
    if is_split_course and related_sessions:
        # Check user permissions: admin/head can see all parts, regular teachers only their own
        user_roles = set(parse_roles(getattr(current_user, 'role', '')))
        if getattr(current_user, 'active_role', None):
            user_roles = set(parse_roles(current_user.active_role))
        can_view_all = is_admin(current_user) or 'head' in user_roles or 'dean' in user_roles
        
        # Get current teacher if not admin/head
        current_teacher = None
        if not can_view_all:
            current_teacher = _ensure_current_teacher()
        
        # Get all students from all related sessions (for marks calculation)
        all_session_ids = [s.id for s in related_sessions]
        all_students = ClassStudent.query.filter(ClassStudent.session_id.in_(all_session_ids)).order_by(ClassStudent.student_id).all()
        student_lookup = {stu.student_id: stu for stu in all_students}
        
        # Get combined attendance summary for marks (this stays the same)
        agg_student_map = attendance_summary.get('per_student', {})
        agg_total_classes = attendance_summary.get('total_classes', 0)
        
        # Prepare data for each teacher/session (filtered by permissions)
        for related_session in related_sessions:
            # Skip if user doesn't have permission to view this session
            if not can_view_all:
                if not current_teacher or related_session.teacher_id != current_teacher.id:
                    continue
            session_students = ClassStudent.query.filter_by(session_id=related_session.id).order_by(ClassStudent.student_id).all()
            session_attendance_records = ClassAttendance.query.filter_by(session_id=related_session.id).order_by(ClassAttendance.date, ClassAttendance.id).all()
            
            if not session_attendance_records:
                continue
            
            # Build attendance data for this session
            attendance_by_date = defaultdict(list)
            for record in session_attendance_records:
                attendance_by_date[record.date].append(record)
            
            daily_class_counts = {}
            for date, records in attendance_by_date.items():
                student_counts_on_date = defaultdict(int)
                for record in records:
                    student_counts_on_date[record.student_id] += 1
                daily_class_counts[date] = max(student_counts_on_date.values()) if student_counts_on_date else 0
            
            headers_with_meta = []
            sorted_dates = sorted(daily_class_counts.keys())
            for date in sorted_dates:
                count = daily_class_counts.get(date, 0)
                if count == 1:
                    headers_with_meta.append({'label': date.strftime('%b %d, %Y'), 'date': date.strftime('%Y-%m-%d'), 'slot': 1})
                else:
                    for i in range(1, count + 1):
                        headers_with_meta.append({'label': f"{date.strftime('%b %d')} ({i})", 'date': date.strftime('%Y-%m-%d'), 'slot': i})
            
            student_report_data = []
            for student in session_students:
                student_records = [r for r in session_attendance_records if r.student_id == student.id]
                # Use combined stats for marks calculation
                agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})
                
                attendance_row = []
                student_attendance_by_date = defaultdict(list)
                for r in student_records:
                    student_attendance_by_date[r.date].append(r)
                
                for date in sorted_dates:
                    records_for_date = student_attendance_by_date[date]
                    num_classes_on_day = daily_class_counts.get(date, 0)
                    for i in range(num_classes_on_day):
                        cell = {
                            'status': '-',
                            'date': date.strftime('%Y-%m-%d'),
                            'slot': i + 1
                        }
                        if i < len(records_for_date):
                            record = records_for_date[i]
                            cell['status'] = 'P' if record.is_present else 'A'
                            cell['record_id'] = record.id
                        attendance_row.append(cell)
                
                student_data = {
                    'info': student,
                    'attendance_row': attendance_row,
                    'total_classes': agg_total_classes,  # Combined total for marks
                    'present_count': agg_stats['present'],  # Combined present count
                    'percentage': f"{agg_stats['percentage']:.2f}%",  # Combined percentage
                    'marks': agg_stats['marks'],  # Combined marks
                    'marks_manual': agg_stats.get('marks_manual', False)  # Whether marks are manually set
                }
                student_report_data.append(student_data)
            
            teacher_attendance_data.append({
                'session': related_session,
                'teacher_name': related_session.teacher.name if related_session.teacher else 'Unknown',
                'teacher_short': related_session.teacher.short_name if related_session.teacher else '',
                'scope_label': COURSE_SCOPE_LABELS.get(related_session.course_scope, 'Part'),
                'headers_with_meta': headers_with_meta,
                'student_report_data': student_report_data,
                'part_total_classes': attendance_summary.get('per_session_totals', {}).get(related_session.id, 0),
                'unique_dates': sorted(attendance_by_date.keys(), reverse=True)  # For delete modal
            })
        
        # Get unique dates from all sessions for delete modal
        all_attendance_records = ClassAttendance.query.filter(ClassAttendance.session_id.in_(all_session_ids)).all()
        attendance_by_date_all = defaultdict(list)
        for record in all_attendance_records:
            attendance_by_date_all[record.date].append(record)
        unique_dates_for_modal = sorted(attendance_by_date_all.keys(), reverse=True)
        
        return render_template(
            'class_management/view_attendance.html',
            session=session,
            headers=[],
            headers_with_meta=[],
            student_report_data=[],
            unique_dates=unique_dates_for_modal,
            split_meta=_build_split_context(session, attendance_summary),
            attendance_summary=attendance_summary,
            part_total_classes=part_total_classes,
            is_split_course=True,
            teacher_attendance_data=teacher_attendance_data
        )
    
    # Non-split course: original logic
    students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()
    all_attendance_records = ClassAttendance.query.filter_by(session_id=session_id).order_by(ClassAttendance.date, ClassAttendance.id).all()
    student_lookup = {stu.id: stu for stu in students}

    headers_with_meta = []
    if not all_attendance_records:
        return render_template(
            'class_management/view_attendance.html',
            session=session,
            students=students,
            headers=[],
            headers_with_meta=headers_with_meta,
            student_report_data=[],
            unique_dates=[],
            split_meta=_build_split_context(session, attendance_summary),
            attendance_summary=attendance_summary,
            part_total_classes=part_total_classes,
            is_split_course=False,
            teacher_attendance_data=[]
        )

    attendance_by_date = defaultdict(list)
    for record in all_attendance_records:
        attendance_by_date[record.date].append(record)

    unique_dates_for_modal = sorted(attendance_by_date.keys(), reverse=True)

    daily_class_counts = {}
    for date, records in attendance_by_date.items():
        student_counts_on_date = defaultdict(int)
        for record in records:
            student_counts_on_date[record.student_id] += 1
        daily_class_counts[date] = max(student_counts_on_date.values()) if student_counts_on_date else 0

    headers = []
    headers_with_meta = []
    sorted_dates = sorted(daily_class_counts.keys())
    for date in sorted_dates:
        count = daily_class_counts.get(date, 0)
        if count == 1:
            headers.append(date.strftime('%b %d, %Y'))
            headers_with_meta.append({'label': date.strftime('%b %d, %Y'), 'date': date.strftime('%Y-%m-%d'), 'slot': 1})
        else:
            for i in range(1, count + 1):
                headers.append(f"{date.strftime('%b %d')} ({i})")
                headers_with_meta.append({'label': f"{date.strftime('%b %d')} ({i})", 'date': date.strftime('%Y-%m-%d'), 'slot': i})

    student_report_data = []
    agg_student_map = attendance_summary.get('per_student', {})
    agg_total_classes = attendance_summary.get('total_classes', sum(daily_class_counts.values()))

    for student in students:
        student_records = [r for r in all_attendance_records if r.student_id == student.id]
        agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})

        attendance_row = []
        student_attendance_by_date = defaultdict(list)
        for r in student_records:
            student_attendance_by_date[r.date].append(r)

        for date in sorted_dates:
            records_for_date = student_attendance_by_date[date]
            num_classes_on_day = daily_class_counts.get(date, 0)
            for i in range(num_classes_on_day):
                cell = {
                    'status': '-',
                    'date': date.strftime('%Y-%m-%d'),
                    'slot': i + 1
                }
                if i < len(records_for_date):
                    record = records_for_date[i]
                    cell['status'] = 'P' if record.is_present else 'A'
                    cell['record_id'] = record.id
                attendance_row.append(cell)

        student_data = {
            'info': student,
            'attendance_row': attendance_row,
            'total_classes': agg_total_classes,
            'present_count': agg_stats['present'],
            'percentage': f"{agg_stats['percentage']:.2f}%",
            'marks': agg_stats['marks']
        }
        student_report_data.append(student_data)

    unique_headers_with_metadata = []
    header_index = 0
    for date in sorted_dates:
        count = daily_class_counts.get(date, 0)
        if count <= 1:
            unique_headers_with_metadata.append({'label': date.strftime('%b %d, %Y'), 'date': date.strftime('%Y-%m-%d'), 'slot': 1})
        else:
            for i in range(1, count + 1):
                unique_headers_with_metadata.append({'label': f"{date.strftime('%b %d')} ({i})", 'date': date.strftime('%Y-%m-%d'), 'slot': i})
    return render_template(
        'class_management/view_attendance.html',
        session=session,
        headers=headers,
        headers_with_meta=unique_headers_with_metadata,
        student_report_data=student_report_data,
        unique_dates=unique_dates_for_modal,
        split_meta=_build_split_context(session, attendance_summary),
        attendance_summary=attendance_summary,
        part_total_classes=part_total_classes,
        is_split_course=False,
        teacher_attendance_data=[]
    )


@class_management_bp.route('/toggle_attendance_record/<int:record_id>', methods=['POST'])
@login_required
def toggle_attendance_record(record_id):
    record = ClassAttendance.query.get_or_404(record_id)
    session = Session.query.get_or_404(record.session_id)
    teacher = _ensure_current_teacher()
    if session.teacher_id != teacher.id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403

    record.is_present = not record.is_present
    db.session.commit()

    attendance_summary = _build_attendance_summary(session)
    student_public_id = record.student.student_id
    student_stats = attendance_summary.get('per_student', {}).get(student_public_id, {'present': 0, 'percentage': 0, 'marks': 0})

    # Emit WebSocket event for live update
    try:
        from utils.websocket_events import emit_attendance_update
        emit_attendance_update(record.session_id, {
            'record_id': record_id,
            'status': 'P' if record.is_present else 'A',
            'student_id': student_public_id,
            'student_db_id': record.student_id,
            'present_count': student_stats.get('present', 0),
            'percentage': student_stats.get('percentage', 0),
            'marks': student_stats.get('marks', 0)
        })
    except Exception as e:
        current_app.logger.warning(f'Failed to emit attendance update event: {e}')

    return jsonify({
        'success': True,
        'status': 'P' if record.is_present else 'A',
        'present_count': student_stats.get('present', 0),
        'percentage': f"{student_stats.get('percentage', 0):.2f}%",
        'marks': student_stats.get('marks', 0),
        'marks_manual': student_stats.get('marks_manual', False),
        'student_db_id': record.student_id
    })

@class_management_bp.route('/save_attendance_marks_manual/<int:session_id>', methods=['POST'])
@login_required
def save_attendance_marks_manual(session_id):
    """Save manual attendance marks override for students"""
    session = Session.query.get_or_404(session_id)
    teacher = _ensure_current_teacher()
    if session.teacher_id != teacher.id:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    try:
        data = request.get_json()
        student_id = data.get('student_id')  # Public student_id (string)
        marks_value = data.get('marks')
        
        if not student_id:
            return jsonify({'success': False, 'message': 'Student ID required'}), 400
        
        # Find the student in this session
        student = ClassStudent.query.filter_by(
            session_id=session_id,
            student_id=student_id
        ).first()
        
        if not student:
            return jsonify({'success': False, 'message': 'Student not found'}), 404
        
        # For split courses, set manual marks on all related student records with the same student_id
        # This ensures manual marks are preserved across all sessions
        related_sessions = _get_related_sessions(session)
        session_ids = [s.id for s in related_sessions if s]
        all_students_with_same_id = ClassStudent.query.filter(
            ClassStudent.student_id == student_id,
            ClassStudent.session_id.in_(session_ids)
        ).all()
        
        # Set manual marks (None to clear manual override)
        if marks_value is not None and marks_value != '':
            try:
                marks_float = float(marks_value)
                # Set manual marks on all related student records
                for s in all_students_with_same_id:
                    s.attendance_marks_manual = marks_float
            except (ValueError, TypeError):
                return jsonify({'success': False, 'message': 'Invalid marks value'}), 400
        else:
            # Clear manual override from all related student records
            for s in all_students_with_same_id:
                s.attendance_marks_manual = None
        
        db.session.commit()
        
        # Rebuild attendance summary to get updated marks
        attendance_summary = _build_attendance_summary(session)
        student_stats = attendance_summary.get('per_student', {}).get(student_id, {'present': 0, 'percentage': 0, 'marks': 0, 'marks_manual': False})
        
        return jsonify({
            'success': True,
            'marks': student_stats.get('marks', 0),
            'marks_manual': student_stats.get('marks_manual', False),
            'present_count': student_stats.get('present', 0),
            'percentage': f"{student_stats.get('percentage', 0):.2f}%"
        })
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving manual attendance marks: {e}", exc_info=True)
        return jsonify({'success': False, 'message': str(e)}), 500

@class_management_bp.route('/students/<int:session_id>')
@login_required
def students_list(session_id):
    """View students list for a session"""
    session = Session.query.get_or_404(session_id)
    students = ClassStudent.query.filter_by(session_id=session_id).all()
    
    # Get all batches for filter dropdown
    batches = []
    if Student:
        try:
            all_batches = db.session.query(Student.batch).distinct().filter(Student.batch.isnot(None)).order_by(Student.batch.desc()).all()
            batches = [batch[0] for batch in all_batches]
        except Exception:
            batches = []
    
    return render_template('class_management/students_list.html', 
                         session=session, students=students,
                         split_meta=_build_split_context(session),
                         batches=batches)

@class_management_bp.route('/api/students', methods=['GET'])
@login_required
def get_students_for_selection():
    """Get students from Students Management for selection (AJAX)"""
    try:
        if not Student:
            current_app.logger.warning('Student model not available in get_students_for_selection')
            return jsonify({'success': False, 'message': 'Students Management module not available'}), 503
        
        batch_filter = request.args.get('batch', '').strip()
        search = request.args.get('search', '').strip()
        
        current_app.logger.info(f'Fetching students - batch: {batch_filter}, search: {search}')
        
        query = Student.query
        
        if batch_filter:
            query = query.filter(Student.batch == batch_filter)
        
        if search:
            query = query.filter(
                or_(
                    Student.name.ilike(f'%{search}%'),
                    Student.student_id.ilike(f'%{search}%')
                )
            )
        
        # Increase limit to show more students (or remove limit entirely)
        students = query.order_by(Student.student_id.asc()).limit(500).all()
        
        current_app.logger.info(f'Found {len(students)} students')
        
        return jsonify({
            'success': True,
            'students': [{
                'id': s.id,
                'student_id': s.student_id,
                'name': s.name,
                'batch': s.batch,
                'email': s.email,
                'phone': s.phone
            } for s in students]
        })
    except Exception as e:
        current_app.logger.error(f'Error in get_students_for_selection: {str(e)}', exc_info=True)
        return jsonify({'success': False, 'message': f'Error loading students: {str(e)}'}), 500

@class_management_bp.route('/api/curricula', methods=['GET'])
@login_required
def get_curricula_by_batch():
    """Get curricula applicable to a specific batch (AJAX)"""
    if not Curriculum:
        current_app.logger.error('Curriculum model not available')
        return jsonify({'success': False, 'message': 'Curriculum Management module not available'}), 503
    
    batch = request.args.get('batch', '').strip()
    if not batch:
        return jsonify({'success': False, 'message': 'Batch is required'}), 400
    
    try:
        # Find curricula where the batch is in applicable_batches
        all_curricula = Curriculum.query.all()
        applicable_curricula = []
        
        # Normalize the input batch
        normalized_batch = str(batch).strip()
        current_app.logger.info(f'Searching curricula for batch: {normalized_batch}, Total curricula: {len(all_curricula)}')
        
        for curriculum in all_curricula:
            batches_list = curriculum.get_batches_list()
            current_app.logger.debug(f'Curriculum {curriculum.id} ({curriculum.name}) has batches: {batches_list}')
            
            # Check if the batch matches any batch in the list
            for b in batches_list:
                if str(b).strip() == normalized_batch:
                    applicable_curricula.append({
                        'id': curriculum.id,
                        'name': curriculum.name,
                        'date': curriculum.date
                    })
                    current_app.logger.info(f'Found matching curriculum: {curriculum.name} (ID: {curriculum.id})')
                    break  # Found a match, no need to check other batches for this curriculum
        
        current_app.logger.info(f'Found {len(applicable_curricula)} applicable curricula for batch {normalized_batch}')
        
        return jsonify({
            'success': True,
            'curricula': applicable_curricula,
            'batch_searched': normalized_batch,
            'total_curricula_checked': len(all_curricula),
            'debug': {
                'batch_received': batch,
                'normalized_batch': normalized_batch,
                'applicable_count': len(applicable_curricula)
            }
        })
    except Exception as e:
        current_app.logger.error(f'Error in get_curricula_by_batch: {str(e)}', exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error fetching curricula: {str(e)}',
            'error': str(e)
        }), 500

@class_management_bp.route('/api/curriculum/<int:curriculum_id>/years-terms', methods=['GET'])
@login_required
def get_years_terms_by_curriculum(curriculum_id):
    """Get distinct (display) years and terms from courses in a curriculum (AJAX)"""
    if not Course:
        return jsonify({'success': False, 'message': 'Course Management module not available'}), 503
    
    Curriculum.query.get_or_404(curriculum_id)
    courses = Course.query.filter_by(curriculum_id=curriculum_id, offered=True).all()  # Only offered courses
    
    years = sorted({c.display_year for c in courses if getattr(c, 'display_year', None)}, key=lambda x: x or '')
    terms = sorted({c.display_term for c in courses if getattr(c, 'display_term', None)}, key=lambda x: x or '')
    
    return jsonify({
        'success': True,
        'years': years,
        'terms': terms
    })

@class_management_bp.route('/api/courses', methods=['GET'])
@login_required
def get_courses_by_filters():
    """Get courses filtered by curriculum, year, and term (AJAX)"""
    if not Course:
        return jsonify({'success': False, 'message': 'Course Management module not available'}), 503
    
    curriculum_id = request.args.get('curriculum_id', type=int)
    year = request.args.get('year', '').strip()
    term = request.args.get('term', '').strip()
    
    query = Course.query.filter_by(offered=True)  # Only show offered courses
    if curriculum_id:
        query = query.filter_by(curriculum_id=curriculum_id)
    
    courses = query.order_by(Course.course_name.asc()).all()
    
    if year:
        courses = [c for c in courses if c.display_year == year]
    if term:
        courses = [c for c in courses if c.display_term == term]

    # Filter only Theory courses
    courses = [c for c in courses if (c.course_type or '').lower() == 'theory']
    
    return jsonify({
        'success': True,
        'courses': [{
            'id': c.id,
            'course_code': c.course_code,
            'course_name': c.course_name,
            'year': c.display_year,
            'term': c.display_term,
            'credit': c.credit,
            'course_type': c.course_type,
            'category': c.category
        } for c in courses]
    })

@class_management_bp.route('/api/academic-sessions', methods=['GET'])
@login_required
def get_academic_sessions():
    """Get academic sessions from CurriculumYearTerm based on curriculum, year, and term (AJAX)"""
    if not CurriculumYearTerm:
        return jsonify({'success': False, 'message': 'Course Management module not available'}), 503
    
    curriculum_id = request.args.get('curriculum_id', type=int)
    year = request.args.get('year', '').strip()
    term = request.args.get('term', '').strip()
    
    if not curriculum_id or not year or not term:
        return jsonify({
            'success': True,
            'academic_sessions': []
        })
    
    # Fetch academic session from CurriculumYearTerm
    config = CurriculumYearTerm.query.filter_by(
        curriculum_id=curriculum_id,
        year=year,
        term=term
    ).first()
    
    academic_sessions = []
    if config and config.academic_session:
        academic_sessions.append(config.academic_session)
    
    # Also fetch distinct academic sessions from all CurriculumYearTerm for this curriculum/year/term
    # in case there are multiple entries (shouldn't happen due to unique constraint, but for safety)
    all_sessions = db.session.query(CurriculumYearTerm.academic_session).filter_by(
        curriculum_id=curriculum_id,
        year=year,
        term=term
    ).filter(CurriculumYearTerm.academic_session.isnot(None)).distinct().all()
    
    unique_sessions = sorted({s[0] for s in all_sessions if s[0]})
    
    return jsonify({
        'success': True,
        'academic_sessions': unique_sessions
    })

@class_management_bp.route('/add_student/<int:session_id>', methods=['POST'])
@login_required
def add_student(session_id):
    """Add students to a session from Students Management"""
    session = Session.query.get_or_404(session_id)
    teacher = _ensure_current_teacher()
    
    # Handle AJAX request (multiple students)
    if request.is_json:
        if not Student:
            return jsonify({'success': False, 'message': 'Students Management module not available'}), 503
        
        data = request.get_json()
        student_ids = data.get('student_ids', [])
        
        if not student_ids:
            return jsonify({'success': False, 'message': 'No students selected!'}), 400
        
        added_count = 0
        skipped_count = 0
        not_registered_count = 0
        
        # Get existing student IDs in this session
        existing_student_ids = {s.student_id for s in ClassStudent.query.filter_by(session_id=session_id).all()}
        
        for student_id in student_ids:
            # Get student from Students Management
            student = Student.query.get(student_id)
            if not student:
                continue
            
            # Check if already in session
            if student.student_id in existing_student_ids:
                skipped_count += 1
                continue
            
            # Check if student is registered for this course (finalized registration only)
            if StudentCourseRegistration and session.course_code and session.academic_session and session.year and session.term:
                registration = StudentCourseRegistration.query.filter_by(
                    student_id=student.id,
                    course_code=session.course_code,
                    academic_session=session.academic_session,
                    year=session.year,
                    term=session.term,
                    status='finalized'
                ).first()
                
                if not registration:
                    not_registered_count += 1
                    current_app.logger.info(f'Student {student.student_id} ({student.name}) not registered for course {session.course_code}, skipping...')
                    continue
            
            class_student = ClassStudent(
                student_id=student.student_id,
                name=student.name,
                session_id=session.id,
                teacher_id=session.teacher_id or teacher.id
            )
            db.session.add(class_student)
            db.session.flush()  # Flush to get class_student.id before carry on
            
            # Carry on assessment marks if enabled in registration
            _carry_on_assessment_marks(class_student, session)
            
            _replicate_student_to_peers(session, class_student)
            existing_student_ids.add(student.student_id)
            added_count += 1
        
        try:
            db.session.commit()
            message = f'Successfully added {added_count} student(s).'
            if skipped_count > 0:
                message += f' Skipped {skipped_count} existing student(s).'
            if not_registered_count > 0:
                message += f' Skipped {not_registered_count} student(s) not registered for this course.'
            return jsonify({'success': True, 'message': message})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': f'Error adding students: {str(e)}'}), 500
    
    # Handle form submission (backward compatibility)
    student_id = request.form.get('student_id')
    name = request.form.get('name')
    
    if not student_id or not name:
        flash('Student ID and name are required!', 'error')
        return redirect(url_for('class_management.students_list', session_id=session_id))
    
    # Check if already exists
    existing = ClassStudent.query.filter_by(session_id=session_id, student_id=student_id).first()
    if existing:
        flash('Student already exists in this session!', 'error')
        return redirect(url_for('class_management.students_list', session_id=session_id))
    
    student = ClassStudent(
        student_id=student_id,
        name=name,
        session_id=session.id,
        teacher_id=session.teacher_id or teacher.id
    )
    db.session.add(student)
    db.session.flush()  # Flush to get student.id before carry on
    
    # Carry on assessment marks if enabled in registration
    _carry_on_assessment_marks(student, session)
    
    _replicate_student_to_peers(session, student)
    db.session.commit()
    flash('Student added successfully!', 'success')
    return redirect(url_for('class_management.students_list', session_id=session_id))

@class_management_bp.route('/edit_student/<int:student_id>', methods=['POST'])
@login_required
def edit_student(student_id):
    """Edit a student"""
    student = ClassStudent.query.get_or_404(student_id)
    session = student.session
    old_identifier = student.student_id
    student.student_id = request.form.get('student_id')
    student.name = request.form.get('name')
    _replicate_student_to_peers(session, student, old_identifier=old_identifier)
    db.session.commit()
    flash('Student updated successfully!', 'success')
    return redirect(url_for('class_management.students_list', session_id=student.session_id))

@class_management_bp.route('/delete_student/<int:student_id>', methods=['POST'])
@login_required
def delete_student(student_id):
    """Delete a student from a session"""
    try:
        student = ClassStudent.query.get_or_404(student_id)
        session_id = student.session_id
        session = student.session
        student_identifier = student.student_id
        
        # Delete student from peer sessions (split courses)
        _delete_student_from_peers(session, student_identifier)
        
        # Delete the student (cascade will handle related attendance records)
        db.session.delete(student)
        db.session.commit()
        flash('Student deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error deleting student {student_id}: {e}', exc_info=True)
        flash(f'Error deleting student: {str(e)}', 'danger')
        # Try to get session_id for redirect even if deletion failed
        try:
            student = ClassStudent.query.get(student_id)
            session_id = student.session_id if student else None
        except:
            session_id = None
    
    if session_id:
        return redirect(url_for('class_management.students_list', session_id=session_id))
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/delete_session/<int:session_id>', methods=['POST'])
@login_required
def delete_session(session_id):
    """Delete a session and all related records using SQLAlchemy ORM (database-agnostic)"""
    try:
        # Check authorization: admin/head can delete any session, regular teachers only their own
        user_roles = set(parse_roles(getattr(current_user, 'role', '')))
        if getattr(current_user, 'active_role', None):
            user_roles = set(parse_roles(current_user.active_role))
        can_delete_all = is_admin(current_user) or 'head' in user_roles or 'dean' in user_roles
        
        session = Session.query.get_or_404(session_id)
        
        # If not admin/head, check if this session belongs to the current teacher
        if not can_delete_all:
            current_teacher = _ensure_current_teacher()
            if not current_teacher or session.teacher_id != current_teacher.id:
                flash('You do not have permission to delete this session.', 'danger')
                return redirect(url_for('class_management.index'))
        
        # Import BatchCustomEvent if available
        try:
            from blueprints.academic_calendar.models import BatchCustomEvent
        except ImportError:
            BatchCustomEvent = None
        
        # Delete all related records in correct order (respecting foreign key constraints)
        try:
            # 1. Delete student feedback responses first (before feedback links)
            feedback_link_ids = [link.id for link in StudentFeedbackLink.query.filter_by(session_id=session_id).all()]
            if feedback_link_ids:
                StudentFeedbackResponse.query.filter(
                    StudentFeedbackResponse.feedback_link_id.in_(feedback_link_ids)
                ).delete(synchronize_session=False)
            
            # 2. Delete student feedback links
            StudentFeedbackLink.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 3. Delete batch custom events (if model exists)
            if BatchCustomEvent:
                BatchCustomEvent.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 4. Delete course outline
            CourseOutline.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 5. Delete evaluation submissions
            EvaluationSubmission.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 6. Delete evaluation invites
            EvaluationInvite.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 7. Delete course reviews
            CourseReview.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 8. Delete split course invites (where this session is the inviter)
            ClassSplitInvite.query.filter_by(inviter_session_id=session_id).delete(synchronize_session=False)
            
            # 9. Delete class attendance (cascade will handle if relationship is set up)
            ClassAttendance.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 10. Delete class students (cascade will handle if relationship is set up)
            ClassStudent.query.filter_by(session_id=session_id).delete(synchronize_session=False)
            
            # 11. Finally delete the session itself
            db.session.delete(session)
            db.session.commit()
            
            current_app.logger.info(f'Session {session_id} and all related data deleted successfully')
            flash('Session deleted successfully!', 'success')
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error deleting session {session_id}: {e}', exc_info=True)
            flash(f'Error deleting session: {str(e)}', 'danger')
            
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error deleting session {session_id}: {e}', exc_info=True)
        flash(f'Error deleting session: {str(e)}', 'danger')
    
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/course_file/<int:session_id>')
@login_required
def course_file(session_id):
    """Course file management page"""
    session = Session.query.get_or_404(session_id)
    
    # Get or create course outline
    course_outline = None
    try:
        course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
        if not course_outline:
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if teacher:
                # Ensure all columns exist before creating
                try:
                    db.create_all()
                except Exception as e:
                    current_app.logger.warning(f"Could not create all tables/columns: {e}")
                
                course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
                db.session.add(course_outline)
                db.session.commit()
    except Exception as e:
        # If table doesn't exist, create it
        current_app.logger.warning(f"CourseOutline table might not exist: {e}")
        try:
            db.create_all()
            # Try again after creating tables
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if teacher:
                course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
                db.session.add(course_outline)
                db.session.commit()
        except Exception as e2:
            current_app.logger.error(f"Error creating CourseOutline: {e2}")
            flash('Course outline feature is not available. Please ensure database migration is complete.', 'warning')
    
    # Get course data from curriculum if available
    course_data = find_course_from_curriculum(session.course_code, session.course_name)
    
    # Get uploaded files for this session
    from blueprints.class_management.models import CourseFileUpload
    uploaded_files = CourseFileUpload.query.filter_by(session_id=session_id).order_by(CourseFileUpload.created_at.desc()).all()
    
    return render_template('class_management/course_file.html', 
                         session=session, 
                         course_outline=course_outline,
                         course_data=course_data,
                         uploaded_files=uploaded_files)

@class_management_bp.route('/course_file/<int:session_id>/save', methods=['POST'])
@login_required
def save_course_outline(session_id):
    """Save course outline data"""
    try:
        session = Session.query.get_or_404(session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        if not teacher or teacher.id != session.teacher_id:
            if request.is_json:
                return jsonify({'success': False, 'message': 'You are not authorized to edit this course outline.'}), 403
            flash('You are not authorized to edit this course outline.', 'danger')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
        if not course_outline:
            course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
            db.session.add(course_outline)
        
        # Ensure all columns exist (for database migrations)
        try:
            db.create_all()  # This will add missing columns if database supports it
        except Exception as e:
            current_app.logger.warning(f"Could not create all tables/columns: {e}")
        
        # Save all form data as JSON
        data = request.get_json() if request.is_json else request.form.to_dict()
        
        if not data:
            if request.is_json:
                return jsonify({'success': False, 'message': 'No data received.'}), 400
            flash('No data received.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        if 'course_objectives' in data:
            course_outline.course_objectives = json.dumps(data.get('course_objectives', [])) if isinstance(data.get('course_objectives'), list) else data.get('course_objectives')
        if 'course_summary' in data:
            course_outline.course_summary = data.get('course_summary')
        if 'prerequisites' in data:
            course_outline.prerequisites = data.get('prerequisites')
        if 'contact_hours' in data:
            course_outline.contact_hours = data.get('contact_hours')
        if 'cie_marks' in data:
            course_outline.cie_marks = data.get('cie_marks')
        if 'smee_marks' in data:
            course_outline.smee_marks = data.get('smee_marks')
        # Try to set new fields, but handle if columns don't exist yet
        try:
            if 'credit_value' in data:
                course_outline.credit_value = data.get('credit_value')
            if 'course_type' in data:
                course_outline.course_type = data.get('course_type')
            if 'level_term_section' in data:
                course_outline.level_term_section = data.get('level_term_section')
            if 'clo_data' in data:
                course_outline.clo_data = json.dumps(data.get('clo_data', []))
            if 'plo_mapping' in data:
                course_outline.plo_mapping = json.dumps(data.get('plo_mapping', {}))
        except AttributeError:
            current_app.logger.warning("Some new course_outline columns don't exist yet. Please run migration.")
        
        if 'course_content_summary' in data:
            # Store as JSON format
            content_summary = data.get('course_content_summary')
            if content_summary:
                try:
                    # Try to parse as JSON
                    content_data = json.loads(content_summary)
                    if isinstance(content_data, dict):
                        course_outline.course_content_summary = json.dumps(content_data)
                    else:
                        # Store as text if not valid JSON
                        course_outline.course_content_summary = content_summary
                except (json.JSONDecodeError, TypeError):
                    # Not JSON, store as text
                    course_outline.course_content_summary = content_summary
            else:
                course_outline.course_content_summary = None
        if 'clo_plo_mapping' in data:
            course_outline.clo_plo_mapping = data.get('clo_plo_mapping')
        if 'evaluation_policy' in data:
            course_outline.evaluation_policy = json.dumps(data.get('evaluation_policy', {})) if isinstance(data.get('evaluation_policy'), dict) else data.get('evaluation_policy')
        if 'cie_breakdown' in data:
            cie_breakdown = data.get('cie_breakdown', [])
            if isinstance(cie_breakdown, (list, dict)):
                course_outline.cie_breakdown = json.dumps(cie_breakdown)
            else:
                course_outline.cie_breakdown = cie_breakdown
        if 'smee_breakdown' in data:
            smee_breakdown = data.get('smee_breakdown', [])
            if isinstance(smee_breakdown, (list, dict)):
                course_outline.smee_breakdown = json.dumps(smee_breakdown)
            else:
                course_outline.smee_breakdown = smee_breakdown
        
        if 'lesson_plan' in data:
            course_outline.lesson_plan = json.dumps(data.get('lesson_plan', [])) if isinstance(data.get('lesson_plan'), list) else data.get('lesson_plan')
        if 'assessment_strategy' in data:
            course_outline.assessment_strategy = json.dumps(data.get('assessment_strategy', {})) if isinstance(data.get('assessment_strategy'), dict) else data.get('assessment_strategy')
        if 'assessment_techniques' in data:
            course_outline.assessment_techniques = json.dumps(data.get('assessment_techniques', [])) if isinstance(data.get('assessment_techniques'), list) else data.get('assessment_techniques')
        if 'rubrics' in data:
            course_outline.rubrics = json.dumps(data.get('rubrics', [])) if isinstance(data.get('rubrics'), list) else data.get('rubrics')
        if 'grading_policy' in data:
            course_outline.grading_policy = json.dumps(data.get('grading_policy', [])) if isinstance(data.get('grading_policy'), list) else data.get('grading_policy')
        if 'textbooks' in data:
            course_outline.textbooks = json.dumps(data.get('textbooks', [])) if isinstance(data.get('textbooks'), list) else data.get('textbooks')
        if 'reference_books' in data:
            course_outline.reference_books = json.dumps(data.get('reference_books', [])) if isinstance(data.get('reference_books'), list) else data.get('reference_books')
        if 'other_resources' in data:
            course_outline.other_resources = json.dumps(data.get('other_resources', [])) if isinstance(data.get('other_resources'), list) else data.get('other_resources')
        try:
            if 'course_file_components' in data:
                course_outline.course_file_components = json.dumps(data.get('course_file_components', [])) if isinstance(data.get('course_file_components'), list) else data.get('course_file_components')
        except AttributeError:
            current_app.logger.warning("course_file_components column doesn't exist yet.")
        
        if 'make_up_procedures' in data:
            course_outline.make_up_procedures = data.get('make_up_procedures')
        if 'other_issues' in data:
            course_outline.other_issues = json.dumps(data.get('other_issues', {})) if isinstance(data.get('other_issues'), dict) else data.get('other_issues')
        
        try:
            # Always process student_access_enabled
            if 'student_access_enabled' in data:
                # Handle both boolean and string values
                student_access = data.get('student_access_enabled', False)
                if isinstance(student_access, str):
                    student_access = student_access.lower() in ('true', '1', 'on', 'yes')
                course_outline.student_access_enabled = bool(student_access)
                current_app.logger.info(f"Setting student_access_enabled to {course_outline.student_access_enabled} for session {session_id} (received value: {data.get('student_access_enabled')}, type: {type(data.get('student_access_enabled'))})")
            else:
                # If not provided, keep existing value (don't reset to False)
                current_app.logger.info(f"student_access_enabled not in data for session {session_id}, keeping existing value: {course_outline.student_access_enabled}")
        except AttributeError:
            current_app.logger.warning("student_access_enabled column doesn't exist yet.")
        except Exception as e:
            current_app.logger.error(f"Error setting student_access_enabled: {e}", exc_info=True)
        
        try:
            db.session.commit()
            current_app.logger.info(f"Course outline saved successfully for session {session_id}")
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error committing course outline for session {session_id}: {e}", exc_info=True)
            if request.is_json:
                return jsonify({'success': False, 'message': f'Database error: {str(e)}'}), 500
            flash(f'Error saving course outline: {str(e)}', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error saving course outline for session {session_id}: {e}", exc_info=True)
        if request.is_json:
            return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500
        flash(f'Error saving course outline: {str(e)}', 'error')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    if request.is_json:
        return jsonify({'success': True, 'message': 'Course outline saved successfully!'})
    flash('Course outline saved successfully!', 'success')
    return redirect(url_for('class_management.course_file', session_id=session_id))

@class_management_bp.route('/course_file/<int:session_id>/upload', methods=['POST'])
@login_required
def upload_course_file(session_id):
    """Upload course file"""
    try:
        import os
        from werkzeug.utils import secure_filename
        from blueprints.class_management.models import CourseFileUpload
        
        session = Session.query.get_or_404(session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        if not teacher:
            flash('Teacher not found.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Check authorization
        if teacher.id != session.teacher_id:
            flash('You are not authorized to upload files for this course.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Get form data
        file_name = request.form.get('file_name', '').strip()
        file_category = request.form.get('file_category', 'other')
        description = request.form.get('description', '').strip()
        
        if not file_name:
            flash('File name is required.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Get uploaded file
        if 'file' not in request.files:
            flash('No file selected.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        file = request.files['file']
        if file.filename == '':
            flash('No file selected.', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Validate file extension
        allowed_extensions = {'.pdf', '.doc', '.docx', '.ppt', '.pptx', '.xls', '.xlsx', 
                            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_extensions:
            flash(f'File type not allowed. Supported formats: PDF, DOC, DOCX, PPT, PPTX, XLS, XLSX, images', 'error')
            return redirect(url_for('class_management.course_file', session_id=session_id))
        
        # Create upload directory if it doesn't exist
        upload_dir = os.path.join(UPLOAD_FOLDER, 'course_files', str(session_id))
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate secure filename
        secure_name = secure_filename(file.filename)
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{secure_name}"
        file_path = os.path.join(upload_dir, filename)
        
        # Save file
        file.save(file_path)
        file_size = os.path.getsize(file_path)
        
        # Get file type (MIME type or extension)
        file_type = file_ext[1:] if file_ext else 'unknown'
        
        # Create CourseFileUpload record
        uploaded_file = CourseFileUpload(
            session_id=session_id,
            teacher_id=teacher.id,
            file_name=file_name,
            file_path=file_path,
            file_size=file_size,
            file_type=file_type,
            description=description,
            student_access_enabled=True
        )
        
        db.session.add(uploaded_file)
        db.session.commit()
        
        current_app.logger.info(f"Course file uploaded successfully: {file_name} for session {session_id}")
        flash(f'File "{file_name}" uploaded successfully!', 'success')
        return redirect(url_for('class_management.course_file', session_id=session_id))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error uploading course file for session {session_id}: {e}", exc_info=True)
        flash(f'Error uploading file: {str(e)}', 'error')
        return redirect(url_for('class_management.course_file', session_id=session_id))

@class_management_bp.route('/course_file/<int:file_id>/download')
@login_required
def download_course_file(file_id):
    """Download course file"""
    try:
        import os
        from flask import send_file
        from blueprints.class_management.models import CourseFileUpload
        
        uploaded_file = CourseFileUpload.query.get_or_404(file_id)
        session = Session.query.get_or_404(uploaded_file.session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        # Check authorization
        if not teacher or teacher.id != session.teacher_id:
            flash('You are not authorized to download this file.', 'error')
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
        # Check if file exists
        if not os.path.exists(uploaded_file.file_path):
            flash('File not found.', 'error')
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
        return send_file(
            uploaded_file.file_path,
            as_attachment=True,
            download_name=uploaded_file.file_name
        )
    except Exception as e:
        current_app.logger.error(f"Error downloading course file {file_id}: {e}", exc_info=True)
        flash(f'Error downloading file: {str(e)}', 'error')
        if 'uploaded_file' in locals():
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/course_file/<int:file_id>/delete', methods=['POST'])
@login_required
def delete_course_file(file_id):
    """Delete course file"""
    try:
        import os
        from blueprints.class_management.models import CourseFileUpload
        
        uploaded_file = CourseFileUpload.query.get_or_404(file_id)
        session = Session.query.get_or_404(uploaded_file.session_id)
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        # Check authorization
        if not teacher or teacher.id != session.teacher_id:
            flash('You are not authorized to delete this file.', 'error')
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
        # Delete file from filesystem
        if os.path.exists(uploaded_file.file_path):
            try:
                os.remove(uploaded_file.file_path)
            except Exception as e:
                current_app.logger.warning(f"Could not delete file from filesystem: {e}")
        
        # Delete record from database
        db.session.delete(uploaded_file)
        db.session.commit()
        
        flash(f'File "{uploaded_file.file_name}" deleted successfully.', 'success')
        return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting course file {file_id}: {e}", exc_info=True)
        flash(f'Error deleting file: {str(e)}', 'error')
        if 'uploaded_file' in locals():
            return redirect(url_for('class_management.course_file', session_id=uploaded_file.session_id))
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/course_file/<int:session_id>/outline/generate-ai', methods=['POST'])
@login_required
def generate_weekly_plan_ai(session_id):
    """Generate weekly plan using AI based on Course Content (section 14), Credit Value, and Academic Calendar"""
    from datetime import timedelta
    
    session = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    
    # Check if teacher is authorized (either main teacher or part of split group)
    is_authorized = False
    if teacher:
        if teacher.id == session.teacher_id:
            is_authorized = True
        elif session.split_group_id:
            related_sessions = Session.query.filter_by(split_group_id=session.split_group_id).all()
            for related_session in related_sessions:
                if related_session.teacher_id == teacher.id:
                    is_authorized = True
                    break
    
    if not is_authorized:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    
    data = request.get_json()
    course_content_summary = data.get('course_content_summary', '')  # JSON string from section 14
    credit_value = data.get('credit_value', '')  # Credit value from form
    part = data.get('part', None)  # 'A' or 'B' for split courses, None for full courses
    
    try:
        # Import Academic Calendar model
        try:
            from blueprints.academic_calendar.models import AcademicCalendarEvent
        except ImportError:
            AcademicCalendarEvent = None
            return jsonify({'success': False, 'message': 'Academic Calendar module not available'}), 503
        
        # Get credit value - try from data, then course outline, then course data
        credit = None
        if credit_value:
            try:
                credit = float(credit_value)
            except:
                pass
        
        if not credit:
            # Try to get from course outline
            course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
            if course_outline and course_outline.credit_value:
                try:
                    credit = float(course_outline.credit_value)
                except:
                    pass
        
        if not credit:
            # Try to get from course data
            course_data = find_course_from_curriculum(session.course_code, session.course_name)
            if course_data and course_data.credit:
                try:
                    credit = float(course_data.credit)
                except:
                    pass
        
        if not credit:
            return jsonify({'success': False, 'message': 'Credit value is required. Please fill in the Credit Value field.'}), 400
        
        # Parse Course Content (section 14) - Only use selected topics
        # For split courses, filter by part (A or B)
        course_contents = []
        if course_content_summary:
            try:
                content_data = json.loads(course_content_summary) if isinstance(course_content_summary, str) else course_content_summary
                if isinstance(content_data, dict):
                    # For split courses, use only the specified part (A or B)
                    # For full courses, combine both sections
                    if part == 'A':
                        section_a = content_data.get('sectionA', [])
                        all_contents = section_a
                    elif part == 'B':
                        section_b = content_data.get('sectionB', [])
                        all_contents = section_b
                    else:
                        # Full course: combine both sections
                        section_a = content_data.get('sectionA', [])
                        section_b = content_data.get('sectionB', [])
                        all_contents = section_a + section_b
                    
                    # Filter to only include items where selected is True (or undefined for backward compatibility)
                    course_contents = [
                        item for item in all_contents 
                        if item.get('selected', True)  # Default to True for backward compatibility
                    ]
                elif isinstance(content_data, list):
                    # Filter to only include selected topics
                    course_contents = [
                        item for item in content_data 
                        if item.get('selected', True)  # Default to True for backward compatibility
                    ]
            except Exception as e:
                current_app.logger.warning(f"Error parsing course content: {e}")
        
        # If no course content from section 14, try to get from curriculum
        if not course_contents:
            if course_data := find_course_from_curriculum(session.course_code, session.course_name):
                # For split courses, use only the specified part (A or B)
                # For full courses, combine both sections
                if part == 'A':
                    if course_data.content_section_a:
                        try:
                            content_a = json.loads(course_data.content_section_a) if isinstance(course_data.content_section_a, str) else course_data.content_section_a
                            if isinstance(content_a, list):
                                course_contents.extend(content_a)
                        except:
                            pass
                elif part == 'B':
                    if course_data.content_section_b:
                        try:
                            content_b = json.loads(course_data.content_section_b) if isinstance(course_data.content_section_b, str) else course_data.content_section_b
                            if isinstance(content_b, list):
                                course_contents.extend(content_b)
                        except:
                            pass
                else:
                    # Full course: combine both sections
                    if course_data.content_section_a:
                        try:
                            content_a = json.loads(course_data.content_section_a) if isinstance(course_data.content_section_a, str) else course_data.content_section_a
                            if isinstance(content_a, list):
                                course_contents.extend(content_a)
                        except:
                            pass
                    if course_data.content_section_b:
                        try:
                            content_b = json.loads(course_data.content_section_b) if isinstance(course_data.content_section_b, str) else course_data.content_section_b
                            if isinstance(content_b, list):
                                course_contents.extend(content_b)
                        except:
                            pass
        
        # Get CLOs
        clos = []
        course_data = find_course_from_curriculum(session.course_code, session.course_name)
        if course_data:
            clos = course_data.get_clos_list()
            
        # Normalize session year and term for matching
        session_year = str(session.year).strip().lower() if session.year else ''
        session_term = str(session.term).strip().lower() if session.term else ''
        
        # Map year/term to common formats for matching
        year_mapping = {
            '1': 'first', 'first': 'first', '1st': 'first',
            '2': 'second', 'second': 'second', '2nd': 'second',
            '3': 'third', 'third': 'third', '3rd': 'third',
            '4': 'fourth', 'fourth': 'fourth', '4th': 'fourth'
        }
        term_mapping = {
            '1': 'first', 'first': 'first', '1st': 'first',
            '2': 'second', 'second': 'second', '2nd': 'second'
        }
        
        normalized_year = year_mapping.get(session_year, session_year)
        normalized_term = term_mapping.get(session_term, session_term)
        
        # Get Academic Calendar events
        holidays = set()
        semester_start_date = None
        semester_end_date = None
        
        # Get current year and next year for broader search
        current_year = datetime.now().year
        year_start = date(current_year, 1, 1)
        year_end = date(current_year + 1, 12, 31)
        
        calendar_events = AcademicCalendarEvent.query.filter(
            AcademicCalendarEvent.event_date >= year_start,
            AcademicCalendarEvent.event_date <= year_end
        ).order_by(AcademicCalendarEvent.event_date.asc()).all()
        
        # Collect all holidays
        for event in calendar_events:
            if event.event_type == 'holiday':
                # Add all dates in the range if end_date exists
                if event.end_date and event.end_date > event.event_date:
                    current_date = event.event_date
                    while current_date <= event.end_date:
                        holidays.add(current_date)
                        current_date += timedelta(days=1)
        else:
                    holidays.add(event.event_date)
        
        # Add recurring Friday and Saturday holidays
        current_date = year_start
        while current_date <= year_end:
            if current_date.weekday() == 4:  # Friday
                holidays.add(current_date)
            elif current_date.weekday() == 5:  # Saturday
                holidays.add(current_date)
            current_date += timedelta(days=1)
        
        # Find semester start and end dates
        # First, check which session, year, and term this subject belongs to
        session_academic_session = session.academic_session or ''
        session_year = normalized_year
        session_term = normalized_term
        
        current_app.logger.info(f"Looking for semester dates for Session: {session_academic_session}, Year: {session_year}, Term: {session_term}")
        
        semester_start_events = []
        semester_end_events = []
        
        for event in calendar_events:
            if event.event_type == 'semester_start':
                semester_start_events.append(event)
            elif event.event_type == 'semester_end':
                semester_end_events.append(event)
        
        # Try to find matching semester based on session, year, term, and academic_session
        # Priority: 1. Match by academic_session + year + term in title/description
        #           2. Match by year + term in title/description
        #           3. Use most appropriate date based on session academic_session or current date
        if semester_start_events:
            matched_start = None
            
            # Priority 1: Match by academic_session + year + term
            if session_academic_session:
                for event in semester_start_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    # Check if academic_session, year, and term all appear in the event
                    if (session_academic_session.lower() in event_text and 
                        normalized_year in event_text and 
                        normalized_term in event_text):
                        matched_start = event
                        current_app.logger.info(f"Matched semester start by academic_session+year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 2: Match by year + term in title/description
            if not matched_start:
                for event in semester_start_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    if normalized_year in event_text and normalized_term in event_text:
                        matched_start = event
                        current_app.logger.info(f"Matched semester start by year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 3: Use the most recent past or upcoming start date
            if not matched_start:
                today = date.today()
                upcoming_starts = [e for e in semester_start_events if e.event_date >= today]
                if upcoming_starts:
                    matched_start = min(upcoming_starts, key=lambda x: x.event_date)
                    current_app.logger.info(f"Using upcoming semester start: {matched_start.title} on {matched_start.event_date}")
                else:
                    # Use most recent past
                    matched_start = max(semester_start_events, key=lambda x: x.event_date)
                    current_app.logger.info(f"Using most recent semester start: {matched_start.title} on {matched_start.event_date}")
            
            if matched_start:
                semester_start_date = matched_start.event_date
        
        if semester_end_events:
            matched_end = None
            
            # Priority 1: Match by academic_session + year + term
            if session_academic_session:
                for event in semester_end_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    # Check if academic_session, year, and term all appear in the event
                    if (session_academic_session.lower() in event_text and 
                        normalized_year in event_text and 
                        normalized_term in event_text):
                        matched_end = event
                        current_app.logger.info(f"Matched semester end by academic_session+year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 2: Match by year + term in title/description
            if not matched_end:
                for event in semester_end_events:
                    title_lower = (event.title or '').lower()
                    description_lower = (event.description or '').lower()
                    event_text = title_lower + ' ' + description_lower
                    
                    if normalized_year in event_text and normalized_term in event_text:
                        matched_end = event
                        current_app.logger.info(f"Matched semester end by year+term: {event.title} on {event.event_date}")
                        break
            
            # Priority 3: Use the end date that comes after the start date
            if not matched_end:
                if semester_start_date:
                    future_ends = [e for e in semester_end_events if e.event_date > semester_start_date]
                    if future_ends:
                        matched_end = min(future_ends, key=lambda x: x.event_date)
                        current_app.logger.info(f"Using future semester end after start: {matched_end.title} on {matched_end.event_date}")
                else:
                    # Use most recent past or upcoming
                    today = date.today()
                    upcoming_ends = [e for e in semester_end_events if e.event_date >= today]
                    if upcoming_ends:
                        matched_end = min(upcoming_ends, key=lambda x: x.event_date)
                        current_app.logger.info(f"Using upcoming semester end: {matched_end.title} on {matched_end.event_date}")
                    else:
                        matched_end = max(semester_end_events, key=lambda x: x.event_date)
                        current_app.logger.info(f"Using most recent semester end: {matched_end.title} on {matched_end.event_date}")
            
            if matched_end:
                semester_end_date = matched_end.event_date
        
        # Validate dates
        if not semester_start_date:
            return jsonify({
                'success': False,
                'message': 'Semester Start Date not found in Academic Calendar. Please add a "Semester Start Date" event.'
            }), 400
        
        if not semester_end_date:
            return jsonify({
                'success': False,
                'message': 'Semester End Date not found in Academic Calendar. Please add a "Semester End Date" event.'
            }), 400
        
        if semester_end_date <= semester_start_date:
            return jsonify({
                'success': False,
                'message': 'Semester End Date must be after Semester Start Date.'
            }), 400
        
        # Calculate weekly classes based on credit (1 credit = 1 class per week)
        # For split courses, divide classes equally between Part A and Part B
        if part in ['A', 'B']:
            # Split course: each part gets half the classes
            classes_per_week = int(credit) // 2
            if classes_per_week == 0:
                classes_per_week = 1  # Minimum 1 class per week
        else:
            # Full course: all classes
            classes_per_week = int(credit)
        
        # Calculate total working days (excluding holidays)
        working_days = []
        check_date = semester_start_date
        while check_date <= semester_end_date:
            if check_date not in holidays:
                working_days.append(check_date)
            check_date += timedelta(days=1)
        
        # Group working days into weeks
        week_groups = []
        current_week_days = []
        current_week_start = None
        
        for day in working_days:
            if day.weekday() == 0 or current_week_start is None:
                if current_week_days:
                    week_groups.append(current_week_days)
                current_week_days = [day]
                current_week_start = day
            elif (day - current_week_start).days >= 7:
                if current_week_days:
                    week_groups.append(current_week_days)
                current_week_days = [day]
                current_week_start = day
            else:
                current_week_days.append(day)
        
        if current_week_days:
            week_groups.append(current_week_days)
        
        total_weeks = len(week_groups)
        
        # Limit to maximum 14 weeks for generation
        MAX_WEEKS = 14
        if total_weeks > MAX_WEEKS:
            current_app.logger.info(f"Semester has {total_weeks} weeks, limiting to {MAX_WEEKS} weeks for class plan generation")
            total_weeks = MAX_WEEKS
            # Actually limit the week_groups array to 14 weeks
            week_groups = week_groups[:MAX_WEEKS]
        
        # Prepare course content topics (parse semicolons) - Only from selected topics
        all_topics = []
        for content_item in course_contents:
            if isinstance(content_item, dict):
                # Only process if selected (should already be filtered, but double-check)
                if content_item.get('selected', True):
                    content_text = content_item.get('content', '')
                    if content_text:
                        topics = [t.strip() for t in content_text.split(';') if t.strip()]
                        all_topics.extend(topics)
        
        # Generate lesson plan using rule-based logic (always within 14 weeks)
        # For split courses, both parts start from Week 1 and run simultaneously
        lesson_plan = _generate_rule_based_plan(
            session, credit, course_contents, week_groups, 
            holidays, semester_start_date, semester_end_date, 
            classes_per_week, total_weeks, all_topics, clos, part=part
        )
        
        part_text = f" (Part {part})" if part else ""
        return jsonify({
            'success': True,
            'lesson_plan': lesson_plan,
            'message': f'Generated {len(lesson_plan)} classes{part_text} within {total_weeks} weeks (maximum 14 weeks) based on {credit} credits, Course Content, and Academic Calendar (Semester: {semester_start_date.strftime("%d-%b-%Y")} to {semester_end_date.strftime("%d-%b-%Y")})'
        })
    except Exception as e:
        current_app.logger.error(f"Error generating AI plan: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error generating plan: {str(e)}'
        }), 500


def _generate_rule_based_plan(session, credit, course_contents, week_groups, holidays, 
                               semester_start_date, semester_end_date, classes_per_week, 
                               total_weeks, all_topics, clos, part=None):
    """Fallback rule-based plan generation when AI is not available - Always within 14 weeks maximum
    For split courses, both Part A and Part B start from Week 1 and run simultaneously"""
    from datetime import timedelta
    
    # Limit to maximum 14 weeks
    MAX_WEEKS = 14
    limited_weeks = min(total_weeks, MAX_WEEKS)
    limited_week_groups = week_groups[:limited_weeks]
    
    lesson_plan = []
    content_index = 0
    total_classes = limited_weeks * classes_per_week
    
    # For split courses, both parts start from Week 1
    # The classes_per_week is already divided (half for each part)
    
    # Distribute course contents across all classes
    if all_topics and total_classes > 0:
        topics_per_class = max(1, len(all_topics) // total_classes)
    else:
        topics_per_class = 1
    
    # Track which weeks have assessments (for 3-4 credit courses)
    # For split courses, divide assessments equally between Part A and Part B
    assessment_weeks = set()
    if credit in [3, 4]:
        # Distribute 4 assessments across 14 weeks
        if limited_weeks >= 14:
            all_assessment_week_indices = [3, 6, 9, 12]  # Weeks 4, 7, 10, 13
        elif limited_weeks >= 10:
            all_assessment_week_indices = [2, 4, 6, 8]
        else:
            all_assessment_week_indices = [1, 2, 3, 4]
        
        # For split courses, divide assessments equally
        if part in ['A', 'B']:
            # Split the assessments: Part A gets first half, Part B gets second half
            mid_point = len(all_assessment_week_indices) // 2
            if part == 'A':
                assessment_week_indices = all_assessment_week_indices[:mid_point]  # First half
            else:  # part == 'B'
                assessment_week_indices = all_assessment_week_indices[mid_point:]  # Second half
        else:
            # Full course: use all assessments
            assessment_week_indices = all_assessment_week_indices
        
        for week_idx in assessment_week_indices:
            if week_idx < limited_weeks:
                assessment_weeks.add(week_idx)
    
    # Generate plan for each week (maximum 14 weeks)
    week_num = 1
    for week_idx, week_days in enumerate(limited_week_groups):
        # Safety check: never exceed 14 weeks
        if week_num > MAX_WEEKS:
            break
            
        if not week_days:
            continue
        
        week_start = week_days[0]
        week_end = week_days[-1]
        week_classes_count = min(classes_per_week, len(week_days))
        
        # Check if this week has an assessment (for 3-4 credit courses)
        has_assessment = week_idx in assessment_weeks
        
        # If this week has an assessment, reduce regular classes by 1 to make room
        regular_classes = week_classes_count - 1 if has_assessment else week_classes_count
        
        # Generate regular classes for this week
        for class_num in range(regular_classes):
            if regular_classes == 1:
                class_date = week_days[0]
            else:
                day_index = int((class_num / regular_classes) * len(week_days))
                class_date = week_days[min(day_index, len(week_days) - 1)]
            
            # Get topics for this class
            week_contents = []
            week_clos = set()
            
            if all_topics:
                for i in range(topics_per_class):
                    if content_index < len(all_topics):
                        week_contents.append(all_topics[content_index])
                        content_index += 1
            
            # Get CLO from course contents if available
            if course_contents and content_index < len(course_contents):
                content_item = course_contents[min(content_index, len(course_contents) - 1)]
                if isinstance(content_item, dict):
                    clo_value = content_item.get('clo', '')
                    if clo_value:
                        if isinstance(clo_value, str):
                            for clo in clo_value.split(','):
                                clo = clo.strip()
                                if clo:
                                    week_clos.add(clo)
                        elif isinstance(clo_value, (int, float)):
                            week_clos.add(str(int(clo_value)))
            
            date_str = f"{week_start.strftime('%d-%b-%Y')} to {week_end.strftime('%d-%b-%Y')}"
            
            if week_contents:
                topic = ', '.join(week_contents[:2])
                if len(week_contents) > 2:
                    topic += f' and {len(week_contents) - 2} more'
            else:
                topic = f'Week {week_num} - Class {class_num + 1} Content'
            
            # Specific Outcome and Teaching & Assessment should be empty
            outcome = ''
            teaching_assessment = ''
            
            if week_clos:
                clo_alignment = ', '.join(sorted(week_clos, key=lambda x: int(x) if x.isdigit() else 999))
            elif clos:
                clo_alignment = ', '.join([str(i+1) for i in range(min(2, len(clos)))])
            else:
                clo_alignment = '1'
            
            lesson_plan.append({
                'week': f'Week {week_num}',
                'date': date_str,
                'topic': topic,
                'outcome': outcome,
                'teaching_assessment': teaching_assessment,
                'clo_alignment': clo_alignment
            })
        
        # Add assessment for this week if scheduled
        if has_assessment:
            # For split courses, renumber assessments starting from 1 for each part
            # Part A: Assessment 1, 2; Part B: Assessment 1, 2
            # For full course: Assessment 1, 2, 3, 4
            sorted_assessment_weeks = sorted(list(assessment_weeks))
            assessment_num = sorted_assessment_weeks.index(week_idx) + 1
            if part in ['A', 'B']:
                # For split courses, add part label
                topic_text = f'Assessment {assessment_num} (Part {part})'
            else:
                topic_text = f'Assessment {assessment_num}'
            
            date_str = f"{week_start.strftime('%d-%b-%Y')} to {week_end.strftime('%d-%b-%Y')}"
            
            # Specific Outcome and Teaching & Assessment should be empty
            lesson_plan.append({
                'week': f'Week {week_num}',
                'date': date_str,
                'topic': topic_text,
                'outcome': '',
                'teaching_assessment': '',
                'clo_alignment': '1, 2, 3, 4'
            })
        
        week_num += 1
        # Ensure we never exceed 14 weeks
        if week_num > MAX_WEEKS:
            break
    
    # Final safety check: ensure we never return more than 14 weeks worth of classes
    # Calculate max entries: 14 weeks * classes_per_week + assessments (max 4)
    max_entries = (MAX_WEEKS * classes_per_week) + (4 if credit in [3, 4] else 0)
    if len(lesson_plan) > max_entries:
        current_app.logger.warning(f"Generated {len(lesson_plan)} entries, limiting to {max_entries} (14 weeks max)")
        lesson_plan = lesson_plan[:max_entries]
    
    # Final safety: ensure we never exceed max_entries (14 weeks worth of classes)
    if len(lesson_plan) > max_entries:
        current_app.logger.warning(f"Generated {len(lesson_plan)} entries, limiting to {max_entries} (14 weeks max)")
        lesson_plan = lesson_plan[:max_entries]
    
    return lesson_plan

@class_management_bp.route('/course_file/<int:session_id>/outline/edit')
@login_required
def edit_course_outline(session_id):
    """Edit course outline page"""
    session = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    
    # Check if teacher is authorized (either main teacher or part of split group)
    is_authorized = False
    if teacher:
        if teacher.id == session.teacher_id:
            is_authorized = True
        elif session.split_group_id:
            # Check if teacher is part of the split group
            related_sessions = Session.query.filter_by(split_group_id=session.split_group_id).all()
            for related_session in related_sessions:
                if related_session.teacher_id == teacher.id:
                    is_authorized = True
                    break
    
    if not is_authorized:
        flash('You are not authorized to edit this course outline.', 'danger')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    # Get all teachers for this course (if split group exists)
    course_teachers = [session.teacher]
    if session.split_group_id:
        related_sessions = Session.query.filter_by(split_group_id=session.split_group_id).all()
        for related_session in related_sessions:
            if related_session.teacher and related_session.teacher not in course_teachers:
                course_teachers.append(related_session.teacher)
    
    course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
    if not course_outline:
        course_outline = CourseOutline(session_id=session_id, teacher_id=teacher.id)
        db.session.add(course_outline)
        db.session.commit()
    
    # Get course data from curriculum using improved matching
    # IMPORTANT: Same course can exist in multiple curricula, so we need to find the right one
    course_data = None
    try:
        if Course:
            import re
            current_app.logger.info(f"Searching for course - session.course_code: '{session.course_code}', session.course_name: '{session.course_name}', session.year: '{session.year}', session.term: '{session.term}', session.academic_session: '{session.academic_session}'")
            
            def extract_core_code(code_str):
                """Extract the core code pattern (e.g., 'Law4103' without space) from various formats"""
                if not code_str:
                    return None
                # Try to find pattern like "Law 4103", "Law4103", "CSE 1101", etc.
                match = re.search(r'([A-Za-z]+)\s*(\d{4})', code_str)
                if match:
                    return f"{match.group(1)}{match.group(2)}"  # No space: "Law4103"
                return None
            
            def has_course_info(course):
                """Check if a course has Course Information (rationale, CLO, or content)"""
                return bool(course.rationale or course.clo or course.content_section_a or course.content_section_b)
            
            def normalize_year_term(value):
                """Normalize year/term values for comparison"""
                if not value:
                    return ''
                value = str(value).strip().lower()
                # Map numeric to text
                year_map = {'1': 'first', '2': 'second', '3': 'third', '4': 'fourth', '5': 'fifth'}
                term_map = {'1': 'first', '2': 'second'}
                if value in year_map:
                    return year_map[value]
                if value in term_map:
                    return term_map[value]
                # Remove "year" or "term" suffix
                for suffix in [' year', ' term']:
                    if value.endswith(suffix):
                        value = value[:-len(suffix)]
                return value
            
            # Extract core code pattern from session course code
            session_core_code = extract_core_code(session.course_code)
            if session_core_code:
                current_app.logger.info(f"Extracted core code: '{session_core_code}' from '{session.course_code}'")
            
            # Normalize session year/term for matching
            session_year_norm = normalize_year_term(session.year)
            session_term_norm = normalize_year_term(session.term)
            
            def find_matching_courses(code_filter_func):
                """Find all courses matching the code filter, then prioritize by curriculum match and course info"""
                all_courses = Course.query.all()
                matching_courses = [c for c in all_courses if c.course_code and code_filter_func(c.course_code)]
                
                if not matching_courses:
                    return None
                
                current_app.logger.info(f"Found {len(matching_courses)} courses matching code filter")
                
                # Score courses: higher score = better match
                scored_courses = []
                for course in matching_courses:
                    score = 0
                    reasons = []
                    
                    # Priority 1: Has Course Information (most important)
                    if has_course_info(course):
                        score += 100
                        reasons.append("has_course_info")
                    
                    # Priority 2: Year/Term match from course's derived or stored year/term
                    course_year_norm = normalize_year_term(course.year or course.derived_year)
                    course_term_norm = normalize_year_term(course.term or course.derived_term)
                    
                    if session_year_norm and course_year_norm and session_year_norm == course_year_norm:
                        score += 20
                        reasons.append(f"year_match({course_year_norm})")
                    if session_term_norm and course_term_norm and session_term_norm == course_term_norm:
                        score += 10
                        reasons.append(f"term_match({course_term_norm})")
                    
                    # Priority 3: Academic session match via CurriculumYearTerm
                    if CurriculumYearTerm and session.academic_session and course.curriculum_id:
                        try:
                            config = CurriculumYearTerm.query.filter_by(
                                curriculum_id=course.curriculum_id,
                                academic_session=session.academic_session
                            ).first()
                            if config:
                                score += 50
                                reasons.append(f"academic_session_match({session.academic_session})")
                        except Exception as e:
                            current_app.logger.warning(f"Error checking CurriculumYearTerm: {e}")
                    
                    scored_courses.append((course, score, reasons))
                    current_app.logger.info(f"  Course: {course.course_code} (curriculum_id={course.curriculum_id}), score={score}, reasons={reasons}")
                
                # Sort by score (highest first)
                scored_courses.sort(key=lambda x: x[1], reverse=True)
                
                if scored_courses:
                    best_course, best_score, best_reasons = scored_courses[0]
                    current_app.logger.info(f"Selected best match: {best_course.course_code} with score {best_score} ({best_reasons})")
                    return best_course
                
                return None
            
            # Try exact match by course code
            if session.course_code:
                course_data = find_matching_courses(
                    lambda code: code == session.course_code
                )
                if course_data:
                    current_app.logger.info(f"Found by exact course_code match: {course_data.course_code}")
            
            # If not found, try case-insensitive match by course code
            if not course_data and session.course_code:
                course_data = find_matching_courses(
                    lambda code: code.lower() == session.course_code.lower()
                )
                if course_data:
                    current_app.logger.info(f"Found by case-insensitive course_code match: {course_data.course_code}")
            
            # If not found, try whitespace-normalized match
            if not course_data and session.course_code:
                session_code_normalized = ' '.join(session.course_code.strip().split()).lower()
                course_data = find_matching_courses(
                    lambda code: ' '.join(code.strip().split()).lower() == session_code_normalized
                )
                if course_data:
                    current_app.logger.info(f"Found by whitespace-normalized match: {course_data.course_code}")
            
            # If not found, try with extracted course code pattern (with space)
            if not course_data and session_core_code:
                extracted_with_space = re.sub(r'([A-Za-z]+)(\d{4})', r'\1 \2', session_core_code).lower()
                course_data = find_matching_courses(
                    lambda code: code.lower() == extracted_with_space
                )
                if course_data:
                    current_app.logger.info(f"Found by extracted code with space: {course_data.course_code}")
            
            # If not found, try with extracted course code pattern (without space)
            if not course_data and session_core_code:
                course_data = find_matching_courses(
                    lambda code: code.lower() == session_core_code.lower()
                )
                if course_data:
                    current_app.logger.info(f"Found by extracted code without space: {course_data.course_code}")
            
            # If not found, try normalized core code matching
            if not course_data and session_core_code:
                session_core_lower = session_core_code.lower()
                course_data = find_matching_courses(
                    lambda code: extract_core_code(code) and extract_core_code(code).lower() == session_core_lower
                )
                if course_data:
                    current_app.logger.info(f"Found by normalized core code match: {course_data.course_code}")
            
            # If not found, try partial match - check if curriculum course code is contained in session code
            if not course_data and session.course_code:
                session_code_lower = session.course_code.lower()
                session_code_no_space = session_code_lower.replace(' ', '')
                course_data = find_matching_courses(
                    lambda code: code.lower() in session_code_lower or code.lower().replace(' ', '') in session_code_no_space
                )
                if course_data:
                    current_app.logger.info(f"Found by partial code match: {course_data.course_code}")
            
            # If not found, try exact match by course name
            if not course_data and session.course_name:
                course_data = find_matching_courses(
                    lambda code: True  # Match all, but filter by name below
                )
                # Re-filter by name since find_matching_courses filters by code
                if not course_data:
                    all_courses = Course.query.filter_by(course_name=session.course_name).all()
                    if all_courses:
                        # Pick the one with course info
                        for c in all_courses:
                            if has_course_info(c):
                                course_data = c
                                break
                        if not course_data:
                            course_data = all_courses[0]
                        if course_data:
                            current_app.logger.info(f"Found by exact course_name match: {course_data.course_name}")
            
            # If not found, try case-insensitive partial match by course name
            if not course_data and session.course_name:
                all_courses = Course.query.filter(func.lower(Course.course_name).like(f'%{session.course_name.lower()}%')).all()
                if all_courses:
                    # Pick the one with course info
                    for c in all_courses:
                        if has_course_info(c):
                            course_data = c
                            break
                    if not course_data:
                        course_data = all_courses[0]
                    if course_data:
                        current_app.logger.info(f"Found by partial course_name match: {course_data.course_name}")
            
            # If still not found, try reverse match (session name contains course name)
            if not course_data and session.course_name:
                all_courses = Course.query.all()
                session_name_lower = session.course_name.lower()
                matching = []
                for course in all_courses:
                    if course.course_name:
                        course_name_lower = course.course_name.lower()
                        if session_name_lower in course_name_lower or course_name_lower in session_name_lower:
                            matching.append(course)
                if matching:
                    # Pick the one with course info
                    for c in matching:
                        if has_course_info(c):
                            course_data = c
                            break
                    if not course_data:
                        course_data = matching[0]
                    if course_data:
                        current_app.logger.info(f"Found by reverse course_name match: {course_data.course_name}")
            
            # Log the result for debugging
            if course_data:
                current_app.logger.info(f"✓ Found course_data: {course_data.course_code} - {course_data.course_name}, core_optional: {course_data.core_optional}, course_type: {course_data.course_type}, category: {course_data.category}")
                # Log additional fields for debugging
                current_app.logger.info(f"  - rationale: {course_data.rationale[:100] if course_data.rationale else 'EMPTY'}")
                current_app.logger.info(f"  - clo: {course_data.clo[:100] if course_data.clo else 'EMPTY'}")
                current_app.logger.info(f"  - content_section_a: {course_data.content_section_a[:100] if course_data.content_section_a else 'EMPTY'}")
                current_app.logger.info(f"  - content_section_b: {course_data.content_section_b[:100] if course_data.content_section_b else 'EMPTY'}")
            else:
                current_app.logger.warning(f"✗ Course data NOT found for session - course_code: '{session.course_code}', course_name: '{session.course_name}'")
                # List all courses for debugging with their core codes
                all_courses = Course.query.all()
                current_app.logger.info(f"Available courses in database ({len(all_courses)} total):")
                for c in all_courses[:15]:  # Show first 15
                    c_core = extract_core_code(c.course_code)
                    current_app.logger.info(f"  - Code: '{c.course_code}', Core: '{c_core}', Name: '{c.course_name}'")
    except Exception as e:
        current_app.logger.error(f"Error fetching course data: {e}", exc_info=True)
        current_app.logger.warning(f"Session course_code: {session.course_code}, course_name: {session.course_name}")
        course_data = None
    
    # Get CLO data from course if not already saved in outline
    clo_data_from_course = []
    if course_data and hasattr(course_data, 'get_clos_list'):
        try:
            course_clos = course_data.get_clos_list()
            if course_clos:
                # Convert course CLO format to outline CLO format
                for idx, clo in enumerate(course_clos, 1):
                    # Parse PLO from curriculum
                    plos_list = []
                    if clo.get('plo'):
                        plo_value = clo.get('plo', '')
                        if isinstance(plo_value, str) and plo_value.strip():
                            plos_list = [p.strip() for p in plo_value.split(',') if p.strip()]
                        elif isinstance(plo_value, list):
                            plos_list = [str(p).strip() for p in plo_value if p]
                    
                    clo_data_from_course.append({
                        'number': idx,
                        'description': clo.get('text', ''),
                        'plos': plos_list
                    })
        except Exception as e:
            current_app.logger.warning(f"Error parsing course CLOs: {e}")
    
    # Parse course contents from curriculum for import
    course_contents_a = []
    course_contents_b = []
    if course_data:
        try:
            if course_data.content_section_a:
                content_a = course_data.content_section_a
                try:
                    content_a_data = json.loads(content_a) if isinstance(content_a, str) else content_a
                    if isinstance(content_a_data, list):
                        course_contents_a = content_a_data
                except:
                    pass
            if course_data.content_section_b:
                content_b = course_data.content_section_b
                try:
                    content_b_data = json.loads(content_b) if isinstance(content_b, str) else content_b
                    if isinstance(content_b_data, list):
                        course_contents_b = content_b_data
                except:
                    pass
        except Exception as e:
            current_app.logger.warning(f"Error parsing course contents: {e}")
    
    # Parse existing course content summary
    existing_content_summary = None
    if course_outline.course_content_summary:
        try:
            existing_content_summary = json.loads(course_outline.course_content_summary)
        except:
            existing_content_summary = course_outline.course_content_summary
    
    # Parse JSON fields
    outline_data = {
        'course_objectives': json.loads(course_outline.course_objectives) if course_outline.course_objectives else [],
        'course_content_summary': existing_content_summary or '',
        'clo_plo_mapping': course_outline.clo_plo_mapping or '',
        'clo_data': json.loads(course_outline.clo_data) if hasattr(course_outline, 'clo_data') and course_outline.clo_data else (clo_data_from_course if clo_data_from_course else []),
        'plo_mapping': json.loads(course_outline.plo_mapping) if hasattr(course_outline, 'plo_mapping') and course_outline.plo_mapping else {},
        'lesson_plan': json.loads(course_outline.lesson_plan) if course_outline.lesson_plan else [],
        'assessment_strategy': json.loads(course_outline.assessment_strategy) if course_outline.assessment_strategy else {},
        'assessment_techniques': json.loads(course_outline.assessment_techniques) if course_outline.assessment_techniques else [],
        'rubrics': json.loads(course_outline.rubrics) if course_outline.rubrics else [],
        'grading_policy': json.loads(course_outline.grading_policy) if course_outline.grading_policy else [],
        'cie_breakdown': json.loads(course_outline.cie_breakdown) if hasattr(course_outline, 'cie_breakdown') and course_outline.cie_breakdown else [],
        'smee_breakdown': json.loads(course_outline.smee_breakdown) if hasattr(course_outline, 'smee_breakdown') and course_outline.smee_breakdown else [],
        'textbooks': json.loads(course_outline.textbooks) if course_outline.textbooks else [],
        'reference_books': json.loads(course_outline.reference_books) if course_outline.reference_books else [],
        'other_resources': json.loads(course_outline.other_resources) if course_outline.other_resources else [],
        'course_file_components': json.loads(course_outline.course_file_components) if course_outline.course_file_components else [],
        'other_issues': json.loads(course_outline.other_issues) if course_outline.other_issues else {},
    }
    
    return render_template('class_management/edit_course_outline.html',
                         session=session,
                         course_outline=course_outline,
                         course_data=course_data,
                         outline_data=outline_data,
                         course_contents_a=course_contents_a,
                         course_contents_b=course_contents_b,
                         course_teachers=course_teachers)

def _generate_course_outline_docx(session_id):
    """Generate course outline as DOCX document"""
    session = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    
    if not teacher or teacher.id != session.teacher_id:
        flash('You are not authorized to download this course outline.', 'danger')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
    if not course_outline:
        flash('Course outline not found. Please create it first.', 'warning')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    # Get course data
    course_data = find_course_from_curriculum(session.course_code, session.course_name)
    
    # Parse JSON fields
    course_objectives = json.loads(course_outline.course_objectives) if course_outline.course_objectives else []
    lesson_plan = json.loads(course_outline.lesson_plan) if course_outline.lesson_plan else []
    textbooks = json.loads(course_outline.textbooks) if course_outline.textbooks else []
    reference_books = json.loads(course_outline.reference_books) if course_outline.reference_books else []
    other_resources = json.loads(course_outline.other_resources) if course_outline.other_resources else []
    
    # Lazy import docx to prevent startup hang
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Create DOCX document
    doc = Document()
    
    # Cover Page
    cover_para = doc.add_paragraph()
    cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_para.add_run('Course Outline')
    run.font.size = Pt(18)
    run.font.bold = True
    
    if session.course_code:
        code_para = doc.add_paragraph()
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = code_para.add_run(session.course_code + ':')
        run.font.size = Pt(16)
        run.font.bold = True
    
    if session.course_name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(session.course_name.upper())
        run.font.size = Pt(20)
        run.font.bold = True
    
    doc.add_page_break()
    
    # Part A: Introduction
    doc.add_heading('PART A: INTRODUCTION', level=1)
    
    # Course Identification Table
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = 'Course No:'
    info_table.cell(0, 1).text = session.course_code or '—'
    # Add more rows...
    
    # Course Objectives
    if course_objectives:
        doc.add_heading('Course Objectives', level=2)
        for obj in course_objectives:
            doc.add_paragraph(obj, style='List Bullet')
    
    # Course Summary
    if course_outline.course_summary:
        doc.add_heading('Course Summary', level=2)
        doc.add_paragraph(course_outline.course_summary)
    
    # Lesson Plan
    if lesson_plan:
        doc.add_page_break()
        doc.add_heading('Class Schedule/Lesson Plan/Weekly plan', level=1)
        lesson_table = doc.add_table(rows=1, cols=7)
        lesson_table.style = 'Table Grid'
        headers = ['Week', 'Date', 'Topic', 'Specific Outcome', 'Suggested Activities', 'Teaching and Assessment', 'Alignment with CLO']
        for i, header in enumerate(headers):
            lesson_table.cell(0, i).text = header
            lesson_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
        
        for lesson in lesson_plan:
            row = lesson_table.add_row().cells
            row[0].text = lesson.get('week', '')
            row[1].text = lesson.get('date', '')
            row[2].text = lesson.get('topic', '')
            row[3].text = lesson.get('outcome', '')
            row[4].text = lesson.get('activities', '')
            row[5].text = lesson.get('teaching_assessment', '')
            row[6].text = lesson.get('clo_alignment', '')
    
    # Learning Resources
    if textbooks or reference_books or other_resources:
        doc.add_page_break()
        doc.add_heading('PART D: LEARNING RESOURCES', level=1)
        
        if textbooks:
            doc.add_heading('Textbooks', level=2)
            for book in textbooks:
                doc.add_paragraph(book, style='List Bullet')
        
        if reference_books:
            doc.add_heading('Reference Books', level=2)
            for book in reference_books:
                doc.add_paragraph(book, style='List Bullet')
        
        if other_resources:
            doc.add_heading('Other Resources', level=2)
            for resource in other_resources:
                doc.add_paragraph(resource, style='List Bullet')
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_data = buffer.getvalue()
    buffer.close()
    
    filename = f"course_outline_{session.course_code or 'course'}.docx"
    return Response(
        docx_data,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(docx_data)),
        },
    )

def _generate_course_outline_pdf(session_id, skip_auth_check=False):
    """Generate comprehensive course outline as PDF document with cover page and page numbers using WeasyPrint
    
    Args:
        session_id: The session ID for the course outline
        skip_auth_check: If True, skip authorization checks (for student downloads)
    """
    try:
        from weasyprint import HTML, CSS
    except ImportError:
        flash('WeasyPrint is not installed. Please install it to generate PDFs.', 'error')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    session = Session.query.get_or_404(session_id)
    
    # Skip authorization check if requested (for student downloads)
    if not skip_auth_check:
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        
        # Check if user is authorized (teacher or part of split group)
        is_authorized = False
        if teacher and teacher.id == session.teacher_id:
            is_authorized = True
        elif session.split_group_id:
            related_sessions = Session.query.filter_by(split_group_id=session.split_group_id).all()
            for related_session in related_sessions:
                if related_session.teacher and related_session.teacher.id == teacher.id:
                    is_authorized = True
                    break
        
        if not is_authorized:
            flash('You are not authorized to download this course outline.', 'danger')
            return redirect(url_for('class_management.course_file', session_id=session_id))
    
    course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
    if not course_outline:
        flash('Course outline not found. Please create it first.', 'warning')
        return redirect(url_for('class_management.course_file', session_id=session_id))
    
    # Get course data from curriculum if available
    course_data = find_course_from_curriculum(session.course_code, session.course_name)
    
    # Parse all JSON fields
    def safe_json_parse(data, default=None):
        if not data:
            return default if default is not None else []
        try:
            return json.loads(data) if isinstance(data, str) else data
        except:
            return default if default is not None else []
    
    course_objectives = safe_json_parse(course_outline.course_objectives, [])
    lesson_plan = safe_json_parse(course_outline.lesson_plan, [])
    clo_data_raw = safe_json_parse(course_outline.clo_data, [])
    # Ensure plos is always a list in each CLO entry
    clo_data = []
    for clo in clo_data_raw:
        clo_entry = dict(clo)
        plos = clo_entry.get('plos', [])
        if not isinstance(plos, list):
            if isinstance(plos, str):
                clo_entry['plos'] = [plos] if plos else []
            else:
                clo_entry['plos'] = []
        clo_data.append(clo_entry)
    course_content_summary = safe_json_parse(course_outline.course_content_summary, {})
    assessment_strategy = safe_json_parse(course_outline.assessment_strategy, {})
    assessment_techniques = safe_json_parse(course_outline.assessment_techniques, [])
    cie_breakdown = safe_json_parse(course_outline.cie_breakdown, []) if hasattr(course_outline, 'cie_breakdown') else []
    smee_breakdown = safe_json_parse(course_outline.smee_breakdown, []) if hasattr(course_outline, 'smee_breakdown') else []
    rubrics = safe_json_parse(course_outline.rubrics, [])
    # Group rubrics by type for easier template rendering
    rubrics_by_type = {}
    for rubric in rubrics:
        rubric_type = rubric.get('type', '') or ''
        if rubric_type not in rubrics_by_type:
            rubrics_by_type[rubric_type] = []
        rubrics_by_type[rubric_type].append(rubric)
    
    grading_policy = safe_json_parse(course_outline.grading_policy, [])
    evaluation_policy = safe_json_parse(course_outline.evaluation_policy, {})
    textbooks = safe_json_parse(course_outline.textbooks, [])
    reference_books = safe_json_parse(course_outline.reference_books, [])
    other_resources = safe_json_parse(course_outline.other_resources, [])
    course_file_components = safe_json_parse(course_outline.course_file_components, [])
    other_issues = safe_json_parse(course_outline.other_issues, {})
    
    # Get all teachers for this course (if split group exists)
    course_teachers = [session.teacher]
    course_teachers_pdf = [session.teacher]
    if session.split_group_id:
        related_sessions = Session.query.filter_by(split_group_id=session.split_group_id).all()
        for related_session in related_sessions:
            if related_session.teacher and related_session.teacher not in course_teachers:
                course_teachers.append(related_session.teacher)
            if related_session.teacher and related_session.teacher not in course_teachers_pdf:
                course_teachers_pdf.append(related_session.teacher)
    
    # Render HTML template
    html_content = render_template(
        'class_management/course_outline_pdf.html',
        session=session,
        course_outline=course_outline,
        course_data=course_data,
        course_objectives=course_objectives,
        lesson_plan=lesson_plan,
        clo_data=clo_data,
        course_content_summary=course_content_summary,
        assessment_strategy=assessment_strategy,
        assessment_techniques=assessment_techniques,
        cie_breakdown=cie_breakdown,
        smee_breakdown=smee_breakdown,
        rubrics=rubrics,
        rubrics_by_type=rubrics_by_type,
        grading_policy=grading_policy,
        evaluation_policy=evaluation_policy,
        textbooks=textbooks,
        reference_books=reference_books,
        other_resources=other_resources,
        course_file_components=course_file_components,
        other_issues=other_issues,
        course_teachers=course_teachers,
        course_teachers_pdf=course_teachers_pdf
    )
    
    # Generate PDF using WeasyPrint
    buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(buffer)
    buffer.seek(0)
    
    filename = f"course_outline_{session.course_code or 'course'}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return Response(
        buffer.getvalue(),
        mimetype='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(buffer.getvalue())),
        },
    )

@class_management_bp.route('/course_file/<int:session_id>/outline/download/docx')
@login_required
def download_course_outline_docx(session_id):
    """Download course outline as DOCX"""
    return _generate_course_outline_docx(session_id)

@class_management_bp.route('/course_file/<int:session_id>/outline/download/pdf')
@login_required
def download_course_outline_pdf(session_id):
    """Download course outline as PDF"""
    return _generate_course_outline_pdf(session_id)

@class_management_bp.route('/archive_session/<int:session_id>', methods=['POST'])
@login_required
def archive_session(session_id):
    """Archive a session"""
    session = Session.query.get_or_404(session_id)
    session.archived = True
    db.session.commit()
    flash('Session archived successfully!', 'success')
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/unarchive_session/<int:session_id>', methods=['POST'])
@login_required
def unarchive_session(session_id):
    """Unarchive a session"""
    session = Session.query.get_or_404(session_id)
    session.archived = False
    db.session.commit()
    flash('Session unarchived successfully!', 'success')
    return redirect(url_for('class_management.index'))

@class_management_bp.route('/edit_session/<int:session_id>', methods=['GET', 'POST'])
@login_required
def edit_session(session_id):
    """Edit a session"""
    session = Session.query.get_or_404(session_id)
    
    if request.method == 'POST':
        try:
            session.year = request.form.get('year')
            session.term = request.form.get('term')
            session.academic_session = request.form.get('academic_session')
            session.course_code = request.form.get('course_code')
            session.course_name = request.form.get('course_name')
            session.course_type = request.form.get('course_type', 'theory')
            session.category = request.form.get('category', 'ug')
            
            if not session.year or not session.term:
                flash('Year and term are required!', 'error')
                return redirect(url_for('class_management.edit_session', session_id=session_id))
            
            db.session.commit()
            flash('Session updated successfully!', 'success')
            return redirect(url_for('class_management.index'))
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error updating session {session_id}: {e}")
            flash(f'Error updating session: {str(e)}', 'error')
            return redirect(url_for('class_management.edit_session', session_id=session_id))
    
    return render_template('class_management/edit_session.html', session=session)

@class_management_bp.route('/archive')
@login_required
def archive():
    """View archived sessions"""
    # Get or create teacher for current user
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher:
        teacher = Teacher(name=current_user.full_name)
        db.session.add(teacher)
        db.session.commit()
    
    archived_sessions = Session.query.filter_by(teacher_id=teacher.id, archived=True).order_by(Session.created_at.desc()).all()
    
    # Build assignment map for template to access batch and academic_session from CourseSessionAssignment
    assignment_map = {}
    if CourseSessionAssignment and Course:
        try:
            for session in archived_sessions:
                # Try to find assignment by session_id first
                assignment = CourseSessionAssignment.query.filter_by(session_id=session.id).first()
                
                # If not found by session_id, try to find by course_code, teacher_id, year, term
                if not assignment and session.course_code and session.teacher_id and session.year and session.term:
                    try:
                        # Try to match by course_code, teacher_id, year, term
                        assignment = CourseSessionAssignment.query.filter_by(
                            teacher_id=session.teacher_id,
                            year=session.year,
                            term=session.term
                        ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                            Course.course_code == session.course_code
                        ).first()
                        
                        # If not found, try without section matching (for full course sessions)
                        if not assignment:
                            assignment = CourseSessionAssignment.query.filter_by(
                                teacher_id=session.teacher_id,
                                year=session.year,
                                term=session.term
                            ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                                Course.course_code == session.course_code
                            ).filter(
                                or_(
                                    CourseSessionAssignment.section.is_(None),
                                    CourseSessionAssignment.section == ''
                                )
                            ).first()
                    except Exception as query_error:
                        current_app.logger.warning(f'Error querying assignment for archived session {session.id}: {query_error}')
                
                if assignment:
                    # If assignment doesn't have batch/academic_session, try to get from curriculum year-term config
                    batch = assignment.batch
                    academic_session = assignment.academic_session
                    
                    if Curriculum and CurriculumYearTerm and (not batch or not academic_session):
                        try:
                            if assignment.curriculum_id:
                                curriculum = Curriculum.query.get(assignment.curriculum_id)
                                if curriculum:
                                    year_term_config = curriculum.get_year_term_config(assignment.year, assignment.term)
                                    if year_term_config:
                                        if not batch and year_term_config.batch and year_term_config.batch != 'None':
                                            batch = year_term_config.batch
                                        if not academic_session and year_term_config.academic_session:
                                            academic_session = year_term_config.academic_session
                        except Exception as config_error:
                            current_app.logger.warning(f'Error getting year-term config for assignment {assignment.id}: {config_error}')
                    
                    assignment_map[session.id] = {
                        'batch': batch or '',
                        'academic_session': academic_session or ''
                    }
        except Exception as e:
            current_app.logger.error(f'Error building assignment map for archived sessions: {str(e)}', exc_info=True)
    
    return render_template('class_management/archive.html', 
                         sessions=archived_sessions, 
                         assignment_map=assignment_map if assignment_map else {})

@class_management_bp.route('/delete_attendance/<int:session_id>/<string:date_str>', methods=['POST'])
@login_required
def delete_attendance_by_date(session_id, date_str):
    """Delete all attendance records for a specific date."""
    try:
        attendance_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        
        session = Session.query.get_or_404(session_id)
        
        if session.teacher.id != current_user.id:
            flash('You are not authorized to delete attendance for this session.', 'danger')
            return redirect(url_for('class_management.index'))
        
        # Count records before deletion
        records_to_delete = ClassAttendance.query.filter_by(
            session_id=session_id,
            date=attendance_date
        )
        
        if records_to_delete.count() == 0:
            flash(f'No attendance records found for {attendance_date.strftime("%b %d, %Y")}.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
            
        # Delete the records
        deleted_count = records_to_delete.delete()
        
        db.session.commit()
        
        flash(f'Successfully deleted {deleted_count} attendance records for {attendance_date.strftime("%b %d, %Y")}.', 'success')
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting attendance for session {session_id} on date {date_str}: {str(e)}")
        flash(f'Error deleting attendance: {str(e)}', 'danger')
        
    return redirect(url_for('class_management.view_attendance', session_id=session_id))

@class_management_bp.route('/download_attendance_excel/<int:session_id>')
@login_required
def download_attendance_excel(session_id):
    """Generate and download an Excel report of the attendance."""
    try:
        # Import error handler for detailed logging
        from error_handler import log_error
        
        current_app.logger.info(f"Starting Excel generation for session {session_id}")
        
        # Check if required modules are available
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
            current_app.logger.info("Required modules available for Excel")
        except ImportError as e:
            current_app.logger.error(f"Missing required module for Excel: {e}")
            flash(f'Missing required module for Excel: {e}', 'error')
            return redirect(url_for('class_management.index'))
            
        session = Session.query.get_or_404(session_id)
        students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()
        attendance_summary = _build_attendance_summary(session)
        combined_assessment_map = _collect_combined_assessment_marks(session)
        attendance_summary = _build_attendance_summary(session)
        combined_assessment_map = _collect_combined_assessment_marks(session)
        attendance_summary = _build_attendance_summary(session)
        all_attendance_records = ClassAttendance.query.filter_by(session_id=session_id).order_by(ClassAttendance.date, ClassAttendance.id).all()

        if not all_attendance_records:
            flash('No attendance data to download.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))

        # This logic is similar to view_attendance, consider refactoring in a real app
        attendance_by_date = defaultdict(list)
        for record in all_attendance_records:
            attendance_by_date[record.date].append(record)
        
        daily_class_counts = {}
        for date, records in attendance_by_date.items():
            student_counts_on_date = defaultdict(int)
            for record in records:
                student_counts_on_date[record.student_id] += 1
            if student_counts_on_date:
                daily_class_counts[date] = max(student_counts_on_date.values())
                
        headers = []
        sorted_dates = sorted(daily_class_counts.keys())
        for dt in sorted_dates:
            count = daily_class_counts.get(dt, 0)
            if count == 1:
                headers.append(dt.strftime('%b %d, %Y'))
            else:
                for i in range(1, count + 1):
                    headers.append(f"{dt.strftime('%b %d, %Y')} ({i})")

        # Prepare structured data for Excel
        local_classes_held = sum(daily_class_counts.values())
        data_rows = []

        agg_student_map = attendance_summary.get('per_student', {})
        agg_total_classes = attendance_summary.get('total_classes', local_classes_held)

        for index, student in enumerate(students, start=1):
            student_attendance_records = [r for r in all_attendance_records if r.student_id == student.id]
            agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})
            present_count = agg_stats['present']
            percentage = agg_stats['percentage']
            marks = agg_stats['marks']
            
            # Initialize attendance row with placeholders
            attendance_statuses = ['-'] * len(headers)
            
            col_idx = 0
            for dt in sorted_dates:
                records_on_date = [r for r in student_attendance_records if r.date == dt]
                num_classes_on_date = daily_class_counts.get(dt, 0)
                
                for class_num in range(num_classes_on_date):
                    if class_num < len(records_on_date):
                        attendance_statuses[col_idx] = 'P' if records_on_date[class_num].is_present else 'A'
                    col_idx += 1

            data_rows.append({
                'serial': index,
                'student_id': student.student_id,
                'name': student.name,
                'attendance': attendance_statuses,
                'total_classes': agg_total_classes,
                'present': present_count,
                'percentage': percentage,
                'marks': marks
            })

        # Create workbook with styled layout
        wb = Workbook()
        ws = wb.active
        ws.title = 'Attendance Report'

        total_columns = 3 + len(headers) + 4  # SL + ID + Name + attendance + summary columns
        last_column_letter = get_column_letter(total_columns)

        # Title row
        ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=total_columns)
        title_cell = ws.cell(row=1, column=2)
        title_cell.value = 'Attendance Sheet'
        title_cell.font = Font(size=16, bold=True)
        title_cell.alignment = Alignment(horizontal='center', vertical='center')

        # Subject information
        subject_text = 'Subject: '
        if session.course_code and session.course_name:
            subject_text += f"{session.course_code} {session.course_name}"
        elif session.course_name:
            subject_text += session.course_name
        elif session.course_code:
            subject_text += session.course_code
        else:
            subject_text += 'N/A'

        ws.merge_cells(start_row=2, start_column=2, end_row=2, end_column=total_columns)
        subject_cell = ws.cell(row=2, column=2)
        subject_cell.value = subject_text
        subject_cell.font = Font(bold=True)
        subject_cell.alignment = Alignment(horizontal='left', vertical='center')

        # Additional metadata row
        teacher_name = None
        if session.teacher and getattr(session.teacher, 'name', None):
            teacher_name = session.teacher.name
        else:
            teacher_name = getattr(current_user, 'full_name', None) or getattr(current_user, 'username', 'N/A')

        metadata_items = [
            f"Year: {session.year or 'N/A'}",
            f"Term: {session.term or 'N/A'}",
            f"Session: {session.academic_session or 'N/A'}",
            f"Course Teacher: {teacher_name}"
        ]

        metadata_start_col = 2
        for item in metadata_items:
            if metadata_start_col > total_columns:
                break
            cell = ws.cell(row=3, column=metadata_start_col)
            cell.value = item
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='left', vertical='center')
            metadata_start_col += 4

        # Header row for table
        header_row = 5
        headers_for_sheet = ['Sl', 'Student ID', 'Name'] + headers + ['Total Classes', 'Present', 'Percentage', 'Marks']

        header_fill = PatternFill(start_color='D8E4BC', end_color='D8E4BC', fill_type='solid')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        for col_index, header_label in enumerate(headers_for_sheet, start=1):
            cell = ws.cell(row=header_row, column=col_index)
            cell.value = header_label
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=False, text_rotation=90, shrink_to_fit=False)
            cell.border = thin_border

        # Data rows
        data_start_row = header_row + 1
        attendance_start_col = 4
        percentage_col = attendance_start_col + len(headers) + 2  # Serial(1) + ID(2) + Name(3) + attendance + Total + Present => +2, Percentage column index
        marks_col = percentage_col + 1

        for row_offset, row_data in enumerate(data_rows):
            row_number = data_start_row + row_offset
            ws.cell(row=row_number, column=1, value=row_data['serial']).alignment = Alignment(horizontal='center')
            ws.cell(row=row_number, column=2, value=row_data['student_id']).alignment = Alignment(horizontal='center')
            ws.cell(row=row_number, column=3, value=row_data['name']).alignment = Alignment(horizontal='left')

            for idx, status in enumerate(row_data['attendance']):
                cell = ws.cell(row=row_number, column=attendance_start_col + idx, value=status)
                cell.alignment = Alignment(horizontal='center', vertical='center')

            total_col = attendance_start_col + len(headers)
            ws.cell(row=row_number, column=total_col, value=row_data['total_classes']).alignment = Alignment(horizontal='center')
            ws.cell(row=row_number, column=total_col + 1, value=row_data['present']).alignment = Alignment(horizontal='center')

            percentage_cell = ws.cell(row=row_number, column=percentage_col, value=(row_data['percentage'] / 100 if agg_total_classes else 0))
            percentage_cell.number_format = '0.00%'
            percentage_cell.alignment = Alignment(horizontal='center')

            ws.cell(row=row_number, column=marks_col, value=row_data['marks']).alignment = Alignment(horizontal='center')

        # Apply borders to data cells
        max_row = max(header_row, data_start_row + len(data_rows) - 1)
        for row in ws.iter_rows(min_row=header_row, max_row=max_row, min_col=1, max_col=total_columns):
            for cell in row:
                cell.border = thin_border

        # Auto fit all columns based on content - exact fit without extra space
        for col_idx in range(1, total_columns + 1):
            column_letter = get_column_letter(col_idx)
            max_length = 0
            
            # Check all cells in this column to find maximum content length
            for row in ws.iter_rows(min_col=col_idx, max_col=col_idx, values_only=False):
                for cell in row:
                    if cell.value is not None:
                        # Calculate exact length of cell content
                        cell_value = str(cell.value)
                        # Count actual character length
                        actual_length = len(cell_value)
                        max_length = max(max_length, actual_length)
            
            # Set column width to exact content width with minimal padding
            # openpyxl column width is in character units (approximate)
            if max_length > 0:
                # Very minimal padding: 0.5-1 character for readability
                # For narrow columns (like attendance P/A), use less padding
                if max_length <= 2:  # Single character columns (P, A, -)
                    adjusted_width = max_length + 0.5
                else:
                    adjusted_width = max_length + 1.0
            else:
                # If column is completely empty, set minimal width
                adjusted_width = 2
            
            ws.column_dimensions[column_letter].width = adjusted_width

        ws.freeze_panes = ws.cell(row=data_start_row, column=attendance_start_col)

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        current_app.logger.info(f"Excel generated successfully for session {session_id}")
        
        # Use Response instead of send_file for better cPanel compatibility
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={
                'Content-Disposition': f'attachment; filename="attendance_report_{session.course_name}_{date.today()}.xlsx"',
                'Content-Length': str(len(output.getvalue()))
            }
        )
        
    except Exception as e:
        # Log detailed error information
        log_error(e, {
            'session_id': session_id,
            'function': 'download_attendance_excel',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        
        current_app.logger.error(f"Error generating Excel for session {session_id}: {e}")
        flash(f'Error generating Excel: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/download_pdf_report/<int:session_id>')
@login_required
def download_pdf_report(session_id):
    """Generate and download a PDF summary report with headers and page numbers."""
    try:
        # Import error handler for detailed logging
        from error_handler import log_error
        from flask import Response
        
        current_app.logger.info(f"Starting PDF generation for session {session_id}")
        
        # Check if required modules are available
        try:
            import reportlab
            import pandas
            current_app.logger.info("Required modules available")
        except ImportError as e:
            current_app.logger.error(f"Missing required module: {e}")
            flash(f'Missing required module: {e}', 'error')
            return redirect(url_for('class_management.index'))
        session = Session.query.get_or_404(session_id)
        students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()
        combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        attendance_summary = _build_attendance_summary(session)
        combined_assessment_map = _collect_combined_assessment_marks(session)

        buffer = io.BytesIO()
        
        # Define margins
        top_margin = 2.5 * inch
        bottom_margin = 0.7 * inch
        
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=0.5*inch, leftMargin=0.5*inch,
                                topMargin=top_margin, bottomMargin=bottom_margin)
        
        elements = []
        
        # --- Data Calculation ---
        student_data_for_pdf = []
        per_student_attendance = attendance_summary.get('per_student', {})
        for student in students:
            attendance_stats = per_student_attendance.get(student.student_id, {'marks': 0})
            attendance_marks = attendance_stats.get('marks', 0)

            if session.course_type == 'sessional':
                # For sessional courses, keep report and viva separate
                sessional_report = int(round(student.sessional_report or 0))
                sessional_viva = int(round(student.sessional_viva or 0))
                student_data_for_pdf.append({
                    'id': student.student_id,
                    'attendance': attendance_marks,
                    'sessional_report': sessional_report,
                    'sessional_viva': sessional_viva
                })
            else:
                # For theory courses, use combined assessment
                assessment_marks_display = 0
                if session.course_type == 'theory':
                    if session.category == 'pg':
                        assessment_marks_display = combined_pg_total.get(student.student_id) or 0  # Already rounded
                    else:
                        # UG: Keep fraction in assessment page, but round for result generation
                        ug_total = combined_best3.get(student.student_id) or 0
                        assessment_marks_display = int(round(ug_total)) if ug_total else 0  # Round for result generation

                student_data_for_pdf.append({
                    'id': student.student_id,
                    'attendance': attendance_marks,
                    'assessment': assessment_marks_display
                })

        # --- Table Creation ---
        if session.course_type == 'sessional':
            # For sessional courses: separate columns for Report and Viva
            table_data = [['ID', 'Attendance (10)', 'Sessional Report (60)', 'Sessional Viva (30)']]
            for s_data in student_data_for_pdf:
                table_data.append([
                    s_data['id'], 
                    s_data['attendance'], 
                    s_data['sessional_report'],
                    s_data['sessional_viva']
                ])
            table = Table(table_data, colWidths=[1.5*inch, 1.5*inch, 2*inch, 2*inch], repeatRows=0)
        else:
            # For theory courses: single assessment column
            assessment_header_text = "Continuous Assessment (30)"
            if session.course_type == 'theory' and session.category == 'pg':
                assessment_header_text = "Continuous Assessment (40)"
            
            table_data = [['ID', 'Attendance (10)', assessment_header_text]]
            for s_data in student_data_for_pdf:
                table_data.append([s_data['id'], s_data['attendance'], s_data['assessment']])
            table = Table(table_data, colWidths=[2*inch, 2*inch, 2.5*inch], repeatRows=0)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F4F4F')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F2F2F2')),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
        ]))
        
        elements.append(table)
        
        # --- PDF Build with header/footer ---
        def _draw_header_footer(canvas_obj, doc_obj, include_header=False):
            canvas_obj.saveState()
            width, height = doc_obj.pagesize
            if include_header:
                canvas_obj.setFont('Helvetica-Bold', 18)
                canvas_obj.drawCentredString(width / 2.0, height - 1.0 * inch, "Khulna University")

                canvas_obj.setFont('Helvetica', 12)
                canvas_obj.drawCentredString(width / 2.0, height - 1.35 * inch, "Law Discipline")
                canvas_obj.drawCentredString(width / 2.0, height - 1.60 * inch, "Continuous Assessment and Attendance Marks")

                canvas_obj.setFont('Helvetica-Bold', 10)
                course_info = f"Course: {session.course_name} ({session.course_code or 'N/A'})  |  Type: {session.course_type.capitalize()}"
                year_term_info = f"Year: {session.year}, Term: {session.term}  |  Session: {session.academic_session}"
                canvas_obj.drawCentredString(width / 2.0, height - 1.95 * inch, course_info)
                canvas_obj.drawCentredString(width / 2.0, height - 2.15 * inch, year_term_info)

            canvas_obj.setFont('Helvetica', 9)
            page_text = f"Page {doc_obj.page}"
            canvas_obj.drawRightString(width - 0.5 * inch, 0.5 * inch, page_text)
            canvas_obj.restoreState()

        def first_page(canvas_obj, doc_obj):
            _draw_header_footer(canvas_obj, doc_obj, include_header=True)

        def later_pages(canvas_obj, doc_obj):
            _draw_header_footer(canvas_obj, doc_obj, include_header=False)

        doc.build(elements, onFirstPage=first_page, onLaterPages=later_pages)
        
        buffer.seek(0)
        filename = f"Report_{session.course_name}_{session.year}.pdf"
        
        current_app.logger.info(f"PDF generated successfully for session {session_id}")
        
        # Use Response instead of send_file for better cPanel compatibility
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
    except Exception as e:
        # Log detailed error information
        log_error(e, {
            'session_id': session_id,
            'function': 'download_pdf_report',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        
        current_app.logger.error(f"Error generating PDF for session {session_id}: {e}")
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))


@class_management_bp.route('/download_attendance_sheet/<int:session_id>')
@login_required
def download_attendance_sheet(session_id):
    try:
        from error_handler import log_error
        from reportlab.lib.pagesizes import landscape, letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        session = Session.query.get_or_404(session_id)
        attendance_summary = _build_attendance_summary(session)

        related_sessions = _get_related_sessions(session, include_archived=True)
        session_ids = [s.id for s in related_sessions if s]
        attendance_records = ClassAttendance.query.filter(
            ClassAttendance.session_id.in_(session_ids)
        ).order_by(ClassAttendance.date.asc(), ClassAttendance.id.asc()).all()

        attendance_by_date = defaultdict(list)
        for record in attendance_records:
            attendance_by_date[record.date].append(record)

        daily_class_counts = {}
        for date, records in attendance_by_date.items():
            student_counts = defaultdict(int)
            for r in records:
                student_counts[r.student_id] += 1
            daily_class_counts[date] = max(student_counts.values()) if student_counts else 0

        sorted_dates = sorted(daily_class_counts.keys())
        headers = []
        header_keys = []
        for date in sorted_dates:
            count = daily_class_counts.get(date, 0)
            if count <= 1:
                # Full date format: "Nov 20, 2025" or "20 Nov 2025"
                headers.append(date.strftime('%b %d, %Y'))
                header_keys.append((date, 1))
            else:
                for i in range(1, count + 1):
                    # Full date format with session: "Nov 20, 2025 (1)"
                    headers.append(f"{date.strftime('%b %d, %Y')} ({i})")
                    header_keys.append((date, i))

        if not headers:
            flash('No attendance data to download.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))

        students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()

        # Prepare data for template
        data_rows = []
        agg_student_map = attendance_summary.get('per_student', {})
        total_classes = attendance_summary.get('total_classes', sum(daily_class_counts.values()))

        for idx, student in enumerate(students, start=1):
            student_records = [r for r in attendance_records if r.student_id == student.id]
            student_attendance_by_date = defaultdict(list)
            for r in student_records:
                student_attendance_by_date[r.date].append(r)
            attendance_list = []
            for date, slot in header_keys:
                records_for_date = student_attendance_by_date[date]
                if len(records_for_date) >= slot:
                    attendance_list.append('P' if records_for_date[slot-1].is_present else 'A')
                else:
                    attendance_list.append('-')
            agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})
            data_rows.append({
                'idx': idx,
                'student_id': str(student.student_id),
                'name': student.name,
                'attendance': attendance_list,
                'total_classes': str(total_classes),
                'present_days': str(agg_stats['present']),
                'percentage': f"{agg_stats['percentage']:.2f}",
                'marks': str(agg_stats['marks'])
            })

        # Format course scope label
        scope_label = 'Full'
        if session.course_scope == 'part_a':
            scope_label = 'Part A'
        elif session.course_scope == 'part_b':
            scope_label = 'Part B'
        elif session.course_scope:
            scope_label = session.course_scope.replace('_', ' ').title()

        buffer = io.BytesIO()
        # Landscape orientation with minimal margins for maximum space
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=landscape(letter), 
            leftMargin=0.15*inch, 
            rightMargin=0.15*inch, 
            topMargin=0.2*inch, 
            bottomMargin=0.15*inch
        )
        styles = getSampleStyleSheet()
        elements = []

        # Create custom centered styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=14,
            textColor=colors.HexColor('#000000'),
            spaceAfter=8,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        meta_style = ParagraphStyle(
            'CustomMeta',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#000000'),
            spaceAfter=4,
            alignment=TA_LEFT,
            fontName='Helvetica',
            leftIndent=5
        )

        # Build header text - Excel style
        header_text = "Attendance Sheet"
        
        # Build metadata - Excel style (separate lines)
        course_code = session.course_code or ''
        course_name = session.course_name or ''
        if course_code and course_name:
            subject_text = f"Subject: {course_code} {course_name}"
        elif course_name:
            subject_text = f"Subject: {course_name}"
        elif course_code:
            subject_text = f"Subject: {course_code}"
        else:
            subject_text = "Subject: N/A"
        
        teacher_name = session.teacher.name if session.teacher else 'N/A'
        year_text = f"Year: {session.year}" if session.year else "Year: N/A"
        term_text = f"Term: {session.term}" if session.term else "Term: N/A"
        
        session_text = ""
        if session.academic_session:
            session_text = f"Session: {session.academic_session}"
        session_text += f" Course Teacher: {teacher_name}" if session_text else f"Course Teacher: {teacher_name}"

        elements.append(Spacer(1, 5))
        elements.append(Paragraph(header_text, title_style))
        elements.append(Paragraph(subject_text, meta_style))
        elements.append(Paragraph(year_text, meta_style))
        elements.append(Paragraph(term_text, meta_style))
        elements.append(Paragraph(session_text, meta_style))
        elements.append(Spacer(1, 5))

        # Create styles for text wrapping
        name_style = ParagraphStyle(
            'NameStyle',
            parent=styles['Normal'],
            fontSize=6,
            leading=7,
            alignment=0,  # LEFT
            wordWrap='CJK'
        )
        
        # Header row
        header_row = ['SI', 'Student ID', 'Name']
        for h in headers:
            header_row.append(h)
        header_row.extend(['Total Classes', 'Present', 'Percentage', 'Marks'])
        
        # Convert name column to Paragraph objects for text wrapping
        table_data = [header_row]
        for row in data_rows:
            wrapped_row = [str(row['idx']), str(row['student_id']), Paragraph(row['name'], name_style)]
            wrapped_row.extend(row['attendance'])
            wrapped_row.extend([row['total_classes'], row['present_days'], f"{row['percentage']}%", row['marks']])
            table_data.append(wrapped_row)

        # Calculate dynamic column widths - optimized for ONE page with rotated headers
        available_width = doc.width
        si_width = 0.15*inch
        student_id_width = 0.35*inch
        name_width = 0.5*inch
        summary_col_width = 0.3*inch
        summary_total_width = summary_col_width * 4

        # Calculate remaining width for attendance columns
        fixed_width = si_width + student_id_width + name_width + summary_total_width
        remaining_width = max(0.15*inch, available_width - fixed_width)
        
        # Distribute remaining width among attendance columns - narrow for rotated headers
        if len(headers) > 0:
            attendance_col_width = max(0.1*inch, remaining_width / len(headers))
        else:
            attendance_col_width = 0.1*inch

        column_widths = (
            [si_width, student_id_width, name_width] +
            [attendance_col_width] * len(headers) +
            [summary_col_width] * 4
        )

        # Ensure total width fits within available width
        total_width = sum(column_widths)
        if total_width > available_width * 0.95:
            scale_factor = (available_width * 0.95) / total_width
            column_widths = [w * scale_factor for w in column_widths]
        
        table = Table(table_data, repeatRows=1, colWidths=column_widths, hAlign='CENTER')
        
        # Build table style - optimized for ONE page with rotated headers
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D8E4BC')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTSIZE', (0, 0), (-1, 0), 6),  # Header font
            ('FONTSIZE', (0, 1), (-1, -1), 5),  # Data font
            ('ALIGN', (2, 1), (2, -1), 'LEFT'),
            ('LEFTPADDING', (0, 0), (-1, -1), 1),
            ('RIGHTPADDING', (0, 0), (-1, -1), 1),
            ('TOPPADDING', (0, 0), (-1, -1), 1),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
        ])
        
        # Rotate attendance header columns (-90 degrees) to save horizontal space
        if len(headers) > 0:
            attendance_start_col = 3
            attendance_end_col = 3 + len(headers) - 1
            table_style.add('TEXTROTATION', (attendance_start_col, 0), (attendance_end_col, 0), -90)
            table_style.add('FONTSIZE', (attendance_start_col, 0), (attendance_end_col, 0), 5.5)
            table_style.add('TOPPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 50)
            table_style.add('BOTTOMPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 50)
            table_style.add('LEFTPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 1)
            table_style.add('RIGHTPADDING', (attendance_start_col, 0), (attendance_end_col, 0), 1)
        
        table.setStyle(table_style)
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)

        filename = f"attendance_sheet_{session.course_code or session.id}.pdf"
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename=\"{filename}\"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
    except Exception as e:
        log_error(e, {
            'session_id': session_id,
            'function': 'download_attendance_sheet',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        current_app.logger.error(f"Error generating attendance sheet PDF for session {session_id}: {e}")
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.view_attendance', session_id=session_id))

@class_management_bp.route('/download_attendance_sheet_weasyprint/<int:session_id>')
@login_required
def download_attendance_sheet_weasyprint(session_id):
    """Generate attendance sheet PDF using WeasyPrint in legal landscape format."""
    # Lazy import WeasyPrint - only when actually needed
    HTML = _get_weasyprint_html()
    if HTML is None:
        from error_handler import log_error
        error_msg = 'Error generating PDF: WeasyPrint is not available. '
        error_msg += 'Please ensure WeasyPrint dependencies are installed. '
        error_msg += 'On macOS, run: brew install cairo pango gdk-pixbuf gobject-introspection'
        flash(error_msg, 'error')
        current_app.logger.error("WeasyPrint not available for PDF generation")
        current_app.logger.error(f"Current availability status: {_WEASYPRINT_AVAILABLE}")
        return redirect(url_for('class_management.view_attendance', session_id=session_id))
    
    try:
        from error_handler import log_error
        
        session = Session.query.get_or_404(session_id)
        attendance_summary = _build_attendance_summary(session)
        
        # Get related sessions for split courses
        related_sessions = _get_related_sessions(session, include_archived=True)
        session_ids = [s.id for s in related_sessions if s]
        attendance_records = ClassAttendance.query.filter(
            ClassAttendance.session_id.in_(session_ids)
        ).order_by(ClassAttendance.date.asc(), ClassAttendance.id.asc()).all()
        
        attendance_by_date = defaultdict(list)
        for record in attendance_records:
            attendance_by_date[record.date].append(record)
        
        daily_class_counts = {}
        for date, records in attendance_by_date.items():
            student_counts = defaultdict(int)
            for r in records:
                student_counts[r.student_id] += 1
            daily_class_counts[date] = max(student_counts.values()) if student_counts else 0
        
        sorted_dates = sorted(daily_class_counts.keys())
        headers = []
        header_keys = []
        for date in sorted_dates:
            count = daily_class_counts.get(date, 0)
            if count <= 1:
                headers.append(date.strftime('%b %d'))
                header_keys.append((date, 1))
            else:
                for i in range(1, count + 1):
                    headers.append(f"{date.strftime('%b %d')} ({i})")
                    header_keys.append((date, i))
        
        if not headers:
            flash('No attendance data to download.', 'warning')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
        
        students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()
        
        # Prepare data for template
        data_rows = []
        agg_student_map = attendance_summary.get('per_student', {})
        total_classes = attendance_summary.get('total_classes', sum(daily_class_counts.values()))
        
        for idx, student in enumerate(students, start=1):
            # Filter records for this specific student
            student_records = [r for r in attendance_records if r.student_id == student.id]
            # Sort student records by date and id to ensure correct order
            student_records.sort(key=lambda x: (x.date, x.id))
            
            student_attendance_by_date = defaultdict(list)
            for r in student_records:
                student_attendance_by_date[r.date].append(r)
            
            # Sort records by id for each date to ensure correct slot order
            for date in student_attendance_by_date:
                student_attendance_by_date[date].sort(key=lambda x: x.id)
            
            attendance_list = []
            for date, slot in header_keys:
                records_for_date = student_attendance_by_date.get(date, [])
                if len(records_for_date) >= slot:
                    # Records are already sorted by id, so slot-1 index is correct
                    record = records_for_date[slot-1]
                    # Explicitly check is_present value - use bool() to ensure proper conversion
                    is_present_value = bool(record.is_present) if record.is_present is not None else False
                    status = 'P' if is_present_value else 'A'
                    attendance_list.append(status)
                else:
                    attendance_list.append('-')
            agg_stats = agg_student_map.get(student.student_id, {'present': 0, 'percentage': 0, 'marks': 0})
            data_rows.append({
                'idx': idx,
                'student_id': str(student.student_id),
                'name': student.name,
                'attendance': attendance_list,
                'total_classes': str(total_classes),
                'present_days': str(agg_stats['present']),
                'percentage': f"{agg_stats['percentage']:.2f}",
                'marks': str(agg_stats['marks'])
            })
        
        # Get assignment data for Session, Year, Term, Course Teacher
        assignment = None
        if CourseSessionAssignment:
            assignment = CourseSessionAssignment.query.filter_by(session_id=session_id).first()
            # If not found by session_id, try to find by course_code, teacher_id, year, term
            if not assignment and session.course_code and session.teacher_id and session.year and session.term:
                try:
                    if Course:
                        assignment = CourseSessionAssignment.query.filter_by(
                            teacher_id=session.teacher_id,
                            year=session.year,
                            term=session.term
                        ).join(Course, CourseSessionAssignment.course_id == Course.id).filter(
                            Course.course_code == session.course_code
                        ).first()
                except Exception as query_error:
                    current_app.logger.warning(f'Error querying CourseSessionAssignment: {query_error}')
        
        # Get course information
        course_code = session.course_code or ''
        course_name = session.course_name or ''
        
        # Get session, year, term, course teacher
        academic_session = ''
        if assignment and assignment.academic_session:
            academic_session = assignment.academic_session
        elif session.academic_session:
            academic_session = session.academic_session
        
        year = session.year or ''
        term = session.term or ''
        
        course_teacher = 'N/A'
        if assignment and assignment.teacher:
            course_teacher = assignment.teacher.name
        elif session.teacher:
            course_teacher = session.teacher.name
        
        # Render template
        html_content = render_template(
            'class_management/attendance_sheet_weasyprint.html',
            course_code=course_code,
            course_name=course_name,
            academic_session=academic_session,
            year=year,
            term=term,
            course_teacher=course_teacher,
            headers=headers,
            data_rows=data_rows
        )
        
        # Generate PDF with WeasyPrint (lazy import already done above)
        try:
            pdf_buffer = io.BytesIO()
            HTML(string=html_content).write_pdf(pdf_buffer)
            pdf_buffer.seek(0)
        except Exception as e:
            current_app.logger.error(f"Error generating PDF with WeasyPrint: {e}", exc_info=True)
            flash(f'Error generating PDF: {str(e)}', 'error')
            return redirect(url_for('class_management.view_attendance', session_id=session_id))
        
        filename = f"attendance_sheet_{course_code or session.id}.pdf"
        
        current_app.logger.info(f"WeasyPrint PDF generated successfully for session {session_id}")
        return Response(
            pdf_buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(pdf_buffer.getvalue()))
            }
        )
        
    except Exception as e:
        log_error(e, {
            'session_id': session_id,
            'function': 'download_attendance_sheet_weasyprint',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        current_app.logger.error(f"Error generating WeasyPrint attendance sheet PDF for session {session_id}: {e}", exc_info=True)
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.view_attendance', session_id=session_id))

@class_management_bp.route('/assessment/<int:session_id>', methods=['GET', 'POST'])
@login_required
def assessment(session_id):
    """Assessment management for a session"""
    try:
        session = Session.query.get_or_404(session_id)
        students = ClassStudent.query.filter_by(session_id=session_id).all()
        is_split_theory = session.course_type == 'theory' and session.course_scope in SPLIT_PARTS
        editable_indices = _get_editable_assessment_indices(session)
        current_teacher = _ensure_current_teacher()
        
        # Load reveal status
        import json
        reveal_status = {}
        if session.assessment_revealed:
            try:
                reveal_status = json.loads(session.assessment_revealed)
            except:
                reveal_status = {}
        current_teacher_reveals = reveal_status.get(str(current_teacher.id), {})
        
        # Default attendance marks to revealed if not set
        if 'attendance' not in current_teacher_reveals:
            current_teacher_reveals['attendance'] = True
            # Save the default if session doesn't have reveal status yet
            if str(current_teacher.id) not in reveal_status:
                reveal_status[str(current_teacher.id)] = current_teacher_reveals
                session.assessment_revealed = json.dumps(reveal_status)
                db.session.commit()
        
        if request.method == 'POST':
            try:
                import json
                if session.course_type == 'theory':
                    for student in students:
                        # Load existing absent status
                        absent_status = {}
                        if student.assessment_absent:
                            try:
                                absent_status = json.loads(student.assessment_absent)
                            except:
                                absent_status = {}
                        
                        for i in range(1, 5):
                            if i in editable_indices:
                                absent_key = f'absent_{i}_{student.id}'
                                is_absent = absent_key in request.form
                                absent_status[f'assessment{i}'] = is_absent
                                
                                if is_absent:
                                    setattr(student, f'assessment{i}', None)
                                else:
                                    value = request.form.get(f'assessment{i}_{student.id}')
                                    setattr(student, f'assessment{i}', float(value) if value else None)
                        
                        # Save absent status
                        student.assessment_absent = json.dumps(absent_status) if absent_status else None
                    _recalculate_assessment_totals(session)

                elif session.course_type == 'sessional' and session.category == 'ug':
                    for student in students:
                        # Load existing absent status
                        absent_status = {}
                        if student.assessment_absent:
                            try:
                                absent_status = json.loads(student.assessment_absent)
                            except:
                                absent_status = {}
                        
                        report_absent_key = f'sessional_absent_report_{student.id}'
                        viva_absent_key = f'sessional_absent_viva_{student.id}'
                        
                        report_absent = report_absent_key in request.form
                        viva_absent = viva_absent_key in request.form
                        
                        absent_status['sessional_report'] = report_absent
                        absent_status['sessional_viva'] = viva_absent
                        
                        if report_absent:
                            student.sessional_report = None
                        else:
                            report = request.form.get(f'sessional_report_{student.id}')
                            student.sessional_report = float(report) if report else None
                        
                        if viva_absent:
                            student.sessional_viva = None
                        else:
                            viva = request.form.get(f'sessional_viva_{student.id}')
                            student.sessional_viva = float(viva) if viva else None
                        
                        # Save absent status
                        student.assessment_absent = json.dumps(absent_status) if absent_status else None
                else:
                    flash('Unsupported course type for assessment entry.', 'error')
                    return redirect(url_for('class_management.assessment', session_id=session_id))


                db.session.commit()
                flash('Assessment marks saved successfully!', 'success')
                return redirect(url_for('class_management.assessment', session_id=session_id))
                
            except Exception as e:
                db.session.rollback()
                flash(f'Error saving assessment: {str(e)}', 'error')
                return redirect(url_for('class_management.assessment', session_id=session_id))
        
        # Build combined assessment values from all related sessions (for split courses)
        combined_assessment_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        
        # Debug logging for split courses
        if session.split_group_id:
            current_app.logger.info(f"Assessment page - Session {session_id}: split_group_id={session.split_group_id}, course_scope={session.course_scope}, editable_indices={editable_indices}")
            if students:
                first_student = students[0]
                first_vals = combined_assessment_values.get(first_student.student_id, {})
                current_app.logger.info(f"First student {first_student.student_id} combined values: A1={first_vals.get(1)}, A2={first_vals.get(2)}, A3={first_vals.get(3)}, A4={first_vals.get(4)}")
        
        # Build absent status map for template - combine from all related sessions for split courses
        import json
        absent_status_map = {}
        
        # Get all related sessions for split courses
        related_sessions, student_map = _gather_split_student_map(session)
        all_student_records = []
        
        # Collect all student records from related sessions
        for student_id, entries in student_map.items():
            for entry in entries:
                all_student_records.append(entry)
        
        # Build absent status map combining from all sessions
        for student in students:
            absent_status = {}
            # Start with current student's absent status
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            # For split courses, also check absent status from related sessions for the same student_id
            if session.split_group_id:
                for entry in all_student_records:
                    if entry.student_id == student.student_id and entry.id != student.id:
                        if entry.assessment_absent:
                            try:
                                other_absent = json.loads(entry.assessment_absent)
                                # Merge absent status (if any assessment is absent in any session, mark as absent)
                                for key, value in other_absent.items():
                                    if key not in absent_status or not absent_status.get(key):
                                        absent_status[key] = value
                            except:
                                pass
            
            absent_status_map[student.id] = absent_status
        
        return render_template(
            'class_management/assessment.html',
            session=session,
            students=students,
            split_meta=_build_split_context(session),
            editable_indices=editable_indices,
            combined_assessment_values=combined_assessment_values,
            combined_best3=combined_best3,
            combined_pg_avg=combined_pg_avg,
            combined_pg_total=combined_pg_total,
            absent_status_map=absent_status_map,
            current_teacher_id=current_teacher.id,
            reveal_status=current_teacher_reveals,
            assessment_slot_count=len(editable_indices)
        )
        
    except Exception as e:
        flash(f'Error loading assessment page: {str(e)}', 'error')
        return redirect(url_for('class_management.index'))

@class_management_bp.route('/assessment/<int:session_id>/toggle-reveal', methods=['POST'])
@login_required
def toggle_assessment_reveal(session_id):
    """Toggle reveal status for assessment scores"""
    try:
        import json
        session = Session.query.get_or_404(session_id)
        current_teacher = _ensure_current_teacher()
        
        # Ensure current teacher owns this session or is part of split course
        if session.teacher_id != current_teacher.id:
            # Check if it's a split course and teacher is partner
            split_meta = _build_split_context(session)
            if not split_meta or not split_meta.peers:
                return jsonify({'success': False, 'message': 'Unauthorized'}), 403
            partner_ids = [peer.teacher_id for peer in split_meta.peers]
            if current_teacher.id not in partner_ids:
                return jsonify({'success': False, 'message': 'Unauthorized'}), 403
        
        data = request.get_json()
        assessment_type = data.get('assessment_type')
        revealed = data.get('revealed', False)
        teacher_id = data.get('teacher_id', current_teacher.id)
        
        # Load existing reveal status
        reveal_status = {}
        if session.assessment_revealed:
            try:
                reveal_status = json.loads(session.assessment_revealed)
            except:
                reveal_status = {}
        
        # Initialize teacher's reveal status if not exists
        teacher_key = str(teacher_id)
        if teacher_key not in reveal_status:
            reveal_status[teacher_key] = {}
        
        # Update reveal status for this assessment type
        reveal_status[teacher_key][assessment_type] = revealed
        
        # Save to database
        session.assessment_revealed = json.dumps(reveal_status)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Reveal status updated'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

@class_management_bp.route('/assessment/<int:session_id>/auto-save', methods=['POST'])
@login_required
def auto_save_assessment(session_id):
    """Auto-save assessment marks via AJAX"""
    try:
        import json
        session = Session.query.get_or_404(session_id)
        students = ClassStudent.query.filter_by(session_id=session_id).all()
        editable_indices = _get_editable_assessment_indices(session)
        
        if request.is_json:
            data = request.get_json()
        else:
            data = request.form.to_dict()
        
        try:
            if session.course_type == 'theory':
                for student in students:
                    # Load existing absent status
                    absent_status = {}
                    if student.assessment_absent:
                        try:
                            absent_status = json.loads(student.assessment_absent)
                        except:
                            absent_status = {}
                    
                    for i in range(1, 5):
                        if i in editable_indices:
                            key = f'assessment{i}_{student.id}'
                            absent_key = f'absent_{i}_{student.id}'
                            
                            # Check if absent checkbox is checked
                            is_absent = data.get(absent_key) == 'on' or data.get(absent_key) == True
                            absent_status[f'assessment{i}'] = is_absent
                            
                            # If absent, set mark to None, otherwise save the value
                            if is_absent:
                                setattr(student, f'assessment{i}', None)
                            else:
                                value = data.get(key, '')
                                setattr(student, f'assessment{i}', float(value) if value else None)
                    
                    # Save absent status
                    student.assessment_absent = json.dumps(absent_status) if absent_status else None
                    
                _recalculate_assessment_totals(session)

            elif session.course_type == 'sessional' and session.category == 'ug':
                for student in students:
                    # Load existing absent status
                    absent_status = {}
                    if student.assessment_absent:
                        try:
                            absent_status = json.loads(student.assessment_absent)
                        except:
                            absent_status = {}
                    
                    report_absent_key = f'sessional_absent_report_{student.id}'
                    viva_absent_key = f'sessional_absent_viva_{student.id}'
                    
                    report_absent = data.get(report_absent_key) == 'on' or data.get(report_absent_key) == True
                    viva_absent = data.get(viva_absent_key) == 'on' or data.get(viva_absent_key) == True
                    
                    absent_status['sessional_report'] = report_absent
                    absent_status['sessional_viva'] = viva_absent
                    
                    if report_absent:
                        student.sessional_report = None
                    else:
                        report = data.get(f'sessional_report_{student.id}', '')
                        student.sessional_report = float(report) if report else None
                    
                    if viva_absent:
                        student.sessional_viva = None
                    else:
                        viva = data.get(f'sessional_viva_{student.id}', '')
                        student.sessional_viva = float(viva) if viva else None
                    
                    # Save absent status
                    student.assessment_absent = json.dumps(absent_status) if absent_status else None
            else:
                return jsonify({'success': False, 'message': 'Unsupported course type'}), 400

            db.session.commit()
            return jsonify({'success': True, 'message': 'Assessment marks saved automatically'})
            
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error auto-saving assessment for session {session_id}: {e}", exc_info=True)
            return jsonify({'success': False, 'message': f'Error saving: {str(e)}'}), 500
            
    except Exception as e:
        current_app.logger.error(f"Error in auto_save_assessment for session {session_id}: {e}", exc_info=True)
        return jsonify({'success': False, 'message': 'Server error'}), 500

@class_management_bp.route('/student/view-scores')
@login_required
def student_view_scores():
    """Student view for revealed assessment and attendance scores"""
    try:
        import json
        from flask_login import current_user
        
        # Get student ID from current user (assuming username is student_id)
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('index'))
        
        # Find all ClassStudent records for this student
        student_records = ClassStudent.query.filter_by(student_id=student_id).all()
        
        # First pass: collect all sessions and identify split groups
        # Filter: Only show Theory courses (course_type == 'theory')
        session_records_map = {}  # session_id -> Session
        split_group_processed = set()  # Track processed split_group_ids to avoid duplicate processing
        processed_sessions = set()  # Track which sessions we've processed
        course_map = {}  # unique_course_key -> {'session_ids': set(), 'student_record_ids': set(), 'student_records': []}
        
        # Collect all sessions first
        for student_record in student_records:
            session_obj = Session.query.get(student_record.session_id)
            if not session_obj or session_obj.archived:
                continue
            
            # Filter: Only process Theory courses
            # Also check course_name to catch cases where course_type might be incorrectly set
            if session_obj.course_type != 'theory':
                continue
            
            # Additional check: Filter out courses with "Sessional" in the name (case-insensitive)
            # This catches cases where course_type might be incorrectly set to 'theory'
            if session_obj.course_name and 'sessional' in session_obj.course_name.lower():
                continue
            
            session_records_map[session_obj.id] = session_obj
        
        # Second pass: Group sessions by course (handle split courses properly)
        for session_id, session_obj in session_records_map.items():
            if session_id in processed_sessions:
                continue
            
            # Create unique key for grouping by course
            if session_obj.split_group_id:
                # Split course: group by split_group_id
                # Process all related sessions at once to avoid duplicates
                if session_obj.split_group_id in split_group_processed:
                    continue  # Already processed this split group
                split_group_processed.add(session_obj.split_group_id)
                
                course_key = f"split_{session_obj.split_group_id}"
                related_sessions = _get_related_sessions(session_obj)
                
                if course_key not in course_map:
                    course_map[course_key] = {
                        'session_ids': set(),
                        'student_record_ids': set(),
                        'student_records': []
                    }
                
                # Add all related sessions and their student records
                for related_session in related_sessions:
                    if related_session.course_type != 'theory' or related_session.archived:
                        continue
                    # Additional check: Filter out courses with "Sessional" in the name
                    if related_session.course_name and 'sessional' in related_session.course_name.lower():
                        continue
                    if related_session.id in processed_sessions:
                        continue
                    
                    session_records_map[related_session.id] = related_session
                    course_map[course_key]['session_ids'].add(related_session.id)
                    processed_sessions.add(related_session.id)
                    
                    # Add student records from this related session
                    related_student_records = ClassStudent.query.filter_by(
                        session_id=related_session.id,
                        student_id=student_id
                    ).all()
                    for related_rec in related_student_records:
                        if related_rec.id not in course_map[course_key]['student_record_ids']:
                            course_map[course_key]['student_record_ids'].add(related_rec.id)
                            course_map[course_key]['student_records'].append(related_rec)
            else:
                # Regular course: group by course_name, course_code, year, term, academic_session
                course_key = (
                    str(session_obj.course_name or '').strip().lower(),
                    str(session_obj.course_code or '').strip().lower(),
                    str(session_obj.year or '').strip(),
                    str(session_obj.term or '').strip(),
                    str(session_obj.academic_session or '').strip()
                )
                
                if course_key not in course_map:
                    course_map[course_key] = {
                        'session_ids': set(),
                        'student_record_ids': set(),
                        'student_records': []
                    }
                
                course_map[course_key]['session_ids'].add(session_obj.id)
                processed_sessions.add(session_obj.id)
                
                # Add student records for this session
                session_student_records = ClassStudent.query.filter_by(
                    session_id=session_obj.id,
                    student_id=student_id
                ).all()
                for rec in session_student_records:
                    if rec.id not in course_map[course_key]['student_record_ids']:
                        course_map[course_key]['student_record_ids'].add(rec.id)
                        course_map[course_key]['student_records'].append(rec)
        
        # Second pass: build reveal status by combining all sessions with same course key
        for course_key, course_data in course_map.items():
            reveal_status = {}
            # Default attendance to revealed
            reveal_status['attendance'] = True
            
            # Combine reveal status from all sessions with this course key
            for session_id in course_data['session_ids']:
                session_obj = session_records_map.get(session_id)
                if session_obj and session_obj.assessment_revealed:
                    try:
                        all_reveals = json.loads(session_obj.assessment_revealed)
                        for teacher_id, teacher_reveals in all_reveals.items():
                            for assessment_type, is_revealed in teacher_reveals.items():
                                if is_revealed:
                                    reveal_status[assessment_type] = True
                    except:
                        pass
            
            course_data['reveal_status'] = reveal_status
            # Convert session_ids set to list and get primary session
            session_ids_list = sorted(list(course_data['session_ids']))
            course_data['primary_session'] = session_records_map[session_ids_list[0]]  # Use first session as primary
        
        # Build combined data for each unique course
        courses_data = []
        for course_key, course_data in course_map.items():
            session_obj = course_data['primary_session']  # Use primary session for display
            
            # Double-check: Only process Theory courses (skip Sessional courses)
            if session_obj.course_type != 'theory':
                continue
            
            # Additional check: Filter out courses with "Sessional" in the name (case-insensitive)
            # This catches cases where course_type might be incorrectly set to 'theory'
            if session_obj.course_name and 'sessional' in session_obj.course_name.lower():
                continue
            
            student_records = course_data['student_records']
            reveal_status = course_data['reveal_status']
            all_session_ids = course_data['session_ids']
            
            # Use the first student record as primary (prioritize non-null values)
            primary_record = student_records[0]
            
            # Build assessment scores based on reveal status
            # For split courses, use _build_combined_assessment_values to properly combine from all related sessions
            assessment_scores = {}
            best3_total = None
            pg_total = None
            
            if session_obj.course_type == 'theory' and session_obj.category == 'ug':
                # Use _build_combined_assessment_values which properly handles split courses
                # This function combines assessments from Part A (assessments 1-2) and Part B (assessments 3-4)
                combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session_obj)
                
                # Get combined assessment values for this student
                student_combined_values = combined_values.get(student_id, {})
                
                # Build assessment_scores dict from combined values based on reveal status
                for i in range(1, 5):
                    assessment_key = f'assessment{i}'
                    if reveal_status.get(assessment_key, False):
                        # Use combined value which includes assessments from all related sessions
                        value = student_combined_values.get(i, None)
                        assessment_scores[assessment_key] = value
                
                # Get best3_total from combined calculation (includes all assessments from split sessions)
                best3_total = combined_best3.get(student_id, None)
            elif session_obj.course_type == 'theory' and session_obj.category == 'pg':
                # Use _build_combined_assessment_values which properly handles split courses
                # This function combines assessments from Part A (assessments 1-2) and Part B (assessments 3-4)
                combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session_obj)
                
                # Get combined assessment values for this student
                student_combined_values = combined_values.get(student_id, {})
                
                # Build assessment_scores dict from combined values based on reveal status
                for i in range(1, 5):
                    assessment_key = f'assessment{i}'
                    if reveal_status.get(assessment_key, False):
                        # Use combined value which includes assessments from all related sessions
                        value = student_combined_values.get(i, None)
                        assessment_scores[assessment_key] = value
                
                # Get pg_total from combined calculation (includes all assessments from split sessions)
                # This is calculated as: (best 3 sum / 30) * 40
                pg_total = combined_pg_total.get(student_id, None)
            # Note: Sessional courses are filtered out - only Theory courses should reach this point
            
            
            # Get attendance marks if revealed
            # For split courses, aggregate attendance from all related sessions
            attendance_data = None
            if reveal_status.get('attendance', True):  # Default to True if not set
                # Use _build_attendance_summary which already handles split courses correctly
                # It aggregates attendance from all related sessions automatically
                attendance_summary = _build_attendance_summary(session_obj)
                
                # Find the student's attendance stats from the summary
                # Use student_id (string) for lookup
                student_stats = attendance_summary.get('per_student', {}).get(student_id, {})
                
                if student_stats:
                    attendance_data = {
                        'present_count': student_stats.get('present', 0),
                        'total_classes': attendance_summary.get('total_classes', 0),
                        'percentage': student_stats.get('percentage', 0),
                        'marks': student_stats.get('marks', 0)
                    }
            
            # Build teacher options (split courses may have multiple teachers)
            teacher_options = []
            try:
                related_sessions = _get_related_sessions(session_obj)
                for related_session in related_sessions:
                    if related_session and related_session.teacher:
                        teacher_options.append({
                            'session_id': related_session.id,
                            'teacher_id': related_session.teacher.id,
                            'teacher_name': related_session.teacher.name,
                            'teacher_short': related_session.teacher.short_name,
                            'scope_label': COURSE_SCOPE_LABELS.get(related_session.course_scope, 'Part')
                        })
            except Exception as teacher_error:
                current_app.logger.warning(f"Error building teacher options: {teacher_error}")

            # Load Q&A threads for this course (student-specific)
            qa_threads = []
            qa_new_reply_count = 0
            try:
                from sqlalchemy.orm import selectinload
                session_ids = list(course_data['session_ids'])
                if session_ids:
                    qa_threads = CourseQuestionThread.query.options(
                        selectinload(CourseQuestionThread.messages).selectinload(CourseQuestionMessage.attachments)
                    ).filter(
                        CourseQuestionThread.session_id.in_(session_ids),
                        CourseQuestionThread.student_id == student_id
                    ).order_by(CourseQuestionThread.created_at.desc()).all()
                    # Annotate threads with latest message role for notifications
                    for thread in qa_threads:
                        last_message = None
                        if thread.messages:
                            last_message = max(
                                thread.messages,
                                key=lambda m: m.created_at or datetime.min
                            )
                        thread.last_message_role = last_message.sender_role if last_message else None
                        if thread.last_message_role == 'teacher':
                            qa_new_reply_count += 1
            except Exception as qa_error:
                current_app.logger.warning(f"Error loading Q&A threads for student {student_id}: {qa_error}")

            courses_data.append({
                'session': session_obj,
                'student_record': primary_record,
                'assessment_scores': assessment_scores,
                'best3_total': best3_total,
                'pg_total': pg_total,
                'attendance_data': attendance_data,
                'reveal_status': reveal_status,
                'qa_threads': qa_threads,
                'qa_new_reply_count': qa_new_reply_count,
                'teacher_options': teacher_options
            })
        
        # Sort courses by course_name, then by year-term for consistent display
        courses_data.sort(key=lambda x: (
            x['session'].course_name or '',
            x['session'].year or '',
            x['session'].term or ''
        ))
        
        return render_template(
            'class_management/student_view_scores.html',
            student_id=student_id,
            courses_data=courses_data
        )
        
    except Exception as e:
        current_app.logger.error(f"Error loading student view scores: {e}", exc_info=True)
        flash(f'Error loading scores: {str(e)}', 'error')
        return redirect(url_for('index'))


@class_management_bp.route('/student/course-questions/<int:session_id>/create', methods=['POST'])
@login_required
def student_create_course_question(session_id):
    """Create a new Q&A thread from student."""
    student_id = current_user.username if hasattr(current_user, 'username') else None
    if not student_id:
        flash('Student ID not found.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    session_obj = Session.query.get_or_404(session_id)
    related_sessions = _get_related_sessions(session_obj)
    related_session_ids = [s.id for s in related_sessions if s]

    selected_session_id = request.form.get('selected_session_id', type=int) or session_id
    if selected_session_id not in related_session_ids:
        flash('Invalid teacher selection for this course.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    # Ensure student is enrolled in any related session
    student_record = ClassStudent.query.filter(
        ClassStudent.student_id == student_id,
        ClassStudent.session_id.in_(related_session_ids)
    ).first()
    if not student_record:
        flash('You are not enrolled in this course.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    selected_session = Session.query.get_or_404(selected_session_id)

    subject = request.form.get('subject', '').strip()
    message_body = request.form.get('message', '').strip()
    files = request.files.getlist('attachments')

    if not subject:
        flash('Please provide a subject for your question.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    if not message_body and not files:
        flash('Please write a message or attach a file.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    try:
        thread = CourseQuestionThread(
            session_id=selected_session_id,
            student_id=student_id,
            student_name=student_record.name or getattr(current_user, 'full_name', '') or student_id,
            teacher_id=selected_session.teacher_id,
            subject=subject,
            status='open'
        )
        db.session.add(thread)
        db.session.flush()

        message = CourseQuestionMessage(
            thread_id=thread.id,
            sender_role='student',
            sender_user_id=None,
            body=message_body
        )
        db.session.add(message)
        db.session.flush()

        saved_attachments = _save_qa_attachments(files, thread.id)
        for attachment in saved_attachments:
            db.session.add(CourseQuestionAttachment(
                message_id=message.id,
                file_name=attachment['file_name'],
                file_path=attachment['file_path'],
                file_size=attachment.get('file_size'),
                file_type=attachment.get('file_type')
            ))

        db.session.commit()
        flash('Your question has been sent to the teacher.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating question thread: {e}", exc_info=True)
        flash('Failed to send your question. Please try again.', 'error')

    return redirect(url_for('class_management.student_view_scores'))


@class_management_bp.route('/student/course-questions/<int:thread_id>/reply', methods=['POST'])
@login_required
def student_reply_course_question(thread_id):
    """Reply to an existing Q&A thread from student."""
    student_id = current_user.username if hasattr(current_user, 'username') else None
    if not student_id:
        flash('Student ID not found.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    thread = CourseQuestionThread.query.get_or_404(thread_id)
    if thread.student_id != student_id:
        flash('You are not authorized to reply to this thread.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    message_body = request.form.get('message', '').strip()
    files = request.files.getlist('attachments')

    if not message_body and not files:
        flash('Please write a message or attach a file.', 'error')
        return redirect(url_for('class_management.student_view_scores'))

    try:
        message = CourseQuestionMessage(
            thread_id=thread.id,
            sender_role='student',
            sender_user_id=None,
            body=message_body
        )
        db.session.add(message)
        db.session.flush()

        saved_attachments = _save_qa_attachments(files, thread.id)
        for attachment in saved_attachments:
            db.session.add(CourseQuestionAttachment(
                message_id=message.id,
                file_name=attachment['file_name'],
                file_path=attachment['file_path'],
                file_size=attachment.get('file_size'),
                file_type=attachment.get('file_type')
            ))

        thread.updated_at = datetime.utcnow()
        db.session.commit()
        flash('Reply sent successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error replying to question thread: {e}", exc_info=True)
        flash('Failed to send reply. Please try again.', 'error')

    return redirect(url_for('class_management.student_view_scores'))


@class_management_bp.route('/course-questions/<int:session_id>')
@login_required
def course_questions(session_id):
    """Teacher view of course Q&A threads for a session."""
    session_obj = Session.query.get_or_404(session_id)
    teacher = _ensure_current_teacher()

    if not teacher or (teacher.id != session_obj.teacher_id and not is_admin(current_user)):
        flash('You are not authorized to view these questions.', 'error')
        return redirect(url_for('class_management.index'))

    try:
        from sqlalchemy.orm import selectinload
        threads = CourseQuestionThread.query.options(
            selectinload(CourseQuestionThread.messages).selectinload(CourseQuestionMessage.attachments)
        ).filter_by(session_id=session_id).order_by(CourseQuestionThread.created_at.desc()).all()
        for thread in threads:
            last_message = None
            if thread.messages:
                last_message = max(
                    thread.messages,
                    key=lambda m: m.created_at or datetime.min
                )
            thread.last_message_role = last_message.sender_role if last_message else None
    except Exception as e:
        current_app.logger.error(f"Error loading course questions: {e}", exc_info=True)
        threads = []

    return render_template(
        'class_management/course_questions.html',
        session=session_obj,
        threads=threads
    )


@class_management_bp.route('/course-questions/<int:thread_id>/reply', methods=['POST'])
@login_required
def teacher_reply_course_question(thread_id):
    """Teacher reply to a Q&A thread."""
    thread = CourseQuestionThread.query.get_or_404(thread_id)
    session_obj = Session.query.get_or_404(thread.session_id)
    teacher = _ensure_current_teacher()

    if not teacher or (teacher.id != session_obj.teacher_id and not is_admin(current_user)):
        flash('You are not authorized to reply to this thread.', 'error')
        return redirect(url_for('class_management.index'))

    message_body = request.form.get('message', '').strip()
    files = request.files.getlist('attachments')

    if not message_body and not files:
        flash('Please write a message or attach a file.', 'error')
        return redirect(url_for('class_management.course_questions', session_id=session_obj.id))

    try:
        message = CourseQuestionMessage(
            thread_id=thread.id,
            sender_role='teacher',
            sender_user_id=teacher.id,
            body=message_body
        )
        db.session.add(message)
        db.session.flush()

        saved_attachments = _save_qa_attachments(files, thread.id)
        for attachment in saved_attachments:
            db.session.add(CourseQuestionAttachment(
                message_id=message.id,
                file_name=attachment['file_name'],
                file_path=attachment['file_path'],
                file_size=attachment.get('file_size'),
                file_type=attachment.get('file_type')
            ))

        thread.updated_at = datetime.utcnow()
        db.session.commit()
        flash('Reply sent successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error replying to question thread: {e}", exc_info=True)
        flash('Failed to send reply. Please try again.', 'error')

    return redirect(url_for('class_management.course_questions', session_id=session_obj.id))


@class_management_bp.route('/course-questions/attachments/<int:attachment_id>')
@login_required
def download_course_question_attachment(attachment_id):
    """Download attachment for course Q&A."""
    try:
        attachment = CourseQuestionAttachment.query.get_or_404(attachment_id)
        message = CourseQuestionMessage.query.get_or_404(attachment.message_id)
        thread = CourseQuestionThread.query.get_or_404(message.thread_id)
        session_obj = Session.query.get_or_404(thread.session_id)

        student_id = current_user.username if hasattr(current_user, 'username') else None
        teacher = Teacher.query.filter_by(name=current_user.full_name).first()

        is_student_allowed = bool(student_id and thread.student_id == student_id)
        is_teacher_allowed = bool(teacher and teacher.id == session_obj.teacher_id)
        if not (is_student_allowed or is_teacher_allowed or is_admin(current_user)):
            flash('You are not authorized to access this file.', 'error')
            return redirect(url_for('index'))

        if not os.path.exists(attachment.file_path):
            flash('File not found.', 'error')
            return redirect(url_for('index'))

        return send_file(
            attachment.file_path,
            as_attachment=True,
            download_name=attachment.file_name
        )
    except Exception as e:
        current_app.logger.error(f"Error downloading question attachment: {e}", exc_info=True)
        flash('Error downloading file.', 'error')
        return redirect(url_for('index'))


@class_management_bp.route('/course-questions/<int:thread_id>/delete', methods=['POST'])
@login_required
def delete_course_question_thread(thread_id):
    """Delete a Q&A thread (question)."""
    thread = CourseQuestionThread.query.get_or_404(thread_id)
    session_obj = Session.query.get_or_404(thread.session_id)
    student_id = current_user.username if hasattr(current_user, 'username') else None
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    is_student_owner = bool(student_id and thread.student_id == student_id)
    is_teacher_owner = bool(teacher and teacher.id == session_obj.teacher_id)
    if not (is_student_owner or is_teacher_owner or is_admin(current_user)):
        flash('You are not authorized to delete this question.', 'error')
        return redirect(url_for('index'))

    try:
        # Delete files for all attachments
        for msg in thread.messages:
            _delete_qa_attachments(msg.attachments)
        db.session.delete(thread)
        db.session.commit()
        flash('Question deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting question thread: {e}", exc_info=True)
        flash('Failed to delete question.', 'error')

    if is_teacher_owner:
        return redirect(url_for('class_management.course_questions', session_id=session_obj.id))
    return redirect(url_for('class_management.student_view_scores'))


@class_management_bp.route('/course-questions/message/<int:message_id>/delete', methods=['POST'])
@login_required
def delete_course_question_message(message_id):
    """Delete a single message in a Q&A thread."""
    message = CourseQuestionMessage.query.get_or_404(message_id)
    thread = CourseQuestionThread.query.get_or_404(message.thread_id)
    session_obj = Session.query.get_or_404(thread.session_id)
    student_id = current_user.username if hasattr(current_user, 'username') else None
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    is_student_owner = bool(student_id and thread.student_id == student_id and message.sender_role == 'student')
    is_teacher_owner = bool(teacher and message.sender_role == 'teacher' and message.sender_user_id == teacher.id)
    if not (is_student_owner or is_teacher_owner or is_admin(current_user)):
        flash('You are not authorized to delete this message.', 'error')
        return redirect(url_for('index'))

    try:
        _delete_qa_attachments(message.attachments)
        db.session.delete(message)
        db.session.commit()
        flash('Message deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting question message: {e}", exc_info=True)
        flash('Failed to delete message.', 'error')

    if is_teacher_owner:
        return redirect(url_for('class_management.course_questions', session_id=session_obj.id))
    return redirect(url_for('class_management.student_view_scores'))

@class_management_bp.route('/download_assessment_excel/<int:session_id>')
@login_required
def download_assessment_excel(session_id):
    """Download assessment data as Excel file"""
    try:
        import json
        session = Session.query.get_or_404(session_id)
        students = ClassStudent.query.filter_by(session_id=session_id).all()
        combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        
        # Helper function to get assessment value with absent check
        def get_assessment_value(student, assessment_num):
            """Get assessment value, showing 'Absent' if marked absent"""
            absent_status = {}
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            is_absent = absent_status.get(f'assessment{assessment_num}', False)
            if is_absent:
                return 'Absent'
            
            combined = combined_values.get(student.student_id, {})
            value = combined.get(assessment_num)
            return value if value is not None else ''
        
        # Helper function to get sessional value with absent check
        def get_sessional_value(student, sessional_type):
            """Get sessional value, showing 'Absent' if marked absent"""
            absent_status = {}
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            is_absent = absent_status.get(f'sessional_{sessional_type}', False)
            if is_absent:
                return 'Absent'
            
            if sessional_type == 'report':
                value = student.sessional_report
            else:  # viva
                value = student.sessional_viva
            
            return int(round(value)) if value is not None else ''
        
        # Build data for DataFrame
        data = []
        if session.course_type == 'theory' and session.category == 'ug':
            columns = ['Student ID', 'Name', 'Assessment 1', 'Assessment 2', 'Assessment 3', 'Assessment 4', 'Total of Best 3']
            for s in students:
                best3_total = combined_best3.get(s.student_id) if combined_best3 else None
                data.append([
                    s.student_id,
                    s.name,
                    get_assessment_value(s, 1),
                    get_assessment_value(s, 2),
                    get_assessment_value(s, 3),
                    get_assessment_value(s, 4),
                    best3_total
                ])
        elif session.course_type == 'theory' and session.category == 'pg':
            columns = ['Student ID', 'Name', 'Assessment 1', 'Assessment 2', 'Assessment 3', 'Assessment 4', 'Total (40)']
            for s in students:
                data.append([
                    s.student_id,
                    s.name,
                    get_assessment_value(s, 1),
                    get_assessment_value(s, 2),
                    get_assessment_value(s, 3),
                    get_assessment_value(s, 4),
                    combined_pg_total.get(s.student_id)
                ])
        elif session.course_type == 'sessional' and session.category == 'ug':
            columns = ['Student ID', 'Name', 'Sessional Report (60)', 'Sessional Viva (30)', 'Total (Sessional: 90)']
            for s in students:
                report_val = get_sessional_value(s, 'report')
                viva_val = get_sessional_value(s, 'viva')
                
                # Calculate total (skip if either is 'Absent')
                if report_val == 'Absent' or viva_val == 'Absent':
                    total = 'Absent'
                else:
                    total = (s.sessional_report or 0) + (s.sessional_viva or 0)
                    total = int(round(total)) if total else ''
                
                data.append([
                    s.student_id,
                    s.name,
                    report_val,
                    viva_val,
                    total
                ])
        else:
            flash('Unsupported course type for assessment export', 'error')
            return redirect(url_for('class_management.assessment', session_id=session_id))
        
        # Create DataFrame and Excel file
        df = pd.DataFrame(data, columns=columns)
        
        # Create Excel file in memory
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Assessment', index=False)
            
            # Auto-adjust column widths
            worksheet = writer.sheets['Assessment']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        
        filename = f"assessment_{session.course_name or session.term}_{session.year}_{session.term}.xlsx"
        
        # Use Response instead of send_file for better cPanel compatibility
        return Response(
            output.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(output.getvalue()))
            }
        )
        
    except Exception as e:
        flash(f'Error downloading assessment Excel: {str(e)}', 'error')
        return redirect(url_for('class_management.assessment', session_id=session_id))

@class_management_bp.route('/download_assessment_pdf/<int:session_id>')
@login_required
def download_assessment_pdf(session_id):
    """Download assessment marks as PDF"""
    try:
        import json
        session = Session.query.get_or_404(session_id)
        students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()
        combined_values, combined_best3, combined_pg_avg, combined_pg_total = _build_combined_assessment_values(session)
        
        # Helper function to get assessment value with absent check
        def get_assessment_value(student, assessment_num):
            """Get assessment value, showing 'Absent' if marked absent"""
            absent_status = {}
            if student.assessment_absent:
                try:
                    absent_status = json.loads(student.assessment_absent)
                except:
                    absent_status = {}
            
            is_absent = absent_status.get(f'assessment{assessment_num}', False)
            if is_absent:
                return 'Absent'
            
            combined = combined_values.get(student.student_id, {})
            value = combined.get(assessment_num)
            return value if value is not None else '-'
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=0.3*inch,
            rightMargin=0.3*inch,
            topMargin=0.3*inch,
            bottomMargin=0.3*inch
        )
        styles = getSampleStyleSheet()
        elements = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.black,
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        title = Paragraph("Continuous Assessment Marks", title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.15*inch))
        
        # Course Information
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=5,
            alignment=TA_LEFT,
            fontName='Helvetica'
        )
        
        # Course Information with spacing between items
        info_items = []
        if session.course_code:
            info_items.append(f"<b>Subject Code:</b> {session.course_code}")
        if session.course_name:
            info_items.append(f"<b>Subject Name:</b> {session.course_name}")
        if session.academic_session:
            info_items.append(f"<b>Session:</b> {session.academic_session}")
        if session.year:
            info_items.append(f"<b>Year:</b> {session.year}")
        if session.term:
            info_items.append(f"<b>Term:</b> {session.term}")
        
        # Add items with spacing between them
        for i, info_line in enumerate(info_items):
            elements.append(Paragraph(info_line, info_style))
            if i < len(info_items) - 1:  # Add space between items, but not after the last one
                elements.append(Spacer(1, 0.1*inch))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Table data
        if session.course_type == 'theory' and session.category == 'ug':
            # UG Theory: Assessment 1-4 and Total of Best 3
            table_data = [['SI', 'Student ID', 'Assessment 1 (10)', 'Assessment 2 (10)', 'Assessment 3 (10)', 'Assessment 4 (10)', 'Total of Best 3 (30)']]
            
            for idx, student in enumerate(students, start=1):
                best3_total = combined_best3.get(student.student_id) if combined_best3 else '-'
                # Format total: Keep fraction for UG (show decimals if not whole number)
                if best3_total != '-' and best3_total is not None:
                    if best3_total == int(best3_total):
                        formatted_total = str(int(best3_total))
                    else:
                        formatted_total = f"{best3_total:.2f}".rstrip('0').rstrip('.')
                else:
                    formatted_total = '-'
                row = [
                    str(idx),
                    str(student.student_id),
                    str(get_assessment_value(student, 1)),
                    str(get_assessment_value(student, 2)),
                    str(get_assessment_value(student, 3)),
                    str(get_assessment_value(student, 4)),
                    formatted_total
                ]
                table_data.append(row)
        
        elif session.course_type == 'theory' and session.category == 'pg':
            # PG Theory: Assessment 1-4 and Total (40)
            table_data = [['SI', 'Student ID', 'Assessment 1 (10)', 'Assessment 2 (10)', 'Assessment 3 (10)', 'Assessment 4 (10)', 'Total (40)']]
            
            for idx, student in enumerate(students, start=1):
                pg_total = combined_pg_total.get(student.student_id) if combined_pg_total else '-'
                # Format total: Round for PG (integer)
                if pg_total != '-' and pg_total is not None:
                    formatted_total = str(int(round(pg_total)))
                else:
                    formatted_total = '-'
                row = [
                    str(idx),
                    str(student.student_id),
                    str(get_assessment_value(student, 1)),
                    str(get_assessment_value(student, 2)),
                    str(get_assessment_value(student, 3)),
                    str(get_assessment_value(student, 4)),
                    formatted_total
                ]
                table_data.append(row)
        
        else:
            flash('Assessment PDF is only available for theory courses.', 'error')
            return redirect(url_for('class_management.assessment', session_id=session_id))
        
        # Create table
        table = Table(table_data, repeatRows=1)
        
        # Table style - Black and white
        table_style = TableStyle([
            # Header row - White background, black text
            ('BACKGROUND', (0, 0), (-1, 0), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('TOPPADDING', (0, 0), (-1, 0), 8),
            
            # Data rows
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1.5, colors.black),  # Thicker borders
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            
            # All rows white background (black and white)
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            
            # Column widths - optimized for one page (without Name column)
            ('COLWIDTH', (0, 0), (0, -1), 0.4*inch),  # SI
            ('COLWIDTH', (1, 0), (1, -1), 1.1*inch),  # Student ID
            ('COLWIDTH', (2, 0), (2, -1), 1.0*inch),  # Assessment 1 (10)
            ('COLWIDTH', (3, 0), (3, -1), 1.0*inch),  # Assessment 2 (10)
            ('COLWIDTH', (4, 0), (4, -1), 1.0*inch),  # Assessment 3 (10)
            ('COLWIDTH', (5, 0), (5, -1), 1.0*inch),  # Assessment 4 (10)
            ('COLWIDTH', (6, 0), (6, -1), 1.1*inch),  # Total
        ])
        
        table.setStyle(table_style)
        elements.append(table)
        
        # Page number callback function
        def add_page_number(canvas_obj, doc_obj):
            """Add page number to each page"""
            canvas_obj.saveState()
            page_num = canvas_obj.getPageNumber()
            text = f"Page {page_num}"
            width, height = letter
            canvas_obj.setFont('Helvetica', 9)
            canvas_obj.drawRightString(width - 0.3*inch, 0.2*inch, text)
            canvas_obj.restoreState()
        
        # Build PDF with page numbers
        doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
        buffer.seek(0)
        
        filename = f"assessment_{session.course_code or 'marks'}_{session.year}_{session.term}.pdf"
        
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
        
    except Exception as e:
        current_app.logger.error(f"Error generating assessment PDF for session {session_id}: {e}", exc_info=True)
        flash(f'Error generating assessment PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.assessment', session_id=session_id))

# Template filter for dynamic attribute access
@class_management_bp.app_template_filter('getattr')
def jinja_getattr(obj, name):
    return getattr(obj, name) 

@class_management_bp.route('/evaluation/<int:session_id>')
@login_required
def evaluation(session_id):
    """Placeholder Evaluation page for a session. Forms will be added later."""
    session = Session.query.get_or_404(session_id)
    students = ClassStudent.query.filter_by(session_id=session_id).order_by(ClassStudent.student_id).all()
    return render_template('class_management/evaluation.html', session=session, students=students)

@class_management_bp.route('/evaluation/<int:session_id>/course-assessment', methods=['GET', 'POST'])
@login_required
def course_assessment(session_id):
    """Invite other teachers to evaluate this course; show existing invitations."""
    session = Session.query.get_or_404(session_id)
    inviter_teacher = Teacher.query.filter_by(id=session.teacher_id).first()

    if request.method == 'POST':
        try:
            evaluator_teacher_id = int(request.form.get('evaluator_teacher_id'))
            if evaluator_teacher_id == session.teacher_id:
                flash('You cannot invite yourself to evaluate.', 'warning')
                return redirect(url_for('class_management.course_assessment', session_id=session_id))

            # prevent duplicate invites
            existing = EvaluationInvite.query.filter_by(
                session_id=session_id,
                evaluator_teacher_id=evaluator_teacher_id
            ).first()
            if existing:
                if existing.status == 'cancelled':
                    # remove any previous submission
                    EvaluationSubmission.query.filter_by(invite_id=existing.id).delete()
                    existing.status = 'invited'
                    existing.created_at = datetime.utcnow()
                    db.session.commit()
                    flash('Invitation re-activated.', 'success')
                else:
                    flash('This teacher is already invited for this course.', 'info')
            else:
                invite = EvaluationInvite(
                    session_id=session_id,
                    inviter_teacher_id=session.teacher_id,
                    evaluator_teacher_id=evaluator_teacher_id,
                    status='invited'
                )
                db.session.add(invite)
                db.session.commit()
                flash('Invitation created successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating invitation: {str(e)}', 'danger')
        return redirect(url_for('class_management.course_assessment', session_id=session_id))

    # List of other teachers to invite (exclude self, head, and teaching assistants)
    from role_utils import get_teachers_excluding_head
    all_teachers = get_teachers_excluding_head()
    other_teachers = [t for t in all_teachers if t.id != session.teacher_id]
    # Existing invites for this session
    invites = EvaluationInvite.query.filter_by(session_id=session_id).all()
    session_ids = [inv.session_id for inv in invites]
    sessions_by_id = {s.id: s for s in Session.query.filter(Session.id.in_(session_ids)).all()} if session_ids else {}
    teacher_ids = {inv.evaluator_teacher_id for inv in invites}
    teachers_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(teacher_ids)).all()} if teacher_ids else {}

    return render_template(
        'class_management/evaluation_course_assessment.html',
        session=session,
        inviter_teacher=inviter_teacher,
        other_teachers=other_teachers,
        invites=invites,
        sessions_by_id=sessions_by_id,
        teachers_by_id=teachers_by_id
    )

# Context processor to inject pending invites count into all templates
@class_management_bp.app_context_processor
def inject_invites_count():
    try:
        if current_user.is_authenticated and has_teacher_privileges(current_user):
            teacher = Teacher.query.filter_by(name=current_user.full_name).first()
            if teacher:
                count = EvaluationInvite.query.filter_by(evaluator_teacher_id=teacher.id, status='invited').count()
                exam_count = ExamScrutinizerInvite.query.filter_by(scrutinizer_teacher_id=teacher.id, status='invited').count()
                split_count = ClassSplitInvite.query.filter_by(invited_teacher_id=teacher.id, status='pending').count()
                return { 'pending_invites_count': count + exam_count + split_count }
    except Exception:
        pass
    return { 'pending_invites_count': 0 }

@class_management_bp.route('/invitations')
@login_required
def my_invitations():
    """List invitations for the logged-in teacher."""
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher:
        flash('No teacher profile found.', 'warning')
        return redirect(url_for('index'))
    invites = EvaluationInvite.query.filter_by(evaluator_teacher_id=teacher.id).order_by(EvaluationInvite.created_at.desc()).all()
    session_ids = [inv.session_id for inv in invites]
    sessions_by_id = {s.id: s for s in Session.query.filter(Session.id.in_(session_ids)).all()} if session_ids else {}
    inviter_ids = [inv.inviter_teacher_id for inv in invites]
    inviter_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(inviter_ids)).all()} if inviter_ids else {}

    exam_invites = ExamScrutinizerInvite.query.filter_by(scrutinizer_teacher_id=teacher.id).order_by(ExamScrutinizerInvite.created_at.desc()).all()
    exam_entry_ids = [inv.exam_entry_id for inv in exam_invites]
    exam_entries_by_id = {e.id: e for e in ExamPaperEvaluation.query.filter(ExamPaperEvaluation.id.in_(exam_entry_ids)).all()} if exam_entry_ids else {}
    exam_inviter_ids = [inv.inviter_teacher_id for inv in exam_invites]
    exam_inviter_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(exam_inviter_ids)).all()} if exam_inviter_ids else {}

    split_invites = ClassSplitInvite.query.filter_by(invited_teacher_id=teacher.id).order_by(ClassSplitInvite.created_at.desc()).all()
    split_sessions_ids = {inv.inviter_session_id for inv in split_invites}
    split_sessions_by_id = {s.id: s for s in Session.query.filter(Session.id.in_(split_sessions_ids)).all()} if split_sessions_ids else {}
    split_inviter_ids = {inv.inviter_teacher_id for inv in split_invites}
    split_inviter_by_id = {t.id: t for t in Teacher.query.filter(Teacher.id.in_(split_inviter_ids)).all()} if split_inviter_ids else {}

    return render_template(
        'class_management/invitations.html',
        invites=invites,
        sessions_by_id=sessions_by_id,
        inviter_by_id=inviter_by_id,
        exam_invites=exam_invites,
        exam_entries_by_id=exam_entries_by_id,
        exam_inviter_by_id=exam_inviter_by_id,
        split_invites=split_invites,
        split_sessions_by_id=split_sessions_by_id,
        split_inviter_by_id=split_inviter_by_id,
        course_scope_labels=COURSE_SCOPE_LABELS
    )

@class_management_bp.route('/evaluation/<int:session_id>/course-assessment/open/<int:invite_id>', methods=['GET', 'POST'])
@login_required
def course_assessment_form(session_id, invite_id):
    """Invitee fills the class observation report form."""
    invite = EvaluationInvite.query.get_or_404(invite_id)
    if invite.session_id != session_id or invite.status == 'cancelled':
        flash('Invalid invitation.', 'danger')
        return redirect(url_for('index'))

    # Current user must be evaluator
    evaluator = Teacher.query.filter_by(name=current_user.full_name).first()
    if not evaluator or evaluator.id != invite.evaluator_teacher_id:
        flash('You are not authorized for this form.', 'danger')
        return redirect(url_for('index'))

    session = Session.query.get_or_404(session_id)
    current_teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not current_teacher or current_teacher.id not in {invite.inviter_teacher_id, invite.evaluator_teacher_id}:
        flash('You are not authorized to view this submission.', 'danger')
        return redirect(url_for('index'))

    submission = EvaluationSubmission.query.filter_by(invite_id=invite.id).first()
    general_data = {}
    score_data = {}
    section_totals = {}
    if submission:
        import json
        try:
            general_data = json.loads(submission.general_info or '{}')
            score_data = json.loads(submission.scores or '{}')
            section_totals = {
                'b1': sum(score_data.get(k, 0) for k in ['b1_a','b1_b','b1_c']),
                'b2': sum(score_data.get(k, 0) for k in ['b2_a','b2_b','b2_c','b2_d','b2_e','b2_f']),
                'b3': sum(score_data.get(k, 0) for k in ['b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h']),
                'b4': sum(score_data.get(k, 0) for k in ['b4_a','b4_b','b4_c'])
            }
            # Backward compatibility for older data
            if 'venue' not in general_data and general_data.get('venue_time'):
                general_data.setdefault('venue', general_data.get('venue_time'))
            general_data.setdefault('session_date', '')
            general_data.setdefault('session_time', '')
        except Exception:
            general_data = {}
            score_data = {}
            section_totals = {}

    general_data.setdefault('program_name', '')
    general_data['teacher_name'] = session.teacher.name if session.teacher else ''
    general_data.setdefault('observer_name', current_user.full_name)
    general_data.setdefault('course_name', session.course_name or '')
    general_data.setdefault('course_code', session.course_code or '')
    general_data.setdefault('course_year', session.year or '')
    general_data.setdefault('course_term', session.term or '')
    general_data.setdefault('academic_session', session.academic_session or '')
    general_data.setdefault('venue', '')
    general_data.setdefault('session_date', '')
    general_data.setdefault('session_time', '')

    if request.method == 'POST':
        try:
            import json
            general = {
                'program_name': request.form.get('program_name'),
                'teacher_name': session.teacher.name if session.teacher else '',
                'observer_name': request.form.get('observer_name'),
                'course_name': request.form.get('course_name') or session.course_name,
                'course_code': request.form.get('course_code') or session.course_code,
                'course_year': request.form.get('course_year') or session.year,
                'course_term': request.form.get('course_term') or session.term,
                'academic_session': request.form.get('academic_session') or session.academic_session,
                'session_date': request.form.get('session_date'),
                'session_time': request.form.get('session_time'),
                'venue': request.form.get('venue')
            }
            # Collect scores
            score_keys = [
                'b1_a','b1_b','b1_c',
                'b2_a','b2_b','b2_c','b2_d','b2_e','b2_f',
                'b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h',
                'b4_a','b4_b','b4_c'
            ]
            scores = {}
            total = 0
            for k in score_keys:
                v = request.form.get(k)
                if v:
                    scores[k] = int(v)
                    total += int(v)
            comments_observer = request.form.get('comments_observer')
            comments_presenter = submission.comments_presenter if submission else None

            if not submission:
                submission = EvaluationSubmission(
                    invite_id=invite.id,
                    session_id=session_id,
                    evaluator_teacher_id=evaluator.id
                )
                db.session.add(submission)

            submission.general_info = json.dumps(general)
            submission.scores = json.dumps(scores)
            submission.comments_observer = comments_observer
            submission.comments_presenter = comments_presenter
            submission.total_score = total
            invite.status = 'submitted'
            db.session.commit()
            flash('Assessment form submitted.', 'success')
            return redirect(url_for('class_management.my_invitations'))
        except Exception as e:
            db.session.rollback()
            flash(f'Error submitting form: {str(e)}', 'danger')

    current_date_str = date.today().isoformat()
    current_time_str = datetime.now().strftime('%H:%M')
    return render_template(
        'class_management/evaluation_course_assessment_form.html',
        session=session,
        invite=invite,
        submission=submission,
        general_data=general_data,
        score_data=score_data,
        section_totals=section_totals,
        current_date=current_date_str,
        current_time=current_time_str
    )


@class_management_bp.route('/evaluation/<int:session_id>/student-feedback', methods=['GET', 'POST'])
@login_required
def student_feedback_manage(session_id):
    """Manage anonymous student feedback link and view submissions."""
    session = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session.teacher_id:
        flash('শুধুমাত্র কোর্সটির শিক্ষকই ফিডব্যাক সেটআপ করতে পারবেন।', 'danger')
        return redirect(url_for('class_management.evaluation', session_id=session_id))

    feedback_link = (
        StudentFeedbackLink.query.filter_by(session_id=session_id)
        .order_by(StudentFeedbackLink.created_at.desc())
        .first()
    )

    if request.method == 'POST':
        action = request.form.get('action')
        expires_at = None
        expires_raw = request.form.get('expires_at') or ''
        if expires_raw:
            try:
                expires_at = datetime.strptime(expires_raw, '%Y-%m-%d')
            except ValueError:
                flash('ভ্যালিড মেয়াদ তারিখ দিন (YYYY-MM-DD).', 'warning')
                return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

        allow_multiple = bool(request.form.get('allow_multiple'))
        title = request.form.get('title') or f"{session.course_name or 'Course'} Feedback"
        description = request.form.get('description') or ''

        try:
            if action == 'create':
                if feedback_link:
                    flash('ফিডব্যাক লিংক ইতিমধ্যে বিদ্যমান।', 'info')
                else:
                    feedback_link = StudentFeedbackLink(
                        session_id=session_id,
                        access_code=_generate_feedback_code(),
                        title=title,
                        description=description,
                        expires_at=expires_at,
                        allow_multiple=allow_multiple,
                    )
                    db.session.add(feedback_link)
                    db.session.commit()
                    flash('ফিডব্যাক লিংক তৈরি হয়েছে।', 'success')
            elif action == 'update' and feedback_link:
                feedback_link.title = title
                feedback_link.description = description
                feedback_link.expires_at = expires_at
                feedback_link.allow_multiple = allow_multiple
                db.session.commit()
                flash('সেটিংস আপডেট হয়েছে।', 'success')
            elif action == 'regenerate' and feedback_link:
                feedback_link.access_code = _generate_feedback_code()
                db.session.commit()
                flash('নতুন অ্যাক্সেস কোড তৈরি হয়েছে।', 'info')
            elif action == 'delete' and feedback_link:
                db.session.delete(feedback_link)
                db.session.commit()
                flash('ফিডব্যাক লিংক ও সমস্ত উত্তর মুছে ফেলা হয়েছে।', 'info')
                return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))
            else:
                flash('অজানা অ্যাকশন।', 'danger')
        except Exception as exc:
            db.session.rollback()
            flash(f'ফিডব্যাক সেটআপ পরিবর্তন ব্যর্থ: {exc}', 'danger')

        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_responses = []
    if feedback_link:
        responses = (
            StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
            .order_by(StudentFeedbackResponse.submitted_at.desc())
            .all()
        )
        for item in responses:
            try:
                answers = json.loads(item.payload or '{}')
            except json.JSONDecodeError:
                answers = {}
            feedback_responses.append(
                {
                    'submitted_at': item.submitted_at,
                    'data': answers,
                }
            )

    feedback_url = (
        url_for('student_feedback_form', code=feedback_link.access_code, _external=True)
        if feedback_link
        else None
    )

    return render_template(
        'class_management/student_feedback_manage.html',
        session=session,
        feedback_link=feedback_link,
        feedback_url=feedback_url,
        feedback_responses=feedback_responses,
        section_a_labels=FEEDBACK_SECTION_A,
        section_b_likert=FEEDBACK_SECTION_B_LIKERT,
        method_options=FEEDBACK_METHOD_OPTIONS,
        effort_options=FEEDBACK_EFFORT_OPTIONS,
    )


@class_management_bp.route('/evaluation/<int:session_id>/course-review', methods=['GET', 'POST'])
@login_required
def course_review_form(session_id):
    """Course teacher documents reflections and improvement plans."""
    session = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    if not teacher:
        flash('You must be registered as a teacher to access the course review form.', 'danger')
        return redirect(url_for('class_management.index'))

    if session.teacher_id != teacher.id:
        flash('Only the course teacher can access the course review form.', 'danger')
        return redirect(url_for('class_management.evaluation', session_id=session_id))

    review = CourseReview.query.filter_by(session_id=session_id, teacher_id=teacher.id).first()
    review_data = {}
    if review and review.data:
        try:
            review_data = json.loads(review.data)
        except Exception:
            review_data = {}
    if not isinstance(review_data, dict):
        review_data = {}

    if 'discipline' not in review_data and review_data.get('department'):
        review_data['discipline'] = review_data.pop('department')
    if 'school' not in review_data and review_data.get('faculty'):
        review_data['school'] = review_data.pop('faculty')
    if 'course_term' not in review_data and review_data.get('semester_term'):
        review_data['course_term'] = review_data.pop('semester_term')
    if 'course_year' not in review_data and review_data.get('level'):
        review_data['course_year'] = review_data.pop('level')

    review_data.setdefault('discipline', '')
    review_data.setdefault('school', '')
    review_data.setdefault('course_title', session.course_name or '')
    review_data.setdefault('course_code', session.course_code or '')
    review_data.setdefault('session_name', session.academic_session or session.year or '')
    review_data.setdefault('course_year', session.year or '')
    review_data.setdefault('course_term', session.term or '')
    review_data.setdefault('credit_value', '')
    review_data.setdefault('instructor_name', session.teacher.name if session.teacher else '')
    review_data.setdefault('enrollment_count', '')
    for row in COURSE_REVIEW_GRADE_ROWS:
        review_data.setdefault(f"{row['key']}_number", '')
        review_data.setdefault(f"{row['key']}_percentage", '')
    review_data.setdefault('grade_total_number', '')
    review_data.setdefault('grade_total_percentage', '')
    if request.method == 'POST':
        try:
            def _clean(field):
                return (request.form.get(field) or '').strip()
            def _to_number(value):
                try:
                    if value is None or value == '':
                        return None
                    return float(value)
                except (TypeError, ValueError):
                    return None

            def _format_number(value, decimals=None):
                if value is None:
                    return ''
                if decimals is not None:
                    return f"{value:.{decimals}f}"
                if float(value).is_integer():
                    return str(int(round(value)))
                return f"{value:.2f}".rstrip('0').rstrip('.')

            form_data = {
                'discipline': _clean('discipline'),
                'school': _clean('school'),
                'course_code': _clean('course_code'),
                'course_title': _clean('course_title'),
                'session_name': _clean('session_name'),
                'course_year': _clean('course_year'),
                'course_term': _clean('course_term'),
                'credit_value': _clean('credit_value'),
                'instructor_name': _clean('instructor_name'),
                'contact_hours': _clean('contact_hours'),
                'lecture_hours': _clean('lecture_hours'),
                'seminar_hours': _clean('seminar_hours'),
                'other_instruction': _clean('other_instruction'),
                'assessment_methods': _clean('assessment_methods'),
                'enrollment_count': _clean('enrollment_count'),
            }

            enrollment_value = _to_number(form_data['enrollment_count'])
            total_number = 0.0
            has_grade_values = False

            for row in COURSE_REVIEW_GRADE_ROWS:
                num_key = f"{row['key']}_number"
                pct_key = f"{row['key']}_percentage"
                number_value = _to_number(_clean(num_key))
                if number_value is not None:
                    has_grade_values = True
                    total_number += number_value
                    form_data[num_key] = _format_number(number_value)
                    if enrollment_value and enrollment_value > 0:
                        percentage_value = (number_value / enrollment_value) * 100
                        form_data[pct_key] = _format_number(percentage_value, decimals=2)
                    else:
                        form_data[pct_key] = ''
                else:
                    form_data[num_key] = ''
                    form_data[pct_key] = ''

            if has_grade_values:
                form_data['grade_total_number'] = _format_number(total_number)
                if enrollment_value and enrollment_value > 0:
                    total_percentage = (total_number / enrollment_value) * 100
                    form_data['grade_total_percentage'] = _format_number(total_percentage, decimals=2)
                else:
                    form_data['grade_total_percentage'] = ''
            else:
                form_data['grade_total_number'] = ''
                form_data['grade_total_percentage'] = ''

            for item in COURSE_REVIEW_COMMENT_FIELDS:
                form_data[item['key']] = _clean(item['key'])

            if review is None:
                review = CourseReview(session_id=session_id, teacher_id=teacher.id)
                db.session.add(review)

            review.data = json.dumps(form_data)
            db.session.commit()
            flash('Course review saved successfully.', 'success')
            return redirect(url_for('class_management.course_review_form', session_id=session_id))
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Error saving course review for session {session_id}: {e}")
            flash(f'Error saving course review: {str(e)}', 'danger')

    has_saved_review = bool(review and review.data)

    return render_template(
        'class_management/evaluation_course_review_form.html',
        session=session,
        review_data=review_data,
        grade_rows=COURSE_REVIEW_GRADE_ROWS,
        comment_fields=COURSE_REVIEW_COMMENT_FIELDS,
        has_saved_review=has_saved_review
    )


@class_management_bp.route('/evaluation/<int:session_id>/course-review/pdf')
@login_required
def course_review_pdf(session_id):
    """Download the Faculty Course Review Report as a PDF."""
    session = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()

    if not teacher or teacher.id != session.teacher_id:
        flash('Only the course teacher can download the course review PDF.', 'danger')
        return redirect(url_for('class_management.evaluation', session_id=session_id))

    review = CourseReview.query.filter_by(session_id=session_id, teacher_id=teacher.id).first()
    if not review or not review.data:
        flash('No saved course review found to generate PDF.', 'warning')
        return redirect(url_for('class_management.course_review_form', session_id=session_id))

    try:
        stored_data = json.loads(review.data)
    except Exception:
        flash('Stored course review data is invalid.', 'danger')
        return redirect(url_for('class_management.course_review_form', session_id=session_id))

    from xml.sax.saxutils import escape

    def get_value(key):
        return escape(str(stored_data.get(key, '') or ''))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name='CourseReviewTitle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        fontSize=14,
        leading=16,
        spaceAfter=12
    )
    subtitle_style = ParagraphStyle(
        name='CourseReviewSubtitle',
        parent=styles['BodyText'],
        alignment=TA_CENTER,
        fontSize=9,
        leading=12,
        spaceAfter=18
    )
    label_style = ParagraphStyle(
        name='CourseReviewLabel',
        parent=styles['BodyText'],
        fontSize=9,
        leading=11,
        spaceAfter=4,
        spaceBefore=2
    )
    comment_style = ParagraphStyle(
        name='CourseReviewComment',
        parent=styles['BodyText'],
        fontSize=9,
        leading=12
    )

    elements = []
    elements.append(Paragraph('FACULTY COURSE REVIEW REPORT', title_style))
    elements.append(Paragraph('(To be filled by each teacher at the time of Course Completion)', subtitle_style))

    table_width = doc.width

    other_paragraph = Paragraph(get_value('other_instruction') or '&nbsp;', comment_style)
    assessment_paragraph = Paragraph(get_value('assessment_methods') or '&nbsp;', comment_style)

    enrollment_paragraph = Paragraph(
        f'Number of Enrolled Students: {get_value("enrollment_count") or "&nbsp;"}',
        label_style
    )

    info_col_widths = [table_width * 0.2, table_width * 0.3, table_width * 0.2, table_width * 0.3]
    info_data = [
        [Paragraph('Discipline', label_style), Paragraph(get_value('discipline'), comment_style), Paragraph('School', label_style), Paragraph(get_value('school'), comment_style)],
        [Paragraph('Course Code', label_style), Paragraph(get_value('course_code'), comment_style), Paragraph('Title', label_style), Paragraph(get_value('course_title'), comment_style)],
        [Paragraph('Session', label_style), Paragraph(get_value('session_name'), comment_style), Paragraph('Year', label_style), Paragraph(get_value('course_year'), comment_style)],
        [Paragraph('Term', label_style), Paragraph(get_value('course_term'), comment_style), Paragraph('Credit Value', label_style), Paragraph(get_value('credit_value'), comment_style)],
        [Paragraph('Name of Course Instructor', label_style), Paragraph(get_value('instructor_name'), comment_style), Paragraph('No. of Students Contact Hour', label_style), Paragraph(get_value('contact_hours'), comment_style)],
        [Paragraph('Lectures', label_style), Paragraph(get_value('lecture_hours'), comment_style), Paragraph('Seminar', label_style), Paragraph(get_value('seminar_hours'), comment_style)],
        [Paragraph('Number of Enrolled Students', label_style), Paragraph(get_value('enrollment_count'), comment_style), '', ''],
        [Paragraph('Other (Please State)', label_style), other_paragraph, '', ''],
        [Paragraph('Assessment Methods: give precise details (no & length of assignments, exams, weightings etc)', label_style), assessment_paragraph, '', '']
    ]

    info_table = Table(info_data, colWidths=info_col_widths, hAlign='CENTER', repeatRows=1)
    info_table.setStyle(TableStyle([
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.whitesmoke),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('SPAN', (0, 6), (1, 6)),
        ('SPAN', (0, 7), (1, 7)),
        ('SPAN', (0, 8), (1, 8)),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 16))

    grade_table_data = [['Scale', 'Letter Grade', 'Number of Students', '%']]
    for row in COURSE_REVIEW_GRADE_ROWS:
        grade_table_data.append([
            row['scale'],
            row['letter'],
            get_value(f"{row['key']}_number"),
            get_value(f"{row['key']}_percentage"),
        ])
    grade_table_data.append([
        'Total', '',
        get_value('grade_total_number'),
        get_value('grade_total_percentage'),
    ])

    grade_col_widths = [
        table_width * 0.36,
        table_width * 0.14,
        table_width * 0.25,
        table_width * 0.25,
    ]
    grade_table = Table(grade_table_data, colWidths=grade_col_widths, hAlign='CENTER')
    grade_table.setStyle(TableStyle([
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('BACKGROUND', (0, -1), (-1, -1), colors.whitesmoke),
        ('ALIGN', (1, 1), (-1, -2), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))
    elements.append(grade_table)
    elements.append(Spacer(1, 18))

    elements.append(Paragraph('Overview/Evaluation (Course Coordinator\'s Comments)', label_style))

    comment_table_data = []
    for item in COURSE_REVIEW_COMMENT_FIELDS:
        comment_table_data.append([
            Paragraph(item['label'], label_style),
            Paragraph(get_value(item['key']) or '&nbsp;', comment_style)
        ])

    comment_table = Table(comment_table_data, colWidths=[table_width / 2, table_width / 2], hAlign='CENTER')
    comment_table.setStyle(TableStyle([
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BACKGROUND', (0, 0), (-1, -1), colors.whitesmoke),
        ('BACKGROUND', (1, 0), (1, -1), colors.white),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('WORDWRAP', (0, 0), (-1, -1), True),
    ]))
    elements.append(comment_table)
    elements.append(Spacer(1, 20))

    signature_width = table_width / 2
    signature_table_data = [
        [
            Paragraph('Head of the Discipline Signature', label_style),
            Paragraph('Course Instructor (s) Signature', label_style)
        ],
        [
            Paragraph('&nbsp;' * 4, comment_style),
            Paragraph('&nbsp;' * 4, comment_style)
        ],
        [
            Paragraph('Date: ___________________', label_style),
            Paragraph('Date: ___________________', label_style)
        ]
    ]

    signature_table = Table(
        signature_table_data,
        colWidths=[signature_width, signature_width],
        hAlign='CENTER'
    )
    signature_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('LEFTPADDING', (0, 0), (-1, -1), 24),
        ('RIGHTPADDING', (0, 0), (-1, -1), 24),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LINEABOVE', (0, 1), (0, 1), 0.8, colors.black),
        ('LINEABOVE', (1, 1), (1, 1), 0.8, colors.black),
        ('BOTTOMPADDING', (0, 1), (1, 1), 24),
    ]))

    elements.append(signature_table)

    def _add_page_number(canvas_obj, doc_obj):
        canvas_obj.saveState()
        page_num = canvas_obj.getPageNumber()
        text = f"Page {page_num}"
        canvas_obj.setFont('Helvetica', 9)
        canvas_obj.drawRightString(
            doc_obj.pagesize[0] - doc_obj.rightMargin,
            doc_obj.bottomMargin - 20,
            text
        )
        canvas_obj.restoreState()

    doc.build(elements, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    buffer.seek(0)

    filename = f"course_review_{session.course_code or session.id}.pdf"
    return Response(
        buffer.getvalue(),
        mimetype='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(buffer.getvalue()))
        }
    )


@class_management_bp.route('/evaluation/<int:session_id>/course-assessment/view/<int:invite_id>', methods=['GET', 'POST'])
@login_required
def course_assessment_view(session_id, invite_id):
    """Inviter views a submission and can download PDF."""
    invite = EvaluationInvite.query.get_or_404(invite_id)
    if invite.session_id != session_id:
        flash('Invalid invitation.', 'danger')
        return redirect(url_for('index'))
    session = Session.query.get_or_404(session_id)
    current_teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not current_teacher or current_teacher.id not in {invite.inviter_teacher_id, invite.evaluator_teacher_id}:
        flash('You are not authorized to view this submission.', 'danger')
        return redirect(url_for('index'))

    submission = EvaluationSubmission.query.filter_by(invite_id=invite.id).first()
    if not submission:
        flash('No submission yet.', 'info')
        return redirect(url_for('class_management.course_assessment', session_id=session_id))
    import json
    score_data = json.loads(submission.scores or '{}')
    general_data = json.loads(submission.general_info or '{}')
    if 'venue' not in general_data and general_data.get('venue_time'):
        general_data.setdefault('venue', general_data.get('venue_time'))
    general_data.setdefault('session_date', '')
    general_data.setdefault('session_time', '')
    general_data.setdefault('program_name', '')
    general_data['teacher_name'] = session.teacher.name if session.teacher else ''
    general_data.setdefault('observer_name', '')
    general_data.setdefault('course_name', session.course_name or '')
    general_data.setdefault('course_code', session.course_code or '')
    general_data.setdefault('course_year', session.year or '')
    general_data.setdefault('course_term', session.term or '')
    general_data.setdefault('academic_session', session.academic_session or '')
    section_totals = {
        'b1': sum(score_data.get(k, 0) for k in ['b1_a','b1_b','b1_c']),
        'b2': sum(score_data.get(k, 0) for k in ['b2_a','b2_b','b2_c','b2_d','b2_e','b2_f']),
        'b3': sum(score_data.get(k, 0) for k in ['b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h']),
        'b4': sum(score_data.get(k, 0) for k in ['b4_a','b4_b','b4_c'])
    }
    can_edit_presenter = current_teacher and current_teacher.id == session.teacher_id

    if request.method == 'POST':
        if not can_edit_presenter:
            flash('You are not authorized to update presenter comments.', 'danger')
            return redirect(url_for('class_management.course_assessment_view', session_id=session_id, invite_id=invite_id))
        submission.comments_presenter = request.form.get('comments_presenter')
        try:
            db.session.commit()
            flash("Presenter's comments updated.", 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating comments: {str(e)}', 'danger')
        return redirect(url_for('class_management.course_assessment_view', session_id=session_id, invite_id=invite_id))

    return render_template('class_management/evaluation_course_assessment_view.html', session=session, invite=invite, submission=submission, general_data=general_data, score_data=score_data, section_totals=section_totals, can_edit_presenter=can_edit_presenter)

@class_management_bp.route('/evaluation/<int:session_id>/course-assessment/pdf/<int:invite_id>')
@login_required
def course_assessment_pdf(session_id, invite_id):
    """Generate PDF of the submitted assessment."""
    try:
        import json
        from xml.sax.saxutils import escape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        invite = EvaluationInvite.query.get_or_404(invite_id)
        if invite.session_id != session_id:
            flash('Invalid invitation.', 'danger')
            return redirect(url_for('index'))
        session = Session.query.get_or_404(session_id)
        current_teacher = Teacher.query.filter_by(name=current_user.full_name).first()
        if not current_teacher or current_teacher.id not in {invite.inviter_teacher_id, invite.evaluator_teacher_id}:
            flash('You are not authorized to download this report.', 'danger')
            return redirect(url_for('index'))

        submission = EvaluationSubmission.query.filter_by(invite_id=invite.id).first()
        if not submission:
            flash('No submission to export.', 'warning')
            return redirect(url_for('class_management.course_assessment', session_id=session_id))

        general = json.loads(submission.general_info or '{}')
        if 'venue' not in general and general.get('venue_time'):
            general.setdefault('venue', general.get('venue_time'))
        general.setdefault('session_date', '')
        general.setdefault('session_time', '')
        general.setdefault('program_name', '')
        general.setdefault('teacher_name', session.teacher.name if session.teacher else '')
        general.setdefault('observer_name', '')
        general.setdefault('course_name', session.course_name or '')
        general.setdefault('course_code', session.course_code or '')
        general.setdefault('course_year', session.year or '')
        general.setdefault('course_term', session.term or '')
        general.setdefault('academic_session', session.academic_session or '')
        scores = json.loads(submission.scores or '{}')
        section_totals = {
            'b1': sum(scores.get(k, 0) for k in ['b1_a','b1_b','b1_c']),
            'b2': sum(scores.get(k, 0) for k in ['b2_a','b2_b','b2_c','b2_d','b2_e','b2_f']),
            'b3': sum(scores.get(k, 0) for k in ['b3_a','b3_b','b3_c','b3_d','b3_e','b3_f','b3_g','b3_h']),
            'b4': sum(scores.get(k, 0) for k in ['b4_a','b4_b','b4_c'])
        }

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('ReportTitle', parent=styles['Title'], fontSize=16, spaceAfter=12)
        label_style = ParagraphStyle('Label', parent=styles['BodyText'], fontName='Helvetica-Bold', fontSize=10)
        value_style = ParagraphStyle('Value', parent=styles['BodyText'], fontSize=10)
        center_value_style = ParagraphStyle('CenterValue', parent=value_style, alignment=1)
        section_header_style = ParagraphStyle('SectionHeader', parent=styles['BodyText'], fontName='Helvetica-Bold', fontSize=10)
        subtotal_style = ParagraphStyle('Subtotal', parent=styles['BodyText'], fontName='Helvetica-Oblique', fontSize=10)
        comment_header_style = ParagraphStyle('CommentHeader', parent=styles['Heading4'], fontSize=11, spaceBefore=12, spaceAfter=6)

        elements = []
        elements.append(Paragraph('Classroom Teaching Observation Report', title_style))

        info_data = [
            [Paragraph('Program', label_style), Paragraph(escape(general.get('program_name') or '-') , value_style),
             Paragraph('Teacher', label_style), Paragraph(escape(general.get('teacher_name') or '-') , value_style)],
            [Paragraph('Observer', label_style), Paragraph(escape(general.get('observer_name') or '-') , value_style),
             Paragraph('Course Name', label_style), Paragraph(escape(general.get('course_name') or '-') , value_style)],
            [Paragraph('Course Code', label_style), Paragraph(escape(general.get('course_code') or '-') , value_style),
             Paragraph('Academic Session', label_style), Paragraph(escape(general.get('academic_session') or '-') , value_style)],
            [Paragraph('Year', label_style), Paragraph(escape(general.get('course_year') or '-') , value_style),
             Paragraph('Term', label_style), Paragraph(escape(general.get('course_term') or '-') , value_style)],
            [Paragraph('Date', label_style), Paragraph(escape(general.get('session_date') or '-') , value_style),
             Paragraph('Time', label_style), Paragraph(escape(general.get('session_time') or '-') , value_style)],
            [Paragraph('Venue', label_style), Paragraph(escape(general.get('venue') or '-') , value_style),
             Paragraph('', label_style), Paragraph('', value_style)]
        ]
        info_table = Table(info_data, colWidths=[80, 180, 80, 180])
        info_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f7f7f7')),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('SPAN', (2,5), (3,5)),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 10))
        elements.append(Paragraph('Score Scale: 5 = Excellent, 4 = Very Good, 3 = Good, 2 = Fair, 1 = Poor', styles['BodyText']))
        elements.append(Spacer(1, 8))

        score_rows = [[Paragraph('<b>Description</b>', label_style), Paragraph('<b>Score</b>', center_value_style)]]
        section_headers = []
        span_rows = []
        subtotal_rows = []

        def add_section(title, items, subtotal_key):
            header_idx = len(score_rows)
            score_rows.append([Paragraph(f'<b>{escape(title)}</b>', section_header_style), ''])
            section_headers.append(header_idx)
            span_rows.append(header_idx)
            for text, key in items:
                score_rows.append([
                    Paragraph(escape(text), value_style),
                    Paragraph(str(scores.get(key)) if key in scores else '-', center_value_style)
                ])
            subtotal_idx = len(score_rows)
            score_rows.append([
                Paragraph(f'<i>Subtotal for {escape(title)}</i>', subtotal_style),
                Paragraph(f"<b>{section_totals.get(subtotal_key, 0)}</b>", center_value_style)
            ])
            subtotal_rows.append(subtotal_idx)

        add_section('Section 1: Set Induction', [
            ('a) Clarity of objectives', 'b1_a'),
            ('b) Relevance to topic', 'b1_b'),
            ('c) Appropriateness of introduction', 'b1_c')
        ], 'b1')

        add_section('Section 2: Content', [
            ('a) Knowledge', 'b2_a'),
            ('b) Extend of coverage', 'b2_b'),
            ('c) Level of interest generated', 'b2_c'),
            ('d) Logical flow of presentation', 'b2_d'),
            ('e) Correctness of language used', 'b2_e'),
            ('f) Clear and relevant use of analogies/examples', 'b2_f')
        ], 'b2')

        add_section('Section 3: Presentation', [
            ('a) Appropriate pacing', 'b3_a'),
            ('b) Confidence', 'b3_b'),
            ('c) Enthusiasm', 'b3_c'),
            ('d) Provoking students to think', 'b3_d'),
            ('e) Clarity of presentation', 'b3_e'),
            ('f) Interaction with students', 'b3_f'),
            ('g) Effective use of teaching/learning aids', 'b3_g'),
            ('h) Effective class management', 'b3_h')
        ], 'b3')

        add_section('Section 4: Closure', [
            ('a) Appropriateness of closure', 'b4_a'),
            ('b) Effective questions for feedback', 'b4_b'),
            ('c) Appropriate links to the next lesson', 'b4_c')
        ], 'b4')

        score_rows.append([
            Paragraph('<b>Total Score</b>', label_style),
            Paragraph(f"<b>{submission.total_score or 0}</b>", center_value_style)
        ])

        score_table = Table(score_rows, colWidths=[360, 80])
        score_style = TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0f0')),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (1,0), (1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE')
        ])
        for idx in section_headers:
            score_style.add('SPAN', (0, idx), (1, idx))
            score_style.add('BACKGROUND', (0, idx), (-1, idx), colors.HexColor('#e5e5e5'))
            score_style.add('FONTNAME', (0, idx), (-1, idx), 'Helvetica-Bold')
        for idx in subtotal_rows:
            score_style.add('BACKGROUND', (0, idx), (-1, idx), colors.HexColor('#f9f9f9'))
        score_style.add('BACKGROUND', (0, len(score_rows)-1), (-1, len(score_rows)-1), colors.HexColor('#e8f4ff'))
        score_table.setStyle(score_style)
        elements.append(score_table)

        elements.append(Spacer(1, 12))

        interpret_data = [
            [Paragraph('<b>Score</b>', label_style), Paragraph('<b>Interpretation</b>', label_style)],
            ['90 to 100', 'Excellent'],
            ['80 to less than 90', 'Very Good'],
            ['70 to less than 80', 'Good'],
            ['60 to less than 70', 'Fair'],
            ['50 to less than 60', 'Poor'],
            ['40 to less than 50', 'Very Poor']
        ]
        interpret_table = Table(interpret_data, colWidths=[140, 200])
        interpret_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0f0')),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER')
        ]))
        elements.append(interpret_table)

        observer_comment = (submission.comments_observer or '').strip()
        presenter_comment = (submission.comments_presenter or '').strip()
        elements.append(Paragraph("Observer's Overall Comments and Suggestions for Improvement", comment_header_style))
        elements.append(Paragraph(escape(observer_comment).replace('\n', '<br/>') or '-', styles['BodyText']))
        elements.append(Paragraph("Presenter's Comments", comment_header_style))
        elements.append(Paragraph(escape(presenter_comment).replace('\n', '<br/>') or '-', styles['BodyText']))

        elements.append(Spacer(1, 60))
        line_width = (A4[0] - 72) / 2
        signature_table = Table(
            [['', ''], ["Presenter's Signature", "Observer's Signature"]],
            colWidths=[line_width, line_width]
        )
        signature_table.setStyle(TableStyle([
            ('LINEABOVE', (0,0), (0,0), 0.7, colors.black),
            ('LINEABOVE', (1,0), (1,0), 0.7, colors.black),
            ('TOPPADDING', (0,0), (-1,0), 30),
            ('ALIGN', (0,1), (-1,1), 'CENTER'),
            ('TOPPADDING', (0,1), (-1,1), 8)
        ]))
        elements.append(signature_table)

        class NumberedCanvas(canvas.Canvas):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self._saved_page_states = []

            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                num_pages = len(self._saved_page_states)
                for state in self._saved_page_states:
                    self.__dict__.update(state)
                    self.draw_page_number(num_pages)
                    super().showPage()
                super().save()

            def draw_page_number(self, page_count):
                self.setFont('Helvetica', 9)
                text = f"Page {self._pageNumber} of {page_count}"
                self.drawRightString(A4[0] - 36, 30, text)

        doc.build(elements, canvasmaker=NumberedCanvas)
        buffer.seek(0)
        filename = f"course_assessment_{session.course_code or session.id}.pdf"
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
    except Exception as e:
        flash(f'Error generating PDF: {str(e)}', 'danger')
        return redirect(url_for('class_management.course_assessment', session_id=session_id))

@class_management_bp.route('/invitation/<int:invite_id>/cancel', methods=['POST'])
@login_required
def cancel_invitation(invite_id):
    """Allow either inviter or invitee to cancel an invitation."""
    invite = EvaluationInvite.query.get_or_404(invite_id)

    # Resolve current teacher for the logged-in user
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher:
        flash('No teacher profile found.', 'warning')
        return redirect(url_for('index'))

    if teacher.id not in [invite.inviter_teacher_id, invite.evaluator_teacher_id]:
        flash('You are not authorized to cancel this invitation.', 'danger')
        return redirect(url_for('index'))

    try:
        invite.status = 'cancelled'
        EvaluationSubmission.query.filter_by(invite_id=invite.id).delete()
        db.session.commit()
        flash('Invitation has been cancelled.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error cancelling invitation: {str(e)}', 'danger')

    # Redirect back to sensible page
    if teacher.id == invite.inviter_teacher_id:
        return redirect(url_for('class_management.course_assessment', session_id=invite.session_id))
    else:
        return redirect(url_for('class_management.my_invitations'))

FEEDBACK_SECTION_A = [
    ('course_structure', 'Course content is structured in a comprehensible manner.'),
    ('course_goals', 'The goals of the course are clear.'),
    ('course_content_guidance', 'The course contents are explained in an understandable fashion.'),
    ('course_interest', 'The course fosters my interest in the discussed topics.'),
]

FEEDBACK_SECTION_B_LIKERT = [
    ('course_plan_discussed', 'Course plan (assessment criteria/ content) was discussed in advance.'),
    ('guidelines_received', 'Received oral instruction and written guidelines for continuous assessment.'),
    ('assessment_helpful', 'Course projects/assignments/tests were helpful to demonstrate an understanding of the course material.'),
    ('feedback_timely', 'Received feedback and grades of continuous assessments from course teacher in due time.'),
]

FEEDBACK_METHOD_OPTIONS = [
    'Lectures (including online lectures)',
    'Class discussions (including online discussion boards)',
    'In-class learning activities (other than discussion)',
    'In-class clickers or other quick response methods',
    'Homework (readings and assignments)',
    'Labs',
    'Projects or portfolios',
    'Teamwork or group activities',
    'Student presentations',
    'Guest lecturers',
    'Fieldwork/field trips',
    'Mentoring outside of the classroom',
    'Support from Teaching/ Research Assistants',
    'Others'
]

FEEDBACK_EFFORT_OPTIONS = [
    'Memorizing facts and repeating ideas from the readings and lectures.',
    'Making judgments about the value of information, arguments, or methods.',
    'Applying basic elements of an idea, experience, or theory.',
    'Applying theories or concepts to practical problems or in new situations.',
    'Synthesizing and organizing ideas, information, or experiences.',
    'Solving problems.',
    'Thinking creatively or critically.',
    'Teamwork or group activities.',
    'Doing lab work.',
    'Presenting in person or via a recording.',
    'Reading and writing for deep understanding.',
    'Others'
]

@class_management_bp.route('/evaluation/<int:session_id>/student-feedback/responses/pdf')
@login_required
def student_feedback_responses_pdf(session_id):
    session_obj = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session_obj.teacher_id:
        flash('You are not authorized to access this download.', 'danger')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_link = StudentFeedbackLink.query.filter_by(session_id=session_id).first()
    if not feedback_link:
        flash('No feedback responses found.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    responses = (
        StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
        .order_by(StudentFeedbackResponse.submitted_at.asc())
        .all()
    )
    if not responses:
        flash('No feedback responses to download.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    buffer = io.BytesIO()
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    response_font = 'Helvetica'
    response_bold_font = 'Helvetica-Bold'
    kalpurush_available = False
    try:
        font_root = os.path.join(current_app.root_path, 'static', 'fonts')
        regular_candidates = ['Kalpurush.ttf', 'Kalpurush-Regular.ttf']
        bold_candidates = ['Kalpurush-Bold.ttf', 'Kalpurush Bold.ttf']
        regular_path = next((os.path.join(font_root, f) for f in regular_candidates if os.path.exists(os.path.join(font_root, f))), None)
        bold_path = next((os.path.join(font_root, f) for f in bold_candidates if os.path.exists(os.path.join(font_root, f))), None)
        if regular_path:
            pdfmetrics.registerFont(TTFont('Kalpurush', regular_path))
            response_font = 'Kalpurush'
            kalpurush_available = True
        if bold_path:
            pdfmetrics.registerFont(TTFont('Kalpurush-Bold', bold_path))
            response_bold_font = 'Kalpurush-Bold'
        elif regular_path:
            response_bold_font = 'Kalpurush'
    except Exception as exc:  # pragma: no cover
        current_app.logger.warning('Kalpurush font registration failed: %s', exc)

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=28,
        bottomMargin=28,
        leftMargin=32,
        rightMargin=32,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], alignment=1, fontSize=16, leading=18, spaceAfter=6, textTransform='uppercase', fontName='Helvetica-Bold')
    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Heading2'], alignment=1, fontSize=13, leading=16, spaceAfter=12, textTransform='uppercase', fontName='Helvetica-Bold')
    section_header_style = ParagraphStyle('SectionHeader', parent=styles['Heading3'], fontSize=11, leading=13, spaceBefore=8, spaceAfter=4, textTransform='uppercase', fontName='Helvetica-Bold')
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontSize=9, leading=11, fontName='Helvetica-Bold', wordWrap='CJK')
    instruction_style = ParagraphStyle('Instruction', parent=styles['Normal'], fontSize=9, leading=11, alignment=1, textTransform='uppercase', spaceBefore=6, spaceAfter=6, wordWrap='CJK', fontName='Helvetica-Bold')
    value_style = ParagraphStyle('Value', parent=styles['Normal'], fontSize=9.5, leading=11, fontName=response_font, wordWrap='CJK')
    value_bold_style = ParagraphStyle('ValueBold', parent=value_style, fontName=response_bold_font)
    # Style for Praise and Suggestions section - always use Kalpurush if available
    praise_suggestions_font = 'Kalpurush' if kalpurush_available else response_font
    praise_suggestions_style = ParagraphStyle('PraiseSuggestions', parent=styles['Normal'], fontSize=9.5, leading=11, fontName=praise_suggestions_font, wordWrap='CJK')

    likert_header = ['Statement', 'Strongly disagree', 'Disagree', 'Neutral', 'Agree', 'Strongly Agree']

    def likert_table(section_values, labels):
        data = [[Paragraph(text, label_style) for text in likert_header]]
        for key, question in labels:
            selected = (section_values or {}).get(key)
            row = [Paragraph(question, value_style)]
            for option in likert_header[1:]:
                mark = '✓' if selected and selected.lower() == option.lower() else ''
                row.append(Paragraph(mark, value_style))
            data.append(row)
        table = Table(data, colWidths=[240, 52, 52, 52, 52, 52])
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('ALIGN', (1, 1), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEADING', (0, 0), (-1, -1), 11),
        ]))
        return table

    def checklist_table(options, selected):
        rows = []
        selected = selected or []
        for idx in range(0, len(options), 2):
            row = []
            for offset in (0, 1):
                if idx + offset < len(options):
                    option = options[idx + offset]
                    mark = '✓' if option in selected else ''
                    row.append(Paragraph(f"{mark} {option}", value_style))
                else:
                    row.append('')
            rows.append(row)
        table = Table(rows, colWidths=[255, 255])
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        return table

    elements = []
    for idx, item in enumerate(responses, start=1):
        try:
            data = json.loads(item.payload or '{}')
        except json.JSONDecodeError:
            data = {}

        academic = data.get('academic_info', {}) or {}
        section_a = data.get('section_a', {}) or {}
        section_b = data.get('section_b', {}) or {}
        section_c = data.get('section_c', {}) or {}
        section_d = data.get('section_d', {}) or {}

        if idx > 1:
            elements.append(PageBreak())

        elements.append(Paragraph('Student Feedback Form', title_style))
        elements.append(Paragraph('Khulna University', subtitle_style))
        elements.append(Paragraph(f"Response {idx} - {item.submitted_at.strftime('%Y-%m-%d %H:%M')}", value_bold_style))
        elements.append(Spacer(1, 6))

        info_data = [
            [Paragraph('Academic Session', label_style), Paragraph(academic.get('academic_session') or '—', value_style)],
            [Paragraph('Title of the Course', label_style), Paragraph(academic.get('course_title') or session_obj.course_name or '—', value_style)],
            [Paragraph('Course Code', label_style), Paragraph(academic.get('course_code') or session_obj.course_code or '—', value_style)],
        ]
        info_table = Table(info_data, colWidths=[140, 360])
        info_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (-1, -1), colors.whitesmoke),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
        ]))
        elements.append(info_table)
        elements.append(Paragraph('Please tick/cross in the blank space which best describes how much you agree with the following statements', instruction_style))

        elements.append(Paragraph('A. Satisfaction with the Course', section_header_style))
        elements.append(likert_table(section_a, FEEDBACK_SECTION_A))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('B. Teaching-Learning Methods', section_header_style))
        elements.append(likert_table(section_b, [FEEDBACK_SECTION_B_LIKERT[0]]))
        elements.append(Spacer(1, 4))

        methods = section_b.get('teaching_methods') or []
        if methods:
            elements.append(Paragraph('Teaching methods that contributed significantly (selected):', value_bold_style))
            elements.append(checklist_table(FEEDBACK_METHOD_OPTIONS, methods))
            elements.append(Spacer(1, 4))

        elements.append(likert_table(section_b, FEEDBACK_SECTION_B_LIKERT[1:]))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('C. Engagement and Workload', section_header_style))
        engagement_table = Table(
            [
                [Paragraph('How much time do you devote to this course before and after each lecture?', label_style), Paragraph(section_c.get('study_time') or '—', value_style)],
                [Paragraph('About what percent of the class meetings (including discussions) did you attend?', label_style), Paragraph(section_c.get('attendance_percent') or '—', value_style)],
            ],
            colWidths=[360, 140]
        )
        engagement_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(engagement_table)
        elements.append(Spacer(1, 4))

        efforts = section_c.get('effort_focus') or []
        elements.append(Paragraph('Significant aspects of your effort (selected):', value_bold_style))
        elements.append(checklist_table(FEEDBACK_EFFORT_OPTIONS, efforts))
        elements.append(Spacer(1, 6))

        elements.append(Paragraph('D. Praise and Suggestions', section_header_style))
        def open_block(title, content):
            # Use Kalpurush font for Praise and Suggestions content
            content_style = praise_suggestions_style
            table = Table(
                [
                    [Paragraph(title, label_style)],
                    [Paragraph(content or '—', content_style)]
                ],
                colWidths=[500]
            )
            table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 12)
            ]))
            return table

        elements.append(open_block('What did you like especially about this course?', section_d.get('likes')))
        elements.append(Spacer(1, 4))
        elements.append(open_block('What are the challenges you have faced in attending the course?', section_d.get('challenges')))
        elements.append(Spacer(1, 4))
        elements.append(open_block('Suggestions on how to improve the course:', section_d.get('suggestions')))

    doc.build(elements)
    pdf_data = buffer.getvalue()
    buffer.close()

    filename = f"student_feedback_responses_{session_obj.course_code or 'course'}.pdf"
    return Response(
        pdf_data,
        mimetype='application/pdf',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(pdf_data)),
        },
    )

@class_management_bp.route('/evaluation/<int:session_id>/student-feedback/responses/pdf-weasyprint')
@login_required
def student_feedback_responses_pdf_weasyprint(session_id):
    """Generate student feedback PDF using WeasyPrint with Kalpurush for Praise and Suggestions."""
    # Lazy import WeasyPrint - only when actually needed
    HTML = _get_weasyprint_html()
    if HTML is None:
        error_msg = 'Error generating PDF: WeasyPrint is not available. '
        error_msg += 'Please ensure WeasyPrint dependencies are installed. '
        error_msg += 'On macOS, run: brew install cairo pango gdk-pixbuf gobject-introspection'
        flash(error_msg, 'error')
        current_app.logger.error("WeasyPrint not available for PDF generation")
        current_app.logger.error(f"Current availability status: {_WEASYPRINT_AVAILABLE}")
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))
    
    session_obj = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session_obj.teacher_id:
        flash('You are not authorized to access this download.', 'danger')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_link = StudentFeedbackLink.query.filter_by(session_id=session_id).first()
    if not feedback_link:
        flash('No feedback responses found.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    responses = (
        StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
        .order_by(StudentFeedbackResponse.submitted_at.asc())
        .all()
    )
    if not responses:
        flash('No feedback responses to download.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    try:
        from error_handler import log_error
        import os
        
        # Prepare data for template
        feedback_data = []
        for idx, item in enumerate(responses, start=1):
            try:
                data = json.loads(item.payload or '{}')
            except json.JSONDecodeError:
                data = {}
            
            academic = data.get('academic_info', {}) or {}
            section_a = data.get('section_a', {}) or {}
            section_b = data.get('section_b', {}) or {}
            section_c = data.get('section_c', {}) or {}
            section_d = data.get('section_d', {}) or {}
            
            feedback_data.append({
                'index': idx,
                'submitted_at': item.submitted_at.strftime('%Y-%m-%d'),
                'academic': academic,
                'section_a': section_a,
                'section_b': section_b,
                'section_c': section_c,
                'section_d': section_d,
                'course_name': academic.get('course_title') or session_obj.course_name or '—',
                'course_code': academic.get('course_code') or session_obj.course_code or '—',
                'academic_session': academic.get('academic_session') or '—',
            })
        
        # Get font path for WeasyPrint
        font_path = os.path.join(current_app.root_path, 'static', 'Fonts', 'kalpurush.ttf')
        if not os.path.exists(font_path):
            # Try alternative path
            font_path = os.path.join(current_app.root_path, 'static', 'fonts', 'kalpurush.ttf')
        
        # Render template
        html_content = render_template(
            'class_management/student_feedback_weasyprint.html',
            feedback_data=feedback_data,
            feedback_section_a=FEEDBACK_SECTION_A,
            feedback_section_b_likert=FEEDBACK_SECTION_B_LIKERT,
            feedback_method_options=FEEDBACK_METHOD_OPTIONS,
            feedback_effort_options=FEEDBACK_EFFORT_OPTIONS,
            likert_options=['Strongly disagree', 'Disagree', 'Neutral', 'Agree', 'Strongly Agree'],
            kalpurush_font_path=font_path if os.path.exists(font_path) else None
        )
        
        # Generate PDF with WeasyPrint (lazy import already done above)
        try:
            pdf_buffer = io.BytesIO()
            HTML(string=html_content, base_url=request.url_root).write_pdf(pdf_buffer)
            pdf_buffer.seek(0)
        except Exception as e:
            current_app.logger.error(f"Error generating PDF with WeasyPrint: {e}", exc_info=True)
            flash(f'Error generating PDF: {str(e)}', 'error')
            return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))
        
        filename = f"student_feedback_responses_{session_obj.course_code or 'course'}.pdf"
        
        current_app.logger.info(f"WeasyPrint PDF generated successfully for student feedback session {session_id}")
        return Response(
            pdf_buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Length': str(len(pdf_buffer.getvalue()))
            }
        )
        
    except Exception as e:
        log_error(e, {
            'session_id': session_id,
            'function': 'student_feedback_responses_pdf_weasyprint',
            'user': current_user.username if current_user.is_authenticated else 'Anonymous'
        })
        current_app.logger.error(f"Error generating WeasyPrint student feedback PDF for session {session_id}: {e}", exc_info=True)
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

@class_management_bp.route('/evaluation/<int:session_id>/student-feedback/responses/docx')
@login_required
def student_feedback_responses_docx(session_id):
    session_obj = Session.query.get_or_404(session_id)
    teacher = Teacher.query.filter_by(name=current_user.full_name).first()
    if not teacher or teacher.id != session_obj.teacher_id:
        flash('You are not authorized to access this download.', 'danger')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    feedback_link = StudentFeedbackLink.query.filter_by(session_id=session_id).first()
    if not feedback_link:
        flash('No feedback responses found.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    responses = (
        StudentFeedbackResponse.query.filter_by(feedback_link_id=feedback_link.id)
        .order_by(StudentFeedbackResponse.submitted_at.asc())
        .all()
    )
    if not responses:
        flash('No feedback responses to download.', 'info')
        return redirect(url_for('class_management.student_feedback_manage', session_id=session_id))

    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    document = Document()
    normal_style = document.styles['Normal']
    normal_style.font.name = 'Kalpurush'
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Kalpurush')

    heading1 = document.styles['Heading 1']
    heading1.font.name = 'Helvetica'
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = document.styles['Heading 2']
    heading2.font.name = 'Helvetica'
    heading2.font.size = Pt(13)
    heading2.font.bold = True

    for idx, item in enumerate(responses, start=1):
        try:
            data = json.loads(item.payload or '{}')
        except json.JSONDecodeError:
            data = {}

        academic = data.get('academic_info', {}) or {}
        section_a = data.get('section_a', {}) or {}
        section_b = data.get('section_b', {}) or {}
        section_c = data.get('section_c', {}) or {}
        section_d = data.get('section_d', {}) or {}

        if idx > 1:
            document.add_page_break()

        title_para = document.add_paragraph('STUDENT FEEDBACK FORM')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(16)

        uni_para = document.add_paragraph('KHULNA UNIVERSITY')
        uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        uni_para.runs[0].bold = True
        uni_para.runs[0].font.size = Pt(12)

        meta_para = document.add_paragraph(f"Response {idx} - {item.submitted_at.strftime('%Y-%m-%d %H:%M')}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.runs[0].font.size = Pt(10)

        add_header_table([
            ('Academic Session', academic.get('academic_session') or '—'),
            ('Title of the Course', academic.get('course_title') or session_obj.course_name or '—'),
            ('Course Code', academic.get('course_code') or session_obj.course_code or '—'),
        ])

        instruction_para = document.add_paragraph(
            'PLEASE TICK/CROSS IN THE BLANK SPACE WHICH BEST DESCRIBES HOW MUCH YOU AGREE WITH THE FOLLOWING STATEMENTS'
        )
        instruction_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        instruction_para.runs[0].font.size = Pt(9)

        document.add_paragraph('A. SATISFACTION WITH THE COURSE', style='Heading 3')
        section_a_rows = []
        for key, question in FEEDBACK_SECTION_A:
            selected = section_a.get(key)
            row = [question]
            for option in likert_header[1:]:
                row.append('✓' if selected and selected.lower() == option.lower() else '')
            section_a_rows.append(row)
        add_likert_table(section_a_rows, likert_header)

        document.add_paragraph('B. TEACHING-LEARNING METHODS', style='Heading 3')
        first_question = FEEDBACK_SECTION_B_LIKERT[0]
        first_selected = section_b.get(first_question[0])
        first_rows = [[first_question[1]] + [('✓' if first_selected and first_selected.lower() == opt.lower() else '') for opt in likert_header[1:]]]
        add_likert_table(first_rows, likert_header)

        methods_heading = document.add_paragraph('Teaching methods that contributed significantly (selected):')
        methods_heading.runs[0].bold = True
        methods = section_b.get('teaching_methods') or []
        if methods:
            for method in methods:
                bullet = document.add_paragraph(method, style='List Bullet')
                bullet.paragraph_format.space_after = Pt(1)
        else:
            document.add_paragraph('—', style='List Bullet')

        remaining_rows = []
        for key, question in FEEDBACK_SECTION_B_LIKERT[1:]:
            selected = section_b.get(key)
            row = [question]
            for option in likert_header[1:]:
                row.append('✓' if selected and selected.lower() == option.lower() else '')
            remaining_rows.append(row)
        add_likert_table(remaining_rows, likert_header)

        document.add_paragraph('C. ENGAGEMENT AND WORKLOAD', style='Heading 3')
        engagement_rows = [
            ('How much time do you devote to this course before and after each lecture?', section_c.get('study_time') or '—'),
            ('About what percent of the class meetings (including discussions) did you attend?', section_c.get('attendance_percent') or '—'),
        ]
        engagement_table = document.add_table(rows=len(engagement_rows), cols=2)
        engagement_table.style = 'Table Grid'
        for row_idx, (label, value) in enumerate(engagement_rows):
            set_cell_shading(engagement_table.cell(row_idx, 0), 'D9D9D9')
            set_cell_text(engagement_table.cell(row_idx, 0), label, bold=True, align='left')
            set_cell_text(engagement_table.cell(row_idx, 1), value, align='left')

        effort_heading = document.add_paragraph('Significant aspects of your effort (selected):')
        effort_heading.runs[0].bold = True
        efforts = section_c.get('effort_focus') or []
        if efforts:
            for effort in efforts:
                bullet = document.add_paragraph(effort, style='List Bullet')
                bullet.paragraph_format.space_after = Pt(1)
        else:
            document.add_paragraph('—', style='List Bullet')

        document.add_paragraph('D. PRAISE AND SUGGESTIONS', style='Heading 3')
        add_open_block('What did you like especially about this course?', section_d.get('likes'))
        add_open_block('What are the challenges you have faced in attending the course?', section_d.get('challenges'))
        add_open_block('Suggestions on how to improve the course:', section_d.get('suggestions'))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    docx_data = buffer.getvalue()
    buffer.close()

    filename = f"student_feedback_responses_{session_obj.course_code or 'course'}.docx"
    return Response(
        docx_data,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"',
            'Content-Length': str(len(docx_data)),
        },
    )

@class_management_bp.route('/student/course-files')
@login_required
def student_course_files():
    """Student view for course files - course outlines and teacher-uploaded files"""
    try:
        # Get student ID from current user (assuming username is student_id)
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('index'))
        
        # Find all ClassStudent records for this student
        student_records = ClassStudent.query.filter_by(student_id=student_id).all()
        
        # Get all sessions where student is enrolled
        enrolled_sessions = []
        for student_record in student_records:
            session_obj = Session.query.get(student_record.session_id)
            if session_obj and not session_obj.archived:
                enrolled_sessions.append(session_obj)
        
        # Import CourseFileUpload model
        from blueprints.class_management.models import CourseFileUpload
        
        # Get course outlines that are enabled for student access
        course_files = []
        for session in enrolled_sessions:
            course_outline = CourseOutline.query.filter_by(session_id=session.id).first()
            if course_outline and course_outline.student_access_enabled:
                course_files.append({
                    'session_id': session.id,
                    'course_code': session.course_code,
                    'course_name': session.course_name,
                    'teacher_name': session.teacher.name if session.teacher else 'Unknown',
                    'academic_session': session.academic_session,
                    'year': session.year,
                    'term': session.term,
                    'type': 'course_outline',
                    'file_name': f"{session.course_code or 'Course'}_Outline.pdf"
                })
            
            # Get teacher-uploaded files for this session
            uploaded_files = CourseFileUpload.query.filter_by(
                session_id=session.id,
                student_access_enabled=True
            ).all()
            
            for uploaded_file in uploaded_files:
                course_files.append({
                    'session_id': session.id,
                    'course_code': session.course_code,
                    'course_name': session.course_name,
                    'teacher_name': session.teacher.name if session.teacher else 'Unknown',
                    'academic_session': session.academic_session,
                    'year': session.year,
                    'term': session.term,
                    'type': 'uploaded_file',
                    'file_id': uploaded_file.id,
                    'file_name': uploaded_file.file_name,
                    'description': uploaded_file.description
                })
        
        # Sort by academic session, year, term
        course_files.sort(key=lambda x: (
            x.get('academic_session', ''),
            x.get('year', ''),
            x.get('term', '')
        ), reverse=True)
        
        return render_template('class_management/student_course_files.html',
                             course_files=course_files)
    except Exception as e:
        current_app.logger.error(f"Error in student_course_files: {e}", exc_info=True)
        flash('An error occurred while loading course files.', 'error')
        return redirect(url_for('index'))

@class_management_bp.route('/student/course-files/<int:session_id>/download-pdf')
@login_required
def student_download_course_outline_pdf(session_id):
    """Student download course outline PDF"""
    try:
        # Get student ID from current user
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if student is enrolled in this session
        student_record = ClassStudent.query.filter_by(
            session_id=session_id,
            student_id=student_id
        ).first()
        
        if not student_record:
            flash('You are not enrolled in this course.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Get course outline
        course_outline = CourseOutline.query.filter_by(session_id=session_id).first()
        if not course_outline:
            flash('Course outline not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if student access is enabled
        if not course_outline.student_access_enabled:
            flash('This course outline is not available for download.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Generate and return PDF (skip auth check since we already verified student access)
        return _generate_course_outline_pdf(session_id, skip_auth_check=True)
    except Exception as e:
        current_app.logger.error(f"Error in student_download_course_outline_pdf: {e}", exc_info=True)
        flash('An error occurred while downloading the course outline.', 'error')
        return redirect(url_for('class_management.student_course_files'))

@class_management_bp.route('/student/course-files/<int:file_id>/download')
@login_required
def student_download_uploaded_file(file_id):
    """Student download teacher-uploaded file"""
    try:
        from blueprints.class_management.models import CourseFileUpload
        import os
        from flask import send_file
        
        # Get student ID from current user
        student_id = current_user.username if hasattr(current_user, 'username') else None
        if not student_id:
            flash('Student ID not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Get uploaded file
        uploaded_file = CourseFileUpload.query.get_or_404(file_id)
        
        # Check if student access is enabled
        if not uploaded_file.student_access_enabled:
            flash('This file is not available for download.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if student is enrolled in this session
        student_record = ClassStudent.query.filter_by(
            session_id=uploaded_file.session_id,
            student_id=student_id
        ).first()
        
        if not student_record:
            flash('You are not enrolled in this course.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Check if file exists
        file_path = uploaded_file.file_path
        if not os.path.exists(file_path):
            flash('File not found.', 'error')
            return redirect(url_for('class_management.student_course_files'))
        
        # Send file for download
        return send_file(
            file_path,
            as_attachment=True,
            download_name=uploaded_file.file_name
        )
    except Exception as e:
        current_app.logger.error(f"Error in student_download_uploaded_file: {e}", exc_info=True)
        flash('An error occurred while downloading the file.', 'error')
        return redirect(url_for('class_management.student_course_files'))