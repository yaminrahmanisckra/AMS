from flask import Blueprint, render_template, redirect, url_for, flash, request, send_file, Response, current_app, jsonify
from flask_login import login_required, current_user
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Frame, PageTemplate
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.lib.units import inch
from io import BytesIO
from .models import db, RSession, RStudent, RSubject, RMark, RCourseRegistration
from role_utils import parse_roles, is_admin
from blueprints.class_management.models import Teacher, Session as ClassSession, ExamPaperEvaluation
from blueprints.course_management.models import StudentCourseRegistration, Course
from blueprints.student_management.models import Student as StudentProfile
from blueprints.course_management.models import DutyAssignment
import json
from openpyxl import load_workbook
import io
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, A4
import zipfile
# docx imports moved to lazy imports (only when needed) to prevent startup hang
# from docx import Document
# from docx.shared import Inches, Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH

result_management_bp = Blueprint('result_management', __name__, template_folder='templates/result_management')


def _is_head_user():
    if not current_user.is_authenticated:
        return False
    roles = parse_roles(current_user.role)
    return 'head' in roles or 'dean' in roles


def _current_teacher():
    if not current_user.is_authenticated:
        return None
    return Teacher.query.filter_by(name=current_user.full_name).first()


def _is_tabulator_user():
    teacher = _current_teacher()
    if not teacher:
        return False
    return DutyAssignment.query.filter_by(
        assigned_teacher_id=teacher.id,
        duty_type='tabulator',
        status='active'
    ).count() > 0

def _get_tabulator_assignments():
    """Get all tabulator assignments for the current user"""
    teacher = _current_teacher()
    if not teacher:
        return []
    return DutyAssignment.query.filter_by(
        assigned_teacher_id=teacher.id,
        duty_type='tabulator',
        status='active'
    ).all()

def _can_access_session(session):
    """Check if current user can access a specific result session"""
    if is_admin(current_user):
        return True
    if _is_head_user():
        return True
    
    # For tabulators, check if session matches their assignment
    if _is_tabulator_user():
        assignments = _get_tabulator_assignments()
        for assignment in assignments:
            # Strict matching: all specified fields in assignment must match session exactly
            # If assignment has academic_session, year, or term, they must match
            
            # Check academic_session match (RSession.name corresponds to academic_session)
            if assignment.academic_session:
                if not session.name or session.name.strip() != assignment.academic_session.strip():
                    continue  # No match, check next assignment
            
            # Check year match
            if assignment.year:
                if not session.year or str(session.year).strip() != str(assignment.year).strip():
                    continue  # No match, check next assignment
            
            # Check term match
            if assignment.term:
                if not session.term or str(session.term).strip() != str(assignment.term).strip():
                    continue  # No match, check next assignment
            
            # All specified fields match - allow access
            return True
        return False
    
    return False


def _has_result_access():
    if not current_user.is_authenticated:
        return False
    if is_admin(current_user):
        return True
    if _is_head_user():
        return True
    return _is_tabulator_user()


def _can_manage_sessions():
    if not current_user.is_authenticated:
        return False
    return is_admin(current_user) or _is_head_user()


@result_management_bp.before_request
def restrict_result_module():
    """Allow only Head/Dean/admin or assigned tabulators to access Result Management."""
    if not current_user.is_authenticated:
        return
    if _has_result_access():
        return
    flash('Result Management is restricted to the Head or assigned Tabulators.', 'danger')
    return redirect(url_for('index'))

def _determine_subject_type(course):
    """Return the subject type label used in Result Management based on course data."""
    course_type = (course.course_type or '').strip()
    category = (getattr(course, 'category', '') or '').strip().lower()

    # Handle new Thesis types
    if course_type == 'Thesis (UG)':
        return 'Thesis (UG)'
    if course_type == 'Thesis I (UG)':
        return 'Thesis I (UG)'
    if course_type == 'Thesis II (UG)':
        return 'Thesis II (UG)'
    
    # Handle Dissertation types
    if course_type == 'Dissertation Proposal (PG)':
        return 'Dissertation'  # Will set dissertation_type='Type1' when creating subject
    if course_type == 'Dissertation Defence (PG)':
        return 'Dissertation'  # Will set dissertation_type='Type2' when creating subject
    
    # Handle existing types (case-insensitive for backward compatibility)
    course_type_lower = course_type.lower()
    if course_type_lower == 'theory':
        if category == 'ug':
            return 'Theory (UG)'
        if category == 'pg':
            return 'Theory (PG)'
        return 'Theory'
    if course_type_lower == 'sessional':
        return 'Sessional'
    if course_type_lower == 'viva':
        return 'Viva'
    if course_type_lower == 'dissertation':
        return 'Dissertation'
    
    return course.course_type or 'Theory'

def calculate_grade(total_marks, is_retake=False):
    grade_point = 0.0
    grade_letter = 'F'
    if total_marks >= 80: grade_point, grade_letter = 4.0, 'A+'
    elif total_marks >= 75: grade_point, grade_letter = 3.75, 'A'
    elif total_marks >= 70: grade_point, grade_letter = 3.5, 'A-'
    elif total_marks >= 65: grade_point, grade_letter = 3.25, 'B+'
    elif total_marks >= 60: grade_point, grade_letter = 3.0, 'B'
    elif total_marks >= 55: grade_point, grade_letter = 2.75, 'B-'
    elif total_marks >= 50: grade_point, grade_letter = 2.5, 'C+'
    elif total_marks >= 45: grade_point, grade_letter = 2.25, 'C'
    elif total_marks >= 40: grade_point, grade_letter = 2.0, 'D'
    if is_retake and grade_letter != 'F':
        if grade_letter == 'A+': grade_point, grade_letter = 3.75, 'A'
        elif grade_letter == 'A': grade_point, grade_letter = 3.5, 'A-'
        elif grade_letter == 'A-': grade_point, grade_letter = 3.25, 'B+'
        elif grade_letter == 'B+': grade_point, grade_letter = 3.0, 'B'
        elif grade_letter == 'B': grade_point, grade_letter = 2.75, 'B-'
        elif grade_letter == 'B-': grade_point, grade_letter = 2.5, 'C+'
        elif grade_letter == 'C+': grade_point, grade_letter = 2.25, 'C'
        elif grade_letter == 'C': grade_point, grade_letter = 2.0, 'D'
    return grade_point, grade_letter

def convert_to_roman(num):
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num

# --- Page Numbering Function ---
def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 9)
    page_number_text = f"Page {doc.page} of {doc.doc.page_count}"
    canvas.drawRightString(letter[0] - 40, 30, page_number_text)
    canvas.restoreState()

@result_management_bp.route('/')
@login_required
def index():
    all_sessions = RSession.query.filter_by(is_archived=False).order_by(RSession.created_at.desc()).all()
    
    # Filter sessions based on user access
    if is_admin(current_user) or _is_head_user():
        sessions = all_sessions
    elif _is_tabulator_user():
        # Tabulators can only see sessions matching their assignment
        sessions = [s for s in all_sessions if _can_access_session(s)]
    else:
        sessions = []
    
    return render_template(
        'rm_index.html',
        sessions=sessions,
        can_manage_sessions=_can_manage_sessions()
    )

@result_management_bp.route('/archived')
@login_required
def archived_sessions():
    all_sessions = RSession.query.filter_by(is_archived=True).order_by(RSession.created_at.desc()).all()
    
    # Filter sessions based on user access
    if is_admin(current_user) or _is_head_user():
        sessions = all_sessions
    elif _is_tabulator_user():
        # Tabulators can only see sessions matching their assignment
        sessions = [s for s in all_sessions if _can_access_session(s)]
    else:
        sessions = []
    
    return render_template(
        'rm_archive.html',
        sessions=sessions,
        can_manage_sessions=_can_manage_sessions()
    )

@result_management_bp.route('/archive_session/<int:session_id>', methods=['POST'])
@login_required
def archive_session(session_id):
    if not _can_manage_sessions():
        flash('Only the Head can archive sessions.', 'danger')
        return redirect(url_for('result_management.index'))
    session = RSession.query.get_or_404(session_id)
    session.is_archived = True
    db.session.commit()
    flash(f'Session "{session.name}" has been archived.', 'success')
    return redirect(url_for('result_management.index'))

@result_management_bp.route('/unarchive_session/<int:session_id>', methods=['POST'])
@login_required
def unarchive_session(session_id):
    if not _can_manage_sessions():
        flash('Only the Head can unarchive sessions.', 'danger')
        return redirect(url_for('result_management.archived_sessions'))
    session = RSession.query.get_or_404(session_id)
    session.is_archived = False
    db.session.commit()
    flash(f'Session "{session.name}" has been unarchived.', 'success')
    return redirect(url_for('result_management.archived_sessions'))

@result_management_bp.route('/add_session', methods=['GET', 'POST'])
@login_required
def add_session():
    if not _can_manage_sessions():
        flash('Only the Head can create result sessions.', 'danger')
        return redirect(url_for('result_management.index'))

    # Fetch distinct values from class sessions for dropdowns
    session_options = [
        value[0] for value in db.session.query(ClassSession.academic_session)
        .filter(ClassSession.academic_session.isnot(None))
        .distinct()
        .order_by(ClassSession.academic_session.asc())
        .all()
        if value[0]
    ]
    year_options = [
        value[0] for value in db.session.query(ClassSession.year)
        .filter(ClassSession.year.isnot(None))
        .distinct()
        .order_by(ClassSession.year.asc())
        .all()
        if value[0]
    ]
    term_options = [
        value[0] for value in db.session.query(ClassSession.term)
        .filter(ClassSession.term.isnot(None))
        .distinct()
        .order_by(ClassSession.term.asc())
        .all()
        if value[0]
    ]

    context = {
        'session_options': session_options,
        'year_options': year_options,
        'term_options': term_options,
        'selected_session': None,
        'selected_year': None,
        'selected_term': None,
    }

    if request.method == 'POST':
        name = request.form.get('name')
        term = request.form.get('term')
        year = request.form.get('year')

        context['selected_session'] = name
        context['selected_year'] = year
        context['selected_term'] = term

        if name and term:
            new_session = RSession(name=name, term=term, year=year)
            db.session.add(new_session)
            db.session.commit()
            flash('Session added successfully!', 'success')
            return redirect(url_for('result_management.index'))
        else:
            flash('Session name and term are required.', 'danger')

    return render_template('rm_add_session.html', **context)

@result_management_bp.route('/add_student/<int:session_id>', methods=['GET', 'POST'])
@login_required
def add_student(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        flash('You do not have access to this session.', 'danger')
        return redirect(url_for('result_management.index'))
    if request.method == 'POST':
        if 'excel_file' in request.files and request.files['excel_file'].filename != '':
            file = request.files['excel_file']
            if file and file.filename.endswith('.xlsx'):
                try:
                    wb = load_workbook(file)
                    ws = wb.active
                    added_count = 0
                    skipped_count = 0

                    # --- Optimization Start ---
                    # 1. Read all student IDs from the Excel file
                    all_student_ids_from_excel = []
                    rows_to_process = []
                    for row in ws.iter_rows(min_row=2, values_only=True):
                        student_data = (list(row) + [None]*5)[:5]
                        student_id = student_data[0]
                        if student_id:
                            all_student_ids_from_excel.append(str(student_id))
                            rows_to_process.append(student_data)

                    # 2. Find which of these students already exist in the DB in a single query
                    existing_students = db.session.query(RStudent.student_id).filter(
                        RStudent.session_id == session_id,
                        RStudent.student_id.in_(all_student_ids_from_excel)
                    ).all()
                    existing_student_ids = {str(s_id[0]) for s_id in existing_students}

                    # 3. Iterate and add only the new students
                    students_to_add = []
                    for student_id, name, year, discipline, school in rows_to_process:
                        if str(student_id) not in existing_student_ids:
                            students_to_add.append(RStudent(
                                student_id=str(student_id), name=name, year=year,
                                discipline=discipline, school=school, session_id=session_id
                            ))
                            added_count += 1
                        else:
                            skipped_count += 1
                    
                    if students_to_add:
                        db.session.bulk_save_objects(students_to_add)
                    # --- Optimization End ---
                    
                    db.session.commit()
                    flash(f'Successfully added {added_count} new students. Skipped {skipped_count} existing students.', 'success')
                except Exception as e:
                    flash(f'Error processing Excel file: {e}', 'danger')
                return redirect(url_for('result_management.add_student', session_id=session_id))
            else:
                flash('Invalid file type. Please upload a .xlsx file.', 'danger')
        elif request.is_json:
            # Handle AJAX request for batch or ID-based addition
            data = request.get_json()
            student_ids = data.get('student_ids', [])
            
            if not student_ids:
                return jsonify({'success': False, 'message': 'No students selected!'}), 400
            
            try:
                from blueprints.student_management.models import Student
                if not Student:
                    return jsonify({'success': False, 'message': 'Students Management module not available'}), 503
                
                added_count = 0
                skipped_count = 0
                
                # Get existing student IDs for this session
                existing_student_ids = {s.student_id for s in RStudent.query.filter_by(session_id=session_id).all()}
                
                for student_db_id in student_ids:
                    student = Student.query.get(student_db_id)
                    if student:
                        if student.student_id in existing_student_ids:
                            skipped_count += 1
                            continue
                        
                        r_student = RStudent(
                            student_id=student.student_id,
                            name=student.name,
                            session_id=session_id
                        )
                        db.session.add(r_student)
                        existing_student_ids.add(student.student_id)
                        added_count += 1
                
                db.session.commit()
                message = f'Successfully added {added_count} student(s).'
                if skipped_count > 0:
                    message += f' Skipped {skipped_count} existing student(s).'
                return jsonify({'success': True, 'message': message})
            except Exception as e:
                db.session.rollback()
                current_app.logger.error(f"Error adding students to result session {session_id}: {e}", exc_info=True)
                return jsonify({'success': False, 'message': f'Error adding students: {str(e)}'}), 500
        else:
            student_id = request.form.get('student_id')
            name = request.form.get('name')
            if student_id and name:
                exists = RStudent.query.filter_by(student_id=student_id, session_id=session_id).first()
                if exists:
                    flash('Student with this ID already exists in this session.', 'danger')
                else:
                    student = RStudent(
                        student_id=student_id, name=name, session_id=session_id
                    )
                    db.session.add(student)
                    db.session.commit()
                    flash('Student added successfully!', 'success')
            else:
                flash('Student ID and Name are required for single add.', 'warning')
        return redirect(url_for('result_management.add_student', session_id=session_id))
    
    session = RSession.query.get_or_404(session_id)
    
    # Auto-load students from Class Management based on session/year/term
    try:
        from blueprints.class_management.models import Session as ClassSession, ClassStudent
        from blueprints.student_management.models import Student
        
        # Find all ClassSessions matching the result session's academic_session, year, and term
        query_class_sessions = ClassSession.query.filter(
            ClassSession.academic_session == session.name,
            ClassSession.term == session.term
        )
        # Only filter by year if it's provided in the result session
        if session.year:
            query_class_sessions = query_class_sessions.filter(ClassSession.year == session.year)
        
        matching_class_sessions = query_class_sessions.all()
        
        if matching_class_sessions:
            class_session_ids = [cs.id for cs in matching_class_sessions]
            # Get all student_ids from those class sessions
            class_students = ClassStudent.query.filter(
                ClassStudent.session_id.in_(class_session_ids)
            ).all()
            
            # Get existing RStudent IDs for this session
            existing_rstudent_ids = {s.student_id for s in RStudent.query.filter_by(session_id=session_id).all()}
            
            # Auto-add students that don't exist in RStudent
            students_to_add = []
            for cs in class_students:
                if cs.student_id not in existing_rstudent_ids:
                    # Get student details from Student Management if available
                    student_detail = None
                    if Student:
                        student_detail = Student.query.filter_by(student_id=cs.student_id).first()
                    
                    r_student = RStudent(
                        student_id=cs.student_id,
                        name=cs.name,
                        session_id=session_id,
                        year=student_detail.batch if student_detail and student_detail.batch else None,
                        discipline=None,
                        school=None
                    )
                    students_to_add.append(r_student)
                    existing_rstudent_ids.add(cs.student_id)
            
            if students_to_add:
                db.session.bulk_save_objects(students_to_add)
                db.session.commit()
                flash(f'Auto-loaded {len(students_to_add)} student(s) from Class Management.', 'success')
    except Exception as e:
        current_app.logger.error(f'Error auto-loading students: {str(e)}', exc_info=True)
        # Don't show error to user, just log it
    
    students = RStudent.query.filter_by(session_id=session_id).order_by(RStudent.student_id).all()
    
    # Get batches for dropdown
    batches = []
    try:
        from blueprints.student_management.models import Student
        if Student:
            all_batches = db.session.query(Student.batch).distinct().filter(Student.batch.isnot(None)).order_by(Student.batch.desc()).all()
            batches = [batch[0] for batch in all_batches]
    except:
        batches = []
    
    return render_template('rm_add_student.html', session=session, students=students, batches=batches)

@result_management_bp.route('/edit_student/<int:student_id>', methods=['GET', 'POST'])
@login_required
def edit_student(student_id):
    student = RStudent.query.get_or_404(student_id)
    if request.method == 'POST':
        student.student_id = request.form['student_id']
        student.name = request.form['name']
        student.year = request.form.get('year')
        student.discipline = request.form.get('discipline')
        student.school = request.form.get('school')
        db.session.commit()
        flash('Student updated successfully!', 'success')
        return redirect(url_for('result_management.add_student', session_id=student.session_id))
    return render_template('rm_edit_student.html', student=student)

@result_management_bp.route('/delete_student/<int:student_id>', methods=['POST'])
@login_required
def delete_student(student_id):
    """Delete a student from result session"""
    try:
        student = RStudent.query.get_or_404(student_id)
        session_id = student.session_id
        student_name = student.name
        
        # Cascade delete will handle marks and registrations
        # due to cascade="all, delete-orphan" in the model relationships
        db.session.delete(student)
        db.session.commit()
        flash(f'Student "{student_name}" deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error deleting result student {student_id}: {e}', exc_info=True)
        flash(f'Error deleting student: {str(e)}', 'danger')
        # Try to get session_id for redirect even if deletion failed
        try:
            student = RStudent.query.get(student_id)
            session_id = student.session_id if student else None
        except:
            session_id = None
    
    if session_id:
        return redirect(url_for('result_management.add_student', session_id=session_id))
    return redirect(url_for('result_management.index'))

@result_management_bp.route('/api/students', methods=['GET'])
@login_required
def get_students_for_result():
    """Get students from Students Management for selection (AJAX)"""
    try:
        from blueprints.student_management.models import Student
        from blueprints.class_management.models import Session as ClassSession, ClassStudent
        from sqlalchemy import or_
        
        if not Student:
            return jsonify({'success': False, 'message': 'Students Management module not available'}), 503
        
        batch_filter = request.args.get('batch', '').strip()
        search = request.args.get('search', '').strip()
        session_id = request.args.get('session_id', type=int)
        
        # If session_id is provided, filter students by that session's academic_session, year, and term
        student_ids_from_class_sessions = None
        if session_id:
            session = RSession.query.get(session_id)
            
            if session:
                # Find all ClassSessions matching the result session's academic_session, year, and term
                query_class_sessions = ClassSession.query.filter(
                    ClassSession.academic_session == session.name,
                    ClassSession.term == session.term
                )
                # Only filter by year if it's provided in the result session
                if session.year:
                    query_class_sessions = query_class_sessions.filter(ClassSession.year == session.year)
                
                matching_class_sessions = query_class_sessions.all()
                
                if matching_class_sessions:
                    class_session_ids = [cs.id for cs in matching_class_sessions]
                    # Get all student_ids from those class sessions
                    class_students = ClassStudent.query.filter(
                        ClassStudent.session_id.in_(class_session_ids)
                    ).all()
                    student_ids_from_class_sessions = {cs.student_id for cs in class_students}
        
        query = Student.query
        
        # Filter by student_ids from class sessions if available
        if student_ids_from_class_sessions:
            query = query.filter(Student.student_id.in_(student_ids_from_class_sessions))
        elif session_id:
            # If session_id provided but no matching class sessions, return empty result
            return jsonify({
                'success': True,
                'students': [],
                'message': 'No students found for this session. Please ensure class sessions exist for this academic session, year, and term.'
            })
        
        if batch_filter:
            query = query.filter(Student.batch == batch_filter)
        
        if search:
            query = query.filter(
                or_(
                    Student.name.ilike(f'%{search}%'),
                    Student.student_id.ilike(f'%{search}%')
                )
            )
        
        students = query.order_by(Student.student_id.asc()).limit(500).all()
        
        return jsonify({
            'success': True,
            'students': [{
                'id': s.id,
                'student_id': s.student_id,
                'name': s.name,
                'batch': s.batch,
                'email': s.email,
                'phone': s.phone
            } for s in students]
        })
    except Exception as e:
        current_app.logger.error(f'Error in get_students_for_result: {str(e)}', exc_info=True)
        return jsonify({'success': False, 'message': f'Error loading students: {str(e)}'}), 500

@result_management_bp.route('/add_subject/<int:session_id>', methods=['GET', 'POST'])
@login_required
def add_subject(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        flash('You do not have access to this session.', 'danger')
        return redirect(url_for('result_management.index'))
    
    if request.method == 'POST':
        # Handle batch selection and auto-load courses
        if request.is_json:
            data = request.get_json()
            batch = data.get('batch', '').strip()
            curriculum_id = data.get('curriculum_id')
            
            # Convert curriculum_id to int if it's a string
            if curriculum_id:
                try:
                    curriculum_id = int(curriculum_id)
                except (ValueError, TypeError):
                    curriculum_id = None
            
            if batch and curriculum_id:
                # Update session with batch and curriculum
                session.batch = batch
                session.curriculum_id = curriculum_id
                db.session.commit()
                
                # Auto-load offered courses from curriculum
                try:
                    from blueprints.course_management.models import Course
                    
                    if not Course:
                        return jsonify({'success': False, 'message': 'Course Management module not available'}), 503
                    
                    courses = Course.query.filter_by(
                        curriculum_id=curriculum_id,
                        offered=True
                    ).all()
                    
                    if not courses:
                        return jsonify({
                            'success': True,
                            'message': 'No offered courses found in this curriculum.',
                            'added_count': 0,
                            'skipped_count': 0
                        })
                    
                    added_count = 0
                    skipped_count = 0
                    
                    # Get existing subject codes to avoid duplicates
                    existing_codes = {s.code for s in RSubject.query.filter_by(session_id=session_id).all()}
                    
                    for course in courses:
                        if course.course_code in existing_codes:
                            skipped_count += 1
                            continue
                        
                        subject_type = _determine_subject_type(course)
                        
                        # Determine dissertation_type if applicable
                        dissertation_type = None
                        if course.course_type == 'Dissertation Proposal (PG)':
                            dissertation_type = 'Type1'
                        elif course.course_type == 'Dissertation Defence (PG)':
                            dissertation_type = 'Type2'
                        
                        subject = RSubject(
                            code=course.course_code,
                            name=course.course_name,
                            credit=course.credit,
                            subject_type=subject_type,
                            dissertation_type=dissertation_type,
                            session_id=session_id
                        )
                        db.session.add(subject)
                        existing_codes.add(course.course_code)
                        added_count += 1
                    
                    db.session.commit()
                    return jsonify({
                        'success': True,
                        'message': f'Successfully loaded {added_count} offered course(s). {skipped_count} already existed.',
                        'added_count': added_count,
                        'skipped_count': skipped_count
                    })
                except ImportError as e:
                    db.session.rollback()
                    current_app.logger.error(f"Import error: {e}", exc_info=True)
                    return jsonify({'success': False, 'message': 'Course Management module not available. Please ensure it is properly configured.'}), 503
                except Exception as e:
                    db.session.rollback()
                    current_app.logger.error(f"Error auto-loading courses: {e}", exc_info=True)
                    return jsonify({'success': False, 'message': f'Error loading courses: {str(e)}'}), 500
            
            return jsonify({'success': False, 'message': 'Batch and curriculum are required'}), 400
        
        # Handle manual subject addition
        code = request.form['code']
        name = request.form['name']
        credit = float(request.form['credit'])
        subject_type = request.form['subject_type']
        dissertation_type = request.form.get('dissertation_type') if subject_type == 'Dissertation' else None
        
        subject = RSubject(
            code=code, name=name, credit=credit, subject_type=subject_type,
            dissertation_type=dissertation_type, session_id=session_id
        )
        db.session.add(subject)
        db.session.commit()
        flash('Subject added successfully!', 'success')
        return redirect(url_for('result_management.add_subject', session_id=session_id))
    
    # GET request - show form
    # Auto-load subjects from curriculum based on session/year/term
    # First verify session, year, and term before loading
    try:
        from blueprints.course_management.models import CurriculumYearTerm, Course
        
        # Verify that session has required fields
        if not session.name or not session.term:
            current_app.logger.warning(f'Cannot auto-load subjects: session missing name or term. Session ID: {session_id}')
        else:
            # Find CurriculumYearTerm matching the session's academic_session, year, and term
            # ALL three must match: academic_session, year, and term
            query_cyt = CurriculumYearTerm.query.filter(
                CurriculumYearTerm.academic_session == session.name,
                CurriculumYearTerm.term == session.term
            )
            # Year is required - filter by it
            if session.year:
                query_cyt = query_cyt.filter(CurriculumYearTerm.year == session.year)
            else:
                # If year is not provided, we cannot match accurately
                current_app.logger.warning(f'Cannot auto-load subjects: session missing year. Session ID: {session_id}')
                query_cyt = None
            
            matching_cyt = query_cyt.first() if query_cyt else None
            
            if matching_cyt and matching_cyt.curriculum_id:
                # Verify that the matching CurriculumYearTerm actually matches all three fields
                cyt_matches = (
                    matching_cyt.academic_session == session.name and
                    matching_cyt.term == session.term and
                    (not session.year or matching_cyt.year == session.year)
                )
                
                if not cyt_matches:
                    current_app.logger.warning(f'CurriculumYearTerm does not match session criteria. Session: {session.name}/{session.year}/{session.term}, CYT: {matching_cyt.academic_session}/{matching_cyt.year}/{matching_cyt.term}')
                else:
                    # Get existing subject codes to avoid duplicates
                    existing_codes = {s.code for s in RSubject.query.filter_by(session_id=session_id).all()}
                    
                    # Load offered courses from the curriculum
                    # Then filter strictly by year/term
                    all_courses = Course.query.filter_by(
                        curriculum_id=matching_cyt.curriculum_id,
                        offered=True
                    ).all()
                    
                    # Strictly filter courses by year/term - must match exactly
                    courses = []
                    session_year_normalized = str(session.year).strip().lower() if session.year else None
                    session_term_normalized = str(session.term).strip().lower() if session.term else None
                    
                    for course in all_courses:
                        # Check stored year/term first
                        course_year = course.year
                        course_term = course.term
                        
                        # If not stored, use derived from course code
                        if not course_year:
                            course_year = course.derived_year
                        if not course_term:
                            course_term = course.derived_term
                        
                        # If course has no year/term at all (neither stored nor derived), skip it
                        if not course_year or not course_term:
                            continue
                        
                        # Normalize for comparison
                        course_year_normalized = str(course_year).strip().lower() if course_year else None
                        course_term_normalized = str(course_term).strip().lower() if course_term else None
                        
                        # Both year and term must match EXACTLY - if session has year/term, course must match
                        year_match = False
                        term_match = False
                        
                        if session_year_normalized:
                            year_match = (course_year_normalized and 
                                         session_year_normalized == course_year_normalized)
                        else:
                            # If session has no year, we can't match accurately
                            year_match = False
                        
                        if session_term_normalized:
                            term_match = (course_term_normalized and 
                                         session_term_normalized == course_term_normalized)
                        else:
                            # If session has no term, we can't match accurately
                            term_match = False
                        
                        if year_match and term_match:
                            courses.append(course)
                    
                    if courses:
                        subjects_to_add = []
                        for course in courses:
                            if course.course_code in existing_codes:
                                continue
                            
                            subject_type = _determine_subject_type(course)
                            
                            # Determine dissertation_type if applicable
                            dissertation_type = None
                            if course.course_type == 'Dissertation Proposal (PG)':
                                dissertation_type = 'Type1'
                            elif course.course_type == 'Dissertation Defence (PG)':
                                dissertation_type = 'Type2'
                            
                            subject = RSubject(
                                code=course.course_code,
                                name=course.course_name,
                                credit=course.credit,
                                subject_type=subject_type,
                                dissertation_type=dissertation_type,
                                session_id=session_id
                            )
                            subjects_to_add.append(subject)
                            existing_codes.add(course.course_code)
                        
                        if subjects_to_add:
                            db.session.bulk_save_objects(subjects_to_add)
                            db.session.commit()
                            flash(f'Auto-loaded {len(subjects_to_add)} subject(s) from curriculum for {session.name} - Year {session.year} - Term {session.term}.', 'success')
    except Exception as e:
        current_app.logger.error(f'Error auto-loading subjects: {str(e)}', exc_info=True)
        # Don't show error to user, just log it
    
    subjects = RSubject.query.filter_by(session_id=session_id).all()
    
    # Get batches for dropdown (for manual addition if needed)
    batches = []
    try:
        from blueprints.student_management.models import Student
        if Student:
            all_batches = db.session.query(Student.batch).distinct().filter(Student.batch.isnot(None)).order_by(Student.batch.desc()).all()
            batches = [batch[0] for batch in all_batches]
    except:
        batches = []
    
    return render_template('rm_add_subject.html', session=session, subjects=subjects, batches=batches)

@result_management_bp.route('/api/curricula', methods=['GET'])
@login_required
def get_curricula_for_batch():
    """Get curricula applicable to a specific batch (AJAX)"""
    try:
        from blueprints.course_management.models import Curriculum
        if not Curriculum:
            return jsonify({'success': False, 'message': 'Curriculum Management module not available'}), 503
        
        batch = request.args.get('batch', '').strip()
        if not batch:
            return jsonify({'success': False, 'message': 'Batch is required'}), 400
        
        all_curricula = Curriculum.query.all()
        applicable_curricula = []
        
        normalized_batch = str(batch).strip()
        
        for curriculum in all_curricula:
            batches_list = curriculum.get_batches_list()
            for b in batches_list:
                if str(b).strip() == normalized_batch:
                    applicable_curricula.append({
                        'id': curriculum.id,
                        'name': curriculum.name,
                        'date': curriculum.date
                    })
                    break
        
        return jsonify({
            'success': True,
            'curricula': applicable_curricula
        })
    except Exception as e:
        current_app.logger.error(f"Error fetching curricula for batch: {e}", exc_info=True)
        return jsonify({'success': False, 'message': f'Error loading curricula: {str(e)}'}), 500

@result_management_bp.route('/refresh_marks/<int:session_id>', methods=['POST'])
@login_required
def refresh_marks(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        return jsonify({'success': False, 'message': 'You do not have access to this session.'}), 403
    """Refresh marks from Class Management and Exam Paper Evaluation"""
    try:
        subject_id = request.args.get('subject_id', type=int)
        if not subject_id:
            return jsonify({'success': False, 'message': 'Subject ID is required'}), 400
        
        selected_subject = RSubject.query.get_or_404(subject_id)
        
        # Get students from Course Management (same logic as add_marks)
        # This ensures we refresh marks for all registered students, not just those in RCourseRegistration
        from blueprints.course_management.models import StudentCourseRegistration
        from blueprints.student_management.models import Student as StudentProfile
        
        # Get all RStudents in this session
        all_rstudents = RStudent.query.filter_by(session_id=session_id).all()
        student_id_to_rstudent_id = {rs.student_id: rs.id for rs in all_rstudents}
        
        # Get Student profiles for these student_ids
        student_profiles = StudentProfile.query.filter(
            StudentProfile.student_id.in_(student_id_to_rstudent_id.keys())
        ).all()
        student_profile_id_to_rstudent_id = {
            profile.id: student_id_to_rstudent_id.get(profile.student_id)
            for profile in student_profiles
            if profile.student_id in student_id_to_rstudent_id
        }
        
        # Find registered students from Course Management
        # Get ALL registrations for this course_code (regardless of year/term)
        # This ensures students who deregistered and re-registered are still included
        if not student_profile_id_to_rstudent_id:
            students = []
        else:
            course_filters = [
                StudentCourseRegistration.student_id.in_(student_profile_id_to_rstudent_id.keys()),
                StudentCourseRegistration.course_code == selected_subject.code,
                StudentCourseRegistration.status.in_(['finalized', 'pending', 'archived'])
            ]
            
            # Get all registrations for this course_code
            registered_regs = StudentCourseRegistration.query.filter(*course_filters).all()
            current_app.logger.debug(f'Refresh marks: Found {len(registered_regs)} total registrations for course_code={selected_subject.code}')
            
            # Map registration student_id (Student profile ID) back to RStudent.id
            registered_rstudent_ids = set()
            for reg in registered_regs:
                rstudent_id = student_profile_id_to_rstudent_id.get(reg.student_id)
                if rstudent_id:
                    registered_rstudent_ids.add(rstudent_id)
            
            if registered_rstudent_ids:
                students = RStudent.query.filter(RStudent.id.in_(registered_rstudent_ids)).order_by(RStudent.student_id).all()
                current_app.logger.debug(f'Refresh marks: Loaded {len(students)} students from {len(registered_regs)} registrations')
            else:
                students = []
        
        updated_count = 0
        error_count = 0
        
        for student in students:
            try:
                # Get or create mark
                mark = RMark.query.filter_by(student_id=student.id, subject_id=subject_id).first()
                if not mark:
                    mark = RMark(student_id=student.id, subject_id=subject_id)
                    db.session.add(mark)
                
                # Import from Class Management
                try:
                    from blueprints.class_management.models import Session, ClassStudent
                    
                    # Find matching session in Class Management by course code
                    # Try multiple matching strategies:
                    # 1. Exact match: course_code + year + term + academic_session
                    # 2. Partial match: course_code + year + term (ignore academic_session)
                    # 3. Fallback: course_code only (for backward compatibility)
                    class_session = None
                    
                    if session.year and session.term and session.name:
                        class_session = Session.query.filter_by(
                            course_code=selected_subject.code,
                            year=session.year,
                            term=session.term,
                            academic_session=session.name
                        ).first()
                    
                    if not class_session and session.year and session.term:
                        class_session = Session.query.filter_by(
                            course_code=selected_subject.code,
                            year=session.year,
                            term=session.term
                        ).first()
                    
                    if not class_session:
                        class_session = Session.query.filter_by(
                            course_code=selected_subject.code
                        ).order_by(Session.created_at.desc()).first()  # Get most recent if multiple
                    
                    if class_session:
                        # Find student in Class Management
                        class_student = ClassStudent.query.filter_by(
                            session_id=class_session.id,
                            student_id=student.student_id
                        ).first()
                        
                        if class_student:
                            # Import attendance marks
                            # Only auto-populate if attendance_manual is False or None (not manually entered)
                            if not mark.attendance_manual:
                                try:
                                    from blueprints.class_management.routes import _build_attendance_summary
                                    
                                    attendance_summary = _build_attendance_summary(class_session)
                                    student_attendance = attendance_summary.get('per_student', {}).get(student.student_id, {})
                                    
                                    if student_attendance and 'marks' in student_attendance:
                                        attendance_marks = student_attendance.get('marks', 0)
                                        if attendance_marks is not None:
                                            mark.attendance = float(attendance_marks)
                                        else:
                                            mark.attendance = None
                                    else:
                                        mark.attendance = None
                                except Exception as e:
                                    current_app.logger.error(f"Error importing attendance for student {student.student_id}: {e}", exc_info=True)
                            
                            # Import marks based on subject type
                            if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                                # Import continuous assessment for Theory courses
                                if class_student.assessment_total_40 is not None:
                                    # PG: Already rounded in assessment calculation
                                    mark.continuous_assessment = float(class_student.assessment_total_40)
                                elif class_student.assessment_total is not None:
                                    assessment_total = float(class_student.assessment_total)
                                    if assessment_total <= 40:
                                        # UG: Round for result generation
                                        mark.continuous_assessment = round(assessment_total)
                                    else:
                                        # Scale from 30 to 40 if needed, then round for result generation
                                        mark.continuous_assessment = round(min(40.0, (assessment_total / 30) * 40))
                                else:
                                    mark.continuous_assessment = None
                            
                            elif selected_subject.subject_type == 'Sessional':
                                # Import sessional report and viva for Sessional courses
                                if class_student.sessional_report is not None:
                                    mark.sessional_report = float(class_student.sessional_report)
                                else:
                                    mark.sessional_report = None
                                
                                if class_student.sessional_viva is not None:
                                    mark.sessional_viva = float(class_student.sessional_viva)
                                else:
                                    mark.sessional_viva = None
                        else:
                            # class_student not found - try to import attendance from Class Management anyway
                            # This ensures attendance and assessment are imported even if student is not in Class Management session
                            if not mark.attendance_manual:
                                try:
                                    from blueprints.class_management.routes import _build_attendance_summary
                                    
                                    attendance_summary = _build_attendance_summary(class_session)
                                    student_attendance = attendance_summary.get('per_student', {}).get(student.student_id, {})
                                    
                                    if student_attendance and 'marks' in student_attendance:
                                        attendance_marks = student_attendance.get('marks', 0)
                                        if attendance_marks is not None:
                                            mark.attendance = float(attendance_marks)
                                except Exception as e:
                                    current_app.logger.debug(f"Could not import attendance for student {student.student_id} (not in Class Management): {e}")
                            
                            # Try to find student in other Class Management sessions for this course
                            if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                                if mark.continuous_assessment is None:
                                    # Try to find student in any Class Management session for this course
                                    alternative_class_student = ClassStudent.query.join(
                                        Session, ClassStudent.session_id == Session.id
                                    ).filter(
                                        Session.course_code == selected_subject.code,
                                        ClassStudent.student_id == student.student_id
                                    ).first()
                                    
                                    if alternative_class_student:
                                        if alternative_class_student.assessment_total_40 is not None:
                                            mark.continuous_assessment = float(alternative_class_student.assessment_total_40)
                                        elif alternative_class_student.assessment_total is not None:
                                            assessment_total = float(alternative_class_student.assessment_total)
                                            if assessment_total <= 40:
                                                mark.continuous_assessment = round(assessment_total)
                                            else:
                                                mark.continuous_assessment = round(min(40.0, (assessment_total / 30) * 40))
                except Exception as e:
                    current_app.logger.error(f"Error importing from Class Management for student {student.student_id}: {e}", exc_info=True)
                    # Continue even if Class Management import fails
                
                # Import Section A and B from Exam Paper Evaluation (only for Theory courses)
                # SIMPLE DIRECT APPROACH - No complex logic, just get the data
                if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                    try:
                        from blueprints.class_management.models import ExamPaperEvaluation
                        import json
                        
                        # Find ExamPaperEvaluation - try course_code first, then course_name
                        exam_entry = ExamPaperEvaluation.query.filter_by(
                            course_code=selected_subject.code,
                            archived=False
                        ).first()
                        
                        if not exam_entry:
                            exam_entry = ExamPaperEvaluation.query.filter(
                                ExamPaperEvaluation.course_name.ilike(f'%{selected_subject.name}%'),
                                ExamPaperEvaluation.archived == False
                            ).first()
                        
                        if exam_entry and exam_entry.marks_data:
                            try:
                                exam_data = json.loads(exam_entry.marks_data)
                                questions = exam_data.get('questions', [])
                                rows = exam_data.get('rows', [])
                                
                                # Find student in rows - try both student_id and code
                                student_row = None
                                student_id_str = str(student.student_id).strip()
                                
                                for row in rows:
                                    row_id = str(row.get('student_id', '')).strip()
                                    row_code = str(row.get('code', '')).strip()
                                    
                                    if row_id == student_id_str or row_code == student_id_str:
                                        student_row = row
                                        break
                                
                                if student_row:
                                    marks_dict = student_row.get('marks', {})
                                    
                                    # SIMPLE APPROACH: Look for ANY key containing "a" or "b" or "part" or "section"
                                    part_a_value = None
                                    part_b_value = None
                                    
                                    # First, try to find by question labels
                                    for question in questions:
                                        q_label = question.get('label', '')
                                        q_label_lower = q_label.lower()
                                        
                                        if ('a' in q_label_lower or 'part a' in q_label_lower or 'section a' in q_label_lower) and not part_a_value:
                                            q_marks = marks_dict.get(q_label, {})
                                            if isinstance(q_marks, dict):
                                                total = sum(float(v) for k, v in q_marks.items() if v and str(v).strip() and str(v).strip() != '')
                                                if total > 0:
                                                    part_a_value = min(25.0, total)
                                            elif isinstance(q_marks, (int, float)):
                                                if q_marks > 0:
                                                    part_a_value = min(25.0, float(q_marks))
                                        
                                        if ('b' in q_label_lower or 'part b' in q_label_lower or 'section b' in q_label_lower) and not part_b_value:
                                            q_marks = marks_dict.get(q_label, {})
                                            if isinstance(q_marks, dict):
                                                total = sum(float(v) for k, v in q_marks.items() if v and str(v).strip() and str(v).strip() != '')
                                                if total > 0:
                                                    part_b_value = min(25.0, total)
                                            elif isinstance(q_marks, (int, float)):
                                                if q_marks > 0:
                                                    part_b_value = min(25.0, float(q_marks))
                                    
                                    # If not found, try ALL keys in marks_dict
                                    if not part_a_value or not part_b_value:
                                        for key, value in marks_dict.items():
                                            key_lower = key.lower()
                                            
                                            if not part_a_value and ('a' in key_lower or 'part a' in key_lower or 'section a' in key_lower or key_lower.endswith(' a)')):
                                                if isinstance(value, dict):
                                                    total = sum(float(v) for k, v in value.items() if v and str(v).strip() and str(v).strip() != '')
                                                    if total > 0:
                                                        part_a_value = min(25.0, total)
                                                elif isinstance(value, (int, float, str)):
                                                    try:
                                                        val = float(value)
                                                        if val > 0:
                                                            part_a_value = min(25.0, val)
                                                    except:
                                                        pass
                                            
                                            if not part_b_value and ('b' in key_lower or 'part b' in key_lower or 'section b' in key_lower or key_lower.endswith(' b)')):
                                                if isinstance(value, dict):
                                                    total = sum(float(v) for k, v in value.items() if v and str(v).strip() and str(v).strip() != '')
                                                    if total > 0:
                                                        part_b_value = min(25.0, total)
                                                elif isinstance(value, (int, float, str)):
                                                    try:
                                                        val = float(value)
                                                        if val > 0:
                                                            part_b_value = min(25.0, val)
                                                    except:
                                                        pass
                                    
                                    # If still not found and we have exactly 2 questions, split total
                                    if not part_a_value and not part_b_value and len(questions) == 2:
                                        total_str = student_row.get('total', '')
                                        if total_str:
                                            try:
                                                total = float(total_str)
                                                part_a_value = total / 2
                                                part_b_value = total / 2
                                            except:
                                                pass
                                    
                                    # SET THE VALUES
                                    if part_a_value is not None:
                                        mark.part_a = part_a_value
                                        current_app.logger.info(f'✅ Set part_a={mark.part_a} for student {student.student_id}')
                                    
                                    if part_b_value is not None:
                                        mark.part_b = part_b_value
                                        current_app.logger.info(f'✅ Set part_b={mark.part_b} for student {student.student_id}')
                                    
                                    if part_a_value is None and part_b_value is None:
                                        current_app.logger.warning(f'⚠️ Could not find Part A/B for student {student.student_id}. marks_dict keys: {list(marks_dict.keys())}')
                                
                                else:
                                    current_app.logger.warning(f'⚠️ Student {student.student_id} not found in exam data rows')
                            
                            except Exception as e:
                                current_app.logger.error(f'Error processing exam marks: {e}', exc_info=True)
                        else:
                            if not exam_entry:
                                current_app.logger.warning(f'No ExamPaperEvaluation found for course {selected_subject.code}')
                            elif not exam_entry.marks_data:
                                current_app.logger.warning(f'ExamPaperEvaluation found but marks_data is empty')
                    
                    except Exception as e:
                        current_app.logger.error(f'Error fetching exam marks: {e}', exc_info=True)
                
                # Calculate total_marks, grade_point, and grade_letter
                try:
                    # Get is_retake status
                    registration = RCourseRegistration.query.filter_by(student_id=student.id, subject_id=subject_id).first()
                    is_retake = registration.is_retake if registration else False
                    mark.is_retake = is_retake
                    
                    total_marks = 0
                    if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                        total_marks = sum(filter(None, [mark.attendance, mark.continuous_assessment, mark.part_a, mark.part_b]))
                    elif selected_subject.subject_type == 'Sessional':
                        total_marks = sum(filter(None, [mark.attendance, mark.sessional_report, mark.sessional_viva]))
                    elif selected_subject.subject_type in ('Thesis (UG)', 'Thesis I (UG)', 'Thesis II (UG)'):
                        total_marks = sum(filter(None, [mark.attendance, mark.thesis_evaluation, mark.presentation]))
                    elif selected_subject.subject_type == 'Dissertation':
                        if selected_subject.dissertation_type == 'Type1':
                            total_marks = sum(filter(None, [mark.supervisor_assessment, mark.proposal_presentation]))
                        elif selected_subject.dissertation_type == 'Type2':
                            total_marks = sum(filter(None, [mark.supervisor_assessment, mark.project_report, mark.defense]))
                        else:
                            # Fallback for existing Dissertation subjects without type
                            total_marks = sum(filter(None, [mark.supervisor_assessment, mark.proposal_presentation, mark.project_report, mark.defense]))
                    elif selected_subject.subject_type == 'Viva':
                        total_marks = mark.viva or 0
                    
                    mark.total_marks = total_marks
                    mark.grade_point, mark.grade_letter = calculate_grade(total_marks, is_retake=is_retake)
                except Exception as e:
                    current_app.logger.error(f"Error calculating grades for student {student.student_id}: {e}", exc_info=True)
                    # Continue even if calculation fails
                
                updated_count += 1
                
            except Exception as e:
                current_app.logger.error(f"Error refreshing marks for student {student.student_id}: {e}", exc_info=True)
                error_count += 1
                continue
        
        db.session.commit()
        
        message = f'Successfully refreshed marks for {updated_count} student(s).'
        if error_count > 0:
            message += f' {error_count} student(s) had errors.'
        
        return jsonify({
            'success': True,
            'message': message,
            'updated_count': updated_count,
            'error_count': error_count
        })
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error refreshing marks: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error refreshing marks: {str(e)}'
        }), 500

@result_management_bp.route('/delete_subject/<int:subject_id>', methods=['POST'])
@login_required
def delete_subject(subject_id):
    """Delete a subject from result session"""
    try:
        subject = RSubject.query.get_or_404(subject_id)
        session_id = subject.session_id
        subject_name = subject.name
        
        # Cascade delete will handle marks and registrations
        # due to cascade="all, delete-orphan" in the model relationships
        db.session.delete(subject)
        db.session.commit()
        flash(f'Subject "{subject_name}" deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error deleting subject {subject_id}: {e}', exc_info=True)
        flash(f'Error deleting subject: {str(e)}', 'danger')
        # Try to get session_id for redirect even if deletion failed
        try:
            subject = RSubject.query.get(subject_id)
            session_id = subject.session_id if subject else None
        except:
            session_id = None
    
    if session_id:
        return redirect(url_for('result_management.add_subject', session_id=session_id))
    return redirect(url_for('result_management.index'))


def clear_exam_marks_from_result_management(exam_entry_id):
    """
    Clear Part A or Part B marks from Result Management when an exam entry is marked as incomplete.
    This function is called when a scrutinizer marks an exam entry as incomplete.
    
    Args:
        exam_entry_id: The ID of the ExamPaperEvaluation entry
        
    Returns:
        dict: {'success': bool, 'message': str, 'marks_cleared': int, 'errors': list}
    """
    # #region agent log
    import json as json_module
    try:
        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
            f.write(json_module.dumps({'location': 'clear_exam_marks_from_result_management:entry', 'message': 'Clear function called', 'data': {'exam_entry_id': exam_entry_id}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'M'}) + '\n')
    except: pass
    # #endregion
    
    result = {
        'success': False,
        'message': '',
        'marks_cleared': 0,
        'errors': []
    }
    
    try:
        # Get the exam entry
        exam_entry = ExamPaperEvaluation.query.get(exam_entry_id)
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'clear_exam_marks_from_result_management:exam_entry', 'message': 'Exam entry fetched', 'data': {'found': exam_entry is not None, 'academic_session': exam_entry.academic_session if exam_entry else None, 'year': exam_entry.year if exam_entry else None, 'term': exam_entry.term if exam_entry else None, 'course_code': exam_entry.course_code if exam_entry else None, 'section': exam_entry.section if exam_entry else None}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'M'}) + '\n')
        except: pass
        # #endregion
        
        if not exam_entry:
            result['message'] = f'Exam entry {exam_entry_id} not found'
            result['errors'].append(result['message'])
            return result
        
        # Validate required fields
        if not exam_entry.academic_session or not exam_entry.term:
            result['message'] = 'Exam entry missing required fields (academic_session, term)'
            result['errors'].append(result['message'])
            return result
        
        if not exam_entry.course_code:
            result['message'] = 'Exam entry missing course_code'
            result['errors'].append(result['message'])
            return result
        
        # Determine which part to clear based on section field
        clear_part_a = False
        clear_part_b = False
        
        if exam_entry.section:
            section_lower = str(exam_entry.section).lower().strip()
            if 'part a' in section_lower or section_lower == 'a' or section_lower.startswith('a '):
                clear_part_a = True
            elif 'part b' in section_lower or section_lower == 'b' or section_lower.startswith('b '):
                clear_part_b = True
        
        # If section is not specified, we cannot determine which part to clear
        # In this case, we'll skip clearing (safer approach - don't clear if unsure)
        if not clear_part_a and not clear_part_b:
            result['message'] = 'Cannot clear marks: exam entry section field is not specified'
            result['errors'].append(result['message'])
            return result
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'clear_exam_marks_from_result_management:clear_strategy', 'message': 'Clear strategy determined', 'data': {'section': exam_entry.section, 'clear_part_a': clear_part_a, 'clear_part_b': clear_part_b}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'M'}) + '\n')
        except: pass
        # #endregion
        
        # Find the corresponding RSession
        r_session = RSession.query.filter_by(
            name=exam_entry.academic_session,
            year=exam_entry.year,
            term=exam_entry.term
        ).first()
        
        if not r_session:
            result['message'] = f'No Result Management session found for {exam_entry.academic_session} {exam_entry.year} {exam_entry.term}'
            result['errors'].append(result['message'])
            # Don't return error - session might not exist yet, which is fine
            return result
        
        # Find the corresponding RSubject
        r_subject = RSubject.query.filter_by(
            code=exam_entry.course_code,
            session_id=r_session.id
        ).first()
        
        if not r_subject:
            result['message'] = f'No Result Management subject found for course {exam_entry.course_code}'
            result['errors'].append(result['message'])
            # Don't return error - subject might not exist yet, which is fine
            return result
        
        # Get all RMark entries for this subject
        r_marks = RMark.query.filter_by(subject_id=r_subject.id).all()
        
        # Clear marks based on section
        for r_mark in r_marks:
            updated = False
            if clear_part_a and r_mark.part_a is not None:
                r_mark.part_a = None
                updated = True
            if clear_part_b and r_mark.part_b is not None:
                r_mark.part_b = None
                updated = True
            
            if updated:
                result['marks_cleared'] += 1
        
        # Commit changes
        db.session.commit()
        
        result['success'] = True
        result['message'] = f'Successfully cleared marks: {result["marks_cleared"]} marks cleared'
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'clear_exam_marks_from_result_management:commit', 'message': 'Changes committed', 'data': {'marks_cleared': result['marks_cleared']}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'M'}) + '\n')
        except: pass
        # #endregion
        
    except Exception as e:
        db.session.rollback()
        error_msg = f'Error clearing marks for exam_entry {exam_entry_id}: {str(e)}'
        result['message'] = error_msg
        result['errors'].append(error_msg)
        current_app.logger.error(error_msg, exc_info=True)
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'clear_exam_marks_from_result_management:error', 'message': 'Error occurred', 'data': {'error': str(e)}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'M'}) + '\n')
        except: pass
        # #endregion
    
    return result


def sync_exam_marks_to_result_management(exam_entry_id):
    """
    Automatically sync Part A and Part B marks from Exam Paper Evaluation to Result Management.
    This function is called when a scrutinizer marks an exam entry as complete.
    
    Args:
        exam_entry_id: The ID of the ExamPaperEvaluation entry
        
    Returns:
        dict: {'success': bool, 'message': str, 'session_created': bool, 'subject_created': bool, 
               'students_created': int, 'marks_updated': int, 'errors': list}
    """
    # #region agent log
    import json as json_module
    try:
        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
            f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:entry', 'message': 'Sync function called', 'data': {'exam_entry_id': exam_entry_id}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'A'}) + '\n')
    except: pass
    # #endregion
    
    result = {
        'success': False,
        'message': '',
        'session_created': False,
        'subject_created': False,
        'students_created': 0,
        'marks_updated': 0,
        'errors': []
    }
    
    try:
        # Get the exam entry
        exam_entry = ExamPaperEvaluation.query.get(exam_entry_id)
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:exam_entry', 'message': 'Exam entry fetched', 'data': {'found': exam_entry is not None, 'academic_session': exam_entry.academic_session if exam_entry else None, 'year': exam_entry.year if exam_entry else None, 'term': exam_entry.term if exam_entry else None, 'course_code': exam_entry.course_code if exam_entry else None, 'section': exam_entry.section if exam_entry else None, 'has_marks_data': bool(exam_entry.marks_data) if exam_entry else False}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run3', 'hypothesisId': 'B'}) + '\n')
        except: pass
        # #endregion
        if not exam_entry:
            result['message'] = f'Exam entry {exam_entry_id} not found'
            result['errors'].append(result['message'])
            return result
        
        # Validate required fields
        if not exam_entry.academic_session or not exam_entry.term:
            result['message'] = 'Exam entry missing required fields: academic_session and term are required'
            result['errors'].append(result['message'])
            return result
        
        if not exam_entry.marks_data:
            result['message'] = 'Exam entry has no marks data'
            result['errors'].append(result['message'])
            return result
        
        # Step 1: Find or create RSession
        # Build query with proper handling of None year values
        query = RSession.query.filter_by(
            name=exam_entry.academic_session,
            term=exam_entry.term
        )
        if exam_entry.year:
            query = query.filter_by(year=exam_entry.year)
        else:
            query = query.filter(RSession.year.is_(None))
        r_session = query.first()
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:rsession', 'message': 'RSession lookup', 'data': {'found': r_session is not None, 'session_id': r_session.id if r_session else None, 'search_name': exam_entry.academic_session, 'search_year': exam_entry.year, 'search_term': exam_entry.term}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'C'}) + '\n')
        except: pass
        # #endregion
        
        if not r_session:
            # Auto-create RSession if Head of Discipline hasn't created it
            r_session = RSession(
                name=exam_entry.academic_session,
                year=exam_entry.year,
                term=exam_entry.term,
                batch=exam_entry.batch,
                curriculum_id=None,  # Can be set later by Head of Discipline
                is_archived=False
            )
            db.session.add(r_session)
            db.session.flush()  # Get the ID
            result['session_created'] = True
            current_app.logger.info(f'Auto-created RSession: {r_session.name} - Year {r_session.year} - Term {r_session.term}')
        
        # Step 2: Find or create RSubject
        r_subject = RSubject.query.filter_by(
            code=exam_entry.course_code,
            session_id=r_session.id
        ).first()
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:rsubject', 'message': 'RSubject lookup', 'data': {'found': r_subject is not None, 'subject_id': r_subject.id if r_subject else None, 'search_code': exam_entry.course_code, 'session_id': r_session.id}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'D'}) + '\n')
        except: pass
        # #endregion
        
        if not r_subject:
            # Try to get credit from Course model
            credit = 3.0  # Default
            course = Course.query.filter_by(course_code=exam_entry.course_code).first()
            if course:
                credit = course.credit
            
            # Determine subject type and dissertation type from course if available
            subject_type = 'Theory'  # Default
            dissertation_type = None
            if course:
                subject_type = _determine_subject_type(course)
                if course.course_type == 'Dissertation Proposal (PG)':
                    dissertation_type = 'Type1'
                elif course.course_type == 'Dissertation Defence (PG)':
                    dissertation_type = 'Type2'
            
            r_subject = RSubject(
                code=exam_entry.course_code,
                name=exam_entry.course_name,
                credit=credit,
                subject_type=subject_type,
                dissertation_type=dissertation_type,
                session_id=r_session.id
            )
            db.session.add(r_subject)
            db.session.flush()  # Get the ID
            result['subject_created'] = True
            current_app.logger.info(f'Auto-created RSubject: {r_subject.code} - {r_subject.name}')
        
        # Step 3: Parse marks_data JSON
        try:
            exam_marks = json.loads(exam_entry.marks_data)
            questions = exam_marks.get('questions', [])
            rows = exam_marks.get('rows', [])
            
            # #region agent log
            try:
                with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                    f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:parse_json', 'message': 'JSON parsed', 'data': {'questions_count': len(questions), 'rows_count': len(rows), 'question_labels': [q.get('label', '') for q in questions[:5]]}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'E'}) + '\n')
            except: pass
            # #endregion
        except json.JSONDecodeError as e:
            result['message'] = f'Error parsing marks_data JSON: {str(e)}'
            result['errors'].append(result['message'])
            return result
        
        if not rows:
            result['message'] = 'No student rows found in marks_data'
            result['errors'].append(result['message'])
            return result
        
        # Step 4: Process each student row
        for row_idx, row in enumerate(rows):
            try:
                student_id_str = str(row.get('student_id', '')).strip()
                if not student_id_str:
                    # Skip rows without student_id
                    continue
                
                # #region agent log
                try:
                    with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                        f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:student_row', 'message': 'Processing student row', 'data': {'row_idx': row_idx, 'student_id': student_id_str, 'has_marks': 'marks' in row, 'marks_keys': list(row.get('marks', {}).keys()) if isinstance(row.get('marks'), dict) else None}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'F'}) + '\n')
                except: pass
                # #endregion
                
                # Find or create RStudent
                r_student = RStudent.query.filter_by(
                    student_id=student_id_str,
                    session_id=r_session.id
                ).first()
                
                if not r_student:
                    # Try to get student name from Student model
                    student_name = 'Unknown Student'
                    student_profile = StudentProfile.query.filter_by(student_id=student_id_str).first()
                    if student_profile:
                        student_name = student_profile.name
                    
                    r_student = RStudent(
                        student_id=student_id_str,
                        name=student_name,
                        session_id=r_session.id,
                        year=exam_entry.year,
                        discipline=exam_entry.discipline,
                        school=exam_entry.school
                    )
                    db.session.add(r_student)
                    db.session.flush()  # Get the ID
                    result['students_created'] += 1
                
                # Find or create RMark
                r_mark = RMark.query.filter_by(
                    student_id=r_student.id,
                    subject_id=r_subject.id
                ).first()
                
                if not r_mark:
                    r_mark = RMark(
                        student_id=r_student.id,
                        subject_id=r_subject.id
                    )
                    db.session.add(r_mark)
                
                # CRITICAL: Store existing Part A/B marks before processing
                # If this entry is for Part B, we should NOT overwrite existing Part A marks
                # If this entry is for Part A, we should NOT overwrite existing Part B marks
                existing_part_a = r_mark.part_a
                existing_part_b = r_mark.part_b
                
                # #region agent log
                try:
                    with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                        f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:existing_marks', 'message': 'Existing marks before sync', 'data': {'student_id': student_id_str, 'existing_part_a': existing_part_a, 'existing_part_b': existing_part_b}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run7', 'hypothesisId': 'L'}) + '\n')
                except: pass
                # #endregion
                
                # Extract Part A and Part B marks (reuse logic from refresh_marks)
                marks_dict = row.get('marks', {})
                section_a_found = False
                section_b_found = False
                
                # First, check ExamPaperEvaluation.section field to determine Part A/B
                # If section indicates Part A, all questions go to Part A
                # If section indicates Part B, all questions go to Part B
                section_field_assigned = False  # Track if section field assigned anything
                if exam_entry.section:
                    section_lower = str(exam_entry.section).lower().strip()
                    
                    # #region agent log
                    try:
                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                            f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:section_check', 'message': 'Checking section field', 'data': {'student_id': student_id_str, 'section_original': exam_entry.section, 'section_lower': section_lower}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run4', 'hypothesisId': 'I'}) + '\n')
                    except: pass
                    # #endregion
                    
                    # Use strict matching: check for "part a" or "parta" (not just 'a' which matches both A and B)
                    if 'part a' in section_lower or section_lower == 'a' or section_lower.startswith('a '):
                        # All questions in this entry are Part A marks
                        part_a_total = 0.0
                        for q_label in marks_dict.keys():
                            q_marks = marks_dict.get(q_label, {})
                            if isinstance(q_marks, dict):
                                for part_key, part_value in q_marks.items():
                                    try:
                                        val = float(part_value) if part_value and str(part_value).strip() else 0.0
                                        part_a_total += val
                                    except (ValueError, TypeError):
                                        pass
                            elif isinstance(q_marks, (int, float, str)):
                                try:
                                    val = float(q_marks) if q_marks and str(q_marks).strip() else 0.0
                                    part_a_total += val
                                except (ValueError, TypeError):
                                    pass
                        if part_a_total > 0:
                            r_mark.part_a = min(25.0, part_a_total)
                            # CRITICAL: If Part A is assigned, ensure Part B is NOT overwritten (keep existing)
                            # BUT: Only preserve if Part B was NOT set from a previous Part A sync
                            # If existing Part A equals existing Part B, it might be from a wrong sync, so don't preserve
                            if existing_part_b is not None and existing_part_a != existing_part_b:
                                r_mark.part_b = existing_part_b
                            else:
                                # If Part A equals Part B, it was likely set incorrectly, so clear it
                                r_mark.part_b = None
                            section_a_found = True
                            section_field_assigned = True
                            
                            # #region agent log
                            try:
                                with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                    f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:part_a_assigned', 'message': 'Part A assigned from section', 'data': {'student_id': student_id_str, 'section': exam_entry.section, 'part_a_total': part_a_total, 'part_a_final': r_mark.part_a, 'preserved_part_b': r_mark.part_b}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run4', 'hypothesisId': 'I'}) + '\n')
                            except: pass
                            # #endregion
                    # Use strict matching: check for "part b" or "partb" (not just 'b' which might match other things)
                    elif 'part b' in section_lower or section_lower == 'b' or section_lower.startswith('b '):
                        # All questions in this entry are Part B marks
                        part_b_total = 0.0
                        for q_label in marks_dict.keys():
                            q_marks = marks_dict.get(q_label, {})
                            if isinstance(q_marks, dict):
                                for part_key, part_value in q_marks.items():
                                    try:
                                        val = float(part_value) if part_value and str(part_value).strip() else 0.0
                                        part_b_total += val
                                    except (ValueError, TypeError):
                                        pass
                            elif isinstance(q_marks, (int, float, str)):
                                try:
                                    val = float(q_marks) if q_marks and str(q_marks).strip() else 0.0
                                    part_b_total += val
                                except (ValueError, TypeError):
                                    pass
                        if part_b_total > 0:
                            r_mark.part_b = min(25.0, part_b_total)
                            # CRITICAL: If Part B is assigned, ensure Part A is NOT overwritten (keep existing)
                            # BUT: Only preserve if Part A was NOT set from a previous Part B sync
                            # If existing Part A equals existing Part B, it might be from a wrong sync, so don't preserve
                            if existing_part_a is not None and existing_part_a != existing_part_b:
                                r_mark.part_a = existing_part_a
                            else:
                                # If Part A equals Part B, it was likely set incorrectly, so clear it
                                r_mark.part_a = None
                            section_b_found = True
                            section_field_assigned = True
                            
                            # #region agent log
                            try:
                                with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                    f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:part_b_assigned', 'message': 'Part B assigned from section', 'data': {'student_id': student_id_str, 'section': exam_entry.section, 'part_b_total': part_b_total, 'part_b_final': r_mark.part_b, 'preserved_part_a': r_mark.part_a}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run4', 'hypothesisId': 'I'}) + '\n')
                            except: pass
                            # #endregion
                
                # #region agent log
                try:
                    # Log detailed structure of questions and marks
                    questions_detail = []
                    for q in questions:
                        q_detail = {'label': q.get('label', ''), 'parts': []}
                        if 'parts' in q:
                            for p in q.get('parts', []):
                                q_detail['parts'].append({'label': p.get('label', '')})
                        questions_detail.append(q_detail)
                    
                    marks_detail = {}
                    for k, v in marks_dict.items():
                        if isinstance(v, dict):
                            marks_detail[k] = {'type': 'dict', 'keys': list(v.keys()), 'sample_values': {k2: v2 for k2, v2 in list(v.items())[:2]}}
                        else:
                            marks_detail[k] = {'type': type(v).__name__, 'value': str(v)[:50]}
                    
                    with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                        f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:marks_extraction_start', 'message': 'Starting marks extraction with detailed structure', 'data': {'student_id': student_id_str, 'questions_detail': questions_detail, 'marks_detail': marks_detail, 'total': row.get('total')}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run2', 'hypothesisId': 'G'}) + '\n')
                except Exception as log_err:
                    try:
                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                            f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:marks_extraction_start', 'message': 'Logging error', 'data': {'error': str(log_err)}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run2', 'hypothesisId': 'G'}) + '\n')
                    except: pass
                # #endregion
                
                # Check question labels for Section A/B (ONLY if section field didn't assign anything)
                # CRITICAL: Skip this entire section if section field already assigned Part A or Part B
                if not section_field_assigned:
                    # #region agent log
                    try:
                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                            f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:check_question_labels', 'message': 'Checking question labels (section_field_assigned=False)', 'data': {'student_id': student_id_str, 'section_field_assigned': section_field_assigned}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run6', 'hypothesisId': 'K'}) + '\n')
                    except: pass
                    # #endregion
                    for question in questions:
                        question_label = question.get('label', '').lower()
                        
                        if ('section a' in question_label or 'part a' in question_label) and not section_a_found:
                            question_marks = marks_dict.get(question.get('label', ''), {})
                            if isinstance(question_marks, dict):
                                section_a_total = sum(float(v) for k, v in question_marks.items() if v and str(v).strip())
                                if section_a_total > 0:
                                    r_mark.part_a = min(25.0, section_a_total)
                                    section_a_found = True
                                    # #region agent log
                                    try:
                                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                            f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:part_a_from_labels', 'message': 'Part A assigned from question labels', 'data': {'student_id': student_id_str, 'question_label': question.get('label'), 'part_a': r_mark.part_a}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run6', 'hypothesisId': 'K'}) + '\n')
                                    except: pass
                                    # #endregion
                            elif isinstance(question_marks, (int, float)):
                                r_mark.part_a = min(25.0, float(question_marks))
                                section_a_found = True
                                # #region agent log
                                try:
                                    with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                        f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:part_a_from_labels', 'message': 'Part A assigned from question labels', 'data': {'student_id': student_id_str, 'question_label': question.get('label'), 'part_a': r_mark.part_a}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run6', 'hypothesisId': 'K'}) + '\n')
                                except: pass
                                # #endregion
                        
                        elif ('section b' in question_label or 'part b' in question_label) and not section_b_found:
                            question_marks = marks_dict.get(question.get('label', ''), {})
                            if isinstance(question_marks, dict):
                                section_b_total = sum(float(v) for k, v in question_marks.items() if v and str(v).strip())
                                if section_b_total > 0:
                                    r_mark.part_b = min(25.0, section_b_total)
                                    section_b_found = True
                                    # #region agent log
                                    try:
                                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                            f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:part_b_from_labels', 'message': 'Part B assigned from question labels', 'data': {'student_id': student_id_str, 'question_label': question.get('label'), 'part_b': r_mark.part_b}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run6', 'hypothesisId': 'K'}) + '\n')
                                    except: pass
                                    # #endregion
                            elif isinstance(question_marks, (int, float)):
                                r_mark.part_b = min(25.0, float(question_marks))
                                section_b_found = True
                                # #region agent log
                                try:
                                    with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                        f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:part_b_from_labels', 'message': 'Part B assigned from question labels', 'data': {'student_id': student_id_str, 'question_label': question.get('label'), 'part_b': r_mark.part_b}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run6', 'hypothesisId': 'K'}) + '\n')
                                except: pass
                                # #endregion
                else:
                    # #region agent log
                    try:
                        with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                            f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:skip_question_labels', 'message': 'Skipping question labels check (section_field_assigned=True)', 'data': {'student_id': student_id_str, 'section_field_assigned': section_field_assigned}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run6', 'hypothesisId': 'K'}) + '\n')
                    except: pass
                    # #endregion
                
                # Alternative patterns (only if section field didn't assign anything)
                # Skip if section field already assigned Part A or Part B
                if not section_field_assigned and (not section_a_found or not section_b_found):
                    for question_label, question_data in marks_dict.items():
                        q_label_lower = question_label.lower()
                        
                        if not section_a_found and ('section a' in q_label_lower or 'part a' in q_label_lower or 'a)' in q_label_lower):
                            if isinstance(question_data, dict):
                                section_a_total = sum(float(v) for k, v in question_data.items() if v and str(v).strip())
                                if section_a_total > 0:
                                    r_mark.part_a = min(25.0, section_a_total)
                                    section_a_found = True
                            elif isinstance(question_data, (int, float, str)):
                                try:
                                    val = float(question_data)
                                    if val > 0:
                                        r_mark.part_a = min(25.0, val)
                                        section_a_found = True
                                except (ValueError, TypeError):
                                    pass
                        
                        if not section_b_found and ('section b' in q_label_lower or 'part b' in q_label_lower or 'b)' in q_label_lower):
                            if isinstance(question_data, dict):
                                section_b_total = sum(float(v) for k, v in question_data.items() if v and str(v).strip())
                                if section_b_total > 0:
                                    r_mark.part_b = min(25.0, section_b_total)
                                    section_b_found = True
                            elif isinstance(question_data, (int, float, str)):
                                try:
                                    val = float(question_data)
                                    if val > 0:
                                        r_mark.part_b = min(25.0, val)
                                        section_b_found = True
                                except (ValueError, TypeError):
                                    pass
                
                # If no sections found, try alternative strategies
                # IMPORTANT: Only execute if section field didn't assign anything
                # If section field assigned Part A or Part B, don't use fallback strategies
                
                # #region agent log
                try:
                    with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                        f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:fallback_check', 'message': 'Checking fallback strategy conditions', 'data': {'student_id': student_id_str, 'section_field_assigned': section_field_assigned, 'section_a_found': section_a_found, 'section_b_found': section_b_found, 'will_execute_fallback': not section_field_assigned and not section_a_found and not section_b_found}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run5', 'hypothesisId': 'J'}) + '\n')
                except: pass
                # #endregion
                
                if not section_field_assigned and not section_a_found and not section_b_found:
                    # Strategy 1: If 2 questions, split total equally
                    if len(questions) == 2:
                        total_marks_str = row.get('total', '')
                        if total_marks_str:
                            try:
                                total = float(total_marks_str)
                                r_mark.part_a = min(25.0, total / 2)
                                r_mark.part_b = min(25.0, total / 2)
                                section_a_found = True
                                section_b_found = True
                            except (ValueError, TypeError):
                                pass
                    # Strategy 2: If 4 questions, assume Questions 1-2 = Part A, Questions 3-4 = Part B
                    elif len(questions) == 4:
                        # #region agent log
                        try:
                            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                                f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:fallback_4q', 'message': 'Executing 4-question fallback strategy', 'data': {'student_id': student_id_str}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run5', 'hypothesisId': 'J'}) + '\n')
                        except: pass
                        # #endregion
                        try:
                            part_a_total = 0.0
                            part_b_total = 0.0
                            
                            # Sum Questions 1 and 2 for Part A
                            for i in [0, 1]:
                                if i < len(questions):
                                    q_label = questions[i].get('label', '')
                                    q_marks = marks_dict.get(q_label, {})
                                    if isinstance(q_marks, dict):
                                        # Sum all parts of this question
                                        for part_key, part_value in q_marks.items():
                                            try:
                                                val = float(part_value) if part_value and str(part_value).strip() else 0.0
                                                part_a_total += val
                                            except (ValueError, TypeError):
                                                pass
                                    elif isinstance(q_marks, (int, float, str)):
                                        try:
                                            val = float(q_marks) if q_marks and str(q_marks).strip() else 0.0
                                            part_a_total += val
                                        except (ValueError, TypeError):
                                            pass
                            
                            # Sum Questions 3 and 4 for Part B
                            for i in [2, 3]:
                                if i < len(questions):
                                    q_label = questions[i].get('label', '')
                                    q_marks = marks_dict.get(q_label, {})
                                    if isinstance(q_marks, dict):
                                        # Sum all parts of this question
                                        for part_key, part_value in q_marks.items():
                                            try:
                                                val = float(part_value) if part_value and str(part_value).strip() else 0.0
                                                part_b_total += val
                                            except (ValueError, TypeError):
                                                pass
                                    elif isinstance(q_marks, (int, float, str)):
                                        try:
                                            val = float(q_marks) if q_marks and str(q_marks).strip() else 0.0
                                            part_b_total += val
                                        except (ValueError, TypeError):
                                            pass
                            
                            if part_a_total > 0:
                                r_mark.part_a = min(25.0, part_a_total)
                                section_a_found = True
                            if part_b_total > 0:
                                r_mark.part_b = min(25.0, part_b_total)
                                section_b_found = True
                        except Exception as e:
                            current_app.logger.error(f'Error in 4-question split strategy: {str(e)}', exc_info=True)
                
                # Note: ExamPaperEvaluation.section field is informational
                # We sync both Part A and Part B marks regardless of section field
                # The section field might indicate which part this entry represents,
                # but we still sync all available marks
                
                # #region agent log
                try:
                    with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                        f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:marks_extraction_end', 'message': 'Marks extraction completed', 'data': {'student_id': student_id_str, 'part_a': r_mark.part_a, 'part_b': r_mark.part_b, 'section_a_found': section_a_found, 'section_b_found': section_b_found}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'G'}) + '\n')
                except: pass
                # #endregion
                
                result['marks_updated'] += 1
                
            except Exception as e:
                error_msg = f'Error processing student row {row.get("student_id", "unknown")}: {str(e)}'
                result['errors'].append(error_msg)
                current_app.logger.error(error_msg, exc_info=True)
                continue
        
        # Commit all changes
        db.session.commit()
        
        # #region agent log
        try:
            with open('/Users/isckra/Documents/App Projects/Academic Management System/.cursor/debug.log', 'a') as f:
                f.write(json_module.dumps({'location': 'sync_exam_marks_to_result_management:commit', 'message': 'Changes committed', 'data': {'marks_updated': result['marks_updated'], 'students_created': result['students_created'], 'session_created': result['session_created'], 'subject_created': result['subject_created']}, 'timestamp': int(__import__('time').time() * 1000), 'sessionId': 'debug-session', 'runId': 'run1', 'hypothesisId': 'H'}) + '\n')
        except: pass
        # #endregion
        
        result['success'] = True
        result['message'] = f'Successfully synced marks: {result["marks_updated"]} marks updated'
        if result['session_created']:
            result['message'] += f', session created'
        if result['subject_created']:
            result['message'] += f', subject created'
        if result['students_created'] > 0:
            result['message'] += f', {result["students_created"]} students created'
        
        current_app.logger.info(f'Sync completed for exam_entry {exam_entry_id}: {result["message"]}')
        
    except Exception as e:
        db.session.rollback()
        result['message'] = f'Error syncing exam marks: {str(e)}'
        result['errors'].append(result['message'])
        current_app.logger.error(f'Error syncing exam marks for entry {exam_entry_id}: {str(e)}', exc_info=True)
    
    return result

@result_management_bp.route('/add_marks/<int:session_id>', methods=['GET', 'POST'])
@login_required
def add_marks(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        flash('You do not have access to this session.', 'danger')
        return redirect(url_for('result_management.index'))
    
    # Get all subjects for this session, ordered by code for consistent display
    subjects = RSubject.query.filter_by(session_id=session_id).order_by(RSubject.code).all()
    current_app.logger.info(f'Result Management: Found {len(subjects)} subjects for session {session_id} ({session.name})')
    
    selected_subject_id = request.args.get('subject_id', type=int)
    selected_subject = RSubject.query.get(selected_subject_id) if selected_subject_id else None
    
    if selected_subject:
        current_app.logger.info(f'Result Management: Selected subject {selected_subject.code} ({selected_subject.name}) for session {session_id}')

    # Only show students who are registered for the selected subject in Course Management
    if selected_subject:
        try:
            from blueprints.course_management.models import StudentCourseRegistration
            from blueprints.student_management.models import Student as StudentProfile
            
            # Get all RStudents in this session
            all_rstudents = RStudent.query.filter_by(session_id=session_id).all()
            student_id_to_rstudent_id = {rs.student_id: rs.id for rs in all_rstudents}
            existing_student_ids = set(student_id_to_rstudent_id.keys())
            
            # Get ALL registrations for this course_code first (to find missing RStudents)
            all_registrations = StudentCourseRegistration.query.filter(
                StudentCourseRegistration.course_code == selected_subject.code,
                StudentCourseRegistration.status.in_(['finalized', 'pending', 'archived'])
            ).all()
            
            # Auto-create RStudent records for registered students who don't exist yet
            students_to_create = []
            for reg in all_registrations:
                # Get StudentProfile to get student_id (public ID)
                student_profile = StudentProfile.query.get(reg.student_id)
                if not student_profile:
                    current_app.logger.warning(f'StudentProfile not found for registration: reg.student_id={reg.student_id}, course_code={reg.course_code}')
                    continue
                
                # Check if RStudent already exists for this session
                if student_profile.student_id not in existing_student_ids:
                    # Create new RStudent record
                    r_student = RStudent(
                        student_id=student_profile.student_id,
                        name=student_profile.name,
                        session_id=session_id,
                        year=session.year if session.year else None,
                        discipline=None,
                        school=None
                    )
                    students_to_create.append(r_student)
                    existing_student_ids.add(student_profile.student_id)
                    current_app.logger.debug(f'Auto-creating RStudent: student_id={student_profile.student_id}, name={student_profile.name}')
            
            # Bulk create missing RStudent records
            if students_to_create:
                try:
                    db.session.bulk_save_objects(students_to_create)
                    db.session.commit()
                    current_app.logger.info(f'Auto-created {len(students_to_create)} RStudent records for session {session_id}')
                    
                    # Refresh the mapping after creating new students
                    all_rstudents = RStudent.query.filter_by(session_id=session_id).all()
                    student_id_to_rstudent_id = {rs.student_id: rs.id for rs in all_rstudents}
                except Exception as e:
                    current_app.logger.error(f'Error auto-creating RStudent records: {e}', exc_info=True)
                    db.session.rollback()
            
            # Get Student profiles for these student_ids
            student_profiles = StudentProfile.query.filter(
                StudentProfile.student_id.in_(student_id_to_rstudent_id.keys())
            ).all()
            student_profile_id_to_rstudent_id = {
                profile.id: student_id_to_rstudent_id.get(profile.student_id)
                for profile in student_profiles
                if profile.student_id in student_id_to_rstudent_id
            }
            
            # Find registered students from Course Management
            # Only show students who have finalized registration for this specific course
            if not student_profile_id_to_rstudent_id:
                # No student profiles found - cannot match registrations
                current_app.logger.warning(f'No student profiles found for session {session_id}. Cannot load registered students.')
                students = []
            else:
                # Build filters - make matching more flexible
                # Include 'finalized', 'pending', and 'archived' statuses
                course_filters = [
                    StudentCourseRegistration.student_id.in_(student_profile_id_to_rstudent_id.keys()),
                    StudentCourseRegistration.course_code == selected_subject.code,
                    StudentCourseRegistration.status.in_(['finalized', 'pending', 'archived'])
                ]
                
                # Get ALL registrations for this course_code (regardless of year/term)
                # This ensures students who deregistered and re-registered are still included
                registered_regs = StudentCourseRegistration.query.filter(*course_filters).all()
                current_app.logger.debug(f'Found {len(registered_regs)} total registrations for course_code={selected_subject.code}')
                
                # If we have year/term, prefer exact matches but don't exclude others
                # This handles cases where students re-registered with different year/term
                if session.year and session.term and registered_regs:
                    # Count exact matches for logging
                    exact_matches = [reg for reg in registered_regs 
                                   if reg.year == session.year and reg.term == session.term]
                    current_app.logger.debug(f'Exact match (year={session.year}, term={session.term}): {len(exact_matches)} out of {len(registered_regs)} registrations')
                
                # Map registration student_id (Student profile ID) back to RStudent.id
                registered_rstudent_ids = set()
                for reg in registered_regs:
                    rstudent_id = student_profile_id_to_rstudent_id.get(reg.student_id)
                    if rstudent_id:
                        registered_rstudent_ids.add(rstudent_id)
                    else:
                        current_app.logger.warning(f'Registration found but RStudent mapping failed: reg.student_id={reg.student_id}, course_code={reg.course_code}, status={reg.status}')
                
                if registered_rstudent_ids:
                    students = RStudent.query.filter(RStudent.id.in_(registered_rstudent_ids)).order_by(RStudent.student_id).all()
                    current_app.logger.info(f'Loaded {len(students)} registered students for subject {selected_subject.code} (from {len(registered_regs)} registrations)')
                else:
                    # No registered students found for this subject
                    current_app.logger.warning(f'No registered students found for subject {selected_subject.code} (code={selected_subject.code}, session={session.name}/{session.year}/{session.term})')
                    students = []
        except Exception as e:
            current_app.logger.error(f'Error loading registered students: {str(e)}', exc_info=True)
            # Fallback to empty list if error
            students = []
    else:
        # No subject selected - show empty list (subject must be selected)
        students = []
        current_app.logger.debug(f'Result Management: No subject selected, students list is empty')
    
    current_app.logger.info(f'Result Management: Loaded {len(students)} students for subject {selected_subject.code if selected_subject else "N/A"}')
    
    marks_data = {}
    registrations_data = {} # To store retake status from Course Management
    if selected_subject:
        # Get retake status from Course Management's StudentCourseRegistration
        try:
            from blueprints.course_management.models import StudentCourseRegistration
            from blueprints.student_management.models import Student as StudentProfile
            
            # Get all RStudents in this session
            all_rstudents = RStudent.query.filter_by(session_id=session_id).all()
            student_id_to_rstudent_id = {rs.student_id: rs.id for rs in all_rstudents}
            
            # Get Student profiles
            student_profiles = StudentProfile.query.filter(
                StudentProfile.student_id.in_(student_id_to_rstudent_id.keys())
            ).all()
            student_profile_id_to_rstudent_id = {
                profile.id: student_id_to_rstudent_id.get(profile.student_id)
                for profile in student_profiles
                if profile.student_id in student_id_to_rstudent_id
            }
            
            # Get registrations from Course Management
            course_filters = [
                StudentCourseRegistration.student_id.in_(student_profile_id_to_rstudent_id.keys()),
                StudentCourseRegistration.course_code == selected_subject.code,
                StudentCourseRegistration.status == 'finalized'
            ]
            if session.name:
                course_filters.append(StudentCourseRegistration.academic_session == session.name)
            if session.year:
                course_filters.append(StudentCourseRegistration.year == session.year)
            if session.term:
                course_filters.append(StudentCourseRegistration.term == session.term)
            
            registered_regs = StudentCourseRegistration.query.filter(*course_filters).all()
            
            # Build registrations_data mapping RStudent.id to registration info
            for reg in registered_regs:
                rstudent_id = student_profile_id_to_rstudent_id.get(reg.student_id)
                if rstudent_id:
                    remark_text = (reg.remark or '').strip().lower()
                    is_retake = remark_text in {'retake', 're-retake', 're retake', 'reretake'}
                    registrations_data[rstudent_id] = {
                        'is_retake': is_retake,
                        'remark': reg.remark
                    }
        except Exception as e:
            current_app.logger.error(f'Error loading registration data: {str(e)}', exc_info=True)
        
        current_app.logger.debug(f'Result Management: Processing {len(students)} students for auto-populate')
        
        for student in students:
            mark = RMark.query.filter_by(student_id=student.id, subject_id=selected_subject.id).first()
            
            if mark:
                current_app.logger.info(f'Existing mark found for student {student.student_id}: attendance={mark.attendance}, assessment={mark.continuous_assessment}, part_a={mark.part_a}, part_b={mark.part_b}, attendance_manual={mark.attendance_manual}')
            else:
                current_app.logger.info(f'No existing mark for student {student.student_id}, will create new mark')
            
            # If mark doesn't exist, try to auto-populate from Class Management
            if not mark:
                current_app.logger.debug(f'Result Management: No existing mark found for student {student.student_id}, attempting auto-populate')
                try:
                    from blueprints.class_management.models import Session, ClassStudent, ClassAttendance
                    
                    # Find matching session in Class Management by course code
                    # Try multiple strategies to find the session
                    class_session = None
                    
                    # Strategy 1: Exact match (course_code + year + term + academic_session)
                    if session.year and session.term and session.name:
                        class_session = Session.query.filter_by(
                            course_code=selected_subject.code,
                            year=session.year,
                            term=session.term,
                            academic_session=session.name
                        ).first()
                        if class_session:
                            current_app.logger.debug(f'Found class_session (exact match) for {selected_subject.code}')
                    
                    # Strategy 2: Partial match (course_code + year + term)
                    if not class_session and session.year and session.term:
                        class_session = Session.query.filter_by(
                            course_code=selected_subject.code,
                            year=session.year,
                            term=session.term
                        ).first()
                        if class_session:
                            current_app.logger.debug(f'Found class_session (year+term match) for {selected_subject.code}')
                    
                    # Strategy 3: Fallback (course_code only, most recent)
                    if not class_session:
                        class_session = Session.query.filter_by(
                            course_code=selected_subject.code
                        ).order_by(Session.created_at.desc()).first()
                        if class_session:
                            current_app.logger.debug(f'Found class_session (course_code only) for {selected_subject.code}')
                    
                    if class_session:
                        # Find student in Class Management
                        class_student = ClassStudent.query.filter_by(
                            session_id=class_session.id,
                            student_id=student.student_id
                        ).first()
                        
                        if class_student:
                            # Create new RMark and auto-populate from Class Management
                            mark = RMark(student_id=student.id, subject_id=selected_subject.id)
                            current_app.logger.info(f'Creating new mark for student {student.student_id}, class_student found: assessment_total_40={class_student.assessment_total_40}, assessment_total={class_student.assessment_total}')
                            
                            # Import attendance calculation functions from Class Management
                            # Only auto-populate if attendance_manual is False or None (not manually entered)
                            if not mark.attendance_manual:
                                try:
                                    from blueprints.class_management.routes import _build_attendance_summary
                                    
                                    # Build attendance summary for the session (same as Class Management uses)
                                    attendance_summary = _build_attendance_summary(class_session)
                                    current_app.logger.debug(f'Built attendance summary for class_session {class_session.id}, total_classes={attendance_summary.get("total_classes", 0)}')
                                    
                                    # Get student's attendance data using student_id (public ID)
                                    # The summary uses student_id (public) as key, not database ID
                                    student_attendance = attendance_summary.get('per_student', {}).get(student.student_id, {})
                                    current_app.logger.debug(f'Student attendance data for {student.student_id}: {student_attendance}')
                                    
                                    if student_attendance and 'marks' in student_attendance:
                                        # Get marks directly from the summary (already calculated by Class Management)
                                        attendance_marks = student_attendance.get('marks', 0)
                                        current_app.logger.info(f'Found attendance marks for student {student.student_id}: {attendance_marks}')
                                        if attendance_marks is not None:
                                            mark.attendance = float(attendance_marks)
                                            current_app.logger.info(f'Set attendance={mark.attendance} for student {student.student_id}')
                                        else:
                                            mark.attendance = None
                                            current_app.logger.warning(f'Attendance marks is None for student {student.student_id}')
                                    else:
                                        # If no attendance data found, set to None
                                        mark.attendance = None
                                        current_app.logger.warning(f"No attendance data found for student {student.student_id} in session {class_session.id}. student_attendance={student_attendance}")
                                except ImportError as e:
                                    # If import fails, log warning
                                    current_app.logger.warning(f"Could not import _build_attendance_summary: {e}")
                                    mark.attendance = None
                                except Exception as e:
                                    current_app.logger.error(f"Error calculating attendance marks from Class Management for student {student.student_id}: {e}", exc_info=True)
                                    mark.attendance = None
                            
                            # Get continuous assessment (sum of best 3 assessments or assessment_total_40)
                            if class_student.assessment_total_40 is not None:
                                # PG: Already rounded in assessment calculation
                                mark.continuous_assessment = float(class_student.assessment_total_40)
                                current_app.logger.info(f'Set continuous_assessment={mark.continuous_assessment} from assessment_total_40 for student {student.student_id}')
                            elif class_student.assessment_total is not None:
                                # assessment_total might be out of 30, convert to out of 40 if needed
                                assessment_total = float(class_student.assessment_total)
                                current_app.logger.info(f'Found assessment_total={assessment_total} for student {student.student_id}')
                                # If it's already out of 40, use as is; otherwise scale it
                                if assessment_total <= 40:
                                    # UG: Round for result generation
                                    mark.continuous_assessment = round(assessment_total)
                                else:
                                    # Scale from 30 to 40 if needed, then round for result generation
                                    mark.continuous_assessment = round(min(40.0, (assessment_total / 30) * 40))
                                current_app.logger.info(f'Set continuous_assessment={mark.continuous_assessment} for student {student.student_id}')
                            else:
                                mark.continuous_assessment = None
                                current_app.logger.warning(f'No assessment data found for student {student.student_id}: assessment_total_40={class_student.assessment_total_40}, assessment_total={class_student.assessment_total}')
                            
                            # Get Section A and B marks from Exam Paper Evaluation
                            try:
                                from blueprints.class_management.models import ExamPaperEvaluation
                                import json
                                
                                # Find Exam Paper Evaluation entry for this course
                                # Try exact match first, then partial match
                                exam_entry = ExamPaperEvaluation.query.filter_by(
                                    course_code=selected_subject.code,
                                    archived=False
                                ).first()
                                
                                # If not found, try searching by course name
                                if not exam_entry:
                                    exam_entry = ExamPaperEvaluation.query.filter(
                                        ExamPaperEvaluation.course_name.ilike(f'%{selected_subject.name}%'),
                                        ExamPaperEvaluation.archived == False
                                    ).first()
                                
                                if exam_entry and exam_entry.marks_data:
                                    try:
                                        exam_marks = json.loads(exam_entry.marks_data)
                                        questions = exam_marks.get('questions', [])
                                        rows = exam_marks.get('rows', [])
                                        
                                        # Find student's marks in exam data
                                        student_row = None
                                        for row in rows:
                                            row_student_id = str(row.get('student_id', '')).strip()
                                            row_code = str(row.get('code', '')).strip()
                                            student_id_str = str(student.student_id).strip()
                                            
                                            if row_student_id == student_id_str or row_code == student_id_str:
                                                student_row = row
                                                break
                                        
                                        if student_row:
                                            marks_dict = student_row.get('marks', {})
                                            
                                            # Look for Section A and B in question labels
                                            section_a_found = False
                                            section_b_found = False
                                            
                                            # First, check question labels for Section A/B
                                            for question in questions:
                                                question_label = question.get('label', '').lower()
                                                
                                                if ('section a' in question_label or 'part a' in question_label) and not section_a_found:
                                                    # Get marks for this question
                                                    question_marks = marks_dict.get(question.get('label', ''), {})
                                                    if isinstance(question_marks, dict):
                                                        # Sum all parts
                                                        section_a_total = 0
                                                        for part_label, part_mark in question_marks.items():
                                                            try:
                                                                if part_mark and str(part_mark).strip():
                                                                    section_a_total += float(part_mark)
                                                            except (ValueError, TypeError):
                                                                pass
                                                        if section_a_total > 0:
                                                            mark.part_a = min(25.0, section_a_total)
                                                            section_a_found = True
                                                    elif isinstance(question_marks, (int, float)):
                                                        mark.part_a = min(25.0, float(question_marks))
                                                        section_a_found = True
                                                
                                                elif ('section b' in question_label or 'part b' in question_label) and not section_b_found:
                                                    # Get marks for this question
                                                    question_marks = marks_dict.get(question.get('label', ''), {})
                                                    if isinstance(question_marks, dict):
                                                        # Sum all parts
                                                        section_b_total = 0
                                                        for part_label, part_mark in question_marks.items():
                                                            try:
                                                                if part_mark and str(part_mark).strip():
                                                                    section_b_total += float(part_mark)
                                                            except (ValueError, TypeError):
                                                                pass
                                                        if section_b_total > 0:
                                                            mark.part_b = min(25.0, section_b_total)
                                                            section_b_found = True
                                                    elif isinstance(question_marks, (int, float)):
                                                        mark.part_b = min(25.0, float(question_marks))
                                                        section_b_found = True
                                            
                                            # If still not found, try alternative patterns
                                            if not section_a_found or not section_b_found:
                                                for question_label, question_data in marks_dict.items():
                                                    q_label_lower = question_label.lower()
                                                    
                                                    if not section_a_found and ('section a' in q_label_lower or 'part a' in q_label_lower or 'a)' in q_label_lower):
                                                        if isinstance(question_data, dict):
                                                            section_a_total = sum(float(v) for k, v in question_data.items() if v and str(v).strip())
                                                            if section_a_total > 0:
                                                                mark.part_a = min(25.0, section_a_total)
                                                                section_a_found = True
                                                        elif isinstance(question_data, (int, float, str)):
                                                            try:
                                                                val = float(question_data)
                                                                if val > 0:
                                                                    mark.part_a = min(25.0, val)
                                                                    section_a_found = True
                                                            except (ValueError, TypeError):
                                                                pass
                                                    
                                                    if not section_b_found and ('section b' in q_label_lower or 'part b' in q_label_lower or 'b)' in q_label_lower):
                                                        if isinstance(question_data, dict):
                                                            section_b_total = sum(float(v) for k, v in question_data.items() if v and str(v).strip())
                                                            if section_b_total > 0:
                                                                mark.part_b = min(25.0, section_b_total)
                                                                section_b_found = True
                                                        elif isinstance(question_data, (int, float, str)):
                                                            try:
                                                                val = float(question_data)
                                                                if val > 0:
                                                                    mark.part_b = min(25.0, val)
                                                                    section_b_found = True
                                                            except (ValueError, TypeError):
                                                                pass
                                            
                                            # If we have total but not individual sections, and there are exactly 2 questions, split them
                                            if not section_a_found and not section_b_found and len(questions) == 2:
                                                total_marks_str = student_row.get('total', '')
                                                if total_marks_str:
                                                    try:
                                                        total = float(total_marks_str)
                                                        # Split 50-50 between two questions
                                                        mark.part_a = total / 2
                                                        mark.part_b = total / 2
                                                    except (ValueError, TypeError):
                                                        pass
                                            
                                    except json.JSONDecodeError as e:
                                        current_app.logger.error(f"Error parsing exam marks JSON: {e}", exc_info=True)
                                    except Exception as e:
                                        current_app.logger.error(f"Error processing exam marks: {e}", exc_info=True)
                            except Exception as e:
                                current_app.logger.error(f"Error fetching exam marks: {e}", exc_info=True)
                            
                            # Save the auto-populated mark
                            db.session.add(mark)
                            db.session.commit()
                        else:
                            # class_student not found - try to import attendance from Class Management anyway
                            # Create mark even if class_student doesn't exist
                            mark = RMark(student_id=student.id, subject_id=selected_subject.id)
                            
                            # Try to import attendance from Class Management session
                            if not mark.attendance_manual:
                                try:
                                    from blueprints.class_management.routes import _build_attendance_summary
                                    
                                    attendance_summary = _build_attendance_summary(class_session)
                                    student_attendance = attendance_summary.get('per_student', {}).get(student.student_id, {})
                                    
                                    if student_attendance and 'marks' in student_attendance:
                                        attendance_marks = student_attendance.get('marks', 0)
                                        if attendance_marks is not None:
                                            mark.attendance = float(attendance_marks)
                                except Exception as e:
                                    current_app.logger.debug(f"Could not import attendance for student {student.student_id} (not in Class Management): {e}")
                            
                            # Try to find student in other Class Management sessions for this course
                            if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                                # Try to find student in any Class Management session for this course
                                alternative_class_student = ClassStudent.query.join(
                                    Session, ClassStudent.session_id == Session.id
                                ).filter(
                                    Session.course_code == selected_subject.code,
                                    ClassStudent.student_id == student.student_id
                                ).first()
                                
                                if alternative_class_student:
                                    if alternative_class_student.assessment_total_40 is not None:
                                        mark.continuous_assessment = float(alternative_class_student.assessment_total_40)
                                        current_app.logger.debug(f'Imported assessment_total_40={mark.continuous_assessment} for student {student.student_id}')
                                    elif alternative_class_student.assessment_total is not None:
                                        assessment_total = float(alternative_class_student.assessment_total)
                                        if assessment_total <= 40:
                                            mark.continuous_assessment = round(assessment_total)
                                        else:
                                            mark.continuous_assessment = round(min(40.0, (assessment_total / 30) * 40))
                                        current_app.logger.debug(f'Imported assessment_total={mark.continuous_assessment} for student {student.student_id}')
                            
                            # Import exam marks (Part A/B) even if class_student not found
                            try:
                                from blueprints.class_management.models import ExamPaperEvaluation
                                import json
                                
                                exam_entry = ExamPaperEvaluation.query.filter_by(
                                    course_code=selected_subject.code,
                                    archived=False
                                ).first()
                                
                                if not exam_entry:
                                    exam_entry = ExamPaperEvaluation.query.filter(
                                        ExamPaperEvaluation.course_name.ilike(f'%{selected_subject.name}%'),
                                        ExamPaperEvaluation.archived == False
                                    ).first()
                                
                                if exam_entry and exam_entry.marks_data:
                                    try:
                                        exam_marks = json.loads(exam_entry.marks_data)
                                        rows = exam_marks.get('rows', [])
                                        
                                        student_row = None
                                        for row in rows:
                                            row_student_id = str(row.get('student_id', '')).strip()
                                            row_code = str(row.get('code', '')).strip()
                                            student_id_str = str(student.student_id).strip()
                                            
                                            if row_student_id == student_id_str or row_code == student_id_str:
                                                student_row = row
                                                break
                                        
                                        if student_row:
                                            # Import Part A and B from exam marks (same logic as above)
                                            marks_dict = student_row.get('marks', {})
                                            questions = exam_marks.get('questions', [])
                                            section_a_found = False
                                            section_b_found = False
                                            
                                            for question in questions:
                                                question_label = question.get('label', '').lower()
                                                question_marks = marks_dict.get(question.get('label', ''), {})
                                                
                                                if ('section a' in question_label or 'part a' in question_label) and not section_a_found:
                                                    if isinstance(question_marks, dict):
                                                        section_a_total = sum(float(v) for k, v in question_marks.items() if v and str(v).strip())
                                                        if section_a_total > 0:
                                                            mark.part_a = min(25.0, section_a_total)
                                                            section_a_found = True
                                                    elif isinstance(question_marks, (int, float)):
                                                        mark.part_a = min(25.0, float(question_marks))
                                                        section_a_found = True
                                                
                                                elif ('section b' in question_label or 'part b' in question_label) and not section_b_found:
                                                    if isinstance(question_marks, dict):
                                                        section_b_total = sum(float(v) for k, v in question_marks.items() if v and str(v).strip())
                                                        if section_b_total > 0:
                                                            mark.part_b = min(25.0, section_b_total)
                                                            section_b_found = True
                                                    elif isinstance(question_marks, (int, float)):
                                                        mark.part_b = min(25.0, float(question_marks))
                                                        section_b_found = True
                                    except Exception as e:
                                        current_app.logger.debug(f"Error importing exam marks for student {student.student_id}: {e}")
                            except Exception as e:
                                current_app.logger.debug(f"Error fetching exam marks for student {student.student_id}: {e}")
                            
                            # Save the mark even if no data was imported
                            db.session.add(mark)
                            db.session.commit()
                            current_app.logger.debug(f'Created mark for student {student.student_id} (class_student not found)')
                    else:
                        # class_session not found - create empty mark and try to import from alternative sources
                        mark = RMark(student_id=student.id, subject_id=selected_subject.id)
                        
                        # Try to find student in any Class Management session for this course
                        if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                            alternative_class_student = ClassStudent.query.join(
                                Session, ClassStudent.session_id == Session.id
                            ).filter(
                                Session.course_code == selected_subject.code,
                                ClassStudent.student_id == student.student_id
                            ).first()
                            
                            if alternative_class_student:
                                # Import assessment
                                if alternative_class_student.assessment_total_40 is not None:
                                    mark.continuous_assessment = float(alternative_class_student.assessment_total_40)
                                elif alternative_class_student.assessment_total is not None:
                                    assessment_total = float(alternative_class_student.assessment_total)
                                    if assessment_total <= 40:
                                        mark.continuous_assessment = round(assessment_total)
                                    else:
                                        mark.continuous_assessment = round(min(40.0, (assessment_total / 30) * 40))
                                
                                # Try to import attendance from the alternative session
                                try:
                                    from blueprints.class_management.routes import _build_attendance_summary
                                    alt_session = Session.query.get(alternative_class_student.session_id)
                                    if alt_session:
                                        attendance_summary = _build_attendance_summary(alt_session)
                                        student_attendance = attendance_summary.get('per_student', {}).get(student.student_id, {})
                                        if student_attendance and 'marks' in student_attendance:
                                            attendance_marks = student_attendance.get('marks', 0)
                                            if attendance_marks is not None:
                                                mark.attendance = float(attendance_marks)
                                except Exception as e:
                                    current_app.logger.debug(f"Could not import attendance from alternative session: {e}")
                        
                        db.session.add(mark)
                        db.session.commit()
                        current_app.logger.debug(f'Created mark for student {student.student_id} (class_session not found)')
                            
                except Exception as e:
                    current_app.logger.error(f"Error auto-populating marks from Class Management: {e}", exc_info=True)
                    # Continue without auto-population if there's an error
                    # Create empty mark so student still shows in the list
                    try:
                        mark = RMark(student_id=student.id, subject_id=selected_subject.id)
                        db.session.add(mark)
                        db.session.commit()
                    except Exception as create_error:
                        current_app.logger.error(f"Error creating empty mark: {create_error}", exc_info=True)
                        mark = None
            else:
                # Mark exists - ALWAYS try to sync from ClassManagement (unless manually entered)
                # This ensures marks are always up-to-date when page loads
                current_app.logger.info(f'🔄 Processing EXISTING mark for student {student.student_id}, subject {selected_subject.code}')
                if mark:
                    needs_update = False
                    current_app.logger.info(f'📝 Existing mark for student {student.student_id}: attendance={mark.attendance}, assessment={mark.continuous_assessment}, part_a={mark.part_a}, part_b={mark.part_b}, attendance_manual={mark.attendance_manual}')
                    current_app.logger.info(f'📝 Subject info: code="{selected_subject.code}", name="{selected_subject.name}", type="{selected_subject.subject_type}"')
                    
                    try:
                        from blueprints.class_management.models import Session as CMSession, ClassStudent
                        from blueprints.class_management.routes import _build_attendance_summary
                        
                        # Find ANY ClassManagement session for this course (most flexible approach)
                        class_sessions = CMSession.query.filter_by(course_code=selected_subject.code).order_by(CMSession.created_at.desc()).all()
                        
                        if class_sessions:
                            current_app.logger.info(f'Found {len(class_sessions)} ClassManagement session(s) for course {selected_subject.code}')
                            
                            # Try each session until we find the student
                            class_student = None
                            class_session = None
                            
                            for cs in class_sessions:
                                cs_student = ClassStudent.query.filter_by(
                                    session_id=cs.id,
                                    student_id=student.student_id
                                ).first()
                                
                                if cs_student:
                                    class_student = cs_student
                                    class_session = cs
                                    current_app.logger.info(f'Found class_student in session {cs.id} for student {student.student_id}')
                                    break
                            
                            # If student found, sync attendance and assessment
                            if class_student and class_session:
                                # Sync attendance (if not manually entered)
                                if not mark.attendance_manual:
                                    try:
                                        attendance_summary = _build_attendance_summary(class_session)
                                        student_attendance = attendance_summary.get('per_student', {}).get(student.student_id, {})
                                        
                                        if student_attendance and 'marks' in student_attendance:
                                            attendance_marks = student_attendance.get('marks', 0)
                                            if attendance_marks is not None:
                                                new_attendance = float(attendance_marks)
                                                if mark.attendance != new_attendance:
                                                    mark.attendance = new_attendance
                                                    needs_update = True
                                                    current_app.logger.info(f'✅ Updated attendance: {mark.attendance} for student {student.student_id}')
                                    except Exception as e:
                                        current_app.logger.error(f"Error syncing attendance: {e}", exc_info=True)
                                
                                # Sync continuous assessment (for Theory subjects)
                                if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                                    if class_student.assessment_total_40 is not None:
                                        new_assessment = float(class_student.assessment_total_40)
                                        if mark.continuous_assessment != new_assessment:
                                            mark.continuous_assessment = new_assessment
                                            needs_update = True
                                            current_app.logger.info(f'✅ Updated assessment_total_40: {mark.continuous_assessment} for student {student.student_id}')
                                    elif class_student.assessment_total is not None:
                                        assessment_total = float(class_student.assessment_total)
                                        if assessment_total <= 40:
                                            new_assessment = round(assessment_total)
                                        else:
                                            new_assessment = round(min(40.0, (assessment_total / 30) * 40))
                                        
                                        if mark.continuous_assessment != new_assessment:
                                            mark.continuous_assessment = new_assessment
                                            needs_update = True
                                            current_app.logger.info(f'✅ Updated assessment_total: {mark.continuous_assessment} for student {student.student_id}')
                                    else:
                                        current_app.logger.warning(f'No assessment data for student {student.student_id} in ClassManagement')
                            else:
                                # Student not found in any ClassManagement session - try to get attendance from any session
                                if not mark.attendance_manual:
                                    try:
                                        first_session = class_sessions[0]
                                        attendance_summary = _build_attendance_summary(first_session)
                                        student_attendance = attendance_summary.get('per_student', {}).get(student.student_id, {})
                                        
                                        if student_attendance and 'marks' in student_attendance:
                                            attendance_marks = student_attendance.get('marks', 0)
                                            if attendance_marks is not None:
                                                new_attendance = float(attendance_marks)
                                                if mark.attendance != new_attendance:
                                                    mark.attendance = new_attendance
                                                    needs_update = True
                                                    current_app.logger.info(f'✅ Updated attendance from any session: {mark.attendance} for student {student.student_id}')
                                    except Exception as e:
                                        current_app.logger.debug(f"Could not sync attendance from any session: {e}")
                                
                                # Try to find student in any session for assessment
                                if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                                    for cs in class_sessions:
                                        alt_student = ClassStudent.query.filter_by(
                                            session_id=cs.id,
                                            student_id=student.student_id
                                        ).first()
                                        
                                        if alt_student:
                                            if alt_student.assessment_total_40 is not None:
                                                new_assessment = float(alt_student.assessment_total_40)
                                                if mark.continuous_assessment != new_assessment:
                                                    mark.continuous_assessment = new_assessment
                                                    needs_update = True
                                                    current_app.logger.info(f'✅ Updated assessment from alternative session: {mark.continuous_assessment} for student {student.student_id}')
                                                    break
                                            elif alt_student.assessment_total is not None:
                                                assessment_total = float(alt_student.assessment_total)
                                                if assessment_total <= 40:
                                                    new_assessment = round(assessment_total)
                                                else:
                                                    new_assessment = round(min(40.0, (assessment_total / 30) * 40))
                                                
                                                if mark.continuous_assessment != new_assessment:
                                                    mark.continuous_assessment = new_assessment
                                                    needs_update = True
                                                    current_app.logger.info(f'✅ Updated assessment from alternative session: {mark.continuous_assessment} for student {student.student_id}')
                                                    break
                        else:
                            current_app.logger.warning(f'No ClassManagement sessions found for course {selected_subject.code}')
                        
                        # Sync Part A and Part B from Exam Paper Evaluation (for Theory subjects)
                        # This is independent of ClassManagement, so always try to sync
                        current_app.logger.info(f'🔍 Checking Part A/B sync for student {student.student_id}')
                        current_app.logger.info(f'📋 Subject details: code="{selected_subject.code}", name="{selected_subject.name}", subject_type="{selected_subject.subject_type}" (type: {type(selected_subject.subject_type)})')
                        current_app.logger.info(f'🔍 Checking if subject_type in ("Theory", "Theory (UG)", "Theory (PG)")')
                        is_theory = selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)')
                        current_app.logger.info(f'🔍 is_theory={is_theory}')
                        if is_theory:
                            current_app.logger.info(f'✅ Subject type is Theory, proceeding with Part A/B sync for student {student.student_id}')
                            try:
                                from blueprints.class_management.models import ExamPaperEvaluation
                                import json
                                current_app.logger.info(f'🔍 Looking for ExamPaperEvaluation: course_code="{selected_subject.code}", course_name="{selected_subject.name}"')
                                
                                # Find Exam Paper Evaluation entry for this course
                                exam_entry = ExamPaperEvaluation.query.filter_by(
                                    course_code=selected_subject.code,
                                    archived=False
                                ).first()
                                
                                if exam_entry:
                                    current_app.logger.info(f'✅ Found ExamPaperEvaluation by course_code: ID={exam_entry.id}, course_code="{exam_entry.course_code}", course_name="{exam_entry.course_name}"')
                                else:
                                    current_app.logger.info(f'⚠️ No ExamPaperEvaluation found by course_code="{selected_subject.code}", trying course_name...')
                                
                                # If not found, try searching by course name
                                if not exam_entry:
                                    exam_entry = ExamPaperEvaluation.query.filter(
                                        ExamPaperEvaluation.course_name.ilike(f'%{selected_subject.name}%'),
                                        ExamPaperEvaluation.archived == False
                                    ).first()
                                    if exam_entry:
                                        current_app.logger.info(f'✅ Found ExamPaperEvaluation by course_name: ID={exam_entry.id}, course_code="{exam_entry.course_code}", course_name="{exam_entry.course_name}"')
                                    else:
                                        current_app.logger.warning(f'⚠️ No ExamPaperEvaluation found by course_name="{selected_subject.name}"')
                                
                                if exam_entry and exam_entry.marks_data:
                                    current_app.logger.info(f'🔍 Found ExamPaperEvaluation entry for course {selected_subject.code} (ID: {exam_entry.id})')
                                    try:
                                        exam_marks = json.loads(exam_entry.marks_data)
                                        questions = exam_marks.get('questions', [])
                                        rows = exam_marks.get('rows', [])
                                        
                                        current_app.logger.info(f'📊 Exam data: {len(questions)} questions, {len(rows)} student rows')
                                        current_app.logger.info(f'📋 Question labels: {[q.get("label", "") for q in questions]}')
                                        
                                        # Find student's marks in exam data
                                        student_row = None
                                        for row in rows:
                                            row_student_id = str(row.get('student_id', '')).strip()
                                            row_code = str(row.get('code', '')).strip()
                                            student_id_str = str(student.student_id).strip()
                                            
                                            current_app.logger.debug(f'🔎 Comparing: row_student_id="{row_student_id}", row_code="{row_code}", student_id="{student_id_str}"')
                                            
                                            if row_student_id == student_id_str or row_code == student_id_str:
                                                student_row = row
                                                current_app.logger.info(f'✅ Found student row in exam data for {student.student_id}')
                                                break
                                        
                                        if student_row:
                                            marks_dict = student_row.get('marks', {})
                                            current_app.logger.info(f'📝 Student marks_dict keys: {list(marks_dict.keys())}')
                                            current_app.logger.info(f'📝 Student marks_dict: {marks_dict}')
                                            
                                            section_a_found = False
                                            section_b_found = False
                                            
                                            # First, check question labels for Section A/B
                                            for question in questions:
                                                question_label = question.get('label', '')
                                                question_label_lower = question_label.lower()
                                                
                                                current_app.logger.debug(f'🔍 Checking question: "{question_label}" (lower: "{question_label_lower}")')
                                                
                                                if ('section a' in question_label_lower or 'part a' in question_label_lower) and not section_a_found:
                                                    # Try exact match first
                                                    question_marks = marks_dict.get(question_label, {})
                                                    current_app.logger.info(f'🔍 Part A: question_label="{question_label}", question_marks type={type(question_marks)}, value={question_marks}')
                                                    
                                                    if isinstance(question_marks, dict):
                                                        section_a_total = 0
                                                        for k, v in question_marks.items():
                                                            try:
                                                                if v and str(v).strip():
                                                                    section_a_total += float(v)
                                                                    current_app.logger.debug(f'  Part A sub-part "{k}": {v} (total so far: {section_a_total})')
                                                            except (ValueError, TypeError):
                                                                pass
                                                        
                                                        if section_a_total > 0:
                                                            new_part_a = min(25.0, section_a_total)
                                                            if mark.part_a != new_part_a:
                                                                mark.part_a = new_part_a
                                                                needs_update = True
                                                                current_app.logger.info(f'✅ Updated part_a: {mark.part_a} for student {student.student_id} (from question "{question_label}")')
                                                            section_a_found = True
                                                    elif isinstance(question_marks, (int, float)):
                                                        new_part_a = min(25.0, float(question_marks))
                                                        if mark.part_a != new_part_a:
                                                            mark.part_a = new_part_a
                                                            needs_update = True
                                                            current_app.logger.info(f'✅ Updated part_a: {mark.part_a} for student {student.student_id} (from question "{question_label}")')
                                                        section_a_found = True
                                                    else:
                                                        current_app.logger.warning(f'⚠️ Part A question_marks is not dict or number: {type(question_marks)} = {question_marks}')
                                                
                                                elif ('section b' in question_label_lower or 'part b' in question_label_lower) and not section_b_found:
                                                    # Try exact match first
                                                    question_marks = marks_dict.get(question_label, {})
                                                    current_app.logger.info(f'🔍 Part B: question_label="{question_label}", question_marks type={type(question_marks)}, value={question_marks}')
                                                    
                                                    if isinstance(question_marks, dict):
                                                        section_b_total = 0
                                                        for k, v in question_marks.items():
                                                            try:
                                                                if v and str(v).strip():
                                                                    section_b_total += float(v)
                                                                    current_app.logger.debug(f'  Part B sub-part "{k}": {v} (total so far: {section_b_total})')
                                                            except (ValueError, TypeError):
                                                                pass
                                                        
                                                        if section_b_total > 0:
                                                            new_part_b = min(25.0, section_b_total)
                                                            if mark.part_b != new_part_b:
                                                                mark.part_b = new_part_b
                                                                needs_update = True
                                                                current_app.logger.info(f'✅ Updated part_b: {mark.part_b} for student {student.student_id} (from question "{question_label}")')
                                                            section_b_found = True
                                                    elif isinstance(question_marks, (int, float)):
                                                        new_part_b = min(25.0, float(question_marks))
                                                        if mark.part_b != new_part_b:
                                                            mark.part_b = new_part_b
                                                            needs_update = True
                                                            current_app.logger.info(f'✅ Updated part_b: {mark.part_b} for student {student.student_id} (from question "{question_label}")')
                                                        section_b_found = True
                                                    else:
                                                        current_app.logger.warning(f'⚠️ Part B question_marks is not dict or number: {type(question_marks)} = {question_marks}')
                                            
                                            # If still not found, try alternative patterns
                                            if not section_a_found or not section_b_found:
                                                current_app.logger.info(f'🔍 Trying alternative patterns. section_a_found={section_a_found}, section_b_found={section_b_found}')
                                                for question_label, question_data in marks_dict.items():
                                                    q_label_lower = question_label.lower()
                                                    
                                                    if not section_a_found and ('section a' in q_label_lower or 'part a' in q_label_lower or 'a)' in q_label_lower):
                                                        current_app.logger.info(f'🔍 Found Part A pattern in key: "{question_label}"')
                                                        if isinstance(question_data, dict):
                                                            section_a_total = 0
                                                            for k, v in question_data.items():
                                                                try:
                                                                    if v and str(v).strip():
                                                                        section_a_total += float(v)
                                                                except (ValueError, TypeError):
                                                                    pass
                                                            if section_a_total > 0:
                                                                new_part_a = min(25.0, section_a_total)
                                                                if mark.part_a != new_part_a:
                                                                    mark.part_a = new_part_a
                                                                    needs_update = True
                                                                    current_app.logger.info(f'✅ Updated part_a (alt): {mark.part_a} for student {student.student_id} from key "{question_label}"')
                                                                section_a_found = True
                                                        elif isinstance(question_data, (int, float, str)):
                                                            try:
                                                                val = float(question_data)
                                                                if val > 0:
                                                                    new_part_a = min(25.0, val)
                                                                    if mark.part_a != new_part_a:
                                                                        mark.part_a = new_part_a
                                                                        needs_update = True
                                                                        current_app.logger.info(f'✅ Updated part_a (alt): {mark.part_a} for student {student.student_id} from key "{question_label}"')
                                                                    section_a_found = True
                                                            except (ValueError, TypeError):
                                                                pass
                                                    
                                                    if not section_b_found and ('section b' in q_label_lower or 'part b' in q_label_lower or 'b)' in q_label_lower):
                                                        current_app.logger.info(f'🔍 Found Part B pattern in key: "{question_label}"')
                                                        if isinstance(question_data, dict):
                                                            section_b_total = 0
                                                            for k, v in question_data.items():
                                                                try:
                                                                    if v and str(v).strip():
                                                                        section_b_total += float(v)
                                                                except (ValueError, TypeError):
                                                                    pass
                                                            if section_b_total > 0:
                                                                new_part_b = min(25.0, section_b_total)
                                                                if mark.part_b != new_part_b:
                                                                    mark.part_b = new_part_b
                                                                    needs_update = True
                                                                    current_app.logger.info(f'✅ Updated part_b (alt): {mark.part_b} for student {student.student_id} from key "{question_label}"')
                                                                section_b_found = True
                                                        elif isinstance(question_data, (int, float, str)):
                                                            try:
                                                                val = float(question_data)
                                                                if val > 0:
                                                                    new_part_b = min(25.0, val)
                                                                    if mark.part_b != new_part_b:
                                                                        mark.part_b = new_part_b
                                                                        needs_update = True
                                                                        current_app.logger.info(f'✅ Updated part_b (alt): {mark.part_b} for student {student.student_id} from key "{question_label}"')
                                                                    section_b_found = True
                                                            except (ValueError, TypeError):
                                                                pass
                                            
                                            # If we have total but not individual sections, and there are exactly 2 questions, split them
                                            if not section_a_found and not section_b_found:
                                                current_app.logger.info(f'⚠️ Part A/B not found. Trying split from total. Questions count: {len(questions)}')
                                                if len(questions) == 2:
                                                    total_marks_str = student_row.get('total', '')
                                                    current_app.logger.info(f'📊 Total marks string: "{total_marks_str}"')
                                                    if total_marks_str:
                                                        try:
                                                            total = float(total_marks_str)
                                                            # Split 50-50 between two questions
                                                            new_part_a = total / 2
                                                            new_part_b = total / 2
                                                            if mark.part_a != new_part_a or mark.part_b != new_part_b:
                                                                mark.part_a = new_part_a
                                                                mark.part_b = new_part_b
                                                                needs_update = True
                                                                current_app.logger.info(f'✅ Updated part_a/b (split): {mark.part_a}/{mark.part_b} for student {student.student_id}')
                                                        except (ValueError, TypeError):
                                                            current_app.logger.warning(f'⚠️ Could not parse total marks "{total_marks_str}"')
                                                else:
                                                    current_app.logger.warning(f'⚠️ Cannot split: {len(questions)} questions (expected 2)')
                                            
                                            if not section_a_found:
                                                current_app.logger.warning(f'⚠️ Part A not found for student {student.student_id}')
                                            if not section_b_found:
                                                current_app.logger.warning(f'⚠️ Part B not found for student {student.student_id}')
                                        else:
                                            current_app.logger.warning(f'⚠️ Student row not found in exam data for student {student.student_id}. Searched in {len(rows)} rows.')
                                    except json.JSONDecodeError as e:
                                        current_app.logger.error(f"❌ Error parsing exam marks JSON: {e}", exc_info=True)
                                    except Exception as e:
                                        current_app.logger.error(f"❌ Error processing exam marks: {e}", exc_info=True)
                                else:
                                    if not exam_entry:
                                        current_app.logger.warning(f'⚠️ No ExamPaperEvaluation entry found for course {selected_subject.code} (archived=False)')
                                    elif not exam_entry.marks_data:
                                        current_app.logger.warning(f'⚠️ ExamPaperEvaluation entry found (ID={exam_entry.id}) but marks_data is empty for course {selected_subject.code}')
                            except Exception as e:
                                current_app.logger.error(f"❌ Error fetching exam marks for student {student.student_id}: {e}", exc_info=True)
                        else:
                            current_app.logger.info(f'⏭️ Skipping Part A/B sync: subject_type="{selected_subject.subject_type}" is not Theory')
                        
                        # Commit all sync changes at once
                        if needs_update:
                            try:
                                db.session.commit()
                                current_app.logger.info(f'✅ Committed sync updates for student {student.student_id}')
                            except Exception as commit_error:
                                current_app.logger.error(f'Error committing sync: {commit_error}', exc_info=True)
                                db.session.rollback()
                    except Exception as e:
                        current_app.logger.error(f'Error syncing marks for student {student.student_id}: {e}', exc_info=True)
                        # Continue with existing mark even if sync fails
            
            # Add mark to marks_data (even if None) to ensure all registered students are shown
            # Template will handle None marks by showing empty input fields
            # Always add mark to marks_data, even if sync failed
            if mark is None:
                # If mark is None for some reason, try to load it from database first
                current_app.logger.warning(f'Mark is None for student {student.student_id}, attempting to load from database')
                mark = RMark.query.filter_by(student_id=student.id, subject_id=selected_subject.id).first()
                if mark is None:
                    # Still None - create an empty one
                    current_app.logger.warning(f'Mark still None after query, creating new mark for student {student.student_id}')
                    mark = RMark(student_id=student.id, subject_id=selected_subject.id)
                    db.session.add(mark)
                    try:
                        db.session.commit()
                    except Exception as e:
                        current_app.logger.error(f'Error creating empty mark for student {student.student_id}: {e}', exc_info=True)
                        db.session.rollback()
            
            marks_data[student.id] = mark
            # registrations_data is already populated above from StudentCourseRegistration
        
        # Final safety check: Ensure all students have marks in marks_data
        if selected_subject:
            current_app.logger.info(f'Final check: {len(students)} students, {len(marks_data)} marks in marks_data')
            for student in students:
                if student.id not in marks_data:
                    current_app.logger.warning(f'Student {student.student_id} (ID: {student.id}) missing from marks_data, loading from database')
                    mark = RMark.query.filter_by(student_id=student.id, subject_id=selected_subject.id).first()
                    if mark is None:
                        current_app.logger.warning(f'Creating new mark for student {student.student_id}')
                        mark = RMark(student_id=student.id, subject_id=selected_subject.id)
                        db.session.add(mark)
                        try:
                            db.session.commit()
                            current_app.logger.info(f'Created new mark for student {student.student_id}')
                        except Exception as e:
                            current_app.logger.error(f'Error creating mark for missing student {student.student_id}: {e}', exc_info=True)
                            db.session.rollback()
                    else:
                        current_app.logger.info(f'Loaded existing mark for student {student.student_id}: attendance={mark.attendance}, assessment={mark.continuous_assessment}')
                    marks_data[student.id] = mark
            
            current_app.logger.info(f'After final check: {len(marks_data)} marks in marks_data for {len(students)} students')

    if request.method == 'POST':
        subject_id = request.form.get('subject_id', type=int)
        if not subject_id:
            flash('Please select a subject.', 'danger')
            return redirect(url_for('result_management.add_marks', session_id=session_id))

        subject = RSubject.query.get_or_404(subject_id)

        for student in students:
            existing_mark = RMark.query.filter_by(student_id=student.id, subject_id=subject.id).first()
            if existing_mark is None:
                existing_mark = RMark(student_id=student.id, subject_id=subject.id)
                db.session.add(existing_mark)
            
            # Check if the student is registered for this course as a retake
            registration = RCourseRegistration.query.filter_by(student_id=student.id, subject_id=subject.id).first()
            is_retake = registration.is_retake if registration else False
            existing_mark.is_retake = is_retake

            total_marks = 0
            if subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                attendance = request.form.get(f'attendance_{student.id}')
                continuous_assessment = request.form.get(f'continuous_assessment_{student.id}')
                part_a = request.form.get(f'part_a_{student.id}')
                part_b = request.form.get(f'part_b_{student.id}')
                
                if attendance:
                    existing_mark.attendance = float(attendance)
                    existing_mark.attendance_manual = True
                else:
                    existing_mark.attendance = None
                existing_mark.continuous_assessment = float(continuous_assessment) if continuous_assessment else None
                existing_mark.part_a = float(part_a) if part_a else None
                existing_mark.part_b = float(part_b) if part_b else None
                
                total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.continuous_assessment, existing_mark.part_a, existing_mark.part_b]))
            
            elif subject.subject_type == 'Sessional':
                attendance = request.form.get(f'attendance_{student.id}')
                sessional_report = request.form.get(f'sessional_report_{student.id}')
                sessional_viva = request.form.get(f'sessional_viva_{student.id}')

                if attendance:
                    existing_mark.attendance = float(attendance)
                    existing_mark.attendance_manual = True
                else:
                    existing_mark.attendance = None
                existing_mark.sessional_report = float(sessional_report) if sessional_report else None
                existing_mark.sessional_viva = float(sessional_viva) if sessional_viva else None

                total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.sessional_report, existing_mark.sessional_viva]))
            
            elif subject.subject_type in ('Thesis (UG)', 'Thesis I (UG)', 'Thesis II (UG)'):
                # Contact (10), Thesis Evaluation (60), Presentation (30)
                attendance = request.form.get(f'attendance_{student.id}')  # Contact uses attendance field
                thesis_evaluation = request.form.get(f'thesis_evaluation_{student.id}')
                presentation = request.form.get(f'presentation_{student.id}')
                
                if attendance:
                    existing_mark.attendance = float(attendance)
                    existing_mark.attendance_manual = True
                else:
                    existing_mark.attendance = None
                existing_mark.thesis_evaluation = float(thesis_evaluation) if thesis_evaluation else None
                existing_mark.presentation = float(presentation) if presentation else None
                
                total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.thesis_evaluation, existing_mark.presentation]))

            elif subject.subject_type == 'Dissertation':
                if subject.dissertation_type == 'Type1':
                    # Dissertation Proposal (PG): Supervisor Assessment (30), Proposal Presentation (70)
                    supervisor_assessment = request.form.get(f'supervisor_assessment_{student.id}')
                    proposal_presentation = request.form.get(f'proposal_presentation_{student.id}')
                    
                    existing_mark.supervisor_assessment = float(supervisor_assessment) if supervisor_assessment else None
                    existing_mark.proposal_presentation = float(proposal_presentation) if proposal_presentation else None
                    existing_mark.project_report = None
                    existing_mark.defense = None
                    
                    total_marks = sum(filter(None, [existing_mark.supervisor_assessment, existing_mark.proposal_presentation]))
                elif subject.dissertation_type == 'Type2':
                    # Dissertation Defence (PG): Supervisor Assessment (20), Project Report (50), Defense (30)
                    supervisor_assessment = request.form.get(f'supervisor_assessment_{student.id}')
                    project_report = request.form.get(f'project_report_{student.id}')
                    defense = request.form.get(f'defense_{student.id}')
                    
                    existing_mark.supervisor_assessment = float(supervisor_assessment) if supervisor_assessment else None
                    existing_mark.project_report = float(project_report) if project_report else None
                    existing_mark.defense = float(defense) if defense else None
                    existing_mark.proposal_presentation = None
                    
                    total_marks = sum(filter(None, [existing_mark.supervisor_assessment, existing_mark.project_report, existing_mark.defense]))
                else:
                    # Fallback for existing Dissertation subjects without type
                    supervisor_assessment = request.form.get(f'supervisor_assessment_{student.id}')
                    proposal_presentation = request.form.get(f'proposal_presentation_{student.id}')
                    project_report = request.form.get(f'project_report_{student.id}')
                    defense = request.form.get(f'defense_{student.id}')
                    
                    existing_mark.supervisor_assessment = float(supervisor_assessment) if supervisor_assessment else None
                    existing_mark.proposal_presentation = float(proposal_presentation) if proposal_presentation else None
                    existing_mark.project_report = float(project_report) if project_report else None
                    existing_mark.defense = float(defense) if defense else None

                    total_marks = sum(filter(None, [existing_mark.supervisor_assessment, existing_mark.proposal_presentation, existing_mark.project_report, existing_mark.defense]))
            
            elif subject.subject_type == 'Viva':
                viva = request.form.get(f'viva_{student.id}')
                existing_mark.viva = float(viva) if viva else None
                total_marks = existing_mark.viva or 0
            
            existing_mark.total_marks = total_marks
            existing_mark.grade_point, existing_mark.grade_letter = calculate_grade(total_marks, is_retake=is_retake)

        db.session.commit()
        # Emit WebSocket event for live update
        try:
            from utils.websocket_events import emit_marks_update
            emit_marks_update(session_id, {
                'subject_id': subject.id,
                'updated_at': datetime.utcnow().isoformat()
            })
        except Exception as e:
            current_app.logger.warning(f'Failed to emit marks update event: {e}')
        flash(f'Marks for {subject.name} saved successfully!', 'success')
        return redirect(url_for('result_management.add_marks', session_id=session_id, subject_id=subject.id))

    # Debug logging before rendering
    current_app.logger.info(f'Rendering add_marks template: session_id={session_id}, subjects_count={len(subjects)}, students_count={len(students)}')
    if subjects:
        current_app.logger.debug(f'Subjects being passed to template: {[s.code for s in subjects[:5]]}...')
    else:
        current_app.logger.warning(f'No subjects found for session {session_id} - dropdown will be empty!')
    
    # Log marks_data status
    if selected_subject:
        current_app.logger.info(f'Marks data status: {len(marks_data)} marks for {len(students)} students')
        if marks_data:
            sample_mark = list(marks_data.values())[0]
            current_app.logger.debug(f'Sample mark: attendance={sample_mark.attendance}, assessment={sample_mark.continuous_assessment}, part_a={sample_mark.part_a}, part_b={sample_mark.part_b}')
        else:
            current_app.logger.warning(f'⚠️ marks_data is EMPTY! This will cause no marks to display!')
    
    return render_template('rm_add_marks.html', 
                           session=session, 
                           subjects=subjects,
                           students=students,
                           selected_subject=selected_subject,
                           marks_data=marks_data,
                           registrations_data=registrations_data)

@result_management_bp.route('/auto_save_marks/<int:session_id>', methods=['POST'])
@login_required
def auto_save_marks(session_id):
    """Auto-save marks via AJAX"""
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        return jsonify({'success': False, 'message': 'You do not have access to this session.'}), 403
    
    try:
        data = request.get_json()
        subject_id = data.get('subject_id')
        if not subject_id:
            return jsonify({'success': False, 'message': 'Subject ID is required'}), 400
        
        subject = RSubject.query.get_or_404(subject_id)
        
        # Get students for this subject (same logic as add_marks)
        from blueprints.course_management.models import StudentCourseRegistration
        from blueprints.student_management.models import Student as StudentProfile
        
        all_rstudents = RStudent.query.filter_by(session_id=session_id).all()
        student_id_to_rstudent_id = {rs.student_id: rs.id for rs in all_rstudents}
        
        student_profiles = StudentProfile.query.filter(
            StudentProfile.student_id.in_(student_id_to_rstudent_id.keys())
        ).all()
        student_profile_id_to_rstudent_id = {
            profile.id: student_id_to_rstudent_id.get(profile.student_id)
            for profile in student_profiles
            if profile.student_id in student_id_to_rstudent_id
        }
        
        if not student_profile_id_to_rstudent_id:
            return jsonify({'success': False, 'message': 'No students found'}), 400
        
        # Get ALL registrations for this course_code (regardless of year/term)
        # This ensures students who deregistered and re-registered are still included
        course_filters = [
            StudentCourseRegistration.student_id.in_(student_profile_id_to_rstudent_id.keys()),
            StudentCourseRegistration.course_code == subject.code,
            StudentCourseRegistration.status.in_(['finalized', 'pending', 'archived'])
        ]
        
        registered_regs = StudentCourseRegistration.query.filter(*course_filters).all()
        registered_rstudent_ids = set()
        for reg in registered_regs:
            rstudent_id = student_profile_id_to_rstudent_id.get(reg.student_id)
            if rstudent_id:
                registered_rstudent_ids.add(rstudent_id)
        
        students = RStudent.query.filter(RStudent.id.in_(registered_rstudent_ids)).all()
        student_id_map = {s.id: s for s in students}
        
        # Process marks from request
        marks_data = data.get('marks', {})
        updated_count = 0
        
        for student_id_str, student_marks in marks_data.items():
            try:
                student_id = int(student_id_str)
                if student_id not in student_id_map:
                    continue
                
                student = student_id_map[student_id]
                existing_mark = RMark.query.filter_by(student_id=student.id, subject_id=subject.id).first()
                if existing_mark is None:
                    existing_mark = RMark(student_id=student.id, subject_id=subject.id)
                    db.session.add(existing_mark)
                
                # Check if the student is registered for this course as a retake
                registration = RCourseRegistration.query.filter_by(student_id=student.id, subject_id=subject.id).first()
                is_retake = registration.is_retake if registration else False
                existing_mark.is_retake = is_retake
                
                total_marks = 0
                if subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                    if student_marks.get('attendance'):
                        existing_mark.attendance = float(student_marks.get('attendance'))
                        existing_mark.attendance_manual = True
                    else:
                        existing_mark.attendance = None
                    existing_mark.continuous_assessment = float(student_marks.get('continuous_assessment')) if student_marks.get('continuous_assessment') else None
                    existing_mark.part_a = float(student_marks.get('part_a')) if student_marks.get('part_a') else None
                    existing_mark.part_b = float(student_marks.get('part_b')) if student_marks.get('part_b') else None
                    total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.continuous_assessment, existing_mark.part_a, existing_mark.part_b]))
                
                elif subject.subject_type == 'Sessional':
                    if student_marks.get('attendance'):
                        existing_mark.attendance = float(student_marks.get('attendance'))
                        existing_mark.attendance_manual = True
                    else:
                        existing_mark.attendance = None
                    existing_mark.sessional_report = float(student_marks.get('sessional_report')) if student_marks.get('sessional_report') else None
                    existing_mark.sessional_viva = float(student_marks.get('sessional_viva')) if student_marks.get('sessional_viva') else None
                    total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.sessional_report, existing_mark.sessional_viva]))
                
                elif subject.subject_type in ('Thesis (UG)', 'Thesis I (UG)', 'Thesis II (UG)'):
                    if student_marks.get('attendance'):
                        existing_mark.attendance = float(student_marks.get('attendance'))
                        existing_mark.attendance_manual = True
                    else:
                        existing_mark.attendance = None
                    existing_mark.thesis_evaluation = float(student_marks.get('thesis_evaluation')) if student_marks.get('thesis_evaluation') else None
                    existing_mark.presentation = float(student_marks.get('presentation')) if student_marks.get('presentation') else None
                    total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.thesis_evaluation, existing_mark.presentation]))
                
                elif subject.subject_type == 'Dissertation':
                    if subject.dissertation_type == 'Type1':
                        existing_mark.supervisor_assessment = float(student_marks.get('supervisor_assessment')) if student_marks.get('supervisor_assessment') else None
                        existing_mark.proposal_presentation = float(student_marks.get('proposal_presentation')) if student_marks.get('proposal_presentation') else None
                        existing_mark.project_report = None
                        existing_mark.defense = None
                        total_marks = sum(filter(None, [existing_mark.supervisor_assessment, existing_mark.proposal_presentation]))
                    elif subject.dissertation_type == 'Type2':
                        existing_mark.supervisor_assessment = float(student_marks.get('supervisor_assessment')) if student_marks.get('supervisor_assessment') else None
                        existing_mark.project_report = float(student_marks.get('project_report')) if student_marks.get('project_report') else None
                        existing_mark.defense = float(student_marks.get('defense')) if student_marks.get('defense') else None
                        existing_mark.proposal_presentation = None
                        total_marks = sum(filter(None, [existing_mark.supervisor_assessment, existing_mark.project_report, existing_mark.defense]))
                    else:
                        # Fallback for existing Dissertation subjects without type
                        existing_mark.supervisor_assessment = float(student_marks.get('supervisor_assessment')) if student_marks.get('supervisor_assessment') else None
                        existing_mark.proposal_presentation = float(student_marks.get('proposal_presentation')) if student_marks.get('proposal_presentation') else None
                        existing_mark.project_report = float(student_marks.get('project_report')) if student_marks.get('project_report') else None
                        existing_mark.defense = float(student_marks.get('defense')) if student_marks.get('defense') else None
                        total_marks = sum(filter(None, [existing_mark.supervisor_assessment, existing_mark.proposal_presentation, existing_mark.project_report, existing_mark.defense]))
                
                elif subject.subject_type == 'Viva':
                    existing_mark.viva = float(student_marks.get('viva')) if student_marks.get('viva') else None
                    total_marks = existing_mark.viva or 0
                
                existing_mark.total_marks = total_marks
                existing_mark.grade_point, existing_mark.grade_letter = calculate_grade(total_marks, is_retake=is_retake)
                updated_count += 1
                
            except (ValueError, TypeError) as e:
                current_app.logger.error(f"Error processing marks for student {student_id_str}: {e}", exc_info=True)
                continue
        
        db.session.commit()
        # Emit WebSocket event for live update
        try:
            from utils.websocket_events import emit_marks_update
            from datetime import datetime
            emit_marks_update(session_id, {
                'subject_id': subject_id,
                'updated_count': updated_count
            })
        except Exception as e:
            current_app.logger.warning(f'Failed to emit marks update event: {e}')
        return jsonify({
            'success': True,
            'message': f'Auto-saved marks for {updated_count} student(s)',
            'updated_count': updated_count
        })
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error auto-saving marks: {e}", exc_info=True)
        return jsonify({
            'success': False,
            'message': f'Error auto-saving marks: {str(e)}'
        }), 500

@result_management_bp.route('/view_results/<int:session_id>')
@login_required
def view_results(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        flash('You do not have access to this session.', 'danger')
        return redirect(url_for('result_management.index'))
    return render_template('rm_view_results.html', session=session)

@result_management_bp.route('/course_wise_result/<int:session_id>')
@login_required
def course_wise_result(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        flash('You do not have access to this session.', 'danger')
        return redirect(url_for('result_management.index'))
    
    subjects = RSubject.query.filter_by(session_id=session_id).order_by(RSubject.code).all()
    selected_subject_id = request.args.get('subject_id', type=int)
    results = []
    selected_subject = None
    if selected_subject_id:
        selected_subject = RSubject.query.get(selected_subject_id)
        if not selected_subject or selected_subject.session_id != session_id:
            flash('Invalid subject selected.', 'danger')
            return redirect(url_for('result_management.course_wise_result', session_id=session_id))
        
        # Check if marks exist for this subject
        marks_count = RMark.query.filter_by(subject_id=selected_subject_id).count()
        registrations_count = RCourseRegistration.query.filter_by(subject_id=selected_subject_id).count()
        
        # Define columns to select
        base_columns = [
            RStudent.student_id, RStudent.name, RMark.total_marks,
            RMark.grade_letter, RMark.grade_point, RMark.is_retake
        ]
        extra_columns = []
        if selected_subject:
            if selected_subject.subject_type in ['Theory', 'Theory (UG)', 'Theory (PG)']:
                extra_columns = [RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b]
            elif selected_subject.subject_type == 'Sessional':
                extra_columns = [RMark.attendance, RMark.sessional_report, RMark.sessional_viva]
            elif selected_subject.subject_type in ['Thesis (UG)', 'Thesis I (UG)', 'Thesis II (UG)']:
                extra_columns = [RMark.attendance, RMark.thesis_evaluation, RMark.presentation]
            elif selected_subject.subject_type == 'Dissertation':
                if selected_subject.dissertation_type == 'Type1':
                    extra_columns = [RMark.supervisor_assessment, RMark.proposal_presentation]
                else:  # Type2
                    extra_columns = [RMark.supervisor_assessment, RMark.project_report, RMark.defense]
            elif selected_subject.subject_type == 'Viva':
                extra_columns = [RMark.viva]
        all_columns = base_columns + extra_columns
        
        try:
            # If registrations exist, use them. Otherwise, query marks directly (for backward compatibility)
            if registrations_count > 0:
                # Use outerjoin to include students who are registered even if they don't have marks yet
                results = db.session.query(*all_columns)\
                    .select_from(RStudent)\
                    .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == selected_subject_id))\
                    .outerjoin(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == selected_subject_id))\
                    .filter(RStudent.session_id == session_id)\
                    .order_by(RStudent.student_id).all()
            else:
                # No registrations found - query marks directly (marks exist but registrations don't)
                results = db.session.query(*all_columns)\
                    .select_from(RStudent)\
                    .join(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == selected_subject_id))\
                    .filter(RStudent.session_id == session_id)\
                    .order_by(RStudent.student_id).all()
            
            # Recalculate ALL total_marks, grade_point, grade_letter to ensure accuracy
            # This ensures totals are always correct even if marks were updated
            marks_to_update = RMark.query.filter_by(subject_id=selected_subject_id).all()
            needs_update = False
            for mark in marks_to_update:
                # ALWAYS recalculate - don't check if total_marks is None
                try:
                    registration = RCourseRegistration.query.filter_by(student_id=mark.student_id, subject_id=selected_subject_id).first()
                    is_retake = registration.is_retake if registration else False
                    mark.is_retake = is_retake
                    
                    # Calculate total_marks based on subject type
                    total_marks = 0
                    if selected_subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                        # Sum all values (None = 0): attendance + continuous_assessment + part_a + part_b
                        # Don't use filter(None) because it skips 0 values too
                        attendance = mark.attendance if mark.attendance is not None else 0
                        ca = mark.continuous_assessment if mark.continuous_assessment is not None else 0
                        part_a = mark.part_a if mark.part_a is not None else 0
                        part_b = mark.part_b if mark.part_b is not None else 0
                        total_marks = attendance + ca + part_a + part_b
                    elif selected_subject.subject_type == 'Sessional':
                        attendance = mark.attendance if mark.attendance is not None else 0
                        report = mark.sessional_report if mark.sessional_report is not None else 0
                        viva = mark.sessional_viva if mark.sessional_viva is not None else 0
                        total_marks = attendance + report + viva
                    elif selected_subject.subject_type in ('Thesis (UG)', 'Thesis I (UG)', 'Thesis II (UG)'):
                        attendance = mark.attendance if mark.attendance is not None else 0
                        evaluation = mark.thesis_evaluation if mark.thesis_evaluation is not None else 0
                        presentation = mark.presentation if mark.presentation is not None else 0
                        total_marks = attendance + evaluation + presentation
                    elif selected_subject.subject_type == 'Dissertation':
                        if selected_subject.dissertation_type == 'Type1':
                            supervisor = mark.supervisor_assessment if mark.supervisor_assessment is not None else 0
                            proposal = mark.proposal_presentation if mark.proposal_presentation is not None else 0
                            total_marks = supervisor + proposal
                        elif selected_subject.dissertation_type == 'Type2':
                            supervisor = mark.supervisor_assessment if mark.supervisor_assessment is not None else 0
                            report = mark.project_report if mark.project_report is not None else 0
                            defense = mark.defense if mark.defense is not None else 0
                            total_marks = supervisor + report + defense
                        else:
                            supervisor = mark.supervisor_assessment if mark.supervisor_assessment is not None else 0
                            proposal = mark.proposal_presentation if mark.proposal_presentation is not None else 0
                            report = mark.project_report if mark.project_report is not None else 0
                            defense = mark.defense if mark.defense is not None else 0
                            total_marks = supervisor + proposal + report + defense
                    elif selected_subject.subject_type == 'Viva':
                        total_marks = mark.viva if mark.viva is not None else 0
                    
                    # Update total_marks and grades
                    mark.total_marks = total_marks
                    mark.grade_point, mark.grade_letter = calculate_grade(total_marks, is_retake=is_retake)
                    needs_update = True
                except Exception as e:
                    current_app.logger.error(f"Error calculating grades for mark {mark.id}: {e}", exc_info=True)
            
            if needs_update:
                db.session.commit()
                # Re-query results with updated values
                if registrations_count > 0:
                    results = db.session.query(*all_columns)\
                        .select_from(RStudent)\
                        .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == selected_subject_id))\
                        .outerjoin(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == selected_subject_id))\
                        .filter(RStudent.session_id == session_id)\
                        .order_by(RStudent.student_id).all()
                else:
                    results = db.session.query(*all_columns)\
                        .select_from(RStudent)\
                        .join(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == selected_subject_id))\
                        .filter(RStudent.session_id == session_id)\
                        .order_by(RStudent.student_id).all()
        except Exception as e:
            current_app.logger.error(f'Error fetching course-wise results: {e}', exc_info=True)
            flash('Error loading results. Please try again.', 'danger')
            results = []
    return render_template('rm_course_wise_result.html',
                           session=session,
                           subjects=subjects,
                           selected_subject_id=selected_subject_id,
                           selected_subject=selected_subject,
                           results=results)

@result_management_bp.route('/student_wise_result/<int:session_id>')
@login_required
def student_wise_result(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        flash('You do not have access to this session.', 'danger')
        return redirect(url_for('result_management.index'))
    students = RStudent.query.filter_by(session_id=session_id).order_by(RStudent.student_id).all()

    selected_student_id = request.args.get('student_id', type=int)
    selected_student = None
    results = []
    term_assessment = {}

    if selected_student_id:
        selected_student = RStudent.query.get(selected_student_id)
        if not selected_student or selected_student.session_id != session_id:
            flash('Invalid student selected.', 'danger')
            return redirect(url_for('result_management.student_wise_result', session_id=session_id))
        
        # Check if marks exist for this student
        marks_count = RMark.query.filter_by(student_id=selected_student_id).count()
        registrations_count = RCourseRegistration.query.filter_by(student_id=selected_student_id).count()
        
        try:
            # If registrations exist, use them. Otherwise, query marks directly (for backward compatibility)
            if registrations_count > 0:
                # Fetch results for the selected student (ONLY registered courses)
                # Use outerjoin to show registered courses even if marks haven't been entered yet
                results = db.session.query(
                    RSubject.code.label('subject_code'),
                    RSubject.name.label('subject_name'),
                    RSubject.credit.label('registered_credits'),
                    RMark.grade_letter,
                    RMark.grade_point,
                    RMark.is_retake,
                    RSubject.subject_type
                ).select_from(RStudent)\
                 .join(RCourseRegistration, RCourseRegistration.student_id == RStudent.id)\
                 .join(RSubject, RSubject.id == RCourseRegistration.subject_id)\
                 .outerjoin(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == RSubject.id))\
                 .filter(RStudent.id == selected_student_id)\
                 .filter(RSubject.session_id == session_id)\
                 .order_by(RSubject.code).all()
            else:
                # No registrations found - query marks directly (marks exist but registrations don't)
                results = db.session.query(
                    RSubject.code.label('subject_code'),
                    RSubject.name.label('subject_name'),
                    RSubject.credit.label('registered_credits'),
                    RMark.grade_letter,
                    RMark.grade_point,
                    RMark.is_retake,
                    RSubject.subject_type
                ).select_from(RStudent)\
                 .join(RMark, RMark.student_id == RStudent.id)\
                 .join(RSubject, RSubject.id == RMark.subject_id)\
                 .filter(RStudent.id == selected_student_id)\
                 .filter(RSubject.session_id == session_id)\
                 .order_by(RSubject.code).all()
            
            # Calculate missing total_marks, grade_point, grade_letter on-the-fly
            # Fetch all marks for this student to update missing calculations
            marks_to_update = RMark.query.filter_by(student_id=selected_student_id).all()
            needs_update = False
            for mark in marks_to_update:
                if mark.total_marks is None:
                    # Get subject to determine calculation method
                    subject = RSubject.query.get(mark.subject_id)
                    if not subject:
                        continue
                    
                    # Calculate total_marks, grade_point, grade_letter
                    try:
                        registration = RCourseRegistration.query.filter_by(student_id=mark.student_id, subject_id=mark.subject_id).first()
                        is_retake = registration.is_retake if registration else False
                        mark.is_retake = is_retake
                        
                        total_marks = 0
                        if subject.subject_type in ('Theory', 'Theory (UG)', 'Theory (PG)'):
                            total_marks = sum(filter(None, [mark.attendance, mark.continuous_assessment, mark.part_a, mark.part_b]))
                        elif subject.subject_type == 'Sessional':
                            total_marks = sum(filter(None, [mark.attendance, mark.sessional_report, mark.sessional_viva]))
                        elif subject.subject_type in ('Thesis (UG)', 'Thesis I (UG)', 'Thesis II (UG)'):
                            total_marks = sum(filter(None, [mark.attendance, mark.thesis_evaluation, mark.presentation]))
                        elif subject.subject_type == 'Dissertation':
                            if subject.dissertation_type == 'Type1':
                                total_marks = sum(filter(None, [mark.supervisor_assessment, mark.proposal_presentation]))
                            elif subject.dissertation_type == 'Type2':
                                total_marks = sum(filter(None, [mark.supervisor_assessment, mark.project_report, mark.defense]))
                            else:
                                total_marks = sum(filter(None, [mark.supervisor_assessment, mark.proposal_presentation, mark.project_report, mark.defense]))
                        elif subject.subject_type == 'Viva':
                            total_marks = mark.viva or 0
                        
                        mark.total_marks = total_marks
                        mark.grade_point, mark.grade_letter = calculate_grade(total_marks, is_retake=is_retake)
                        needs_update = True
                    except Exception as e:
                        current_app.logger.error(f"Error calculating grades for mark {mark.id}: {e}", exc_info=True)
            
            if needs_update:
                db.session.commit()
                # Re-query results with updated values
                if registrations_count > 0:
                    results = db.session.query(
                        RSubject.code.label('subject_code'),
                        RSubject.name.label('subject_name'),
                        RSubject.credit.label('registered_credits'),
                        RMark.grade_letter,
                        RMark.grade_point,
                        RMark.is_retake,
                        RSubject.subject_type
                    ).select_from(RStudent)\
                     .join(RCourseRegistration, RCourseRegistration.student_id == RStudent.id)\
                     .join(RSubject, RSubject.id == RCourseRegistration.subject_id)\
                     .outerjoin(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == RSubject.id))\
                     .filter(RStudent.id == selected_student_id)\
                     .filter(RSubject.session_id == session_id)\
                     .order_by(RSubject.code).all()
                else:
                    results = db.session.query(
                        RSubject.code.label('subject_code'),
                        RSubject.name.label('subject_name'),
                        RSubject.credit.label('registered_credits'),
                        RMark.grade_letter,
                        RMark.grade_point,
                        RMark.is_retake,
                        RSubject.subject_type
                    ).select_from(RStudent)\
                     .join(RMark, RMark.student_id == RStudent.id)\
                     .join(RSubject, RSubject.id == RMark.subject_id)\
                     .filter(RStudent.id == selected_student_id)\
                     .filter(RSubject.session_id == session_id)\
                     .order_by(RSubject.code).all()
        except Exception as e:
            current_app.logger.error(f'Error fetching student-wise results: {e}', exc_info=True)
            flash('Error loading results. Please try again.', 'danger')
            results = []

        total_registered_credits = 0
        total_earned_credits = 0
        total_earned_credit_points = 0
        
        processed_results = []
        for res in results:
            earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
            earned_credit_points = (res.grade_point or 0) * res.registered_credits
            
            processed_results.append({
                'subject_code': res.subject_code,
                'subject_name': res.subject_name,
                'registered_credits': res.registered_credits,
                'grade_letter': res.grade_letter,
                'grade_point': res.grade_point,
                'earned_credits': earned_credits,
                'earned_credit_points': earned_credit_points,
                'remarks': 'Retake' if res.is_retake else ''
            })

            total_registered_credits += res.registered_credits
            total_earned_credits += earned_credits
            total_earned_credit_points += earned_credit_points

        tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
        
        term_assessment = {
            'total_registered_credits': total_registered_credits,
            'total_earned_credits': total_earned_credits,
            'total_earned_credit_points': total_earned_credit_points,
            'tgpa': tgpa
        }
        results = processed_results
        
    return render_template('rm_student_wise_result.html',
                           session=session,
                           students=students,
                           selected_student_id=selected_student_id,
                           selected_student=selected_student,
                           results=results,
                           term_assessment=term_assessment)


@result_management_bp.route('/course_registration/<int:session_id>', methods=['GET', 'POST'])
@login_required
def course_registration(session_id):
    session = RSession.query.get_or_404(session_id)
    if not _can_access_session(session):
        flash('You do not have access to this session.', 'danger')
        return redirect(url_for('result_management.index'))
    subjects = RSubject.query.filter_by(session_id=session_id).order_by(RSubject.code).all()
    students = RStudent.query.filter_by(session_id=session_id).order_by(RStudent.student_id).all()

    selected_subject_id = request.args.get('subject_id', type=int)
    selected_subject = RSubject.query.get(selected_subject_id) if selected_subject_id else None

    registrations = {}
    auto_registered_student_ids = set()
    auto_retake_student_ids = set()
    if selected_subject:
        # Load existing registrations for the selected subject
        existing_regs = db.session.query(RCourseRegistration).filter_by(subject_id=selected_subject.id).all()
        for reg in existing_regs:
            registrations[reg.student_id] = reg

        # Auto-select students who registered the course in the Course Registration module
        student_number_to_result_id = {
            student.student_id: student.id
            for student in students
            if student.student_id
        }

        if student_number_to_result_id:
            student_profiles = StudentProfile.query.filter(
                StudentProfile.student_id.in_(student_number_to_result_id.keys())
            ).all()
            student_profile_id_to_result_id = {
                profile.id: student_number_to_result_id.get(profile.student_id)
                for profile in student_profiles
            }

            if student_profile_id_to_result_id:
                course_filters = [
                    StudentCourseRegistration.student_id.in_(student_profile_id_to_result_id.keys()),
                    StudentCourseRegistration.course_code == selected_subject.code
                ]
                if session.name:
                    course_filters.append(StudentCourseRegistration.academic_session == session.name)
                if session.year:
                    course_filters.append(StudentCourseRegistration.year == session.year)
                if session.term:
                    course_filters.append(StudentCourseRegistration.term == session.term)

                allowed_statuses = ['draft', 'pending', 'finalized']
                auto_regs = StudentCourseRegistration.query.filter(
                    *course_filters,
                    StudentCourseRegistration.status.in_(allowed_statuses)
                ).all()

                for auto_reg in auto_regs:
                    result_student_id = student_profile_id_to_result_id.get(auto_reg.student_id)
                    if result_student_id:
                        auto_registered_student_ids.add(result_student_id)
                        remark_text = (auto_reg.remark or '').strip().lower()
                        if remark_text in {'retake', 're-retake', 're retake', 'reretake'}:
                            auto_retake_student_ids.add(result_student_id)

    if request.method == 'POST':
        subject_id = int(request.form.get('subject_id'))
        if not subject_id:
            flash('A subject must be selected.', 'danger')
            return redirect(url_for('result_management.course_registration', session_id=session_id))
        
        # Get all student IDs that were submitted (i.e., whose checkboxes could have been checked)
        students_on_page = RStudent.query.filter_by(session_id=session_id).all()
        student_ids_on_page = {s.id for s in students_on_page}

        # First, delete all existing registrations for this subject for the students shown on the page
        db.session.query(RCourseRegistration).filter(
            RCourseRegistration.subject_id == subject_id,
            RCourseRegistration.student_id.in_(student_ids_on_page)
        ).delete(synchronize_session=False)

        # Now, add back the ones that were checked in the form
        for student_id in student_ids_on_page:
            if f'reg_{student_id}' in request.form:
                is_retake = f'retake_{student_id}' in request.form
                new_reg = RCourseRegistration(
                    student_id=student_id,
                    subject_id=subject_id,
                    is_retake=is_retake
                )
                db.session.add(new_reg)

                # Sync the is_retake flag with the corresponding RMark record
                mark = RMark.query.filter_by(student_id=student_id, subject_id=subject_id).first()
                if mark:
                    mark.is_retake = is_retake
                    # Recalculate grade if total_marks exists
                    if mark.total_marks is not None:
                        mark.grade_point, mark.grade_letter = calculate_grade(mark.total_marks, is_retake=is_retake)
            else:
                 # If a student is unregistered, ensure their mark record also has is_retake as False
                mark = RMark.query.filter_by(student_id=student_id, subject_id=subject_id).first()
                if mark and mark.is_retake:
                    mark.is_retake = False
                    if mark.total_marks is not None:
                        mark.grade_point, mark.grade_letter = calculate_grade(mark.total_marks, is_retake=False)

        db.session.commit()
        flash('Course registration updated successfully!', 'success')
        return redirect(url_for('result_management.course_registration', session_id=session_id, subject_id=subject_id))

    return render_template('rm_course_registration.html', 
                           session=session, 
                           students=students, 
                           subjects=subjects,
                           selected_subject=selected_subject,
                           registrations=registrations,
                           auto_registered_student_ids=auto_registered_student_ids,
                           auto_retake_student_ids=auto_retake_student_ids)


@result_management_bp.route('/delete_session/<int:session_id>', methods=['POST'])
@login_required
def delete_session(session_id):
    """Delete a result session and all related data"""
    if not _can_manage_sessions():
        flash('Only the Head can delete sessions.', 'danger')
        return redirect(url_for('result_management.index'))
    
    try:
        session = RSession.query.get_or_404(session_id)
        session_name = session.name
        
        # Cascade delete will handle students, subjects, marks, and registrations
        # due to cascade="all, delete-orphan" in the model relationships
        db.session.delete(session)
        db.session.commit()
        
        flash(f'Session "{session_name}" and all related data deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error deleting result session {session_id}: {e}', exc_info=True)
        flash(f'Error deleting session: {str(e)}', 'danger')
    
    return redirect(url_for('result_management.index'))


class PDFGenerator:
    def __init__(self, buffer, pagesize):
        from reportlab.platypus import SimpleDocTemplate
        from reportlab.lib.units import inch
        self.buffer = buffer
        self.doc = SimpleDocTemplate(
            buffer,
            pagesize=pagesize,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch,
            leftMargin=0.5*inch,
            rightMargin=0.5*inch
        )
        self.story = []

    def _footer(self, canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        page_number_text = f"Page {doc.page} of {doc.doc.page_count}"
        canvas.drawRightString(letter[0] - 40, 30, page_number_text)
        canvas.restoreState()

    def build(self, elements):
        # A two-pass approach to get total page numbers for the footer
        
        # First pass
        doc_temp = SimpleDocTemplate(BytesIO(), pagesize=letter)
        frame = Frame(doc_temp.leftMargin, doc_temp.bottomMargin, doc_temp.width, doc_temp.height, id='normal')
        template = PageTemplate(id='main_temp', frames=[frame])
        doc_temp.addPageTemplates([template])
        doc_temp.build(elements)
        self.total_pages = doc_temp.page

        # Second pass (actual build)
        frame = Frame(self.doc.leftMargin, self.doc.bottomMargin, self.doc.width, self.doc.height, id='normal')
        template = PageTemplate(id='main', frames=[frame], onPage=self._footer)
        self.doc.addPageTemplates([template])
        self.doc.build(elements)

class CourseTabulationPDF(PDFGenerator):
    def __init__(self, buffer, subject, session):
        from reportlab.lib.pagesizes import A4
        super().__init__(buffer, pagesize=A4)
        self.subject = subject
        self.session = session
        self.doc.title = f"Course_Result_{subject.code}"

    def generate_elements(self, results):
        styles = getSampleStyleSheet()
        # Custom styles
        styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='Right', alignment=TA_RIGHT))
        styles.add(ParagraphStyle(name='Left', alignment=TA_LEFT))
        styles.add(ParagraphStyle(name='Line_Data', parent=styles['Normal'], alignment=TA_CENTER, leading=14))
        styles.add(ParagraphStyle(name='TableCellCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_CENTER, leading=8))
        styles.add(ParagraphStyle(name='TableCellLeftCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_LEFT, leading=8))
        styles.add(ParagraphStyle(name='InfoCompact', parent=styles['Normal'], fontSize=9, leading=10))
        
        # Center align headers
        styles['h1'].alignment = TA_CENTER
        styles['h2'].alignment = TA_CENTER
        
        elements = []
        
        # Header
        elements.append(Paragraph("Khulna University", styles['h1']))
        elements.append(Paragraph("Course-wise Tabulation Sheet", styles['h2']))
        elements.append(Spacer(1, 0.08*inch))

        # Info Table (compact)
        info_data = [
            [
                Paragraph(f"<b>Year:</b> {self.session.year or 'N/A'}<br/><b>Discipline:</b> Law<br/><b>Course No.:</b> {self.subject.code}<br/><b>Course Title:</b> {self.subject.name}", styles['InfoCompact']),
                Paragraph(f"<b>Term:</b> {self.session.term}<br/><b>School:</b> Law<br/><b>CH:</b> {self.subject.credit:.1f}<br/><br/><b>Session:</b> {self.session.name}", styles['InfoCompact'])
            ]
        ]
        info_table = Table(info_data, colWidths=[3.6*inch, 2.0*inch])
        info_table.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('FONTSIZE', (0,0), (-1,-1), 9)]))
        elements.append(info_table)
        elements.append(Spacer(1, 0.08*inch))

        # Base headers
        base_headers = ['Student\nNo.']
        specific_headers = []
        # Specific headers based on subject type
        if self.subject.subject_type in ['Theory', 'Theory (UG)']:
            specific_headers = ['Attendance\n(10)', 'C.A.\n(40)', 'Sec. A\n(25)', 'Sec. B\n(25)']
        elif self.subject.subject_type == 'Sessional':
            specific_headers = ['Attendance\n(10)', 'Report\n(60)', 'Viva\n(30)']
        elif self.subject.subject_type == 'Dissertation':
            if self.subject.dissertation_type == 'Type1':
                specific_headers = ['Supervisor\n(70)', 'Presentation\n(30)']
            else:
                # Type2 distribution updated per requirement: 20 + 50 + 30
                specific_headers = ['Supervisor\n(20)', 'Report\n(50)', 'Defense\n(30)']
        end_headers = ['Total\nMarks\n(100)', 'Grade\nPoint', 'Grade\nLetter', 'Remarks']
        table_headers = base_headers + specific_headers + end_headers
        data = [[Paragraph(h.replace('\n', '<br/>'), styles['TableCellCompact']) for h in table_headers]]
        for res in results:
            row_data = [Paragraph(str(res.student_id), styles['TableCellCompact'])]
            marks_data = list(res)[6:]
            for mark in marks_data:
                row_data.append(Paragraph(f"{mark:.1f}" if mark is not None else '', styles['TableCellCompact']))
            total_marks_rounded = round(res.total_marks) if res.total_marks is not None else None
            row_data.extend([
                Paragraph(f"{total_marks_rounded}" if total_marks_rounded is not None else '', styles['TableCellCompact']),
                Paragraph(f"{res.grade_point:.2f}" if res.grade_point is not None else '', styles['TableCellCompact']),
                Paragraph(res.grade_letter or '', styles['TableCellCompact']),
                Paragraph('Retake' if res.is_retake else '', styles['TableCellCompact'])
            ])
            data.append(row_data)
        # Dynamic column widths, reduce by 10%
        base_widths = [1.1*inch] + [0.8*inch]*len(specific_headers) + [0.7*inch, 0.7*inch, 0.7*inch, 0.8*inch]
        col_widths = [w * 0.9 for w in base_widths]
        table = Table(data, colWidths=col_widths, rowHeights=0.35*inch)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 3),
            ('TOPPADDING', (0, 0), (-1, 0), 3),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])
        table.setStyle(style)
        elements.append(table)
        elements.append(Spacer(1, 0.5*inch))  # Add more space before signature
        
        # Signature section (wider, with space)
        signature_table = Table([
            [Paragraph('-----------------------', styles['Center']), '', Paragraph('-----------------------', styles['Center']), '', Paragraph('-----------------------', styles['Center'])],
            [
                Paragraph('Signature of the First Tabulator<br/>Date:', styles['Center']), '',
                Paragraph('Signature of the Chairman of the Examination Committee<br/>Date:', styles['Center']), '',
                Paragraph('Signature of the Second Tabulator<br/>Date:', styles['Center'])
            ]
        ], colWidths=[2.2*inch, 0.7*inch, 2.2*inch, 0.7*inch, 2.2*inch])
        signature_table.setStyle(TableStyle([
            ('VALIGN', (0, 1), (0, 1), 'CENTER'),
            ('ALIGN', (2, 1), (2, 1), 'CENTER'),
            ('ALIGN', (4, 1), (4, 1), 'CENTER'),
            ('FONTNAME', (0, 1), (4, 1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (4, 1), 8),
            ('TOPPADDING', (0, 1), (4, 1), 10),
        ]))
        elements.append(signature_table)
        
        return elements

class StudentTabulationPDF(PDFGenerator):
    def __init__(self, buffer, student, session):
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate
        pagesize = landscape(A4)
        super().__init__(buffer, pagesize=pagesize)
        self.student = student
        self.session = session
        self.page_count = 1
        self.doc.title = f"Tabulation_{student.student_id}"
        # Restore default margins
        # (no aggressive margin reduction)

    def generate_elements(self, results, term_assessment):
        from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.units import inch
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Bold', parent=styles['Normal'], fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='TableCell', parent=styles['Normal'], fontSize=10, wordWrap='CJK', alignment=TA_LEFT))
        styles.add(ParagraphStyle(name='TableCellCenter', parent=styles['Normal'], fontSize=10, wordWrap='CJK', alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='TableCellCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_LEFT, leading=8))
        styles.add(ParagraphStyle(name='TableCellCenterCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_CENTER, leading=8))
        styles.add(ParagraphStyle(name='InfoCompact', parent=styles['Normal'], fontSize=9, leading=10))
        elements = []

        # Title
        elements.append(Paragraph("Khulna University", styles['Title']))
        elements.append(Paragraph("Student-wise Tabulation Sheet", styles['Title']))
        elements.append(Spacer(1, 0.08*inch))

        # Two-column info table (compact)
        info_data = [
            [Paragraph('<b>Year:</b>', styles['InfoCompact']), str(self.session.year or self.student.year or 'N/A'),
             Paragraph('<b>Term:</b>', styles['InfoCompact']), str(self.session.term)],
            [Paragraph('<b>Student No.:</b>', styles['InfoCompact']), str(self.student.student_id),
             Paragraph('<b>Name of Student:</b>', styles['InfoCompact']), str(self.student.name)],
            [Paragraph('<b>Discipline:</b>', styles['InfoCompact']), str(self.student.discipline or 'Law'),
             Paragraph('<b>Session:</b>', styles['InfoCompact']), str(self.session.name)],
            [Paragraph('<b>School:</b>', styles['InfoCompact']), 'Law', '', '']
        ]
        info_table = Table(info_data, colWidths=[0.9*inch, 1.7*inch, 0.9*inch, 1.7*inch])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('TOPPADDING', (0,0), (-1,-1), 1),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 0.08*inch))

        # If no results, show message and return
        if not results:
            elements.append(Paragraph("No results found for this student.", styles['Normal']))
            return elements

        # Results Table (compact)
        headers = ['Course No.', 'Course Title', 'Registered Credit Hours', 'Letter Grade', 'Grade Point (GP)', 'Earned Credit Hours (CH)', 'Earned Credit Points (GP*CH)', 'Remarks']
        data = [[Paragraph(h, styles['TableCellCenterCompact']) for h in headers]]
        total_registered_credits = 0
        total_earned_credits = 0
        total_earned_credit_points = 0
        for res in results:
            # Always show 0 for GP, Earned CH, Earned Credit Points if missing/None
            grade_point = res.get('grade_point', 0)
            earned_credits = res.get('earned_credits', 0)
            earned_credit_points = res.get('earned_credit_points', 0)
            row = [
                Paragraph(str(res.get('subject_code', '') or ''), styles['TableCellCompact']),
                Paragraph(str(res.get('subject_name', '') or ''), styles['TableCellCompact']),
                Paragraph(str(res.get('registered_credits', '') or ''), styles['TableCellCenterCompact']),
                Paragraph(str(res.get('grade_letter', '') or ''), styles['TableCellCenterCompact']),
                Paragraph(f"{grade_point}" if grade_point is not None else "0", styles['TableCellCenterCompact']),
                Paragraph(f"{earned_credits}" if earned_credits is not None else "0", styles['TableCellCenterCompact']),
                Paragraph(f"{earned_credit_points}" if earned_credit_points is not None else "0", styles['TableCellCenterCompact']),
                Paragraph(str(res.get('remarks', '') or ''), styles['TableCellCenterCompact'])
            ]
            # If mark distribution fields exist, show them (even for F)
            extra_fields = []
            for key in ['attendance', 'continuous_assessment', 'part_a', 'part_b', 'sessional_report', 'sessional_viva', 'supervisor_assessment', 'proposal_presentation', 'project_report', 'defense']:
                if key in res:
                    val = res.get(key, None)
                    extra_fields.append(Paragraph(str(val) if val is not None else '', styles['TableCellCenterCompact']))
            if extra_fields:
                row = row[:3] + extra_fields + row[3:]
            data.append(row)
            total_registered_credits += res['registered_credits'] if res['registered_credits'] else 0
            total_earned_credits += res['earned_credits'] if res['earned_credits'] else 0
            total_earned_credit_points += res['earned_credit_points'] if res['earned_credit_points'] else 0
        # Total row
        data.append([
            '', '',
            Paragraph(f"Total = {total_registered_credits}", styles['TableCellCenterCompact']), '', '',
            Paragraph(f"{total_earned_credits}", styles['TableCellCenterCompact']),
            Paragraph(f"{total_earned_credit_points}", styles['TableCellCenterCompact']),
            ''
        ])
        # Reduce column widths by 20% for compactness
        base_widths = [1.5*inch, 2.5*inch, 1.1*inch, 1.1*inch, 1.2*inch, 1.3*inch, 1.5*inch, 1.1*inch]
        col_widths = [w * 0.8 for w in base_widths]
        table = Table(data, colWidths=col_widths)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 0.7, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ])
        table.setStyle(style)
        elements.append(table)
        elements.append(Spacer(1, 0.08*inch))

        # Term Assessment (compact)
        elements.append(Paragraph('<b>Term Assessment</b>', styles['InfoCompact']))
        elements.append(Paragraph(f"Total Earned Credit Hours in this Term (TCH) = {term_assessment['total_earned_credits']}", styles['InfoCompact']))
        elements.append(Paragraph(f"Total Registered Credit Hours in this Term (RCH) = {term_assessment['total_registered_credits']}", styles['InfoCompact']))
        elements.append(Paragraph(f"Total Earned Credit Points in this Term (TCP) = {term_assessment['total_earned_credit_points']}", styles['InfoCompact']))
        elements.append(Paragraph(f"TGPA = TCP/TCH = {term_assessment['tgpa']:.2f}", styles['InfoCompact']))
        elements.append(Spacer(1, 0.18*inch))

        # Signature lines (compact, with spacing)
        sig_col_width = 1.7*inch
        spacer_col_width = 0.4*inch
        sig_table = Table([
            ['', '', '', '', ''],
            [
                Paragraph('Signature of the First Tabulator<br/>Date:', styles['TableCellCenterCompact']), '',
                Paragraph('Signature of the Chairman of the Examination Committee<br/>Date:', styles['TableCellCenterCompact']), '',
                Paragraph('Signature of the Second Tabulator<br/>Date:', styles['TableCellCenterCompact'])
            ]
        ], colWidths=[sig_col_width, spacer_col_width, sig_col_width, spacer_col_width, sig_col_width])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0, 1), (0, 1), 'CENTER'),
            ('ALIGN', (2, 1), (2, 1), 'CENTER'),
            ('ALIGN', (4, 1), (4, 1), 'CENTER'),
            ('FONTNAME', (0, 1), (4, 1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (4, 1), 8),
            ('TOPPADDING', (0, 1), (4, 1), 10),
        ]))
        elements.append(sig_table)
        return elements


@result_management_bp.route('/test_download/<int:session_id>/<int:subject_id>')
@login_required
def test_download(session_id, subject_id):
    """Simple test endpoint for debugging download issues"""
    try:
        current_app.logger.info(f"Testing download endpoint for session {session_id}, subject {subject_id}")
        
        # Check if subject and session exist
        subject = RSubject.query.get_or_404(subject_id)
        session = RSession.query.get_or_404(session_id)
        
        # Create a simple test PDF
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        elements = []
        elements.append(Paragraph(f"Test PDF for {subject.name}", styles['Title']))
        elements.append(Paragraph(f"Session: {session.name}", styles['Normal']))
        elements.append(Paragraph(f"Subject: {subject.code} - {subject.name}", styles['Normal']))
        elements.append(Paragraph(f"Generated at: {datetime.now()}", styles['Normal']))
        
        doc.build(elements)
        buffer.seek(0)
        
        current_app.logger.info(f"Test PDF generated successfully")
        
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="test_{subject.code}.pdf"',
                'Content-Length': str(len(buffer.getvalue()))
            }
        )
        
    except Exception as e:
        current_app.logger.error(f"Test download failed: {e}")
        return f"Test failed: {str(e)}", 500

@result_management_bp.route('/download/course_result/<int:session_id>/<int:subject_id>')
@login_required
def download_course_result(session_id, subject_id):
    try:
        current_app.logger.info(f"Starting course result PDF generation for session {session_id}, subject {subject_id}")
        
        # Check if required modules are available
        try:
            import reportlab
            current_app.logger.info("Required modules available for PDF")
        except ImportError as e:
            current_app.logger.error(f"Missing required module for PDF: {e}")
            flash(f'Missing required module for PDF: {e}', 'error')
            return redirect(url_for('result_management.course_wise_result', session_id=session_id))
        
        subject = RSubject.query.get_or_404(subject_id)
        session = RSession.query.get_or_404(session_id)

        # Build query based on subject type
        base_columns = [
            RStudent.student_id, RStudent.name, RMark.total_marks,
            RMark.grade_letter, RMark.grade_point, RMark.is_retake
        ]
        extra_columns = []
        if subject.subject_type in ['Theory', 'Theory (UG)']:
            extra_columns = [RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b]
        elif subject.subject_type == 'Sessional':
            extra_columns = [RMark.attendance, RMark.sessional_report, RMark.sessional_viva]
        elif subject.subject_type == 'Dissertation':
            if subject.dissertation_type == 'Type1':
                extra_columns = [RMark.supervisor_assessment, RMark.proposal_presentation]
            else:
                extra_columns = [RMark.supervisor_assessment, RMark.project_report, RMark.defense]
        elif subject.subject_type == 'Viva':
            extra_columns = [RMark.viva]
        all_columns = base_columns + extra_columns
        
        # Check if registrations exist
        registrations_count = RCourseRegistration.query.filter_by(subject_id=subject_id).count()
        
        if registrations_count > 0:
            results = db.session.query(*all_columns)\
                .join(RMark, RStudent.id == RMark.student_id)\
                .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == subject_id))\
                .filter(RMark.subject_id == subject_id)\
                .order_by(RStudent.student_id).all()
        else:
            # No registrations found - query marks directly
            results = db.session.query(*all_columns)\
                .select_from(RStudent)\
                .join(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == subject_id))\
                .filter(RStudent.session_id == session_id)\
                .order_by(RStudent.student_id).all()
        
        buffer = BytesIO()
        pdf = CourseTabulationPDF(buffer, subject, session)
        elements = pdf.generate_elements(results)
        pdf.doc.build(elements)
        buffer.seek(0)
        
        current_app.logger.info(f"Course result PDF generated successfully for session {session_id}, subject {subject_id}")
        
        # Enhanced headers for cPanel compatibility
        pdf_data = buffer.getvalue()
        filename = f"Course_{subject.code}_Result.pdf"
        
        response = Response(
            pdf_data,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"; filename*=UTF-8\'\'{filename}',
                'Content-Length': str(len(pdf_data)),
                'Cache-Control': 'no-cache, no-store, must-revalidate',
                'Pragma': 'no-cache',
                'Expires': '0',
                'X-Content-Type-Options': 'nosniff',
                'X-Frame-Options': 'DENY'
            }
        )
        
        return response
        
    except Exception as e:
        # Simple error logging without error_handler dependency
        current_app.logger.error(f"Error generating course result PDF for session {session_id}, subject {subject_id}: {e}")
        current_app.logger.error(f"Error details: {str(e)}")
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('result_management.course_wise_result', session_id=session_id))

@result_management_bp.route('/download/student_result/<int:session_id>/<int:student_id>')
@login_required
def download_student_result(session_id, student_id):
    try:
        current_app.logger.info(f"Starting student result PDF generation for session {session_id}, student {student_id}")
        
        # Check if required modules are available
        try:
            import reportlab
            current_app.logger.info("Required modules available for PDF")
        except ImportError as e:
            current_app.logger.error(f"Missing required module for PDF: {e}")
            flash(f'Missing required module for PDF: {e}', 'error')
            return redirect(url_for('result_management.student_wise_result', session_id=session_id))
        
        student = RStudent.query.get_or_404(student_id)
        session = RSession.query.get_or_404(session_id)

        # Check if registrations exist for this student
        registrations_count = RCourseRegistration.query.filter_by(student_id=student_id).count()
        
        # If registrations exist, use them. Otherwise, query marks directly (for backward compatibility)
        if registrations_count > 0:
            # Use outerjoin to show registered courses even if marks haven't been entered yet
            results_query = db.session.query(
                RSubject.code.label('subject_code'),
                RSubject.name.label('subject_name'),
                RSubject.credit.label('registered_credits'),
                RMark.grade_letter, RMark.grade_point, RMark.is_retake, RSubject.subject_type,
                RMark.total_marks,
                RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b,
                RMark.sessional_report, RMark.sessional_viva,
                RMark.viva,
                RMark.supervisor_assessment, RMark.proposal_presentation, RMark.project_report, RMark.defense
            ).select_from(RStudent)
            results_query = results_query.join(RCourseRegistration, RCourseRegistration.student_id == RStudent.id)
            results_query = results_query.join(RSubject, RSubject.id == RCourseRegistration.subject_id)
            results_query = results_query.outerjoin(RMark, (RMark.student_id == RStudent.id) & (RMark.subject_id == RSubject.id))
            results_query = results_query.filter(RStudent.id == student_id)
            results_query = results_query.filter(RSubject.session_id == session_id)
            results_query = results_query.order_by(RSubject.code)
        else:
            # No registrations found - query marks directly (marks exist but registrations don't)
            results_query = db.session.query(
                RSubject.code.label('subject_code'),
                RSubject.name.label('subject_name'),
                RSubject.credit.label('registered_credits'),
                RMark.grade_letter, RMark.grade_point, RMark.is_retake, RSubject.subject_type,
                RMark.total_marks,
                RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b,
                RMark.sessional_report, RMark.sessional_viva,
                RMark.viva,
                RMark.supervisor_assessment, RMark.proposal_presentation, RMark.project_report, RMark.defense
            ).select_from(RStudent)
            results_query = results_query.join(RMark, RMark.student_id == RStudent.id)
            results_query = results_query.join(RSubject, RSubject.id == RMark.subject_id)
            results_query = results_query.filter(RStudent.id == student_id)
            results_query = results_query.filter(RSubject.session_id == session_id)
            results_query = results_query.order_by(RSubject.code)
        
        results = results_query.all()

        total_registered_credits, total_earned_credits, total_earned_credit_points = 0, 0, 0
        processed_results = []
        for res in results:
            earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
            earned_credit_points = (res.grade_point or 0) * res.registered_credits
            processed_results.append({
                'subject_code': res.subject_code, 'subject_name': res.subject_name,
                'registered_credits': res.registered_credits, 'grade_letter': res.grade_letter,
                'grade_point': res.grade_point, 'earned_credits': earned_credits,
                'earned_credit_points': earned_credit_points,
                'remarks': 'Retake' if res.is_retake else ''
            })
            total_registered_credits += res.registered_credits
            total_earned_credits += earned_credits
            total_earned_credit_points += earned_credit_points

        tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
        term_assessment = {
            'total_registered_credits': total_registered_credits, 'total_earned_credits': total_earned_credits,
            'total_earned_credit_points': total_earned_credit_points, 'tgpa': tgpa
        }

        buffer = BytesIO()
        pdf = StudentTabulationPDF(buffer, student, session)
        elements = pdf.generate_elements(processed_results, term_assessment)
        pdf.doc.build(elements)
        buffer.seek(0)
        
        current_app.logger.info(f"Student result PDF generated successfully for session {session_id}, student {student_id}")
        
        # Enhanced headers for cPanel compatibility
        pdf_data = buffer.getvalue()
        filename = f"Student_{student.student_id}_Tabulation.pdf"
        
        response = Response(
            pdf_data,
            mimetype='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{filename}"; filename*=UTF-8\'\'{filename}',
                'Content-Length': str(len(pdf_data)),
                'Cache-Control': 'no-cache, no-store, must-revalidate',
                'Pragma': 'no-cache',
                'Expires': '0',
                'X-Content-Type-Options': 'nosniff',
                'X-Frame-Options': 'DENY'
            }
        )
        
        return response
        
    except Exception as e:
        # Simple error logging without error_handler dependency
        current_app.logger.error(f"Error generating student result PDF for session {session_id}, student {student_id}: {e}")
        current_app.logger.error(f"Error details: {str(e)}")
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('result_management.student_wise_result', session_id=session_id))

@result_management_bp.route('/download/student_result_docx/<int:session_id>/<int:student_id>')
@login_required
def download_student_result_docx(session_id, student_id):
    student = RStudent.query.get_or_404(student_id)
    session = RSession.query.get_or_404(session_id)

    # Only include subjects where the student is registered
    results_query = db.session.query(
        RSubject.code.label('subject_code'),
        RSubject.name.label('subject_name'),
        RSubject.credit.label('registered_credits'),
        RMark.grade_letter, RMark.grade_point, RMark.is_retake, RSubject.subject_type
    ).select_from(RMark)\
     .join(RStudent, RStudent.id == RMark.student_id)\
     .join(RSubject, RSubject.id == RMark.subject_id)\
     .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == RSubject.id))\
     .filter(RStudent.id == student_id)\
     .order_by(RSubject.code).all()

    total_registered_credits, total_earned_credits, total_earned_credit_points = 0, 0, 0
    processed_results = []
    for res in results_query:
        earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
        earned_credit_points = (res.grade_point or 0) * res.registered_credits
        processed_results.append({
            'subject_code': res.subject_code, 'subject_name': res.subject_name,
            'registered_credits': res.registered_credits, 'grade_letter': res.grade_letter,
            'grade_point': res.grade_point, 'earned_credits': earned_credits,
            'earned_credit_points': earned_credit_points,
            'remarks': 'Retake' if res.is_retake else ''
        })
        total_registered_credits += res.registered_credits
        total_earned_credits += earned_credits
        total_earned_credit_points += earned_credit_points

    tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
    term_assessment = {
        'total_registered_credits': total_registered_credits, 'total_earned_credits': total_earned_credits,
        'total_earned_credit_points': total_earned_credit_points, 'tgpa': tgpa
    }

    # Lazy import docx to prevent startup hang
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Create DOCX
    doc = Document()
    # Title
    doc.add_heading('Khulna University', 0).alignment = 1
    doc.add_heading('Student-wise Tabulation Sheet', level=1).alignment = 1

    # Two-column info table
    info_table = doc.add_table(rows=2, cols=4)
    info_table.autofit = False
    info_table.columns[0].width = Inches(1.2)
    info_table.columns[1].width = Inches(2.2)
    info_table.columns[2].width = Inches(1.2)
    info_table.columns[3].width = Inches(2.2)
    # Left column
    info_table.cell(0,0).text = 'Year:'
    info_table.cell(0,1).text = str(session.year or student.year or 'N/A')
    info_table.cell(1,0).text = 'Student No.:'
    info_table.cell(1,1).text = str(student.student_id)
    # Right column
    info_table.cell(0,2).text = 'Term:'
    info_table.cell(0,3).text = str(session.term)
    info_table.cell(1,2).text = 'Name of Student:'
    info_table.cell(1,3).text = str(student.name)
    # Next row for discipline/session/school
    row = info_table.add_row().cells
    row[0].text = 'Discipline:'
    row[1].text = str(student.discipline or 'Law')
    row[2].text = 'Session:'
    row[3].text = str(session.name)
    row2 = info_table.add_row().cells
    row2[0].text = 'School:'
    row2[1].text = 'Law'
    row2[2].text = ''
    row2[3].text = ''
    for row in info_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.runs[0].font.size = Pt(11)
    doc.add_paragraph('')

    # Results Table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Course No.', 'Course Title', 'Registered Credit Hours', 'Letter Grade', 'Grade Point (GP)', 'Earned Credit Hours (CH)', 'Earned Credit Points (GP*CH)', 'Remarks']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.runs[0].font.bold = True
            p.alignment = 1
    for result in processed_results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(result['subject_code'])
        row_cells[1].text = str(result['subject_name'])
        row_cells[2].text = f"{result['registered_credits']}"
        row_cells[3].text = result['grade_letter'] or ''
        row_cells[4].text = f"{result['grade_point']}" if result['grade_point'] is not None else ''
        row_cells[5].text = f"{result['earned_credits']}"
        row_cells[6].text = f"{result['earned_credit_points']}"
        row_cells[7].text = result['remarks']
    # Total row
    total_row = table.add_row().cells
    total_row[0].text = ''
    total_row[1].text = ''
    total_row[2].text = f"Total = {total_registered_credits}"
    total_row[3].text = ''
    total_row[4].text = ''
    total_row[5].text = f"{total_earned_credits}"
    total_row[6].text = f"{total_earned_credit_points}"
    total_row[7].text = ''
    for cell in total_row:
        for p in cell.paragraphs:
            p.runs[0].font.bold = True
    doc.add_paragraph('')

    # Term Assessment
    p = doc.add_paragraph()
    p.add_run('Term Assessment').bold = True
    doc.add_paragraph(f"Total Earned Credit Hours in this Term (TCH) = {total_earned_credits}")
    doc.add_paragraph(f"Total Registered Credit Hours in this Term (RCH) = {total_registered_credits}")
    doc.add_paragraph(f"Total Earned Credit Points in this Term (TCP) = {total_earned_credit_points}")
    doc.add_paragraph(f"TGPA = TCP/TCH = {tgpa:.2f}")
    doc.add_paragraph('')

    # Signature lines (3 columns)
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(2.5)
    sig_table.columns[1].width = Inches(2.5)
    sig_table.columns[2].width = Inches(2.5)
    sig_table.cell(0,0).text = ''
    sig_table.cell(0,1).text = ''
    sig_table.cell(0,2).text = ''
    sig_table.cell(1,0).text = 'Signature of the First Tabulator\nDate:'
    sig_table.cell(1,1).text = 'Signature of the Second Tabulator\nDate:'
    sig_table.cell(1,2).text = 'Signature of the Chairman, Examination Committee\nDate:'
    for i in range(3):
        for p in sig_table.cell(1,i).paragraphs:
            p.alignment = 1
    doc.add_paragraph('')

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'Student_{student.student_id}_Tabulation.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@result_management_bp.route('/download/all_student_results/<int:session_id>')
@login_required
def download_all_student_results(session_id):
    session = RSession.query.get_or_404(session_id)
    students = RStudent.query.filter_by(session_id=session_id).all()
    
    if not students:
        flash('No students in this session to generate results for.', 'warning')
        return redirect(url_for('result_management.student_wise_result', session_id=session_id))

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zf:
        for student in students:
            # Re-use logic from download_student_result
            results_query = db.session.query(
                RSubject.code.label('subject_code'), RSubject.name.label('subject_name'), RSubject.credit.label('registered_credits'),
                RMark.grade_letter, RMark.grade_point, RMark.is_retake, RSubject.subject_type
            ).select_from(RMark)
            results_query = results_query.join(RStudent, RStudent.id == RMark.student_id)
            results_query = results_query.join(RSubject, RSubject.id == RMark.subject_id)
            results_query = results_query.join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == RSubject.id))
            results_query = results_query.filter(RStudent.id == student.id).order_by(RSubject.code)
            results_query = results_query.all()

            if not results_query: continue

            total_registered_credits, total_earned_credits, total_earned_credit_points = 0, 0, 0
            processed_results = []
            for res in results_query:
                earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
                earned_credit_points = (res.grade_point or 0) * res.registered_credits
                processed_results.append({
                    'subject_code': res.subject_code, 'subject_name': res.subject_name, 'registered_credits': res.registered_credits,
                    'grade_letter': res.grade_letter, 'grade_point': res.grade_point, 'earned_credits': earned_credits,
                    'earned_credit_points': earned_credit_points, 'remarks': 'Retake' if res.is_retake else ''
                })
                total_registered_credits += res.registered_credits
                total_earned_credits += earned_credits
                total_earned_credit_points += earned_credit_points
            
            tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
            term_assessment = {
                'total_registered_credits': total_registered_credits, 'total_earned_credits': total_earned_credits,
                'total_earned_credit_points': total_earned_credit_points, 'tgpa': tgpa
            }
            
            pdf_buffer = BytesIO()
            pdf = StudentTabulationPDF(pdf_buffer, student, session)
            elements = pdf.generate_elements(processed_results, term_assessment)
            pdf.doc.build(elements)
            pdf_buffer.seek(0)
            zf.writestr(f'Student_{student.student_id}_Tabulation.pdf', pdf_buffer.read())

    zip_buffer.seek(0)
    
    # Enhanced headers for cPanel compatibility
    zip_data = zip_buffer.getvalue()
    filename = f'All_Student_Results_{session.name}.zip'
    
    response = Response(
        zip_data,
        mimetype='application/zip',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"; filename*=UTF-8\'\'{filename}',
            'Content-Length': str(len(zip_data)),
            'Cache-Control': 'no-cache, no-store, must-revalidate',
            'Pragma': 'no-cache',
            'Expires': '0',
            'X-Content-Type-Options': 'nosniff',
            'X-Frame-Options': 'DENY'
        }
    )
    
    return response

@result_management_bp.route('/download/all_course_results/<int:session_id>')
@login_required
def download_all_course_results(session_id):
    session = RSession.query.get_or_404(session_id)
    subjects = RSubject.query.filter_by(session_id=session_id).all()

    if not subjects:
        flash('No subjects in this session to generate results for.', 'warning')
        return redirect(url_for('result_management.course_wise_result', session_id=session_id))
        
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zf:
        for subject in subjects:
            base_columns = [
                RStudent.student_id, RStudent.name, RMark.total_marks,
                RMark.grade_letter, RMark.grade_point, RMark.is_retake
            ]
            extra_columns = []
            if subject.subject_type in ['Theory', 'Theory (UG)']:
                extra_columns = [RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b]
            elif subject.subject_type == 'Sessional':
                extra_columns = [RMark.attendance, RMark.sessional_report, RMark.sessional_viva]
            elif subject.subject_type == 'Dissertation':
                if subject.dissertation_type == 'Type1':
                    extra_columns = [RMark.supervisor_assessment, RMark.proposal_presentation]
                else:
                    extra_columns = [RMark.supervisor_assessment, RMark.project_report, RMark.defense]
            all_columns = base_columns + extra_columns
            results = db.session.query(*all_columns)\
                .join(RMark, RStudent.id == RMark.student_id)\
                .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == subject.id))\
                .filter(RMark.subject_id == subject.id)\
                .order_by(RStudent.student_id).all()

            if not results: continue

            pdf_buffer = BytesIO()
            pdf = CourseTabulationPDF(pdf_buffer, subject, session)
            elements = pdf.generate_elements(results)
            pdf.doc.build(elements)
            pdf_buffer.seek(0)
            zf.writestr(f'Course_{subject.code}_Result.pdf', pdf_buffer.read())

    zip_buffer.seek(0)
    
    # Enhanced headers for cPanel compatibility
    zip_data = zip_buffer.getvalue()
    filename = f'All_Course_Results_{session.name}.zip'
    
    response = Response(
        zip_data,
        mimetype='application/zip',
        headers={
            'Content-Disposition': f'attachment; filename="{filename}"; filename*=UTF-8\'\'{filename}',
            'Content-Length': str(len(zip_data)),
            'Cache-Control': 'no-cache, no-store, must-revalidate',
            'Pragma': 'no-cache',
            'Expires': '0',
            'X-Content-Type-Options': 'nosniff',
            'X-Frame-Options': 'DENY'
        }
    )
    
    return response
