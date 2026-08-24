"""Extract plain text from uploaded course files, with AI OCR for scans."""
import base64
import io
import os
import re


_MIN_NATIVE_CHARS = 80
_MAX_OCR_PAGES = 12
_MAX_PDF_INLINE = 8 * 1024 * 1024


def _clean_text(text):
    if not text:
        return ''
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_text_from_file(file_path, file_type=None):
    """Return extracted plain text or empty string on failure."""
    if not file_path or not os.path.isfile(file_path):
        return ''

    ext = (file_type or os.path.splitext(file_path)[1]).lower().lstrip('.')
    try:
        if ext == 'pdf':
            return _extract_pdf(file_path)
        if ext in ('doc', 'docx'):
            return _extract_docx(file_path)
        if ext in ('txt', 'md'):
            return _extract_plain(file_path)
        if ext in ('jpg', 'jpeg', 'png', 'webp'):
            return _ocr_images([_image_file_to_jpeg_b64(file_path)])
    except Exception:
        return ''
    return ''


def _extract_plain(file_path):
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as handle:
        return _clean_text(handle.read())


def _extract_pdf(file_path):
    native = _extract_pdf_native(file_path)
    if len(native) >= _MIN_NATIVE_CHARS:
        return native
    ocr = _ocr_pdf(file_path)
    return _clean_text(ocr) or native


def _extract_pdf_native(file_path):
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return ''
    reader = PdfReader(file_path)
    parts = []
    pages = getattr(reader, 'pages', None)
    if pages is None:
        return ''
    for page in list(pages)[:40]:
        try:
            parts.append(page.extract_text() or '')
        except Exception:
            continue
    return _clean_text('\n'.join(parts))


def _extract_docx(file_path):
    try:
        from docx import Document
    except ImportError:
        return ''
    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return _clean_text('\n'.join(parts))


def _pil_to_jpeg_b64(image, max_width=1280, quality=55):
    from PIL import Image

    if image.mode not in ('RGB', 'L'):
        image = image.convert('RGB')
    elif image.mode == 'L':
        image = image.convert('RGB')
    if image.width > max_width:
        height = int(image.height * (max_width / float(image.width)))
        image = image.resize((max_width, max(1, height)))
    buf = io.BytesIO()
    image.save(buf, format='JPEG', quality=quality, optimize=True)
    return base64.b64encode(buf.getvalue()).decode('ascii')


def _image_file_to_jpeg_b64(file_path):
    from PIL import Image

    with Image.open(file_path) as image:
        return _pil_to_jpeg_b64(image.copy())


def _render_pdf_page_images(file_path, max_pages=_MAX_OCR_PAGES):
    images = _render_with_pypdfium(file_path, max_pages)
    if images:
        return images
    return _render_with_pymupdf(file_path, max_pages)


def _render_with_pypdfium(file_path, max_pages):
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return []
    try:
        pdf = pdfium.PdfDocument(file_path)
        out = []
        count = min(len(pdf), max_pages)
        for index in range(count):
            page = pdf[index]
            bitmap = page.render(scale=1.6)
            image = bitmap.to_pil()
            out.append(_pil_to_jpeg_b64(image))
        pdf.close()
        return out
    except Exception:
        return []


def _render_with_pymupdf(file_path, max_pages):
    try:
        import fitz
    except ImportError:
        return []
    try:
        doc = fitz.open(file_path)
        out = []
        for index, page in enumerate(doc):
            if index >= max_pages:
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            from PIL import Image

            image = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)
            out.append(_pil_to_jpeg_b64(image))
        doc.close()
        return out
    except Exception:
        return []


def _ocr_images(images_b64):
    images_b64 = [img for img in (images_b64 or []) if img]
    if not images_b64:
        return ''
    from utils.ai.client import generate_text_from_media

    chunks = []
    batch_size = 3
    for start in range(0, len(images_b64), batch_size):
        batch = images_b64[start:start + batch_size]
        page_from = start + 1
        page_to = start + len(batch)
        text = generate_text_from_media(
            'You transcribe scanned exam papers and answer guidelines. '
            'Return plain text only. Keep question numbers, Bangla and English as written. '
            'Do not answer the questions.',
            f'Transcribe pages {page_from}-{page_to} exactly.',
            images=batch,
            max_tokens=4000,
        )
        if text:
            chunks.append(text)
    return _clean_text('\n'.join(chunks))


def _ocr_pdf(file_path):
    images = _render_pdf_page_images(file_path)
    if images:
        return _ocr_images(images)
    try:
        from utils.ai.client import generate_text_from_media, get_active_provider_setting

        cfg = get_active_provider_setting()
        if cfg.get('provider') != 'gemini':
            return ''
        size = os.path.getsize(file_path)
        if size > _MAX_PDF_INLINE:
            return ''
        with open(file_path, 'rb') as handle:
            pdf_bytes = handle.read()
        return generate_text_from_media(
            'You transcribe scanned exam papers and answer guidelines. '
            'Return plain text only. Keep question numbers, Bangla and English as written. '
            'Do not answer the questions.',
            'This PDF may be a scan. Extract all readable text in page order.',
            pdf_bytes=pdf_bytes,
            max_tokens=5000,
        )
    except Exception:
        return ''


def chunk_text(text, chunk_size=1200, overlap=150):
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(0, end - overlap)
    return [c for c in chunks if c]
